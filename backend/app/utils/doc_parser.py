"""
文档格式解析器
支持：Word (.docx) / PDF / Markdown
"""
import re
from difflib import SequenceMatcher


def parse_docx(filepath):
    """
    解析 Word 文档，提取文本和章节结构
    使用 python-docx 库
    """
    try:
        from docx import Document
    except ImportError:
        return {"paragraphs": [], "tables_text": [], "chapters": [], "full_text": "", "error": "python-docx not installed"}

    doc = Document(filepath)

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else "Normal"
        paragraphs.append({
            "text": text,
            "style": style_name
        })

    # 提取表格文本
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                tables_text.append(row_text)

    # 提取章节结构
    chapters = _extract_chapters_from_paragraphs(paragraphs)

    full_text = "\n".join(p["text"] for p in paragraphs)
    full_text += "\n" + "\n".join(tables_text)

    # 如果顶级章节太少，扁平化（同 PDF 逻辑）
    chapters = _flatten_chapters_if_needed(chapters)

    return {
        "paragraphs": paragraphs,
        "tables_text": tables_text,
        "chapters": chapters,
        "full_text": full_text,
        "type": "docx"
    }


def parse_pdf(filepath):
    """
    解析 PDF 文档，提取文本和书签结构
    优先级：PyMuPDF > pdfplumber > pypdf > OCR
    """
    pages_text = []
    toc = []

    # 尝试 PyMuPDF (fitz) - 主提取器
    try:
        import fitz
        doc = fitz.open(filepath)

        toc = doc.get_toc()

        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()
            if text and text.strip():
                pages_text.append(text)

        doc.close()
    except ImportError:
        print("[doc_parser] PyMuPDF (fitz) not available")
    except Exception as e:
        print(f"[doc_parser] PyMuPDF error: {e}")

    fitz_text = "\n".join(pages_text)

    # 如果 fitz 提取到文本，用 fitz 做章节分析
    if fitz_text.strip():
        full_text = fitz_text
        chapters = _build_chapters(fitz_text, pages_text, toc, filepath)
        return {
            "pages_text": pages_text,
            "toc": toc,
            "chapters": chapters,
            "full_text": full_text,
            "type": "pdf"
        }

    # 回退 pdfplumber
    print("[doc_parser] PyMuPDF got empty text, trying pdfplumber")
    plumber_texts = []
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t and t.strip():
                    plumber_texts.append(t)
        if plumber_texts:
            pages_text = plumber_texts
            full_text = "\n".join(plumber_texts)
            print(f"[doc_parser] pdfplumber extracted {len(full_text)} chars")
            chapters = _build_chapters(full_text, pages_text, [], filepath)
            return {
                "pages_text": pages_text,
                "toc": [],
                "chapters": chapters,
                "full_text": full_text,
                "type": "pdf"
            }
    except ImportError:
        print("[doc_parser] pdfplumber not available")
    except Exception as e:
        print(f"[doc_parser] pdfplumber error: {e}")

    # 回退 pypdf
    print("[doc_parser] pdfplumber got empty text, trying pypdf")
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        pypdf_texts = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                pypdf_texts.append(t)
        if pypdf_texts:
            pages_text = pypdf_texts
            full_text = "\n".join(pypdf_texts)
            print(f"[doc_parser] pypdf extracted {len(full_text)} chars")
            chapters = _build_chapters(full_text, pages_text, [], filepath)
            return {
                "pages_text": pages_text,
                "toc": [],
                "chapters": chapters,
                "full_text": full_text,
                "type": "pdf"
            }
    except ImportError:
        print("[doc_parser] pypdf not available")
    except Exception as e:
        print(f"[doc_parser] pypdf error: {e}")

    # OCR 回退（扫描件/图片型PDF）
    print("[doc_parser] pypdf got empty text, trying OCR")
    ocr_text = _ocr_pdf_pages_doc_parser(filepath)
    if ocr_text.strip():
        pages_text = [ocr_text]
        full_text = ocr_text
        print(f"[doc_parser] OCR extracted {len(full_text)} chars")
        chapters = _build_chapters(full_text, pages_text, [], filepath)
        return {
            "pages_text": pages_text,
            "toc": [],
            "chapters": chapters,
            "full_text": full_text,
            "type": "pdf"
        }

    # 所有提取器都失败了
    full_text = ""
    chapters = [{"heading": "全文", "content": full_text, "subsections": []}]
    return {
        "pages_text": pages_text,
        "toc": [],
        "chapters": chapters,
        "full_text": full_text,
        "type": "pdf",
        "error": "无法提取PDF文本内容"
    }


def _build_chapters(full_text, pages_text, toc, filepath):
    """将文本构建为章节结构（使用文本模式匹配）"""
    chapters = []
    text_chapters = detect_headings_from_text(pages_text)
    if text_chapters and len(text_chapters) >= 2:
        chapters = text_chapters
    if not chapters:
        first_page_text = pages_text[0] if pages_text else ""
        first_line = ""
        for line in first_page_text.split('\n'):
            line = line.strip()
            if line and len(line) < 100:
                first_line = line
                break
        heading = first_line or "全文"
        chapters = [{"heading": heading, "content": full_text, "subsections": []}]
    chapters = _flatten_chapters_if_needed(chapters)
    chapters = _filter_cover_chapters(chapters)
    return chapters


def _ocr_pdf_pages_doc_parser(filepath):
    """对扫描件/图片型 PDF 进行 OCR 识别"""
    text_parts = []
    try:
        import fitz
        import io
        from PIL import Image, ImageFilter, ImageEnhance

        doc = fitz.open(filepath)
        zoom = 3.0
        mat = fitz.Matrix(zoom, zoom)

        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = img.convert('L')
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            img = img.filter(ImageFilter.SHARPEN)
            page_text = _ocr_image_doc_parser(img)
            if page_text and page_text.strip():
                text_parts.append(page_text)
                print(f"[doc_parser] OCR Page {i+1}: {len(page_text)} chars")

        doc.close()
    except ImportError as e:
        print(f"[doc_parser] OCR: required library not available - {e}")
    except Exception as e:
        print(f"[doc_parser] OCR error: {e}")

    return "\n".join(text_parts)


def _ocr_image_doc_parser(pil_image):
    """对单张 PIL Image 进行 OCR"""
    import os

    _TESSDATA_CANDIDATES = [
        '/usr/share/tesseract-ocr/5/tessdata',
        '/usr/share/tesseract-ocr/4.00/tessdata',
        '/usr/share/tessdata',
        os.environ.get('TESSDATA_PREFIX', ''),
    ]

    try:
        import tesserocr
        tessdata_path = None
        for candidate in _TESSDATA_CANDIDATES:
            if candidate and os.path.isdir(candidate):
                tessdata_path = candidate
                break
        api = tesserocr.PyTessBaseAPI(path=tessdata_path, lang='chi_sim+eng')
        api.SetImage(pil_image)
        text = api.GetUTF8Text()
        api.End()
        if text and text.strip():
            return _clean_ocr_doc_parser(text)
    except ImportError:
        print("[doc_parser] tesserocr not available for OCR")
    except Exception as e:
        print(f"[doc_parser] tesserocr OCR error: {e}")

    try:
        import pytesseract
        for candidate in _TESSDATA_CANDIDATES:
            if candidate and os.path.isdir(candidate):
                os.environ['TESSDATA_PREFIX'] = candidate
                break
        text = pytesseract.image_to_string(pil_image, lang='chi_sim+eng')
        if text and text.strip():
            return _clean_ocr_doc_parser(text)
    except ImportError:
        print("[doc_parser] pytesseract not available for OCR")
    except Exception as e:
        print(f"[doc_parser] pytesseract OCR error: {e}")

    return ""


def _clean_ocr_doc_parser(text):
    """清理 OCR 文本"""
    import re
    while True:
        cleaned = re.sub(r'([\u4e00-\u9fff\u3400-\u4dbf])\s+([\u4e00-\u9fff\u3400-\u4dbf])', r'\1\2', text)
        if cleaned == text:
            break
        text = cleaned
    text = re.sub(r'([\u4e00-\u9fff])\s+(\d)', r'\1\2', text)
    text = re.sub(r'(\d)\s+([\u4e00-\u9fff])', r'\1\2', text)
    text = re.sub(r'([\u4e00-\u9fff])\s+([a-zA-Z])', r'\1\2', text)
    return text


def parse_markdown(filepath):
    """
    解析 Markdown 文件，按 # 标题提取章节结构
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception:
        return {"chapters": [], "full_text": "", "error": "Failed to read file"}

    lines = content.split('\n')
    chapters = []
    stack = []
    tables_text = []
    in_table = False
    table_buffer = []

    def _add_text_to_current(text):
        if not text.strip():
            return
        if stack:
            node = stack[-1]
            node["content"] += text + " "

    for line in lines:
        stripped = line.strip()

        # 表格检测
        if '|' in stripped and not stripped.startswith('#'):
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(stripped)
            _add_text_to_current(stripped)
            continue
        else:
            if in_table and table_buffer:
                tables_text.append("\n".join(table_buffer))
                in_table = False
                table_buffer = []

        if not stripped:
            continue

        # 跳过分隔线
        if re.match(r'^[-*_]{3,}$', stripped):
            continue

        # 检测标题（支持 h1-h6）
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            new_node = {"heading": text, "content": "", "subsections": []}

            # 弹栈直到找到父级
            while stack and len(stack) >= level:
                stack.pop()

            if stack:
                stack[-1]["subsections"].append(new_node)
            else:
                chapters.append(new_node)

            stack.append(new_node)
            continue

        # 正文
        _add_text_to_current(stripped)

    # 处理最后一个表格
    if in_table and table_buffer:
        tables_text.append("\n".join(table_buffer))

    # 如果顶级章节太少，扁平化
    chapters = _flatten_chapters_if_needed(chapters)

    # 过滤封面/文档标题章节
    chapters = _filter_cover_chapters(chapters)

    full_text = content
    if tables_text:
        full_text += "\n" + "\n".join(tables_text)

    return {
        "chapters": chapters,
        "full_text": full_text,
        "tables_text": tables_text,
        "type": "md"
    }


def _extract_chapters_from_paragraphs(paragraphs):
    """
    从段落列表中提取章节结构
    基于 Heading 样式识别标题层级
    """
    chapters = []
    current_chapter = None
    current_section = None

    for p in paragraphs:
        style = p["style"]
        text = p["text"]

        if style.startswith("Heading 1") or style.startswith("标题 1"):
            if current_chapter:
                if current_section:
                    current_chapter["subsections"].append(current_section)
                chapters.append(current_chapter)
            current_chapter = {"heading": text, "content": "", "subsections": []}
            current_section = None

        elif style.startswith("Heading 2") or style.startswith("标题 2"):
            if current_section and current_chapter:
                current_chapter["subsections"].append(current_section)
            current_section = {"heading": text, "content": ""}

        else:
            if current_section:
                current_section["content"] += text + " "
            elif current_chapter:
                current_chapter["content"] += text + " "

    if current_chapter:
        if current_section:
            current_chapter["subsections"].append(current_section)
        chapters.append(current_chapter)

    return chapters


def _extract_chapters_from_toc(toc, pages_text, doc=None):
    """从 PDF 书签提取章节结构（支持多级）"""
    if not toc:
        return []

    n_pages = len(pages_text)

    # 第一步：用栈构建树结构（先不计算内容）
    roots = []
    stack = []  # [(level, node)]

    for idx, item in enumerate(toc):
        level, title, page_num = item
        title = title.strip()
        start_page_1based = page_num  # 1-indexed

        # 找结束页：下一个同级或更高级书签的前一页
        end_page_1based = n_pages  # 默认到最后一页（1-indexed）
        for j in range(idx + 1, len(toc)):
            next_level, next_title, next_page = toc[j]
            if next_level <= level:
                end_page_1based = next_page - 1
                break

        # 转换为 0-indexed
        start_0 = max(0, start_page_1based - 1)
        end_0 = min(n_pages - 1, end_page_1based - 1)

        # 确保 end >= start（同一页的书签也至少包含那一页）
        if end_0 < start_0:
            end_0 = start_0

        node = {
            "heading": title,
            "content": "",
            "subsections": [],
            "_start": start_0,
            "_end": end_0
        }

        # 弹出栈中 level >= 当前 level 的节点
        while stack and stack[-1][0] >= level:
            stack.pop()

        if not stack:
            roots.append(node)
        else:
            stack[-1][1]["subsections"].append(node)

        stack.append((level, node))

    # 第二步：后序遍历，计算每个节点的内容
    # 父节点内容 = 自己范围内的全部文本（包含子节点，对比时会合并）
    def calc_content(node):
        start = node["_start"]
        end = node["_end"]
        content = ""
        for p in range(start, end + 1):
            if 0 <= p < n_pages:
                content += pages_text[p] + " "
        node["content"] = content.strip()
        # 递归子节点
        for sub in node["subsections"]:
            calc_content(sub)

    for root in roots:
        calc_content(root)

    # 转换内部字段为公开字段（页码，1-indexed）
    def convert_internal(node):
        node["page_start"] = node.pop("_start", 0) + 1
        node["page_end"] = node.pop("_end", 0) + 1
        for sub in node["subsections"]:
            convert_internal(sub)

    for root in roots:
        convert_internal(root)

    return roots


def _flatten_chapters_if_needed(chapters, min_top_level=3):
    """
    如果顶级章节太少但子节很多，把子节提升为一级章节
    这样在对比报告里能看到更细粒度的结果
    """
    if not chapters:
        return chapters

    n_top = len(chapters)
    if n_top >= min_top_level:
        return chapters

    total_subs = sum(len(ch.get("subsections", [])) for ch in chapters)

    if total_subs >= min_top_level:
        new_chapters = []
        for ch in chapters:
            subs = ch.get("subsections", [])
            if subs:
                for sub in subs:
                    new_ch = {
                        "heading": sub["heading"],
                        "content": sub.get("content", ""),
                        "subsections": sub.get("subsections", [])
                    }
                    new_chapters.append(new_ch)
            else:
                new_chapters.append(ch)
        return new_chapters

    return chapters


def _filter_cover_chapters(chapters):
    """
    过滤掉封面/文档标题等没有实质内容的章节
    识别规则：
    1. 第一个章节且内容很少
    2. 标题包含封面、目录、修订历史、版本说明等关键词
    3. 内容主要是元数据（版本号、日期、编号等）
    4. 没有子节且内容远少于其他章节
    """
    if not chapters or len(chapters) <= 1:
        return chapters

    # 封面/非正文关键词（中英文）
    cover_keywords = [
        '封面', '目录', '修订历史', '版本说明', '版本历史', '变更记录',
        '前言', '关于本说明书', '关于本手册', '说明', '目录',
        'cover', 'toc', 'table of contents', 'revision history',
        'version history', 'change log', 'preface', 'about',
        'document number', 'doc no', 'document no',
    ]

    # 计算每个章节的内容长度（包含子节内容）
    chapter_lens = []
    for ch in chapters:
        content = ch.get("content", "") or ""
        for sub in ch.get("subsections", []):
            content += sub.get("content", "") or ""
        chapter_lens.append(len(content.strip()))

    # 找有效章节的平均长度（去掉最短的，去掉最长的，取中间的平均）
    sorted_lens = sorted([l for l in chapter_lens if l > 100], reverse=True)
    if len(sorted_lens) >= 3:
        top_half = sorted_lens[:max(1, len(sorted_lens) * 2 // 3)]
        avg_len = sum(top_half) / len(top_half)
    elif sorted_lens:
        avg_len = max(sorted_lens)
    else:
        avg_len = max(chapter_lens) if chapter_lens else 0

    # 过滤阈值：内容少于平均 10% 且没有子节，或包含封面关键词
    threshold = max(200, avg_len * 0.10)

    def is_cover_chapter(ch, content_len, idx):
        heading = ch.get("heading", "").lower()
        subs = ch.get("subsections", [])
        page_start = ch.get("page_start", 1)

        # 规则1：第一个章节，在第1页，内容很少且没有子节
        if idx == 0 and page_start <= 1 and content_len < threshold and not subs:
            return True

        # 规则2：标题包含封面/目录等关键词
        for kw in cover_keywords:
            if kw.lower() in heading:
                if content_len < avg_len * 0.3 or not subs:
                    return True

        # 规则3：没有子节且内容非常少（少于平均的5%）
        if not subs and content_len < max(100, avg_len * 0.05):
            return True

        return False

    filtered = []
    for i, ch in enumerate(chapters):
        if is_cover_chapter(ch, chapter_lens[i], i):
            continue
        filtered.append(ch)

    # 保护机制：如果过滤后章节数太少（少于原始的一半，或少于2个），返回原始的
    if len(filtered) < max(2, len(chapters) // 2):
        return chapters

    return filtered


def detect_headings_by_font_size(doc):
    """
    基于字体大小检测 PDF 标题
    原理：标题字号通常比正文大，通过字号层级判断标题级别

    返回章节列表，格式同 detect_headings_from_text
    """
    from collections import Counter

    # 收集所有行（按页），每行记录最大字号和文本
    all_lines = []  # [(page_idx, text, max_size, fonts)]

    for page_idx in range(doc.page_count):
        page = doc[page_idx]
        d = page.get_text("dict")
        for block in d.get("blocks", []):
            if block.get("type") != 0:  # 0 = text block
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                if not spans:
                    continue
                line_text = "".join(s.get("text", "") for s in spans).strip()
                if not line_text:
                    continue
                max_size = max(s.get("size", 0) for s in spans)
                all_lines.append({
                    "page": page_idx,
                    "text": line_text,
                    "size": max_size,
                    "fonts": set(s.get("font", "") for s in spans)
                })

    if not all_lines:
        return []

    # 统计字号分布，找正文字号（出现次数最多的）
    size_counter = Counter()
    for line in all_lines:
        # 按四舍五入到 0.5 精度来统计
        size_key = round(line["size"] * 2) / 2
        size_counter[size_key] += 1

    if not size_counter:
        return []

    # 正文字号 = 出现次数最多的字号
    body_size = size_counter.most_common(1)[0][0]

    # 找出所有比正文大的字号（候选标题字号）
    heading_sizes = sorted(
        [s for s in size_counter.keys() if s > body_size + 0.5],
        reverse=True
    )

    # 过滤掉出现次数太少的字号（噪声）
    min_count = max(2, len(all_lines) * 0.005)
    heading_sizes = [s for s in heading_sizes if size_counter.get(s, 0) >= min_count]

    # 只取前 3 个级别（太多级可能是噪声）
    heading_sizes = heading_sizes[:3]

    if not heading_sizes:
        return []

    # 字号到级别的映射：最大=1级，次大=2级，依此类推
    size_to_level = {s: i + 1 for i, s in enumerate(heading_sizes)}

    # 跳过前几页（封面、目录、修订历史等）
    # 找到第一个出现一级标题的页码，从那里开始
    first_heading_page = 0
    for line in all_lines:
        size_key = round(line["size"] * 2) / 2
        if size_key in size_to_level and size_to_level[size_key] == 1:
            first_heading_page = line["page"]
            break

    # 过滤：只保留从第一个一级标题开始的内容
    # 同时排除明显的目录行（含很多点的行）
    filtered_lines = []
    for line in all_lines:
        if line["page"] < first_heading_page:
            continue
        text = line["text"]
        # 排除目录行（含很多省略号）
        if text.count('.') > 5 or '……' in text:
            continue
        # 排除太短的行（页码等）
        if len(text) < 3:
            continue
        # 排除纯数字/页码
        if re.match(r'^[\dIVXLCDM]+$', text, re.IGNORECASE):
            continue
        filtered_lines.append(line)

    if not filtered_lines:
        return []

    # 构建章节结构
    # 用栈来处理多级嵌套
    chapters = []
    stack = []  # [(level, node)]
    current_lines = {1: [], 2: [], 3: []}  # 各级别当前累积的正文行

    def flush_lines_to_current():
        """把累积的正文行加到当前最内层的节点"""
        if not stack:
            return
        level, node = stack[-1]
        # 把该级别及更低级别的正文都加进去
        for l in range(level, 4):
            if current_lines.get(l, []):
                if node["content"]:
                    node["content"] += "\n"
                node["content"] += "\n".join(current_lines[l])
                current_lines[l] = []

    for line in filtered_lines:
        text = line["text"]
        size_key = round(line["size"] * 2) / 2
        level = size_to_level.get(size_key, None)

        if level is not None:
            # 这是一个标题行
            # 先把之前累积的正文加到当前节点
            flush_lines_to_current()

            # 弹出栈中 level >= 当前 level 的节点
            while stack and stack[-1][0] >= level:
                # 先把内容刷到要弹出的节点
                popped_level, popped_node = stack.pop()
                # 把更高级别的正文加到这个节点
                for l in range(popped_level, 4):
                    if current_lines.get(l, []):
                        if popped_node["content"]:
                            popped_node["content"] += "\n"
                        popped_node["content"] += "\n".join(current_lines[l])
                        current_lines[l] = []

            page_num = line["page"] + 1  # 1-indexed
            node = {"heading": text, "content": "", "subsections": [], "page_start": page_num, "page_end": page_num}

            if not stack:
                chapters.append(node)
            else:
                # 加到父节点的 subsections
                parent = stack[-1][1]
                parent["subsections"].append(node)
                # 更新父节点的结束页
                if page_num > parent.get("page_end", 0):
                    parent["page_end"] = page_num

            stack.append((level, node))
            # 清空该级别及更低级别的正文缓存
            for l in range(level, 4):
                current_lines[l] = []
        else:
            # 正文行，加到当前最内层节点
            if stack:
                level = stack[-1][0]
                current_lines[level].append(text)
                # 更新当前节点及其所有祖先的结束页
                page_num = line["page"] + 1
                for l, n in stack:
                    if page_num > n.get("page_end", 0):
                        n["page_end"] = page_num
            # 如果没有任何标题，先不处理（会被 fallback 覆盖）

    # 处理最后一个节点的正文
    flush_lines_to_current()
    # 把栈里剩下的节点的内容都补全
    while stack:
        level, node = stack.pop()
        for l in range(level, 4):
            if current_lines.get(l, []):
                if node["content"]:
                    node["content"] += "\n"
                node["content"] += "\n".join(current_lines[l])
                current_lines[l] = []

    # 额外验证：一级标题数量不能太少也不能太多
    # 如果只有 1 个一级章节，但有多个二级章节，说明一级可能是文档标题
    # 应该把二级提升为一级
    if len(chapters) == 1 and chapters[0]["subsections"]:
        sub_count = len(chapters[0]["subsections"])
        if sub_count >= 3:
            # 很可能一级是文档标题，把二级提升为一级
            title_chapter = chapters[0]
            new_chapters = title_chapter["subsections"]
            # 把原一级标题的内容分配到第一个新章节（或者丢弃）
            # 简单起见，直接用二级作为新的一级
            chapters = new_chapters

    if len(chapters) < 2:
        if len(chapters) == 1 and not chapters[0]["subsections"]:
            return []

    # 补充：在每个一级章节内用文本模式找子节（X.Y / 一、二等格式）
    # 因为子标题字号可能和正文一样，字体大小检测不到
    for ch in chapters:
        if ch.get("subsections"):
            continue  # 已经有子节了，跳过
        content = ch.get("content", "")
        if not content:
            continue
        subsections = _detect_subheadings_from_text(content, max_depth=2)
        if subsections:
            ch["subsections"] = subsections

    return chapters


def _detect_subheadings_from_text(text, max_depth=2):
    """
    从一段文本中检测子标题（用于在章节内补充子节结构）

    支持的标题格式：
    - X.Y / X.Y.Z（数字编号）
    - 第X节 / 第X小节
    - 一、二、 三、（中文数字+顿号）
    - 1. 2. 3.（数字+点+空格）
    - (1) (2) (3) / 1) 2) 3)
    """
    import re as re_mod

    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines:
        return []

    # 表格项目过滤：数字 + 空格 + 单位（ml, mL, μL, L, g, kg 等）
    is_table_item = re_mod.compile(
        r'^\d+(\.\d+)?\s*(ml|mL|μL|uL|L|g|kg|mg|cm|mm|m|nm|µm|V|W|Hz|A|kW|kV|well|Well)\b',
        re_mod.IGNORECASE
    )
    # 表格行标识（含表格列名）
    table_row_pattern = re_mod.compile(
        r'(Cat\.?\s*No\.?|Quantity|Brand|Volume/well|Position\s*$|Tips\s*$|'
        r'^Consumables\s*$|^Reagents\s*$|^Equipment\s*$|^Specification|'
        r'Adaptable\s+tube|型号|规格|数量|单位|备注|参数)',
        re_mod.IGNORECASE
    )
    # 表格标题行
    table_caption_pattern = re_mod.compile(
        r'^(Table|Figure|表|图)\s*\d+[-\u2013\u2014]',
        re_mod.IGNORECASE
    )

    # 标题模式（按优先级排序）
    patterns = [
        # X.Y.Z 三级
        (re_mod.compile(r'^(\d+)\.(\d+)\.(\d+)\s+([\u4e00-\u9fffA-Za-z].{0,80})$'), 3),
        # X.Y 二级
        (re_mod.compile(r'^(\d+)\.(\d+)\s+([\u4e00-\u9fffA-Za-z].{0,80})$'), 2),
        # 第X节 / 第X小节
        (re_mod.compile(r'^第([一二三四五六七八九十\d]+)[节小节]\s+([\u4e00-\u9fffA-Za-z].{0,80})$'), 2),
        # 一、二、三、（中文数字+顿号）
        (re_mod.compile(r'^([一二三四五六七八九十]+)[、\.．]\s*([\u4e00-\u9fffA-Za-z].{0,80})$'), 2),
    ]

    # 收集所有标题行
    headings = []  # [(line_idx, level, text)]
    seen = set()

    for i, line in enumerate(lines):
        if len(line) > 100:
            continue
        # 跳过含很多点的行（目录行）
        if line.count('.') > 5:
            continue
        # 跳过表格项目（数字+单位开头）
        if is_table_item.match(line):
            continue
        # 跳过表格行（含表格列名）
        if table_row_pattern.search(line):
            continue
        # 跳过表格标题行
        if table_caption_pattern.match(line):
            continue
        # 跳过含 | 的表格行
        if '|' in line:
            continue

        for pat, level in patterns:
            m = pat.match(line)
            if m:
                # 去重：同一标题文本只保留第一次
                heading_text = line.strip()
                if heading_text not in seen:
                    seen.add(heading_text)
                    headings.append((i, level, heading_text))
                break

    if not headings:
        return []

    # 额外校验：X.Y 格式的标题编号应该是递增的，如果突然跳很大可能是误报
    # 比如 1.1, 1.2, 1.3, 1.4... 是正常的
    # 但如果 1.1, 1.5, 2.3... 之间夹了很多非标题行，且编号乱序，可能有问题
    xy_headings = [(i, h) for i, lvl, h in headings if lvl == 2 and re_mod.match(r'^\d+\.\d+\s', h)]
    if len(xy_headings) >= 3:
        # 检查编号顺序是否合理
        prev_main = 0
        prev_sub = 0
        bad_count = 0
        for idx, text in xy_headings:
            m = re_mod.match(r'^(\d+)\.(\d+)\s', text)
            if m:
                main = int(m.group(1))
                sub = int(m.group(2))
                # 编号应该递增（main 不变时 sub 递增，main 增加时 sub 重置）
                if main < prev_main or (main == prev_main and sub <= prev_sub):
                    bad_count += 1
                prev_main = main
                prev_sub = sub
        # 如果超过一半都是乱序的，说明检测质量不好
        if bad_count > len(xy_headings) / 2:
            return []

    # 构建章节树（只取前两级）
    # 先确定有哪些级别
    levels_present = sorted(set(h[1] for h in headings))
    if len(levels_present) == 1:
        # 只有一个级别，都作为二级子节
        result = []
        last_idx = 0
        for idx, level, text in headings:
            # 收集从上一个标题到这个标题之间的内容
            content_lines = lines[last_idx + 1:idx]
            content = "\n".join(content_lines).strip()
            if result:
                result[-1]["content"] = content
            result.append({"heading": text, "content": "", "subsections": []})
            last_idx = idx
        # 最后一个标题的内容
        if result:
            content_lines = lines[last_idx + 1:]
            result[-1]["content"] = "\n".join(content_lines).strip()
        return result
    else:
        # 有多个级别，构建嵌套结构
        level_2 = min(levels_present)
        level_3 = max(levels_present)

        result = []
        current_section = None
        current_sub = None
        last_2_idx = 0
        last_3_idx = 0

        for idx, level, text in headings:
            if level == level_2:
                # 二级标题
                # 先把上一个三级的内容补全
                if current_sub and current_section:
                    content_lines = lines[last_3_idx + 1:idx]
                    current_sub["content"] = "\n".join(content_lines).strip()
                    current_section["subsections"].append(current_sub)
                    current_sub = None
                # 把上一个二级的内容补全
                if current_section:
                    content_lines = lines[last_2_idx + 1:idx]
                    current_section["content"] = "\n".join(content_lines).strip()
                    result.append(current_section)
                current_section = {"heading": text, "content": "", "subsections": []}
                last_2_idx = idx
                last_3_idx = idx
            elif level == level_3 and current_section:
                # 三级标题
                if current_sub:
                    content_lines = lines[last_3_idx + 1:idx]
                    current_sub["content"] = "\n".join(content_lines).strip()
                    current_section["subsections"].append(current_sub)
                current_sub = {"heading": text, "content": "", "subsections": []}
                last_3_idx = idx

        # 处理最后一个
        if current_sub and current_section:
            content_lines = lines[last_3_idx + 1:]
            current_sub["content"] = "\n".join(content_lines).strip()
            current_section["subsections"].append(current_sub)
        if current_section:
            if not current_section["content"]:
                content_lines = lines[last_2_idx + 1:]
                current_section["content"] = "\n".join(content_lines).strip()
            result.append(current_section)

        return result


def detect_headings_from_text(pages_text):
    """
    从 PDF 文本中检测标题（当 PDF 没有书签时使用）

    识别规则：
    1. 顶级章节: "Chapter X Title" 或 "第X章 标题"
    2. 二级章节: "X.Y Title"（如 1.1, 1.2, 3.1，支持中文如"1.2 电气安全"）
    3. 三级小节: "X.Y.Z Subtitle"

    排除规则：
    - 目录行（含 "...."/"……"）
    - 单独页码（"1"、"2" 等）和 "Page N"、罗马数字页码
    - 表格项目（数字+单位开头，如 "1.3 ml ..."、"2.2 mL ..."、"220V"）
    - 表格行（含 "Cat. No."、"Quantity"、"型号"、"规格" 等）
    - 常见页眉页脚
    - 重复行（同一行出现多次时只保留第一次）
    """
    import re as re_mod

    # 收集所有行（带全局索引），同时去重
    seen_lines = set()  # 用于去重
    global_lines = []
    for page_idx, page_text in enumerate(pages_text):
        for line in page_text.split('\n'):
            s = line.strip()
            if s and s not in seen_lines:
                seen_lines.add(s)
                global_lines.append(s)

    # 标题模式（支持中英文）
    # 顶级：Chapter X Title（英文）或 第X章 标题（中文）
    chapter_pattern = re_mod.compile(r'^Chapter\s+(\d+)\s+([A-Z\u4e00-\u9fff][A-Za-z\u4e00-\u9fff\s\-]{2,80})$')
    chapter_cn_pattern = re_mod.compile(r'^第([一二三四五六七八九十\d]+)[章节篇]\s+([A-Za-z\u4e00-\u9fff][A-Za-z\u4e00-\u9fff\s\-]{2,80})$')
    # 二级：X.Y Title（支持中文标题，如 "1.2 电气安全"）
    h1_pattern = re_mod.compile(r'^(\d+)\.(\d+)\s+([A-Za-z\u4e00-\u9fff][A-Za-z\u4e00-\u9fff\s\-，、：]{0,80})$')
    # 三级：X.Y.Z Subtitle
    h2_pattern = re_mod.compile(r'^(\d+)\.(\d+)\.(\d+)\s+([A-Za-z\u4e00-\u9fff][A-Za-z0-9\u4e00-\u9fff\s\-\(\)\uff08\uff09]{0,80})$')

    # 表格项目：数字 + 空格 + 单位（ml, mL, μL, L, g, kg 等）
    is_table_item = re_mod.compile(r'^\d+(\.\d+)?\s*(ml|mL|μL|uL|L|g|kg|mg|cm|mm|m|nm|µm|V|W|Hz|A|kW|kV)\b', re_mod.IGNORECASE)
    # 表格行标识（含表格列名，中英文）
    table_row_pattern = re_mod.compile(r'(Cat\.?\s*No\.?|Quantity|Brand|Volume/well|Position\s*$|Tips\s*$|^Consumables\s*$|^Reagents\s*$|^Equipment\s*$|^Specification|Adaptable\s+tube|型号|规格|数量|单位|备注|参数)', re_mod.IGNORECASE)
    # 表格标题行（如 "Table 1-2 ..."、"Figure 3-1 ..."、"表 1-2"、"图 3-1"）
    table_caption_pattern = re_mod.compile(r'^(Table|Figure|表|图)\s*\d+[-\u2013\u2014]\d+', re_mod.IGNORECASE)
    # 页眉页脚（中英文）
    footer_pattern = re_mod.compile(r'^(For Research Use Only\.|MGI Tech Co\.,? Ltd\.|Doc\.?\s*No\.?:?|版权所有|仅供研究使用|技术文档|内部资料)\s*', re_mod.IGNORECASE)
    # 单独页码
    page_num_pattern = re_mod.compile(r'^\d+$|^Page\s*\d+|^第\s*\d+\s*页', re_mod.IGNORECASE)
    # 罗马数字页码
    roman_pattern = re_mod.compile(r'^(I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII)$')

    # 标记哪些行属于表格区（Table/Figure 区域）
    in_table_zone = [False] * len(global_lines)
    i = 0
    while i < len(global_lines):
        line = global_lines[i]
        if table_caption_pattern.match(line):
            # 从这行开始标记为表格区，到下一个标题行或非表格内容行结束
            start = i
            i += 1
            # 找到表格区结束（遇到下一个标题、Table 标题、Chapter 等）
            while i < len(global_lines):
                next_line = global_lines[i]
                # 结束条件：遇到新的标题或长正文
                if (chapter_pattern.match(next_line) or
                    h1_pattern.match(next_line) or
                    h2_pattern.match(next_line) or
                    table_caption_pattern.match(next_line) or
                    (len(next_line) > 60 and not table_row_pattern.search(next_line))):
                    break
                in_table_zone[i] = True
                i += 1
            in_table_zone[start] = True
        else:
            i += 1

    chapters = []
    current_chapter = None
    current_h1 = None  # 当前的 X.Y 章节
    current_section = None  # 当前的 X.Y.Z 小节

    def flush_h1():
        """把当前 H1（X.Y）的所有内容提交到 Chapter 的 subsections"""
        nonlocal current_h1, current_section
        if current_h1 and current_chapter is not None:
            if current_section:
                current_h1["subsections"].append(current_section)
                current_section = None
            current_chapter["subsections"].append(current_h1)
            current_h1 = None

    def flush_chapter():
        """把当前 Chapter 的所有内容提交到 chapters"""
        nonlocal current_chapter, current_h1, current_section
        flush_h1()
        if current_chapter is not None:
            chapters.append(current_chapter)
            current_chapter = None

    for line_idx, line in enumerate(global_lines):
        s = line.strip()
        if not s:
            continue

        # 先检测标题（标题检测优先于表格过滤）
        m_chapter = chapter_pattern.match(s)
        m_chapter_cn = chapter_cn_pattern.match(s)
        m_h1 = h1_pattern.match(s)
        m_h2 = h2_pattern.match(s)

        # 如果是标题，直接处理（不跳过）
        if (m_chapter or m_chapter_cn or m_h1 or m_h2) and len(s) <= 80:
            # 处理标题（见下方）
            pass
        else:
            # 不是标题，检查是否需要跳过
            # 跳过表格区
            if in_table_zone[line_idx]:
                # 把内容加到最近的章节
                if current_section:
                    current_section["content"] += s + " "
                elif current_h1:
                    current_h1["content"] += s + " "
                elif current_chapter:
                    current_chapter["content"] += s + " "
                continue
            # 跳过页眉页脚
            if footer_pattern.match(s):
                continue
            # 跳过页码
            if page_num_pattern.match(s):
                continue
            if roman_pattern.match(s):
                continue
            # 跳过目录行（含省略号）
            if '....' in s or '……' in s:
                continue
            # 跳过表格项目（数字+单位）
            if is_table_item.match(s):
                continue
            # 跳过表格行（含表格列名）- 但标题行已在上方处理，不会被跳过
            if table_row_pattern.search(s):
                continue

        # 额外校验：标题行长度不应过长
        if (m_h1 or m_h2) and len(s) > 80:
            # 作为正文处理
            if current_section:
                current_section["content"] += s + " "
            elif current_h1:
                current_h1["content"] += s + " "
            elif current_chapter:
                current_chapter["content"] += s + " "
            continue

        if m_chapter or m_chapter_cn:
            # 顶级章节（英文 Chapter X 或中文 第X章）
            flush_chapter()
            if m_chapter:
                chapter_num = int(m_chapter.group(1))
            else:
                # 中文数字转换
                cn_num = m_chapter_cn.group(1)
                cn_num_map = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
                chapter_num = cn_num_map.get(cn_num, int(cn_num) if cn_num.isdigit() else 0)
            current_chapter = {
                "heading": s,
                "content": "",
                "subsections": [],
                "_chapter_num": chapter_num
            }
            current_h1 = None
            current_section = None
        elif m_h2:
            # 三级小节（X.Y.Z） - 应该归到当前 X.Y 下
            if current_section and current_h1:
                current_h1["subsections"].append(current_section)
            elif current_section and current_chapter and not current_h1:
                # 没有 X.Y 时，把 X.Y.Z 直接放到 Chapter 下
                current_chapter["subsections"].append(current_section)
            current_section = {"heading": s, "content": "", "subsections": []}
        elif m_h1:
            # 二级章节（X.Y）- 应该归到对应的 Chapter X 下
            chapter_num = int(m_h1.group(1))
            # 如果当前 chapter 不存在或者编号不匹配，创建一个 Chapter 占位
            if not current_chapter or current_chapter.get("_chapter_num") != chapter_num:
                flush_chapter()
                current_chapter = {
                    "heading": f"Chapter {chapter_num}",
                    "content": "",
                    "subsections": [],
                    "_chapter_num": chapter_num
                }
            else:
                # 同 chapter 下的下一个 H1，先 flush 当前的 H1
                flush_h1()
            current_h1 = {"heading": s, "content": "", "subsections": []}
            current_section = None
        else:
            # 正文内容
            if current_section:
                current_section["content"] += s + " "
            elif current_h1:
                current_h1["content"] += s + " "
            elif current_chapter:
                current_chapter["content"] += s + " "

    # 处理最后一个章节
    flush_chapter()

    # 如果检测到的章节太少（< 3），尝试用更多模式补充检测
    if len(chapters) < 3:
        # 尝试用更丰富的模式从全文检测
        all_text = "\n".join(global_lines)
        extra_chapters = _detect_subheadings_from_text(all_text, max_depth=2)
        if len(extra_chapters) > len(chapters):
            chapters = extra_chapters

    return chapters


def flatten_chapters(chapters, parent_heading=""):
    """
    展平章节树为列表，每个子节都带上父章节标题
    如果章节没有子节，就用章节本身；如果有子节，就展开子节
    返回列表：[{"heading": 子节标题, "parent_heading": 父章节标题, "content": 内容, "page_start": 起始页, "page_end": 结束页}, ...]
    """
    flat = []
    for ch in chapters:
        heading = ch.get("heading", "")
        content = ch.get("content", "") or ""
        page_start = ch.get("page_start", 0)
        page_end = ch.get("page_end", 0)
        subs = ch.get("subsections", [])

        if subs:
            # 有子节，展开子节
            for sub in subs:
                sub_heading = sub.get("heading", "")
                sub_content = sub.get("content", "") or ""
                sub_page_start = sub.get("page_start", page_start)
                sub_page_end = sub.get("page_end", page_end)
                flat.append({
                    "heading": sub_heading,
                    "parent_heading": heading,
                    "content": sub_content,
                    "page_start": sub_page_start,
                    "page_end": sub_page_end
                })
                # 递归处理更深层级
                if sub.get("subsections"):
                    sub_flat = flatten_chapters([sub], parent_heading=sub_heading)
                    flat.extend(sub_flat)
        else:
            # 没有子节，用章节本身
            flat.append({
                "heading": heading,
                "parent_heading": parent_heading or heading,
                "content": content,
                "page_start": page_start,
                "page_end": page_end
            })

    return flat


def _align_flat_chapters(flat_a, flat_b):
    """
    对齐两个展平后的子节列表
    优先按标题精确匹配，再按模糊匹配
    """
    from difflib import SequenceMatcher

    aligned = []
    matched_b = set()

    for i, item_a in enumerate(flat_a):
        heading_a = item_a["heading"]
        best_match = None
        best_score = 0

        for j, item_b in enumerate(flat_b):
            if j in matched_b:
                continue
            heading_b = item_b["heading"]
            score = SequenceMatcher(None, heading_a, heading_b).ratio()
            if score > best_score:
                best_score = score
                best_match = j

        if best_match is not None and best_score > 0.5:
            aligned.append({
                "heading_a": heading_a,
                "heading_b": flat_b[best_match]["heading"],
                "parent_heading_a": item_a.get("parent_heading", ""),
                "parent_heading_b": flat_b[best_match].get("parent_heading", ""),
                "content_a": item_a.get("content", ""),
                "content_b": flat_b[best_match].get("content", ""),
                "page_a": item_a.get("page_start", 0),
                "page_b": flat_b[best_match].get("page_start", 0),
                "match_type": "exact" if best_score > 0.9 else "fuzzy",
                "match_score": best_score
            })
            matched_b.add(best_match)
        else:
            aligned.append({
                "heading": heading_a,
                "parent_heading": item_a.get("parent_heading", ""),
                "content": item_a.get("content", ""),
                "page_start": item_a.get("page_start", 0),
                "match_type": "orphan_a",
                "match_score": 0
            })

    for j, item_b in enumerate(flat_b):
        if j not in matched_b:
            aligned.append({
                "heading": item_b["heading"],
                "parent_heading": item_b.get("parent_heading", ""),
                "content": item_b.get("content", ""),
                "page_start": item_b.get("page_start", 0),
                "match_type": "orphan_b",
                "match_score": 0
            })

    return aligned


def align_chapters(chapters_a, chapters_b):
    """
    按标题文本匹配对齐两个文档的章节
    """
    aligned = []
    matched_b = set()

    for ch_a in chapters_a:
        best_match = None
        best_score = 0
        for j, ch_b in enumerate(chapters_b):
            if j in matched_b:
                continue
            score = SequenceMatcher(None, ch_a["heading"], ch_b["heading"]).ratio()
            if score > best_score and score > 0.5:
                best_score = score
                best_match = j

        if best_match is not None:
            aligned.append({
                "heading_a": ch_a["heading"],
                "heading_b": chapters_b[best_match]["heading"],
                "content_a": ch_a["content"],
                "content_b": chapters_b[best_match]["content"],
                "subsections_a": ch_a.get("subsections", []),
                "subsections_b": chapters_b[best_match].get("subsections", []),
                "match_type": "exact" if best_score > 0.9 else "fuzzy",
                "match_score": best_score
            })
            matched_b.add(best_match)
        else:
            aligned.append({
                "heading_a": ch_a["heading"],
                "heading_b": None,
                "content_a": ch_a["content"],
                "content_b": None,
                "match_type": "orphan_a",
                "match_score": 0
            })

    for j, ch_b in enumerate(chapters_b):
        if j not in matched_b:
            aligned.append({
                "heading_a": None,
                "heading_b": ch_b["heading"],
                "content_a": None,
                "content_b": ch_b["content"],
                "match_type": "orphan_b",
                "match_score": 0
            })

    return aligned


def _normalize_text(text, aggressive=False):
    """
    文本规范化：清理 PDF 提取产生的噪声
    1. 去除行尾换行导致的单词拆分（中文单字间的空格）
    2. 合并多余空格
    3. 去除首尾空白
    4. 可选：激进模式（用于匹配比较）- 统一标点、去除多余空格等
    """
    if not text:
        return ""

    # 把换行替换为空格
    text = text.replace('\n', ' ')
    text = text.replace('\r', ' ')
    text = text.replace('\t', ' ')

    # 清理中文单字之间的空格（PDF换行导致的）
    # 匹配：中文字符 + 空格 + 中文字符 的模式，去除中间的空格
    cn_char = r'[\u4e00-\u9fff\uff0c\u3002\u3001\uff1b\uff1a\u201c\u201d\u2018\u2019\uff08\uff09\u3010\u3011]'
    pattern = cn_char + r'(\s+)' + cn_char
    def _remove_space(m):
        return m.group(0).replace(' ', '').replace('\t', '')
    text = re.sub(pattern, _remove_space, text)
    text = re.sub(pattern, _remove_space, text)

    # 英文单词之间保留单个空格（已经通过上面的处理，中文间无空格，英文间保留）

    if aggressive:
        # 激进模式：用于比较时的规范化
        # 统一中英文标点
        text = text.replace('，', ',')
        text = text.replace('。', '.')
        text = text.replace('；', ';')
        text = text.replace('：', ':')
        text = text.replace('（', '(')
        text = text.replace('）', ')')
        text = text.replace('【', '[')
        text = text.replace('】', ']')
        text = text.replace('“', '"')
        text = text.replace('”', '"')
        text = text.replace("'", "'")
        text = text.replace("'", "'")
        # 去除所有空格（包括数字和单位之间的空格）
        text = re.sub(r'\s+', '', text)
        # 转小写
        text = text.lower()
        # 去除句末标点
        text = text.rstrip('.,;:!?')
        # 去除首尾的引号和括号
        text = text.strip('"\'()[]')
        return text

    # 普通模式：合并多余空格
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()

    return text


def _split_sentences(text):
    """
    智能句子切分：
    1. 按句号、问号、感叹号、分号切分（中英文）
    2. 过滤太短（< 5 字符）的片段（页码、噪声）
    3. 保留列表项和正文
    """
    if not text:
        return []

    # 先规范化文本
    text = _normalize_text(text)

    # 按句末标点切分（中英文）
    # 注意：不切分 "1)" "2)" "a)" 等列表项标号，也不切分 "1.3" "X.Y" 等编号
    parts = re.split(r'(?<=[。！？?])\s*|(?<=[;；])\s*|(?<=[.!?])\s+(?=[A-Z])', text)

    sentences = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        # 过滤太短的片段（页码、单字符等）
        if len(p) < 5:
            continue
        # 过滤纯数字/罗马数字
        if re.match(r'^[\dIVXLCDM]+$', p, re.IGNORECASE):
            continue
        # 过滤纯标点
        if re.match(r'^[\s\W]+$', p):
            continue
        sentences.append(p)

    return sentences


def fuzzy_jaccard(sentences_a, sentences_b, threshold=0.70):
    """
    句子级模糊 Jaccard 对比
    返回匹配对数和 Jaccard 值
    
    匹配策略：
    1. 先做激进规范化（统一标点、大小写）后的精确匹配
    2. 再做原文本的精确匹配
    3. 最后做模糊匹配（SequenceMatcher）
    """
    if not sentences_a or not sentences_b:
        return {
            "jaccard": 0.0,
            "coverage": 0.0,
            "avg_score": 1.0,
            "matched_pairs": [],
            "only_a": sentences_a or [],
            "only_b": sentences_b or [],
            "n_sentences_a": len(sentences_a) if sentences_a else 0,
            "n_sentences_b": len(sentences_b) if sentences_b else 0,
            "n_matched": 0
        }

    # 预计算激进规范化后的句子（用于更宽松的精确匹配）
    norm_a = [_normalize_text(s, aggressive=True) for s in sentences_a]
    norm_b = [_normalize_text(s, aggressive=True) for s in sentences_b]

    matched_a = set()
    matched_b = set()
    pairs = []

    # 第1轮：激进规范化后的精确匹配（算 exact）
    # 这种匹配只是标点/大小写不同，内容实质完全一致
    norm_hash_b = {}
    for j, s in enumerate(norm_b):
        if not s:
            continue
        h = hash(s)
        norm_hash_b.setdefault(h, []).append(j)

    for i, s in enumerate(norm_a):
        if not s or i in matched_a:
            continue
        h = hash(s)
        if h in norm_hash_b:
            for j in norm_hash_b[h]:
                if j not in matched_b:
                    matched_a.add(i)
                    matched_b.add(j)
                    pairs.append({"type": "exact", "a": sentences_a[i], "b": sentences_b[j], "score": 1.0})
                    break

    # 第2轮：原文本小写精确匹配
    hash_b = {}
    for j, s in enumerate(sentences_b):
        if j in matched_b:
            continue
        h = hash(s.lower())
        hash_b.setdefault(h, []).append(j)

    for i, s in enumerate(sentences_a):
        if i in matched_a:
            continue
        h = hash(s.lower())
        if h in hash_b:
            for j in hash_b[h]:
                if j not in matched_b:
                    matched_a.add(i)
                    matched_b.add(j)
                    pairs.append({"type": "exact", "a": s, "b": sentences_b[j], "score": 1.0})
                    break

    # 第3轮：模糊匹配（使用规范化后的文本提高准确性）
    candidates = []
    for i in range(len(sentences_a)):
        if i in matched_a:
            continue
        if not norm_a[i]:
            continue
        for j in range(len(sentences_b)):
            if j in matched_b:
                continue
            if not norm_b[j]:
                continue
            # 用规范化后的文本比较，减少标点、空格、大小写的干扰
            score = SequenceMatcher(None, norm_a[i], norm_b[j]).ratio()
            if score >= threshold:
                candidates.append((i, j, score))
    candidates.sort(key=lambda x: -x[2])

    for i, j, score in candidates:
        if i in matched_a or j in matched_b:
            continue
        matched_a.add(i)
        matched_b.add(j)
        pairs.append({"type": "fuzzy", "a": sentences_a[i], "b": sentences_b[j], "score": score})

    only_a = [sentences_a[i] for i in range(len(sentences_a)) if i not in matched_a]
    only_b = [sentences_b[j] for j in range(len(sentences_b)) if j not in matched_b]

    # 计算 Jaccard（改进版）
    # Jaccard = 匹配数 / (A总数 + B总数 - 匹配数)
    union = len(sentences_a) + len(sentences_b) - len(pairs)
    jaccard = len(pairs) / union if union > 0 else 0.0

    # 额外计算：匹配覆盖率（更直观的相似度）
    coverage = len(pairs) / max(len(sentences_a), len(sentences_b), 1)

    # 计算平均匹配分数
    avg_score = sum(p.get("score", 1.0) for p in pairs) / len(pairs) if pairs else 0.0

    return {
        "jaccard": jaccard,
        "coverage": coverage,
        "avg_score": avg_score,
        "matched_pairs": pairs,
        "only_a": only_a,
        "only_b": only_b,
        "n_sentences_a": len(sentences_a),
        "n_sentences_b": len(sentences_b),
        "n_matched": len(pairs)
    }


def compare_documents_by_format(file_a_path, file_b_path):
    """
    自动识别格式并对比两个文档

    支持：.docx / .doc / .pdf / .md / .txt
    """
    ext_a = file_a_path.split('.')[-1].lower().replace(' ', '')
    ext_b = file_b_path.split('.')[-1].lower().replace(' ', '')

    # 选择解析器
    parsed_a = None
    parsed_b = None

    if ext_a in ('docx', 'doc'):
        parsed_a = parse_docx(file_a_path)
    elif ext_a == 'pdf':
        parsed_a = parse_pdf(file_a_path)
    elif ext_a in ('md', 'markdown'):
        parsed_a = parse_markdown(file_a_path)
    else:
        # 尝试纯文本
        try:
            with open(file_a_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            parsed_a = {"chapters": [{"heading": "全文", "content": content, "subsections": []}], "full_text": content, "type": "text"}
        except Exception:
            parsed_a = {"chapters": [], "full_text": "", "type": "unknown"}

    if ext_b in ('docx', 'doc'):
        parsed_b = parse_docx(file_b_path)
    elif ext_b == 'pdf':
        parsed_b = parse_pdf(file_b_path)
    elif ext_b in ('md', 'markdown'):
        parsed_b = parse_markdown(file_b_path)
    else:
        try:
            with open(file_b_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            parsed_b = {"chapters": [{"heading": "全文", "content": content, "subsections": []}], "full_text": content, "type": "text"}
        except Exception:
            parsed_b = {"chapters": [], "full_text": "", "type": "unknown"}

    # 展平章节为子节列表
    chapters_a = parsed_a.get("chapters", [])
    chapters_b = parsed_b.get("chapters", [])
    flat_a = flatten_chapters(chapters_a)
    flat_b = flatten_chapters(chapters_b)

    # 统计文件信息
    n_sentences_a = 0
    n_sentences_b = 0
    for item in flat_a:
        n_sentences_a += len(_split_sentences(item.get("content", "")))
    for item in flat_b:
        n_sentences_b += len(_split_sentences(item.get("content", "")))

    # 子节对齐
    aligned = _align_flat_chapters(flat_a, flat_b)

    # 句子级对比
    results = []
    n_full = 0
    n_high = 0
    n_partial = 0
    n_low = 0
    n_only_a = 0
    n_only_b = 0

    for item in aligned:
        if item["match_type"] == "orphan_a":
            results.append({
                "heading": item["heading"],
                "parent_heading": item.get("parent_heading", ""),
                "content_a": item.get("content", ""),
                "content_b": "",
                "page_a": item.get("page_start", 0),
                "page_b": 0,
                "status": "仅在A中",
                "similarity": 0,
                "diffs": [],
                "only_a": _split_sentences(item.get("content", "")),
                "only_b": []
            })
            n_only_a += 1
        elif item["match_type"] == "orphan_b":
            results.append({
                "heading": item["heading"],
                "parent_heading": item.get("parent_heading", ""),
                "content_a": "",
                "content_b": item.get("content", ""),
                "page_a": 0,
                "page_b": item.get("page_start", 0),
                "status": "仅在B中",
                "similarity": 0,
                "diffs": [],
                "only_a": [],
                "only_b": _split_sentences(item.get("content", ""))
            })
            n_only_b += 1
        else:
            content_a = item.get("content_a", "") or ""
            content_b = item.get("content_b", "") or ""

            # 使用智能句子切分
            sentences_a = _split_sentences(content_a)
            sentences_b = _split_sentences(content_b)
            n_a = len(sentences_a)
            n_b = len(sentences_b)

            # 特殊情况：两边都没有内容，视为完全一致
            if n_a == 0 and n_b == 0:
                sim = 1.0
                avg_score = 1.0
                coverage = 1.0
                n_matched_pairs = 0
                result = {"matched_pairs": [], "only_a": [], "only_b": [], "n_matched": 0}
                status = "完全一致"
                n_full += 1
            elif n_a == 0 or n_b == 0:
                # 一边有内容一边没有
                sim = 0.0
                avg_score = 0.0
                coverage = 0.0
                n_matched_pairs = 0
                result = {"matched_pairs": [], "only_a": sentences_a, "only_b": sentences_b, "n_matched": 0}
                status = "差异较大"
                n_low += 1
            else:
                result = fuzzy_jaccard(sentences_a, sentences_b, threshold=0.70)

                # 相似度 = 覆盖率 × 平均匹配分数（更准确反映真实相似度）
                coverage = result["coverage"]
                avg_score = result.get("avg_score", 1.0)
                n_matched_pairs = result.get("n_matched", 0)
                sim = coverage * avg_score

                # 统计 exact 匹配数量
                n_exact = sum(1 for p in result.get("matched_pairs", []) if p.get("type") == "exact")
                n_fuzzy = n_matched_pairs - n_exact

                # 状态判定
                # 完全一致：所有句子都精确匹配，覆盖率100%，没有fuzzy匹配
                if n_fuzzy == 0 and coverage >= 0.99 and n_a == n_b:
                    status = "完全一致"
                    n_full += 1
                # 高度相似：只有少量措辞或句子差异（相似度>=85%）
                elif sim >= 0.85:
                    status = "高度相似"
                    n_high += 1
                # 部分相似：有部分内容变更（相似度>=60%）
                elif sim >= 0.60:
                    status = "部分相似"
                    n_partial += 1
                else:
                    status = "差异较大"
                    n_low += 1

            results.append({
                "heading": item["heading_a"],
                "parent_heading": item.get("parent_heading_a", ""),
                "content_a": content_a,
                "content_b": content_b,
                "page_a": item.get("page_a", 0),
                "page_b": item.get("page_b", 0),
                "status": status,
                "similarity": sim,
                "avg_score": avg_score,
                "coverage": coverage,
                "n_matched": n_matched_pairs,
                "diffs": result["matched_pairs"],
                "only_a": result["only_a"],
                "only_b": result["only_b"]
            })

    # 计算整体指标
    n_matched = len([r for r in results if r["status"] not in ("仅在A中", "仅在B中")])
    total = n_matched + n_only_a + n_only_b
    overall_sim = sum(r["similarity"] for r in results if r["status"] not in ("仅在A中", "仅在B中")) / max(n_matched, 1)
    matched_avg = overall_sim

    # 加权一致性（按句子数加权）
    total_sentences = n_sentences_a + n_sentences_b
    weighted_sim = 0.0
    if total_sentences > 0:
        weighted_sum = 0.0
        for r in results:
            if r["status"] in ("仅在A中", "仅在B中"):
                continue
            n_s = len(_split_sentences(r.get("content_a", "")))
            weighted_sum += r["similarity"] * n_s
        weighted_sim = weighted_sum / max(n_sentences_a, 1) if n_sentences_a > 0 else 0

    return {
        "results": results,
        "stats": {
            "n_full": n_full,
            "n_high": n_high,
            "n_partial": n_partial,
            "n_low": n_low,
            "n_only_a": n_only_a,
            "n_only_b": n_only_b,
            "n_matched": n_matched,
            "overall_sim": overall_sim,
            "matched_avg": matched_avg,
            "weighted_sim": weighted_sim,
            "total": total,
            "n_sentences_a": n_sentences_a,
            "n_sentences_b": n_sentences_b,
            "n_topics_a": len(flat_a),
            "n_topics_b": len(flat_b)
        },
        "type_a": parsed_a.get("type", "unknown"),
        "type_b": parsed_b.get("type", "unknown"),
        "aligned_chapters": aligned
    }
