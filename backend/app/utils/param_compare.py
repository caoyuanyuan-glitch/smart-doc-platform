import re
import os
from difflib import SequenceMatcher


def normalize_param_name(name):
    name = name.strip()
    name = re.sub(r'[\s\u3000]+', ' ', name)
    name = re.sub(r'[（(][^)）]*[)）]$', '', name)
    return name


def parse_value(raw):
    raw = raw.strip()
    if not raw:
        return raw
    raw = re.sub(r'\s*[±＋]\s*', '±', raw)
    raw = re.sub(r'\s*[×xX]\s*', '×', raw)
    raw = re.sub(r'[−‐‑‒–—―]', '-', raw)
    return raw


_COLON_PATTERN = re.compile(
    r'^\s*'
    r'(?:[\d]+[\.\、\)）]\s*)?'          # 可选编号前缀
    r'([\u4e00-\u9fff\w][\u4e00-\u9fff\w \t/\-\.·\u2103\u2109\u00b0\u03a9\u03bc\u00b5\u2126]{1,60}?)'
    r'\s*[：:]\s*'
    r'(.+)$'
)

_EQUALS_PATTERN = re.compile(
    r'^\s*'
    r'(?:[\d]+[\.\、\)）]\s*)?'
    r'([\u4e00-\u9fff\w][\u4e00-\u9fff\w \t/\-\.·\u2103\u2109\u00b0\u03a9\u03bc\u00b5\u2126]{1,60}?)'
    r'\s*[=＝]\s*'
    r'(.+)$'
)

_TABLE_LIKE_PATTERN = re.compile(
    r'^\s*'
    r'(?:[\d]+[\.\、\)）]\s*)?'
    r'([\u4e00-\u9fff\w][\u4e00-\u9fff\w \t/\-\.·\u2103\u2109\u00b0\u03a9\u03bc\u00b5\u2126]{1,50}?)'
    r'(?:\s{2,}|\t)\s*'
    r'(.+)$'
)

# OCR 特殊模式：中文参数名直接连接字母数字值（如 "额定功率300W"、"产品型号HTS-100"）
_OCR_CONCAT_PATTERN = re.compile(
    r'^\s*'
    r'([\u4e00-\u9fff][\u4e00-\u9fff\s]{1,40}?)'
    r'([\dA-Za-z\u2103\u2109\u00b0\u03a9\u03bc\u00b5\u2126\u00b1\u2264\u2265][^。,，；;]{0,60})$'
)

_NUMERIC_PATTERN = re.compile(r'[-+]?\d+\.?\d*')
_UNIT_PATTERN = re.compile(
    r'(?:mm|cm|m|km|mm\^?[23]|cm\^?[23]|m\^?[23]|'
    r'kg|g|mg|t|ton|'
    r's|ms|min|h|hour|sec|'
    r'V|kV|mV|A|mA|W|kW|MW|Hz|kHz|MHz|GHz|'
    r'℃|°C|°F|℉|℃|℉|K|'
    r'Ω|kΩ|MΩ|μF|pF|nF|'
    r'L|mL|m\^3|L/min|'
    r'Pa|kPa|MPa|bar|psi|'
    r'N|kN|Nm|rpm|RPM|'
    r'%|％|ppm|ppb|'
    r'dB|dBm|dBi|'
    r'bps|kbps|Mbps|Gbps|'
    r'px|pt|em|rem|'
    r'[\u2103\u2109\u00b0\u03a9\u03bc\u00b5\u2126])'
)


def _is_param_like(text):
    text = text.strip()
    if len(text) > 80:
        return False
    if re.match(r'^[\d\.\-\+\s,;，；·]+$', text):
        return False
    if not re.search(r'[\u4e00-\u9fff\w]', text):
        return False
    if re.match(r'^(第|Chapter|Section|Table|图|Figure|注|Note|参见|参考|详见|如|例如|其中|式中)\s', text):
        return False
    if re.match(r'^(目\s*录|目录|前言|引言|概述|简介|总结|参考文献|附录|致谢)', text):
        return False
    if re.match(r'^\d+\.?\d*\s+', text):
        return False
    # 排除长句子（含中文标点的多词文本）
    if len(text) > 20 and re.search(r'[，。、；,]', text):
        return False
    # 排除含中文句号的文本（OCR 多行拼接）
    if re.search(r'[。]', text):
        return False
    # 排除以动词/介词开头的非参数文本
    if re.match(r'^.*(将|按|如|根据|检查|重新|联系|停止)', text) and len(text) > 10:
        return False
    # 排除章节词混杂（如"常见问题处理设备无法开机"）
    if re.match(r'^(常见问题|安全注意|清洁与维护|产品介绍|产品组成|主要技术|操作说明|使用环境|储存与运输)', text):
        return False
    return True


def _is_value_like(text):
    text = text.strip()
    if not text:
        return False
    if len(text) > 200:
        return False
    if re.search(r'[。；，；\n]{2,}', text):
        return False
    words = text.split()
    if len(words) > 20:
        return False
    # 排除纯数字编号（如 "6.2"、"6.4"、"3."）
    if re.match(r'^\d+\.?\d*\s*(开机|设置|停止|启动|清洁|安全|操作|产品|使用|储存)?$', text):
        return False
    has_digit = bool(_NUMERIC_PATTERN.search(text))
    has_unit = bool(_UNIT_PATTERN.search(text.casefold()))
    is_short = len(text) <= 30
    # 值不包含数字且不包含单位时，必须是短文本
    if not has_digit and not has_unit:
        if len(text) > 15:
            return False
        if re.search(r'[。，、；]', text):
            return False
    if has_digit or has_unit:
        return True
    if is_short and not re.search(r'[。，、；]', text):
        return True
    return False


def extract_params_from_text(text):
    params = []
    seen = set()

    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        source = "text"

        m = _COLON_PATTERN.match(line)
        if not m:
            m = _EQUALS_PATTERN.match(line)
        if not m:
            m = _TABLE_LIKE_PATTERN.match(line)
        if not m:
            m = _OCR_CONCAT_PATTERN.match(line)
            if m:
                source = "ocr_concat"
        if not m:
            continue

        name = normalize_param_name(m.group(1))
        value = parse_value(m.group(2).strip())
        if not name or not value:
            continue
        if not _is_param_like(name):
            continue
        if not _is_value_like(value):
            continue

        # OCR 拼接模式额外验证：值必须含数字或单位
        if source == "ocr_concat":
            if not (_NUMERIC_PATTERN.search(value) or _UNIT_PATTERN.search(value.casefold())):
                continue

        norm_name = name.casefold()
        if norm_name in seen:
            continue
        seen.add(norm_name)
        params.append({"name": name, "value": value, "source": source})
    return params


def _is_header_row(cells):
    """判断一行是否为表头（列标题行）

    表头特征：单元格全是短中文词，不含数字/单位，通常为"参数名称"、"规格"、"数值"等
    """
    if not cells:
        return False
    header_keywords = [
        '参数', '名称', '项目', '指标', '规格', '型号', '数值', '值', '单位',
        '要求', '标准', '范围', '条件', '说明', '备注', '描述', '类型', '序号',
        '编号', '代码', 'name', 'value', 'unit', 'spec', 'param', 'item',
    ]
    numeric_count = 0
    for cell in cells:
        cell_text = cell.strip()
        if not cell_text:
            continue
        if _NUMERIC_PATTERN.search(cell_text) or _UNIT_PATTERN.search(cell_text.casefold()):
            numeric_count += 1
        for kw in header_keywords:
            if kw in cell_text:
                return True
    if numeric_count == 0 and len(cells) >= 2:
        return True
    return False


def _detect_header_row_idx(rows, max_check=5):
    """自动检测表头行索引，返回第一个疑似表头的行号"""
    for idx in range(min(len(rows), max_check)):
        cells = [c.strip() for c in rows[idx] if c.strip()]
        if _is_header_row(cells):
            return idx
    return -1


def extract_params_from_table_rows(rows, header_row_idx=0):
    params = []
    seen = set()

    if not rows or len(rows) < 1:
        return params

    skip_idx = max(header_row_idx, _detect_header_row_idx(rows))

    for row_idx, row in enumerate(rows):
        if row_idx == skip_idx:
            continue
        cells = [c.strip() for c in row if c.strip()]
        if len(cells) < 2:
            continue

        for i in range(len(cells) - 1):
            name = cells[i]
            value = cells[i + 1]

            if not name or not value:
                continue
            if not _is_param_like(name):
                name2 = cells[i + 1] if i + 2 < len(cells) else ""
                if not name2:
                    continue
                continue

            name = normalize_param_name(name)
            value = parse_value(value)

            norm_name = name.casefold()
            if norm_name in seen:
                continue

            if not _is_value_like(value):
                continue

            seen.add(norm_name)
            params.append({"name": name, "value": value, "source": f"table_row_{row_idx}"})

    return params


def extract_params_from_docx(filepath):
    try:
        from docx import Document
    except ImportError:
        return []

    try:
        doc = Document(filepath)
    except Exception:
        return []

    text_parts = []
    table_rows = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(cells)

    full_text = "\n".join(text_parts)
    text_params = extract_params_from_text(full_text)
    table_params = extract_params_from_table_rows(table_rows)

    merged = {}
    for p in text_params:
        merged[p["name"].casefold()] = p
    for p in table_params:
        key = p["name"].casefold()
        if key not in merged:
            merged[key] = p

    result = list(merged.values())
    result.sort(key=lambda x: x["name"])
    return result


def extract_params_from_pdf(filepath):
    text_parts = []
    table_rows = []

    # 主提取器：PyMuPDF (fitz) 对 CJK 支持最好
    fitz_text = ""
    try:
        import fitz
        doc = fitz.open(filepath)
        print(f"[param_compare] PyMuPDF: PDF has {len(doc)} pages")
        for i, page in enumerate(doc):
            t = page.get_text("text")
            if t and t.strip():
                text_parts.append(t)
                print(f"[param_compare] PyMuPDF Page {i+1} text length: {len(t)}")
            else:
                t2 = page.get_text("blocks")
                if t2:
                    block_texts = []
                    for b in t2:
                        if len(b) >= 5 and b[4].strip():
                            block_texts.append(b[4].strip())
                    if block_texts:
                        combined = "\n".join(block_texts)
                        text_parts.append(combined)
                        print(f"[param_compare] PyMuPDF Page {i+1} blocks text length: {len(combined)}")
        doc.close()
        fitz_text = "\n".join(text_parts)
        print(f"[param_compare] PyMuPDF total text length: {len(fitz_text)}")
    except ImportError:
        print("[param_compare] PyMuPDF (fitz) not available")
    except Exception as e:
        print(f"[param_compare] PyMuPDF error: {e}")

    # 如果 PyMuPDF 提取到了文本，用它
    if fitz_text.strip():
        text_params = extract_params_from_text(fitz_text)
        multiline_params = _extract_params_multiline(fitz_text)
        print(f"[param_compare] text_params: {len(text_params)}, multiline_params: {len(multiline_params)}")
    else:
        text_params = []
        multiline_params = []

    # 始终尝试用 pdfplumber 提取文本和表格（补充 PyMuPDF 可能遗漏的表格数据）
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            print(f"[param_compare] pdfplumber supplement: PDF has {len(pdf.pages)} pages")
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if tables:
                    print(f"[param_compare] pdfplumber Page {i+1} has {len(tables)} tables")
                for tbl in tables:
                    for row in tbl:
                        if row and any(c for c in row if c):
                            table_rows.append([(c or "").strip() for c in row])
                if not fitz_text.strip():
                    t = page.extract_text()
                    if t and t.strip():
                        text_parts.append(t)
                        print(f"[param_compare] pdfplumber Page {i+1} text length: {len(t)}")
        if not fitz_text.strip():
            full_text = "\n".join(text_parts)
            print(f"[param_compare] pdfplumber total text length: {len(full_text)}")
            if full_text.strip():
                text_params = extract_params_from_text(full_text)
                multiline_params = _extract_params_multiline(full_text)
    except ImportError:
        print("[param_compare] pdfplumber not available for supplement")
    except Exception as e:
        print(f"[param_compare] pdfplumber supplement error: {e}")

    # 回退到 pypdf
    if not fitz_text.strip() and not text_params and not multiline_params:
        try:
            from pypdf import PdfReader
            reader = PdfReader(filepath)
            text_parts2 = []
            print(f"[param_compare] pypdf: PDF has {len(reader.pages)} pages")
            for page in reader.pages:
                t = page.extract_text()
                if t and t.strip():
                    text_parts2.append(t)
            full_text = "\n".join(text_parts2)
            print(f"[param_compare] pypdf total text length: {len(full_text)}")
            text_params = extract_params_from_text(full_text)
            multiline_params = _extract_params_multiline(full_text)
        except ImportError:
            print("[param_compare] pypdf not available")
        except Exception as e:
            print(f"[param_compare] pypdf error: {e}")

    # OCR 回退：扫描件/图片型PDF（无文本层时自动触发）
    if not fitz_text.strip() and not text_params and not multiline_params:
        ocr_text = _ocr_pdf_pages(filepath)
        if ocr_text.strip():
            print(f"[param_compare] OCR text length: {len(ocr_text)}")
            text_params = extract_params_from_text(ocr_text)
            multiline_params = _extract_params_multiline(ocr_text)

    # 提取表格参数
    table_params = extract_params_from_table_rows(table_rows)
    print(f"[param_compare] table_params: {len(table_params)}")

    # 合并：单行 > 多行 > 表格，同名优先保留先出现的
    merged = {}
    for p in text_params:
        merged[p["name"].casefold()] = p
    for p in multiline_params:
        key = p["name"].casefold()
        if key not in merged:
            merged[key] = p
    for p in table_params:
        key = p["name"].casefold()
        if key not in merged:
            merged[key] = p

    result = list(merged.values())
    result.sort(key=lambda x: x["name"])
    print(f"[param_compare] PDF total merged params: {len(result)}")
    return result


def _ocr_pdf_pages(filepath):
    """对扫描件/图片型 PDF 进行 OCR 识别，返回提取的文本"""
    text_parts = []
    try:
        import fitz
        import io
        from PIL import Image, ImageFilter, ImageEnhance

        doc = fitz.open(filepath)
        zoom = 3.0
        mat = fitz.Matrix(zoom, zoom)

        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            img = img.convert('L')
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(1.5)
            img = img.filter(ImageFilter.SHARPEN)
            page_text = _ocr_image(img)
            if page_text and page_text.strip():
                text_parts.append(page_text)
                print(f"[param_compare] OCR Page {i+1}: {len(page_text)} chars")

        doc.close()
    except ImportError as e:
        print(f"[param_compare] OCR: required library not available - {e}")
    except Exception as e:
        print(f"[param_compare] OCR error: {e}")

    return "\n".join(text_parts)


def _ocr_image(pil_image):
    """对单张 PIL Image 进行 OCR (tesserocr > pytesseract)，并对结果进行清洗"""
    import os

    _TESSDATA_CANDIDATES = [
        '/usr/share/tesseract-ocr/5/tessdata',
        '/usr/share/tesseract-ocr/4.00/tessdata',
        '/usr/share/tessdata',
        os.environ.get('TESSDATA_PREFIX', ''),
    ]

    # 尝试 tesserocr（直接 C API，速度快）
    try:
        import tesserocr
        tessdata_path = None
        for candidate in _TESSDATA_CANDIDATES:
            if candidate and os.path.isdir(candidate):
                tessdata_path = candidate
                break
        api = tesserocr.PyTessBaseAPI(path=tessdata_path, lang='chi_sim+eng')
        api.SetImage(pil_image)
        text = api.GetUTF8Text()
        api.End()
        if text and text.strip():
            print(f"[param_compare] tesserocr OCR success: {len(text)} chars")
            return _clean_ocr_text(text)
        else:
            print("[param_compare] tesserocr OCR returned empty text")
    except ImportError:
        print("[param_compare] tesserocr not available, trying pytesseract")
    except Exception as e:
        print(f"[param_compare] tesserocr OCR error: {e}, trying pytesseract")

    # 回退到 pytesseract（命令行 wrapper，需安装 tesseract-ocr）
    try:
        import pytesseract
        for candidate in _TESSDATA_CANDIDATES:
            if candidate and os.path.isdir(candidate):
                os.environ['TESSDATA_PREFIX'] = candidate
                break
        text = pytesseract.image_to_string(pil_image, lang='chi_sim+eng')
        if text and text.strip():
            print(f"[param_compare] pytesseract OCR success: {len(text)} chars")
            return _clean_ocr_text(text)
        else:
            print("[param_compare] pytesseract OCR returned empty text")
    except ImportError:
        print("[param_compare] pytesseract not available for OCR")
    except Exception as e:
        print(f"[param_compare] pytesseract OCR error: {e}")

    return ""


def _clean_ocr_text(text):
    """清理 OCR 文本：去除中文间多余空格、修复常见 OCR 错误"""
    import re
    # 去除中文（含中日韩）字符之间的空格（迭代清理，直到无变化）
    while True:
        cleaned = re.sub(
            r'([\u4e00-\u9fff\u3400-\u4dbf])\s+([\u4e00-\u9fff\u3400-\u4dbf])',
            r'\1\2', text
        )
        if cleaned == text:
            break
        text = cleaned
    # 去除中文与数字/单位之间的多余空格
    text = re.sub(
        r'([\u4e00-\u9fff])\s+(\d)',
        r'\1\2', text
    )
    text = re.sub(
        r'(\d)\s+([\u4e00-\u9fff])',
        r'\1\2', text
    )
    # 去除中文与字母/单位之间的多余空格
    text = re.sub(
        r'([\u4e00-\u9fff])\s+([a-zA-Z])',
        r'\1\2', text
    )
    text = re.sub(
        r'([a-zA-Z])\s+([\u4e00-\u9fff])',
        r'\1\2', text
    )

    # 修复字母/单位间多余空格（如 "HTS- 100" → "HTS-100"）
    text = re.sub(r'([A-Z]{2,})\s+-', r'\1-', text)
    text = re.sub(r'-\s+(\d)', r'-\1', text)

    # 常见 OCR 数字/符号错误修复
    text = re.sub(r'\b8096\b', '80%', text)
    text = re.sub(r'\b2096\b', '20%', text)
    text = re.sub(r'\b109%\b', '10%', text)
    text = re.sub(r'\b1096\b', '10%', text)
    text = re.sub(r'\b90%6\b', '90%', text)
    text = re.sub(r'\b1006\b', '100', text)
    # OCR 将百分比数字拆开：200%6 → 20%, 809%6 → 80%
    text = re.sub(r'200%6', '20%', text)
    text = re.sub(r'809%6', '80%', text)
    text = re.sub(r'10%6', '10%', text)
    text = re.sub(r'90%6', '90%', text)
    # 也处理末尾带6的变体
    text = re.sub(r'200%\b', '20%', text)
    text = re.sub(r'809%\b', '80%', text)

    # 百分比 OCR 误识别修复
    text = re.sub(r'(\d{1,3})%(\d)', r'\1%\2', text)

    # 修复 OCR 把 ~ 识别成其他字符
    text = re.sub(r'[一—―]', '~', text)

    # 修复常见 OCR 单位/符号错误
    text = re.sub(r'\bKE\b', 'kg', text)
    text = re.sub(r'\bke\b', 'kg', text)
    text = re.sub(r'kKPa', 'kPa', text)
    text = re.sub(r'逞0O', '200', text)
    text = re.sub(r'\b([\d])O([\d])', r'\1O\2', text)  # 谨慎处理

    # 修复常见中文 OCR 错误
    corrections = {
        '温度控制范图': '温度控制范围',
        '定时范图': '定时范围',
        '分劲率': '分辨率',
        '均匀性万': '均匀性≤',
        '圆周振济': '圆周振荡',
        '控制精度土': '控制精度±',
        'HTS-100恒温热振烫器': 'HTS-100恒温热振荡器',
        '热振烫器': '热振荡器',
    }
    for wrong, right in corrections.items():
        text = text.replace(wrong, right)

    # 修复 ± 符号（土 → ±）
    text = re.sub(r'(?<=[\d\s])土(?=[\d])', '±', text)
    # 修复 ≤ 符号（万 → ≤）
    text = re.sub(r'(?<=[性\s])万(?=[土±\d])', '≤', text)
    text = re.sub(r'([<>])万(?=[土±\d])', r'\1≤', text)

    # 拆分 OCR 拼接的参数行：在单位/数值后紧跟参数名称处插入换行
    param_keywords = [
        '电源', '额定功率', '额定电压', '额定频率', '输入功率',
        '外形尺寸', '净重', '毛重',
        '温度控制范围', '温度控制精度', '温度显示分辨率', '温度均匀性',
        '温度上升时间', '温度下降时间', '升温时间', '降温时间',
        '振荡方式', '振荡转速', '振荡转速范围', '转速调节间隔', '转速调节步进',
        '振荡幅度', '振荡频率', '定时范围', '定时时间',
        '适配模块', '模块材料', '模块更换方式',
        '产品型号', '产品名称', '工作条件', '操作环境',
        '环境温度', '相对湿度', '大气压力',
        '使用环境', '储存温度', '储存湿度', '运输温度', '运输湿度',
        '连续运行', '文件编号', '版本', '关键参数',
    ]
    for kw in sorted(param_keywords, key=len, reverse=True):
        text = re.sub(
            r'([WVAHzΩkgmLCsminrp℃°\u2103\u2109%％\d])\s*' + re.escape(kw),
            r'\1\n' + kw, text
        )
        # 也处理中文→中文参数名的边界（如 "约8.5kg温度控制" → "约8.5kg\n温度控制"）
        text = re.sub(
            r'([\u4e00-\u9fff])\s*' + re.escape(kw),
            r'\1\n' + kw, text
        )

    return text


def _extract_params_multiline(text):
    """提取参数名和值在不同行的情况，如：
    额定电压
    220V
    """
    params = []
    seen = set()
    lines = [l.strip() for l in text.split('\n')]
    lines = [l for l in lines if l and len(l) > 0]

    if len(lines) < 2:
        return params

    i = 0
    while i < len(lines) - 1:
        name_line = lines[i]
        value_line = lines[i + 1]

        if _COLON_PATTERN.match(value_line) or _EQUALS_PATTERN.match(value_line):
            i += 1
            continue

        # 参数名不应过长，且不应包含数字（纯中文/英文关键词）
        if _is_param_like(name_line) and _is_value_like(value_line):
            if not _NUMERIC_PATTERN.search(name_line) and len(name_line) <= 7:
                name = normalize_param_name(name_line)
                value = parse_value(value_line)
                norm_name = name.casefold()
                if norm_name not in seen:
                    seen.add(norm_name)
                    params.append({"name": name, "value": value, "source": "text_ml"})
                i += 2
                continue
        i += 1

    return params


def extract_params_from_xlsx(filepath):
    try:
        from openpyxl import load_workbook
    except ImportError:
        return []

    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
    except Exception:
        return []

    all_params = {}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            str_cells = [(str(c).strip() if c is not None else "") for c in row]
            if any(str_cells):
                rows.append(str_cells)

        sheet_params = extract_params_from_table_rows(rows)
        for p in sheet_params:
            key = p["name"].casefold()
            if key not in all_params:
                p["source"] = f"sheet_{sheet_name}"
                all_params[key] = p

    wb.close()
    result = list(all_params.values())
    result.sort(key=lambda x: x["name"])
    return result


def _extract_raw_text(filepath):
    """提取文件中的原始文本（用于调试）"""
    ext = os.path.splitext(filepath)[1].lower()

    if ext in ('.docx', '.doc'):
        try:
            from docx import Document
            doc = Document(filepath)
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        except Exception:
            return ""

    if ext == '.pdf':
        parts = []
        try:
            import fitz
            doc = fitz.open(filepath)
            for page in doc:
                t = page.get_text("text")
                if t and t.strip():
                    parts.append(t)
                else:
                    blocks = page.get_text("blocks")
                    for b in blocks:
                        if len(b) >= 5 and b[4].strip():
                            parts.append(b[4].strip())
            doc.close()
            if parts:
                return "\n---PAGE---\n".join(parts)
        except Exception:
            pass

        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t and t.strip():
                        parts.append(t)
            if parts:
                return "\n---PAGE---\n".join(parts)
        except Exception:
            pass

        ocr_text = _ocr_pdf_pages(filepath)
        if ocr_text.strip():
            return ocr_text

        return ""

    if ext in ('.xlsx', '.xls'):
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath, read_only=True, data_only=True)
            parts = []
            for name in wb.sheetnames:
                ws = wb[name]
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c).strip() if c is not None else "" for c in row]
                    if any(cells):
                        parts.append(" | ".join(cells))
            wb.close()
            return "\n".join(parts)
        except Exception:
            return ""

    if ext in ('.txt', '.md'):
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception:
            return ""

    return ""


def extract_params(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ('.docx', '.doc'):
        return extract_params_from_docx(filepath)
    elif ext == '.pdf':
        return extract_params_from_pdf(filepath)
    elif ext in ('.xlsx', '.xls'):
        return extract_params_from_xlsx(filepath)
    elif ext in ('.txt', '.md'):
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            return sorted(extract_params_from_text(text), key=lambda x: x["name"])
        except Exception:
            return []
    else:
        return []


def _fuzzy_match_key(key, candidates, threshold=0.70):
    best_score = 0
    best_key = None
    key_lower = key.casefold()
    for cand in candidates:
        cand_lower = cand.casefold()
        if key_lower == cand_lower:
            return 1.0, cand
        sim = SequenceMatcher(None, key_lower, cand_lower).ratio()
        if sim < threshold:
            edit_sim = _edit_similarity(key_lower, cand_lower)
            sim = max(sim, edit_sim)
        if sim < threshold:
            sim = max(sim, _common_prefix_similarity(key_lower, cand_lower))
        if sim > best_score:
            best_score = sim
            best_key = cand
    if best_score >= threshold:
        return best_score, best_key
    return 0, None


def _common_prefix_similarity(s1, s2):
    """公共前缀加权相似度，用于匹配同义词（如"使用地点"≈"使用场所"）"""
    n = min(len(s1), len(s2))
    prefix_len = 0
    for i in range(n):
        if s1[i] == s2[i]:
            prefix_len += 1
        else:
            break
    if prefix_len < 2:
        return 0
    prefix_ratio = prefix_len / max(len(s1), len(s2))
    if prefix_ratio < 0.3:
        return 0
    return 0.60 + prefix_ratio * 0.25


def _edit_similarity(s1, s2):
    """字符级编辑距离相似度，对中文短文本（<=15字）更敏感"""
    if not s1 or not s2:
        return 0
    if len(s1) > 15 or len(s2) > 15:
        return 0
    max_len = max(len(s1), len(s2))
    if max_len == 0:
        return 1.0
    d = [[0] * (len(s2) + 1) for _ in range(len(s1) + 1)]
    for i in range(len(s1) + 1):
        d[i][0] = i
    for j in range(len(s2) + 1):
        d[0][j] = j
    for i in range(1, len(s1) + 1):
        for j in range(1, len(s2) + 1):
            cost = 0 if s1[i-1] == s2[j-1] else 1
            d[i][j] = min(d[i-1][j] + 1, d[i][j-1] + 1, d[i-1][j-1] + cost)
    return 1.0 - d[len(s1)][len(s2)] / max_len


def _values_equal(v1, v2):
    v1c = v1.strip().casefold()
    v2c = v2.strip().casefold()
    if v1c == v2c:
        return True

    nums1 = _NUMERIC_PATTERN.findall(v1)
    nums2 = _NUMERIC_PATTERN.findall(v2)

    if nums1 and nums2:
        try:
            n1 = float(nums1[0])
            n2 = float(nums2[0])
            return abs(n1 - n2) < 1e-9
        except ValueError:
            pass

    sim = SequenceMatcher(None, v1c, v2c).ratio()
    return sim >= 0.98


def _translate_source(source):
    """将内部 source 编码翻译为可读的中文标签"""
    if not source:
        return ""
    if source == "text":
        return "正文"
    if source == "text_ml":
        return "正文(多行)"
    if source == "ocr_concat":
        return "OCR提取"
    if source.startswith("table_row_"):
        n = source.replace("table_row_", "")
        return f"表格第{n}行" if n.isdigit() else f"表格行"
    if source.startswith("sheet_"):
        name = source.replace("sheet_", "")
        return f"工作表「{name}」"
    return source


def compare_params(params_a, params_b):
    keys_a = {p["name"].casefold(): p for p in params_a}
    keys_b = {p["name"].casefold(): p for p in params_b}

    matched_b = set()
    results = []

    for pa in params_a:
        key = pa["name"].casefold()
        if key in keys_b:
            pb = keys_b[key]
            matched_b.add(key)
            equal = _values_equal(pa["value"], pb["value"])
            results.append({
                "param": pa["name"],
                "value_a": pa["value"],
                "value_b": pb["value"],
                "match": "一致" if equal else "不一致",
                "source_a": _translate_source(pa.get("source", "")),
                "source_b": _translate_source(pb.get("source", "")),
            })
        else:
            score, matched_key = _fuzzy_match_key(key, [k for k in keys_b if k not in matched_b])
            if matched_key:
                pb = keys_b[matched_key]
                matched_b.add(matched_key)
                equal = _values_equal(pa["value"], pb["value"])
                results.append({
                    "param": pa["name"],
                    "value_a": pa["value"],
                    "value_b": pb["value"],
                    "match": "一致" if equal else "不一致",
                    "fuzzy": True,
                    "fuzzy_score": round(score, 3),
                    "matched_param": pb["name"],
                    "source_a": _translate_source(pa.get("source", "")),
                    "source_b": _translate_source(pb.get("source", "")),
                })
            else:
                results.append({
                    "param": pa["name"],
                    "value_a": pa["value"],
                    "value_b": "",
                    "match": "仅A有",
                    "source_a": _translate_source(pa.get("source", "")),
                    "source_b": "",
                })

    for pb in params_b:
        key = pb["name"].casefold()
        if key not in matched_b:
            results.append({
                "param": pb["name"],
                "value_a": "",
                "value_b": pb["value"],
                "match": "仅B有",
                "source_a": "",
                "source_b": _translate_source(pb.get("source", "")),
            })

    return results


def run_param_compare(file_a_path, file_b_path):
    params_a = extract_params(file_a_path)
    params_b = extract_params(file_b_path)

    results = compare_params(params_a, params_b)

    n_total = len(results)
    n_match = sum(1 for r in results if r["match"] == "一致")
    n_diff = sum(1 for r in results if r["match"] == "不一致")
    n_only_a = sum(1 for r in results if r["match"] == "仅A有")
    n_only_b = sum(1 for r in results if r["match"] == "仅B有")

    return {
        "file_a": os.path.basename(file_a_path),
        "file_b": os.path.basename(file_b_path),
        "params_a_count": len(params_a),
        "params_b_count": len(params_b),
        "total": n_total,
        "match_count": n_match,
        "diff_count": n_diff,
        "only_a_count": n_only_a,
        "only_b_count": n_only_b,
        "results": results,
    }
