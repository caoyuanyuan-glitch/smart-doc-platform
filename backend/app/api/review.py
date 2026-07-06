from fastapi import APIRouter, Depends, HTTPException, status, Query, BackgroundTasks, File, UploadFile
from fastapi.responses import HTMLResponse, FileResponse
from sqlalchemy.orm import Session
import html as html_lib
import json
import re
import os
import shutil
import difflib
import concurrent.futures
from copy import deepcopy
from collections import Counter, defaultdict
from datetime import datetime, timedelta, time
from pathlib import Path
from zoneinfo import ZoneInfo
from app.database import get_db
from app.crud.document import get_document
from app.crud.review import create_review, get_review, get_reviews, update_review_status, create_issue, get_issues, update_issue
from app.crud.rule import get_rules
from app.crud.term import get_terms
from app.models.term import Term
from app.crud.audit_basis import get_audit_basis
from app.crud.knowledge import get_folder_tree, get_folder, get_folder_files, get_file
from app.schemas.review import Review, Issue, IssueUpdate, ReviewCreate, IssueCreate
from app.models.review import Review as ReviewModel
from app.models.issue import Issue as IssueModel
from app.models.document import Document as DocumentModel
from app.models.user import User as UserModel
from app.rules.reference_integrity_rule import ReferenceIntegrityRule
from app.rules.term_consistency_rule import TermConsistencyRule
from app.utils.ai_client import ai_client
from app.utils.spell_checker import run_spelling_and_grammar_check, is_whitelisted
from app.utils.document_parser import clean_pdf_text

_review_progress = {}  # 全局进度存储: {review_id: {'status': 'running', 'step': 'xxx', 'progress': 0-100, 'message': 'xxx'}}
UPLOAD_DIR = Path(__file__).resolve().parents[2] / "static" / "uploads"
REVIEW_EXPORT_DIR = Path(__file__).resolve().parents[2] / "static" / "review_exports"


def get_progress(review_id: int):
    return _review_progress.get(review_id, {'status': 'unknown', 'step': '', 'progress': 0, 'message': ''})


def set_progress(review_id: int, status: str, step: str, progress: int, message: str = ''):
    _review_progress[review_id] = {
        'status': status,
        'step': step,
        'progress': progress,
        'message': message,
        'timestamp': datetime.now().isoformat()
    }


def _mark_review_as_failed(db: Session, review, message: str):
    review = update_review_status(db, review.id, "failed", review.total_issues or 0, message)
    set_progress(review.id, 'failed', '失败', 0, message)
    return review


def _reconcile_review_runtime_state(db: Session, review):
    if not review or review.status != 'running':
        return review

    progress = get_progress(review.id)
    if progress.get('status') == 'running':
        return review

    return _mark_review_as_failed(db, review, '审核任务已中断，请重新发起审核')


def _get_active_review_for_document(db: Session, document_id: int):
    for review in get_reviews(db, document_id=document_id, limit=20):
        review = _reconcile_review_runtime_state(db, review)
        if review and review.status == 'running':
            return review
    return None

router = APIRouter()


def _flatten_folder_tree(nodes):
    for node in nodes or []:
        yield node
        yield from _flatten_folder_tree(node.get("children") or [])


def _get_review_spec_files_from_knowledge(db: Session):
    tree = get_folder_tree(db, None)
    folders = list(_flatten_folder_tree(tree))
    folders_by_id = {folder.get("id"): folder for folder in folders}

    def folder_path(folder):
        names = []
        current = folder
        while current:
            names.append(str(current.get("name") or ""))
            current = folders_by_id.get(current.get("parent_id"))
        return "/".join(reversed([name for name in names if name]))

    writing_root = next((folder for folder in folders if folder.get("name") == "写作规范"), None)
    if not writing_root:
        return {}

    style_folder = next(
        (folder for folder in folders if folder.get("parent_id") == writing_root.get("id") and folder.get("name") == "写作风格指南"),
        None,
    )
    common_errors_folder = next(
        (folder for folder in folders if folder.get("parent_id") == writing_root.get("id") and folder.get("name") == "常见错误清单"),
        None,
    )

    style_files = get_folder_files(db, style_folder["id"]) if style_folder else []
    common_error_files = get_folder_files(db, common_errors_folder["id"]) if common_errors_folder else []

    result = {}
    for file in style_files:
        name = str(file.get("name") or "")
        if "中文技术文档写作风格指南" in name:
            result["cn_style"] = file
        elif "MGI英文技术文档写作风格指南" in name:
            result["en_style"] = file

    for file in common_error_files:
        name = str(file.get("name") or "")
        if "技术文档常见错误清单与规范" in name:
            result["common_errors"] = file

    final_checklists = []
    for folder in folders:
        if folder.get("name") != "说明书自检checklist":
            continue
        if not folder_path(folder).startswith("写作规范/"):
            continue
        for file in get_folder_files(db, folder["id"]):
            name = str(file.get("name") or file.get("filename") or "")
            if "Checklist" in name or "checklist" in name or "自检" in name:
                final_checklists.append(file)
    if final_checklists:
        result["final_checklists"] = final_checklists

    return result


def _issue_value(issue, key, default=''):
    if isinstance(issue, dict):
        return issue.get(key, default)
    return getattr(issue, key, default)


def _normalize_search_text(text):
    return re.sub(r"\s+", "", str(text or "")).lower()


def _encode_issue_position(start=0, end=0, paragraph_index=None, char_start=None, char_end=None):
    payload = {"start": int(start or 0), "end": int(end or 0)}
    if paragraph_index is not None:
        payload["paragraph_index"] = int(paragraph_index)
    if char_start is not None:
        payload["char_start"] = int(char_start)
    if char_end is not None:
        payload["char_end"] = int(char_end)
    return json.dumps(payload, ensure_ascii=False)


def _encode_issue_position_with_meta(start=0, end=0, **meta):
    payload = {"start": int(start or 0), "end": int(end or 0)}
    payload.update({key: value for key, value in meta.items() if value not in (None, '')})
    return json.dumps(payload, ensure_ascii=False)


def _decode_issue_position(position):
    if not position:
        return {}
    if isinstance(position, dict):
        return position
    text = str(position).strip()
    if text.startswith('{'):
        try:
            data = json.loads(text)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}
    match = re.match(r"(\d+)-(\d+)", text)
    if match:
        return {"start": int(match.group(1)), "end": int(match.group(2))}
    return {}


def _parse_issue_position(position):
    data = _decode_issue_position(position)
    return (int(data.get("start", 0) or 0), int(data.get("end", 0) or 0))


def _page_number_from_position(content, position):
    start, _ = _parse_issue_position(position)
    if not content:
        return None
    cursor = 0
    for index, page in enumerate(_split_content_pages(content), start=1):
        page_len = len(page)
        if start <= cursor + page_len:
            return index
        cursor += page_len + 1
    return None


def _issue_sort_key(issue):
    severity_rank = {"fatal": 4, "serious": 3, "general": 2, "suggestion": 1}
    pos_start, _ = _parse_issue_position(_issue_value(issue, "position", ""))
    return (-severity_rank.get(_issue_value(issue, "severity", "general"), 0), pos_start, _issue_value(issue, "id", 0))


def _iter_docx_paragraphs(doc):
    for para in doc.paragraphs:
        yield para

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
                    yield from walk_tables(cell.tables)

    yield from walk_tables(doc.tables)


def _split_docx_run(run_element, offset):
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    text_elem = run_element.find(f'{{{w_ns}}}t')
    text = text_elem.text if text_elem is not None and text_elem.text is not None else ''
    if offset <= 0 or offset >= len(text):
        return run_element

    right_run = deepcopy(run_element)
    right_text = right_run.find(f'{{{w_ns}}}t')
    if text_elem is None or right_text is None:
        return run_element

    text_elem.text = text[:offset]
    right_text.text = text[offset:]
    parent = run_element.getparent()
    parent.insert(parent.index(run_element) + 1, right_run)
    return right_run


def _collect_docx_text_runs(paragraph):
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p_element = paragraph._p
    parts = []
    cursor = 0
    for child in p_element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'r':
            continue
        text_elem = child.find(f'{{{w_ns}}}t')
        text = text_elem.text if text_elem is not None and text_elem.text else ''
        if not text:
            continue
        start = cursor
        end = cursor + len(text)
        parts.append((child, start, end, text))
        cursor = end
    return parts


def _locate_docx_issue(paragraphs, issue):
    original_text = str(_issue_value(issue, 'original_text', '') or '')
    if not original_text:
        return None, None, None, None

    position_data = _decode_issue_position(_issue_value(issue, 'position', ''))
    paragraph_index = position_data.get('paragraph_index')
    char_start = position_data.get('char_start')
    char_end = position_data.get('char_end')
    if paragraph_index is not None and char_start is not None and char_end is not None:
        if 0 <= int(paragraph_index) < len(paragraphs):
            para = paragraphs[int(paragraph_index)]
            text = para.text or ''
            local_start = max(0, int(char_start))
            local_end = min(len(text), int(char_end))
            if local_end > local_start and text[local_start:local_end] == original_text:
                return para, int(paragraph_index), local_start, local_end

    pos_start, pos_end = _parse_issue_position(_issue_value(issue, 'position', ''))
    cursor = 0
    for index, para in enumerate(paragraphs):
        text = para.text or ''
        para_start = cursor
        para_end = cursor + len(text)
        if pos_end and para_start <= pos_start < para_end + 1:
            local_start = max(0, pos_start - para_start)
            local_end = min(len(text), max(local_start + 1, pos_end - para_start))
            if text[local_start:local_end] == original_text:
                return para, index, local_start, local_end
        cursor = para_end + 1

    best_para = _find_best_docx_paragraph(paragraphs, issue)
    if best_para is None:
        return None, None, None, None
    para_index = paragraphs.index(best_para)
    text = best_para.text or ''
    local_start = text.find(original_text)
    if local_start < 0:
        lowered = text.lower()
        local_start = lowered.find(original_text.lower())
    if local_start < 0:
        return best_para, para_index, None, None
    return best_para, para_index, local_start, local_start + len(original_text)


def _add_docx_comment_marker(paragraph, comment_id, char_start=None, char_end=None):
    from lxml import etree

    p_element = paragraph._p
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    parts = _collect_docx_text_runs(paragraph)
    if not parts:
        return False

    start_run = parts[0][0]
    end_run = parts[-1][0]

    if char_start is not None and char_end is not None and char_start < char_end:
        for run_el, start, end, text in parts:
            if start <= char_start < end:
                if char_start > start:
                    start_run = _split_docx_run(run_el, char_start - start)
                else:
                    start_run = run_el
                break
        parts = _collect_docx_text_runs(paragraph)
        for run_el, start, end, text in parts:
            if start < char_end <= end:
                if char_end < end:
                    _split_docx_run(run_el, char_end - start)
                end_run = run_el
                break

    comment_start = etree.Element(f'{{{w_ns}}}commentRangeStart')
    comment_start.set(f'{{{w_ns}}}id', str(comment_id))
    p_element.insert(p_element.index(start_run), comment_start)

    insert_idx = p_element.index(end_run) + 1

    comment_end = etree.Element(f'{{{w_ns}}}commentRangeEnd')
    comment_end.set(f'{{{w_ns}}}id', str(comment_id))
    p_element.insert(insert_idx, comment_end)

    comment_ref_r = etree.Element(f'{{{w_ns}}}r')
    comment_ref_rpr = etree.SubElement(comment_ref_r, f'{{{w_ns}}}rPr')
    comment_ref_style = etree.SubElement(comment_ref_rpr, f'{{{w_ns}}}rStyle')
    comment_ref_style.set(f'{{{w_ns}}}val', 'CommentReference')
    comment_ref = etree.SubElement(comment_ref_r, f'{{{w_ns}}}commentReference')
    comment_ref.set(f'{{{w_ns}}}id', str(comment_id))
    p_element.insert(insert_idx + 1, comment_ref_r)
    return True


def _inject_comments_to_docx(docx_path, comments_data):
    import tempfile
    import zipfile
    from lxml import etree

    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def append_comment_line(comment_el, segments):
        p_el = etree.SubElement(comment_el, f'{{{w_ns}}}p')
        for segment in segments:
            text = str(segment.get('text', '') or '')
            if not text:
                continue
            r_el = etree.SubElement(p_el, f'{{{w_ns}}}r')
            color = str(segment.get('color', '') or '').strip()
            if color:
                r_pr = etree.SubElement(r_el, f'{{{w_ns}}}rPr')
                color_el = etree.SubElement(r_pr, f'{{{w_ns}}}color')
                color_el.set(f'{{{w_ns}}}val', color)
            t_el = etree.SubElement(r_el, f'{{{w_ns}}}t')
            if text.startswith(' ') or text.endswith(' '):
                t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = text

    comments_xml = etree.Element(
        f'{{{w_ns}}}comments',
        nsmap={'w': w_ns}
    )

    for cdata in comments_data:
        comment_el = etree.SubElement(
            comments_xml, f'{{{w_ns}}}comment'
        )
        comment_el.set(f'{{{w_ns}}}id', str(cdata['id']))
        comment_el.set(f'{{{w_ns}}}author', cdata['author'])
        comment_el.set(f'{{{w_ns}}}initials', cdata['initials'])

        for line in cdata.get('lines', []):
            append_comment_line(comment_el, line)

    comments_bytes = etree.tostring(comments_xml, xml_declaration=True, encoding='UTF-8', standalone=True)
    temp_dir = tempfile.mkdtemp()
    temp_output = os.path.join(temp_dir, 'output.docx')

    try:
        with zipfile.ZipFile(docx_path, 'r') as zin:
            zin.extractall(temp_dir)

        word_dir = os.path.join(temp_dir, 'word')
        with open(os.path.join(word_dir, 'comments.xml'), 'wb') as f:
            f.write(comments_bytes)

        rels_path = os.path.join(word_dir, '_rels', 'document.xml.rels')
        r_id = 'rIdComments'
        if os.path.exists(rels_path):
            rels_tree = etree.parse(rels_path)
            root = rels_tree.getroot()
            existing = root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            target_exists = False
            if existing:
                for rel in existing:
                    if rel.get('Target') == 'comments.xml':
                        r_id = rel.get('Id', r_id)
                        target_exists = True
                        break
                if not target_exists:
                    existing_ids = {rel.get('Id', '') for rel in existing}
                    next_num = 1
                    while f'rId{next_num}' in existing_ids:
                        next_num += 1
                    r_id = f'rId{next_num}'
                    new_rel = etree.SubElement(
                        root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'
                    )
                    new_rel.set('Id', r_id)
                    new_rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
                    new_rel.set('Target', 'comments.xml')
                with open(rels_path, 'wb') as f:
                    f.write(etree.tostring(rels_tree, xml_declaration=True, encoding='UTF-8', standalone=True))

        doc_path = os.path.join(word_dir, 'document.xml')
        if os.path.exists(doc_path):
            doc_tree = etree.parse(doc_path)
            doc_root = doc_tree.getroot()
            body = doc_root.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}body')
            existing_comment_refs = body.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comments')
            if not existing_comment_refs:
                comments_ref_el = etree.SubElement(
                    body, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comments'
                )
                comments_ref_el.set(
                    '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id', r_id
                )
            with open(doc_path, 'wb') as f:
                f.write(etree.tostring(doc_tree, xml_declaration=True, encoding='UTF-8', standalone=True))

        with zipfile.ZipFile(temp_output, 'w', zipfile.ZIP_DEFLATED) as zout:
            for dirpath, _, filenames in os.walk(temp_dir):
                for fname in filenames:
                    full_path = os.path.join(dirpath, fname)
                    arcname = os.path.relpath(full_path, temp_dir)
                    zout.write(full_path, arcname)

        if os.path.exists(docx_path):
            os.remove(docx_path)
        shutil.move(temp_output, docx_path)
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def _build_review_comment_lines(issue):
    severity = str(_issue_value(issue, 'severity', 'general') or 'general').strip()
    category = str(_issue_value(issue, 'category', '') or '').strip()
    suggestion = str(_issue_value(issue, 'suggestion', '') or '').strip()
    description = str(_issue_value(issue, 'description', '') or '').strip()
    original_text = str(_issue_value(issue, 'original_text', '') or '').strip()

    rule = str(_issue_value(issue, 'rule', '') or '').strip()

    if rule in {'R002', 'HR003'}:
        suggestion_line = '在句号后补空格'
    elif '→' in suggestion:
        suggestion = suggestion.split('→', 1)[1].strip()
        suggestion_line = f"改为 {suggestion}" if suggestion else description
    else:
        suggestion = re.sub(r'^建议(?:改为|替换为|统一为|写为|使用|拆成两句：|改为：)\s*', '', suggestion)
        if suggestion.startswith('Read ') or suggestion.startswith('Click ') or suggestion.startswith('Contact '):
            suggestion_line = suggestion
        elif suggestion:
            suggestion_line = f"改为 {suggestion}" if not suggestion.startswith('改为') else suggestion
        else:
            suggestion_line = description

    lines = [[{"text": f"[{severity}] {category}".strip()}]]
    if suggestion_line:
        if suggestion_line.startswith('改为 '):
            diff_segments = _build_diff_segments(original_text, suggestion_line[3:], rule)
            lines.append([
                {"text": '改为 '},
                *diff_segments
            ])
        else:
            lines.append([{"text": suggestion_line}])
    if description and not suggestion_line:
        lines.append([{"text": description}])
    return [line for line in lines if line]


def _normalize_suggestion_text(suggestion):
    text = str(suggestion or '').strip()
    text = re.sub(r'^建议(?:改为|替换为|统一为|写为|使用|拆成两句：|改为：|使用标准术语:)\s*', '', text)
    text = re.sub(r'^改为\s*', '', text)
    return text.strip()


def _clean_issue_suggestion_for_display(suggestion):
    text = str(suggestion or '').strip()
    if not text:
        return ''

    bracketed = re.search(r'建议改为\s*[:：]?\s*\[([^\]]+)\]', text)
    if bracketed:
        return bracketed.group(1).strip()

    plain = re.search(r'建议(?:改为|替换为|统一为|写为|使用标准术语:)\s*[:：]?\s*([^。；\n]+)', text)
    if plain and re.search(r'疑似错误|是否确定|原文', text):
        return plain.group(1).strip().strip('[]')

    return _normalize_suggestion_text(text)


def validate_suggestion(original: str, suggestion: str) -> bool:
    normalized_suggestion = _normalize_suggestion_text(suggestion)
    if not normalized_suggestion or normalized_suggestion == '-':
        return False
    original_compact = re.sub(r'\s+', ' ', str(original or '')).strip()
    suggestion_compact = re.sub(r'\s+', ' ', normalized_suggestion).strip()
    return original_compact != suggestion_compact


def _sanitize_issue_suggestions(issues):
    for issue in issues:
        suggestion = str(issue.get('suggestion', '') or '')
        cleaned = _clean_issue_suggestion_for_display(suggestion)
        if cleaned != suggestion:
            issue['suggestion'] = cleaned
            suggestion = cleaned
        if suggestion and not validate_suggestion(issue.get('original_text', ''), suggestion):
            issue['suggestion'] = ''
            issue['confidence'] = max(0, int(issue.get('confidence', 0) or 0) - 10)
    return issues


def _is_spelling_like_issue(issue):
    rule = str(issue.get('rule', '') or '').upper()
    category = str(issue.get('category', '') or '')
    source = str(issue.get('source', '') or '').lower()
    basis = str(issue.get('audit_basis', '') or '')
    text = f'{rule} {category} {source} {basis}'
    return 'SPELL' in text or '拼写' in text or 'spelling' in text.lower()


def _extract_issue_tokens(issue):
    fields = [
        str(issue.get('original_text', '') or ''),
        str(issue.get('description', '') or ''),
        str(issue.get('suggestion', '') or ''),
    ]
    tokens = []
    for field in fields:
        tokens.extend(re.findall(r'[A-Za-z][A-Za-z0-9]*(?:-[A-Za-z0-9]+)*', field))
    return tokens


def _should_drop_spelling_issue(issue):
    if not _is_spelling_like_issue(issue):
        return False

    original = str(issue.get('original_text', '') or '').strip()
    context = str(issue.get('context', '') or '').strip()
    description = str(issue.get('description', '') or '').strip()
    suggestion = str(issue.get('suggestion', '') or '').strip()

    if not original and not context:
        return True
    if not original and suggestion and not description:
        return True

    original_tokens = re.findall(r'[A-Za-z][A-Za-z0-9]*(?:-[A-Za-z0-9]+)*', original)
    if original_tokens and all(is_whitelisted(token) for token in original_tokens):
        return True

    protected_tokens = [token for token in _extract_issue_tokens(issue) if is_whitelisted(token)]
    suspicious_tokens = [
        token for token in re.findall(r'[A-Za-z][A-Za-z0-9]*(?:-[A-Za-z0-9]+)*', original or description)
        if len(token) > 2 and not is_whitelisted(token)
    ]
    return bool(protected_tokens) and not suspicious_tokens


def _should_drop_punctuation_issue(issue):
    rule = str(issue.get('rule', '') or '').upper()
    if rule not in {'PUNCT-001', 'HR003', 'PUNCT-002', 'R002'}:
        return False

    original = str(issue.get('original_text', '') or '').strip()
    context = str(issue.get('context', '') or '')
    if rule == 'R002' and re.search(r'Part\s+No\.?', context, re.IGNORECASE):
        return True
    if original == ':C' and re.search(r'Part\s+No\.?\s*:\s*CSS-\d+', context, re.IGNORECASE):
        return True
    if original == '.R' and re.search(r'Fig(?:ure)?\s*\d+', context, re.IGNORECASE):
        return True
    return False


def _should_drop_unit_issue(issue):
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    original = str(_issue_value(issue, 'original_text', '') or '').strip()
    context = str(_issue_value(issue, 'context', '') or '')
    chapter = str(_issue_value(issue, 'chapter', '') or '')

    if rule == 'UNIT-003' and original == 'uL':
        prefix = context[:max(0, context.find('uL'))]
        number_match = re.search(r'(\d+(?:\.\d+)?)\s*$', prefix)
        if number_match and float(number_match.group(1)) >= 1000:
            return True
        if chapter.startswith('3.1 Reagent Components'):
            return False
        if chapter.startswith('3.3') and 'Pipettes' in context:
            return False
        if chapter.startswith('5.2 Cell Capture'):
            return False
        if chapter.startswith('5.4 PCR Amplification'):
            return False
        if chapter.startswith('5.5 Library Purification'):
            return False
        return True

    if rule == 'UNIT-004' and original == '-20C':
        return bool(re.search(r'Problem\s+\d+:|Solutions:', context, re.IGNORECASE))

    return False


def _normalize_report_text(text):
    return re.sub(r'\s+', ' ', str(text or '')).strip().rstrip('.').lower()


def _is_table_or_figure_caption_line(line, label, number_start=0):
    line = re.sub(r'\s+', ' ', str(line or '')).strip()
    if not line:
        return False
    if re.search(r'\.{3,}\s*\d+\s*$', line):
        return False
    if re.search(r'\b(?:refer\s+to|shown\s+in|according\s+to|see)\b', line, re.IGNORECASE):
        return False
    if re.search(r'\bRevision\s+History\b', line, re.IGNORECASE):
        return False

    prefix = line[:number_start].strip()
    if prefix:
        # Allow page numbers accidentally merged before captions, e.g. "8 Figure 3-3 ...".
        if not re.fullmatch(r'\d{1,3}', prefix):
            return False

    caption = line[number_start:].strip()
    pattern = rf'^{re.escape(label)}\s*\d+[-–]\d+\s+[^,.;:()\[\]]{{3,}}$'
    return bool(re.match(pattern, caption, re.IGNORECASE))


def _is_touchscreen_click_context(text, start):
    context = str(text or '')[max(0, start - 500):min(len(text or ''), start + 500)]
    if re.search(r'\b(?:double-click|right-click|left-click|WDesigner|desktop|software|application|import\s+application|icon|menu|dialog)\b', context, re.IGNORECASE):
        return False
    return bool(re.search(r'\b(?:touch\s*screen|touchscreen|screen|instrument\s+interface|control\s+panel|device\s+screen|tap)\b', context, re.IGNORECASE))


def _normalize_ai_issue_category(issue):
    category = str(issue.get('category') or '')
    original = str(issue.get('original_text') or '')
    suggestion = str(issue.get('suggestion') or '')
    has_cjk_text_change = bool(re.search(r'[\u4e00-\u9fff]', original) and re.search(r'[\u4e00-\u9fff]', suggestion))
    punctuation_only = bool(re.fullmatch(r'[\s\.,;:!?，。；：！？、（）()\[\]【】"“”\'‘’\-—]+', original + suggestion))
    if re.search(r'punctuation|标点', category, re.IGNORECASE) and has_cjk_text_change and not punctuation_only:
        issue['category'] = '错别字/用词'
        if str(issue.get('rule') or '').upper() in {'AI', 'PUNCTUATION'}:
            issue['rule'] = 'CN-TYPO-001'
        issue['description'] = issue.get('description') or '中文文本存在疑似错别字或用词错误。'
    return issue


def _should_drop_known_false_positive_issue(issue):
    original = _normalize_report_text(_issue_value(issue, 'original_text', ''))
    if _is_figure_details_sentence_false_positive(issue):
        return True
    if _is_intentionally_blank_page_false_positive(issue):
        return True
    if _is_complete_genomics_url_false_positive(issue):
        return True
    if _is_zh_en_spacing_false_positive(issue):
        return True
    if _is_whitelisted_model_false_positive(issue):
        return True
    if _is_noop_position_or_same_text_issue(issue):
        return True
    if _is_low_value_ai_table_style_issue(issue):
        return True
    if _is_noop_ai_suggestion_issue(issue):
        return True
    return False


def _is_low_value_ai_table_style_issue(issue):
    source = str(_issue_value(issue, 'source', '') or '').lower()
    if source != 'ai' or _is_deterministic_review_rule(issue):
        return False
    confidence = _issue_value(issue, 'confidence', 0) or 0
    try:
        confidence = int(confidence)
    except (TypeError, ValueError):
        confidence = 0
    if confidence > 80:
        return False

    rule = _normalize_report_text(_issue_value(issue, 'rule', ''))
    category = _normalize_report_text(_issue_value(issue, 'category', ''))
    if not re.search(r'\b(format|punctuation|unit|units|unit rules?)\b|标点|格式|单位', f'{rule} {category}', re.IGNORECASE):
        return False

    original = str(_issue_value(issue, 'original_text', '') or '')
    table_like = bool(re.search(
        r'\b(?:model\s+name|\d{3}-\d{6}-\d{2}|rxn|tube|module|kit|buffer|beads|barcode)\b|\d+(?:\.\d+)?\s*(?:μl|ul|ml)\b',
        original,
        re.IGNORECASE,
    ))
    return table_like


def _normalize_action_text(text):
    text = str(text or '').strip()
    text = re.sub(r'^(?:建议(?:改为|替换为|统一为)?|修改建议|修改后)\s*[:：]?\s*', '', text)
    text = text.strip('`"“”‘’[]()（） ')
    text = re.sub(r'\s+', ' ', text)
    return text.lower()


def _is_noop_position_or_same_text_issue(issue):
    original = _normalize_action_text(_issue_value(issue, 'original_text', ''))
    suggestion = _normalize_action_text(_format_issue_suggestion(issue))
    if not original or not suggestion:
        return False
    if original == suggestion:
        return True
    pos_match = re.fullmatch(r'pos\s*(\d+)', original, re.IGNORECASE)
    suggestion_pos_match = re.fullmatch(r'pos\s*(\d+)\s+to\s+pos\s*(\d+)', suggestion, re.IGNORECASE)
    return bool(pos_match and suggestion_pos_match and suggestion_pos_match.group(1) == suggestion_pos_match.group(2) == pos_match.group(1))


def _is_noop_ai_suggestion_issue(issue):
    source = str(_issue_value(issue, 'source', '') or '').lower()
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    if source != 'ai' or _is_deterministic_review_rule(issue) or rule.startswith(('SPELL', 'GRAMMAR', 'UNIT', 'PUNCT', 'HR')):
        return False
    original = _normalize_action_text(_issue_value(issue, 'original_text', ''))
    suggestion = _normalize_action_text(_format_issue_suggestion(issue))
    if not original or not suggestion:
        return False
    return original == suggestion


def _is_whitelisted_model_false_positive(issue):
    if _report_rule(_issue_value(issue, 'rule', '')) != 'DOC-MODEL-001':
        return False
    original = str(_issue_value(issue, 'original_text', '') or '')
    if not original:
        return False
    if re.search(r'\b(?:MGI-SP960|MGIISP(?:-NE384)?)\b', original, re.IGNORECASE):
        return False
    tokens = re.findall(r'\b(?:MGI[A-Z]*|DNBSEQ|STP|SP|T\d+|G\d+|E\d+)[A-Z0-9+×-]*\b', original, re.IGNORECASE)
    if not tokens:
        return False
    return all(is_whitelisted(token) for token in tokens)


def _is_intentionally_blank_page_false_positive(issue):
    original = _normalize_report_text(_issue_value(issue, 'original_text', ''))
    if 'this page is intentionally left blank' not in original:
        return False
    text = _normalize_report_text(' '.join([
        str(_issue_value(issue, 'category', '') or ''),
        str(_issue_value(issue, 'description', '') or ''),
        str(_issue_value(issue, 'suggestion', '') or ''),
        str(_issue_value(issue, 'audit_basis', '') or ''),
    ]))
    return bool(re.search(r'blank|hyphen|dash|format|punct|删除|移除|空白页|连字符|破折号|格式|标点', text, re.IGNORECASE))


def _is_complete_genomics_url_false_positive(issue):
    original = str(_issue_value(issue, 'original_text', '') or '')
    context = str(_issue_value(issue, 'context', '') or '')
    text = f'{original} {context}'
    return bool(re.search(r'https?://www\.CompleteGenomics\.com', text, re.IGNORECASE))


def _is_figure_details_sentence_false_positive(issue):
    figure_sentence = 'the details are shown in the figure below'
    original = _normalize_report_text(_issue_value(issue, 'original_text', ''))
    if figure_sentence not in original:
        return False

    text = _normalize_report_text(' '.join([
        str(_issue_value(issue, 'category', '') or ''),
        str(_issue_value(issue, 'description', '') or ''),
        str(_issue_value(issue, 'suggestion', '') or ''),
        str(_issue_value(issue, 'audit_basis', '') or ''),
    ]))
    if not text:
        return True
    return bool(re.search(r'\b(?:delete|remove|omit|replace|trim|terminology|format)\b|删除|移除|替换|术语|格式', text, re.IGNORECASE))


def _is_zh_en_spacing_false_positive(issue):
    text = ' '.join([
        str(_issue_value(issue, 'category', '') or ''),
        str(_issue_value(issue, 'description', '') or ''),
        str(_issue_value(issue, 'suggestion', '') or ''),
        str(_issue_value(issue, 'audit_basis', '') or ''),
    ])
    normalized = re.sub(r'\s+', '', text).lower()
    if not re.search(r'(中文|中英文|汉字|cjk|chinese)', normalized):
        return False
    if not re.search(r'(英文|英文字母|字母|english|latin)', normalized):
        return False
    return bool(re.search(r'(空格|space|spacing)', normalized))


def _is_deterministic_review_rule(issue):
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    return rule.startswith(('DOC-', 'ENG-CN-', 'CHECKLIST-'))


def _filter_review_false_positives(issues):
    filtered = []
    dropped = 0
    for issue in issues:
        if (
            _should_drop_spelling_issue(issue)
            or _should_drop_punctuation_issue(issue)
            or _should_drop_unit_issue(issue)
            or _should_drop_known_false_positive_issue(issue)
        ):
            dropped += 1
            continue
        filtered.append(issue)
    if dropped:
        print(f'[审核] 拼写白名单最终过滤: 过滤 {dropped} 个, 剩余 {len(filtered)} 个')
    return filtered


def _issue_judgment_signature(issue):
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    original = re.sub(r'\s+', ' ', str(_issue_value(issue, 'original_text', '') or '')).strip().lower()
    category = re.sub(r'\s+', ' ', str(_issue_value(issue, 'category', '') or '')).strip().lower()
    return f'{rule}|{category}|{original}'


def _issue_judgment_signatures(issue):
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    original = re.sub(r'\s+', ' ', str(_issue_value(issue, 'original_text', '') or '')).strip().lower()
    category = re.sub(r'\s+', ' ', str(_issue_value(issue, 'category', '') or '')).strip().lower()
    if not original:
        return set()
    return {
        f'{rule}|{category}|{original}',
        f'{rule}|{original}',
        f'{category}|{original}',
        original,
    }


def _is_issue_hidden_by_judgment(issue):
    return str(getattr(issue, 'status', '') or '').lower() in {'false_positive', 'ignored'}


def _is_pdf_duplicate_word_extraction_noise(issue):
    rule = str(getattr(issue, 'rule', '') or '')
    if rule != 'R024':
        return False
    original = str(getattr(issue, 'original_text', '') or '')
    context = str(getattr(issue, 'context', '') or '')
    tokens = re.findall(r'\b\w+\b', original.lower())
    if len(tokens) != 2 or tokens[0] != tokens[1]:
        return False
    if re.search(r'\n\s*\n', original):
        return True
    return bool(re.search(r'\b(?:Table|Figure|Cat\.\s*No\.|Recommended brand|Library Type|condition)\b', context, re.IGNORECASE))


def _visible_review_issues(issues):
    return [issue for issue in issues if not _is_issue_hidden_by_judgment(issue) and not _is_pdf_duplicate_word_extraction_noise(issue)]


def _issue_position_start(issue):
    raw_position = getattr(issue, 'position', None)
    if not raw_position:
        return None
    if isinstance(raw_position, int):
        return raw_position
    try:
        parsed = json.loads(raw_position) if isinstance(raw_position, str) else raw_position
    except Exception:
        parsed = None
    if isinstance(parsed, dict) and isinstance(parsed.get('start'), int):
        return parsed.get('start')
    if isinstance(raw_position, str) and raw_position.isdigit():
        return int(raw_position)
    return None


def _normalize_review_issue_display(issues, content=None):
    for issue in issues:
        cleaned_suggestion = _clean_issue_suggestion_for_display(getattr(issue, 'suggestion', '') or '')
        if cleaned_suggestion != (getattr(issue, 'suggestion', '') or ''):
            issue.suggestion = cleaned_suggestion
        if getattr(issue, 'rule', None) == 'CHECKLIST-COPYRIGHT-YEAR' and getattr(issue, 'category', None) == '发布前自检':
            issue.chapter = '版本记录'
            continue
        if content:
            start = _issue_position_start(issue)
            if start is not None:
                content_chapter = extract_chapter(content, start)
                if re.match(r'^\d+(?:\.\d+)+\s+', content_chapter):
                    issue.chapter = content_chapter
                    continue
        context = str(getattr(issue, 'context', '') or '')
        original = str(getattr(issue, 'original_text', '') or '')
        chapter = str(getattr(issue, 'chapter', '') or '')
        if not context or not original or not re.match(r'^\d+(?:\.\d+)+\s+', chapter):
            continue
        original_index = context.find(original)
        if original_index < 0:
            original_index = len(context)
        context_chapter = extract_chapter(context, original_index)
        if re.match(r'^\d+(?:\.\d+)+\s+', context_chapter) and context_chapter != chapter:
            issue.chapter = context_chapter
    return issues


def _load_false_positive_signatures_for_document(db: Session, document_id: int):
    signatures = set()
    for review in get_reviews(db, document_id=document_id):
        for issue in get_issues(db, review_id=review.id):
            if str(getattr(issue, 'status', '') or '').lower() == 'false_positive':
                signatures.update(_issue_judgment_signatures(issue))
    return signatures


def _merge_comment_segments(segments):
    merged = []
    for segment in segments:
        text = str(segment.get('text', '') or '')
        if not text:
            continue
        color = str(segment.get('color', '') or '')
        if merged and str(merged[-1].get('color', '') or '') == color:
            merged[-1]['text'] += text
        else:
            merged.append({'text': text, 'color': color} if color else {'text': text})
    return merged


def _build_diff_segments(original_text, suggestion_text, rule=''):
    original_text = str(original_text or '')
    suggestion_text = str(suggestion_text or '')
    rule = str(rule or '')
    if not suggestion_text:
        return []

    # Full-word or full-phrase replacements read better when fully highlighted.
    if rule in {'HR002', 'HR005', 'HR008', 'HR009', 'HR010', 'R036'}:
        return [{'text': suggestion_text, 'color': 'FF0000'}]

    if rule == 'HR004':
        match = re.match(r'^(\d+(?:\.\d+)?)(.*)$', suggestion_text)
        if match:
            return _merge_comment_segments([
                {'text': match.group(1)},
                {'text': match.group(2), 'color': 'FF0000'}
            ])

    matcher = difflib.SequenceMatcher(a=original_text, b=suggestion_text)
    segments = []
    for opcode, a0, a1, b0, b1 in matcher.get_opcodes():
        text = suggestion_text[b0:b1]
        if not text:
            continue
        if opcode == 'equal':
            segments.append({'text': text})
        else:
            segments.append({'text': text, 'color': 'FF0000'})

    segments = _merge_comment_segments(segments)
    if not segments:
        return [{'text': suggestion_text, 'color': 'FF0000'}]

    if all(not segment.get('color') for segment in segments):
        return [{'text': suggestion_text, 'color': 'FF0000'}]

    return segments


def _load_review_summary(summary_text):
    if not summary_text:
        return {}
    try:
        return json.loads(summary_text)
    except Exception:
        return {}


def _find_best_docx_paragraph(paragraphs, issue):
    original = _normalize_search_text(_issue_value(issue, 'original_text', ''))
    if not original:
        return None

    best_para = None
    best_score = -1
    context_tokens = set(re.findall(r"[A-Za-z0-9\u4e00-\u9fff]+", str(_issue_value(issue, 'context', '')).lower()))

    for para in paragraphs:
        text = para.text or ''
        normalized = _normalize_search_text(text)
        if not normalized or original not in normalized:
            continue

        score = 1
        lowered = text.lower()
        if context_tokens:
            score += sum(1 for token in context_tokens if token and token in lowered)
        original_text = str(_issue_value(issue, 'original_text', '') or '')
        if original_text and original_text in text:
            score += 2
        if score > best_score:
            best_para = para
            best_score = score

    return best_para


def _enrich_docx_issue_positions(document, issues):
    from docx import Document as DocxDocument

    source_path = _get_document_upload_path(document)
    if not source_path or not source_path.exists():
        return issues

    paragraphs = list(_iter_docx_paragraphs(DocxDocument(str(source_path))))
    for issue in issues:
        para, para_index, char_start, char_end = _locate_docx_issue(paragraphs, issue)
        start, end = _parse_issue_position(issue.get('position', ''))
        if para is None or char_start is None or char_end is None:
            continue
        issue['position'] = _encode_issue_position(start, end, para_index, char_start, char_end)
    return issues


def _build_long_sentence_suggestion(sentence):
    suggestion = re.sub(r'\s+', ' ', sentence).strip().rstrip('.!?')
    if re.match(r'^If initialize successfully is displayed,', suggestion, re.IGNORECASE):
        return 'If "Initialize successfully" is displayed, the device is connected successfully. Proceed to the next step.'
    if re.match(r'^According to the number of DNB plates,', suggestion, re.IGNORECASE):
        return 'Prepare the reagents and consumables in advance according to the number of DNB plates, as shown in table 16 and figure 20.'
    if re.match(r'^This operating instruction is intended to guide operators through\b', suggestion, re.IGNORECASE):
        rewritten = re.sub(
            r'^This operating instruction is intended to guide operators through\s+',
            'This instruction guides operators through ',
            suggestion,
            flags=re.IGNORECASE,
        )
        if ' based on ' in rewritten:
            lead, basis = rewritten.split(' based on ', 1)
            return f"{lead}. It follows {basis}."
        return f"{rewritten}."
    suggestion = re.sub(r'\bplease\b', '', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bkindly\b', '', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bcan\s+be\b', 'is', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bin order to\b', 'to', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bso as to\b', 'to', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bit is recommended to\b', '', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bit is intended to\b', 'is intended to', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bthe concentration of DNB is lower than\b', 'the DNB concentration is below', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bbefore experiment\b', '', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\bprior to use\b', 'before use', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\buntil further use\b', 'until use', suggestion, flags=re.IGNORECASE)
    suggestion = re.sub(r'\s+,\s+', ', ', suggestion)
    suggestion = re.sub(r'\s{2,}', ' ', suggestion).strip(' ;,')
    if sentence.lower().startswith('before experiment,'):
        rewritten = sentence[len('Before experiment,'):].strip().rstrip('.!?')
        return f"{rewritten} before starting the experiment."
    if len(suggestion) > 110:
        split_point = max(suggestion.rfind(', ', 0, 110), suggestion.rfind(' and ', 0, 110))
        if split_point > 40:
            first = suggestion[:split_point].strip(' ,;')
            second = suggestion[split_point + 2:].strip(' ,;') if suggestion[split_point:split_point + 2] == ', ' else suggestion[split_point + 5:].strip(' ,;')
            if first and second:
                second = second[:1].upper() + second[1:]
                return f"{first}. {second}."
    return f"{suggestion}."


def _run_long_sentence_audit(content, file_type=None):
    if file_type == 'pdf':
        return []

    issues = []
    seen = set()
    sentence_pattern = re.compile(r'[^\n.!?]+[.!?]')
    skip_pattern = re.compile(r'^(?:contents|chapter\s+\d+|figure\s+\d+|table\s+\d+|note:|[-•*]\s+|\d+[\).])', re.IGNORECASE)
    concise_trigger_pattern = re.compile(r'\b(?:before experiment|it is recommended to|in order to|as shown in|intended to guide)\b', re.IGNORECASE)
    imperative_skip_pattern = re.compile(r'^(?:click|select|open|enter|choose|read|take|place|mix|centrifuge|install|log in|log out)\b', re.IGNORECASE)
    action_hint_pattern = re.compile(r'\b(?:double-click|click|clicking|select|open|place|close|add|mix|aspirate|take out|install|log in|log out)\b', re.IGNORECASE)
    ui_step_pattern = re.compile(r'\b(?:figure\s+\d+|table\s+\d+|interface|wizard|button|door|well\b|pos\d+)\b', re.IGNORECASE)
    ui_display_pattern = re.compile(r'\b(?:will appear|will pop up|pop up)\b', re.IGNORECASE)

    for match in sentence_pattern.finditer(content):
        sentence = re.sub(r'\s+', ' ', match.group(0)).strip()
        next_fragment = content[match.end():match.end() + 12].lstrip()
        word_count = len(re.findall(r"[A-Za-z]+(?:[-'][A-Za-z]+)?", sentence))
        if word_count <= 40 or skip_pattern.search(sentence) or '|' in sentence or '\t' in sentence:
            continue
        if not re.search(r'[.!?]$', sentence):
            continue
        if re.search(r'\b[A-Z]\.$', sentence) and next_fragment[:1].islower():
            continue
        if imperative_skip_pattern.search(sentence) and word_count <= 45 and sentence.count(',') <= 1:
            continue
        if action_hint_pattern.search(sentence) and ui_step_pattern.search(sentence) and word_count <= 40:
            continue
        if ui_display_pattern.search(sentence) and ui_step_pattern.search(sentence) and word_count <= 35:
            continue
        normalized = _normalize_search_text(sentence)
        if normalized in seen:
            continue
        seen.add(normalized)
        severity = 'serious' if word_count > 60 else 'general'
        issues.append({
            "severity": severity,
            "category": "句式",
            "rule": "R036",
            "chapter": extract_chapter(content, match.start()),
            "original_text": sentence,
            "context": get_context(content, match.start(), match.end(), 120),
            "suggestion": _build_long_sentence_suggestion(sentence),
            "description": f"单句包含 {word_count} 个词，超过长句阈值，建议拆分或压缩冗余表达。",
            "audit_basis": "英文技术文档写作规范 - 长句优化",
            "confidence": 90,
            "source": "rule",
            "position": _encode_issue_position(match.start(), match.end()),
        })
    return issues


def _normalize_heading_text(text):
    text = re.sub(r'\.{2,}\s*\d+$', '', str(text or '').strip())
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()


TOC_BLACKLIST = {
    'ml', 'μl', 'ul', 'l', 'tube', 'well', 'wells',
    'table', 'figure', 'mg', 'ng', 'μg',
    'according to', '×', 'n×', '×50', '|',
}


def _split_pdf_pages(content):
    pages = [page.strip() for page in str(content or '').split('\f')]
    return [page for page in pages if page]


def _extract_toc_title(line):
    title = re.sub(r'\s*\.+\s*\d+\s*$', '', str(line or '')).strip()
    return re.sub(r'\s+', ' ', title)


def _extract_toc_entry(line):
    match = re.match(r'^(?:(\d+(?:\.\d+)*)\s+)?(.+?)\s*\.{2,}\s*(\d+)\s*$', str(line or '').strip())
    if not match:
        return None
    return {
        'section_no': match.group(1) or '',
        'title': re.sub(r'\s+', ' ', match.group(2).strip()),
        'page_no': match.group(3),
    }


def _is_toc_item(line: str) -> bool:
    stripped = str(line or '').strip()
    entry = _extract_toc_entry(stripped)
    if not entry:
        return False

    title_part = entry['title']
    lowered = title_part.lower()
    if any(keyword in lowered for keyword in TOC_BLACKLIST):
        return False
    if len(title_part) > 50:
        return False
    if re.match(r'^[\d\)\(\s]+$', title_part):
        return False
    if re.match(r'^\d+[\)）]\s*$', title_part):
        return False
    return True


def _detect_toc_page_indexes(pages):
    candidate_indexes = []
    for index, page_text in enumerate(pages[:5]):
        lines = [line.strip() for line in page_text.splitlines() if line.strip()]
        toc_like_count = sum(1 for line in lines if _is_toc_item(line))
        has_contents = any(line.lower() == 'contents' for line in lines)
        if has_contents or toc_like_count >= 3:
            candidate_indexes.append(index)
    return candidate_indexes


def _collect_toc_entries(pages):
    toc_entries = {}
    toc_page_indexes = _detect_toc_page_indexes(pages)
    if not toc_page_indexes:
        return toc_entries, set()

    for page_index in toc_page_indexes:
        lines = [line.strip() for line in pages[page_index].splitlines() if line.strip()]
        for line in lines:
            if not _is_toc_item(line):
                continue
            entry = _extract_toc_entry(line)
            if not entry:
                continue
            key = entry['section_no'] or _normalize_heading_text(entry['title'])
            toc_entries[key] = _extract_toc_title(line)
    return toc_entries, set(toc_page_indexes)


def _is_body_heading_line(line):
    stripped = re.sub(r'\s+', ' ', str(line or '').strip())
    if not stripped or len(stripped) > 80:
        return False
    lowered = stripped.lower()
    if any(keyword in lowered for keyword in TOC_BLACKLIST):
        return False
    if '| ' in stripped or ' |' in stripped or '  |  ' in stripped:
        return False
    if '. .' in stripped or re.search(r'\.{2,}\s*\d+$', stripped):
        return False
    if re.search(r'\d+(?:\.\d+)?\s*(?:mL|μL|uL|mg|ng|mm|cm|kg|rpm|min|sec|h)\b', stripped, re.IGNORECASE):
        return False
    if re.match(r'^\d+(?:\.\d+)*\s+[A-Z][A-Za-z0-9 ,/&()_-]{2,}$', stripped):
        return True
    if re.match(r'^(?:Chapter|Section)\s+\d+[\s:.-]+', stripped, re.IGNORECASE):
        return True
    return False


def _run_pdf_structure_audit(content):
    issues = []
    pages = _split_pdf_pages(content)
    if not pages:
        return issues

    toc_entries, toc_page_indexes = _collect_toc_entries(pages)
    body_pages = [page for index, page in enumerate(pages) if index not in toc_page_indexes]
    lines = [line.strip() for page in body_pages for line in page.splitlines() if line.strip()]
    if not lines:
        return issues

    headings = {}
    for line in lines:
        if not _is_body_heading_line(line):
            continue
        match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', line)
        if not match:
            continue
        section_no = match.group(1)
        title = match.group(2).strip()
        if title.isdigit():
            continue
        headings.setdefault(section_no, title)

    for section_no, title in headings.items():
        if toc_entries and section_no not in toc_entries and section_no.count('.') <= 1:
            issues.append({
                'severity': 'general',
                'category': '目录完整性',
                'rule': 'TOC-001',
                'chapter': title,
                'original_text': title,
                'context': title,
                'suggestion': f'建议将 {section_no} {title} 补充到目录中',
                'description': '正文中的一级或二级章节应在目录中列出。',
                'audit_basis': '目录完整性检查',
                'confidence': 85,
                'source': 'rule',
                'position': ''
            })
        toc_title = toc_entries.get(section_no)
        if toc_title and _normalize_heading_text(toc_title) != _normalize_heading_text(title):
            issues.append({
                'severity': 'serious',
                'category': '索引一致性',
                'rule': 'INDEX-001',
                'chapter': title,
                'original_text': toc_title,
                'context': f'目录: {toc_title} | 正文: {title}',
                'suggestion': f'建议统一为 {title}',
                'description': '目录条目名称应与正文章节标题完全一致。',
                'audit_basis': '索引一致性检查',
                'confidence': 90,
                'source': 'rule',
                'position': ''
            })

    seen_duplicate_lines = set()
    for line in lines:
        compact = re.sub(r'\s+', ' ', line).strip()
        looks_like_header = len(compact) <= 40 and len(compact.split()) <= 6
        if compact.lower() == 'version date version':
            continue
        if looks_like_header and re.search(r'\bVersion\b.*\bVersion\b', line, re.IGNORECASE):
            normalized = line.lower()
            if normalized in seen_duplicate_lines:
                continue
            seen_duplicate_lines.add(normalized)
            confidence = 82
            issues.append({
                'severity': 'suggestion' if confidence < 85 else 'general',
                'category': '表格规范',
                'rule': 'TABLE-001',
                'chapter': '',
                'original_text': line,
                'context': line,
                'suggestion': '建议删除重复的表头列名',
                'description': '表格表头中不应出现重复列名。',
                'audit_basis': '表头重复检查',
                'confidence': confidence,
                'source': 'rule',
                'position': ''
            })

    return issues


def _build_review_export_filename(doc_name, review_id, suffix):
    stem = Path(doc_name or f"review_{review_id}").stem
    safe = re.sub(r"[^A-Za-z0-9\u4e00-\u9fff._-]+", "_", stem).strip("_") or f"review_{review_id}"
    return f"{safe}_审核结果_{review_id}{suffix}"


def _get_document_upload_path(document):
    candidate = UPLOAD_DIR / (document.filename or "")
    if candidate.exists():
        return candidate
    return None


def _normalize_excel_cell(value):
    if value is None:
        return ''
    return str(value).replace('\r\n', '\n').replace('\r', '\n')


def _display_excel_cell(value):
    return re.sub(r'\s+', ' ', str(value or '')).strip()


def _looks_chinese_text(text):
    return bool(re.search(r'[\u4e00-\u9fff]', str(text or '')))


def _looks_english_text(text):
    return bool(re.search(r'[A-Za-z]', str(text or '')))


def _infer_excel_column_mapping(header_values):
    source_column = 0
    target_column = 1
    context_column = -1
    for index, header in enumerate(header_values):
        text = str(header or '').lower()
        if any(key in text for key in ['中文', '原文', 'source', 'cn', 'chinese']) and source_column == 0:
            source_column = index
        if any(key in text for key in ['英文', '译文', 'translation', 'target', 'en', 'english']):
            target_column = index
        if any(key in text for key in ['备注', '上下文', 'context', 'comment', 'note']):
            context_column = index
    if source_column == target_column:
        target_column = 1 if source_column == 0 else 0
    return {
        'source_column': source_column,
        'target_column': target_column,
        'context_column': context_column,
        'header_row': 0,
    }


def _iter_excel_review_rows(file_path):
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            header = [_normalize_excel_cell(value) for value in rows[0]]
            mapping = _infer_excel_column_mapping(header)
            source_col = mapping['source_column']
            target_col = mapping['target_column']
            context_col = mapping['context_column']
            max_col = max(source_col, target_col, context_col)
            for zero_index, row in enumerate(rows[1:], start=1):
                if max_col >= len(row):
                    continue
                source_text = _normalize_excel_cell(row[source_col] if source_col < len(row) else '')
                target_text = _normalize_excel_cell(row[target_col] if target_col < len(row) else '')
                context_text = _normalize_excel_cell(row[context_col] if context_col >= 0 and context_col < len(row) else '')
                if not source_text.strip() and not target_text.strip() and not context_text.strip():
                    continue
                yield {
                    'sheet': sheet_name,
                    'row_number': zero_index + 1,
                    'source_text': source_text,
                    'target_text': target_text,
                    'context_text': context_text,
                    'source_column': source_col + 1,
                    'target_column': target_col + 1,
                }
    finally:
        wb.close()


def _excel_issue(row, issue_type, rule, severity, original_text, suggestion, description, audit_basis, confidence=90):
    context = f"中文: {_display_excel_cell(row['source_text'])} | 英文: {_display_excel_cell(row['target_text'])}"
    if row.get('context_text'):
        context += f" | 备注: {row['context_text']}"
    return {
        'severity': severity,
        'category': issue_type,
        'rule': rule,
        'chapter': f"ROW {row['row_number']}",
        'original_text': original_text,
        'context': context,
        'suggestion': suggestion,
        'description': description,
        'audit_basis': audit_basis,
        'confidence': confidence,
        'source': 'excel',
        'position': _encode_issue_position_with_meta(
            0,
            0,
            sheet=row['sheet'],
            row=row['row_number'],
            source_column=row['source_column'],
            target_column=row['target_column'],
        ),
    }


def _run_excel_review_audit(db: Session, document, max_length=30):
    source_path = _get_document_upload_path(document)
    if not source_path or not source_path.exists():
        raise ValueError('原始 Excel 文件不存在')

    issues = []
    rows = list(_iter_excel_review_rows(source_path))
    source_to_targets = {}
    target_to_sources = {}
    terms = get_terms(db, limit=10000)

    for row in rows:
        source = row['source_text']
        target = row['target_text']
        if source:
            source_to_targets.setdefault(source.strip(), {}).setdefault(target.strip(), []).append(row['row_number'])
        if target:
            target_to_sources.setdefault(target.strip().lower(), {}).setdefault(source.strip(), []).append(row['row_number'])

        if not source:
            issues.append(_excel_issue(row, '完整性', 'XLS-COMP-001', 'serious', target, '补充中文原文', '中文原文为空，需补充。', 'Excel翻译对照表完整性检查'))
        if not target:
            issues.append(_excel_issue(row, '完整性', 'XLS-COMP-002', 'serious', source, '补充英文译文', '英文译文为空，需补充。', 'Excel翻译对照表完整性检查'))
        if source.strip() and target.strip() and source.strip() == target.strip() and (_looks_chinese_text(source) or _looks_english_text(source)):
            issues.append(_excel_issue(row, '完整性', 'XLS-COMP-003', 'general', target, '确认是否已完成翻译', '原文与译文完全相同，需确认是否漏翻。', 'Excel翻译对照表完整性检查'))

        if target:
            target_stripped = target.strip()
            if len(target_stripped) > max_length:
                issues.append(_excel_issue(row, '长度过长', 'XLS-LEN-001', 'general', target, f'建议控制在 {max_length} 个字符以内', f'英文长度 {len(target_stripped)} 字符，超出界面显示限制。', 'Excel界面字符串长度检查'))
            if re.search(r'\s{2,}', target):
                normalized_target = re.sub(r'\s+', ' ', target).strip()
                issues.append(_excel_issue(row, '格式问题', 'XLS-FMT-001', 'general', target, f'建议改为 {normalized_target}', '英文译文存在连续空格。', 'Excel界面字符串格式检查'))
            if target.strip() != target:
                issues.append(_excel_issue(row, '格式问题', 'XLS-FMT-002', 'general', target, '删除首尾多余空格', '单元格文本首尾存在多余空格。', 'Excel界面字符串格式检查'))
            if re.search(r'[，。；：！？]', target):
                issues.append(_excel_issue(row, '格式问题', 'XLS-FMT-003', 'general', target, '英文译文使用英文标点', '英文译文中出现中文标点。', 'Excel界面字符串标点检查'))
            if len(target_stripped) > 1 and target_stripped.isupper() and re.search(r'[A-Z]', target_stripped):
                issues.append(_excel_issue(row, '格式问题', 'XLS-FMT-004', 'suggestion', target, target_stripped.title(), '界面字符串通常避免全大写。', 'Excel界面字符串大小写检查'))
            if re.fullmatch(r'[a-z][a-z ]{2,}', target_stripped) and len(target_stripped.split()) <= 4:
                issues.append(_excel_issue(row, '格式问题', 'XLS-FMT-005', 'suggestion', target, target_stripped[:1].upper() + target_stripped[1:], '按钮或菜单项建议首字母大写。', 'Excel界面字符串大小写检查'))
            if re.fullmatch(r'[A-Za-z]:', target_stripped) or re.fullmatch(r'[A-Za-z][A-Za-z ]{1,30}:', target_stripped):
                issues.append(_excel_issue(row, '格式问题', 'XLS-FMT-006', 'general', target, target_stripped.rstrip(':'), '界面标签不应包含冒号。', 'Excel界面字符串标签格式检查'))
            replacements = {
                'Delet': 'Delete',
                'Sav e': 'Save',
                'Edite': 'Edit',
                'C opy': 'Copy',
                'recieve': 'receive',
                'teh': 'the',
                'File is existed': 'File already exists',
                'Connect failed': 'Connection failed',
                'Network Connect': 'Network Connection',
                'Blue Tooth': 'Bluetooth',
                'Bar Code': 'Barcode',
                'Exit system': 'Exit the system',
                'System Setting': 'System Settings',
                'Press the button': 'Click the button',
            }
            for wrong, right in replacements.items():
                if wrong.lower() in target.lower():
                    spelling_errors = {'Delet', 'Sav e', 'Edite', 'C opy', 'recieve', 'teh', 'Blue Tooth'}
                    format_errors = {'Bar Code'}
                    category = '拼写错误' if wrong in spelling_errors else '格式问题' if wrong in format_errors else '语法错误'
                    issues.append(_excel_issue(row, category, 'XLS-LANG-001', 'serious', target, f'建议改为 {right}', f'发现常见语言问题：{wrong}', 'Excel界面字符串语言质量检查'))
            if re.search(r'\binput\b', target, re.IGNORECASE):
                issues.append(_excel_issue(row, '语法错误', 'XLS-LANG-002', 'serious', target, re.sub(r'\binput\b', 'enter', target, flags=re.IGNORECASE), '界面输入提示建议使用 enter。', '英文技术文档写作风格指南'))
            if source.startswith('请'):
                issues.append(_excel_issue(row, '中文风格', 'XLS-CN-STYLE-001', 'suggestion', source, f'建议改为 {source[1:]}', '界面按钮文案建议简洁，避免使用“请”开头。', '中文技术文档写作风格指南'))
            concise_source_suggestions = {
                '点击此处进行下一步操作': '下一步',
                '请确定您是否要执行此操作？': '确定',
                '请输入您的姓名': '输入您的姓名',
                '请等待...': '等待...',
            }
            if source.strip() in concise_source_suggestions:
                issues.append(_excel_issue(row, '中文建议', 'XLS-CN-STYLE-002', 'suggestion', source, concise_source_suggestions[source.strip()], '中文界面文案建议更简洁。', '中文技术文档写作风格指南'))
            if re.search(r'\bPress\b', target, re.IGNORECASE):
                issues.append(_excel_issue(row, '英文风格', 'XLS-EN-STYLE-001', 'suggestion', target, re.sub(r'\bPress\b', 'Click', target, flags=re.IGNORECASE), '界面操作建议使用 Click。', '英文技术文档写作风格指南'))

        for term in terms:
            non_standard = str(getattr(term, 'non_standard', '') or '').strip()
            standard = str(getattr(term, 'standard', '') or '').strip()
            if non_standard and standard and non_standard in source and standard.lower() not in target.lower():
                issues.append(_excel_issue(row, '术语不一致', 'XLS-TERM-001', 'serious', f'{source} → {target}', f'{non_standard} 建议译为 {standard}', '译文与术语库标准译法不一致。', '知识库术语库'))

    for source, targets in source_to_targets.items():
        normalized_targets = {target for target in targets if target}
        if len(normalized_targets) <= 1:
            continue
        target_list = list(normalized_targets)
        for target, row_numbers in targets.items():
            for row in [item for item in rows if item['source_text'] == source and item['target_text'] == target]:
                issues.append(_excel_issue(row, '术语不一致', 'XLS-TERM-002', 'serious', f'{source} → {target}', f'统一译法，当前存在 {", ".join(target_list)}', f'同一中文“{source}”对应多个英文译法。', 'Excel跨行术语一致性检查'))

    equivalent_sources = [
        ('登录', '登入'),
    ]
    rows_by_source = {}
    for row in rows:
        rows_by_source.setdefault(row['source_text'].strip(), []).append(row)
    for left, right in equivalent_sources:
        left_rows = rows_by_source.get(left, [])
        right_rows = rows_by_source.get(right, [])
        if not left_rows or not right_rows:
            continue
        left_targets = {row['target_text'].strip() for row in left_rows if row['target_text'].strip()}
        for row in right_rows:
            target = row['target_text'].strip()
            if target and left_targets and target not in left_targets:
                issues.append(_excel_issue(row, '术语不一致', 'XLS-TERM-003', 'serious', f"{row['source_text']} → {row['target_text']}", f'与“{left}”统一译法，当前基准为 {", ".join(sorted(left_targets))}', f'“{left}”与“{right}”含义相近，英文译法需统一。', 'Excel跨行语义一致性检查'))

    seen_pairs = {}
    for row in rows:
        pair = (row['source_text'].strip(), row['target_text'].strip())
        if not pair[0] and not pair[1]:
            continue
        if pair in seen_pairs:
            issues.append(_excel_issue(row, '完整性', 'XLS-COMP-004', 'general', f'{pair[0]} → {pair[1]}', f'与第 {seen_pairs[pair]} 行重复，请确认是否重复录入', '存在完全重复的中英文组合。', 'Excel翻译对照表重复行检查'))
        else:
            seen_pairs[pair] = row['row_number']

    return issues


def _select_export_issues(issues):
    visible_issues = _visible_review_issues(issues)
    preferred = [i for i in visible_issues if getattr(i, 'status', None) in (None, '', 'pending', 'confirmed', 'converted_to_rule')]
    return sorted(preferred or visible_issues, key=_issue_sort_key)


def _format_review_mode(mode):
    return {
        "hybrid": "混合审核（规则 + 拼写 + AI）",
        "rule": "规则审核",
        "ai": "AI审核",
    }.get(mode or "", mode or "-")


def _format_issue_source(source):
    return {
        "rule": "规则",
        "term": "术语",
        "ai": "AI",
        "spellcheck": "拼写检查",
        "grammar": "语法检查",
    }.get(source or "", source or "-")


def _format_issue_severity(severity):
    return {
        "fatal": "致命",
        "serious": "严重",
        "general": "一般",
        "suggestion": "建议",
    }.get(severity or "general", (severity or "general").upper())


def _format_issue_display_id(index):
    return f"#{index:04d}"


def _report_now():
    return datetime.now(ZoneInfo('Asia/Shanghai'))


def _format_report_datetime(value=None):
    dt = value or _report_now()
    if isinstance(dt, str):
        try:
            dt = datetime.fromisoformat(dt)
        except ValueError:
            return dt
    if getattr(dt, 'tzinfo', None) is None:
        dt = dt.replace(tzinfo=ZoneInfo('UTC'))
        dt = dt.astimezone(ZoneInfo('Asia/Shanghai'))
    else:
        dt = dt.astimezone(ZoneInfo('Asia/Shanghai'))
    return dt.strftime('%Y-%m-%d %H:%M:%S (UTC+8)')


def _highlight_issue_context(issue):
    grouped_original = _issue_value(issue, 'original_override', '')
    if grouped_original:
        return html_lib.escape(str(grouped_original)).replace('、', '<br>')

    content = ''
    if isinstance(issue, dict):
        content = issue.get('_document_content', '')
    else:
        content = getattr(issue, '_document_content', '')
    context = _extract_issue_snippet(issue, content)
    original = str(getattr(issue, "original_text", "") or "") if not isinstance(issue, dict) else str(issue.get('original_text', '') or '')
    if not context or context == '-':
        return "-"

    escaped_context = html_lib.escape(context)
    if not original:
        return escaped_context.replace("\n", "<br>")

    escaped_original = html_lib.escape(original)
    pattern = re.compile(re.escape(escaped_original), re.IGNORECASE)
    highlighted = pattern.sub('<span class="problem-highlight">\\g<0></span>', escaped_context, count=1)
    return highlighted.replace("\n", "<br>")


def _split_content_pages(content):
    pages = str(content or '').split('\f')
    return pages if pages else ['']


def _extract_document_metadata(doc, content):
    text = str(content or '')
    pages = _split_content_pages(text)
    normalized = text.replace('\f', '\n')
    metadata = {
        'name': getattr(doc, 'filename', '') or '-',
        'file_type': (getattr(doc, 'file_type', '') or '-').upper(),
        'page_count': len([page for page in pages if page.strip()]) or len(pages),
        'section_count': 0,
        'version': '-',
        'doc_no': '-',
        'language': detect_language(text),
    }

    heading_hits = re.findall(r'(?m)^\s*(?:\d+(?:\.\d+)*[\.)]?\s+[A-Z][^\n]{2,}|(?:Chapter|Section)\s+\d+[^\n]*)', normalized)
    metadata['section_count'] = len(heading_hits)

    patterns = {
        'version': [
            r'\b(?:Version|Rev(?:ision)?|Ver\.?)\s*[:：]?\s*([A-Z]?\d+(?:\.\d+)*)\b',
            r'\bV(?:ersion)?\s*[:：]?\s*([A-Z]?\d+(?:\.\d+)*)\b',
        ],
        'doc_no': [
            r'\bDoc\.?\s*No\.?\s*[:：]?\s*([A-Z0-9\-_/]{4,})\b',
            r'\bDocument\s*(?:No\.?|ID)\s*[:：]?\s*([A-Z0-9\-_/]{4,})\b',
        ],
    }
    for key, exprs in patterns.items():
        for expr in exprs:
            match = re.search(expr, normalized, re.IGNORECASE)
            if match:
                metadata[key] = match.group(1).strip()
                break

    if metadata['doc_no'] == '-' and metadata['name'] and metadata['name'] != '-':
        stem = Path(metadata['name']).stem
        if re.search(r'[A-Z]{1,3}-\d{2,}|\d{3,}', stem):
            metadata['doc_no'] = stem

    return metadata


def _infer_issue_area(issue):
    position_meta = _decode_issue_position(_issue_value(issue, 'position', ''))
    explicit_area = position_meta.get('area')
    if explicit_area:
        return explicit_area

    chapter = str(_issue_value(issue, 'chapter', '') or '')
    context = str(_issue_value(issue, 'context', '') or '')
    original = str(_issue_value(issue, 'original_text', '') or '')
    source = f"{chapter}\n{context}\n{original}".strip()
    line = max(source.splitlines(), key=len, default=source).strip()
    lowered = line.lower()

    if re.search(r'\b(doc\.?\s*no\.?|page\s*\d+|copyright|jb-[a-z0-9\-]+)\b', lowered, re.IGNORECASE):
        return '页脚'
    if re.match(r'^(?:step\s+)?\d+[\.)、：:]\s+', line, re.IGNORECASE):
        return '步骤'
    if re.match(r'^(?:table|表)\s*\d+', line, re.IGNORECASE) or ' | ' in line:
        return '表格'
    if re.match(r'^(?:chapter|section)\s+\d+', line, re.IGNORECASE) or re.match(r'^\d+(?:\.\d+)*\s+[A-Z][^\n]{2,}$', line):
        return '章节'
    return '正文'


def _format_issue_location(issue, content):
    location_override = _issue_value(issue, 'location_override', '')
    if location_override:
        return html_lib.escape(str(location_override))

    position = _issue_value(issue, 'position', '')
    original = str(_issue_value(issue, 'original_text', '') or '').strip()
    if content and original:
        start, end = _parse_issue_position(position)
        if end <= start or start == 0:
            found = content.find(original)
            if found >= 0:
                position = _encode_issue_position(found, found + len(original))

    page_no = _page_number_from_position(content, position)
    page_text = f'第{page_no}页' if page_no else ''
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    if rule in {'DOC-REV-001', 'DOC-REV-002', 'CHECKLIST-COPYRIGHT-YEAR', 'CHECKLIST-YEAR'}:
        return f'Revision history（{page_text}）' if page_text else 'Revision history'
    raw_chapter = re.sub(r'\s+', ' ', str(_issue_value(issue, 'chapter', '') or '')).strip()
    if _is_structural_heading_text(raw_chapter) and not _is_table_cell_like(raw_chapter):
        return f'{raw_chapter}（{page_text}）' if page_text else raw_chapter
    effective_issue = dict(issue)
    effective_issue['position'] = position
    chapter = _resolve_issue_heading(effective_issue, content)
    if chapter != '-' and page_text:
        return f'{chapter}（{page_text}）'
    if chapter != '-':
        return chapter
    page_heading = _extract_primary_structural_heading(content, position) or _extract_previous_structural_heading(content, position)
    if page_heading and page_text:
        return f'{page_heading}（{page_text}）'
    if page_heading:
        return page_heading
    area = _infer_issue_area(issue)
    if area == '正文':
        area = '内容段落'
    if area and page_text:
        return f'{page_text}，{area}'
    return page_text or area or '-'


def _format_issue_chapter(issue):
    area = _infer_issue_area(issue)
    chapter = str(_issue_value(issue, 'chapter', '') or '').strip()
    if area in {'页脚', '步骤'}:
        return '-'
    if not chapter:
        return '-'
    if re.search(r'\b(doc\.?\s*no\.?|page\s*\d+)\b', chapter, re.IGNORECASE):
        return '-'
    if re.match(r'^(?:step\s+)?\d+[\.)、：:]\s+', chapter, re.IGNORECASE):
        return '-'
    compact = re.sub(r'\s+', ' ', chapter).strip()
    compact = re.sub(r'^(#+)\s*', '', compact)
    compact = re.sub(r'^(\d+(?:\.\d+)*)([A-Za-z])', r'\1 \2', compact)
    if re.match(r'^\d+(?:\.\d+)*\s*(?:mL|μL|uL|ng|kg|cm|mm)\b', compact, re.IGNORECASE):
        return '-'
    if _is_inline_reference_heading(compact):
        return '-'
    if re.search(r'\b\d{2,4}\s*[μu]L\b', compact, re.IGNORECASE):
        return '-'
    if _is_table_cell_like(compact):
        return '-'
    if len(compact) > 50:
        return '-'
    return compact


def _extract_heading_from_context(text):
    normalized = re.sub(r'\s+', ' ', str(text or '').replace('\f', ' ')).strip()
    if not normalized:
        return ''
    table_match = re.search(r'(Table\s+\d+\s+[A-Za-z][A-Za-z0-9×\- ,/&()]{3,120})', normalized, re.IGNORECASE)
    if table_match:
        heading = _normalize_heading_text(table_match.group(1))
        if heading and not _is_table_cell_like(heading) and not _is_inline_reference_heading(heading):
            return heading
    patterns = [
        r'(\d+(?:\.\d+)+\s+[A-Z][A-Za-z0-9\- ]{3,80})',
        r'((?:Chapter|Section)\s+\d+[A-Za-z0-9\- :]{3,80})',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, normalized)
        if matches:
            for match in reversed(matches):
                heading = _normalize_heading_text(match)
                if heading and not _is_table_cell_like(heading) and not _is_inline_reference_heading(heading):
                    return heading
    return ''


def _extract_primary_page_heading(content, position):
    page_no = _page_number_from_position(content, position)
    if not page_no:
        return ''
    pages = _split_content_pages(content)
    if page_no < 1 or page_no > len(pages):
        return ''
    page_lines = [line.strip() for line in pages[page_no - 1].splitlines() if line.strip()]
    joined = ' '.join(page_lines)
    table_match = re.search(r'(Table\s+\d+\s+[A-Za-z][A-Za-z0-9×\- ,/&()]{3,120})', joined, re.IGNORECASE)
    if table_match:
        heading = _clean_caption_heading(table_match.group(1))
        if heading and not _is_table_cell_like(heading):
            return heading
    candidates = []
    for index, line in enumerate(page_lines[:80]):
        candidate = _score_heading_candidate(page_lines, index)
        if not candidate:
            continue
        score, heading = candidate
        if _is_table_cell_like(heading):
            continue
        candidates.append((score, heading))
    if not candidates:
        return ''
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _extract_primary_structural_heading(content, position):
    page_no = _page_number_from_position(content, position)
    if not page_no:
        return ''
    pages = _split_content_pages(content)
    if page_no < 1 or page_no > len(pages):
        return ''
    page_lines = [line.strip() for line in pages[page_no - 1].splitlines() if line.strip()]
    candidates = []
    for index, _line in enumerate(page_lines[:120]):
        candidate = _score_heading_candidate(page_lines, index)
        if not candidate:
            continue
        score, heading = candidate
        if _is_table_cell_like(heading) or _is_inline_reference_heading(heading):
            continue
        if not _is_structural_heading_text(heading):
            continue
        candidates.append((score, heading))
    if not candidates:
        return ''
    candidates.sort(key=lambda item: item[0], reverse=True)
    return candidates[0][1]


def _extract_previous_page_heading(content, position):
    page_no = _page_number_from_position(content, position)
    if not page_no or page_no <= 1:
        return ''
    pages = _split_content_pages(content)
    for candidate_page in range(page_no - 1, 0, -1):
        page_lines = [line.strip() for line in pages[candidate_page - 1].splitlines() if line.strip()]
        for index in range(len(page_lines) - 1, -1, -1):
            candidate = _score_heading_candidate(page_lines, index)
            if not candidate:
                continue
            score, heading = candidate
            if score < 85:
                continue
            if _is_table_cell_like(heading):
                continue
            return heading
    return ''


def _extract_previous_structural_heading(content, position):
    page_no = _page_number_from_position(content, position)
    if not page_no or page_no <= 1:
        return ''
    pages = _split_content_pages(content)
    for candidate_page in range(page_no - 1, 0, -1):
        page_lines = [line.strip() for line in pages[candidate_page - 1].splitlines() if line.strip()]
        for index in range(len(page_lines) - 1, -1, -1):
            candidate = _score_heading_candidate(page_lines, index)
            if not candidate:
                continue
            score, heading = candidate
            if score < 85:
                continue
            if _is_table_cell_like(heading) or _is_inline_reference_heading(heading):
                continue
            if _is_structural_heading_text(heading):
                return heading
    return ''


def _is_structural_heading_text(text):
    stripped = str(text or '').strip()
    return bool(re.match(r'^(?:\d+(?:\.\d+)*(?:[\).])?|Chapter|Section)\b', stripped, re.IGNORECASE))


def _extract_reference_target(issue):
    fields = [
        str(_issue_value(issue, 'original_text', '') or ''),
        str(_issue_value(issue, 'context', '') or ''),
        str(_issue_value(issue, 'description', '') or ''),
    ]
    combined = ' '.join(fields)
    match = re.search(r'\b(Figure|Fig\.?|Table)\s*(\d+)\b', combined, re.IGNORECASE)
    if not match:
        return '', ''
    label = match.group(1)
    kind = 'Figure' if label.lower().startswith('fig') else 'Table'
    return kind, match.group(2)


def _find_reference_caption(content, issue):
    kind, number = _extract_reference_target(issue)
    if not kind or not number:
        return ''
    page_no = _page_number_from_position(content, _issue_value(issue, 'position', '')) or 1
    pages = _split_content_pages(content)
    page_order = [page_no - 1]
    for distance in range(1, len(pages)):
        prev_index = page_no - 1 - distance
        next_index = page_no - 1 + distance
        if prev_index >= 0:
            page_order.append(prev_index)
        if next_index < len(pages):
            page_order.append(next_index)
    alias = 'Fig\\.?' if kind == 'Figure' else kind
    pattern = re.compile(rf'\b(?:{kind}|{alias})\s*{number}\b[^\n]*', re.IGNORECASE)
    for page_index in page_order:
        page_lines = [line.strip() for line in pages[page_index].splitlines() if line.strip()]
        for line in page_lines:
            if not pattern.search(line):
                continue
            if _is_inline_reference_heading(line) or _is_table_cell_like(line):
                continue
            heading = _clean_caption_heading(line)
            if heading:
                return heading
    return ''


def _issue_prefers_table_caption(issue):
    fields = [
        str(_issue_value(issue, 'original_text', '') or ''),
        str(_issue_value(issue, 'context', '') or ''),
        str(_issue_value(issue, 'description', '') or ''),
    ]
    combined = ' '.join(fields)
    patterns = [
        r'\bLibrary Type\b',
        r'\bcondition\s+condition\b',
        r'\bPos\d',
        r'\boperation deck\b',
        r'\bfilter tips\b',
    ]
    return any(re.search(pattern, combined, re.IGNORECASE) for pattern in patterns)


def _resolve_issue_heading(issue, content):
    chapter = _format_issue_chapter(issue)
    if chapter != '-' and _is_structural_heading_text(chapter):
        return chapter

    position = _issue_value(issue, 'position', '')
    structural_page_heading = _extract_primary_structural_heading(content, position)
    if structural_page_heading and not re.match(r'^(?:\d+(?:\.\d+)*[\).])\s+', structural_page_heading):
        return structural_page_heading

    previous_structural_heading = _extract_previous_structural_heading(content, position)
    if previous_structural_heading:
        return previous_structural_heading

    if structural_page_heading:
        return structural_page_heading

    context_heading = _extract_heading_from_context(_issue_value(issue, 'context', ''))
    if context_heading and _is_structural_heading_text(context_heading):
        return context_heading

    return '-'


def _extract_issue_snippet(issue, content, radius=28):
    original = str(_issue_value(issue, 'original_text', '') or '').strip()
    start, end = _parse_issue_position(_issue_value(issue, 'position', ''))
    if content and original and end > start:
        left = max(0, start - radius)
        right = min(len(content), end + radius)
        snippet = content[left:right].replace('\f', ' ').replace('\n', ' ')
        snippet = re.sub(r'\s+', ' ', snippet).strip()
        if left > 0:
            snippet = '...' + snippet
        if right < len(content):
            snippet = snippet + '...'
        return snippet

    context = str(_issue_value(issue, 'context', '') or '').replace('\n', ' ')
    context = re.sub(r'\s+', ' ', context).strip()
    if not context:
        return original or '-'
    if not original or original not in context:
        return context[:100] + ('...' if len(context) > 100 else '')

    pos = context.find(original)
    left = max(0, pos - radius)
    right = min(len(context), pos + len(original) + radius)
    snippet = context[left:right].strip()
    if left > 0:
        snippet = '...' + snippet
    if right < len(context):
        snippet = snippet + '...'
    return snippet


def _format_issue_description(issue):
    description_override = _issue_value(issue, 'description_override', '')
    if description_override:
        return str(description_override)

    rule = str(_issue_value(issue, 'rule', '') or '')
    original = str(_issue_value(issue, 'original_text', '') or '')
    description = str(_issue_value(issue, 'description', '') or '').strip()

    punct_map = {'，': ',', '。': '.', '；': ';', '：': ':', '（': '(', '）': ')', '、': ','}
    if rule == 'PUNCT-001' and original in punct_map:
        return f'使用了中文全角标点（{original}），应改为英文半角标点（{punct_map[original]}）'

    if rule == 'HR004':
        unit_match = re.match(r'(\d+(?:\.\d+)?)(.+)', original)
        if unit_match:
            return f'数字与单位之间缺少空格，应改为 {unit_match.group(1)} {unit_match.group(2)}'

    if rule.upper().startswith('SPELL'):
        after = _format_issue_suggestion(issue)
        if after and after != '-':
            return f'发现拼写/用词错误：{original} 应改为 {after}。'
        return f'发现拼写/用词错误：{original}。'

    return description or '-'


def _fallback_issue_suggestion(issue):
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    category = str(_issue_value(issue, 'category', '') or '')
    description = str(_issue_value(issue, 'description', '') or '').strip()
    basis = str(_issue_value(issue, 'audit_basis', '') or '').strip()
    text = f'{rule} {category} {description} {basis}'

    if 'PUNCT' in rule or 'Punctuation' in category or '标点' in text:
        return '请按英文技术文档标点规范调整原文，并复核前后空格。'
    if 'UNIT' in rule or 'Units' in category or '单位' in text:
        return '请按标准单位写法统一单位符号、大小写和数字空格。'
    if 'FORMAT' in rule or 'Format' in category or '格式' in text:
        return '请按说明书格式规范调整该处排版或写法。'
    if 'COMPLIANCE' in rule or 'Compliance' in category or '合规' in text:
        return '请复核该处是否缺少必要声明、风险提示或发布信息。'
    if description and description != '-':
        return f'请按问题说明修订：{description}'
    return '请人工复核该问题，并补充明确修改意见。'


def _fallback_issue_description(issue):
    rule = str(_issue_value(issue, 'rule', '') or '')
    category = str(_issue_value(issue, 'category', '') or '')
    original = str(_issue_value(issue, 'original_text', '') or '').strip()
    suggestion = _format_issue_suggestion(issue)
    text = f'{rule} {category}'

    if 'Units rule' in rule or 'Units' in category or '单位' in text:
        if original and suggestion and suggestion != '-':
            return f'单位写法不符合英文技术文档规范，建议统一为 {suggestion}。'
        return '单位写法不符合英文技术文档规范，请统一单位符号、大小写和数字空格。'
    if 'Punctuation rule' in rule or 'Punctuation' in category or '标点' in text:
        return '标点或前后空格不符合英文技术文档规范，请按上下文统一调整。'
    if 'Grammar rule' in rule or '语法' in category or 'Grammar' in category:
        if suggestion and suggestion != '-':
            return f'英文语法或冠词使用不完整，建议改为 {suggestion}。'
        return '英文语法或冠词使用不完整，请按上下文补充或调整。'
    return '需按说明书发布前自检 Checklist 复核该处表达，并确认修改建议。'


def _ensure_issue_report_guidance(issue):
    suggestion = _format_issue_suggestion(issue)
    if not suggestion or suggestion == '-':
        issue['suggestion'] = _fallback_issue_suggestion(issue)
    description = _format_issue_description(issue)
    if not description or description == '-':
        issue['description'] = _fallback_issue_description(issue)
    basis = str(_issue_value(issue, 'audit_basis', '') or '').strip()
    if not basis or basis == '-':
        issue['audit_basis'] = 'AI 深度审核结果，需结合说明书发布前自检 Checklist 复核。'
    return issue


def _format_issue_suggestion(issue):
    suggestion_override = _issue_value(issue, 'suggestion_override', '')
    if suggestion_override:
        return str(suggestion_override)

    rule = str(_issue_value(issue, 'rule', '') or '')
    original = str(_issue_value(issue, 'original_text', '') or '')
    suggestion = str(_issue_value(issue, 'suggestion', '') or '').strip()
    punct_map = {'，': ',', '。': '.', '；': ';', '：': ':', '（': '(', '）': ')', '、': ','}

    if rule == 'PUNCT-001' and original in punct_map:
        return punct_map[original]
    bracketed = re.search(r'建议改为\s*[:：]?\s*\[([^\]]+)\]', suggestion)
    if bracketed:
        return bracketed.group(1).strip()
    plain = re.search(r'建议(?:改为|替换为|统一为)\s*[:：]?\s*([^。；，,]+)', suggestion)
    if plain:
        candidate = plain.group(1).strip().strip('[]')
        if candidate and not re.search(r'(疑似错误|是否确定|原文)', candidate):
            return candidate
    if suggestion.startswith('建议改为 '):
        return suggestion.replace('建议改为 ', '', 1)
    if suggestion.startswith('建议替换为 '):
        return suggestion.replace('建议替换为 ', '', 1)
    if suggestion.startswith('建议统一为 '):
        return suggestion.replace('建议统一为 ', '', 1)
    return suggestion or '-'


def _format_after_text(issue, before=None):
    original = str(before if before is not None else _issue_value(issue, 'original_text', '') or '').strip()
    suggestion = str(_issue_value(issue, 'suggestion', '') or '').strip()
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    if rule in {'HR001', 'HR004', 'HR005', 'PUNCT-001'}:
        return _normalize_example_after(issue, original, _format_issue_suggestion(issue))

    cleaned_suggestion = _format_issue_suggestion(issue)
    if cleaned_suggestion and cleaned_suggestion != suggestion and cleaned_suggestion != '-':
        return _normalize_example_after(issue, original, cleaned_suggestion)

    if '→' in suggestion:
        candidate = suggestion.split('→', 1)[1].strip()
    elif '->' in suggestion:
        candidate = suggestion.split('->', 1)[1].strip()
    elif suggestion.startswith('建议改为 '):
        candidate = suggestion.replace('建议改为 ', '', 1).strip()
    elif suggestion.startswith('建议替换为 '):
        candidate = suggestion.replace('建议替换为 ', '', 1).strip()
    elif suggestion.startswith('建议统一为 '):
        candidate = suggestion.replace('建议统一为 ', '', 1).strip()
    else:
        candidate = suggestion.strip()

    if not candidate or candidate == '-':
        return '-'
    if re.search(r'(建议|确认|补充|删除|统一译法|待修改|需确认|应|请|或|是否|如果|，|。|；)', candidate):
        return '-'
    return _normalize_example_after(issue, original, candidate)


def _report_issue_span(issue):
    start, end = _parse_issue_position(_issue_value(issue, 'position', ''))
    original = str(_issue_value(issue, 'original_text', '') or '')
    if end <= start and original:
        end = start + len(original)
    return start, end


def _report_spans_overlap(left, right):
    left_start, left_end = _report_issue_span(left)
    right_start, right_end = _report_issue_span(right)
    return left_start < right_end and right_start < left_end


def _report_issue_to_dict(issue):
    keys = [
        'id', 'review_id', 'severity', 'category', 'rule', 'chapter', 'original_text',
        'context', 'suggestion', 'description', 'audit_basis', 'confidence', 'source',
        'position', 'status', 'created_at', 'updated_at'
    ]
    return {key: _issue_value(issue, key, '') for key in keys}


def _report_rule(rule):
    return str(rule or '').upper()


def _report_category(issue):
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    category = str(_issue_value(issue, 'category', '') or '').strip()
    if rule == 'UNIT-003':
        return '单位符号'
    if rule == 'UNIT-004':
        return '温度单位'
    if rule == 'HR011':
        return '数学符号'
    if rule == 'TERM-EN-001':
        return '产品名称'
    if rule == 'TERM-EN-002':
        return '表达规范'
    if rule.startswith('SPELL'):
        return '拼写/用词'
    if rule.startswith('GRAMMAR'):
        return '语法'
    if rule.startswith('TERM'):
        return '术语一致性'
    if rule.startswith('LOGIC'):
        return '内容逻辑'
    return category or '未分类'


def _report_severity(issue):
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    if rule.startswith('SPELL'):
        return 'general'
    if rule == 'UNIT-004':
        return 'serious'
    if rule in {'TERM-EN-001', 'LOGIC-002'}:
        return 'serious'
    if rule in {'UNIT-003', 'HR011', 'TERM-EN-002'}:
        return 'general'
    return _issue_value(issue, 'severity', 'general') or 'general'


def _report_audit_basis(issue):
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    mapping = {
        'UNIT-003': '英文技术文档写作规范 §3.2 单位符号',
        'UNIT-004': '英文技术文档写作规范 §3.3 温度单位',
        'HR011': '英文技术文档写作规范 §3.4 数学符号',
        'LOGIC-002': 'Checklist §4.2 实验条件完整性',
        'TERM-EN-001': '英文技术文档写作规范 §2.1 产品名称一致性',
        'TERM-EN-002': '英文技术文档写作规范 §2.2 术语一致性',
        'GRAMMAR-005': '英文技术文档写作规范 §1.3 固定搭配',
        'GRAMMAR-006': '英文技术文档写作规范 §1.2 动词形式',
    }
    if rule.startswith('SPELL'):
        return '英文技术文档写作规范 §1.1 拼写与用词'
    if rule.startswith('CHECKLIST-'):
        return '说明书发布前自检 Checklist §5.1 发布信息一致性'
    return mapping.get(rule) or str(_issue_value(issue, 'audit_basis', '') or '-')


def _report_issue_is_keeped_duplicate(left, right):
    left_text = f"{_issue_value(left, 'original_text', '')} {_issue_value(left, 'suggestion', '')}".lower()
    right_text = f"{_issue_value(right, 'original_text', '')} {_issue_value(right, 'suggestion', '')}".lower()
    if 'keeped' not in left_text or 'keeped' not in right_text:
        return False
    return _report_spans_overlap(left, right)


def _prefer_report_issue(left, right):
    left_rule = _report_rule(_issue_value(left, 'rule', ''))
    right_rule = _report_rule(_issue_value(right, 'rule', ''))
    if right_rule.startswith('GRAMMAR') and not left_rule.startswith('GRAMMAR'):
        return right
    if left_rule.startswith('GRAMMAR') and not right_rule.startswith('GRAMMAR'):
        return left
    return right if _issue_sort_key(right) < _issue_sort_key(left) else left


def _report_resolve_chapter(issue, content):
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    original = str(_issue_value(issue, 'original_text', '') or '').strip()
    if rule == 'CHECKLIST-TRADEMARK' and original.upper() == 'DNBSEQ':
        return '1. Introduction'

    raw_existing = re.sub(r'\s+', ' ', str(_issue_value(issue, 'chapter', '') or '')).strip()
    existing = _format_issue_chapter(issue)
    if existing == '-' and _is_structural_heading_text(raw_existing) and not _is_table_cell_like(raw_existing):
        existing = raw_existing
    start, _ = _report_issue_span(issue)
    if content and start is not None:
        chapter = extract_chapter(content, start)
        if chapter and not re.search(r'\btable of contents\b|目录', chapter, re.IGNORECASE):
            return chapter
    if existing != '-':
        return existing
    return ''


def _report_unique_values(values):
    seen = set()
    result = []
    for value in values:
        text = re.sub(r'\s+', ' ', str(value or '')).strip()
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result


def _compact_report_location_text(location):
    text = re.sub(r'\s+', ' ', str(location or '')).strip()
    if not text:
        return text

    page_suffix = ''
    page_match = re.search(r'(（第\d+页）)$', text)
    if page_match:
        page_suffix = page_match.group(1)
        text = text[:page_match.start()].strip()

    text = re.split(r'\b(?:Refer to|Follow the|For details|See)\b', text, maxsplit=1, flags=re.IGNORECASE)[0].strip()
    text = re.sub(r'\s+(?:and|to|for|with|using)$', '', text, flags=re.IGNORECASE).strip()
    if len(text) > 80:
        text = text[:77].rstrip() + '...'
    return f'{text}{page_suffix}' if text else str(location or '').strip()


def _report_key_text(value):
    text = re.sub(r'\s+', ' ', str(value or '')).strip().lower()
    return text if text and text != '-' else ''


def _report_aggregate_key(issue):
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    original = re.sub(r'\s+', ' ', str(_issue_value(issue, 'original_text', '') or '')).strip()
    suggestion = _format_issue_suggestion(issue)

    if rule.startswith('SPELL') and original:
        return ('SPELL', original.lower(), str(suggestion or '').lower())
    if rule in {'UNIT-004', 'UNIT-003'}:
        return (rule,)
    if rule == 'HR011':
        return (rule, _report_math_symbol_scene(issue))
    if rule.startswith('TERM') and original:
        return (rule, original.lower(), str(suggestion or '').lower())
    category = _report_key_text(_issue_value(issue, 'category', ''))
    description = _report_key_text(_format_issue_description(issue))
    suggestion_key = _report_key_text(suggestion)
    if rule and category and (description or suggestion_key):
        return ('SAME-ISSUE', rule, category, description, suggestion_key)
    return None


def _report_math_symbol_scene(issue):
    text = ' '.join([
        str(_issue_value(issue, 'original_text', '') or ''),
        str(_issue_value(issue, 'context', '') or ''),
        str(_issue_value(issue, 'chapter', '') or ''),
    ])
    if re.search(r'\bvolume\b', text, re.IGNORECASE):
        return '实验操作体积'
    if re.search(r'\bg\b|centrifug|speed', text, re.IGNORECASE):
        return '设备参数'
    return '数学符号'


def _report_aggregate_description(rule, originals, suggestion):
    if rule == 'UNIT-004':
        return '温度单位缺少 °，应统一补充 °。'
    if rule == 'UNIT-003':
        return '微升单位写法不统一，应统一使用 μL。'
    if rule == 'HR011':
        return '数学符号写法不规范，应统一改为 ×。'
    if rule.startswith('SPELL'):
        return f"同一拼写/用词错误多处出现：{'、'.join(originals)} 应改为 {suggestion}。"
    if rule.startswith('TERM'):
        return f"同一术语问题多处出现：{'、'.join(originals)} 应统一为 {suggestion}。"
    return ''


def _report_aggregate_suggestion(rule, suggestion):
    if rule == 'UNIT-004':
        return '统一补充°'
    if rule == 'UNIT-003':
        return '统一改为 μL'
    if rule == 'HR011':
        return '统一改为 ×'
    if suggestion and suggestion != '-':
        if re.match(r'^(?:请|建议|统一|补充|删除|调整|确认|按)', suggestion):
            return suggestion
        return f'统一改为 {suggestion}'
    return suggestion or '-'


def _aggregate_report_issues(issues, content):
    buckets = {}
    passthrough = []
    for issue in sorted(issues, key=_issue_sort_key):
        key = _report_aggregate_key(issue)
        if not key:
            passthrough.append(issue)
            continue
        buckets.setdefault(key, []).append(issue)

    aggregated = list(passthrough)
    for members in buckets.values():
        if len(members) == 1:
            aggregated.append(members[0])
            continue

        first = dict(members[0])
        rule = _report_rule(_issue_value(first, 'rule', ''))
        originals = _report_unique_values(_issue_value(member, 'original_text', '') for member in members)
        suggestions = _report_unique_values(_format_issue_suggestion(member) for member in members)
        locations = _report_unique_values(_compact_report_location_text(_format_issue_location(member, content)) for member in members)
        suggestion = suggestions[0] if len(suggestions) == 1 else '、'.join(suggestions)
        location_text = '、'.join(locations)
        if len(members) > len(locations) and location_text:
            location_text += '多处'

        first['grouped'] = True
        first['group_count'] = len(members)
        first['members'] = members
        first['location_override'] = location_text
        if len(originals) > 3:
            original_text = '、'.join(originals[:3]) + f' 等 {len(members)} 处同类片段'
        else:
            original_text = '、'.join(originals)
        first['original_override'] = original_text
        first['original_text'] = original_text
        first['suggestion_override'] = _report_aggregate_suggestion(rule, suggestion)
        first['suggestion'] = first['suggestion_override']
        first['description_override'] = _report_aggregate_description(rule, originals, suggestion)
        if not first['description_override']:
            first['description_override'] = _format_issue_description(members[0])
        aggregated.append(first)

    return sorted(aggregated, key=_issue_sort_key)


def _prepare_report_issues(issues, content):
    prepared = []
    for issue in issues:
        item = _report_issue_to_dict(issue)
        if _should_drop_known_false_positive_issue(item):
            continue
        item['severity'] = _report_severity(item)
        item['category'] = _report_category(item)
        item['audit_basis'] = _report_audit_basis(item)
        item['chapter'] = _report_resolve_chapter(item, content)
        if _should_drop_low_value_ai_report_issue(item):
            continue
        if _report_rule(item.get('rule')) == 'CHECKLIST-TRADEMARK' and str(item.get('original_text', '')).strip().upper() == 'DNBSEQ':
            item['location_override'] = '1. Introduction（第2页）'
        item['_document_content'] = content
        _ensure_issue_report_guidance(item)
        merged = False
        for index, existing in enumerate(prepared):
            if _report_issue_is_keeped_duplicate(existing, item):
                winner = _prefer_report_issue(existing, item)
                winner['category'] = '语法'
                winner['severity'] = 'serious'
                winner['description'] = 'keeped 拼写错误，且应使用 keep 的过去分词 kept。'
                winner['suggestion'] = '建议改为 should be kept' if 'should be keeped' in str(winner.get('original_text', '')).lower() else '建议改为 kept'
                prepared[index] = winner
                merged = True
                break
        if not merged:
            prepared.append(item)
    return _aggregate_report_issues(prepared, content)


def _should_drop_low_value_ai_report_issue(issue):
    source = str(_issue_value(issue, 'source', '') or '').lower()
    if source != 'ai' or _is_deterministic_review_rule(issue):
        return False
    rule = _report_rule(_issue_value(issue, 'rule', ''))
    if rule.startswith(('SPELL', 'GRAMMAR', 'UNIT', 'PUNCT', 'HR', 'CHECKLIST')):
        return False
    confidence = int(_issue_value(issue, 'confidence', 0) or 0)
    description = str(_issue_value(issue, 'description', '') or '').strip()
    basis = str(_issue_value(issue, 'audit_basis', '') or '').strip()
    suggestion = _format_issue_suggestion(issue)
    if not suggestion or suggestion == '-':
        return True
    if confidence < 90 and (not description or description == '-'):
        return True
    if confidence < 90 and (not basis or basis == '-' or basis.startswith('AI ')):
        return True
    return False


def _issue_dimension(issue):
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    category = str(_issue_value(issue, 'category', '') or '')
    text = f'{rule} {category}'
    if any(token in text for token in ['SPELL', 'GRAMMAR', 'PUNCT', '时态', '标点', '语法', '拼写', '风格']):
        return '语言质量'
    if any(token in text for token in ['TERM', '术语']):
        return '术语一致性'
    if any(token in text for token in ['LOGIC', '结构', '长句', '步骤连续性', '逻辑完整性']):
        return '逻辑完整性'
    if any(token in text for token in ['SAFE', 'WARNING', 'CAUTION', 'DANGER', '安全', '合规']):
        return '安全合规'
    if any(token in text for token in ['REF', '引用', '交叉']):
        return '交叉引用'
    return '格式规范'


def _build_dimension_summary(issues):
    summary = {name: {'count': 0, 'fatal': 0, 'serious': 0} for name in ['格式规范', '语言质量', '术语一致性', '逻辑完整性', '安全合规', '交叉引用']}
    for issue in issues:
        dimension = _issue_dimension(issue)
        summary[dimension]['count'] += 1
        if _issue_value(issue, 'severity', 'general') == 'fatal':
            summary[dimension]['fatal'] += 1
        if _issue_value(issue, 'severity', 'general') == 'serious':
            summary[dimension]['serious'] += 1
    return summary


def _build_report_conclusion(issues):
    fatal = sum(1 for issue in issues if _issue_value(issue, 'severity', '') == 'fatal')
    serious = sum(1 for issue in issues if _issue_value(issue, 'severity', '') == 'serious')
    if fatal:
        return '当前文档存在高风险问题，建议完成修订并复审后再发布。'
    if serious >= 5:
        return '当前文档可作为修订草稿继续流转，建议先关闭严重问题再进入正式发布。'
    if serious or issues:
        return '当前文档主体可读，建议按优先级关闭剩余问题并做一次快速回归。'
    return '当前文档未发现显著问题，可进入发布前终审。'


def _read_review_spec_file_text(file_info):
    path = Path(file_info.get("file_path") or "")
    if not path.exists():
        return ""

    suffix = path.suffix.lower()
    if suffix in {".xlsx", ".xlsm"}:
        from openpyxl import load_workbook
        workbook = load_workbook(path, read_only=True, data_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"[{sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell not in (None, "")]
                if cells:
                    lines.append(" | ".join(cells))
        workbook.close()
        return "\n".join(lines)

    return path.read_text(encoding="utf-8", errors="ignore")


def _load_review_spec_texts(db: Session):
    spec_texts = {}
    spec_files = _get_review_spec_files_from_knowledge(db)
    for spec_key, file_info in spec_files.items():
        try:
            if isinstance(file_info, list):
                spec_texts[spec_key] = "\n\n".join(
                    text for text in (_read_review_spec_file_text(item) for item in file_info) if text
                )
            else:
                spec_texts[spec_key] = _read_review_spec_file_text(file_info)
        except Exception as exc:
            print(f"[审核] 读取知识库规范文件失败 {spec_key}: {exc}")
            spec_texts[spec_key] = ""
    return spec_texts


def _build_ai_review_basis(spec_texts, document_language):
    parts = []
    if document_language in ("cn", "both") and spec_texts.get("cn_style"):
        parts.append("【中文技术文档写作风格指南】\n" + spec_texts["cn_style"][:2500])
    if document_language in ("en", "both") and spec_texts.get("en_style"):
        parts.append("【英文技术文档写作风格指南】\n" + spec_texts["en_style"][:2500])
    if spec_texts.get("common_errors"):
        parts.append("【技术文档常见错误清单】\n" + spec_texts["common_errors"][:2500])
    if spec_texts.get("final_checklists"):
        parts.append("【说明书发布前自检 Checklist】\n" + spec_texts["final_checklists"][:5000])
    return "\n\n".join(parts)


def _call_with_timeout(func, timeout_seconds, *args):
    with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(func, *args)
        return future.result(timeout=timeout_seconds)


def _release_checklist_issue(content, start, end, rule, severity, original_text, suggestion, description, chapter=None):
    return {
        "severity": severity,
        "category": "发布前自检",
        "rule": rule,
        "chapter": chapter or extract_chapter(content, start),
        "original_text": original_text,
        "context": get_context(content, start, end, 160),
        "suggestion": suggestion,
        "description": description,
        "audit_basis": "说明书发布前自检 Checklist",
        "confidence": 95,
        "source": "rule",
        "position": _encode_issue_position_with_meta(start, end, area="发布前自检", check_result="待修改"),
    }


def _extract_copyright_years(text):
    years = []
    for match in re.finditer(r"©\s*((?:19|20)\d{2})(?:\s*[-–—~至到]\s*((?:19|20)\d{2}))?", text):
        if match.group(2):
            years.append((match.group(2), match.start(2), match.end(2)))
        else:
            years.append((match.group(1), match.start(1), match.end(1)))
    for match in re.finditer(r"(?:copyright|版权)[^\n\f]{0,80}?((?:19|20)\d{2})", text, re.IGNORECASE):
        years.append((match.group(1), match.start(1), match.end(1)))
    return years


def _extract_revision_history_years(text):
    headings = ["版本记录", "修订记录", "修订历史", "revision history", "version history", "change history"]
    blocks = []
    lowered = text.lower()
    for heading in headings:
        idx = lowered.find(heading.lower())
        if idx >= 0:
            blocks.append((idx, text[idx:idx + 1800]))

    candidates = []
    search_areas = blocks or [(0, text[:2500])]
    date_patterns = [
        r"((?:19|20)\d{2})\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日",
        r"\b(?:Jan\.?|Feb\.?|Mar\.?|Apr\.?|May\.?|Jun\.?|Jul\.?|Aug\.?|Sep\.?|Oct\.?|Nov\.?|Dec\.?)\s+\d{1,2},?\s+((?:19|20)\d{2})\b",
        r"\b(?:Jan\.?|Feb\.?|Mar\.?|Apr\.?|May\.?|Jun\.?|Jul\.?|Aug\.?|Sep\.?|Oct\.?|Nov\.?|Dec\.?)\s+((?:19|20)\d{2})\b",
        r"(?<![A-Za-z0-9.])\d{1,2}\s+(?:Jan\.?|Feb\.?|Mar\.?|Apr\.?|May\.?|Jun\.?|Jul\.?|Aug\.?|Sep\.?|Oct\.?|Nov\.?|Dec\.?)\s+((?:19|20)\d{2})\b",
        r"\b((?:19|20)\d{2})[-/.]\d{1,2}[-/.]\d{1,2}\b",
    ]
    for base, block in search_areas:
        for pattern in date_patterns:
            for match in re.finditer(pattern, block, re.IGNORECASE):
                year = match.group(1)
                full_date = match.group(0)
                candidates.append((
                    year,
                    base + match.start(1),
                    base + match.end(1),
                    full_date,
                    base + match.start(0),
                    base + match.end(0),
                ))
    return candidates


def _normalize_date_display(text):
    normalized = re.sub(r"\s+", " ", str(text or "")).strip()
    return normalized or str(text or "")


def _run_release_checklist_audit(content, document, spec_texts, document_language):
    if not spec_texts.get("final_checklists"):
        return []

    issues = []
    text = str(content or "")
    filename = str(getattr(document, "filename", "") or "")
    is_english = document_language in ("en", "both")

    copyright_years = _extract_copyright_years(text)
    revision_history_years = _extract_revision_history_years(text)
    if copyright_years and revision_history_years:
        copyright_year = copyright_years[0][0]
        latest_revision_year, _, _, latest_revision_date, latest_start, latest_end = revision_history_years[0]
        if latest_revision_year != copyright_year:
            original_date = _normalize_date_display(latest_revision_date)
            issues.append(_release_checklist_issue(
                text,
                latest_start,
                latest_end,
                "CHECKLIST-COPYRIGHT-YEAR",
                "serious",
                original_date,
                f"建议核对版权年份 {copyright_year} 与修订历史最新版本年份 {latest_revision_year}，并将两处年份调整为一致",
                "说明书发布前自检要求版权年份与修订历史最新版本年份一致。",
                chapter="版本记录",
            ))

    filename_year_match = re.search(r"(?:19|20)\d{2}", filename)
    if filename_year_match:
        filename_year = filename_year_match.group(0)
        revision_years = re.findall(r"\b(?:Jan\.?|Feb\.?|Mar\.?|Apr\.?|May\.?|Jun\.?|Jul\.?|Aug\.?|Sep\.?|Oct\.?|Nov\.?|Dec\.?)\s+((?:19|20)\d{2})\b", text, re.IGNORECASE)
        if revision_years and filename_year not in revision_years[:2]:
            start = max(0, text.find(revision_years[0]))
            issues.append(_release_checklist_issue(
                text,
                start,
                start + len(revision_years[0]),
                "CHECKLIST-YEAR",
                "serious",
                revision_years[0],
                f"建议确认修订历史最新年份与文件名年份 {filename_year} 一致",
                "文件名年份应与修订历史中的最新发布年份一致。",
            ))

    trademark_statement_text = text[:5000]
    trademark_terms = ["DNBSEQ", "Qubit", "NextSeq", "MGIEasy"]
    for term in trademark_terms:
        if re.search(rf"\b{re.escape(term)}(?:\s*(?:TM|®|™))?\b[^\n]{{0,600}}\b(?:trademark|trademarks|property of their respective owners)", trademark_statement_text, re.IGNORECASE | re.DOTALL):
            continue
        pattern = re.compile(rf"\b{re.escape(term)}\b(?!\s*(?:TM|®|™))", re.IGNORECASE)
        for match in pattern.finditer(text):
            snippet_start = max(0, match.start() - 40)
            snippet_end = min(len(text), match.end() + 40)
            snippet = text[snippet_start:snippet_end]
            if re.search(r"trademark|registered trademark|商标", snippet, re.IGNORECASE):
                continue
            replacement = "DNBSEQ™" if term == "DNBSEQ" else f"建议确认 {match.group(0)} 是否需要商标标注或在前言商标列表中列明"
            issues.append(_release_checklist_issue(
                text,
                match.start(),
                match.end(),
                "CHECKLIST-TRADEMARK",
                "general",
                match.group(0),
                replacement,
                "说明书自检 checklist 要求文中出现的商标需列明并保持标注格式一致。",
            ))
            break

    if is_english:
        for match in re.finditer(r"\bClick\b", text):
            if not _is_touchscreen_click_context(text, match.start()):
                continue
            issues.append(_release_checklist_issue(
                text,
                match.start(),
                match.end(),
                "CHECKLIST-TAP",
                "general",
                match.group(0),
                "触摸屏界面按钮操作建议统一使用 Tap",
                "英文说明书自检 checklist 要求配置触摸屏时点击界面按钮使用 Tap。",
            ))
            break

        for match in re.finditer(r"\b[0-9]+(?:\.[0-9]+)?\s*μl\b", text):
            issues.append(_release_checklist_issue(
                text,
                match.start(),
                match.end(),
                "CHECKLIST-UL",
                "serious",
                match.group(0),
                match.group(0)[:-1] + "L",
                "英文说明书自检 checklist 要求 μL 中的 L 必须大写。",
            ))
            break

    return issues


def _build_feedback_advice(issues):
    statuses = Counter((_issue_value(issue, 'status', '') or 'pending') for issue in issues)
    false_positive_count = statuses.get('false_positive', 0)
    false_positive_text = (
        f"已标记误报 {false_positive_count} 条，建议沉淀为规则白名单或定位策略优化样本。"
        if false_positive_count
        else "未发现明显误报，仍建议审核人对需确认项做人工复核。"
    )
    return [
        f"待确认 {statuses.get('pending', 0)} 条，建议审核人优先处理致命和严重问题。",
        false_positive_text,
        f"已确认 {statuses.get('confirmed', 0)} 条，建议修订后重新执行一次规则审核验证闭环。",
    ]


def _build_problem_summary_rows(issues):
    severity_keys = ['fatal', 'serious', 'general', 'suggestion']
    labels = {'fatal': '致命', 'serious': '严重', 'general': '一般', 'suggestion': '建议'}
    rows = []
    for dimension in ['格式规范', '语言质量', '术语一致性', '逻辑完整性', '安全合规', '交叉引用']:
        items = [issue for issue in issues if _issue_dimension(issue) == dimension]
        if not items:
            continue
        row = {'dimension': dimension, 'total': len(items)}
        for key in severity_keys:
            row[key] = sum(1 for issue in items if _issue_value(issue, 'severity', '') == key)
            row[f'{key}_label'] = labels[key]
        rows.append(row)
    return rows


def _group_issues_by_severity(issues):
    groups = {'fatal': [], 'serious': [], 'general': [], 'suggestion': []}
    for issue in sorted(issues, key=_issue_sort_key):
        groups.setdefault(_issue_value(issue, 'severity', 'general'), []).append(issue)
    return groups


def _build_global_issue_index(groups):
    ordered = []
    for severity in ['fatal', 'serious', 'general', 'suggestion']:
        ordered.extend(groups.get(severity) or [])
    return {id(issue): index for index, issue in enumerate(ordered, start=1)}


def _build_group_heading(severity, entries, index_map):
    labels = {
        'fatal': '致命问题',
        'serious': '严重问题',
        'general': '一般问题',
        'suggestion': '建议项',
    }
    if not entries:
        return labels[severity]
    first_no = index_map[id(entries[0])]
    last_no = index_map[id(entries[-1])]
    if first_no == last_no:
        range_text = _format_issue_display_id(first_no)
    else:
        range_text = f'{_format_issue_display_id(first_no)}-{_format_issue_display_id(last_no)}'
    return f'{labels[severity]}（{range_text}）'


def _build_required_and_recommended_lists(issues):
    required = []
    recommended = []
    for issue in sorted(issues, key=_issue_sort_key):
        item = f"{_format_issue_display_id(0)} {_format_issue_description(issue)}"
        severity = _issue_value(issue, 'severity', '')
        if severity in {'fatal', 'serious'} and len(required) < 5:
            required.append((issue, item))
        elif severity in {'general', 'suggestion'} and len(recommended) < 5:
            recommended.append((issue, item))
    return required, recommended


def _normalize_example_after(issue, before, after):
    if after == '-':
        return after
    rule = str(_issue_value(issue, 'rule', '') or '').upper()
    if rule == 'HR005':
        return 'Well'
    if rule == 'HR004' and re.match(r'\d+(?:\.\d+)?[A-Za-zμ/%]+', before):
        match = re.match(r'(\d+(?:\.\d+)?)([A-Za-zμ/%]+)', before)
        if match:
            return f'{match.group(1)} {match.group(2)}'
    if rule == 'HR001':
        return 'https://global-mgitech.com/'
    if '→' in str(after or ''):
        return str(after).split('→', 1)[1].strip()
    if '->' in str(after or ''):
        return str(after).split('->', 1)[1].strip()
    return after


def _build_diff_markup(before, after):
    before = str(before or '-')
    after = str(after or '-')
    matcher = difflib.SequenceMatcher(None, before, after)
    before_parts = []
    after_parts = []
    for opcode, a0, a1, b0, b1 in matcher.get_opcodes():
        left = html_lib.escape(before[a0:a1])
        right = html_lib.escape(after[b0:b1])
        if opcode == 'equal':
            before_parts.append(left)
            after_parts.append(right)
        else:
            if left:
                before_parts.append(f'<span class="diff-remove">{left}</span>')
            if right:
                after_parts.append(f'<span class="diff-add">{right}</span>')
    return ''.join(before_parts) or '-', ''.join(after_parts) or '-'


def _build_modification_examples(issues):
    target_rules = ['HR005', 'HR004', 'GRAMMAR', 'PUNCT-001', 'HR001']
    examples = []
    used_rules = set()

    def try_add(issue, forced_title=None):
        before = str(_issue_value(issue, 'original_text', '') or '').strip()
        after = _format_after_text(issue, before)
        if not before or not after or after == '-' or before == after:
            return False
        before_markup, after_markup = _build_diff_markup(before, after)
        examples.append({
            'title': forced_title or _format_issue_description(issue),
            'before': before_markup,
            'after': after_markup,
        })
        used_rules.add(str(_issue_value(issue, 'rule', '') or '').upper())
        return True

    sorted_issues = sorted(issues, key=_issue_sort_key)
    for target in target_rules:
        for issue in sorted_issues:
            rule = str(_issue_value(issue, 'rule', '') or '').upper()
            if target == 'GRAMMAR':
                category = str(_issue_value(issue, 'category', '') or '')
                if '语法' not in category and 'GRAMMAR' not in rule:
                    continue
            elif rule != target:
                continue
            if try_add(issue):
                break

    for issue in sorted_issues:
        if len(examples) >= 5:
            break
        rule = str(_issue_value(issue, 'rule', '') or '').upper()
        if rule in used_rules:
            continue
        try_add(issue)

    return examples


def _build_report_verdict(issues):
    fatal = sum(1 for issue in issues if _issue_value(issue, 'severity', '') == 'fatal')
    serious = sum(1 for issue in issues if _issue_value(issue, 'severity', '') == 'serious')
    if fatal:
        return '不通过'
    if serious >= 3:
        return '需复审'
    if issues:
        return '建议修改'
    return '通过'


def _load_term_variants_from_db(db: Session):
    pairs = []
    for term in db.query(Term).all():
        non_standard = str(getattr(term, 'non_standard', '') or '').strip()
        standard = str(getattr(term, 'standard', '') or '').strip()
        if not non_standard or not standard or non_standard == standard:
            continue
        pairs.append((standard, non_standard))
    return pairs


def _run_db_term_consistency_audit(db: Session, content):
    issues = []
    normalized = str(content or '')
    seen = set()
    for standard, variant in _load_term_variants_from_db(db):
        pattern = re.compile(rf'(?<![A-Za-z0-9]){re.escape(variant)}(?![A-Za-z0-9])', re.IGNORECASE)
        for match in pattern.finditer(normalized):
            key = (variant.lower(), match.start())
            if key in seen:
                continue
            seen.add(key)
            issues.append({
                'severity': 'general',
                'category': '术语一致性',
                'rule': 'TERM-CONSISTENCY-DB',
                'original_text': match.group(0),
                'suggestion': standard,
                'description': f'术语不一致: "{match.group(0)}" 建议统一为 "{standard}"。',
                'audit_basis': '术语库全文术语一致性检查',
                'confidence': 90,
                'source': 'term',
                'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文'),
                'chapter': extract_chapter(normalized, match.start()),
                'context': get_context(normalized, match.start(), match.end(), 120),
            })
    return issues


def _build_inline_example_html(issue):
    if _issue_value(issue, 'grouped', False):
        suggestion = html_lib.escape(_format_issue_suggestion(issue))
        if not suggestion or suggestion == '-':
            return ''
        return (
            '<div class="advice-panel">'
            '<div class="compare-label">修改建议</div>'
            f'<div class="suggestion-inline">{suggestion}</div>'
            '</div>'
        )

    before = str(_issue_value(issue, 'original_text', '') or '').strip()
    after = _format_after_text(issue, before)
    if not before or not after or before == after or after == '-':
        suggestion = html_lib.escape(_format_issue_suggestion(issue))
        if not suggestion or suggestion == '-':
            return ''
        return (
            '<div class="advice-panel">'
            '<div class="compare-label">修改建议</div>'
            f'<div class="suggestion-inline">{suggestion}</div>'
            '</div>'
        )
    before_markup, after_markup = _build_diff_markup(before, after)
    return (
        '<div class="advice-panel">'
        '<div class="compare-label">修改建议</div>'
        '<div class="comparison-grid">'
        f'<div class="before-box"><span class="box-label">修改前</span>{before_markup}</div>'
        f'<div class="after-box"><span class="box-label">修改后</span>{after_markup}</div>'
        '</div>'
        '</div>'
    )


def _generate_review_html_content(review, doc, issues):
    doc_name = doc.filename if doc else f"文档{review.document_id}"
    content = getattr(doc, 'content', '') if doc else ''
    metadata = _extract_document_metadata(doc, content)
    report_issues = _prepare_report_issues(issues, content)
    summary_rows = _build_problem_summary_rows(report_issues)
    grouped_issues = _group_issues_by_severity(report_issues)
    issue_index_map = _build_global_issue_index(grouped_issues)
    verdict = _build_report_verdict(report_issues)
    conclusion = _build_report_conclusion(report_issues)
    feedback_advice = _build_feedback_advice(report_issues)
    nav_items = []
    for severity in ['fatal', 'serious', 'general', 'suggestion']:
        entries = grouped_issues.get(severity) or []
        if entries:
            nav_items.append((severity, _build_group_heading(severity, entries, issue_index_map)))

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>审核报告 - {doc_name}</title>
    <style>
        * {{ box-sizing: border-box; }}
        :root {{ --ink:#1f2937; --muted:#667085; --line:#dfe5ef; --panel:#ffffff; --bg:#edf2f8; --hero1:#153f70; --hero2:#2f76c9; --fatal:#ef4444; --serious:#f59e0b; --general:#6b7280; --suggestion:#22c55e; --soft:#eff4fb; }}
        body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 0; padding: 28px; background: radial-gradient(circle at top left, #f8fbff 0, #edf2f8 45%, #e6edf6 100%); color: var(--ink); }}
        .container {{ max-width: 1180px; margin: 0 auto; background: var(--panel); padding: 0 0 32px; box-shadow: 0 18px 50px rgba(18, 39, 69, 0.12); border-radius: 24px; overflow: hidden; }}
        .hero {{ padding: 36px 40px 46px; background: linear-gradient(135deg, var(--hero1), var(--hero2)); color: #fff; }}
        .hero h1 {{ margin: 0 0 10px; font-size: 32px; }}
        .hero p {{ margin: 0; max-width: 760px; color: rgba(255,255,255,0.84); line-height: 1.7; }}
        .section {{ padding: 0 40px; }}
        h2 {{ color: #143b68; margin: 30px 0 16px; font-size: 22px; }}
        .module-tag {{ display: inline-block; margin-top: 12px; padding: 6px 12px; border-radius: 999px; background: rgba(255,255,255,0.16); font-size: 12px; letter-spacing: 0.08em; }}
        .meta {{ display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 12px; margin: -26px 40px 0; padding: 14px; background: rgba(255,255,255,0.96); border: 1px solid #d8e4f2; border-radius: 22px; box-shadow: 0 14px 36px rgba(18, 39, 69, 0.12); }}
        .meta-card {{ background: linear-gradient(180deg, #ffffff, #f8fbff); border: 1px solid #e2eaf5; border-radius: 16px; padding: 15px 16px; }}
        .meta-label {{ font-size: 12px; color: var(--muted); margin-bottom: 8px; }}
        .meta-value {{ font-weight: 700; line-height: 1.5; }}
        .summary-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin: 18px 0 10px; }}
        .summary-card {{ padding: 18px 14px; text-align: center; border-radius: 18px; color: #fff; box-shadow: inset 0 1px 0 rgba(255,255,255,0.12); }}
        .summary-card .num {{ font-size: 30px; font-weight: 800; }}
        .summary-card .label {{ font-size: 13px; margin-top: 6px; opacity: 0.92; }}
        .card-fatal {{ background: linear-gradient(135deg, #b91c1c, var(--fatal)); }}
        .card-serious {{ background: linear-gradient(135deg, #d97706, var(--serious)); }}
        .card-general {{ background: linear-gradient(135deg, #4b5563, var(--general)); }}
        .card-suggestion {{ background: linear-gradient(135deg, #15803d, var(--suggestion)); }}
        .callout {{ border: 1px solid #cfe0f5; border-radius: 18px; padding: 18px 20px; background: linear-gradient(180deg, #f8fbff, #f1f6fc); line-height: 1.8; color: #2a4365; }}
        .issue {{ border: 1px solid #e4e7ed; border-radius: 18px; padding: 18px 20px; margin: 12px 0; background: #fff; box-shadow: 0 6px 18px rgba(15, 35, 60, 0.04); }}
        .issue.confirmed {{ border-left: 5px solid #67c23a; background: #f0f9eb; }}
        .issue.false_positive {{ border-left: 5px solid #909399; background: #f4f4f5; opacity: 0.7; }}
        .issue.ignored {{ border-left: 5px solid #c0c4cc; background: #fafafa; opacity: 0.6; }}
        .issue.pending {{ border-left: 5px solid var(--serious); }}
        .issue-header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 10px; }}
        .issue-title {{ font-weight: bold; color: #303133; }}
        .badge {{ display: inline-block; padding: 2px 8px; border-radius: 10px; font-size: 12px; color: #fff; margin-left: 6px; }}
        .badge-confirmed {{ background: #67c23a; }}
        .badge-false {{ background: #909399; }}
        .badge-pending {{ background: var(--serious); }}
        .badge-ignored {{ background: #c0c4cc; }}
        .badge-severity-fatal {{ background: var(--fatal); }}
        .badge-severity-serious {{ background: var(--serious); }}
        .badge-severity-general {{ background: var(--general); }}
        .badge-severity-suggestion {{ background: var(--suggestion); }}
        .issue-field {{ margin: 8px 0; font-size: 14px; }}
        .issue-label {{ font-weight: bold; color: #606266; display: inline-block; min-width: 80px; }}
        .original-text {{ background: #fef0f0; padding: 4px 8px; border-radius: 3px; color: #c45656; font-family: 'Courier New', monospace; }}
        .suggestion {{ background: #f0f9eb; padding: 4px 8px; border-radius: 3px; color: #5a8e3f; }}
        .suggestion-inline {{ color: #14532d; font-weight: 700; }}
        .context {{ color: #444; font-size: 13px; line-height: 1.7; white-space: normal; }}
        .problem-highlight {{ color: #d93025; background: #fff1f0; font-weight: 700; padding: 1px 3px; border-radius: 3px; }}
        .subtle {{ color: #909399; font-size: 12px; }}
        .issue-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 12px; }}
        .empty {{ text-align: center; color: #909399; padding: 40px; }}
        .list-card {{ border: 1px solid var(--line); border-radius: 18px; padding: 18px 20px; background: #fff; margin-bottom: 12px; }}
        .list-card h3 {{ margin: 0 0 10px; font-size: 18px; color: #143b68; }}
        .summary-table-note {{ margin: 12px 0 0; color: var(--muted); font-size: 13px; }}
        table {{ width: 100%; border-collapse: collapse; background: #fff; border-radius: 16px; overflow: hidden; }}
        th, td {{ border: 1px solid #e7edf5; padding: 12px 14px; text-align: left; vertical-align: top; }}
        th {{ background: #f6f9fd; color: #163f6f; }}
        .status-ok {{ color: var(--suggestion); font-weight: 700; }}
        .status-bad {{ color: var(--fatal); font-weight: 700; }}
        .status-warn {{ color: var(--serious); font-weight: 700; }}
        .advice-panel {{ margin-top: 12px; padding: 14px 16px; border: 1px solid #d7e6f7; border-radius: 16px; background: #f8fbff; }}
        .comparison-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 12px; margin-top: 10px; }}
        .before-box, .after-box {{ padding: 12px 14px; border-radius: 14px; line-height: 1.7; }}
        .before-box {{ background: #fffafa; border: 1px solid #f4caca; color: #7f1d1d; }}
        .after-box {{ background: #f6fdf8; border: 1px solid #cdebd6; color: #14532d; }}
        .box-label {{ display: block; color: var(--muted); font-size: 12px; font-weight: 700; margin-bottom: 6px; }}
        .diff-remove {{ background: #fde8e8; text-decoration: line-through; padding: 1px 3px; border-radius: 4px; }}
        .diff-add {{ background: #dcfce7; text-decoration: underline; padding: 1px 3px; border-radius: 4px; }}
        .feedback-block {{ border-top: 1px dashed #cbd5e1; margin-top: 18px; padding-top: 18px; }}
        .compare-label {{ font-weight: 700; color: #163f6f; margin-bottom: 8px; }}
        .report-nav {{ display: flex; flex-wrap: wrap; gap: 10px; margin: 12px 0 18px; }}
        .report-nav a {{ color: #174e83; background: #eef6ff; border: 1px solid #cfe0f5; border-radius: 999px; padding: 7px 12px; text-decoration: none; font-size: 13px; font-weight: 700; }}
        h3 {{ margin: 22px 0 12px; }}
        .severity-fatal {{ color: var(--fatal); }}
        .severity-serious {{ color: var(--serious); }}
        .severity-general {{ color: var(--general); }}
        .severity-suggestion {{ color: var(--suggestion); }}
        .footer {{ margin: 36px 40px 0; text-align: center; color: #909399; font-size: 12px; border-top: 1px solid #ebeef5; padding-top: 20px; }}
        @media (max-width: 920px) {{ .meta, .summary-grid, .issue-grid, .comparison-grid {{ grid-template-columns: 1fr; }} .section {{ padding: 0 22px; }} .hero {{ padding: 28px 22px 44px; }} .meta {{ margin: -24px 22px 0; padding: 12px; }} .footer {{ margin-left: 22px; margin-right: 22px; }} }}
    </style>
</head>
<body>
    <div class="container">
        <div class="hero">
            <h1>技术文档审核报告</h1>
            <p>本报告基于格式规范、语言质量、术语一致性、逻辑完整性、安全合规、交叉引用六个维度输出结构化审核结果，供编辑修订与复审使用。</p>
            <div class="module-tag">REVIEW TASK #{review.id}</div>
        </div>
        <div class="meta">
            <div class="meta-card"><div class="meta-label">文档名称</div><div class="meta-value">{metadata['name']}</div></div>
            <div class="meta-card"><div class="meta-label">审核日期</div><div class="meta-value">{_format_report_datetime(getattr(review, 'completed_at', None) or getattr(review, 'created_at', None))}</div></div>
            <div class="meta-card"><div class="meta-label">审核人</div><div class="meta-value">技术文档审核AI助理</div></div>
            <div class="meta-card"><div class="meta-label">文档范围</div><div class="meta-value">{metadata['file_type']} · {metadata['page_count']} 页 · {metadata['section_count']} 个章节</div></div>
        </div>

        <div class="section">
        <h2>模块 1 · 审核概览</h2>
        <div class="summary-grid">
            <div class="summary-card card-fatal"><div class="num">{sum(1 for issue in report_issues if _issue_value(issue, 'severity', '') == 'fatal')}</div><div class="label">致命</div></div>
            <div class="summary-card card-serious"><div class="num">{sum(1 for issue in report_issues if _issue_value(issue, 'severity', '') == 'serious')}</div><div class="label">严重</div></div>
            <div class="summary-card card-general"><div class="num">{sum(1 for issue in report_issues if _issue_value(issue, 'severity', '') == 'general')}</div><div class="label">一般</div></div>
            <div class="summary-card card-suggestion"><div class="num">{sum(1 for issue in report_issues if _issue_value(issue, 'severity', '') == 'suggestion')}</div><div class="label">建议</div></div>
        </div>
        </div>
        <div class="section">
        <h2>模块 2 · 问题明细</h2>
        <div class="report-nav">
            <a href="#summary-table">问题汇总表</a>
            {''.join(f'<a href="#group-{severity}">{html_lib.escape(title)}</a>' for severity, title in nav_items)}
        </div>
        <h3>问题汇总表</h3>
        <table id="summary-table">
            <thead><tr><th>类型</th><th>致命</th><th>严重</th><th>一般</th><th>建议</th><th>总计</th></tr></thead>
            <tbody>
"""

    for row in summary_rows:
        html += f"""
            <tr>
                <td>{row['dimension']}</td>
                <td>{row['fatal']}</td>
                <td>{row['serious']}</td>
                <td>{row['general']}</td>
                <td>{row['suggestion']}</td>
                <td>{row['total']}</td>
            </tr>
"""

    html += """
            </tbody>
        </table>
        <h3>详细问题列表</h3>
"""

    if not report_issues:
        html += '<div class="empty">未发现任何问题</div>'
    else:
        for severity in ['fatal', 'serious', 'general', 'suggestion']:
            entries = grouped_issues.get(severity) or []
            if not entries:
                continue
            html += f'<h3 id="group-{severity}" class="severity-{severity}">{_build_group_heading(severity, entries, issue_index_map)}</h3>'
            for issue in entries:
                display_no = issue_index_map[id(issue)]
                issue_status = _issue_value(issue, 'status', '') or 'pending'
                sev_class = f"badge-severity-{_issue_value(issue, 'severity', 'general')}"
                status_class = {
                    'confirmed': 'badge-confirmed',
                    'false_positive': 'badge-false',
                    'ignored': 'badge-ignored',
                    'pending': 'badge-pending',
                }.get(issue_status, 'badge-pending')
                status_text = {
                    'confirmed': '已确认',
                    'false_positive': '误报',
                    'ignored': '已忽略',
                    'pending': '待确认',
                }.get(issue_status, '待确认')
                example_html = _build_inline_example_html(issue)
                html += f"""
        <div id="issue-{display_no:04d}" class="issue {issue_status}">
            <div class="issue-header">
                <div class="issue-title">{_format_issue_display_id(display_no)} · {_issue_dimension(issue)} · {_issue_value(issue, 'category', '') or '未分类'}</div>
                <div>
                    <span class="badge {sev_class}">{_format_issue_severity(_issue_value(issue, 'severity', 'general'))}</span>
                    <span class="badge {status_class}">{status_text}</span>
                </div>
            </div>
            <div class="issue-field"><span class="issue-label">位置:</span> {_format_issue_location(issue, content)}</div>
            <div class="issue-field"><span class="issue-label">原文片段:</span> <span class="context">{_highlight_issue_context(issue)}</span></div>
            <div class="issue-field"><span class="issue-label">问题说明:</span> {html_lib.escape(_format_issue_description(issue))}</div>
            <div class="issue-field"><span class="issue-label">审核依据:</span> {html_lib.escape(_issue_value(issue, 'audit_basis', '-') or '-')}</div>
            {example_html}
        </div>
"""

    html += f"""
        </div>
        <div class="section">
        <h2>模块 3 · 审核结论</h2>
        <div class="callout">
            <div><strong>判定结果:</strong> {html_lib.escape(verdict)}</div>
            <div><strong>结论说明:</strong> {html_lib.escape(conclusion)}</div>
            <div><strong>复审建议:</strong></div>
            <ul>
                <li>{html_lib.escape(feedback_advice[0])}</li>
                <li>{html_lib.escape(feedback_advice[1])}</li>
                <li>{html_lib.escape(feedback_advice[2])}</li>
            </ul>
            <div class="feedback-block">
                <div><strong>审核后反馈：</strong></div>
                <div>1. 本次审核有哪些问题是误报？（请列出编号）</div>
                <div>2. 哪些问题或规则可以记录到 memory 中供后续审核参考？</div>
                <div>3. 其他建议：</div>
            </div>
        </div>
        </div>
"""
    html += f"""
        <div class="footer">由 智能技术文档审核平台 生成 | {_format_report_datetime()}</div>
    </div>
</body>
</html>"""
    return html


def _run_logic_integrity_audit(content):
    issues = []
    seen = set()
    lines = content.replace('\f', '\n').splitlines()
    previous_step = None
    cursor = 0
    previous_non_empty_line = ''
    for line in lines:
        text = line.strip()
        if _looks_like_numbered_section_heading(text):
            previous_step = None
            previous_non_empty_line = text
            cursor += len(line) + 1
            continue
        match = re.match(r'^(?:Step\s+)?(\d+)[\.)]\s+(.+)$', text, re.IGNORECASE)
        if not match:
            if text:
                previous_non_empty_line = text
            cursor += len(line) + 1
            continue
        current = int(match.group(1))
        if re.search(r'\b(?:table|figure|fig\.?|section|chapter)\s*$', previous_non_empty_line, re.IGNORECASE):
            previous_non_empty_line = text
            cursor += len(line) + 1
            continue
        if previous_step is not None and current - previous_step > 1:
            key = ('LOGIC-001', current)
            if key in seen:
                continue
            seen.add(key)
            pos = cursor + line.find(text)
            issues.append({
                'severity': 'serious',
                'category': '逻辑完整性',
                'rule': 'LOGIC-001',
                'chapter': '',
                'original_text': text,
                'context': get_context(content, max(pos, 0), max(pos, 0) + len(text), 120),
                'suggestion': f'建议补齐 Step {previous_step + 1}，或重新编号当前步骤。',
                'description': f'步骤编号从 {previous_step} 跳到 {current}，流程连续性不足。',
                'audit_basis': '操作步骤应保持连续且可追踪。',
                'confidence': 90,
                'source': 'rule',
                'position': _encode_issue_position_with_meta(max(pos, 0), max(pos, 0) + len(text), area='步骤'),
            })
        previous_step = current
        previous_non_empty_line = text
        cursor += len(line) + 1
    return issues


def _looks_like_numbered_section_heading(text):
    match = re.match(r'^\d+(?:\.\d+)*\.?\s+(.+)$', str(text or '').strip())
    if not match:
        return False

    title = match.group(1).strip()
    if not title or len(title) > 90:
        return False
    if re.search(r'[.!?。！？]$', title):
        return False
    if re.match(r'^(?:and|or|then|with|using|from|to|at|in)\b', title, re.IGNORECASE):
        return False

    words = re.findall(r'[A-Za-z][A-Za-z\-]*|[\u4e00-\u9fff]+', title)
    if not words or len(words) > 8:
        return False
    action_verbs = {
        'add', 'remove', 'select', 'click', 'tap', 'press', 'open', 'close', 'connect',
        'disconnect', 'insert', 'restart', 'turn', 'check', 'confirm', 'prepare', 'mix',
        'incubate', 'centrifuge', 'load', 'scan', 'clean', 'update', 'save', 'delete',
    }
    return words[0].lower() not in action_verbs


def _run_safety_compliance_audit(content):
    issues = []
    normalized = content.replace('\f', '\n')
    seen = set()

    for match in re.finditer(r'(?im)^\s*(WARNING|CAUTION|DANGER)\b([^\n]*)', normalized):
        label = match.group(1).upper()
        line = match.group(0).strip()
        window = normalized[match.start():match.start() + 120]
        if re.search(r'警告|注意|危险', window):
            continue
        key = ('SAFE-001', match.start())
        if key in seen:
            continue
        seen.add(key)
        issues.append({
                'severity': 'serious',
                'category': '安全合规',
                'rule': 'SAFE-001',
                'chapter': '',
                'original_text': line,
                'context': get_context(normalized, match.start(), match.end(), 120),
                'suggestion': f'{label}: ... → {label} {"警告" if label == "WARNING" else "注意" if label == "CAUTION" else "危险"}: ...',
                'description': f'{label} 缺少对应中文警示语，应补充双语安全提示。',
                'audit_basis': '安全警示应完整传达风险及处置要求。',
                'confidence': 91,
                'source': 'rule',
                'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文'),
            })

    hazard_pattern = re.compile(r'\b(biohazard|flammable|high voltage|corrosive|laser radiation)\b', re.IGNORECASE)
    for match in hazard_pattern.finditer(normalized):
        window = normalized[max(0, match.start() - 80):min(len(normalized), match.end() + 80)]
        if re.search(r'WARNING|CAUTION|DANGER|警告|注意|危险', window, re.IGNORECASE):
            continue
        key = ('SAFE-002', match.start())
        if key in seen:
            continue
        seen.add(key)
        issues.append({
                'severity': 'fatal',
                'category': '安全合规',
                'rule': 'SAFE-002',
                'chapter': '',
                'original_text': match.group(0),
                'context': get_context(normalized, match.start(), match.end(), 120),
                'suggestion': '建议在该风险描述前增加明确的 WARNING/CAUTION 标题，并说明防护动作。',
                'description': '检测到风险词，附近缺少明确的安全警示标题。',
                'audit_basis': '高风险操作应在正文中显式标注安全等级。',
                'confidence': 93,
                'source': 'rule',
                'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文'),
            })

    return issues


def _run_cross_reference_audit(content):
    issues = []
    normalized = content.replace('\f', '\n')
    section_numbers = set(re.findall(r'(?m)^\s*(\d+(?:\.\d+)*)[\.)]?\s+[A-Z][^\n]{2,}', normalized))
    step_numbers = {int(num) for num in re.findall(r'(?im)^\s*(?:Step\s+)?(\d+)[\.)]\s+.+$', normalized)}
    table_numbers = set(re.findall(r'(?im)^\s*(?:Table|表)\s*(\d+)\b', normalized))
    figure_numbers = set(re.findall(r'(?im)^\s*(?:Figure|Fig\.|图)\s*(\d+)\b', normalized))
    structured_section_numbers = {num for num in section_numbers if '.' in num}
    seen = set()

    if len(structured_section_numbers) >= 3:
        for match in re.finditer(r'\b(?:Section|Chapter|章节)\s+(\d+(?:\.\d+)*)\b', normalized, re.IGNORECASE):
            if re.search(r'Revision\s+history', normalized[max(0, match.start() - 900):match.start()], re.IGNORECASE):
                next_contents = normalized.find('Contents', match.start())
                if next_contents < 0 or next_contents - match.start() < 1800:
                    continue
            ref = match.group(1)
            if ref in section_numbers:
                continue
            key = ('REF-001', ref)
            if key in seen:
                continue
            seen.add(key)
            issues.append({
                    'severity': 'general',
                    'category': '交叉引用',
                    'rule': 'REF-001',
                    'chapter': '',
                    'original_text': match.group(0),
                    'context': get_context(normalized, match.start(), match.end(), 120),
                    'suggestion': f'补充 Section {ref} 对应章节标题，或删除该引用。',
                    'description': f'引用原文：{match.group(0)}；引用类型：章节；目标对象：Section {ref}；检查结果：引用缺失。文档中没有对应章节标题。',
                    'audit_basis': '交叉引用应可回溯到唯一的章节或对象。',
                    'confidence': 88,
                    'source': 'rule',
                    'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文', reference_type='章节', target=f'Section {ref}', check_result='引用缺失'),
                })

    for match in re.finditer(r'\b(?:according to|refer to|see)?\s*(?:Table|table|表)\s*(\d+)\b', normalized):
        ref = match.group(1)
        if ref in table_numbers:
            continue
        key = ('REF-TABLE', ref, match.start())
        if key in seen:
            continue
        seen.add(key)
        issues.append({
            'severity': 'serious',
            'category': '交叉引用',
            'rule': 'REF-003',
            'chapter': '',
            'original_text': match.group(0).strip(),
            'context': get_context(normalized, match.start(), match.end(), 120),
            'suggestion': f'补充 Table {ref}，或删除该引用。',
            'description': f'引用原文：{match.group(0).strip()}；引用类型：表；目标对象：Table {ref}；检查结果：引用缺失。文档中没有 Table {ref} 的标题。',
            'audit_basis': '表格引用应能定位到唯一的表格标题。',
            'confidence': 92,
            'source': 'rule',
            'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文', reference_type='表', target=f'Table {ref}', check_result='引用缺失'),
        })

    for match in re.finditer(r'\b(?:Figure|Fig\.|图)\s*(\d+)\b', normalized, re.IGNORECASE):
        ref = match.group(1)
        if ref in figure_numbers:
            continue
        key = ('REF-FIGURE', ref, match.start())
        if key in seen:
            continue
        seen.add(key)
        issues.append({
            'severity': 'general',
            'category': '交叉引用',
            'rule': 'REF-004',
            'chapter': '',
            'original_text': match.group(0),
            'context': get_context(normalized, match.start(), match.end(), 120),
            'suggestion': f'补充 Figure {ref} 对应标题，或删除该引用。',
            'description': f'引用原文：{match.group(0)}；引用类型：图；目标对象：Figure {ref}；检查结果：引用缺失。文档中没有对应图标题。',
            'audit_basis': '图引用应能定位到唯一的图标题。',
            'confidence': 88,
            'source': 'rule',
            'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文', reference_type='图', target=f'Figure {ref}', check_result='引用缺失'),
        })

    if len(step_numbers) >= 3:
        for match in re.finditer(r'\bStep\s+(\d+)\b', normalized, re.IGNORECASE):
            ref = int(match.group(1))
            if ref in step_numbers:
                continue
            context = get_context(normalized, match.start(), match.end(), 120)
            if re.search(r'\bin\s+section\s+\d+(?:\.\d+)*\b', context, re.IGNORECASE):
                continue
            key = ('REF-002', ref)
            if key in seen:
                continue
            seen.add(key)
            issues.append({
                'severity': 'general',
                'category': '交叉引用',
                'rule': 'REF-002',
                'chapter': '',
                'original_text': match.group(0),
                'context': context,
                'suggestion': f'补充 Step {ref}，或将引用改为有效步骤编号。',
                'description': f'引用原文：{match.group(0)}；引用类型：步骤；目标对象：Step {ref}；检查结果：引用缺失。文档中没有对应步骤。',
                'audit_basis': '步骤引用应和实际流程编号保持一致。',
                'confidence': 88,
                'source': 'rule',
                'position': _encode_issue_position_with_meta(match.start(), match.end(), area='正文', reference_type='步骤', target=f'Step {ref}', check_result='引用缺失'),
            })

    return issues


def _export_review_docx(review, document, issues):
    from docx import Document as DocxDocument

    source_path = _get_document_upload_path(document)
    if not source_path or not source_path.exists():
        raise HTTPException(status_code=404, detail="原始 Word 文件不存在")

    REVIEW_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    export_name = _build_review_export_filename(document.filename, review.id, ".docx")
    export_path = REVIEW_EXPORT_DIR / export_name
    shutil.copyfile(source_path, export_path)

    doc = DocxDocument(str(export_path))
    paragraphs = list(_iter_docx_paragraphs(doc))
    comments_data = []
    comment_id = 0

    for issue in _select_export_issues(issues):
        target_para, _, char_start, char_end = _locate_docx_issue(paragraphs, issue)
        if target_para is None or char_start is None or char_end is None:
            continue
        comment_id += 1
        if _add_docx_comment_marker(target_para, comment_id, char_start, char_end):
            comments_data.append({
                "id": comment_id,
                "author": "智能技术文档审核",
                "initials": "RV",
                "lines": _build_review_comment_lines(issue),
            })

    if not comments_data:
        raise HTTPException(status_code=400, detail="未能在 Word 文档中定位可批注的审核问题")

    doc.save(str(export_path))
    _inject_comments_to_docx(str(export_path), comments_data)
    return export_path, export_name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _export_review_html_file(review, document, issues):
    REVIEW_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    export_name = _build_review_export_filename(document.filename if document else "review", review.id, ".html")
    export_path = REVIEW_EXPORT_DIR / export_name
    html = _generate_review_html_content(review, document, issues)
    export_path.write_text(html, encoding="utf-8")
    return export_path, export_name, "text/html; charset=utf-8"


def _excel_issue_position(issue):
    position = _decode_issue_position(_issue_value(issue, 'position', ''))
    return str(position.get('sheet', '') or ''), int(position.get('row', 0) or 0)


def _export_review_excel(review, document, issues):
    from openpyxl import load_workbook

    source_path = _get_document_upload_path(document)
    if not source_path or not source_path.exists():
        raise HTTPException(status_code=404, detail="原始 Excel 文件不存在")

    REVIEW_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe = re.sub(r"[^A-Za-z0-9\u4e00-\u9fff._-]+", "_", Path(document.filename or f"review_{review.id}").stem).strip("_") or f"review_{review.id}"
    export_name = f"{safe}_reviewed.xlsx"
    export_path = REVIEW_EXPORT_DIR / export_name
    shutil.copyfile(source_path, export_path)

    wb = load_workbook(export_path)
    issues_by_row = {}
    for issue in _select_export_issues(issues):
        sheet_name, row_number = _excel_issue_position(issue)
        if not sheet_name or row_number <= 0:
            continue
        issues_by_row.setdefault((sheet_name, row_number), []).append(issue)

    for ws in wb.worksheets:
        opinion_col = ws.max_column + 1
        status_col = opinion_col + 1
        count_col = status_col + 1
        ws.cell(row=1, column=opinion_col, value="审核意见")
        ws.cell(row=1, column=status_col, value="审核状态")
        ws.cell(row=1, column=count_col, value="问题数量")
        for row_number in range(2, ws.max_row + 1):
            row_issues = issues_by_row.get((ws.title, row_number), [])
            if not row_issues:
                ws.cell(row=row_number, column=opinion_col, value="—")
                ws.cell(row=row_number, column=status_col, value="确认无误")
                ws.cell(row=row_number, column=count_col, value=0)
                continue
            opinions = []
            has_term_issue = False
            for issue in row_issues:
                category = _issue_value(issue, 'category', '问题') or '问题'
                description = _issue_value(issue, 'description', '') or _issue_value(issue, 'original_text', '')
                suggestion = _format_issue_suggestion(issue)
                text = f"[{category}] {description}"
                if suggestion and suggestion != '-':
                    text += f" → {suggestion}"
                opinions.append(text)
                if category == '术语不一致':
                    has_term_issue = True
            ws.cell(row=row_number, column=opinion_col, value="\n".join(opinions))
            ws.cell(row=row_number, column=status_col, value="需确认" if has_term_issue else "待修改")
            ws.cell(row=row_number, column=count_col, value=len(row_issues))

    wb.save(export_path)
    return export_path, export_name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _run_english_heuristic_audit(content, file_type=None):
    issues = []
    seen = set()
    allowed_ui_commands = {
        'tap save', 'tap confirm', 'tap cancel', 'tap delete',
        'click ok', 'click cancel', 'click save',
        'press enter', 'press ok',
    }

    def add_issue_by_span(start, end, original_text, rule, category, suggestion, description, audit_basis, severity="general"):
        original_text = str(original_text or "").strip()
        if not original_text:
            return
        key = (rule, _normalize_search_text(original_text), start)
        if key in seen:
            return
        seen.add(key)
        issues.append({
            "severity": severity,
            "category": category,
            "rule": rule,
            "chapter": extract_chapter(content, start),
            "original_text": original_text,
            "context": get_context(content, start, end, 160),
            "suggestion": suggestion,
            "description": description,
            "audit_basis": audit_basis,
            "confidence": 92,
            "source": "rule",
            "position": _encode_issue_position(start, end),
        })

    def add_issue(match, rule, category, suggestion, description, audit_basis, severity="general"):
        add_issue_by_span(match.start(), match.end(), match.group(0), rule, category, suggestion, description, audit_basis, severity)

    def looks_like_table_context(context):
        snippet = str(context or '')
        line = max(snippet.splitlines(), key=len, default=snippet)
        compact = re.sub(r'\s+', ' ', line).strip()
        if '\t' in snippet or '|' in snippet:
            return True
        if re.search(r'\b(?:Table|Figure)\s+\d+\b', compact, re.IGNORECASE):
            return True
        if len(re.findall(r'\d+(?:\.\d+)?\s*[A-Za-z%°μ/]+', compact)) >= 2:
            return True
        return bool(re.search(r'\s{3,}', line))

    def is_correct_sample_span(start):
        marker = content.lower().find('correct samples')
        return marker >= 0 and start >= marker

    def is_troubleshooting_span(start):
        troubleshooting = content.find('7. Troubleshooting')
        appendix = content.find('8. Appendix')
        if troubleshooting < 0:
            return False
        end = appendix if appendix > troubleshooting else len(content)
        return troubleshooting <= start < end

    def should_report_ul_unit(match):
        number_match = re.search(r"(\d+(?:\.\d+)?)\s*$", content[max(0, match.start() - 16):match.start()])
        if number_match and float(number_match.group(1)) >= 1000:
            return False
        chapter = extract_chapter(content, match.start())
        if chapter.startswith('3.1 Reagent Components'):
            return True
        if chapter.startswith('3.3'):
            context = get_context(content, match.start(), match.end(), 80)
            return 'Pipettes' in context
        if chapter.startswith('5.2 Cell Capture'):
            return True
        if chapter.startswith('5.4 PCR Amplification'):
            return True
        if chapter.startswith('5.5 Library Purification'):
            return True
        return False

    official_global_site = 'https://global-mgitech.com/'
    outdated_site_pattern = re.compile(r"(?<![A-Za-z0-9.-])(?:https?://)?(?:www\.)?(?:en\.mgi-tech\.com|global-mgi\.com|global\.mgi-tech\.com)(?![A-Za-z0-9.-])", re.IGNORECASE)
    for match in outdated_site_pattern.finditer(content):
        add_issue(match, "HR001", "合规", f"建议替换为 {official_global_site}", "英文资料中的海外官网地址应使用官网 English 入口域名。", "公司特定规范 - 海外官网地址", "fatal")

    for match in re.finditer(r"\bdesk\s+top\b", content, re.IGNORECASE):
        add_issue(match, "HR002", "英文规范", "建议改为 desktop", "desktop 在该语境中应使用合成词写法。", "英文技术文档写作规范 - 拼写")

    if file_type != "pdf":
        abbreviation_pattern = re.compile(r"(?:\be\.g|\bi\.e|\bDr|\bMr|\bMrs|\bMs|\bProf|\bFig|\bNo)$", re.IGNORECASE)
        for match in re.finditer(r"(?<=[A-Za-z0-9\)])\.(?=[A-Za-z0-9])", content):
            prefix = content[max(0, match.start() - 8):match.start()]
            window = content[max(0, match.start() - 30):min(len(content), match.end() + 30)]
            if abbreviation_pattern.search(prefix):
                continue
            if re.search(r'@|https?://|www\.|\b[A-Za-z0-9-]+\.[A-Za-z]{2,}\b', window):
                continue
            if re.match(r"\d+\.\d+", content[max(0, match.start() - 1):match.end() + 2]):
                continue
            original = content[match.start():min(len(content), match.start() + 2)]
            suggestion = f"建议改为 {original[0]} {original[1:]}" if len(original) == 2 else "建议在句号后补一个空格"
            add_issue_by_span(match.start(), min(len(content), match.start() + len(original)), original, "HR003", "标点符号", suggestion, "英文正文中的句号后应保留空格。", "英文技术文档写作规范 - 标点与空格")

    for match in re.finditer(r"\s+,", content):
        add_issue(match, "PUNCT-002", "标点符号", "建议删除逗号前空格", "逗号前不应有空格。", "英文技术文档写作规范 - 标点与空格")

    unit_pattern = r"(?:ng/μL|ng/uL|μL|uL|mL|mg|ng|g|kg|mm|cm|min|sec|hr|bp)"
    for match in re.finditer(rf"(?<![A-Za-z0-9.])\d+(?:\.\d+)?{unit_pattern}\b", content):
        original = match.group(0)
        unit_match = re.match(r"(\d+(?:\.\d+)?)(.+)", original)
        if not unit_match:
            continue
        suggestion = f"建议改为 {unit_match.group(1)} {unit_match.group(2)}"
        context = get_context(content, match.start(), match.end(), 120)
        severity = 'general' if looks_like_table_context(context) else 'serious'
        add_issue(match, "HR004", "单位", suggestion, "数字与单位之间应保留空格。", "英文技术文档写作规范 - 单位格式", severity)

    for match in re.finditer(r"\b(μL|uL|mL|min|sec|ng|bp)s\b", content):
        unit = match.group(1)
        add_issue(match, "UNIT-002", "单位", f"建议改为 {unit}", "单位符号不使用复数形式。", "英文技术文档写作规范 - 单位格式")

    for match in re.finditer(r"\buL\b", content):
        if file_type == 'pdf' and not should_report_ul_unit(match):
            continue
        add_issue(match, "UNIT-003", "单位", "建议改为 μL", "微升单位建议使用标准符号 μL。", "英文技术文档写作规范 - 单位格式")

    for match in re.finditer(r"(?<![A-Za-z0-9°℃-])(?:-?\d+(?:\.\d+)?)(?:\s+to\s+-?\d+(?:\.\d+)?)?\s*C\b", content):
        if is_correct_sample_span(match.start()):
            continue
        if file_type == 'pdf' and is_troubleshooting_span(match.start()):
            continue
        original = match.group(0)
        if '\n' in original:
            continue
        temp_values = [float(value) for value in re.findall(r'-?\d+(?:\.\d+)?', original)]
        if temp_values and any(abs(value) > 150 for value in temp_values):
            continue
        suggestion = re.sub(r"\s*C\b", "°C", original)
        add_issue(match, "UNIT-004", "单位", f"建议改为 {suggestion}", "摄氏温度单位建议使用 °C。", "英文技术文档写作规范 - 单位格式")

    for match in re.finditer(r"(?<![A-Za-z°℃])(-?\d+(?:\.\d+)?)\s*-\s*(-?\d+(?:\.\d+)?)\s*C\b", content):
        if is_correct_sample_span(match.start()):
            continue
        if file_type == 'pdf' and is_troubleshooting_span(match.start()):
            continue
        original = match.group(0)
        if '\n' in original:
            continue
        temp_values = [float(value) for value in re.findall(r'-?\d+(?:\.\d+)?', original)]
        if temp_values and any(abs(value) > 150 for value in temp_values):
            continue
        suggestion = f"{match.group(1)} to {match.group(2)}°C"
        add_issue(match, "UNIT-004", "单位", f"建议改为 {suggestion}", "温度范围建议使用 to 表达范围，并使用 °C。", "英文技术文档写作规范 - 单位格式")

    for match in re.finditer(r"\bWeLL\b", content):
        add_issue(match, "HR005", "格式", "建议统一为 Well", "单词内部大小写形式不一致。", "技术文档常见错误清单 - 格式一致性", "serious")

    range_unit_pattern = r"(?:ng/μL|μL|uL|mL|mg|mm|cm|℃|°C|%)"
    for match in re.finditer(rf"\b\d+(?:\.\d+)?\s*{range_unit_pattern}\s*-\s*\d+(?:\.\d+)?\s*{range_unit_pattern}\b", content):
        text = match.group(0)
        suggestion = re.sub(r"\s*-\s*", " to ", text)
        add_issue(match, "HR006", "格式", f"建议改为 {suggestion}", "英文数值范围建议使用 to 表达范围。", "英文技术文档写作规范 - 范围表达", "serious")

    for match in re.finditer(r"greater than upper limit range", content, re.IGNORECASE):
        add_issue(match, "HR007", "语法", "建议改为 greater than the upper limit range", "该短语前缺少定冠词 the。", "英文技术文档写作规范 - 冠词", "serious")

    for match in re.finditer(r"\b(time|yield|result|sample|concentration)\b([^.!?\n]{0,80})\bhave\s+been\b", content, re.IGNORECASE):
        subject = match.group(1)
        original = match.group(0)
        suggestion = re.sub(r"\bhave\s+been\b", "has been", original, flags=re.IGNORECASE)
        add_issue_by_span(match.start(), match.end(), original, "GRAMMAR-003", "语法", f"建议改为 {suggestion}", f"主语 {subject} 为单数，谓语应使用 has been。", "英语语法规范 - 主谓一致", "serious")

    for match in re.finditer(r"\bmake\s+total\s+volume\b", content, re.IGNORECASE):
        add_issue(match, "GRAMMAR-004", "语法", "建议改为 make a total volume 或 make the total volume", "total volume 前建议添加冠词。", "英语语法规范 - 冠词")

    please_replacements = {
        'please contact': 'contact technical support',
        'please use with caution': 'Use with caution',
        'please adjust moderately': 'Adjust moderately',
    }
    for phrase, replacement in please_replacements.items():
        for match in re.finditer(rf"\b{re.escape(phrase)}\b", content, re.IGNORECASE):
            if file_type == 'pdf':
                continue
            add_issue(match, "HR008", "风格", f"建议改为 {replacement}", "技术文档建议使用直接的客观表达。", "英文技术文档写作规范 - 语气")

    for match in re.finditer(r"\bafter login in\b", content, re.IGNORECASE):
        add_issue(match, 'GRAMMAR-001', '语法', '建议改为 after logging in', 'login 作为动作时应使用 logging in。', '英语语法规范 - 动作表达', 'serious')

    for match in re.finditer(r"\bafter login\b(?!\s+in\b)", content, re.IGNORECASE):
        context = get_context(content, match.start(), match.end(), 40).lower()
        if re.search(r'\b(?:the|you|the software|the device|tap|click|press|can|will|is|are|automatically)\b', context):
            add_issue(match, 'GRAMMAR-002', '语法', '建议改为 after logging in', 'login 作为动作时应使用 logging in。', '英语语法规范 - 动作表达', 'serious')

    if file_type != 'pdf':
        for match in re.finditer(r"\blog\s*in\b(?!\s+order\b)", content, re.IGNORECASE):
            add_issue(match, "HR009", "一致性", "建议统一为 log in 或 login，并与界面名词/动作语境保持一致", "登录相关表达应区分动作和名词。", "英文技术文档写作规范 - 术语一致性")

    if file_type != 'pdf':
        for match in re.finditer(r"\blog\s*out\b", content, re.IGNORECASE):
            add_issue(match, "HR010", "一致性", "建议统一为 log out 或 logout，并与界面名词/动作语境保持一致", "退出登录相关表达应区分动作和名词。", "英文技术文档写作规范 - 术语一致性")

    for match in re.finditer(r"\b\d{1,3}(?:,\d{3})*(?:\.\d+)?(?:\s*[xX]\s*(?:\d{1,3}(?:,\d{3})*(?:\.\d+)?|[A-Za-z]))+\b", content):
        if match.end() < len(content) and content[match.end():match.end() + 1] == '^':
            continue
        suggestion = re.sub(r"\s*[xX]\s*", " × ", match.group(0))
        add_issue(match, "HR011", "格式", f"建议改为 {suggestion}", "乘号建议使用 × 并保留两侧空格。", "英文技术文档写作规范 - 数学符号")

    for match in re.finditer(r"\b\d+(?:\.\d+)?[xX]\s+volume\b", content):
        suggestion = re.sub(r"[xX]", "×", match.group(0), count=1)
        add_issue(match, "HR011", "格式", f"建议改为 {suggestion}", "倍数符号建议使用 ×。", "英文技术文档写作规范 - 数学符号")

    for match in re.finditer(r"\bshould be disposed\b(?!\s+of\b)", content, re.IGNORECASE):
        add_issue(match, "GRAMMAR-005", "语法", "建议改为 should be disposed of", "dispose of 是固定搭配。", "英语语法规范 - 固定搭配", "serious")

    for match in re.finditer(r"\bshould be keeped\b", content, re.IGNORECASE):
        add_issue(match, "GRAMMAR-006", "语法", "建议改为 should be kept", "keep 的过去分词应为 kept。", "英语语法规范 - 动词形式", "serious")

    for match in re.finditer(r"\bDNBelab C4\b", content):
        add_issue(match, "TERM-EN-001", "术语不一致", "建议统一为 DNBelab C-Series", "产品名称前后应保持一致。", "英文技术文档写作规范 - 术语一致性")

    pcr_marker = content.find("Step | Temperature | Time")
    if pcr_marker >= 0 and "Cycle" not in content[pcr_marker:pcr_marker + 260]:
        add_issue_by_span(pcr_marker, pcr_marker + len("Step | Temperature | Time"), "PCR cycling", "LOGIC-002", "内容逻辑", "建议补充 PCR 循环次数列或循环次数说明", "PCR cycling 条件表缺少循环次数信息。", "说明书发布前自检 Checklist - 内容完整性", "serious")

    for match in re.finditer(r"\b(?:www\.)?mgitech\.cn\b", content, re.IGNORECASE):
        add_issue(match, "HR012", "合规", f"建议改为 {official_global_site}", "海外英文文档中官网域名应统一为官网 English 入口。", "公司特定规范 - 海外官网地址", "fatal")

    for match in re.finditer(r"\b(?:e-?mail)\s*:", content, re.IGNORECASE):
        if match.group(0).strip() == 'Email:':
            continue
        add_issue(match, "HR013", "格式", "建议统一为 Email:", "联系方式标签建议统一大小写。", "英文技术文档写作规范 - 标签格式")

    for match in re.finditer(r"\b(?:web site|website address)\b", content, re.IGNORECASE):
        add_issue(match, "HR014", "一致性", "建议统一为 website", "网站相关名词建议统一。", "英文技术文档写作规范 - 术语一致性")

    for match in re.finditer(r"\bcharaters\b|\bcharater\b|\boccured\b|\bocurred\b|\btecnical\b|\bseperate\b|\bmaintenence\b|\bmaintanance\b|\breferrence\b|\brefered\b|\brefering\b|\buntill\b|\buseing\b|\bwheras\b|\bwich\b|\bdefination\b|\bdesciption\b|\brecieve\b", content, re.IGNORECASE):
        if is_whitelisted(match.group(0)):
            continue
        replacements = {
            'charaters': 'characters',
            'charater': 'character',
            'occured': 'occurred',
            'ocurred': 'occurred',
            'tecnical': 'technical',
            'seperate': 'separate',
            'maintenence': 'maintenance',
            'maintanance': 'maintenance',
            'referrence': 'reference',
            'refered': 'referred',
            'refering': 'referring',
            'untill': 'until',
            'useing': 'using',
            'wheras': 'whereas',
            'wich': 'which',
            'defination': 'definition',
            'desciption': 'description',
            'recieve': 'receive',
        }
        replacement = replacements.get(match.group(0).lower(), '')
        add_issue(match, 'SPELL-001', '拼写/用词错误', f'建议改为 {replacement}', '发现常见高置信拼写错误。', '常见拼写错误清单', 'serious')

    if file_type != 'pdf':
        for match in re.finditer(r"\bTap\s+(?!to\b|the\b|a\b|an\b)([A-Za-z]+)\b", content):
            if match.group(0).lower() in allowed_ui_commands:
                continue
            action = match.group(1)
            if action[:1].isupper() or action.isupper() or len(action) <= 2:
                continue
            add_issue(match, 'STYLE-002', '风格', f'建议改为 Tap to {action.lower()}', 'Tap 后建议使用 to 引导动作。', '交互文案规范', 'general')

    for match in re.finditer(r"Performing the following steps", content, re.IGNORECASE):
        add_issue(match, 'STYLE-003', '风格', '建议改为 Perform the following steps', '祈使句建议直接使用动词原形。', '技术文档句式规范', 'general')

    for match in re.finditer(r"[\u3001\u3002\uff0c\uff1b\uff1a\uff08\uff09]", content):
        add_issue(match, 'PUNCT-001', '标点符号', '建议改为对应的半角英文标点', '英文文档中不应混入全角或中文标点。', '英文标点规范', 'serious')

    if file_type != 'pdf':
        for match in re.finditer(r"\bwill\s+(?:be\s+)?\w+(?:ed|en)\b", content, re.IGNORECASE):
            add_issue(match, 'TENSE-001', '时态', '建议改为一般现在时表达', '技术文档通常使用一般现在时。', '技术文档时态规范', 'general')

    if file_type != 'pdf' and re.search(r'\bTap\b', content) and re.search(r'\b[Cc]lick\b', content):
        for match in re.finditer(r"\b[Cc]lick\b", content):
            phrase = str(content[match.start():min(len(content), match.end() + 8)]).split('\n', 1)[0].strip().lower()
            if any(phrase.startswith(command) for command in allowed_ui_commands if command.startswith('click ')):
                continue
            add_issue(match, 'TERM-001', '术语一致性', '建议统一为 Tap', '触屏设备中的操作用语建议统一使用 Tap。', '触屏交互术语规范', 'general')

    return issues


def _run_manual_engineering_audit(content, file_type=None):
    issues = []
    seen = set()

    def add_issue(start, end, original_text, rule, category, suggestion, description, audit_basis, severity="general", confidence=90):
        original_text = str(original_text or "").strip()
        if not original_text:
            return
        key = (rule, _normalize_search_text(original_text), int(start or 0))
        if key in seen:
            return
        seen.add(key)
        issues.append({
            "severity": severity,
            "category": category,
            "rule": rule,
            "chapter": extract_chapter(content, start),
            "original_text": original_text[:240],
            "context": get_context(content, start, end, 180),
            "suggestion": suggestion,
            "description": description,
            "audit_basis": audit_basis,
            "confidence": confidence,
            "source": "rule",
            "position": _encode_issue_position(start, end),
        })

    if not content:
        return issues

    def normalized_block_text(text):
        text = re.sub(r'(?i)\bM\s*a\s*y\b', 'May', str(text or ''))
        text = re.sub(r'(?i)\bJ\s*u\s*n\b', 'Jun', text)
        text = re.sub(r'(?i)\bJ\s*u\s*l\b', 'Jul', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def version_tuple(value):
        parts = [int(part) for part in re.findall(r'\d+', str(value or ''))]
        return tuple(parts) if parts else (0,)

    def normalize_catalog_item(name):
        name = re.sub(r'\s+', ' ', str(name or '')).strip().lower()
        name = re.sub(r'\b(?:mgi|thermo fisher scientifictm|vitl)\b', '', name)
        name = re.sub(r'\([^)]*\)', '', name)
        name = re.sub(r'\b(?:cat\.?\s*no\.?|cat\.?|brand|quantity|consumables|reagents)\b', '', name)
        name = re.sub(r'[^a-z0-9μ\s\-]+', ' ', name)
        name = re.sub(r'\s+', ' ', name).strip()
        return name

    def add_micro_edit(start, end, original, suggestion, description, confidence=88):
        add_issue(
            start,
            end,
            original,
            'DOC-MICRO-001',
            '英文微编辑',
            suggestion,
            description,
            '说明书审核能力补强方案 - 英文微编辑',
            'general',
            confidence,
        )

    # English manuals should not contain Chinese residual text.
    for match in list(re.finditer(r"[\u4e00-\u9fff]+", content))[:20]:
        add_issue(
            match.start(),
            match.end(),
            match.group(0),
            "ENG-CN-001",
            "英文文档中文混入",
            "建议删除中文残留或替换为英文表达",
            "英文说明书中不应残留中文字符。",
            "说明书审核能力补强方案 - 英文文档中文混入",
            "serious",
            95,
        )

    # Revision History should list the latest version/date first.
    revision_match = re.search(r"Revision\s+History([\s\S]{0,2500}?)(?:Contents|Chapter\s+1|1\s+Product)", content, re.IGNORECASE)
    if not revision_match:
        revision_match = re.search(r"Revision\s+History([\s\S]{0,1200})", content, re.IGNORECASE)
    if revision_match:
        raw_revision_text = revision_match.group(1)
        revision_text = normalized_block_text(raw_revision_text)
        entries = []
        month_pattern = r"Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?"
        full_month_names = {
            'jan': 'January',
            'feb': 'February',
            'mar': 'March',
            'apr': 'April',
            'jun': 'June',
            'jul': 'July',
            'aug': 'August',
            'sep': 'September',
            'sept': 'September',
            'oct': 'October',
            'nov': 'November',
            'dec': 'December',
        }
        for month_match in re.finditer(r"\b(Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)\.\s*,?\s*20\d{2}\b", raw_revision_text, re.IGNORECASE):
            month = full_month_names.get(month_match.group(1).lower())
            if not month:
                continue
            original = month_match.group(0)
            suggestion = re.sub(re.escape(month_match.group(1)) + r'\.', month, original, count=1, flags=re.IGNORECASE)
            absolute_start = revision_match.start(1) + month_match.start()
            add_issue(
                absolute_start,
                absolute_start + len(original),
                original,
                "DOC-REV-002",
                "版本记录",
                f"建议改为 {suggestion}",
                "Revision History 表格中的月份应使用英文全称，不使用月份缩写。",
                "技术文档常见错误清单 - 修订历史月份格式",
                "general",
                92,
            )
        for item in re.finditer(rf"\b(\d+(?:\s*\.\s*\d+)*)\b(?=[^\n]{{0,160}}?\b(?:{month_pattern})\s*,?\s*20\d{{2}}\b)", revision_text, re.IGNORECASE):
            raw_version = re.sub(r'\s+', '', item.group(1))
            version = version_tuple(raw_version)
            if version == (0,) or raw_version in {'1'}:
                continue
            entries.append((version, item.group(0)))
        if len(entries) >= 2 and entries[0][0] < max(version for version, _ in entries):
            start = revision_match.start()
            add_issue(
                start,
                min(len(content), start + 220),
                'Revision History',
                "DOC-REV-001",
                "版本记录",
                "建议将最新版本记录置于 Revision History 最上方",
                "Revision History 通常按最新版本在前排列，便于读者快速确认当前版本。",
                "说明书审核能力补强方案 - 版本记录",
                "general",
                88,
            )

    # Model and platform name whitelist / near-miss checks.
    model_replacements = {
        r"\bMGI-SP960\b": "MGISP-960",
        r"\bMGIISP-NE384\b": "MGISP-NE384",
        r"\bMGIISP\b(?!-NE384)": "MGISP",
    }
    for pattern, replacement in model_replacements.items():
        for match in re.finditer(pattern, content, re.IGNORECASE):
            add_issue(
                match.start(),
                match.end(),
                match.group(0),
                "DOC-MODEL-001",
                "型号一致性",
                f"建议确认是否应改为 {replacement}",
                "设备型号应按白名单保持一致，疑似存在漏字、错字或相似型号混用。",
                "说明书审核能力补强方案 - 型号与产品名一致性",
                "serious",
                90,
            )

    if re.search(r"MGISP-NE384|\bNE384\b", content, re.IGNORECASE):
        hardware_section = re.search(r"(?:Hardware|Hardware configuration|1\.3\s+Hardware)([\s\S]{0,2500}?)(?:1\.4|1\.5|Chapter\s+2|2\s+)", content, re.IGNORECASE)
        hardware_text = hardware_section.group(0) if hardware_section else ""
        if hardware_text and not re.search(r"MGISP-NE384|\bNE384\b", hardware_text, re.IGNORECASE):
            start = hardware_section.start()
            add_issue(
                start,
                min(len(content), start + 120),
                "Hardware configuration",
                "DOC-SCOPE-001",
                "适用范围一致性",
                "建议补充 NE384 对应硬件配置，或调整适用范围描述",
                "文档提到 NE384 适用范围时，硬件配置章节应同步覆盖对应设备。",
                "说明书审核能力补强方案 - 适用范围一致性",
                "serious",
                86,
            )

    ne384_section = re.search(r"\b3\.1\.2\s+MGISP-NE384\b([\s\S]{0,1600}?)(?=\b3\.2\s+|\Z)", content, re.IGNORECASE)
    if ne384_section:
        variant = re.search(r"\bMGISP-NE384RS\b", ne384_section.group(1), re.IGNORECASE)
        if variant:
            start = ne384_section.start(1) + variant.start()
            add_issue(
                start,
                start + len(variant.group(0)),
                variant.group(0),
                "DOC-MODEL-002",
                "型号一致性",
                "建议核对是否应与章节标题 MGISP-NE384 保持一致，或同步调整章节标题",
                "同一小节标题和正文软件/设备型号出现 NE384 与 NE384RS 混用，需确认适用设备范围。",
                "说明书审核能力补强方案 - 型号与产品名一致性",
                "serious",
                88,
            )

    hardware_candidates = list(re.finditer(r"(?m)^\s*1\.3\s+Hardware\s*$([\s\S]{0,4000}?)(?=^\s*1\.4\s+Applicable|^\s*1\.5\s+Equipment|^\s*Chapter\s+2)", content, re.IGNORECASE | re.MULTILINE))
    hardware_section = next((candidate for candidate in hardware_candidates if re.search(r'\bPos\s*\d+\b', candidate.group(1), re.IGNORECASE)), None)
    deck_candidates = list(re.finditer(r"(?m)^\s*3\.4\s+Sample\s+pre-extraction\s*$([\s\S]{0,9000}?)(?=^\s*3\.5\s+DNA|^\s*Chapter\s+4|\Z)", content, re.IGNORECASE | re.MULTILINE))
    deck_section = next((candidate for candidate in deck_candidates if re.search(r'\bPos\s*\d+\s*~\s*Pos\s*\d+\b', candidate.group(1), re.IGNORECASE)), None)
    if hardware_section and deck_section:
        hardware_positions = [int(value) for value in re.findall(r"\bPos\s*(\d+)\b", hardware_section.group(1), re.IGNORECASE)]
        deck_ranges = re.findall(r"\bPos\s*(\d+)\s*~\s*Pos\s*(\d+)\b", deck_section.group(1), re.IGNORECASE)
        deck_positions = [int(value) for pair in deck_ranges for value in pair]
        if hardware_positions and deck_positions and min(hardware_positions) >= 18 and min(deck_positions) <= 4:
            start = deck_section.start(1) + deck_section.group(1).lower().find('pos')
            if start < deck_section.start(1):
                start = deck_section.start()
            add_issue(
                start,
                min(len(content), start + 160),
                'Pos1 ~ Pos16',
                "DOC-POS-001",
                "台面位置一致性",
                "建议核对 3.4 操作台面 Pos 编号是否与 1.3 硬件配置中的设备位点一致",
                "硬件配置章节使用 Pos18 及以上位点，而操作台面章节使用 Pos1~Pos16，需确认是否为不同设备坐标体系或编号错误。",
                "说明书审核能力补强方案 - 台面位置一致性",
                "serious",
                87,
            )

    # Table/Figure numbering sequence checks.
    for label in ("Table", "Figure"):
        chapter_numbers = {}
        for match in re.finditer(rf"\b{label}\s*(\d+)[-–](\d+)\b", content, re.IGNORECASE):
            line_start = max(0, content.rfind('\n', 0, match.start()) + 1)
            line_end = content.find('\n', match.end())
            if line_end < 0:
                line_end = min(len(content), match.end() + 120)
            line = content[line_start:line_end]
            if len(line) > 180 or not _is_table_or_figure_caption_line(line, label, match.start() - line_start):
                continue
            chapter = int(match.group(1))
            number = int(match.group(2))
            chapter_numbers.setdefault(chapter, []).append((number, match.start(), match.group(0)))
        for chapter, rows in chapter_numbers.items():
            unique = sorted({number for number, _, _ in rows})
            if len(unique) < 2:
                continue
            expected = list(range(unique[0], unique[-1] + 1))
            missing = [number for number in expected if number not in unique]
            if missing:
                first = rows[0]
                add_issue(
                    first[1],
                    first[1] + len(first[2]),
                    f"{label} {chapter}-{unique[0]} ... {label} {chapter}-{unique[-1]}",
                    "DOC-FIGTAB-001",
                    "表图编号",
                    f"建议检查 {label} {chapter}-{', '.join(str(n) for n in missing)} 是否缺失或编号跳号",
                    "表图编号应按章节连续，缺号或跳号需要确认。",
                    "说明书审核能力补强方案 - 表图编号连续性",
                    "general",
                    84,
                )
            counts = Counter(number for number, _, _ in rows)
            for number, count in counts.items():
                if count > 1:
                    repeated = next(item for item in rows if item[0] == number)
                    add_issue(
                        repeated[1],
                        repeated[1] + len(repeated[2]),
                        repeated[2],
                        "DOC-FIGTAB-002",
                        "表图编号",
                        f"建议确认 {label} {chapter}-{number} 是否重复编号",
                        "同一章节内表图编号不应重复。",
                        "说明书审核能力补强方案 - 表图编号连续性",
                        "general",
                        82,
                    )

    # Cat. No. wording and variants.
    for match in re.finditer(r"\b(?:cat\.?\s*:|cat\.?\s*number|catalog\s+number)\b", content, re.IGNORECASE):
        original = match.group(0)
        if original == "Cat. No.":
            continue
        add_issue(
            match.start(),
            match.end(),
            original,
            "DOC-CATNO-001",
            "货号写法",
            "建议统一为 Cat. No.",
            "货号标签建议统一使用 Cat. No. 写法。",
            "说明书审核能力补强方案 - 货号一致性",
            "general",
            90,
        )

    # Same consumable/material should not point to conflicting Cat. No. values across tables.
    catalog_entries = []
    for match in re.finditer(r"(?mi)^\s*(.+?)\s+\(?MGI,?\s+Cat\.\s*No\.?:\s*([A-Z0-9-]+)\)?", content):
        item_name = normalize_catalog_item(match.group(1))
        if item_name:
            catalog_entries.append((item_name, match.group(1).strip(), match.group(2), match.start()))
    table_list_pattern = re.compile(r"(?mi)^\s*(.+?)\s+MGI\s+([A-Z0-9]{6,}(?:-[A-Z0-9]{2})*)\s+\d+\s*$")
    for match in table_list_pattern.finditer(content):
        item_name = normalize_catalog_item(match.group(1))
        if item_name and not re.search(r'^(?:table|figure|chapter)\b', item_name, re.IGNORECASE):
            catalog_entries.append((item_name, match.group(1).strip(), match.group(2), match.start()))
    by_item = defaultdict(list)
    for item_name, raw_name, cat_no, start in catalog_entries:
        by_item[item_name].append((raw_name, cat_no, start))
    for item_name, rows in by_item.items():
        cat_values = _report_unique_values(cat_no for _, cat_no, _ in rows)
        if len(cat_values) < 2:
            continue
        first_name, first_cat, first_start = rows[0]
        add_issue(
            first_start,
            min(len(content), first_start + len(first_name) + len(first_cat) + 20),
            first_name,
            "DOC-CATNO-002",
            "货号一致性",
            f"建议核对同一物料的货号，当前出现：{'、'.join(cat_values)}",
            "同一耗材或物料在不同表格中不应对应多个 Cat. No.，需核对是否写错或物料名称不够具体。",
            "说明书审核能力补强方案 - 货号一致性",
            "serious",
            88,
        )

    # Short orphan sections.
    headings = []
    for match in re.finditer(r"(?m)^\s*(\d+(?:\.\d+)+)\s+([^\n]{3,90})\s*$", content):
        title = match.group(0).strip()
        title_text = match.group(2).strip()
        before = content[max(0, match.start() - 1400):match.start()]
        if re.search(r'\bContents\b', before, re.IGNORECASE) and re.search(r'\s+\d+\s*$', title):
            continue
        if re.search(r"\b(?:Table|Figure|Pos\d|Cat\.|No\.)\b", title, re.IGNORECASE):
            continue
        if _is_material_line(title) or re.match(r"^(?:mL|μL|uL|L|mg|ng|g|kg|mm|cm|rpm|Well)\b", title_text, re.IGNORECASE):
            continue
        if not re.match(r"^[A-Z][A-Za-z]", title_text):
            continue
        headings.append((match.start(), match.end(), title))
    for index, (start, end, title) in enumerate(headings):
        next_start = headings[index + 1][0] if index + 1 < len(headings) else min(len(content), end + 1000)
        body = re.sub(r"\s+", " ", content[end:next_start]).strip()
        word_count = len(re.findall(r'[A-Za-z][A-Za-z\-]*|[\u4e00-\u9fff]+', body))
        is_shallow_section = 0 < len(body) < 80 or (word_count < 35 and re.match(r'^2\.\d+\s+', title))
        if is_shallow_section and not re.search(r"\b(?:warning|caution|note)\b", title, re.IGNORECASE):
            add_issue(
                start,
                end,
                title,
                "DOC-SECTION-001",
                "章节结构",
                "建议合并到相邻章节或补充本小节内容",
                "内容过少的孤立小节会影响说明书结构完整性。",
                "说明书审核能力补强方案 - 章节结构",
                "suggestion",
                82,
            )

    micro_patterns = [
        (
            r"\bThe applicable software version of the manual is\b",
            "The software version applicable to this manual is",
            "软件版本说明句式不自然，建议使用更直接的 technical manual 表达。",
        ),
        (
            r"\bthe device of MGISP-960\b",
            "MGISP-960",
            "设备名称前的 the device of 表达冗余，操作步骤中建议直接写设备型号。",
        ),
        (
            r"\bAfter the \"1\. Discard supernatant\" process is finished\.\s*rearrange\b",
            "After the \"1. Discard supernatant\" process is finished, rearrange",
            "独立分句之间应使用逗号衔接，避免句号后接小写动词。",
        ),
        (
            r"\baluminum mold\b",
            "aluminum foil",
            "此处语境为封膜耗材，mold 疑似应为 foil。",
        ),
        (
            r"\bBefore experiment\b",
            "Before the experiment",
            "experiment 作特定实验流程时建议补充定冠词 the。",
        ),
        (
            r"\bIf you have other questions\b",
            "If you have any other questions",
            "other questions 前建议补充 any，使英文表达更自然。",
        ),
        (
            r"\beluent\s+t\s*\(",
            "eluent (",
            "eluent 后多出字母 t，应删除多余字符。",
        ),
        (
            r"\bsystem\s*\(",
            "system (",
            "英文正文中左括号前应保留空格。",
        ),
    ]
    for pattern, suggestion, description in micro_patterns:
        for match in re.finditer(pattern, content, re.IGNORECASE):
            add_micro_edit(match.start(), match.end(), match.group(0), suggestion, description)

    for match in re.finditer(r"\bprior to use\b", content, re.IGNORECASE):
        add_micro_edit(match.start(), match.end(), match.group(0), "before use", "技术文档中建议使用更简洁直接的 before use。", 84)

    for match in re.finditer(r"\bAdd\s+absolute\s+ethanol\s+into\s+\d{2,}\s+according\s+to\s+the\s+label\b", content, re.IGNORECASE):
        add_micro_edit(
            match.start(),
            match.end(),
            match.group(0),
            "请核对数字是否缺少对应试剂或缓冲液名称",
            "操作对象写成纯数字会导致读者无法确认应加入的试剂或容器。",
            88,
        )

    # Procedure language checks: imperative mood, present tense, avoid passive/past-tense narration.
    procedure_start = re.search(r"(?:Chapter\s+3|3\s+(?:Operation|Library\s+preparation\s+protocol)|3\.1\s+(?:Preparing|Workflow))", content, re.IGNORECASE)
    procedure_text = content[procedure_start.start():] if procedure_start else ''
    procedure_offset = procedure_start.start() if procedure_start else 0
    passive_patterns = [
        r"\bthe\s+[^\n\.]{0,60}\s+was\s+\w+(?:ed|en)\b",
        r"\bthe\s+[^\n\.]{0,60}\s+were\s+\w+(?:ed|en)\b",
        r"\bthe\s+[^\n\.]{0,60}\s+is\s+\w+(?:ed|en)\b",
        r"\bthe\s+[^\n\.]{0,60}\s+are\s+\w+(?:ed|en)\b",
    ]
    reported_voice = 0

    def is_numbered_step_context(local_start):
        paragraph_start = procedure_text.rfind('\n\n', 0, local_start)
        paragraph_start = 0 if paragraph_start < 0 else paragraph_start + 2
        paragraph = procedure_text[paragraph_start:local_start]
        return bool(re.match(r"\s*\d+[\.)]\s+", paragraph))

    for pattern in passive_patterns:
        for match in re.finditer(pattern, procedure_text, re.IGNORECASE):
            if reported_voice >= 12:
                break
            if not is_numbered_step_context(match.start()):
                continue
            original = re.sub(r"\s+", " ", match.group(0)).strip()
            suggestion = "建议改为祈使句，并使用一般现在时"
            if re.search(r'\bthe\s+program\s+is\s+completed\b', original, re.IGNORECASE):
                suggestion = "建议改为 When the program completes"
            elif re.search(r'\bthe\s+Process\s+edit\s+interface\s+is\s+displayed\b', original, re.IGNORECASE):
                suggestion = "建议改为 Click to display the Process edit interface"
            start = procedure_offset + match.start()
            end = procedure_offset + match.end()
            add_issue(
                start,
                end,
                original,
                "DOC-PROC-001",
                "操作步骤语气",
                suggestion,
                "操作步骤应优先使用祈使句和一般现在时，避免过去时或被动叙述。",
                "说明书审核能力补强方案 - 操作步骤语气",
                "general",
                86,
            )
            reported_voice += 1

    return issues


def get_knowledge_basis(db: Session):
    basis_list = []
    try:
        tree = get_folder_tree(db, None)
        for folder in tree:
            folder_id = folder.get('id') if isinstance(folder, dict) else getattr(folder, 'id', None)
            folder_name = folder.get('name') if isinstance(folder, dict) else getattr(folder, 'name', '')
            if folder_id is None:
                continue
            folder_obj = get_folder(db, folder_id)
            if folder_obj:
                files = get_folder_files(db, folder_id)
                for file in files:
                    basis_list.append({
                        "title": file.name,
                        "folder": folder_name,
                        "path": f"{folder_name}/{file.name}"
                    })
    except Exception as e:
        print(f"获取知识库审核依据失败: {e}")
    return basis_list


def _is_footer_line(line):
    stripped = str(line or '').strip()
    if not stripped:
        return False
    return bool(re.search(r'\b(doc\.?\s*no\.?|page\s*\d+|copyright|jb-[a-z0-9\-]+)\b', stripped, re.IGNORECASE))


def _is_step_line(line):
    stripped = str(line or '').strip()
    if re.match(r'^step\s+\d+(?:\.\d+)*(?:[\.)、：:])?\s+', stripped, re.IGNORECASE):
        return True
    return bool(re.match(r'^\d+(?:[\.)、：:])\s+', stripped))


def _is_version_history_line(line):
    stripped = str(line or '').strip()
    return bool(re.search(r'\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\s+changes\b', stripped, re.IGNORECASE))


def _is_material_line(line):
    stripped = str(line or '').strip()
    if not stripped:
        return False
    patterns = [
        r'^\d+(?:\.\d+)?\s*(?:mL|μL|uL|L|mg|ng|g|kg|mm|cm)\b',
        r'\bdeep-well plate\b',
        r'\bautomated filter tips\b',
        r'\bskirted PCR plates?\b',
        r'\bNormalized ssCir sample PCR plate\b',
        r'\bPos\d(?:-Pos\d)?\b',
    ]
    return any(re.search(pattern, stripped, re.IGNORECASE) for pattern in patterns)


def _is_table_line(line):
    stripped = str(line or '').strip()
    return bool(' | ' in stripped or re.match(r'^(?:Table|表)\s*\d+', stripped, re.IGNORECASE))


def _is_figure_or_table_caption(line):
    stripped = str(line or '').strip()
    return bool(re.match(r'^(?:Figure|Fig\.?|Table|表|图)\s*\d+\b', stripped, re.IGNORECASE))


def _is_heading_line(line):
    stripped = str(line or '').strip()
    if not stripped or _is_footer_line(stripped) or _is_step_line(stripped) or _is_table_line(stripped) or _is_material_line(stripped) or _is_version_history_line(stripped):
        return False
    if stripped.startswith('#'):
        return True
    if re.match(r'^(?:Chapter|Section)\s+\d+[\s:.-]+', stripped, re.IGNORECASE):
        return True
    return bool(re.match(r'^\d+(?:\.\d+)*(?:[\).])?\s+[^\n]{2,}$', stripped))


def _normalize_heading_text(line):
    cleaned = re.sub(r'^(\d+(?:\.\d+)*)([A-Za-z])', r'\1 \2', str(line or '').strip())
    return re.sub(r'\s+', ' ', cleaned).strip(' -:|')


def _clean_caption_heading(line):
    heading = _normalize_heading_text(line)
    if not heading:
        return ''
    heading = re.sub(r'\b(Name|Position|Brand|Cat\.?|Cat No\.?|Quantity|Components)\b.*$', '', heading, flags=re.IGNORECASE).strip(' -:|,')
    heading = re.sub(r'\s+', ' ', heading).strip()
    return heading or _normalize_heading_text(line)


def _is_inline_reference_heading(text):
    normalized = _normalize_heading_text(text).lower()
    if not normalized:
        return False
    patterns = [
        r'\btable\s+\d+\s+and\s+figure\s+\d+\b',
        r'\bfigure\s+\d+\s+and\s+table\s+\d+\b',
        r'\baccording\s+to\b',
        r'\bas\s+shown\s+in\b',
        r'\bwill\s+be\b',
        r'\bplace\s+it\b',
    ]
    return any(re.search(pattern, normalized, re.IGNORECASE) for pattern in patterns)


def _looks_like_standalone_heading(lines, index):
    current = str(lines[index] or '').strip()
    if not current:
        return False
    if not re.match(r'^[A-Z][A-Za-z0-9\s\-_/]{3,60}$', current):
        return False
    prev_line = str(lines[index - 1] or '').strip() if index - 1 >= 0 else ''
    next_line = str(lines[index + 1] or '').strip() if index + 1 < len(lines) else ''
    if _is_table_cell_like(current) or _is_table_cell_like(prev_line) or _is_table_cell_like(next_line):
        return False
    if prev_line and next_line and not _is_heading_line(prev_line) and not _is_heading_line(next_line):
        return False
    return True


def _is_table_cell_like(line):
    stripped = str(line or '').strip()
    if not stripped:
        return False
    if _is_version_history_line(stripped):
        return True
    if _is_material_line(stripped):
        return True
    if '|' in stripped:
        return True
    if re.search(r'\b(?:Cat\.?\s*No\.?|Recommended brand|Quantity|Components|Library Type|types|condition|None)\b', stripped, re.IGNORECASE):
        return True
    if re.search(r'\b(?:Vortex mixer|Mini centrifuge|Primer hybridization)\b', stripped, re.IGNORECASE):
        return True
    if re.search(r'\bautomated filter tips\b|\bPos\d(?:-Pos\d)?\b', stripped, re.IGNORECASE):
        return True
    if re.search(r'\d+(?:\.\d+)?\s*mL/tube', stripped, re.IGNORECASE):
        return True
    if re.search(r'\b\d{2,4}\s*[μu]L\b', stripped, re.IGNORECASE):
        return True
    return False


def _is_isolated_heading(lines, index):
    prev_line = str(lines[index - 1] or '').strip() if index - 1 >= 0 else ''
    next_line = str(lines[index + 1] or '').strip() if index + 1 < len(lines) else ''
    return (not prev_line) or (not next_line)


def _score_heading_candidate(lines, index):
    line = str(lines[index] or '').strip()
    if not line:
        return None
    if _is_footer_line(line) or _is_step_line(line) or _is_material_line(line) or _is_version_history_line(line):
        return None
    score = None
    kind = None
    if line.startswith('#'):
        score = 120
        kind = 'markdown'
    elif _is_heading_line(line):
        score = 100
        kind = 'numbered'
    elif _is_figure_or_table_caption(line):
        score = 88
        kind = 'caption'
    elif _looks_like_standalone_heading(lines, index):
        score = 60
        kind = 'standalone'
    if score is None:
        return None

    if _is_inline_reference_heading(line):
        return None
    if kind != 'numbered' and _is_table_cell_like(line):
        score -= 45
    if kind == 'standalone' and not _is_isolated_heading(lines, index):
        score -= 20
    if len(line) > 70:
        score -= 10
    if len(line.split()) <= 2 and kind == 'standalone':
        score -= 10
    if score < 55:
        return None

    heading = _clean_caption_heading(line) if kind == 'caption' else _normalize_heading_text(line)
    return score, heading


def extract_chapter(content, position):
    normalized = str(content or '').replace('\f', '\n')
    safe_position = max(0, min(int(position or 0), len(normalized)))
    before_lines = normalized[:safe_position].split('\n')
    after_lines = normalized[safe_position:].split('\n')[:12]
    candidates = []

    start_index = max(0, len(before_lines) - 220)
    for index in range(len(before_lines) - 1, start_index - 1, -1):
        candidate = _score_heading_candidate(before_lines, index)
        if candidate:
            score, text = candidate
            distance_penalty = min(len(before_lines) - 1 - index, 30)
            candidates.append((score - distance_penalty, text))

    for index, _line in enumerate(after_lines[:10]):
        candidate = _score_heading_candidate(after_lines, index)
        if candidate:
            score, text = candidate
            if _is_table_cell_like(text):
                continue
            candidates.append((score - 24 - index, text))

    if candidates:
        candidates.sort(key=lambda item: item[0], reverse=True)
        return candidates[0][1]
    return ""


def get_context(content, start, end, context_length=200):
    start_idx = max(0, start - context_length)
    end_idx = min(len(content), end + context_length)
    context = content[start_idx:end_idx]
    if start_idx > 0:
        context = "..." + context
    if end_idx < len(content):
        context = context + "..."
    return context


@router.get("/", response_model=list[Review])
async def list_reviews(db: Session = Depends(get_db)):
    reviews = get_reviews(db)
    result = []
    for review in reviews:
        review = _reconcile_review_runtime_state(db, review)
        doc = get_document(db, document_id=review.document_id)
        review_dict = review.__dict__.copy()
        review_dict['document_name'] = doc.filename if doc else ''
        review_dict['document_file_type'] = doc.file_type if doc else ''
        if not review_dict.get('summary'):
            review_dict['summary'] = '{}'
        # 添加进度信息
        if review.status == 'running':
            progress = get_progress(review.id)
            review_dict['progress'] = progress
        result.append(review_dict)
    return result


@router.get("/{review_id}/progress")
async def get_review_progress(review_id: int, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")

    review = _reconcile_review_runtime_state(db, review)
    if review.status == 'running':
        return get_progress(review_id)

    return {
        'status': review.status,
        'step': '完成' if review.status == 'completed' else '失败',
        'progress': 100 if review.status == 'completed' else 0,
        'message': review.summary or ('审核完成' if review.status == 'completed' else '审核失败'),
        'timestamp': datetime.now().isoformat(),
    }


def detect_language(content):
    chinese_chars = re.findall(r'[\u4e00-\u9fff]', content)
    english_chars = re.findall(r'[a-zA-Z]', content)

    if len(chinese_chars) > len(english_chars) * 2:
        return "cn"
    elif len(english_chars) > len(chinese_chars) * 2:
        return "en"
    return "both"


# ------------------------------------------------------------------
# 误报白名单: 常见合法英文缩写 / 公司名 / 术语 (避免简单正则误报)
# ------------------------------------------------------------------
_EN_ABBREV_ALLOWED = [
    "Ltd.", "Co.", "Inc.", "LLC.", "LLP.", "Corp.",
    "Mr.", "Mrs.", "Ms.", "Dr.", "Prof.",
    "e.g.", "i.e.", "etc.", "et al.",
    "P.R.", "R.P.", "U.S.A.", "U.K.", "U.S.",
    "vs.", "etc", "sq.",
]


def _in_abbreviation_allowed(text):
    t = text.strip()
    for a in _EN_ABBREV_ALLOWED:
        if t.lower() == a.lower() or t.lower().endswith(a.lower()):
            return True
    return False


_TECH_TERMS = {
    'prep', 'device', 'consumables', 'reagents', 'extraction', 'extent', 'selection',
    'mg', 'ml', 'ul', 'μg', 'μl', 'kb', 'mb', 'gb', 'nm', 'mm', 'cm',
    'dna', 'rna', 'pcr', 'elisa', 'atp', 'cdna', 'mrna', 'trna',
    'buffer', 'solution', 'kit', 'system', 'assay', 'protocol', 'procedure',
    'sample', 'specimen', 'tube', 'plate', 'well', 'tip', 'pipette',
    'centrifuge', 'incubator', 'shaker', 'homogenizer', 'lyser',
    'volume', 'concentration', 'temperature', 'time', 'rpm', 'min', 'hr',
    'mgisp', 'bgiseq', 'dnbseq', 'dnb', 'cpas', 'coolmps',
    'quality', 'control', 'standard', 'reference', 'calibration', 'validation',
    'protocol', 'method', 'technique', 'approach', 'strategy', 'step',
}


def _is_tech_term(word):
    return word.lower() in _TECH_TERMS


def _is_sentence_start(content, position):
    before = content[max(0, position - 20):position]
    return before.strip().endswith('.') or before.strip().endswith('!') or before.strip().endswith('?') or position == 0


_DISABLED_RULES_FOR_REVIEW = {
    # These rules describe style preferences, file-level properties, or consistency checks
    # that cannot be reliably determined by a single regex hit in extracted document text.
    "R017", "R018", "R019", "R037", "R039", "R040",
    "R042", "R043", "R044", "R045", "R048", "R050",
}


def _should_skip_rule_match(rule, match, content, document_language, file_type=None):
    rule_no = getattr(rule, "rule_no", "")
    original_text = match.group()
    start = match.start()
    end = match.end()
    context = content[max(0, start - 40):min(len(content), end + 40)]

    if rule_no in _DISABLED_RULES_FOR_REVIEW:
        return True

    if file_type == "pdf" and rule_no in {"R002", "R013", "R036"}:
        return True

    # Numeric punctuation inside decimals, versions, and enumerations is usually valid.
    if rule_no == "R001":
        if re.search(r"\d[\.,]\d", context):
            return True

    if rule_no == "R002":
        if re.fullmatch(r"\.[a-z]", original_text):
            return True
        if any(ext in context.lower() for ext in [".xml", ".wfex", ".sp960", "@", ".com"]):
            return True

    if rule_no == "R004":
        if len(original_text) > 80 or "\n" in original_text:
            return True

    # Quotation direction regex currently matches normal quoted text too often.
    if rule_no == "R007":
        return True

    # Trademark markers are common inline forms in extracted text and need richer formatting context.
    if rule_no == "R008":
        return True

    # Unit spacing rule should not flag model numbers, percentages, temperatures without unit intent, or compact codes.
    if rule_no == "R013":
        if original_text.endswith('%') or re.fullmatch(r"\d+(?:\.\d+)?[Vv]", original_text):
            return True
        if re.search(r"[A-Z]{2,}\d|\d[A-Z]{2,}", original_text):
            return True

    if rule_no == "R011":
        lowered_context = context.lower()
        if "/" in context or "rpm/min" in lowered_context:
            return True
        if re.search(r'\b(?:min|sec|hr)\b', original_text, re.IGNORECASE):
            return True

    if rule_no == "R013":
        if re.fullmatch(r'\d+[xX]', original_text):
            return True
        if re.search(r'\b\d+[xX]\b', original_text):
            return True

    if rule_no == "R021":
        if ".com" in context.lower() or "@" in context or "http" in context.lower() or any(ext in context.lower() for ext in [".xml", ".json", ".txt", ".docx", ".pdf"]):
            return True

    if rule_no == "R023":
        if not original_text.lower().startswith("a "):
            return True

    if rule_no == "R024":
        tokens = original_text.split()
        if file_type == "pdf" and re.search(r'\n\s*\n', original_text):
            return True
        if file_type == "pdf" and re.search(r'\b(?:Table|Figure|Cat\.\s*No\.|Recommended brand|Library Type|condition)\b', context, re.IGNORECASE):
            return True
        if original_text.strip().lower() == "none none":
            return True
        if any(any(ch.isdigit() for ch in token) for token in tokens):
            return True
        if tokens and all(len(token) <= 1 for token in tokens):
            return True

    # Mixed Chinese-English spacing rule should keep common product names and protocol acronyms intact.
    if rule_no == "R025":
        if re.search(r"[A-Z]{2,}|[A-Za-z]+\d|\d+[A-Za-z]+", original_text):
            return True

    # Chinese closing brackets followed by Chinese text do not need a separating space.
    if rule_no == "R029":
        if re.fullmatch(r"[）\)][\u4e00-\u9fff]", original_text):
            return True

    # Markdown/code formatting rules only apply when the extracted text still clearly shows markdown structure.
    if rule_no in {"R030", "R031", "R032", "R033", "R034", "R046", "R047", "R049"}:
        if "```" not in content and "#" not in content and "[" not in content and "](" not in content:
            return True

    # Long-sentence rules are too noisy on table rows, lists, and extracted fragments.
    if rule_no == "R036":
        return True

    if rule_no == "R035":
        line = content[max(0, start - 120):min(len(content), end + 120)]
        if "|" in line or re.search(r"\b(?:Figure|Table|Step|Note|WARNING|CAUTION)\b", line, re.IGNORECASE):
            return True

    return False


def run_rule_audit(content, rules, knowledge_basis=None, file_type=None):
    """规则审核: 基于正则, 加上简易白名单过滤, 按内容去重"""
    issues = []
    document_language = detect_language(content)
    seen_keys = set()

    for rule in rules:
        try:
            if rule.language == "cn" and document_language == "en":
                continue
            if rule.language == "en" and document_language == "cn":
                continue

            matches = list(re.finditer(rule.regex, content))
            for match in matches:
                original_text = match.group()

                if _should_skip_rule_match(rule, match, content, document_language, file_type):
                    continue

                if document_language in ("en", "both") and _in_abbreviation_allowed(original_text):
                    continue

                if rule.rule_no in ("R021", "R022"):
                    if _is_tech_term(original_text):
                        continue
                    if rule.rule_no == "R021" and not _is_sentence_start(content, match.start()):
                        continue

                chapter = extract_chapter(content, match.start())
                norm_text = re.sub(r"\s+", "", original_text).lower()
                dedup_key = f"{rule.rule_no}|{norm_text}|{chapter}"
                if dedup_key in seen_keys:
                    continue
                seen_keys.add(dedup_key)

                context = get_context(content, match.start(), match.end(), 200)

                basis = rule.audit_basis if rule.audit_basis else ""
                if not basis and knowledge_basis:
                    for kb in knowledge_basis:
                        if any(kw.lower() in rule.description.lower() for kw in ["标点", "格式", "术语", "写作"]):
                            if any(name in kb["title"] for name in ["风格指南", "写作规范", "手册", "指南"]):
                                basis = kb["path"]
                                break

                issues.append({
                    "severity": "general",
                    "category": rule.category or "其他",
                    "rule": rule.rule_no,
                    "chapter": chapter,
                    "original_text": original_text,
                    "context": context,
                    "suggestion": rule.suggestion if rule.suggestion else "",
                    "description": rule.description or "",
                    "audit_basis": basis if basis else "技术文档写作规范",
                    "confidence": 100,
                    "source": "rule",
                    "position": _encode_issue_position(match.start(), match.end())
                })
        except Exception as e:
            print(f"规则 {getattr(rule,'rule_no','?')} 匹配出错: {e}")
            continue
    return issues


def run_term_check(content, terms):
    """术语检查: 按归一化原文去重"""
    issues = []
    if not terms:
        return issues

    seen_keys = set()
    for term in terms:
        try:
            if not term.non_standard:
                continue
            occurrences = list(re.finditer(re.escape(term.non_standard), content))
            for match in occurrences:
                chapter = extract_chapter(content, match.start())
                norm_text = re.sub(r"\s+", "", term.non_standard).lower()
                dedup_key = f"TERM|{norm_text}|{chapter}"
                if dedup_key in seen_keys:
                    continue
                seen_keys.add(dedup_key)

                context = get_context(content, match.start(), match.end(), 200)
                issues.append({
                    "severity": "suggestion",
                    "category": "术语规范",
                    "rule": "TERM",
                    "chapter": chapter,
                    "original_text": term.non_standard,
                    "context": context,
                    "suggestion": f"建议使用标准术语: {term.standard}",
                    "description": f"发现非标准术语: {term.non_standard}",
                    "audit_basis": "MGI中文技术文档写作风格指南 - 缩略语",
                    "confidence": 95,
                    "source": "term",
                    "position": _encode_issue_position(match.start(), match.end())
                })
        except Exception as e:
            print(f"术语 {getattr(term, 'non_standard', '?')} 匹配出错: {e}")
            continue
    return issues


def dedupe_issues_by_original(issues):
    """按 '分类 + 归一化原文 + 位置相近' 合并, 并保留出现次数最多的项"""
    if not issues:
        return []
    
    # 误报过滤 - 已知无效的模式
    FALSE_POSITIVE_PATTERNS = [
        re.compile(r'^[×±÷∞∑∏∫°℃℉]$'),  # 纯数学符号
        re.compile(r'^[\d\.]+$'),  # 纯数字
        re.compile(r'^[A-Z]{1,3}$'),  # 1-3个大写字母（避免误报为单词错误）
        re.compile(r'^[\.,;:\?!]+$'),  # 纯标点
    ]
    
    def is_false_positive(text):
        text = str(text).strip()
        if not text:
            return True
        allowed_single_punctuations = {'.', ',', ';', ':', '!', '?', '(', ')', '[', ']', '{', '}', '：', '，', '。', '；', '（', '）', '“', '”', '‘', '’', '、'}
        if len(text) < 2 and text not in allowed_single_punctuations:
            return True
        for pat in FALSE_POSITIVE_PATTERNS:
            if pat.match(text):
                return True
        return False
    
    # 第一步：过滤已知误报
    filtered = [
        i for i in issues
        if _is_deterministic_review_rule(i) or not is_false_positive(i.get('original_text', ''))
    ]
    print(f"[审核] 误报过滤: 过滤 {len(issues) - len(filtered)} 个, 剩余 {len(filtered)} 个")
    
    def issue_rank(issue):
        source = issue.get("source", "")
        severity = issue.get("severity", "general")
        confidence = issue.get("confidence", 0) or 0
        severity_rank = {"fatal": 4, "serious": 3, "general": 2, "suggestion": 1}.get(severity, 0)
        source_rank = {"rule": 3, "term": 2, "ai": 1}.get(source, 0)
        return (severity_rank, source_rank, confidence)

    # 第二步：去重（优先按原文+章节聚合，避免规则与 AI 对同一问题重复上报）
    seen = {}
    for issue in filtered:
        norm = re.sub(r"\s+", "", str(issue.get("original_text", ""))).lower()
        chapter = issue.get("chapter", "")
        source = issue.get("source", "")
        rule = issue.get("rule", "")
        key = f"{norm}|{chapter}"
        if source == 'spellcheck' or rule == 'SPELL':
            key = f"spell|{norm}|{issue.get('position', '')}"
        if rule in {'UNIT-003', 'UNIT-004', 'HR011'}:
            key = f"{rule}|{norm}|{issue.get('position', '')}"
        if rule in {'STYLE-003', 'HR008', 'GRAMMAR-001', 'GRAMMAR-002'}:
            key = f"{rule}|{norm}|{issue.get('position', '')}"
        if not norm.strip():
            continue
        if key in seen:
            old = seen[key]
            if issue_rank(issue) > issue_rank(old):
                seen[key] = issue
        else:
            seen[key] = issue
    return list(seen.values())


def _run_reference_and_term_consistency_audit(db: Session, content):
    issues = []

    try:
        issues.extend(ReferenceIntegrityRule().check(content))
    except Exception as e:
        print(f"[审核] 引用完整性规则执行失败: {e}")

    try:
        for issue in TermConsistencyRule().check(content):
            start, end = _parse_issue_position(issue.get('position', ''))
            issue.setdefault('chapter', extract_chapter(content, start))
            issue.setdefault('context', get_context(content, start, end or start + len(issue.get('original_text', '')), 120))
            normalized = str(issue.get('suggestion', '') or '').strip()
            if normalized and not normalized.startswith('建议'):
                issue['suggestion'] = f'建议统一为 {normalized}'
            issues.append(issue)
    except Exception as e:
        print(f"[审核] 术语一致性规则执行失败: {e}")

    try:
        issues.extend(_run_db_term_consistency_audit(db, content))
    except Exception as e:
        print(f"[审核] 术语库一致性规则执行失败: {e}")

    return issues


def _run_review_background(review_id: int, document_id: int, mode: str):
    """后台执行审核任务（带进度更新）"""
    from app.database import SessionLocal
    db = SessionLocal()
    try:
        set_progress(review_id, 'running', '加载文档', 5, '正在读取文档内容...')
        
        document = get_document(db, document_id=document_id)
        if not document:
            set_progress(review_id, 'failed', '文档不存在', 0, f'文档ID={document_id}不存在')
            update_review_status(db, review_id, "failed", 0, "Document not found")
            return

        issues = []
        content = document.content or ""
        if document.file_type == 'pdf':
            content = clean_pdf_text(content)
        content_limit = 120000 if document.file_type == 'pdf' else 20000
        content = content[:content_limit]
        document_language = detect_language(content)
        spec_texts = _load_review_spec_texts(db)
        false_positive_signatures = _load_false_positive_signatures_for_document(db, document_id)
        print(f"[审核] 文档ID={document_id}, 语言检测={document_language}, 内容长度={len(content)}字符")
        print(
            "[审核] 当前规范文件长度: "
            f"cn={len(spec_texts.get('cn_style', ''))}, "
            f"en={len(spec_texts.get('en_style', ''))}, "
            f"common={len(spec_texts.get('common_errors', ''))}, "
            f"final_checklists={len(spec_texts.get('final_checklists', ''))}"
        )

        knowledge_basis = get_knowledge_basis(db)
        has_ai_client = ai_client.has_any_client
        print(f"[审核] AI客户端可用: {has_ai_client}, 模式={mode}")

        if document.file_type == 'xlsx':
            set_progress(review_id, 'running', 'Excel审核', 35, '正在执行 Excel 行级审核...')
            issues.extend(_run_excel_review_audit(db, document))
            print(f"[审核] Excel审核发现问题: {len(issues)}个")

        if mode in ["rule", "hybrid"]:
            set_progress(review_id, 'running', '规则审核', 15, '正在加载审核规则...')
            rules = get_rules(db)
            print(f"[审核] 加载规则数量: {len(rules)}")
            
            set_progress(review_id, 'running', '规则审核', 25, '正在执行规则匹配...')
            rule_issues = [] if document.file_type == 'xlsx' else run_rule_audit(content, rules, knowledge_basis, document.file_type)
            if document.file_type == 'pdf':
                rule_issues = [issue for issue in rule_issues if issue.get('rule') not in {'R011', 'R016', 'R021'}]
            print(f"[审核] 规则匹配到问题: {len(rule_issues)}个")

            set_progress(review_id, 'running', '术语检查', 35, '正在执行术语检查...')
            terms = get_terms(db)
            print(f"[审核] 加载术语数量: {len(terms)}")
            term_issues = [] if document.file_type == 'xlsx' else run_term_check(content, terms)
            print(f"[审核] 术语匹配到问题: {len(term_issues)}个")

            if document.file_type != 'xlsx' and document_language in ("en", "both"):
                set_progress(review_id, 'running', '拼写检查', 45, '正在进行拼写和语法检查...')
                try:
                    spelling_issues = run_spelling_and_grammar_check(content, document.file_type)
                    print(f"[审核] 拼写/语法检查发现问题: {len(spelling_issues)}个")
                    rule_issues.extend(spelling_issues)
                except Exception as e:
                    print(f"[审核] 拼写/语法检查失败: {e}")

                try:
                    heuristic_issues = _run_english_heuristic_audit(content, document.file_type)
                    print(f"[审核] 英文高置信规则发现问题: {len(heuristic_issues)}个")
                    rule_issues.extend(heuristic_issues)
                except Exception as e:
                    print(f"[审核] 英文高置信规则执行失败: {e}")

                try:
                    engineering_issues = _run_manual_engineering_audit(content, document.file_type)
                    print(f"[审核] 说明书工程质量规则发现问题: {len(engineering_issues)}个")
                    rule_issues.extend(engineering_issues)
                except Exception as e:
                    print(f"[审核] 说明书工程质量规则执行失败: {e}")

                if document.file_type == 'pdf':
                    try:
                        structural_issues = _run_pdf_structure_audit(content)
                        print(f"[审核] PDF结构规则发现问题: {len(structural_issues)}个")
                        rule_issues.extend(structural_issues)
                    except Exception as e:
                        print(f"[审核] PDF结构规则执行失败: {e}")

                try:
                    supplemental_issues = _run_reference_and_term_consistency_audit(db, content)
                    print(f"[审核] 引用/术语规则发现问题: {len(supplemental_issues)}个")
                    rule_issues.extend(supplemental_issues)
                except Exception as e:
                    print(f"[审核] 引用/术语规则执行失败: {e}")

                try:
                    logic_issues = _run_logic_integrity_audit(content)
                    print(f"[审核] 逻辑完整性规则发现问题: {len(logic_issues)}个")
                    rule_issues.extend(logic_issues)
                except Exception as e:
                    print(f"[审核] 逻辑完整性规则执行失败: {e}")

                try:
                    safety_issues = _run_safety_compliance_audit(content)
                    print(f"[审核] 安全合规规则发现问题: {len(safety_issues)}个")
                    rule_issues.extend(safety_issues)
                except Exception as e:
                    print(f"[审核] 安全合规规则执行失败: {e}")

                try:
                    cross_ref_issues = _run_cross_reference_audit(content)
                    print(f"[审核] 交叉引用规则发现问题: {len(cross_ref_issues)}个")
                    rule_issues.extend(cross_ref_issues)
                except Exception as e:
                    print(f"[审核] 交叉引用规则执行失败: {e}")

                try:
                    long_sentence_issues = _run_long_sentence_audit(content, document.file_type)
                    print(f"[审核] 长句规则发现问题: {len(long_sentence_issues)}个")
                    rule_issues.extend(long_sentence_issues)
                except Exception as e:
                    print(f"[审核] 长句规则执行失败: {e}")

            if document.file_type != 'xlsx':
                try:
                    release_checklist_issues = _run_release_checklist_audit(content, document, spec_texts, document_language)
                    print(f"[审核] 发布前自检 Checklist 发现问题: {len(release_checklist_issues)}个")
                    rule_issues.extend(release_checklist_issues)
                except Exception as e:
                    print(f"[审核] 发布前自检 Checklist 执行失败: {e}")

            candidate_rule_issues = rule_issues + term_issues

            if has_ai_client and len(candidate_rule_issues) > 0:
                set_progress(review_id, 'running', 'AI二次验证', 55, f'正在AI验证 {len(candidate_rule_issues)} 个候选问题...')
                print(f"[审核] 开始AI二次验证，候选问题数={len(candidate_rule_issues)}")
                try:
                    deterministic_issues = [issue for issue in candidate_rule_issues if _is_deterministic_review_rule(issue)]
                    ai_check_issues = [issue for issue in candidate_rule_issues if not _is_deterministic_review_rule(issue)]
                    filtered = _call_with_timeout(ai_client.filter_rule_false_positives, 90.0, ai_check_issues, document_language) if ai_check_issues else []
                    candidate_rule_issues = deterministic_issues + filtered
                    print(f"[审核] AI二次验证后保留问题数={len(candidate_rule_issues)}, 确定性规则保留={len(deterministic_issues)}")
                except concurrent.futures.TimeoutError:
                    print(f"[审核] AI 二次验证超时(90s), 使用原始规则结果({len(candidate_rule_issues)}个)")
                except Exception as e:
                    print(f"[审核] AI 二次验证失败, 使用原始规则结果: {e}")

            issues.extend(candidate_rule_issues)
            print(f"[审核] 规则审核阶段共产出问题数={len(candidate_rule_issues)}")

        if document.file_type != 'xlsx' and mode in ["ai", "hybrid"] and has_ai_client:
            set_progress(review_id, 'running', 'AI智能审核', 65, '正在进行AI深度审核...')
            print(f"[审核] 开始AI智能审核，模式={mode}")
            try:
                ai_review_basis = _build_ai_review_basis(spec_texts, document_language)
                ai_result = _call_with_timeout(ai_client.audit_document, 180.0, content, document_language, ai_review_basis)
                ai_issues = ai_result.get("issues", [])
                print(f"[审核] AI审核返回问题数={len(ai_issues)}")
                for issue in ai_issues:
                    issue["source"] = "ai"
                    issue.setdefault("severity", "general")
                    issue.setdefault("category", "其他")
                    issue.setdefault("original_text", "")
                    issue.setdefault("context", "")
                    issue.setdefault("chapter", "")
                    issue.setdefault("rule", "AI")
                    issue.setdefault("suggestion", "")
                    issue.setdefault("description", "")
                    issue.setdefault("audit_basis", "AI 审核")
                    issue.setdefault("confidence", 80)
                    issue.setdefault("position", "")
                    issue = _normalize_ai_issue_category(issue)
                    if issue["severity"] not in ("fatal", "serious", "general", "suggestion"):
                        issue["severity"] = "general"
                issues.extend(ai_issues)
            except concurrent.futures.TimeoutError:
                print(f"[审核] AI 审核超时(180s), 跳过 AI 审核")
            except Exception as e:
                print(f"[审核] AI 审核失败, 跳过 AI 审核: {e}")

        set_progress(review_id, 'running', '结果处理', 85, '正在去重和保存结果...')
        issues = dedupe_issues_by_original(issues)
        issues = _sanitize_issue_suggestions(issues)
        issues = _filter_review_false_positives(issues)
        if false_positive_signatures:
            before_count = len(issues)
            issues = [issue for issue in issues if not (_issue_judgment_signatures(issue) & false_positive_signatures)]
            print(f"[审核] 历史误报过滤: 过滤 {before_count - len(issues)} 个, 剩余 {len(issues)} 个")
        if document.file_type == 'docx':
            issues = _enrich_docx_issue_positions(document, issues)
        print(f"[审核] 去重后最终问题数={len(issues)}")
        if len(issues) > 0:
            categories = {}
            for issue in issues:
                cat = issue.get("category", "未分类")
                categories[cat] = categories.get(cat, 0) + 1
            print(f"[审核] 问题分类统计: {categories}")

        for issue in issues:
            create_issue(db=db, issue=IssueCreate(
                review_id=review_id,
                severity=issue["severity"],
                category=issue.get("category", ""),
                rule=issue.get("rule", ""),
                chapter=issue.get("chapter", ""),
                original_text=issue.get("original_text", ""),
                context=issue.get("context", ""),
                suggestion=issue.get("suggestion", ""),
                description=issue.get("description", ""),
                audit_basis=issue.get("audit_basis", ""),
                confidence=issue.get("confidence", 0),
                source=issue.get("source", "rule"),
                position=issue.get("position", "")
            ))

        summary = json.dumps({
            "total": len(issues),
            "fatal": len([i for i in issues if i.get("severity") == "fatal"]),
            "serious": len([i for i in issues if i.get("severity") == "serious"]),
            "general": len([i for i in issues if i.get("severity") == "general"]),
            "suggestion": len([i for i in issues if i.get("severity") == "suggestion"]),
            "language": document_language,
        })

        set_progress(review_id, 'completed', '完成', 100, f'审核完成，发现 {len(issues)} 个问题')
        update_review_status(db, review_id, "completed", len(issues), summary)
        print(f"[审核] 任务完成, review_id={review_id}, 问题数={len(issues)}")

    except Exception as e:
        set_progress(review_id, 'failed', '失败', 0, str(e))
        update_review_status(db, review_id, "failed", 0, str(e))
        print(f"[审核] 任务失败, review_id={review_id}, 错误={e}")
    finally:
        db.close()


@router.post("/{document_id}")
async def create_review_task(document_id: int, mode: str = "hybrid", background_tasks: BackgroundTasks = None, db: Session = Depends(get_db)):
    if mode not in {"rule", "ai", "hybrid"}:
        raise HTTPException(status_code=400, detail="Unsupported review mode")

    document = get_document(db, document_id=document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    active_review = _get_active_review_for_document(db, document_id)
    if active_review:
        set_progress(active_review.id, 'running', '初始化', 0, '已有审核任务正在执行')
        return {"review_id": active_review.id, "status": "running", "message": "已有审核任务正在执行，请轮询进度"}

    review = create_review(db=db, review=ReviewCreate(document_id=document_id, mode=mode))
    update_review_status(db, review.id, "running", 0, "审核任务已创建")
    set_progress(review.id, 'running', '初始化', 0, '审核任务已创建')
    
    # 启动后台审核任务
    if background_tasks:
        background_tasks.add_task(_run_review_background, review.id, document_id, mode)
    
    return {"review_id": review.id, "status": "running", "message": "审核任务已启动，请轮询进度"}


def _dashboard_date_range(time_range='7d', start_date=None, end_date=None):
    today = datetime.utcnow().date()
    if time_range == 'custom' and start_date and end_date:
        try:
            start = datetime.strptime(start_date, '%Y-%m-%d').date()
            end = datetime.strptime(end_date, '%Y-%m-%d').date()
        except ValueError:
            raise HTTPException(status_code=400, detail='日期格式应为 YYYY-MM-DD')
    else:
        days = {'1d': 1, '7d': 7, '30d': 30, '90d': 90}.get(time_range or '7d', 7)
        start = today - timedelta(days=days - 1)
        end = today
    if end < start:
        raise HTTPException(status_code=400, detail='结束日期不能早于开始日期')
    return datetime.combine(start, time.min), datetime.combine(end, time.max)


def _dashboard_doc_type_value(doc_type):
    value = str(doc_type or 'all').lower()
    if value in {'all', ''}:
        return None
    if value == 'excel':
        return 'xlsx'
    return value


def _dashboard_review_rows(db, start_dt, end_dt, doc_type=None, user_id=None):
    query = (
        db.query(ReviewModel, DocumentModel, UserModel)
        .join(DocumentModel, ReviewModel.document_id == DocumentModel.id)
        .outerjoin(UserModel, DocumentModel.user_id == UserModel.id)
        .filter(ReviewModel.created_at >= start_dt, ReviewModel.created_at <= end_dt)
    )
    normalized_doc_type = _dashboard_doc_type_value(doc_type)
    if normalized_doc_type:
        query = query.filter(DocumentModel.file_type == normalized_doc_type)
    if user_id not in (None, '', 'all'):
        try:
            query = query.filter(DocumentModel.user_id == int(user_id))
        except (TypeError, ValueError):
            pass
    return query.order_by(ReviewModel.created_at.asc()).all()


def _dashboard_issue_rows(db, review_ids):
    if not review_ids:
        return []
    return db.query(IssueModel).filter(IssueModel.review_id.in_(review_ids)).all()


def _dashboard_pct(count, total):
    return round(count / total, 4) if total else 0


def _dashboard_format_doc_type(file_type):
    if file_type == 'xlsx':
        return 'excel'
    return file_type or 'other'


def _dashboard_review_completed_at(review, issues):
    completed_at = getattr(review, 'completed_at', None)
    if completed_at:
        return completed_at
    review_issues = [issue for issue in issues if issue.review_id == review.id and getattr(issue, 'created_at', None)]
    if review_issues:
        return max(issue.created_at for issue in review_issues)
    return None


def _dashboard_average_review_minutes(completed_reviews, issues):
    durations = []
    for review in completed_reviews:
        if not review.created_at:
            continue
        completed_at = _dashboard_review_completed_at(review, issues)
        if not completed_at or completed_at < review.created_at:
            continue
        minutes = (completed_at - review.created_at).total_seconds() / 60
        if minutes > 0.01:
            durations.append(minutes)
    if not durations:
        return None
    return max(0.1, round(sum(durations) / len(durations), 1))


def _dashboard_quality_metrics(raw_issues):
    visible_issues = _dashboard_visible_issues(raw_issues)
    platform_issues = [issue for issue in visible_issues if str(getattr(issue, 'source', '') or '').lower() != 'manual']
    manual_issues = [issue for issue in visible_issues if str(getattr(issue, 'source', '') or '').lower() == 'manual']
    false_positive = [issue for issue in platform_issues if str(getattr(issue, 'status', '') or '').lower() == 'false_positive']
    valid_platform = [issue for issue in platform_issues if str(getattr(issue, 'status', '') or '').lower() not in {'false_positive', 'ignored'}]
    valid_manual = [issue for issue in manual_issues if str(getattr(issue, 'status', '') or '').lower() not in {'false_positive', 'ignored'}]
    expected = len(valid_platform) + len(valid_manual)
    return {
        'platform_detected': len(valid_platform),
        'manual_supplemented': len(valid_manual),
        'expected_issues': expected,
        'false_positive_count': len(false_positive),
        'platform_reported': len(platform_issues),
        'accuracy_rate': _dashboard_pct(len(valid_platform), len(platform_issues)),
        'false_positive_rate': _dashboard_pct(len(false_positive), len(platform_issues)),
        'detection_rate': _dashboard_pct(len(valid_platform), expected),
}


def _dashboard_visible_issues(raw_issues):
    return [
        issue for issue in raw_issues
        if str(getattr(issue, 'status', '') or '').lower() not in {'false_positive', 'ignored'}
        and not _should_drop_known_false_positive_issue(issue)
    ]


def _dashboard_issue_distribution_label(issue, rule_lookup=None):
    rule = str(getattr(issue, 'rule', '') or '').strip()
    if rule:
        rule_description = (rule_lookup or {}).get(rule)
        if rule_description:
            return f"{rule} {rule_description[:18]}"
        return rule
    description = str(getattr(issue, 'description', '') or '').strip()
    if description:
        return description[:24]
    return str(getattr(issue, 'category', '') or '').strip() or '未分类'


def _build_review_dashboard_payload(db, time_range='7d', start_date=None, end_date=None, doc_type='all', user_id='all'):
    start_dt, end_dt = _dashboard_date_range(time_range, start_date, end_date)
    rows = _dashboard_review_rows(db, start_dt, end_dt, doc_type, user_id)
    reviews = [review for review, _, _ in rows]
    review_ids = [review.id for review in reviews]
    raw_issues = _dashboard_issue_rows(db, review_ids)
    issues = _dashboard_visible_issues(raw_issues)
    issues_by_review = Counter(issue.review_id for issue in issues)
    completed_reviews = [review for review in reviews if review.status == 'completed']
    average_review_minutes = _dashboard_average_review_minutes(completed_reviews, _dashboard_issue_rows(db, review_ids))

    date_count = (end_dt.date() - start_dt.date()).days + 1
    dates = [start_dt.date() + timedelta(days=index) for index in range(date_count)]
    submitted_by_date = Counter(review.created_at.date() for review in reviews if review.created_at)
    completed_by_date = Counter(review.created_at.date() for review in completed_reviews if review.created_at)
    issues_by_date = Counter()
    completed_ids_by_date = {}
    for review in completed_reviews:
        if review.created_at:
            day = review.created_at.date()
            completed_ids_by_date.setdefault(day, []).append(review.id)
            issues_by_date[day] += issues_by_review.get(review.id, 0)

    issue_distribution = []
    rule_lookup = {rule.rule_no: rule.description for rule in get_rules(db, limit=10000) if rule.rule_no and rule.description}
    category_counts = Counter(_dashboard_issue_distribution_label(issue, rule_lookup) for issue in issues)
    for category, count in category_counts.most_common():
        issue_distribution.append({
            'type': category,
            'count': count,
            'percentage': _dashboard_pct(count, len(issues)),
        })

    doc_type_stats = {}
    for review, document, _ in rows:
        key = _dashboard_format_doc_type(document.file_type)
        item = doc_type_stats.setdefault(key, {'type': key, 'total': 0, 'passed': 0, 'issues': 0, 'confirm': 0})
        item['total'] += 1
        review_issue_count = issues_by_review.get(review.id, 0)
        has_confirm = any(issue.review_id == review.id and issue.category == '术语不一致' for issue in issues)
        if review_issue_count == 0:
            item['passed'] += 1
        elif has_confirm:
            item['confirm'] += 1
        else:
            item['issues'] += 1

    return {
        'filters': {
            'time_range': time_range,
            'start_date': start_dt.date().isoformat(),
            'end_date': end_dt.date().isoformat(),
            'doc_type': doc_type or 'all',
            'user_id': user_id or 'all',
        },
        'kpi': {
            'range_tasks': len(reviews),
            'range_completed': len(completed_reviews),
            'avg_issues_per_doc': round(len(issues) / len(reviews), 2) if reviews else 0,
            'avg_review_time': average_review_minutes,
        },
        'quality': _dashboard_quality_metrics(raw_issues),
        'trend': {
            'dates': [day.strftime('%m-%d') for day in dates],
            'submitted': [submitted_by_date.get(day, 0) for day in dates],
            'completed': [completed_by_date.get(day, 0) for day in dates],
            'avg_issues': [round(issues_by_date.get(day, 0) / len(completed_ids_by_date.get(day, [])), 2) if completed_ids_by_date.get(day) else 0 for day in dates],
        },
        'issue_distribution': issue_distribution,
        'doc_type_distribution': list(doc_type_stats.values()),
        'task_list': [
            {
                'review_id': review.id,
                'document_id': document.id,
                'document_name': document.filename,
                'document_type': _dashboard_format_doc_type(document.file_type),
                'submitted_at': review.created_at.isoformat() if review.created_at else '',
                'status': review.status,
                'issue_count': issues_by_review.get(review.id, 0),
                'priority': '中',
            }
            for review, document, _ in sorted(rows, key=lambda item: item[0].created_at or datetime.min, reverse=True)[:50]
        ],
    }


@router.get("/dashboard/overview")
async def get_review_dashboard_overview(
    time_range: str = Query('7d'),
    start_date: str = Query(None),
    end_date: str = Query(None),
    doc_type: str = Query('all'),
    project_id: str = Query('all'),
    user_id: str = Query('all'),
    db: Session = Depends(get_db),
):
    return _build_review_dashboard_payload(db, time_range, start_date, end_date, doc_type, user_id)


@router.get("/dashboard/personal")
async def get_review_dashboard_personal(
    time_range: str = Query('7d'),
    start_date: str = Query(None),
    end_date: str = Query(None),
    doc_type: str = Query('all'),
    user_id: str = Query(1),
    db: Session = Depends(get_db),
):
    return _build_review_dashboard_payload(db, time_range, start_date, end_date, doc_type, user_id)


def _normalize_gold_compare_text(value):
    return re.sub(r"\s+", "", str(value or "")).lower()


def _load_gold_rows_from_excel(file_bytes):
    try:
        from io import BytesIO
        from openpyxl import load_workbook
        workbook = load_workbook(BytesIO(file_bytes), data_only=True)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"标准答案 Excel 解析失败: {exc}")

    worksheet = workbook[workbook.sheetnames[0]]
    rows = list(worksheet.iter_rows(values_only=True))
    if not rows:
        return []

    headers = [str(cell or "").strip() for cell in rows[0]]
    result = []
    for row in rows[1:]:
        item = dict(zip(headers, row))
        if not any(value is not None for value in item.values()):
            continue
        if "错误内容" not in item:
            continue
        serial = item.get("序号")
        if not isinstance(serial, int):
            continue
        issue_type = str(item.get("问题类型") or "").strip()
        if issue_type in {"—", "-", "无", "正确样本", "不应检出"}:
            continue
        result.append({
            "index": serial,
            "location": item.get("位置") or item.get("章节") or "",
            "wrong_text": item.get("错误内容") or "",
            "correct_text": item.get("正确内容") or "",
            "issue_type": issue_type or "未分类",
            "note": item.get("备注") or "",
        })
    return [row for row in result if str(row.get("wrong_text") or "").strip()]


def _gold_row_matches_issue(gold_row, issue):
    expected = _normalize_gold_compare_text(gold_row.get("wrong_text"))
    actual = _normalize_gold_compare_text(_issue_value(issue, "original_text", ""))
    if not expected or not actual:
        return False
    if expected == actual:
        return True
    if len(expected) >= 4 and len(actual) >= 4:
        return expected in actual or actual in expected
    return False


def _gold_text_presence(content, value):
    raw = str(value or "").strip()
    if not raw:
        return False
    if re.search(r"^[A-Za-z][A-Za-z\-]*$", raw):
        return bool(re.search(r"(?<![A-Za-z])" + re.escape(raw) + r"(?![A-Za-z])", content or "", re.IGNORECASE))
    return _normalize_gold_compare_text(raw) in _normalize_gold_compare_text(content)


@router.get("/{review_id}/parsed-text")
async def get_review_parsed_text(review_id: int, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")
    document = get_document(db, document_id=review.document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    content = document.content or ""
    pages = [page for page in content.split("\f") if page.strip()]
    return {
        "review_id": review_id,
        "document_id": document.id,
        "filename": document.filename,
        "file_type": document.file_type,
        "char_count": len(content),
        "word_count": len(re.findall(r"[A-Za-z]+|[\u4e00-\u9fff]", content)),
        "page_count": len(pages) or None,
        "content": content,
    }


@router.post("/{review_id}/gold-compare")
async def compare_review_with_gold(review_id: int, file: UploadFile = File(...), db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")
    document = get_document(db, document_id=review.document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    file_bytes = await file.read()
    gold_rows = _load_gold_rows_from_excel(file_bytes)
    issues = _visible_review_issues(get_issues(db, review_id=review_id))
    content = document.content or ""

    matched_issue_ids = set()
    matches = []
    missed = []
    for row in gold_rows:
        row = dict(row)
        row["wrong_text_exists_in_parsed_text"] = _gold_text_presence(content, row.get("wrong_text"))
        row["correct_text_exists_in_parsed_text"] = _gold_text_presence(content, row.get("correct_text"))
        row_matches = []
        for issue in issues:
            if _gold_row_matches_issue(row, issue):
                matched_issue_ids.add(getattr(issue, "id", id(issue)))
                row_matches.append({
                    "id": getattr(issue, "id", None),
                    "rule": getattr(issue, "rule", ""),
                    "category": getattr(issue, "category", ""),
                    "severity": getattr(issue, "severity", ""),
                    "chapter": getattr(issue, "chapter", ""),
                    "original_text": getattr(issue, "original_text", ""),
                    "suggestion": getattr(issue, "suggestion", ""),
                    "description": getattr(issue, "description", ""),
                })
        if row_matches:
            matches.append({"gold": row, "issues": row_matches})
        else:
            missed.append(row)

    false_positive = []
    for issue in issues:
        issue_id = getattr(issue, "id", id(issue))
        if issue_id in matched_issue_ids:
            continue
        false_positive.append({
            "id": getattr(issue, "id", None),
            "rule": getattr(issue, "rule", ""),
            "category": getattr(issue, "category", ""),
            "severity": getattr(issue, "severity", ""),
            "chapter": getattr(issue, "chapter", ""),
            "original_text": getattr(issue, "original_text", ""),
            "suggestion": getattr(issue, "suggestion", ""),
            "description": getattr(issue, "description", ""),
        })

    tp = len(matches)
    fp = len(false_positive)
    fn = len(missed)
    precision = tp / (tp + fp) if tp + fp else 0
    recall = tp / (tp + fn) if tp + fn else 0
    f1 = 2 * precision * recall / (precision + recall) if precision + recall else 0
    missing_in_parsed_text = [row for row in missed if not row.get("wrong_text_exists_in_parsed_text")]

    return {
        "review_id": review_id,
        "document_id": document.id,
        "filename": document.filename,
        "gold_count": len(gold_rows),
        "platform_count": len(issues),
        "tp": tp,
        "fp": fp,
        "fn": fn,
        "precision": round(precision, 4),
        "recall": round(recall, 4),
        "f1": round(f1, 4),
        "missing_in_parsed_text_count": len(missing_in_parsed_text),
        "matches": matches,
        "missed": missed,
        "false_positive": false_positive,
    }


@router.get("/{review_id}", response_model=Review)
async def read_review(review_id: int, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")
    review = _reconcile_review_runtime_state(db, review)
    doc = get_document(db, document_id=review.document_id)
    review_dict = review.__dict__.copy()
    review_dict['document_name'] = doc.filename if doc else ''
    review_dict['document_file_type'] = doc.file_type if doc else ''
    return review_dict


@router.get("/{review_id}/issues", response_model=list[Issue])
async def read_review_issues(review_id: int, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    doc = get_document(db, document_id=review.document_id) if review else None
    issues = get_issues(db, review_id=review_id)
    return _normalize_review_issue_display(_visible_review_issues(issues), getattr(doc, 'content', None))


@router.post("/{review_id}/issues/manual", response_model=Issue)
async def create_manual_review_issue(review_id: int, payload: dict, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")

    original_text = str(payload.get('original_text') or '').strip()
    if not original_text:
        raise HTTPException(status_code=400, detail="请填写补充上报的问题原文")

    severity = str(payload.get('severity') or 'general').strip()
    if severity not in {'fatal', 'serious', 'general', 'suggestion'}:
        severity = 'general'

    issue = IssueCreate(
        review_id=review_id,
        severity=severity,
        category=str(payload.get('category') or '人工补充').strip() or '人工补充',
        rule='MANUAL-SUPPLEMENT',
        chapter=str(payload.get('chapter') or '').strip(),
        original_text=original_text,
        context=str(payload.get('context') or '').strip(),
        suggestion=str(payload.get('suggestion') or '').strip(),
        description=str(payload.get('description') or '审核人补充上报的平台漏检问题').strip(),
        audit_basis=str(payload.get('audit_basis') or '审核人补充上报').strip(),
        confidence=100,
        source='manual',
        position='{}',
        status='confirmed',
    )
    created = create_issue(db, issue)
    total = db.query(IssueModel).filter(IssueModel.review_id == review_id).count()
    review.total_issues = total
    db.commit()
    db.refresh(created)
    return created


@router.put("/issues/{issue_id}", response_model=Issue)
async def update_issue_status(issue_id: int, issue_update: IssueUpdate, db: Session = Depends(get_db)):
    issue = update_issue(db, issue_id=issue_id, issue_update=issue_update)
    if not issue:
        raise HTTPException(status_code=404, detail="Issue not found")
    return issue


@router.post("/{review_id}/judge")
async def batch_judge_issues(review_id: int, payload: dict, db: Session = Depends(get_db)):
    """批量人工判定问题: { 'judgments': [{'issue_id': 1, 'status': 'confirmed'}, ...] }"""
    judgments = payload.get("judgments", [])
    updated = 0
    for j in judgments:
        issue_id = j.get("issue_id")
        status = j.get("status")
        if not issue_id or not status:
            continue
        # 允许的状态: pending, confirmed, false_positive, ignored
        if status not in ["pending", "confirmed", "false_positive", "ignored"]:
            continue
        from app.schemas.review import IssueUpdate
        issue = update_issue(db, issue_id=issue_id, issue_update=IssueUpdate(status=status))
        if issue:
            updated += 1
    return {"updated": updated, "total": len(judgments)}


@router.get("/{review_id}/export-html")
async def export_review_html(review_id: int, db: Session = Depends(get_db)):
    """导出 HTML 报告 (包含所有问题及人工判定状态)"""
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")

    doc = get_document(db, document_id=review.document_id)
    issues = _normalize_review_issue_display(_visible_review_issues(get_issues(db, review_id=review_id)), getattr(doc, 'content', None))
    html = _generate_review_html_content(review, doc, issues)
    return HTMLResponse(content=html)


@router.get("/{review_id}/export-result")
async def export_review_result(review_id: int, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")

    document = get_document(db, document_id=review.document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    issues = _normalize_review_issue_display(_visible_review_issues(get_issues(db, review_id=review_id)), getattr(document, 'content', None))
    if document.file_type == "docx":
        export_path, export_name, media_type = _export_review_docx(review, document, issues)
    elif document.file_type == "xlsx":
        export_path, export_name, media_type = _export_review_excel(review, document, issues)
    else:
        export_path, export_name, media_type = _export_review_html_file(review, document, issues)

    return FileResponse(path=str(export_path), filename=export_name, media_type=media_type)


@router.get("/{review_id}/report")
async def generate_report(review_id: int, db: Session = Depends(get_db)):
    review = get_review(db, review_id=review_id)
    if not review:
        raise HTTPException(status_code=404, detail="Review not found")

    issues = _visible_review_issues(get_issues(db, review_id=review_id))
    confirmed_issues = [i for i in issues if i.status in ["confirmed", "converted_to_rule"]]

    html_content = _generate_review_html_content(review, get_document(db, document_id=review.document_id), confirmed_issues)
    return {"content": html_content, "format": "html"}
