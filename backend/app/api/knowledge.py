from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form
from fastapi.responses import FileResponse, JSONResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import Optional
import os
import mimetypes
import docx2txt
import csv
import tempfile
from app.database import get_db
from app.crud.knowledge import (
    get_folder, get_folder_tree, get_folder_files, get_subfolders, create_folder, update_folder, delete_folder,
    create_file, get_file, delete_file, move_folder, move_file
)
from app.schemas.knowledge import FolderCreate, FolderUpdate, FolderMove, FileMove, FileContentRequest, FilePermissionRequest
from app.models.knowledge import KnowledgeFile
from app.api.auth import get_current_user, get_default_user, require_admin
from app.utils.file_utils import read_file_safe

PERMISSION_READ = "read"
PERMISSION_EDIT = "edit"
PERMISSION_DOWNLOAD = "download"
VALID_PERMISSIONS = {PERMISSION_READ, PERMISSION_EDIT, PERMISSION_DOWNLOAD}
VALID_EDIT_SCOPES = {"admin", "owner", "all"}

def check_download_permission(file):
    if file.permission == PERMISSION_READ:
        raise HTTPException(status_code=403, detail="该文件仅限只读，不支持下载")

def check_edit_permission(file, user):
    if file.permission != PERMISSION_EDIT:
        raise HTTPException(status_code=403, detail="该文件不支持编辑操作")
    scope = file.edit_scope or "all"
    if scope == "admin" and user.role != "admin":
        raise HTTPException(status_code=403, detail="该文件仅限管理员编辑")
    if scope == "owner" and user.role != "admin" and user.id != file.created_by:
        raise HTTPException(status_code=403, detail="该文件仅限上传者和管理员编辑")

router = APIRouter()

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "knowledge")


def _format_tabular_preview(rows, max_rows: int = 20, max_columns: int = 12) -> str:
    limited_rows = rows[:max_rows]
    normalized_rows = []
    for row in limited_rows:
        current_row = []
        for cell in list(row)[:max_columns]:
            current_row.append("" if cell is None else str(cell).strip())
        normalized_rows.append(current_row)

    if not normalized_rows:
        return "(表格内容为空)"

    output = []
    for row in normalized_rows:
        output.append("\t".join(row))

    if len(rows) > max_rows:
        output.append(f"... 已截断，仅预览前 {max_rows} 行")

    return "\n".join(output)


def _preview_excel_file(file_path: str) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(file_path, read_only=True, data_only=True)
    try:
        worksheet = workbook.active
        rows = list(worksheet.iter_rows(values_only=True))
        sheet_name = worksheet.title or "Sheet1"
        body = _format_tabular_preview(rows)
        return f"工作表: {sheet_name}\n\n{body}"
    finally:
        workbook.close()


def _preview_csv_file(file_path: str) -> str:
    content = read_file_safe(file_path)
    reader = csv.reader(content.splitlines())
    rows = list(reader)
    return _format_tabular_preview(rows)


def _collect_descendant_ids(folder):
    ids = set()
    for child in folder.children:
        ids.add(child.id)
        ids.update(_collect_descendant_ids(child))
    return ids

@router.get("/tree")
async def get_knowledge_tree(db: Session = Depends(get_db)):
    tree = get_folder_tree(db, None)
    return tree

@router.get("/folders/{folder_id}")
async def get_folder_content(folder_id: int, db: Session = Depends(get_db)):
    folder = get_folder(db, folder_id)
    if not folder:
        raise HTTPException(status_code=404, detail="文件夹不存在")
    
    files = get_folder_files(db, folder_id)
    subfolders = get_subfolders(db, folder_id)
    return {
        "folder": {
            "id": folder.id,
            "name": folder.name,
            "parent_id": folder.parent_id,
            "created_at": folder.created_at.isoformat() if folder.created_at else None
        },
        "subfolders": subfolders,
        "files": files
    }

@router.post("/folders", response_model=dict)
async def create_new_folder(
    folder: FolderCreate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    db_folder = create_folder(db, folder, user.id)
    return {"message": "文件夹创建成功", "id": db_folder.id}

@router.put("/folders/{folder_id}")
async def update_folder_name(
    folder_id: int,
    folder_update: FolderUpdate,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    if not folder_update.name:
        raise HTTPException(status_code=400, detail="文件夹名称不能为空")
    
    folder = update_folder(db, folder_id, folder_update.name)
    if not folder:
        raise HTTPException(status_code=404, detail="文件夹不存在")
    
    return {"message": "文件夹重命名成功"}


@router.put("/folders/{folder_id}/move")
async def move_folder_to_target(
    folder_id: int,
    folder_move: FolderMove,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    folder = get_folder(db, folder_id)
    if not folder:
        raise HTTPException(status_code=404, detail="文件夹不存在")

    target_parent_id = folder_move.parent_id
    if target_parent_id == folder.id:
        raise HTTPException(status_code=400, detail="不能移动到当前文件夹自身")

    if target_parent_id is not None:
        target_folder = get_folder(db, target_parent_id)
        if not target_folder:
            raise HTTPException(status_code=404, detail="目标文件夹不存在")

        descendant_ids = _collect_descendant_ids(folder)
        if target_parent_id in descendant_ids:
            raise HTTPException(status_code=400, detail="不能移动到当前文件夹的子文件夹中")

    move_folder(db, folder_id, target_parent_id)
    return {"message": "文件夹移动完成", "id": folder_id, "parent_id": target_parent_id}

@router.delete("/folders/{folder_id}")
async def delete_folder_by_id(
    folder_id: int,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除文件夹和文件")
    
    folder = get_folder(db, folder_id)
    if not folder:
        raise HTTPException(status_code=404, detail="文件夹不存在")
    
    for file in folder.files:
        if os.path.exists(file.file_path):
            os.remove(file.file_path)
    
    delete_folder(db, folder_id)
    return {"message": "文件夹删除成功"}

@router.post("/folders/{folder_id}/files/upload")
async def upload_file_to_folder(
    folder_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    if not os.path.exists(UPLOAD_DIR):
        os.makedirs(UPLOAD_DIR)
    
    folder = get_folder(db, folder_id) if folder_id else None
    if folder_id and not folder:
        raise HTTPException(status_code=404, detail="文件夹不存在")
    
    file_extension = os.path.splitext(file.filename)[1] if file.filename else ""
    import uuid
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    content = await file.read()
    with open(file_path, "wb") as buffer:
        buffer.write(content)
    
    file_size = os.path.getsize(file_path)
    file_type = file_extension[1:] if file_extension else "unknown"
    
    from app.schemas.knowledge import FileCreate
    file_data = FileCreate(
        name=file.filename or "unknown",
        folder_id=folder_id if folder_id else None
    )
    
    db_file = create_file(
        db=db,
        file=file_data,
        file_path=file_path,
        filename=file.filename or "unknown",
        file_size=file_size,
        file_type=file_type,
        created_by=user.id
    )
    
    return {"message": "文件上传成功", "id": db_file.id}

@router.get("/files/{file_id}")
async def get_file_info(file_id: int, db: Session = Depends(get_db)):
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return {
        "id": file.id,
        "name": file.name,
        "filename": file.filename,
        "file_path": file.file_path,
        "file_size": file.file_size,
        "file_type": file.file_type,
        "folder_id": file.folder_id,
        "permission": file.permission or "edit",
        "edit_scope": file.edit_scope or "all",
        "created_at": file.created_at.isoformat() if file.created_at else None
    }


@router.put("/files/{file_id}/move")
async def move_file_to_target(
    file_id: int,
    file_move: FileMove,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")

    target_folder = get_folder(db, file_move.folder_id)
    if not target_folder:
        raise HTTPException(status_code=404, detail="目标文件夹不存在")

    move_file(db, file_id, file_move.folder_id)
    return {"message": "文件移动完成", "id": file_id, "folder_id": file_move.folder_id}

@router.get("/files/{file_id}/download")
async def download_file(file_id: int, db: Session = Depends(get_db)):
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    check_download_permission(file)
    
    if not os.path.exists(file.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")
    
    return FileResponse(
        path=file.file_path,
        filename=file.filename,
        media_type="application/octet-stream"
    )

@router.get("/files/{file_id}/preview")
async def preview_file(file_id: int, db: Session = Depends(get_db)):
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if not os.path.exists(file.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")
    
    file_type = file.file_type.lower()
    
    # Text-based files - return raw content for syntax highlighting
    if file_type in ["txt", "md", "markdown", "json", "xml", "html", "css", "js", "py", "java", "c", "cpp", "h", "go", "rs"]:
        try:
            content = read_file_safe(file.file_path)
            return {"content": content, "type": "text", "file_name": file.filename}
        except Exception as e:
            return {"content": f"文件解析失败: {str(e)}", "type": "error", "file_name": file.filename}
    
    # Images - return direct URL
    elif file_type in ["jpg", "jpeg", "png", "gif", "bmp", "svg", "webp"]:
        return {"file_path": f"/api/knowledge/files/{file_id}/raw", "type": "image", "file_name": file.filename}
    
    # PDF - extract text and return as plain text
    elif file_type == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(file.file_path)
            content = "\n".join([page.extract_text() or "" for page in reader.pages])
            if not content.strip():
                content = "(PDF 文件内容为空或无法提取文字)"
            return {"content": content, "type": "text", "file_name": file.filename}
        except Exception as e:
            return {"content": f"PDF 解析失败: {str(e)}", "type": "error", "file_name": file.filename}
    
    # DOCX - extract text and return as plain text
    elif file_type == "docx":
        try:
            content = docx2txt.process(file.file_path)
            if not content or not content.strip():
                content = "(DOCX 文件内容为空或无法提取文字)"
            return {"content": content, "type": "text", "file_name": file.filename}
        except Exception as e:
            return {"content": f"DOCX 解析失败: {str(e)}", "type": "error", "file_name": file.filename}

    # Excel / CSV - extract first sheet or rows as plain text table preview
    elif file_type in ["xlsx", "xlsm", "xltx", "xltm"]:
        try:
            content = _preview_excel_file(file.file_path)
            return {"content": content, "type": "text", "file_name": file.filename}
        except Exception as e:
            return {"content": f"Excel 解析失败: {str(e)}", "type": "error", "file_name": file.filename}

    elif file_type == "csv":
        try:
            content = _preview_csv_file(file.file_path)
            return {"content": content, "type": "text", "file_name": file.filename}
        except Exception as e:
            return {"content": f"CSV 解析失败: {str(e)}", "type": "error", "file_name": file.filename}
    
    else:
        return {"content": "此文件类型不支持在线预览，请下载后查看", "type": "unsupported", "file_name": file.filename}

@router.get("/files/{file_id}/raw")
async def get_raw_file(file_id: int, db: Session = Depends(get_db)):
    """Return raw file for direct browser preview"""
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if not os.path.exists(file.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")
    
    media_type = mimetypes.guess_type(file.file_path)[0] or "application/octet-stream"
    
    return FileResponse(
        path=file.file_path,
        filename=file.filename,
        media_type=media_type
    )

class FileContentRequest(BaseModel):
    content: str

@router.get("/files/{file_id}/content")
async def get_file_content(file_id: int, db: Session = Depends(get_db)):
    """获取文件可编辑的文本内容"""
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    if not os.path.exists(file.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")

    ext = os.path.splitext(file.file_path)[1].lower()
    try:
        if ext == '.md':
            content = read_file_safe(file.file_path)
        elif ext == '.docx':
            from docx import Document
            doc = Document(file.file_path)
            content = '\n'.join(p.text for p in doc.paragraphs)
        elif ext == '.xlsx':
            from openpyxl import load_workbook
            wb = load_workbook(file.file_path, data_only=True)
            ws = wb.active
            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append('\t'.join(str(c) if c is not None else '' for c in row))
            content = '\n'.join(rows)
            wb.close()
        else:
            raise HTTPException(status_code=400, detail=f"不支持编辑的文件类型: {ext}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取文件失败: {str(e)}")

    return {"filename": file.filename, "content": content, "type": ext}

@router.put("/files/{file_id}/content")
async def update_file_content(
    file_id: int,
    req: FileContentRequest,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """保存编辑后的文件内容"""
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    check_edit_permission(file, current_user)
    if not os.path.exists(file.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")

    ext = os.path.splitext(file.file_path)[1].lower()
    try:
        if ext == '.md':
            with open(file.file_path, 'w', encoding='utf-8') as f:
                f.write(req.content)
        elif ext == '.docx':
            from docx import Document
            doc = Document(file.file_path)
            # 清空现有段落
            for p in doc.paragraphs:
                p.clear()
            # 写入新内容（每行一个段落）
            lines = req.content.split('\n')
            # 第一个段落复用，其余追加
            for i, line in enumerate(lines):
                if i == 0:
                    doc.paragraphs[0].text = line
                else:
                    doc.add_paragraph(line)
            # 删除多余的原始段落
            while len(doc.paragraphs) > len(lines):
                p = doc.paragraphs[-1]._element
                p.getparent().remove(p)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                doc.save(tmp.name)
            os.replace(tmp.name, file.file_path)
        elif ext == '.xlsx':
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            for row in req.content.split('\n'):
                cells = row.split('\t')
                ws.append(cells)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
                wb.save(tmp.name)
            wb.close()
            os.replace(tmp.name, file.file_path)
        else:
            raise HTTPException(status_code=400, detail=f"不支持编辑的文件类型: {ext}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"保存文件失败: {str(e)}")

    # 更新文件大小
    new_size = os.path.getsize(file.file_path) if os.path.exists(file.file_path) else 0
    db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).update({"file_size": new_size})
    db.commit()

    return {"message": "保存成功", "file_size": new_size}

@router.put("/files/{file_id}/permission")
async def update_file_permission(
    file_id: int,
    req: FilePermissionRequest,
    db: Session = Depends(get_db),
    current_user = Depends(require_admin)
):
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    if req.permission is not None and req.permission not in VALID_PERMISSIONS:
        raise HTTPException(status_code=400, detail=f"无效的权限值，合法值：{', '.join(VALID_PERMISSIONS)}")
    update_data = {}
    if req.permission is not None:
        update_data["permission"] = req.permission
    if req.edit_scope is not None:
        if req.edit_scope not in VALID_EDIT_SCOPES:
            raise HTTPException(status_code=400, detail=f"无效的操作者范围，合法值：{', '.join(VALID_EDIT_SCOPES)}")
        update_data["edit_scope"] = req.edit_scope
    if update_data:
        db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).update(update_data)
        db.commit()
    return {"message": "权限更新成功", "permission": file.permission, "edit_scope": file.edit_scope}

@router.delete("/files/{file_id}")
async def delete_file_by_id(
    file_id: int,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    file = get_file(db, file_id)
    if not file:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if user.role != "admin" and file.created_by != user.id:
        raise HTTPException(status_code=403, detail="仅管理员或创建者可删除文件")
    
    if os.path.exists(file.file_path):
        os.remove(file.file_path)
    
    delete_file(db, file_id)
    return {"message": "文件删除成功"}


@router.post("/export-seed")
async def export_knowledge_to_seed(
    current_user = Depends(get_current_user)
):
    """将知识库导出到种子目录，用于 Git 提交后团队共享"""
    from seed.knowledge_seed import export_kb_to_seed
    try:
        export_kb_to_seed()
        return {"message": "知识库已导出到 seed/knowledge/ 目录，请执行 git add/commit/push 分享给团队"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")
