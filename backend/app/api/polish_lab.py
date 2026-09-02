"""智能润色 AI 调试副本。

本模块从原始 `app.api.polish` 复制而来，供 AI 增强调试使用。
修改本文件不会覆盖正式润色模块；数据、任务进度和导出文件均隔离存储。
"""

from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from typing import List, Optional
from datetime import timedelta
from collections import defaultdict
from difflib import SequenceMatcher
from html import escape as html_escape
import json
import asyncio
import os
import uuid
import mimetypes
import re
import threading
import datetime
import logging
import hashlib
import contextvars
import time
from app.database import get_db
from app.crud.document import get_document
from app.crud.polished_document_lab import (
    get_polished_documents, get_polished_document, create_polished_document, delete_polished_document
)
from app.api.auth import get_current_user, get_current_user_optional, get_default_user, require_admin
from app.models.knowledge import KnowledgeFile, Folder
from app.models.term import Term
from app.models.polish_feedback_lab import PolishFeedbackLab as PolishFeedback
from app.models.polished_document_lab import PolishedDocumentLab as PolishedDocument
from app.models.cat_analysis_session_lab import CatAnalysisSessionLab as CatAnalysisSession
from app.models.cat_decision_record_lab import CatDecisionRecordLab as CatDecisionRecord
from app.utils.file_utils import read_file_safe as _read_file_safe
from app.utils.polish_rules_engine import apply_all_rules, apply_custom_rules
from app.crud.polish_learning_rule import get_enabled_engine_keys, get_enabled_custom_rules, get_enabled_typo_rules, record_rule_triggers
from app.crud.term import bulk_create_terms
from app.utils.typo_dictionary import get_default_typo_dict

router = APIRouter()
logger = logging.getLogger(__name__)
POLISH_AI_ONLY = False
_ai_template_rerank_enabled = contextvars.ContextVar("ai_template_rerank_enabled", default=True)

# 润色任务进度追踪
_polish_tasks: dict = {}  # {task_id: {"status", "progress", "message", "result"}}
_polish_tasks_lock = threading.Lock()


def _persist_lab_diagnoses(db: Session, diagnoses, sentence_items=None, analyze_id=None, source="text", source_name=None):
    if not db or not diagnoses:
        return
    try:
        from app.models.cat_diagnose_record_lab import CatDiagnoseRecordLab

        sentence_by_index = {
            item.get("sentence_index"): item
            for item in (sentence_items or [])
            if isinstance(item, dict)
        }
        rows = []
        stored_source_name = str(source_name or "").strip()
        for diag in diagnoses:
            if not isinstance(diag, dict):
                continue
            sent = sentence_by_index.get(diag.get("sentence_index")) or {}
            original = str(
                sent.get("source_sentence_text")
                or sent.get("text")
                or sent.get("original_text")
                or diag.get("quote")
                or ""
            ).strip()
            rows.append(
                CatDiagnoseRecordLab(
                    analyze_id=analyze_id,
                    source=source,
                    source_name=stored_source_name or None,
                    sentence_index=diag.get("sentence_index"),
                    original_text=original,
                    quote=str(diag.get("quote") or ""),
                    category=str(diag.get("category") or "other"),
                    severity=str(diag.get("severity") or "low"),
                    problem=str(diag.get("problem") or ""),
                    revised=str(diag.get("revised") or ""),
                    rationale=str(diag.get("rationale") or ""),
                    ruleable=bool(diag.get("ruleable")) or bool(str(diag.get("revised") or "").strip()),
                    rule_hint=str(diag.get("rule_hint") or diag.get("quote") or ""),
                    status="pending",
                )
            )
        if not rows:
            return
        db.add_all(rows)
        db.commit()
        logger.info(
            "[CAT_DIAGNOSE] persisted count=%s analyze_id=%s source=%s",
            len(rows),
            analyze_id or "",
            source,
        )
    except Exception as exc:
        try:
            db.rollback()
        except Exception:
            pass
        logger.warning("[CAT_DIAGNOSE] 诊断落库失败: %s", exc)


# ============================================================
# 术语库加载 & 语言检测
# ============================================================

def _load_terms_from_db(db: Session) -> dict:
    """从术语库表中加载所有术语，返回 {非标准用语: 标准用语} 映射（带缓存）。"""
    if '__db_terms__' in _term_cache:
        return dict(_term_cache['__db_terms__'])

    terms = db.query(Term).all()
    term_dict = {}
    for t in terms:
        if t.non_standard and t.standard and t.non_standard.strip() != t.standard.strip():
            term_dict[t.non_standard.strip()] = t.standard.strip()
    _term_cache['__db_terms__'] = dict(term_dict)
    return term_dict


def _load_typo_dict(db: Session) -> dict:
    cache_key = '__db_typo_rules__'
    if cache_key in _typo_cache:
        return dict(_typo_cache[cache_key])

    typo_dict = {}
    try:
        rules = get_enabled_typo_rules(db)
        for rule in rules:
            wrong = str(rule.match_pattern or '').strip()
            correct = str(rule.replacement_text or '').strip()
            if wrong and correct and wrong != correct:
                typo_dict[wrong] = correct
    except Exception:
        logger.exception("加载错别字规则失败")
        typo_dict = {}

    if not typo_dict:
        typo_dict = get_default_typo_dict()

    _typo_cache[cache_key] = dict(typo_dict)
    return typo_dict


def _invalidate_typo_cache() -> None:
    _typo_cache.clear()


def _detect_language(text: str) -> str:
    """检测文本语言。返回 'zh' (中文) 或 'en' (英文)。"""
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    total_chars = len(text.replace(' ', '').replace('\n', ''))
    if total_chars == 0:
        return 'zh'
    ratio = chinese_chars / total_chars
    return 'zh' if ratio > 0.3 else 'en'


def _protect_model_numbers(text: str) -> str:
    """保持产品型号内部连写，同时保留编号与术语之间的空格。"""
    if not text:
        return text

    text = re.sub(r'(?<=[A-Za-z-])\s+(?=\d+[A-Za-z])', '', text)
    text = re.sub(r'(?<=[A-Za-z-]\d)\s+(?=[A-Za-z]{2,}\b)', '', text)
    text = re.sub(r'(?<=(?:表|图)\d)\s*(?=[A-Za-z]{2,})', ' ', text)
    text = re.sub(r'(?<=\d\.\d)\s*(?=[A-Za-z]{2,})', ' ', text)
    return text


_MARKDOWN_INLINE_ESCAPE_RE = re.compile(r'\\([\\`*_{}\[\]()#+\-.!|~])')


def _unescape_markdown_inline(text: str) -> str:
    """还原句式库 Markdown 转义，避免模板把 \\. \\- 带进候选。"""
    value = str(text or '')
    if '\\' not in value:
        return value
    return _MARKDOWN_INLINE_ESCAPE_RE.sub(r'\1', value)


def _use_ai_only() -> bool:
    return POLISH_AI_ONLY


def _term_column_lang(header: str) -> str:
    """根据列表头判断该列的语言倾向：zh / en / None（中性）。"""
    zh_keywords = ['中', '中文', 'zh', '汉语', '汉']
    en_keywords = ['英', '英文', 'en', 'english', 'eng']
    for kw in zh_keywords:
        if kw in header.lower():
            return 'zh'
    for kw in en_keywords:
        if kw in header.lower():
            return 'en'
    return None


def _parse_terminology(terminology_input: str) -> dict:
    """统一术语解析入口：自动识别 Markdown 文本或 Excel 文件路径。"""
    if not terminology_input:
        return {}
    # 如果是 .xlsx 文件路径
    if terminology_input.lower().endswith('.xlsx'):
        return _parse_terminology_xlsx(terminology_input)
    # 否则作为 Markdown 文本解析
    return _parse_terminology_md(terminology_input)


def _parse_terminology_md(md_content: str) -> dict:
    """解析术语库 Markdown 文件，返回 {非标准: 标准} 映射。

    支持的列格式：
    - 简单格式：| 非标准 | 标准 |
    - 分语言格式：| 非标准(中) | 标准(中) | 非标准(英) | 标准(英) |
    - 带语言列：| 非标准 | 标准 | 语言 |
    
    兼容全角竖线 ｜ 和半角竖线 |。
    """
    term_dict = {}
    if not md_content:
        return term_dict

    # 兼容全角竖线和全角横线
    content = md_content
    # 分隔符行全角转半角：｜---｜ → |---|
    content = content.replace('\uff5c', '|')  # fullwidth vertical bar
    content = content.replace('\u2502', '|')  # box drawing light vertical

    lines = content.split('\n')
    header = ''
    col_langs = []
    has_lang_col = False
    lang_col_idx = -1

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 捕获表头行（第一个含 | 且不含分隔符的行）
        if not header and '|' in stripped and '---' not in stripped:
            header = stripped
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            # 检查是否有"语言"列
            for idx, cell in enumerate(cells):
                if any(kw in cell.lower() for kw in ['语言', 'lang', 'language']):
                    has_lang_col = True
                    lang_col_idx = idx
                col_langs.append(_term_column_lang(cell))
            continue

        if '|' not in stripped:
            continue
        if stripped.startswith('#'):
            continue
        if stripped.startswith('|---') or stripped.startswith('| :--') or stripped.startswith('|:--'):
            continue
        if stripped.startswith('|序号') or stripped.startswith('| 序号'):
            continue

        cells = [c.strip() for c in stripped.split('|') if c.strip()]
        clean_cells = [c for c in cells if c.strip() and c.strip() != '---']
        if len(clean_cells) < 2:
            continue

        # 模式 A：带语言列
        if has_lang_col and lang_col_idx >= 0 and lang_col_idx < len(clean_cells):
            lang_val = clean_cells[lang_col_idx].lower()
            is_zh = any(kw in lang_val for kw in ['zh', '中', 'cn', 'chinese'])
            is_en = any(kw in lang_val for kw in ['en', '英', 'english', 'eng'])
            # 构建不包含语言列的数据列
            data_cells = [c for i, c in enumerate(clean_cells) if i != lang_col_idx]
            # 数据列两两配对
            for i in range(0, len(data_cells) - 1, 2):
                old_term = data_cells[i].strip().strip('!')
                new_term = data_cells[i + 1].strip().strip('!')
                if old_term and new_term and old_term != new_term and len(old_term) > 1:
                    # 根据语言列值分配语言标记
                    lang_suffix = ''
                    if is_zh:
                        lang_suffix = '##zh'
                    elif is_en:
                        lang_suffix = '##en'
                    key = f"{old_term}{lang_suffix}" if lang_suffix else old_term
                    if key not in term_dict:
                        term_dict[key] = new_term
            continue

        # 模式 B：无语言列，但有表头指示列语言
        if col_langs and len(col_langs) == len(clean_cells):
            for i in range(0, len(clean_cells) - 1, 2):
                old_term = clean_cells[i].strip().strip('!')
                new_term = clean_cells[i + 1].strip().strip('!')
                if old_term and new_term and old_term != new_term and len(old_term) > 1:
                    lang = col_langs[i] or col_langs[i + 1] or ''
                    lang_suffix = f'##{lang}' if lang else ''
                    key = f"{old_term}{lang_suffix}" if lang_suffix else old_term
                    if key not in term_dict:
                        term_dict[key] = new_term
            continue

        # 模式 C：无语言信息，简单两列配对
        for i in range(0, len(clean_cells) - 1, 2):
            old_term = clean_cells[i].strip().strip('!')
            new_term = clean_cells[i + 1].strip().strip('!')
            if old_term and new_term and old_term != new_term and len(old_term) > 1:
                if old_term not in term_dict:
                    term_dict[old_term] = new_term

    return term_dict


def _parse_terminology_xlsx(file_path: str) -> dict:
    """解析 Excel (.xlsx) 术语文件，返回 {非标准: 标准} 映射。
    
    支持两种格式：
    - 替换表：| 非标准 | 标准 |   → 直接作为 old→new 映射
    - 双语表：| zh-CN | en-US | → 仅提取中文列作为标准术语（不做替换，避免中文→英文错乱）
    """
    term_dict = {}
    try:
        import openpyxl
    except ImportError:
        return term_dict
    
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True)
        ws = wb.active
        
        rows = list(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 500), values_only=True))
        if not rows:
            wb.close()
            return term_dict
        
        headers = [str(h).strip().lower() if h else '' for h in rows[0]]
        is_replacement = any(kw in h for h in headers for kw in ['非标准', '旧', 'old', '非标', 'source'])
        is_bilingual = any(kw in h for h in headers for kw in ['zh', 'cn', '中文', 'en', '英', 'us'])
        
        for row in rows[1:]:
            cells = [str(c).strip() if c else '' for c in row]
            cells = [c for c in cells if c]
            if len(cells) < 2:
                continue
            
            if is_replacement:
                # 替换表：col1=非标准, col2=标准
                old_term = cells[0]
                new_term = cells[1]
                if old_term and new_term and old_term != new_term and len(old_term) > 0:
                    if old_term not in term_dict:
                        term_dict[old_term] = new_term
            elif is_bilingual:
                # 双语表：col1=中文标准术语, col2=英文 —— 仅提取中文列，不做替换
                # 将中文标准术语自身作为 key（标识已知标准术语，供 AI 参考）
                std_cn = cells[0]
                if std_cn and len(std_cn) > 0:
                    term_dict[f"__std__{std_cn}"] = std_cn
            else:
                # 未知格式：假设 col1→col2 替换
                old_term = cells[0]
                new_term = cells[1]
                if old_term and new_term and old_term != new_term and len(old_term) > 0:
                    if old_term not in term_dict:
                        term_dict[old_term] = new_term
        
        wb.close()
        if is_bilingual and not is_replacement:
            std_terms = [v for k, v in term_dict.items() if k.startswith('__std__')]
            print(f"[TERM] Excel 双语对照表: 标准中文术语 {len(std_terms)} 条 ({', '.join(std_terms[:5])})")
    except Exception as e:
        print(f"[TERM] Excel 解析失败: {e}")
    
    return term_dict
    term_dict = {}
    if not md_content:
        return term_dict

    # 统一处理全角竖线
    md_content = md_content.replace('\uFF5C', '|')
    lines = md_content.split('\n')
    header = ''
    col_langs = []
    has_lang_col = False
    lang_col_idx = -1

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 捕获表头行（第一个含 | 且不含分隔符的行）
        if not header and '|' in stripped and '---' not in stripped:
            header = stripped
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            # 检查是否有"语言"列
            for idx, cell in enumerate(cells):
                if any(kw in cell.lower() for kw in ['语言', 'lang', 'language']):
                    has_lang_col = True
                    lang_col_idx = idx
                col_langs.append(_term_column_lang(cell))
            continue

        if '|' not in stripped:
            continue
        if stripped.startswith('#'):
            continue
        if stripped.startswith('|---') or stripped.startswith('| :--') or stripped.startswith('|:--'):
            continue
        if stripped.startswith('|序号') or stripped.startswith('| 序号'):
            continue

        cells = [c.strip() for c in stripped.split('|') if c.strip()]
        clean_cells = [c for c in cells if c.strip() and c.strip() != '---']
        if len(clean_cells) < 2:
            continue

        # 模式 A：带语言列
        if has_lang_col and lang_col_idx >= 0 and lang_col_idx < len(clean_cells):
            lang_val = clean_cells[lang_col_idx].lower()
            is_zh = any(kw in lang_val for kw in ['zh', '中', 'cn', 'chinese'])
            is_en = any(kw in lang_val for kw in ['en', '英', 'english', 'eng'])
            # 构建不包含语言列的数据列
            data_cells = [c for i, c in enumerate(clean_cells) if i != lang_col_idx]
            # 数据列两两配对
            for i in range(0, len(data_cells) - 1, 2):
                old_term = data_cells[i].strip().strip('!')
                new_term = data_cells[i + 1].strip().strip('!')
                if old_term and new_term and old_term != new_term and len(old_term) > 1:
                    # 根据语言列值分配语言标记
                    lang_suffix = ''
                    if is_zh:
                        lang_suffix = '##zh'
                    elif is_en:
                        lang_suffix = '##en'
                    # 存入时带语言标记
                    key = f"{old_term}{lang_suffix}" if lang_suffix else old_term
                    if key not in term_dict:
                        term_dict[key] = new_term
            continue

        # 模式 B：无语言列，但有表头指示列语言
        if col_langs and len(col_langs) == len(clean_cells):
            # 两列配对，每对继承对应表头的语言
            for i in range(0, len(clean_cells) - 1, 2):
                old_term = clean_cells[i].strip().strip('!')
                new_term = clean_cells[i + 1].strip().strip('!')
                if old_term and new_term and old_term != new_term and len(old_term) > 1:
                    lang = col_langs[i] or col_langs[i + 1] or ''
                    lang_suffix = f'##{lang}' if lang else ''
                    key = f"{old_term}{lang_suffix}" if lang_suffix else old_term
                    if key not in term_dict:
                        term_dict[key] = new_term
            continue

        # 模式 C：无语言信息，简单两列配对
        for i in range(0, len(clean_cells) - 1, 2):
            old_term = clean_cells[i].strip().strip('!')
            new_term = clean_cells[i + 1].strip().strip('!')
            if old_term and new_term and old_term != new_term and len(old_term) > 1:
                if old_term not in term_dict:
                    term_dict[old_term] = new_term

    return term_dict


def _filter_terms_by_lang(term_dict: dict, target_lang: str) -> dict:
    """从带语言标记的术语字典中筛选出目标语言的术语。返回纯净的 {非标准: 标准}。"""
    filtered = {}
    for key, val in term_dict.items():
        # 跳过 Excel 双语表的标准术语标记（__std__ 前缀）
        if key.startswith('__std__'):
            continue
        if '##zh' in key:
            if target_lang == 'zh':
                clean_key = key.replace('##zh', '')
                filtered[clean_key] = val
        elif '##en' in key:
            if target_lang == 'en':
                clean_key = key.replace('##en', '')
                filtered[clean_key] = val
        else:
            # 无语言标记，通用术语，适用于所有语言
            filtered[key] = val
    return filtered


def _resolve_terminology(db: Session, terminology_md: str = None, text: str = None) -> dict:
    """加载术语：文件术语优先，自动按文本语言过滤。返回纯净 {非标准: 标准}。"""
    merged = {}

    platform_files = _get_platform_feedback_terminology_targets(db, 1)
    for platform_file in platform_files:
        if platform_file and platform_file.file_path and os.path.exists(platform_file.file_path):
            try:
                platform_terms = _parse_terminology(platform_file.file_path if platform_file.file_path.lower().endswith('.xlsx') else _read_file_safe(platform_file.file_path))
                if platform_terms:
                    if text:
                        lang = _detect_language(text)
                        merged.update(_filter_terms_by_lang(platform_terms, lang))
                    else:
                        merged.update(platform_terms)
            except Exception:
                pass

    if terminology_md:
        parsed = _parse_terminology(terminology_md)
        if parsed:
            if text:
                lang = _detect_language(text)
                merged.update(_filter_terms_by_lang(parsed, lang))
            else:
                merged.update(parsed)

    if merged:
        return merged
    return _load_terms_from_db(db)


_DOC_AI_MAX_TEXT_LEN = 5000
_DOC_AI_MAX_GUIDE_LEN = 9000
_DOC_AI_MAX_TERM_COUNT = 120


def _trim_terms_for_ai(term_dict: Optional[dict], text: str, max_items: int = _DOC_AI_MAX_TERM_COUNT) -> Optional[dict]:
    """仅保留与当前文本相关的部分术语，降低模型上下文体积。"""
    if not term_dict:
        return term_dict

    text = text or ""
    text_lower = text.lower()
    matched = []
    remaining = []

    for source, target in term_dict.items():
        source_text = str(source or "").strip()
        target_text = str(target or "").strip()
        if not source_text or not target_text:
            continue
        source_lower = source_text.lower()
        item = (source_text, target_text)
        if source_text in text or source_lower in text_lower:
            matched.append(item)
        else:
            remaining.append(item)

    selected = matched[:max_items]
    if len(selected) < max_items:
        selected.extend(remaining[:max_items - len(selected)])
    return dict(selected)


def _should_skip_document_ai(text: str, style_guide: Optional[str], terminology: Optional[dict]) -> tuple[bool, str]:
    """文档润色整篇送模前做体积保护，避免 8k 模型直接超限。"""
    text_len = len(text or "")
    guide_len = len(style_guide or "")
    term_count = len(terminology or {})

    if text_len > _DOC_AI_MAX_TEXT_LEN:
        return True, f"文档正文过长({text_len})"
    if guide_len > _DOC_AI_MAX_GUIDE_LEN:
        return True, f"润色规则过长({guide_len})"
    if term_count > _DOC_AI_MAX_TERM_COUNT:
        return True, f"术语条目过多({term_count})"
    return False, ""


def _load_terminology_source(db: Session, terminology_id: int = None) -> Optional[str]:
    """加载术语来源。Excel 返回文件路径，其它文本文件返回文件内容。"""
    if not terminology_id:
        return None

    term_file = db.query(KnowledgeFile).filter(KnowledgeFile.id == terminology_id).first()
    if not term_file or not term_file.file_path or not os.path.exists(term_file.file_path):
        return None

    if term_file.file_path.lower().endswith('.xlsx'):
        return term_file.file_path

    return _read_file_safe(term_file.file_path)


# 句式清单所在知识库文件夹 ID（写作规范 / 句式清单）
SENTENCE_GUIDE_FOLDER_IDS = [8]
SENTENCE_FEEDBACK_FOLDER_IDS = [10]
PLATFORM_FEEDBACK_FILENAME = "平台反馈的句式清单-AI调试.md"
DOCUMENT_PRODUCT_TYPES = (
    "建库试剂",
    "测序试剂",
    "核酸提取",
    "测序仪",
    "自动化",
    "软件",
    "超声",
)
TERMINOLOGY_FEEDBACK_FOLDER_IDS = [21]
PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME = "平台反馈的术语对照表-AI调试.md"
TERMINOLOGY_FOLDER_IDS = [19]
PLATFORM_FEEDBACK_SENTENCE_RELATIVE_PATH = os.path.join("..", "polish-lab", "feedback", PLATFORM_FEEDBACK_FILENAME)
PLATFORM_FEEDBACK_TERMINOLOGY_RELATIVE_PATH = os.path.join("..", "polish-lab", "feedback", PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME)
PRIMARY_SENTENCE_GUIDE_FILENAMES = {
    "DNBelab-D4RS_结构化句式库_操作步骤.md",
}
REFERENCE_SENTENCE_GUIDE_FILENAMES = {
    "句式表达参考手册_建库试剂说明书.md",
}
BUNDLED_SENTENCE_GUIDE_RELATIVE_PATHS = [
    os.path.join("static", "bundled", "structured_sentence_guide_d4rs_operations.md"),
]

# 默认写作风格指南文件 ID（写作规范 / 写作风格指南 / 中文技术文档写作风格指南）
# 内容已升级为 V2 完整规则体系（术语→句式→风格→微调四层 + 前置指令 + 禁用词）
DEFAULT_STYLE_GUIDE_ID = 1




def _load_file_content(db: Session, file_id: int) -> str:
    """加载单个知识库文件的内容"""
    kf = db.query(KnowledgeFile).filter(KnowledgeFile.id == file_id).first()
    if kf and kf.file_path and os.path.exists(kf.file_path):
        try:
            return _read_knowledge_text_content(kf.file_path).strip()
        except Exception:
            pass
    return None


def _read_knowledge_text_content(file_path: str) -> str:
    path = str(file_path or '').strip()
    if not path or not os.path.exists(path):
        return ''
    lower_path = path.lower()
    if lower_path.endswith('.docx'):
        try:
            from docx import Document
            doc = Document(path)
            return '\n'.join(p.text for p in _iter_docx_body_paragraphs(doc) if str(p.text or '').strip())
        except Exception:
            try:
                import docx2txt
                return str(docx2txt.process(path) or '')
            except Exception:
                return ''
    return str(_read_file_safe(path) or '')


def _build_document_polish_guide(
    db: Session,
    sentence_file_id: int = None,
    requirements: str = None,
    product_type: Optional[str] = None,
) -> str:
    """构建文档润色的完整规则指南。

    优先级: 句式匹配 > 术语匹配 > 风格指南 > 数据库规则
    句式文件在前，风格指南在后，AI 按顺序给予优先权重。
    """
    parts = []

    # 1. 句式清单（优先匹配）
    # 显式选择句式库时，仅加载该句式库及平台反馈，避免旧库与新库混用。
    if sentence_file_id:
        selected_guide = _load_selected_sentence_guide_with_feedback(db, sentence_file_id)
        if selected_guide:
            parts.append(f"{_CANDIDATE_RECALL_GUIDE_MARKER}\n\n{selected_guide}")
    else:
        all_guides = _load_sentence_guides(db, product_type=product_type)
        if all_guides:
            parts.append(f"{_CANDIDATE_RECALL_GUIDE_MARKER}\n\n{all_guides}")

    # 2. 用户额外的润色要求
    if requirements and requirements.strip():
        parts.append(f"## 额外润色要求\n\n{requirements.strip()}")

    # 3. 写作风格指南（统一使用 V2 完整规则）
    default_guide = _load_file_content(db, DEFAULT_STYLE_GUIDE_ID)
    if default_guide:
        parts.append(f"{_AI_STYLE_GUIDE_MARKER}\n\n{default_guide}")

    return "\n\n".join(parts) if parts else None


def _load_selected_sentence_guide_with_feedback(db: Session, sentence_file_id: Optional[int]) -> Optional[str]:
    if not sentence_file_id:
        return None
    return _load_sentence_guides(db, style_guide_id=sentence_file_id)


def _candidate_recall_guide_text(guide_text: str) -> str:
    value = str(guide_text or '')
    if not value:
        return ''
    if _AI_STYLE_GUIDE_MARKER in value:
        value = value.split(_AI_STYLE_GUIDE_MARKER, 1)[0]
    if _CANDIDATE_RECALL_GUIDE_MARKER in value:
        value = value.split(_CANDIDATE_RECALL_GUIDE_MARKER, 1)[1]
    return value.strip()


def _ai_style_guide_text(guide_text: str) -> str:
    value = str(guide_text or '')
    if not value:
        return ''
    if _AI_STYLE_GUIDE_MARKER in value:
        value = value.split(_AI_STYLE_GUIDE_MARKER, 1)[1]
    return value.strip()


def _document_requirements_text(guide_text: str) -> str:
    value = str(guide_text or '')
    if not value:
        return ''
    match = re.search(r'##\s+额外润色要求\s*(.+?)(?=\n##\s+|\Z)', value, flags=re.S)
    if not match:
        return ''
    return match.group(1).strip()


def _compact_document_ai_style_guide(guide_text: str) -> str:
    requirements = _document_requirements_text(guide_text)
    style_text = _ai_style_guide_text(guide_text)
    if not style_text:
        return requirements

    hard_rules = []
    front_section = re.search(r'##\s*前置指令.*?(?=\n##\s+|\Z)', style_text, flags=re.S)
    if front_section:
        for line in front_section.group(0).splitlines():
            stripped = line.strip()
            if re.match(r'^\d+\.', stripped):
                hard_rules.append(stripped)

    parts = []
    if requirements:
        parts.append(f"## 额外润色要求\n{requirements}")
    if hard_rules:
        parts.append("## 文档润色硬规则\n" + "\n".join(hard_rules))
    elif style_text:
        parts.append(style_text[:1200].strip())
    return "\n\n".join(part for part in parts if part).strip()


def _looks_like_title_or_noun_phrase(text: str) -> bool:
    value = str(text or '').strip()
    if not value:
        return False
    step_prefix, body = _split_step_prefix(value)
    if body:
        value = body
    normalized = _normalize_sentence_for_match(value)
    if len(value) <= 4:
        return True
    if step_prefix and len(value) <= 12 and len(normalized) <= 8 and not re.search(r'[，。；！？,;!?：:]', value):
        short_title_suffixes = ('取出', '准备', '加载', '安装', '检查', '设置', '操作', '运行', '启动', '吸取', '扫码', '录入', '取样', '拆封', '开机', '关机')
        if any(value.endswith(suffix) for suffix in short_title_suffixes):
            return True
    if value.endswith(('：', ':')):
        if len(normalized) > 10 and any(marker in value for marker in ('对应', '分别', '或者', '以及', '并将', '后点击', '点击', '输入', '选择', '设置', '确认')):
            return False
        if len(normalized) > 12 and _extract_sentence_intent(value).get('actions'):
            return False
        if len(normalized) > 12 and any(marker in value for marker in ('推荐', '使用', '需要', '请', '应', '须')):
            return False
        return True
    if re.match(r'^[\d一二三四五六七八九十]+[\.、\s]', value):
        return True
    verb_markers = [
        '将', '请', '按', '点击', '选择', '输入', '打开', '关闭',
        '启动', '停止', '设置', '检查', '确认', '安装', '连接',
        '使用', '进行', '可以', '需要', '应该', '必须', '确保',
        '按下', '旋转', '调节', '插入', '取出', '放置', '执行', '访问',
        '查看', '显示', '支持', '提供', '包含', '通过', '根据',
        '按照', '用于', '适用于', '分为',
    ]
    if _extract_sentence_intent(value).get('actions'):
        return False
    return not any(marker in value for marker in verb_markers)


def _build_polish_debug_info(
    sentence_file_id: Optional[int],
    sentence_file_name: Optional[str],
    sentence_guide: Optional[str],
    candidate_recall_guide: Optional[str],
    ai_style_guide: Optional[str],
    terminology_file_id: Optional[int],
    terminology_file_name: Optional[str],
    skip_ai: bool,
    skip_reason: str,
    ai_polished: str,
    pre_polished: str,
    total_changes: int,
    visible_changes: Optional[int] = None,
    previous_polish_found: bool = False,
    previous_new_change_count: int = 0,
):
    guide_text = sentence_guide or ""
    candidate_text = candidate_recall_guide or ""
    ai_text = ai_style_guide or ""
    return {
        "sentence_file_id": sentence_file_id,
        "sentence_file_name": sentence_file_name or "",
        "sentence_guide_chars": len(guide_text),
        "candidate_guide_chars": len(candidate_text),
        "ai_style_guide_chars": len(ai_text),
        "sentence_guide_sha1": hashlib.sha1(guide_text.encode("utf-8")).hexdigest()[:12] if guide_text else "",
        "terminology_file_id": terminology_file_id,
        "terminology_file_name": terminology_file_name or "",
        "ai_skipped": skip_ai,
        "ai_skip_reason": skip_reason or "",
        "ai_changed": ai_polished != pre_polished,
        "total_change_count": total_changes,
        "visible_change_count": visible_changes if visible_changes is not None else total_changes,
        "previous_polish_found": previous_polish_found,
        "previous_new_change_count": previous_new_change_count,
    }


# 句式清单缓存
_sentence_guide_cache: dict = {}
_term_cache: dict = {}
_typo_cache: dict = {}
_CANDIDATE_RECALL_GUIDE_MARKER = '## 候选召回句式库'
_AI_STYLE_GUIDE_MARKER = '## 仅供 AI 润色的通用风格指南'
_STEP_PREFIX_PATTERN = re.compile(r'^((?:\d+[.、)]?)+)(?:[.。]+)?(?:\s+|(?=[\u4e00-\u9fffA-Za-z(（]))\s*(.+)$')
_DOC_REVIEW_REFERENCE_CANDIDATE_THRESHOLD = 75


def _invalidate_sentence_guide_cache(style_guide_id: Optional[int] = None):
    """句式文件更新后清理相关缓存，保证新内容立即生效。"""
    _sentence_guide_cache.clear()


def _normalize_sentence_for_match(text: str) -> str:
    """去掉空白和常见标点，用于轻量句式相似度匹配。"""
    if not text:
        return ""
    text = str(text).strip()
    text = re.sub(r'^[\s*•·\-]+', '', text)
    text = re.sub(r'\s*[（(][^()（）]*(?:图|表)[^()（）]*[）)]\s*$', '', text)
    text = _canonicalize_sentence_for_template_match(text)
    return re.sub(r'[\s，。！？；：,.!?;:""''()（）【】\[\]<>《》-]+', '', text)


def _split_step_prefix(text: str) -> tuple[str, str]:
    value = str(text or '').strip()
    if not value:
        return '', ''
    match = _STEP_PREFIX_PATTERN.match(value)
    if not match:
        return '', value
    return match.group(1), (match.group(2) or '').strip().lstrip('.。')


def _split_list_marker_prefix(text: str) -> tuple[str, str]:
    value = str(text or '')
    match = re.match(r'^(\s*[*\-•·]+\s*)(.+)$', value)
    if not match:
        return '', value.strip()
    return match.group(1), match.group(2).strip()


def _split_notice_prefix(text: str) -> tuple[str, str]:
    value = str(text or '').strip()
    match = re.match(r'^((?:请)?注意[：:])\s*(.+)$', value)
    if not match:
        return '', value
    return match.group(1), match.group(2).strip()


def _should_add_contextual_terminal_punctuation(text: str) -> bool:
    value = str(text or '').strip()
    if not value or value.endswith(('。', '.', '！', '!', '？', '?', '；', ';', '：', ':')):
        return False
    if not re.search(r'[\u4e00-\u9fff]', value):
        return False

    step_prefix, body = _split_step_prefix(value)
    list_prefix, list_body = _split_list_marker_prefix(body or value)
    notice_prefix, notice_body = _split_notice_prefix(list_body or body or value)
    candidate = (notice_body or list_body or body or value).strip()
    if len(candidate) < 12:
        return False
    if not (step_prefix or list_prefix or notice_prefix) and _looks_like_title_or_noun_phrase(value):
        return False
    if _looks_like_title_or_noun_phrase(candidate):
        return False

    sentence_markers = (
        '应', '应当', '需', '需要', '必须', '请', '避免', '不得', '禁止', '用于', '以便',
        '可以', '将', '按照', '根据', '确保', '确认', '检查', '点击', '选择', '输入', '打开',
        '关闭', '安装', '连接', '取出', '放置', '加入', '倒入', '观察', '记录',
    )
    has_marker = any(marker in candidate for marker in sentence_markers)
    has_clause = '，' in candidate or ',' in candidate
    has_context_prefix = bool(step_prefix or list_prefix or notice_prefix)
    return has_marker and (has_clause or has_context_prefix)


def _normalize_terminal_sentence_punctuation(text: str) -> str:
    value = str(text or '').strip()
    if not value:
        return value
    match = re.search(r'([。.!！？?]+)$', value)
    if not match or len(match.group(1)) <= 1:
        return value
    return f"{value[:match.start(1)]}{match.group(1)[-1]}"


def _reapply_sentence_prefix(original: str, suggestion: str) -> str:
    result = str(suggestion or '').strip()
    if not result:
        return result

    step_prefix, _ = _split_step_prefix(original)
    if step_prefix and not result.startswith(step_prefix):
        result = f'{step_prefix}{result}'

    list_prefix, _ = _split_list_marker_prefix(original)
    if list_prefix and not result.startswith(list_prefix):
        result = f'{list_prefix}{result}'

    notice_prefix, _ = _split_notice_prefix(original)
    if notice_prefix and not result.startswith(notice_prefix):
        result = f'{notice_prefix}{result}'

    result = _normalize_terminal_sentence_punctuation(result)
    return re.sub(r'^(\d+(?:\.\d+)*)\.{2,}', r'\1.', result)


def _normalize_doc_ai_line(original: str, ai_line: str) -> str:
    value = str(ai_line or '').strip()
    if not value:
        return value
    return _reapply_sentence_prefix(original, value)


def _normalize_doc_polished_text(original: str, polished_text: str, is_title: bool = False) -> str:
    value = str(polished_text or '').strip()
    if not value:
        return value
    if is_title:
        return _normalize_terminal_sentence_punctuation(value)
    return _reapply_sentence_prefix(original, value)


def _normalize_visible_compare_text(text: str) -> str:
    value = _normalize_compare_text(text)
    _, value = _split_list_marker_prefix(value)
    _, value = _split_notice_prefix(value)
    step_prefix, body = _split_step_prefix(value)
    if step_prefix and body:
        value = body
    return value.strip()


def _normalize_document_feedback_key(text: str) -> str:
    value = str(text or '').strip()
    if not value:
        return ''
    document_match = re.match(r'^document:(\d+)(?::.*)?$', value)
    if document_match:
        return f"document:{document_match.group(1)}"
    value = re.sub(r'^【修订标记版】', '', value)
    value = re.sub(r'^\[修订标记版\]', '', value)
    return value.strip()


def _build_document_feedback_record_key(document_id: Optional[int], source_filename: str = '') -> str:
    if document_id:
        normalized_name = _normalize_document_feedback_key(source_filename)
        return f"document:{document_id}:{normalized_name}" if normalized_name else f"document:{document_id}"
    return _normalize_document_feedback_key(source_filename)


def _document_feedback_record_aliases(record: PolishFeedback) -> set[str]:
    raw_key = str(getattr(record, 'original_text', '') or '').strip()
    normalized_key = _normalize_document_feedback_key(raw_key)
    aliases = {normalized_key} if normalized_key else set()
    document_match = re.match(r'^document:(\d+)(?::(.+))?$', raw_key)
    if document_match:
        aliases.add(f"document:{document_match.group(1)}")
        filename_part = _normalize_document_feedback_key(document_match.group(2) or '')
        if filename_part:
            aliases.add(filename_part)
    return {alias for alias in aliases if alias}


def _latest_document_feedback_records(records: list[PolishFeedback]) -> list[PolishFeedback]:
    latest_by_key = {}
    for record in records or []:
        aliases = _document_feedback_record_aliases(record)
        if not aliases:
            aliases = {f'feedback:{record.id}'}

        previous_records = [latest_by_key.get(alias) for alias in aliases if latest_by_key.get(alias) is not None]
        latest_record = record
        for previous in previous_records:
            if previous is not None and (previous.created_at, previous.id) > (latest_record.created_at, latest_record.id):
                latest_record = previous

        merged_aliases = set(aliases)
        for previous in previous_records:
            if previous is None:
                continue
            merged_aliases.update(_document_feedback_record_aliases(previous))

        for alias in merged_aliases:
            latest_by_key[alias] = latest_record

    unique_records = {}
    for record in latest_by_key.values():
        unique_records[record.id] = record
    return list(unique_records.values())


def _document_feedback_stats(records: list[PolishFeedback]) -> dict:
    latest_records = _latest_document_feedback_records(records)
    total_submissions = len(latest_records)
    if total_submissions == 0:
        return {"total_submissions": 0, "average_accuracy": 0}

    submission_ratios = []
    for record in latest_records:
        total_changes = record.processed_count or 0
        if total_changes <= 0:
            submission_ratios.append(0)
            continue
        submission_ratios.append((record.accuracy or 0) / total_changes)

    average_accuracy = round((sum(submission_ratios) / total_submissions) * 100, 1)
    return {"total_submissions": total_submissions, "average_accuracy": average_accuracy}


def _json_dumps(data) -> Optional[str]:
    if data is None:
        return None
    return json.dumps(data, ensure_ascii=False)


def _json_loads(data, default):
    if data in (None, "", b""):
        return default
    if isinstance(data, (dict, list)):
        return data
    try:
        parsed = json.loads(data)
    except Exception:
        return default
    return parsed if parsed is not None else default


def _feedback_accuracy_percent(record: PolishFeedback) -> float:
    accuracy = float(record.accuracy or 0)
    return round(max(0.0, min(100.0, accuracy)), 1)


def _apply_created_at_filters(query, model, start_date: Optional[str] = None, end_date: Optional[str] = None):
    if start_date:
        start_dt = datetime.datetime.strptime(start_date, '%Y-%m-%d')
        query = query.filter(model.created_at >= start_dt)
    if end_date:
        end_dt = datetime.datetime.strptime(end_date, '%Y-%m-%d') + datetime.timedelta(days=1)
        query = query.filter(model.created_at < end_dt)
    return query


def _session_match_rate_percent(original_text: str, polished_text: str) -> float:
    before = str(original_text or '').strip()
    after = str(polished_text or '').strip()
    if not before or not after:
        return 0.0
    return float(_cat_match_score(before, after).get('overall_percent', 0) or 0)


def _build_document_feedback_lookup(records: list[PolishFeedback]) -> dict[str, PolishFeedback]:
    latest_by_alias = {}
    for record in _latest_document_feedback_records(records):
        for alias in _document_feedback_record_aliases(record):
            latest_by_alias[alias] = record
    return latest_by_alias


def _build_document_feedback_lookup_for_source_names(db: Session, source_names: list[str]) -> dict[str, PolishFeedback]:
    normalized_names = [
        _normalize_document_feedback_key(name)
        for name in source_names
        if str(name or '').strip()
    ]
    normalized_names = [name for name in normalized_names if name]
    if not normalized_names:
        return {}

    from sqlalchemy import or_

    conditions = [PolishFeedback.original_text.contains(name) for name in normalized_names]
    records = db.query(PolishFeedback).filter(
        PolishFeedback.target == 'document_sentence_guide'
    ).filter(or_(*conditions)).all()
    return _build_document_feedback_lookup(records)


def _resolve_sentence_file_name(db: Session, sentence_file_id: Optional[int]) -> Optional[str]:
    if not sentence_file_id:
        return None
    file = db.query(KnowledgeFile).filter(KnowledgeFile.id == sentence_file_id).first()
    return file.name if file else None


def _resolve_sentence_scope_name(db: Session, sentence_file_id: Optional[int], product_type: Optional[str] = None) -> Optional[str]:
    sentence_file_name = _resolve_sentence_file_name(db, sentence_file_id)
    if sentence_file_name:
        return sentence_file_name
    normalized_product_type = _normalize_document_product_type(product_type)
    if normalized_product_type:
        return f"{normalized_product_type}（全部文件）"
    return None


def _resolve_terminology_scope_name(db: Session, terminology_file_id: Optional[int], product_type: Optional[str] = None) -> Optional[str]:
    if terminology_file_id:
        file = db.query(KnowledgeFile).filter(KnowledgeFile.id == terminology_file_id).first()
        if file:
            return file.name
    normalized_product_type = _normalize_document_product_type(product_type)
    if normalized_product_type:
        return f"{normalized_product_type}（全部文件）"
    return None


def _build_feedback_correction_items(
    feedback: "FeedbackInput",
    corrections_pairs: list[tuple[str, str]],
    raw_lines: list[str],
) -> list[dict]:
    items = []
    if feedback.target == "terminology":
        for before, after in corrections_pairs:
            items.append({
                "before": before,
                "after": after,
                "category": _cat_report_change_category(before, after, "modify"),
            })
        return items

    baseline = str(feedback.polished_text or feedback.original_text or "").strip()
    for line in raw_lines:
        after = str(line or "").strip()
        if not after:
            continue
        items.append({
            "before": baseline,
            "after": after,
            "category": _cat_report_change_category(baseline, after, "modify"),
        })
    return items


def _summarize_categories(items: list[dict]) -> dict:
    counts = defaultdict(int)
    for item in items or []:
        category = str((item or {}).get("category") or "其他").strip() or "其他"
        counts[category] += 1
    return dict(counts)


def _category_brief_text(category_counts: dict, limit: int = 3) -> str:
    rows = sorted((category_counts or {}).items(), key=lambda item: (-int(item[1] or 0), item[0]))
    return "、".join(f"{name}×{count}" for name, count in rows[:limit])


def _normalize_stat_date(value) -> str:
    if isinstance(value, datetime.datetime):
        return value.strftime('%Y-%m-%d')
    text = str(value or '').strip()
    return text[:10] if len(text) >= 10 else text


def _cat_decision_replacement_text(decision) -> str:
    action = getattr(decision, 'action', '')
    if action == 'accept':
        return getattr(decision, 'accepted_template', '') or ''
    if action == 'modify':
        return getattr(decision, 'modified_text', '') or ''
    if action == 'reject':
        return getattr(decision, 'rejected_template', '') or ''
    return ''


def _normalize_polished_document_key(text: str) -> str:
    return _normalize_document_feedback_key(text)


def _find_previous_polished_document(db: Session, filename: str, created_by: Optional[int], file_type: str = 'docx') -> Optional[PolishedDocument]:
    normalized_name = _normalize_polished_document_key(filename)
    if not normalized_name:
        return None
    query = db.query(PolishedDocument)
    if file_type:
        query = query.filter(PolishedDocument.file_type == file_type)
    if created_by is not None:
        query = query.filter(PolishedDocument.created_by == created_by)
    candidates = query.order_by(PolishedDocument.created_at.desc(), PolishedDocument.id.desc()).limit(50).all()
    for doc in candidates:
        doc_keys = {
            _normalize_polished_document_key(doc.name),
            _normalize_polished_document_key(doc.filename),
        }
        if normalized_name in doc_keys:
            return doc
    return None


def _bigram_set(text: str) -> set:
    """将文本切分为字符 bigram 集合。"""
    if len(text) < 2:
        return {text} if text else set()
    return {text[i:i+2] for i in range(len(text) - 1)}


def _lcs_ratio(a: str, b: str) -> float:
    """最长公共子序列长度与较长字符串长度之比。"""
    if not a or not b:
        return 0.0
    m, n = len(a), len(b)
    if m == 0 or n == 0:
        return 0.0
    # 使用 1D 数组优化 LCS
    prev = [0] * (n + 1)
    for i in range(1, m + 1):
        cur = [0] * (n + 1)
        for j in range(1, n + 1):
            if a[i-1] == b[j-1]:
                cur[j] = prev[j-1] + 1
            else:
                cur[j] = max(prev[j], cur[j-1])
        prev = cur
    lcs_len = prev[n]
    return lcs_len / max(m, n)


def _split_sentence_clauses(sentence: str) -> list[str]:
    """按常见中文停顿符号拆分分句，用于长句与模板句的局部对齐。"""
    if not sentence:
        return []
    text = re.sub(r'[。！？；!?;]+', '，', sentence)
    return [clause.strip() for clause in re.split(r'[，,]+', text) if len(_normalize_sentence_for_match(clause)) >= 2]


def _simplify_clause_for_match(sentence: str) -> str:
    """去掉分句中的弱提示词，保留更有判别力的动作和对象。"""
    text = _normalize_sentence_for_match(sentence)
    if not text:
        return ""

    for pattern in [
        '请注意', '注意', '请', '需要在', '需要', '需', '是否还有', '是否有', '还有', '以及', '且', '并',
    ]:
        text = text.replace(pattern, '')

    text = text.replace('干冰剩余', '剩余干冰')
    return text


def _clause_similarity(a: str, b: str) -> float:
    """计算分句级相似度，对模板扩写场景做更强的包含判断。"""
    a_norm = _simplify_clause_for_match(a)
    b_norm = _simplify_clause_for_match(b)
    if not a_norm or not b_norm:
        return 0.0
    if a_norm == b_norm:
        return 1.0

    if a_norm in b_norm or b_norm in a_norm:
        base = max(0.95, _lcs_ratio(a_norm, b_norm))
    else:
        bigram_score = 0.0
        bigrams_a = _bigram_set(a_norm)
        bigrams_b = _bigram_set(b_norm)
        union = len(bigrams_a | bigrams_b)
        if union > 0:
            bigram_score = len(bigrams_a & bigrams_b) / union
        base = 0.45 * bigram_score + 0.55 * _lcs_ratio(a_norm, b_norm)

        source_ngrams = set()
        for n in range(2, min(4, len(a_norm)) + 1):
            source_ngrams.update(a_norm[i:i+n] for i in range(len(a_norm) - n + 1))
        if source_ngrams:
            covered = sum(1 for gram in source_ngrams if gram in b_norm)
            coverage_score = covered / len(source_ngrams)
            base = max(base, coverage_score)

    if any(verb in a_norm and verb in b_norm for verb in _CORE_VERBS):
        base += 0.05
    return min(base, 1.0)


def _clause_alignment_score(a: str, b: str) -> float:
    """要求源句的每个关键分句都能在模板中找到对应片段。"""
    clauses_a = _split_sentence_clauses(a)
    clauses_b = _split_sentence_clauses(b)
    if not clauses_a or not clauses_b:
        return 0.0

    best_scores = []
    for clause_a in clauses_a:
        best_scores.append(max(_clause_similarity(clause_a, clause_b) for clause_b in clauses_b))
    return sum(best_scores) / len(best_scores)


def _sentence_similarity(a: str, b: str) -> float:
    """
    计算两个中文句子的相似度（分段匹配）。

    综合 bigram Jaccard 相似度（捕捉局部片段重叠）和
    LCS 比率（捕捉整体结构相似度），取加权平均。
    当常规得分不足但存在高重合公共子串时，以子串比例兜底。
    """
    if a == b:
        return 1.0
    a_norm = _normalize_sentence_for_match(a)
    b_norm = _normalize_sentence_for_match(b)
    if not a_norm or not b_norm:
        return 0.0
    if a_norm == b_norm:
        return 1.0
    # Bigram Jaccard
    bigrams_a = _bigram_set(a_norm)
    bigrams_b = _bigram_set(b_norm)
    intersection = len(bigrams_a & bigrams_b)
    union = len(bigrams_a | bigrams_b)
    bigram_score = intersection / union if union > 0 else 0.0
    # LCS ratio
    lcs_score = _lcs_ratio(a_norm, b_norm)
    # 句子结构相似度
    struct_score = _compare_sentence_structure(a, b)
    # 语义关键词重叠
    keyword_score = _compare_semantic_keywords(a, b)
    clause_score = _clause_alignment_score(a, b)
    # 加权：bigram 0.3 + LCS 0.3 + 结构 0.2 + 关键词 0.2
    score = 0.3 * bigram_score + 0.3 * lcs_score + 0.2 * struct_score + 0.2 * keyword_score
    if clause_score > 0:
        score = max(score, 0.55 * score + 0.45 * clause_score)
        if clause_score >= 0.9:
            score = max(score, 0.88)
        elif clause_score >= 0.82:
            score = max(score, 0.84)
    # 高重合子串兜底：常规分不到 0.85 但最长公共子串覆盖 >=75% 时激活
    if score < 0.85:
        s = SequenceMatcher(None, a_norm, b_norm)
        match = s.find_longest_match(0, len(a_norm), 0, len(b_norm))
        shorter = min(len(a_norm), len(b_norm))
        substr_ratio = match.size / shorter if shorter > 0 else 0.0
        if substr_ratio >= 0.75:
            score = 0.80 + (substr_ratio - 0.75) * 0.8  # 0.80~1.00
    return min(score, 1.0)


def _template_replace_guard(sentence: str, template: str, match_level: str, score: float) -> bool:
    """控制句式模板自动替换，避免弱相关句子误命中。"""
    if not template:
        return False
    sentence_norm = _normalize_sentence_for_match(sentence)
    template_norm = _normalize_sentence_for_match(template)
    if sentence_norm and len(sentence_norm) <= 8 and not re.search(r'[，。；！？,;!?]', sentence) and len(template_norm) >= len(sentence_norm) * 2:
        return False
    if match_level == 'L1':
        return True

    clause_score = _clause_alignment_score(sentence, template)
    keyword_score = _compare_semantic_keywords(sentence, template)
    structure_score = _compare_sentence_structure(sentence, template)
    ui_object_score = _ui_action_object_score(sentence, template)
    structured_score = _structured_match_score(sentence, template)
    enumeration_score = _enumeration_overlap_score(sentence, template)
    source_struct = _extract_sentence_structure(sentence)
    template_struct = _extract_sentence_structure(template)
    ui_verbs = ('点击', '选择', '打开', '关闭', '进入', '退出', '切换到', '输入')

    if source_struct.get("pattern") == "操作UI元素" and template_struct.get("pattern") == "操作UI元素":
        if ui_object_score < 0.28:
            return False
    if source_struct.get("pattern") == "操作UI元素" and template_struct.get("pattern") != "操作UI元素":
        if any(verb in template for verb in ui_verbs):
            return False
    if _has_conflicting_critical_literals(sentence, template):
        return False
    if _has_inconsistent_model_spacing(sentence, template):
        return False
    if _looks_like_truncated_template(sentence, template):
        return False

    if structured_score['source']['conditions'] and structured_score['condition'] < 0.35:
        return False
    if structured_score['source']['actions'] and structured_score['action'] < 0.35:
        return False
    strong_condition_action = structured_score['condition'] >= 0.9 and structured_score['action'] >= 0.9 and structured_score['total'] >= 0.42
    if structured_score['source']['pairs'] and structured_score['pair'] < 0.3 and not strong_condition_action:
        return False
    if structured_score['source']['objects'] and structured_score['object'] < 0.25:
        return False
    if _has_conflicting_action_object_pair(structured_score['source']['pairs'], structured_score['target']['pairs']):
        return False
    if _has_conflicting_reference_target(sentence, template):
        return False
    if _has_excessive_template_expansion(sentence, template):
        return False
    if not structured_score['source']['actions'] and len(_extract_enumeration_items(sentence)) >= 2 and enumeration_score < 0.72:
        return False
    if len(_extract_action_sequence(structured_score['source']['pairs'])) >= 3 and len(_extract_action_sequence(structured_score['target']['pairs'])) < len(_extract_action_sequence(structured_score['source']['pairs'])):
        return False
    if len(_extract_action_sequence(structured_score['source']['pairs'])) >= 2 and structured_score['sequence'] < 0.65 and not strong_condition_action:
        return False
    if _extract_action_targets(structured_score['source']['pairs'], '进入') and structured_score['ui_target'] < 0.45:
        return False

    if match_level == 'L2':
        return score >= _POLISH_MATCH_CONFIG['l2_min_confidence'] and (
            structured_score['total'] >= 0.55 or
            clause_score >= 0.72 or
            (clause_score >= 0.62 and keyword_score >= 0.18 and structured_score['total'] >= 0.45) or
            (keyword_score >= 0.3 and structure_score >= 0.5)
        )

    if match_level == 'L3':
        return score >= _POLISH_MATCH_CONFIG['l3_auto_confidence'] and (
            (clause_score >= 0.82 and keyword_score >= 0.18 and structured_score['total'] >= 0.5) or
            (structured_score['sequence'] >= 1.0 and structured_score['ui_target'] >= 0.9 and structured_score['total'] >= 0.68) or
            (structured_score['condition'] >= 0.9 and structured_score['action'] >= 0.9 and structured_score['total'] >= 0.42) or
            (structured_score['total'] >= 0.7 and keyword_score >= 0.35 and structured_score['pair'] >= 0.75)
        )

    return False


# 核心动词列表（技术文档常见操作动词）
_CORE_VERBS = {
    '置于', '放置', '安装', '对应', '匹配', '确保', '检查', '设置', '调整',
    '校准', '测量', '分析', '记录', '保存', '删除', '关闭', '打开', '启动',
    '停止', '连接', '断开', '输入', '输出', '读取', '写入', '扫描', '点击',
    '选择', '确认', '取消', '添加', '移除', '插入', '取出', '转移', '等待',
    '观察', '核对', '撕开', '倒入', '拨至', '装入', '置于', '装填', '取下',
    '录入', '进入', '回顾', '冻融', '混样', '见', '解冻', '混匀', '离心',
}

# 方位词列表
_POSITION_WORDS = {'上', '下', '前', '后', '左', '右', '内', '外', '中', '间', '里', '旁', '侧'}
_PRODUCT_NOUN_MARKERS = ('试剂盒套装', '试剂套装', '试剂盒', '试剂', '产品', '文库')

_STRUCTURE_CONDITION_PREFIXES = ('当', '在', '若', '如果', '如', '待', '等到', '完成', '确认', '参数确认')
_STRUCTURE_CONDITION_SUFFIXES = ('时', '后', '之后', '前', '情况下')
_STRUCTURE_ACTION_VERBS = tuple(sorted(_CORE_VERBS | {'使用', '需使用', '点击', '检查', '确认', '输入', '进入', '组成', '构成', '包含', '包括'}, key=len, reverse=True))
_STRUCTURE_OBJECT_STOPWORDS = {'进行', '完成', '实现', '执行', '操作', '处理', '对应', '相关', '界面上', '界面中'}
_GENERIC_UI_MARKERS = {'按钮', '图标', '弹窗', '窗口', '界面', '页面', '标签', '选项', '菜单'}
_OBJECT_PREFIX_PATTERNS = ('相对应的', '对应的', '适配的', '适用的', '相关的')


def _normalize_product_subject(text: str) -> str:
    value = _normalize_structure_text(text)
    if not value:
        return ''
    value = re.sub(r'^(本|该)', '', value)
    compact = re.sub(r'\s+', '', value)
    if '试剂' in compact and '套装' in compact:
        return '试剂套装'
    if '试剂盒' in compact:
        return '试剂盒'
    if '试剂' in compact:
        return '试剂'
    if '产品' in compact:
        return '产品'
    if '文库' in compact:
        return '文库'
    return value


def _extract_subject_markers(text: str) -> set[str]:
    markers = set()
    raw = str(text or '')
    if not raw:
        return markers
    for match in re.findall(r'(?:本|该)?[\u4e00-\u9fffA-Za-z0-9\-\s]{0,40}(?:试剂盒套装|试剂套装|试剂盒|试剂|产品|文库)', raw):
        normalized = _normalize_product_subject(match)
        if normalized:
            markers.add(normalized)
    return markers


def _extract_critical_literal_markers(text: str) -> set[str]:
    markers = set()
    raw = str(text or '')
    if not raw:
        return markers

    compact = re.sub(r'\s+', '', raw)
    for match in re.findall(r'(?:样本制备套件|试剂盒套装|试剂套装|套件|Box|BOX|盒|孔位|孔|通道|板|卡|槽位)\s*[A-Z](?![A-Za-z])', raw):
        normalized = re.sub(r'\s+', '', match)
        if normalized:
            markers.add(normalized)

    exact_phrases = (
        '扫码枪',
        '二维码',
        '包装上的二维码',
        '制备卡包装上的二维码',
        '制备卡',
        '样本制备系统',
    )
    for phrase in exact_phrases:
        if phrase in compact:
            markers.add(phrase)

    return markers


def _has_conflicting_critical_literals(sentence: str, template: str) -> bool:
    source_markers = _extract_critical_literal_markers(sentence)
    if not source_markers:
        return False
    target_markers = _extract_critical_literal_markers(template)
    return not source_markers.issubset(target_markers)


def _extract_semantic_keywords(sentence: str) -> list[str]:
    """提取更稳健的语义锚点，避免把连续字片段误当术语。"""
    raw = str(sentence or '').strip()
    if not raw:
        return []

    keywords = set()
    intent = _extract_sentence_intent(raw)
    for token in intent.get('conditions', []) + intent.get('actions', []) + intent.get('objects', []):
        normalized = _normalize_structure_text(token)
        if len(normalized) >= 2:
            keywords.add(normalized)

    for token in _extract_result_units(raw) + _extract_additional_units(raw) + _extract_enumeration_items(raw):
        normalized = _normalize_structure_text(token)
        if len(normalized) >= 2:
            keywords.add(normalized)

    for match in re.findall(r'[A-Za-z0-9\-]{1,20}(?:按钮|图标|界面|页面|窗口)|[\u4e00-\u9fffA-Za-z0-9\-]{2,20}?(?:界面|页面|窗口|弹窗|按钮|图标|视频|说明书|指南|仓门|载台|制备卡|样本|试剂|产物|孔|盖板|移液器|油相|水相)', raw):
        normalized = _normalize_structure_text(match)
        if len(normalized) >= 2 and normalized not in _GENERIC_UI_MARKERS:
            keywords.add(normalized)

    for marker in _extract_reference_markers(raw):
        if len(marker) >= 2:
            keywords.add(marker)

    for marker in _extract_subject_markers(raw):
        if len(marker) >= 2:
            keywords.add(marker)

    for marker in _extract_critical_literal_markers(raw):
        if len(marker) >= 2:
            keywords.add(marker)

    for marker in _extract_model_markers(raw).keys():
        if len(marker) >= 4:
            keywords.add(marker)

    for pw in _POSITION_WORDS:
        if pw in raw:
            keywords.add(pw)
    return list(keywords)


def _compare_semantic_keywords(a: str, b: str) -> float:
    """计算两个句子的语义关键词重叠度。"""
    kw_a = _extract_semantic_keywords(a)
    kw_b = _extract_semantic_keywords(b)
    if not kw_a or not kw_b:
        return 0.0
    overlap = len(set(kw_a) & set(kw_b))
    denominator = max(len(kw_a), len(kw_b))
    return overlap / denominator if denominator > 0 else 0.0


def _normalize_structure_text(text: str) -> str:
    value = _normalize_sentence_for_match(text)
    value = re.sub(r'^[\d.]+', '', value)
    value = re.sub(r'^(请|需|需要|应|应该|必须|可|可以)', '', value)
    value = re.sub(r'(即可|开始实验|开始操作|完成操作)$', '', value)
    return value.strip()


def _normalize_structure_object(text: str) -> str:
    value = _normalize_structure_text(text)
    if not value:
        return ''
    value = value.replace('按要求', '按指示')
    value = re.sub(r'^[\d.]+', '', value)
    value = re.sub(r'^(?:在)?[^，。；]*?(?:时|后)', '', value)
    value = re.sub(r'^(是否还有|是否有|还有|有无|是否)', '', value)
    for word in sorted(_STRUCTURE_OBJECT_STOPWORDS, key=len, reverse=True):
        if value.startswith(word):
            value = value[len(word):]
        if value.endswith(word):
            value = value[:-len(word)]
    value = re.sub(r'^(对应|相关|界面上|界面中)', '', value)
    for prefix in _OBJECT_PREFIX_PATTERNS:
        if value.startswith(prefix):
            value = value[len(prefix):]
    if re.search(r'(?:^|见)(?:上|下)?表\d*$', value) or re.fullmatch(r'(?:上|下)?表\d*', value):
        return '表引用'
    if re.search(r'(?:^|见)(?:上|下)?图\d*$', value) or re.fullmatch(r'(?:上|下)?图\d*', value):
        return '图引用'
    value = re.sub(r'(信息|内容)$', '', value)
    return value.strip()


def _is_placeholder_platform_object(text: str) -> bool:
    normalized = _normalize_structure_object(text)
    if not normalized:
        return False
    return (
        ('平台' in normalized and '测序类型' in normalized)
        or ('以下平台' in normalized)
        or ('如下平台' in normalized)
    )


def _is_generic_ui_object(text: str) -> bool:
    normalized = _normalize_structure_object(text)
    if not normalized:
        return False
    return normalized in _GENERIC_UI_MARKERS or normalized.endswith(tuple(_GENERIC_UI_MARKERS))


def _normalize_structure_verb(text: str) -> str:
    value = _normalize_structure_text(text)
    verb_aliases = {
        '观察': '检查',
        '查看': '检查',
        '核对': '检查',
    }
    return verb_aliases.get(value, value)


def _extract_condition_units(sentence: str) -> list[str]:
    conditions = []
    for clause in _split_sentence_clauses(sentence):
        clause_text = clause.strip('，,。；;：: ')
        normalized = _normalize_structure_text(clause_text)
        if not normalized:
            continue
        if normalized.startswith('在') and not re.search(r'(时|后|之前|前|情况下)$', normalized):
            continue
        if any(normalized.startswith(prefix) for prefix in _STRUCTURE_CONDITION_PREFIXES):
            match = re.match(r'^(.+?(?:时|后|之前|前|情况下))', normalized)
            conditions.append(match.group(1) if match else normalized)
            continue
        if re.search(r'(?:时|后|之前|前|情况下)$', normalized):
            conditions.append(normalized)
            continue
        if any(suffix in normalized for suffix in _STRUCTURE_CONDITION_SUFFIXES) and any(prefix in normalized[:4] for prefix in _STRUCTURE_CONDITION_PREFIXES):
            match = re.match(r'^(.+?(?:时|后|之前|前|情况下))', normalized)
            conditions.append(match.group(1) if match else normalized)
    return list(dict.fromkeys(conditions))


def _extract_action_object_units(sentence: str) -> list[dict]:
    pairs = []
    seen = set()

    def _append_pair(verb: str, obj: str):
        verb_norm = _normalize_structure_verb(verb)
        obj_norm = _normalize_structure_object(obj)
        if not verb_norm:
            return
        key = (verb_norm, obj_norm)
        if key in seen:
            return
        seen.add(key)
        pairs.append({'verb': verb_norm, 'object': obj_norm})

    struct = _extract_sentence_structure(sentence)
    if struct.get('verb'):
        _append_pair(struct.get('verb', ''), struct.get('object', '') or struct.get('subject', ''))

    for clause in _split_sentence_clauses(sentence):
        clause_text = clause.strip('，,。；;：: ')
        normalized = _normalize_structure_text(clause_text)
        if not normalized:
            continue
        composition_match = re.search(r'(.+?)由(.+?)(组成|构成|包装组成|包装构成|包装)$', normalized)
        if composition_match:
            _append_pair('组成', composition_match.group(2))
            continue
        ui_enter_match = re.search(r'(进入)(.+?(?:界面|页面|窗口))', normalized)
        if ui_enter_match:
            _append_pair(ui_enter_match.group(1), ui_enter_match.group(2))
        for verb in _STRUCTURE_ACTION_VERBS:
            if verb not in normalized:
                continue
            before, after = normalized.split(verb, 1)
            if verb == '确认' and re.search(r'后(点击|进入|选择|打开|关闭|运行|run)', after):
                continue
            obj = re.sub(r'^(界面上|界面中|对应|相关)', '', after)
            obj = re.sub(r'(开始实验|开始操作|即可)$', '', obj)
            if not _normalize_structure_object(obj):
                if verb in {'组成', '构成', '包装'}:
                    continue
                before_obj = re.sub(r'^[\d.]+', '', before)
                before_obj = re.sub(r'(完成|确认无误|准备|结束)$', '', before_obj)
                if before_obj and not before_obj.endswith(('后', '时', '前')):
                    obj = before_obj
            _append_pair(verb, obj)
            break
    return pairs


def _extract_sentence_intent(sentence: str) -> dict:
    action_pairs = _extract_action_object_units(sentence)
    return {
        'conditions': _extract_condition_units(sentence),
        'actions': [item['verb'] for item in action_pairs if item.get('verb')],
        'objects': [item['object'] for item in action_pairs if item.get('object')],
        'pairs': action_pairs,
    }


def _extract_action_sequence(pairs: list[dict]) -> list[str]:
    sequence = []
    for pair in pairs or []:
        verb = pair.get('verb', '')
        if not verb or verb == '等待':
            continue
        if not sequence or sequence[-1] != verb:
            sequence.append(verb)
    return sequence


def _extract_action_targets(pairs: list[dict], verb: str) -> list[str]:
    targets = []
    for pair in pairs or []:
        if pair.get('verb', '') != verb:
            continue
        obj = pair.get('object', '')
        if obj and obj not in targets:
            targets.append(obj)
    return targets


def _object_item_count(text: str) -> int:
    value = _normalize_structure_object(text)
    if not value:
        return 0
    parts = [part for part in re.split(r'[、,，以及和与及]', value) if part]
    return len(parts) if parts else 0


def _extract_enumeration_items(sentence: str) -> list[str]:
    text = _normalize_sentence_for_match(sentence)
    if not text:
        return []
    if '适用于' in text:
        text = text.split('适用于', 1)[1]
    text = re.split(r'[。；;：:]', text, 1)[0]
    raw_parts = [part.strip() for part in re.split(r'[、,，]', text) if part.strip()]
    items = []
    for part in raw_parts:
        value = re.sub(r'^(以及|和|与|及)', '', part)
        value = re.sub(r'(等|样本|检测)$', '', value)
        value = value.strip()
        if len(value) >= 1:
            items.append(value)
    return items


def _extract_slot_sample_pairs(sentence: str) -> list[tuple[str, str]]:
    raw = str(sentence or '')
    if not raw:
        return []
    pairs = []
    pattern = re.compile(r'(?:孔位\s*)?([A-Za-z]\d)\s*对应\s*(?:sample|样本)\s*([A-Za-z]?\d+)', re.IGNORECASE)
    for slot, sample in pattern.findall(raw):
        slot_key = re.sub(r'\s+', '', slot).upper()
        sample_key = re.sub(r'[^A-Za-z0-9]+', '', sample).lower()
        if slot_key and sample_key:
            pairs.append((slot_key, sample_key))
    return pairs


def _slot_sample_pair_overlap_score(source_pairs: list[tuple[str, str]], target_pairs: list[tuple[str, str]]) -> float:
    if not source_pairs or not target_pairs:
        return 0.0
    source_samples = {sample for _, sample in source_pairs}
    target_samples = {sample for _, sample in target_pairs}
    source_slots = {re.sub(r'[^0-9]+', '', slot) for slot, _ in source_pairs}
    target_slots = {re.sub(r'[^0-9]+', '', slot) for slot, _ in target_pairs}
    sample_overlap = len(source_samples & target_samples) / max(len(source_samples), len(target_samples), 1)
    slot_overlap = len(source_slots & target_slots) / max(len(source_slots), len(target_slots), 1)
    return round(0.65 * sample_overlap + 0.35 * slot_overlap, 4)


def _is_slot_sample_mapping_sentence(sentence: str) -> bool:
    raw = str(sentence or '')
    if '对应' not in raw:
        return False
    if 'sample' not in raw.lower() and '样本' not in raw:
        return False
    return len(_extract_slot_sample_pairs(raw)) >= 3


def _extract_pooling_topic_markers(sentence: str) -> set[str]:
    raw = str(sentence or '')
    if not raw:
        return set()
    markers = set()
    clause_count = len(_split_sentence_clauses(raw))
    if clause_count >= 3:
        markers.add('multi_clause')
    if re.search(r'[；;]', raw):
        markers.add('major_split')
    if _extract_condition_units(raw):
        markers.add('condition_clause')

    actions = set(_extract_sentence_intent(raw).get('actions', []))
    if len(actions) >= 2:
        markers.add('multi_action')

    numeric_markers = _extract_numeric_markers(raw)
    if len(numeric_markers) >= 2:
        markers.add('numeric_dense')
    if len(numeric_markers) >= 4:
        markers.add('numeric_rich')

    model_markers = _extract_model_markers(raw)
    if model_markers:
        markers.add('model_ref')

    reference_markers = _extract_reference_markers(raw)
    if reference_markers:
        markers.add('reference_ref')

    if re.search(r'\d+\s*(?:个)?样本', raw):
        markers.add('sample_count')
    if re.search(r'或者|或|分别|对应|组合|编号|序列|方案|体系', raw):
        markers.add('structured_relation')

    for match in re.findall(r'[A-Za-z]{2,}\s*\d+(?:-\d+)?[A-Za-z]*', raw, flags=re.IGNORECASE):
        normalized = re.sub(r'\s+', '', match).lower()
        if normalized:
            markers.add(normalized)
            markers.add('coded_identifier')
    for match in re.findall(r'[A-Za-z]+(?:-[A-Za-z0-9]+(?:\s+[A-Za-z0-9]+)*)+', raw):
        normalized = re.sub(r'\s+', '', match).lower()
        if normalized:
            markers.add(normalized)
    for value in numeric_markers:
        markers.add(value)
    for value in reference_markers:
        markers.add(value)
    return markers


def _pooling_marker_overlap_score(source_markers: set[str], target_markers: set[str]) -> float:
    if not source_markers or not target_markers:
        return 0.0
    shared = len(source_markers & target_markers)
    return shared / max(min(len(source_markers), len(target_markers)), 1)


def _is_pooling_platform_sentence(sentence: str) -> bool:
    raw = str(sentence or '')
    if not raw:
        return False
    markers = _extract_pooling_topic_markers(raw)
    if len(_normalize_sentence_for_match(raw)) < 32:
        return False
    clause_count = len(_split_sentence_clauses(raw))
    if clause_count < 2:
        return False
    core_markers = {
        'sample_count',
        'model_ref',
        'coded_identifier',
    }
    support_markers = {
        'numeric_dense',
        'numeric_rich',
        'structured_relation',
        'major_split',
        'reference_ref',
        'condition_clause',
        'multi_clause',
    }
    if len(markers & core_markers) < 1:
        return False
    if len(markers & support_markers) < 2:
        return False
    return len(markers) >= 4


def _best_pooling_clause_template(sentence: str, templates: list) -> tuple:
    source_text = str(sentence or '').strip()
    source_markers = _extract_pooling_topic_markers(source_text)
    if not source_text or not source_markers:
        return None, 0.0, 'NONE'

    best_template = None
    best_score = 0.0
    source_actions = set(_extract_sentence_intent(source_text).get('actions', []))
    for template in templates or []:
        template_text = _template_entry_text(template).strip()
        if not template_text:
            continue
        if _has_excessive_template_expansion(source_text, template_text):
            continue
        target_markers = _extract_pooling_topic_markers(template_text)
        shared_markers = source_markers & target_markers
        marker_score = _pooling_marker_overlap_score(source_markers, target_markers)
        if marker_score <= 0 and source_markers:
            continue
        clause_score = _clause_alignment_score(source_text, template_text)
        lexical_score = _edit_distance_score(source_text, template_text)
        keyword_score = _compare_semantic_keywords(source_text, template_text)
        action_overlap = len(source_actions & set(_extract_sentence_intent(template_text).get('actions', [])))
        action_score = 1.0 if action_overlap > 0 else 0.0
        score = 0.4 * marker_score + 0.25 * clause_score + 0.15 * lexical_score + 0.1 * keyword_score + 0.1 * action_score
        if len(shared_markers) >= 2:
            score += 0.08
        if {'numeric_dense', 'coded_identifier', 'structured_relation'} <= shared_markers:
            score += 0.08
        if source_text == template_text:
            score = 1.0
        elif source_text in template_text or template_text in source_text:
            score = max(score, 0.88)
        if score > best_score:
            best_score = score
            best_template = template_text

    if best_template and best_score >= 0.5:
        level = 'POOL_EXACT' if best_score >= 0.88 else 'POOL'
        return best_template, round(best_score, 4), level
    return None, 0.0, 'NONE'


def _split_pooling_sentence_windows(sentence: str) -> list[str]:
    raw = str(sentence or '').strip()
    if not raw:
        return []
    windows = []
    major_segments = [segment.strip() for segment in re.split(r'[；;]', raw) if segment.strip()]
    for segment in major_segments:
        parts = [part.strip() for part in re.split(r'[，,]', segment) if part.strip()]
        if len(parts) <= 1:
            windows.append(segment.strip('。；; '))
            continue
        index = 0
        while index < len(parts):
            current = parts[index].strip('。；; ')
            combined = current
            current_actions = set(_extract_sentence_intent(current).get('actions', []))
            next_part = parts[index + 1].strip('。；; ') if index + 1 < len(parts) else ''
            next_actions = set(_extract_sentence_intent(next_part).get('actions', [])) if next_part else set()
            current_has_condition = bool(_extract_condition_units(current))
            current_is_intro = bool(re.match(r'^(根据|按照|按|若|如|如果|当|在|为|为了)', _normalize_structure_text(current)))
            current_is_context_clause = len(current_actions) == 0 and len(_normalize_sentence_for_match(current)) <= 28
            should_merge_next = bool(next_part) and (
                current_has_condition or
                current_is_intro or
                (current_is_context_clause and next_actions) or
                (re.search(r'\d', current) and re.match(r'^(或者|或|并|及|以及|同时)', next_part))
            )
            if should_merge_next and index + 1 < len(parts):
                combined = f'{current}，{next_part}'
                index += 1
                if index + 1 < len(parts) and re.match(r'^(或者|或|并|及|以及|同时)', parts[index + 1].strip()):
                    combined = f'{combined}，{parts[index + 1].strip("。；; ")}'
                    index += 1
            windows.append(combined)
            index += 1
    return [window for window in windows if window]


def _normalize_internal_sentence_punctuation(text: str) -> str:
    value = str(text or '')
    if not value:
        return value
    value = re.sub(r'[。.!！？?]+(?=[，,；;])', '', value)
    value = re.sub(r'([。.!！？?])\1+', r'\1', value)
    value = re.sub(r'([，,；;])\1+', r'\1', value)
    return value


def _enumeration_overlap_score(source: str, template: str) -> float:
    source_items = _extract_enumeration_items(source)
    template_items = _extract_enumeration_items(template)
    if len(source_items) < 2 or len(template_items) < 2:
        return 0.0
    scores = []
    for source_item in source_items:
        best = 0.0
        for target_item in template_items:
            if source_item == target_item:
                best = 1.0
                break
            if source_item in target_item or target_item in source_item:
                best = max(best, 0.92)
                continue
            best = max(best, _clause_similarity(source_item, target_item))
        scores.append(best)
    return sum(scores) / len(scores) if scores else 0.0


def _normalize_reference_marker(text: str) -> str:
    value = _normalize_structure_text(text)
    compact = re.sub(r'\s+', '', value.lower())
    if re.search(r'(?:上|下)?表\d*', compact):
        return '表引用'
    if re.search(r'(?:上|下)?图\d*', compact):
        return '图引用'
    value = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff]+', '', value)
    return value.lower()


def _extract_reference_markers(text: str) -> set[str]:
    markers = set()
    raw = str(text or '')
    if not raw:
        return markers

    for title in re.findall(r'《([^》]+)》', raw):
        marker = _normalize_reference_marker(title)
        if len(marker) >= 4:
            markers.add(marker)

    for code in re.findall(r'[A-Za-z]-?\d{3}-?\d{3,}-?\d{2,}', raw):
        marker = _normalize_reference_marker(code)
        if len(marker) >= 6:
            markers.add(marker)

    for ref in re.findall(r'(?:见)?(?:上|下)?(?:表|图)\s*\d*', raw):
        marker = _normalize_reference_marker(ref)
        if marker:
            markers.add(marker)

    if any(token in raw for token in ('下表', '见表', '表1', '表 1', '见图', '下图', '图1', '图 1')):
        for token in ('表引用', '图引用'):
            if token[0] in raw:
                markers.add(token)

    return markers


def _has_conflicting_reference_target(sentence: str, template: str) -> bool:
    template_markers = _extract_reference_markers(template)
    if not template_markers:
        return False
    source_markers = _extract_reference_markers(sentence)
    if source_markers & template_markers:
        return False
    source_has_reference = any(keyword in sentence for keyword in ('说明书', '操作指南', '参考'))
    return source_has_reference


def _action_sequence_score(source_pairs: list[dict], target_pairs: list[dict]) -> float:
    source_seq = _extract_action_sequence(source_pairs)
    target_seq = _extract_action_sequence(target_pairs)
    if not source_seq or not target_seq:
        return 0.0
    if source_seq == target_seq:
        return 1.0
    matched = 0
    cursor = 0
    for verb in source_seq:
        for idx in range(cursor, len(target_seq)):
            if target_seq[idx] == verb:
                matched += 1
                cursor = idx + 1
                break
    return matched / max(len(source_seq), len(target_seq))


def _token_overlap_score(source_tokens: list[str], target_tokens: list[str]) -> float:
    if not source_tokens or not target_tokens:
        return 0.0
    scores = []
    for source in source_tokens:
        best = 0.0
        for target in target_tokens:
            if not source or not target:
                continue
            if source == target:
                best = 1.0
                break
            if source in target or target in source:
                best = max(best, 0.92)
                continue
            best = max(best, _clause_similarity(source, target))
        scores.append(best)
    return sum(scores) / len(scores) if scores else 0.0


def _marker_overlap_score(source_markers: set[str], target_markers: set[str]) -> float:
    if not source_markers or not target_markers:
        return 0.0
    intersection = len(source_markers & target_markers)
    denominator = max(len(source_markers), len(target_markers))
    return intersection / denominator if denominator > 0 else 0.0


def _edit_distance_score(sentence: str, template: str) -> float:
    source = _normalize_sentence_for_match(sentence)
    target = _normalize_sentence_for_match(template)
    if not source or not target:
        return 0.0
    if source == target:
        return 1.0
    return SequenceMatcher(None, source, target).ratio()


def _extract_placeholder_markers(text: str) -> set[str]:
    markers = set()
    raw = str(text or '')
    if not raw:
        return markers
    patterns = [
        r'\{[^{}]{1,40}\}',
        r'<[^<>]{1,40}>',
        r'\[[^\[\]]{1,40}\]',
        r'%\([^)]+\)[a-zA-Z]',
        r'%[sdif]',
        r':[A-Za-z_][A-Za-z0-9_]*',
    ]
    for pattern in patterns:
        for match in re.findall(pattern, raw):
            normalized = re.sub(r'\s+', '', str(match).lower())
            if normalized:
                markers.add(normalized)
    return markers


def _extract_term_anchor_markers(text: str, structured_score: dict = None) -> set[str]:
    markers = set()
    raw = str(text or '')
    if not raw:
        return markers
    structured = structured_score['source'] if structured_score and 'source' in structured_score else _extract_sentence_intent(raw)
    for token in _extract_semantic_keywords(raw):
        normalized = _normalize_structure_text(token).lower()
        if normalized:
            markers.add(normalized)
    for group in ('conditions', 'actions', 'objects'):
        for token in structured.get(group, []) or []:
            normalized = _normalize_structure_text(token).lower()
            if normalized:
                markers.add(normalized)
    for token in _extract_result_units(raw):
        normalized = _normalize_structure_text(token).lower()
        if normalized:
            markers.add(normalized)
    for token in _extract_enumeration_items(raw):
        normalized = _normalize_structure_text(token).lower()
        if normalized:
            markers.add(normalized)
    markers.update(_extract_subject_markers(raw))
    markers.update(_extract_ui_markers(raw))
    markers.update(_extract_reference_markers(raw))
    markers.update(_extract_model_markers(raw).keys())
    return {marker for marker in markers if marker}


def _term_anchor_score(sentence: str, template: str, structured_score: dict = None) -> float:
    source_struct = structured_score or _structured_match_score(sentence, template)
    source_markers = _extract_term_anchor_markers(sentence, {'source': source_struct.get('source', {})})
    target_markers = _extract_term_anchor_markers(template, {'source': source_struct.get('target', {})})
    scores = []
    if source_markers or target_markers:
        scores.append(_marker_overlap_score(source_markers, target_markers))
    ui_score = _marker_overlap_score(_extract_ui_markers(sentence), _extract_ui_markers(template))
    if _extract_ui_markers(sentence) or _extract_ui_markers(template):
        scores.append(ui_score)
    reference_score = _marker_overlap_score(_extract_reference_markers(sentence), _extract_reference_markers(template))
    if _extract_reference_markers(sentence) or _extract_reference_markers(template):
        scores.append(reference_score)
    model_score = _marker_overlap_score(set(_extract_model_markers(sentence).keys()), set(_extract_model_markers(template).keys()))
    if _extract_model_markers(sentence) or _extract_model_markers(template):
        scores.append(model_score)
    return sum(scores) / len(scores) if scores else 0.0


def _number_placeholder_score(sentence: str, template: str) -> float:
    scores = []
    source_numbers = _extract_numeric_markers(sentence)
    target_numbers = _extract_numeric_markers(template)
    if source_numbers or target_numbers:
        scores.append(_marker_overlap_score(source_numbers, target_numbers))
    source_placeholders = _extract_placeholder_markers(sentence)
    target_placeholders = _extract_placeholder_markers(template)
    if source_placeholders or target_placeholders:
        scores.append(_marker_overlap_score(source_placeholders, target_placeholders))
    return sum(scores) / len(scores) if scores else 0.0


def _context_match_score(sentence: str, template: str, template_metadata=None, structured_score: dict = None) -> float:
    structured = structured_score or _structured_match_score(sentence, template)
    source_pairs = structured.get('source', {}).get('pairs', [])
    target_pairs = structured.get('target', {}).get('pairs', [])
    scores = []
    source_conditions = structured.get('source', {}).get('conditions') or []
    target_conditions = structured.get('target', {}).get('conditions') or []
    source_sequence = _extract_action_sequence(source_pairs)
    target_sequence = _extract_action_sequence(target_pairs)
    source_ui_targets = _extract_action_targets(source_pairs, '进入')
    target_ui_targets = _extract_action_targets(target_pairs, '进入')

    if source_conditions and target_conditions:
        scores.append(structured.get('condition', 0.0))
    if len(source_sequence) >= 2 and len(target_sequence) >= 2:
        scores.append(structured.get('sequence', 0.0))
    if source_ui_targets and target_ui_targets:
        scores.append(structured.get('ui_target', 0.0))
    if scores and isinstance(template_metadata, dict) and template_metadata.get('total', 0.0) > 0:
        scores.append(template_metadata.get('total', 0.0))
    return sum(scores) / len(scores) if scores else 0.0


def _action_object_pair_score(source_pairs: list[dict], target_pairs: list[dict]) -> float:
    if not source_pairs or not target_pairs:
        return 0.0
    scores = []
    for source in source_pairs:
        best = 0.0
        source_verb = source.get('verb', '')
        source_object = source.get('object', '')
        for target in target_pairs:
            target_verb = target.get('verb', '')
            target_object = target.get('object', '')
            verb_score = 0.0
            if source_verb and target_verb:
                if source_verb == target_verb:
                    verb_score = 1.0
                elif source_verb in target_verb or target_verb in source_verb:
                    verb_score = 0.9
            object_score = 0.0
            if source_object and target_object:
                if source_object == target_object:
                    object_score = 1.0
                elif source_object in target_object or target_object in source_object:
                    object_score = 0.94
                else:
                    object_score = max(_clause_similarity(source_object, target_object), _compare_semantic_keywords(source_object, target_object))
                if _is_generic_ui_object(source_object) or _is_generic_ui_object(target_object):
                    object_score = min(object_score, 0.35)
                if object_score < 0.35:
                    continue
            elif not source_object:
                object_score = 0.7
            if verb_score <= 0:
                continue
            pair_score = 0.6 * verb_score + 0.4 * object_score
            if source_verb == '点击' and (_is_generic_ui_object(source_object) or _is_generic_ui_object(target_object)):
                pair_score = min(pair_score, 0.45)
            best = max(best, pair_score)
        scores.append(best)
    return sum(scores) / len(scores) if scores else 0.0


def _has_conflicting_action_object_pair(source_pairs: list[dict], target_pairs: list[dict]) -> bool:
    if not source_pairs or not target_pairs:
        return False
    for source in source_pairs:
        source_verb = source.get('verb', '')
        source_object = source.get('object', '')
        if not source_verb or not source_object:
            continue
        same_verb_targets = [target for target in target_pairs if target.get('verb', '') == source_verb]
        if not same_verb_targets:
            continue
        source_item_count = _object_item_count(source_object)
        if source_item_count >= 2:
            target_item_counts = [_object_item_count(target.get('object', '')) for target in same_verb_targets if target.get('object', '')]
            if target_item_counts and max(target_item_counts) < source_item_count:
                return True
        object_scores = []
        for target in same_verb_targets:
            target_object = target.get('object', '')
            if not target_object:
                continue
            if source_verb == '使用' and (
                _is_placeholder_platform_object(source_object)
                or _is_placeholder_platform_object(target_object)
            ):
                object_scores.append(0.9)
                continue
            if source_verb == '点击' and ('按钮' in source_object or '按钮' in target_object):
                continue
            if source_object == target_object:
                object_scores.append(1.0)
                continue
            if source_object in target_object or target_object in source_object:
                object_scores.append(0.94)
                continue
            object_scores.append(max(_clause_similarity(source_object, target_object), _compare_semantic_keywords(source_object, target_object)))
        if object_scores and max(object_scores) < 0.35:
            return True
    return False


def _ui_target_score(source_pairs: list[dict], target_pairs: list[dict]) -> float:
    source_targets = _extract_action_targets(source_pairs, '进入')
    target_targets = _extract_action_targets(target_pairs, '进入')
    return _token_overlap_score(source_targets, target_targets)


def _structured_match_score(sentence: str, template: str) -> dict:
    source_intent = _extract_sentence_intent(sentence)
    template_intent = _extract_sentence_intent(template)
    condition_score = _token_overlap_score(source_intent['conditions'], template_intent['conditions'])
    action_score = _token_overlap_score(source_intent['actions'], template_intent['actions'])
    object_score = _token_overlap_score(source_intent['objects'], template_intent['objects'])
    pair_score = _action_object_pair_score(source_intent['pairs'], template_intent['pairs'])
    action_sequence_score = _action_sequence_score(source_intent['pairs'], template_intent['pairs'])
    ui_target_score = _ui_target_score(source_intent['pairs'], template_intent['pairs'])
    if any(_is_generic_ui_object(item) for item in source_intent['objects'] + template_intent['objects']):
        ui_overlap = _marker_overlap_score(_extract_ui_markers(sentence), _extract_ui_markers(template))
        if ui_overlap <= 0:
            object_score = min(object_score, 0.35)
            pair_score = min(pair_score, 0.45)
    has_signal = any(source_intent[key] for key in ('conditions', 'actions', 'objects'))
    total = 0.0 if not has_signal else 0.28 * pair_score + 0.18 * condition_score + 0.14 * action_score + 0.1 * object_score + 0.15 * action_sequence_score + 0.15 * ui_target_score
    return {
        'total': total,
        'condition': condition_score,
        'action': action_score,
        'object': object_score,
        'pair': pair_score,
        'sequence': action_sequence_score,
        'ui_target': ui_target_score,
        'source': source_intent,
        'target': template_intent,
    }


def _extract_numeric_markers(text: str) -> set[str]:
    markers = set()
    raw = str(text or '')
    if not raw:
        return markers
    pattern = rf'\d+(?:\.\d+)?\s*(?:{_NUMBER_SPACE_UNITS})?'
    for match in re.findall(pattern, raw, flags=re.IGNORECASE):
        normalized = re.sub(r'\s+', '', str(match).lower())
        if normalized:
            markers.add(normalized)
    return markers


def _extract_ui_markers(text: str) -> set[str]:
    markers = set()
    raw = str(text or '')
    for match in re.findall(r'【[^】]{1,20}】|[\u4e00-\u9fffA-Za-z0-9\-]{2,20}?(?:界面|页面|窗口|弹窗|按钮|图标|下拉框|下拉箭头|下拉列表|列表)', raw):
        normalized = _normalize_ui_object(match)
        if normalized:
            markers.add(normalized)
    for pair in _extract_action_object_units(text):
        verb = pair.get('verb', '')
        obj = pair.get('object', '')
        if verb not in {'点击', '选择', '打开', '关闭', '进入', '切换到', '输入', '运行', '启动'}:
            continue
        normalized = _normalize_ui_object(obj)
        if normalized:
            markers.add(normalized)
    return markers


def _extract_model_markers(text: str) -> dict[str, set[str]]:
    markers = {}
    raw = str(text or '')
    if not raw:
        return markers
    for match in re.findall(r'[A-Za-z]+(?:-[A-Za-z0-9]+(?:\s+[A-Za-z0-9]+)*)+', raw):
        compact = re.sub(r'\s+', '', _protect_model_numbers(match).strip()).lower()
        if not compact:
            continue
        markers.setdefault(compact, set()).add(' '.join(match.split()))
    return markers


def _has_inconsistent_model_spacing(sentence: str, template: str) -> bool:
    source_markers = _extract_model_markers(sentence)
    target_markers = _extract_model_markers(template)
    if not source_markers or not target_markers:
        return False
    for key in source_markers.keys() & target_markers.keys():
        source_forms = source_markers.get(key, set())
        target_forms = target_markers.get(key, set())
        if source_forms == target_forms:
            continue
        if any(' ' in value for value in source_forms | target_forms):
            return True
    return False


def _looks_like_truncated_template(sentence: str, template: str) -> bool:
    sentence_norm = _normalize_sentence_for_match(sentence)
    template_norm = _normalize_sentence_for_match(template)
    if not sentence_norm or not template_norm:
        return False
    if len(sentence_norm) < 20 or len(template_norm) >= len(sentence_norm):
        return False
    if not sentence_norm.startswith(template_norm):
        return False
    if len(template_norm) >= int(len(sentence_norm) * 0.9):
        return False
    return not re.search(r'[。！？.!?]$', str(template or '').strip())


def _source_marker_coverage_ratio(sentence: str, template: str) -> float:
    source_markers = _extract_candidate_topic_markers(sentence)
    target_markers = _extract_candidate_topic_markers(template)
    if not source_markers or not target_markers:
        return 0.0
    return len(source_markers & target_markers) / max(len(source_markers), 1)


def _has_excessive_template_expansion(sentence: str, template: str) -> bool:
    sentence_norm = _normalize_sentence_for_match(sentence)
    template_norm = _normalize_sentence_for_match(template)
    if not sentence_norm or not template_norm:
        return False
    if len(template_norm) <= len(sentence_norm):
        return False

    source_clauses = _split_sentence_clauses(sentence)
    target_clauses = _split_sentence_clauses(template)
    source_markers = _extract_candidate_topic_markers(sentence)
    target_markers = _extract_candidate_topic_markers(template)
    extra_markers = target_markers - source_markers
    marker_coverage = _source_marker_coverage_ratio(sentence, template)
    structured = _structured_match_score(sentence, template)

    if (
        len(source_clauses) >= 2 and
        len(target_clauses) >= len(source_clauses) + 1 and
        len(extra_markers) >= 2 and
        marker_coverage < 0.6 and
        structured.get('pair', 0.0) < 0.75
    ):
        return True

    if (
        len(source_clauses) == 1 and
        len(target_clauses) >= 2 and
        len(template_norm) >= int(len(sentence_norm) * 1.45) and
        len(extra_markers) >= 2 and
        structured.get('total', 0.0) < 0.7
    ):
        return True

    if (
        len(template_norm) >= int(len(sentence_norm) * 1.35) and
        len(extra_markers) >= 2 and
        marker_coverage < 0.5 and
        structured.get('total', 0.0) < 0.6
    ):
        return True

    return False


def _has_insufficient_step_coverage(sentence: str, template: str, structured: dict) -> bool:
    source_clauses = _split_sentence_clauses(sentence)
    target_clauses = _split_sentence_clauses(template)
    if len(source_clauses) < 2:
        return False
    if len(target_clauses) >= len(source_clauses):
        return False
    source_pairs = structured.get('source', {}).get('pairs', [])
    target_pairs = structured.get('target', {}).get('pairs', [])
    sequence_score = structured.get('sequence', 0.0)
    if len(source_pairs) >= 2 and len(target_pairs) < len(source_pairs) and sequence_score < 0.75:
        return True
    tail_clauses = source_clauses[len(target_clauses):]
    if not tail_clauses:
        tail_clauses = source_clauses[1:]
    if not tail_clauses:
        return False
    best_tail_score = max((_clause_similarity(clause, template) for clause in tail_clauses), default=0.0)
    return best_tail_score < 0.45


def _has_missing_additional_context(sentence: str, template: str) -> bool:
    source_units = _extract_additional_units(sentence)
    if not source_units:
        return False
    target_units = _extract_additional_units(template)
    return _token_overlap_score(source_units, target_units) < 0.4


def _number_difference_severity(source_numbers: set, target_numbers: set) -> float:
    """评估数字/单位差异的严重程度，返回 0.0~1.0 之间的渐变值。"""
    if not source_numbers or not target_numbers:
        return 0.0
    if source_numbers & target_numbers:
        return 0.0
    total = len(source_numbers | target_numbers)
    if total <= 1:
        return 0.4
    if total <= 3:
        return 0.7
    return 1.0


def _collect_match_penalties(sentence: str, template: str, structured_score: dict = None) -> dict:
    structured = structured_score or _structured_match_score(sentence, template)
    penalty = 0.0
    reasons = []

    source_numbers = _extract_numeric_markers(sentence)
    target_numbers = _extract_numeric_markers(template)
    if source_numbers and target_numbers and not (source_numbers & target_numbers):
        severity = _number_difference_severity(source_numbers, target_numbers)
        penalty += 0.15 * severity
        reasons.append('数字/单位不一致')

    source_ui = _extract_ui_markers(sentence)
    target_ui = _extract_ui_markers(template)
    if source_ui and target_ui and not (source_ui & target_ui):
        penalty += 0.05
        reasons.append('UI 控件不一致')

    source_actions = structured['source'].get('actions', [])
    target_actions = structured['target'].get('actions', [])
    source_objects = structured['source'].get('objects', [])
    target_objects = structured['target'].get('objects', [])
    platform_placeholder_match = (
        '使用' in source_actions and
        '使用' in target_actions and
        any(_is_placeholder_platform_object(item) for item in source_objects + target_objects)
    )
    if source_actions and target_actions and structured.get('action', 0.0) < 0.5:
        penalty += 0.12
        reasons.append('动作不一致')

    if source_objects and target_objects and structured.get('object', 0.0) < 0.45 and not platform_placeholder_match:
        penalty += 0.18
        reasons.append('对象不一致')

    if not platform_placeholder_match and _has_conflicting_action_object_pair(structured['source'].get('pairs', []), structured['target'].get('pairs', [])):
        penalty += 0.1
        reasons.append('动作对象冲突')

    if _has_insufficient_step_coverage(sentence, template, structured):
        penalty += 0.16
        reasons.append('步骤覆盖不足')

    if _has_missing_additional_context(sentence, template):
        penalty += 0.08
        reasons.append('附加信息缺失')

    marker_coverage = _source_marker_coverage_ratio(sentence, template)
    if (
        len(_extract_candidate_topic_markers(sentence)) >= 2 and
        marker_coverage < 0.34 and
        structured.get('pair', 0.0) < 0.7 and
        _term_anchor_score(sentence, template, structured) < 0.18 and
        _context_match_score(sentence, template, None, structured) < 0.25
    ):
        penalty += 0.14
        reasons.append('主题锚点覆盖不足')

    if _has_inconsistent_model_spacing(sentence, template):
        penalty += 0.10
        reasons.append('产品型号空格不一致')

    if _looks_like_truncated_template(sentence, template):
        penalty += 0.15
        reasons.append('模板疑似残句')

    return {
        'score_penalty': min(penalty, 0.35),
        'reasons': reasons,
    }


def _extract_sentence_structure(sentence: str) -> dict:
    """提取中文技术文档句子的主谓宾结构。"""
    struct = {"verb": "", "subject": "", "object": "", "pattern": ""}
    text = sentence.strip()
    # 模式0: "X由Y组成/构成/包装" 或 "本产品由Y组成"
    m = re.search(r'(.+?)(?:由)(.+?)(组成|构成|包装组成|包装构成|包装)$', text.rstrip('。；;：: '))
    if m:
        struct["verb"] = '组成'
        struct["subject"] = _normalize_product_subject(m.group(1))
        struct["object"] = m.group(2).strip()
        struct["pattern"] = "X由Y组成"
        return struct
    # 模式1: "将A置于B上/中" 或 "将A放置于B"
    m = re.search(r'将(.+?)(置于|放置于|放置在|放入|装到|插入|转移到|拨至|倒[入进])([^，。；！？]+)', text)
    if m:
        raw_object = m.group(3).strip()
        object_match = re.match(r'(.+?(?:位置|[上中下内外]))(?:[)）】\].]*)?$', raw_object)
        if object_match:
            raw_object = object_match.group(1)
        struct["verb"] = m.group(2)
        struct["subject"] = m.group(1)
        struct["object"] = raw_object
        struct["pattern"] = "将X置于Y"
        return struct
    # 模式2: "确保A与B一致/对应" 或 "检查A与B匹配"
    m = re.search(r'(确保|检查|验证|确认)(.+?)(与|和)(.+?)(一致|匹配|对应|对齐|相同)', text)
    if m:
        struct["verb"] = m.group(1)
        struct["subject"] = m.group(2)
        struct["object"] = m.group(4)
        struct["pattern"] = "X与Y一致"
        return struct
    # 模式3: "用/使用A做B"
    m = re.search(r'(用|使用|利用|通过)(.+?)(进行|完成|实现|执行)(.+)', text)
    if m:
        struct["verb"] = m.group(3)
        struct["subject"] = m.group(2)
        struct["object"] = m.group(4)
        struct["pattern"] = "用X做Y"
        return struct
    # 模式4: "待A后，B" 或 "当A时，B"
    m = re.search(r'(?:待|当|等到)(.+?)(?:后|时|之后)(.+)', text)
    if m:
        struct["verb"] = "等待"
        struct["subject"] = m.group(1)
        struct["object"] = m.group(2)
        struct["pattern"] = "待X后做Y"
        return struct
    # 模式5: "点击/选择/打开/关闭A"
    m = re.search(r'(点击|选择|打开|关闭|进入|退出|切换到)(.+)', text)
    if m:
        struct["verb"] = m.group(1)
        struct["object"] = m.group(2)
        struct["pattern"] = "操作UI元素"
        return struct
    return struct


def _compare_sentence_structure(a: str, b: str) -> float:
    """比较两个句子的主谓宾结构相似度。"""
    sa = _extract_sentence_structure(a)
    sb = _extract_sentence_structure(b)
    if not sa["pattern"] or not sb["pattern"]:
        return 0.5  # 无法提取结构时给中性分数
    if sa["pattern"] != sb["pattern"]:
        return 0.0  # 结构模式不同
    score = 1.0
    if sa["verb"] and sb["verb"]:
        if sa["verb"] != sb["verb"]:
            score -= 0.3
    return max(score, 0.0)


def _normalize_ui_object(text: str) -> str:
    value = re.sub(r'^[\s，。；：、,.!?;:()（）【】\[\]<>《》]+', '', text or '')
    value = re.sub(r'[\s，。；：、,.!?;:()（）【】\[\]<>《》]+$', '', value)
    if any(token in value for token in ('下拉框', '下拉箭头', '下拉列表', '弹出的列表', '弹出列表')):
        return '下拉控件'
    bracket_match = re.search(r'【([^】]{1,20})】', value)
    if bracket_match:
        return _normalize_sentence_for_match(bracket_match.group(1))
    specific_matches = [
        match for match in re.findall(r'[\u4e00-\u9fffA-Za-z0-9\-]{2,20}?(?:界面|页面|窗口|弹窗|按钮|图标|下拉框|下拉箭头|下拉列表|列表)', value)
        if not re.search(r'(点击|选择|打开|关闭|进入|切换|输入|运行|启动)', match)
    ]
    if specific_matches:
        value = max(specific_matches, key=len)
    elif re.search(r'(点击|选择|打开|关闭|进入|切换|输入|运行|启动)', value):
        return ''
    if value.startswith('在') and value.endswith(('界面', '页面', '窗口')):
        value = value[1:]
    normalized = _normalize_sentence_for_match(value)
    return '' if normalized in _GENERIC_UI_MARKERS else normalized


def _ui_action_object_score(sentence: str, template: str) -> float:
    source_struct = _extract_sentence_structure(sentence)
    template_struct = _extract_sentence_structure(template)
    if source_struct.get("pattern") != "操作UI元素":
        return 1.0
    if template_struct.get("pattern") != "操作UI元素":
        return 1.0

    source_object = _normalize_ui_object(source_struct.get("object", ""))
    template_object = _normalize_ui_object(template_struct.get("object", ""))

    if not template_object:
        return 0.0
    if source_struct.get("verb") and template_struct.get("verb") and source_struct.get("verb") != template_struct.get("verb"):
        return 0.0
    if not source_object:
        return 0.0

    return max(
        _lcs_ratio(source_object, template_object),
        _compare_semantic_keywords(source_object, template_object),
    )


def _generate_sentence_variants(sentence: str) -> list[str]:
    """为精确匹配生成句子的多种变体。"""
    variants = [sentence]
    text = sentence.strip()
    text_without_prefix = re.sub(r'^((?:\d+[.、)]?)+(?:\s+|(?=[\u4e00-\u9fffA-Za-z(（])))\s*', '', text)
    if text_without_prefix and text_without_prefix != text:
        variants.append(text_without_prefix)
    canonical = _canonicalize_sentence_for_template_match(text)
    if canonical and canonical != text:
        variants.append(canonical)
    canonical_without_prefix = _canonicalize_sentence_for_template_match(text_without_prefix)
    if canonical_without_prefix and canonical_without_prefix != canonical:
        variants.append(canonical_without_prefix)
    # 去掉标点
    no_punct = re.sub(r'[，。！？；：、,.!?;:""''()（）【】\[\]<>《》]+', '', text)
    variants.append(no_punct)
    canonical_no_punct = re.sub(r'[，。！？；：、,.!?;:""''()（）【】\[\]<>《》]+', '', canonical)
    if canonical_no_punct and canonical_no_punct != no_punct:
        variants.append(canonical_no_punct)
    no_prefix_no_punct = re.sub(r'[，。！？；：、,.!?;:""''()（）【】\[\]<>《》]+', '', text_without_prefix)
    if no_prefix_no_punct and no_prefix_no_punct != no_punct:
        variants.append(no_prefix_no_punct)
    canonical_without_prefix_no_punct = re.sub(r'[，。！？；：、,.!?;:""''()（）【】\[\]<>《》]+', '', canonical_without_prefix)
    if canonical_without_prefix_no_punct and canonical_without_prefix_no_punct != canonical_no_punct:
        variants.append(canonical_without_prefix_no_punct)
    # 统一空格
    unified = ' '.join(text.split())
    if unified != text:
        variants.append(unified)
    # 同义词变体：将常见操作同义词替换后生成变体
    _SYNONYM_VARIANTS = {
        '点击': '单击', '单击': '点击',
        '打开': '开启', '开启': '打开',
        '关闭': '停止', '停止': '关闭',
        '进入': '切换到', '切换到': '进入',
        '放入': '放置于', '放置于': '放入',
        '取出': '拿出', '拿出': '取出',
    }
    for src, dst in _SYNONYM_VARIANTS.items():
        if src in text:
            variant = text.replace(src, dst)
            if variant not in variants:
                variants.append(variant)
    return list(set(variants))


# 约束润色器配置
_CONSTRAINT_TERMINOLOGY = {
    "机器": "仪器",
    "推板": "载台",
    "平置": "水平放置",
    "探测器": "检测器",
    "分离柱": "色谱柱",
    "底线": "基线",
    "注射": "进样",
    "滞留时间": "保留时间",
}

_MATCH_CANONICAL_REPLACEMENTS = {
    "机器推板": "载台",
    "仪器载台": "载台",
    "水平放置于": "水平置于",
    "平置于": "水平置于",
    "平置在": "水平置于",
    "放置于": "置于",
    "放置在": "置于",
    "机器": "载台",
    "仪器": "载台",
}


def _canonicalize_sentence_for_template_match(sentence: str) -> str:
    text = sentence or ''
    for source, target in sorted(_MATCH_CANONICAL_REPLACEMENTS.items(), key=lambda item: len(item[0]), reverse=True):
        text = text.replace(source, target)
    return text

_COLLOQUIAL_PATTERNS = [
    (r'一下', ''),
    (r'的话', ''),
]

_SYNTAX_OPTIMIZATIONS = [
    (r'要与(.+?)对应', r'确保与\1一致'),
]

_NUMBER_SPACE_UNITS = r'r/min|kHz|MHz|GHz|kPa|MPa|kVA|mA|mV|kV|kW|MW|kN|min|rpm|bp|mL|μL|µL|uL|μg|µg|mg|ng|kg|μm|µm|mm|cm|°C|℃|Hz|Pa|VA|mM|μM|µM|nM|m|L|g|V|A|W|N|s|h|M'


def _rewrite_field_style_sentence(sentence: str) -> str:
    text = (sentence or '').strip()
    if not text or any(punct in text for punct in '，,；;！？!?'):
        return sentence

    match = re.match(r'^(?P<left>[^：:。]{1,20}?)\s*为\s*(?P<right>[^。]{1,40})$', text.rstrip('。'))
    if not match:
        return sentence

    left = match.group('left').strip()
    right = match.group('right').strip()
    if any(left.startswith(prefix) for prefix in ('每', '各', '本', '该', '此', '其', '所有', '全部')):
        return sentence
    if left.endswith(('即', '则', '后', '前', '时')):
        return sentence
    if any(token in left for token in ('的', '即', '则', '并', '且')):
        return sentence
    if any(token in left for token in ('将', '把', '用', '可', '能', '会', '需', '应', '必须', '建议', '禁止', '避免', '用于', '可以', '能够')):
        return sentence
    if any(token in left for token in ('如果', '则', '表示', '说明', '是否', '目测', '应', '需', '需要', '必须')):
        return sentence
    if re.fullmatch(r'[A-Za-z](?:区)?', left):
        return sentence
    if re.search(r'[A-Za-z](?:\s*[,/、]\s*[A-Za-z])+', left):
        return sentence
    if any(token in left for token in ('区', '孔', '槽', '位')) and any(token in right for token in ('加载区', '回收区', '样本区', '试剂区')):
        return sentence
    if not re.fullmatch(r'[\u4e00-\u9fffA-Za-z0-9()（）【】\[\]<>《》%/\-\s]{1,20}', left):
        return sentence
    return f'{left}：{right}'


def _extract_trailing_figure_ref(sentence: str) -> tuple[str, str]:
    text = (sentence or '').strip().rstrip('。.!！？?')
    match = re.search(r'(\s*[（(][^()（）]*(?:图|表)[^()（）]*[）)])\s*$', text)
    if not match:
        return text, ''
    return text[:match.start()].rstrip(), match.group(1).strip()


def _extract_result_units(sentence: str) -> list[str]:
    results = []
    raw = str(sentence or '').strip()
    if not raw:
        return results

    clauses = _split_sentence_clauses(raw)
    result_markers = ('确保', '确认', '检查', '验证', '判断', '一致', '对应', '匹配', '对齐', '相同', '稳定', '正常')
    for clause in clauses:
        clause_text = clause.strip('，,。；;：: ')
        normalized = _normalize_structure_text(clause_text)
        if not normalized:
            continue
        if any(marker in normalized for marker in result_markers):
            results.append(normalized)
            continue
        align_match = re.search(r'(.+?)(与|和)(.+?)(一致|对应|匹配|对齐|相同)', normalized)
        if align_match:
            results.append(normalized)
    return list(dict.fromkeys(results))


def _extract_additional_units(sentence: str) -> list[str]:
    results = []
    raw = str(sentence or '')
    if not raw:
        return results

    for match in re.findall(r'[（(][^()（）]*(?:图|表|附录|步骤|页)[^()（）]*[）)]', raw):
        value = _normalize_structure_text(match)
        if value:
            results.append(value)
    return list(dict.fromkeys(results))


def _segment_score_entry(key: str, source_units: list[str], target_units: list[str], score: float) -> dict:
    source_text = '；'.join(source_units)
    target_text = '；'.join(target_units)
    has_source = bool(source_units)
    has_target = bool(target_units)
    applicable = has_source or has_target
    final_score = 0.0 if not applicable else max(0.0, min(score, 1.0))

    if not applicable:
        reason = '无此分段'
    elif has_source and not has_target:
        reason = f'{source_text} -> 未匹配'
    elif not has_source and has_target:
        reason = f'无此分段 -> {target_text}'
    else:
        reason = f'{source_text} -> {target_text}'

    return {
        'key': key,
        'label': _MATCH_SEGMENT_LABELS[key],
        'weight': int(_MATCH_SEGMENT_WEIGHTS[key] * 100),
        'score': round(final_score, 4),
        'percent': int(round(final_score * 100)),
        'applicable': applicable,
        'source_text': source_text,
        'target_text': target_text,
        'reason': reason,
    }


def _match_band_label(overall_percent: int) -> tuple[str, str]:
    if overall_percent >= 100:
        return '100%', '完全匹配'
    if overall_percent >= 95:
        return '95%-99%', '高置信度，建议人工确认'
    if overall_percent >= 85:
        return '85%-94%', '较高匹配，建议人工确认'
    if overall_percent >= 75:
        return '75%-84%', '部分匹配，建议只参考局部表达'
    if overall_percent >= 50:
        return '50%-74%', '弱匹配，仅作候选'
    return '50%以下', '低匹配，默认折叠'


def _build_segment_scores(sentence: str, template: str, structured_score: dict = None) -> tuple[dict, float]:
    structured = structured_score or _structured_match_score(sentence, template)
    source_struct = structured['source']
    target_struct = structured['target']

    condition_entry = _segment_score_entry(
        'condition',
        source_struct.get('conditions', []),
        target_struct.get('conditions', []),
        structured.get('condition', 0.0),
    )

    action_entry = _segment_score_entry(
        'action',
        source_struct.get('actions', []),
        target_struct.get('actions', []),
        max(structured.get('action', 0.0), structured.get('pair', 0.0)),
    )

    object_entry = _segment_score_entry(
        'object',
        source_struct.get('objects', []),
        target_struct.get('objects', []),
        max(structured.get('object', 0.0), structured.get('pair', 0.0)),
    )

    result_entry = _segment_score_entry(
        'result',
        _extract_result_units(sentence),
        _extract_result_units(template),
        _token_overlap_score(_extract_result_units(sentence), _extract_result_units(template)),
    )

    additional_entry = _segment_score_entry(
        'additional',
        _extract_additional_units(sentence),
        _extract_additional_units(template),
        _token_overlap_score(_extract_additional_units(sentence), _extract_additional_units(template)),
    )

    segment_scores = {
        'condition': condition_entry,
        'action': action_entry,
        'object': object_entry,
        'result': result_entry,
        'additional': additional_entry,
    }

    applicable_weight = sum(
        _MATCH_SEGMENT_WEIGHTS[key]
        for key in _MATCH_SEGMENT_WEIGHTS
        if segment_scores[key]['applicable']
    )
    applicable_weight = max(applicable_weight, 0.50)
    structured_overall = 0.0
    if applicable_weight > 0:
        structured_overall = sum(
            segment_scores[key]['score'] * _MATCH_SEGMENT_WEIGHTS[key]
            for key in _MATCH_SEGMENT_WEIGHTS
            if segment_scores[key]['applicable']
        ) / applicable_weight

    return segment_scores, structured_overall


def _cat_match_score(sentence: str, template: str, template_metadata=None, structured_score: dict = None) -> dict:
    source_norm = _normalize_sentence_for_match(sentence)
    target_norm = _normalize_sentence_for_match(template)
    if not source_norm or not target_norm:
        return {
            'overall_score': 0.0,
            'overall_percent': 0,
            'ranking_score': 0.0,
            'raw_score': 0.0,
            'display_raw_score': 0.0,
            'structured_score': 0.0,
            'edit_distance_score': 0.0,
            'term_anchor_score': 0.0,
            'number_placeholder_score': 0.0,
            'context_score': 0.0,
            'segment_scores': _empty_segment_scores(),
            'penalty_reasons': [],
            'penalty_score': 0.0,
        }

    if source_norm == target_norm:
        segment_scores, structured_overall = _build_segment_scores(sentence, template, structured_score or _structured_match_score(sentence, template))
        return {
            'overall_score': 1.0,
            'overall_percent': 100,
            'ranking_score': 1.0,
            'raw_score': 1.0,
            'display_raw_score': 1.0,
            'structured_score': max(structured_overall, 1.0),
            'edit_distance_score': 1.0,
            'term_anchor_score': 1.0,
            'number_placeholder_score': 1.0,
            'context_score': 1.0,
            'segment_scores': segment_scores,
            'penalty_reasons': [],
            'penalty_score': 0.0,
        }

    structured = structured_score or _structured_match_score(sentence, template)
    segment_scores, structured_overall = _build_segment_scores(sentence, template, structured)
    edit_distance_score = _edit_distance_score(sentence, template)
    term_anchor_score = _term_anchor_score(sentence, template, structured)
    number_placeholder_score = _number_placeholder_score(sentence, template)
    context_score = _context_match_score(sentence, template, template_metadata, structured)

    raw_score = (
        _CAT_MATCH_COMPONENT_WEIGHTS['structured'] * structured_overall +
        _CAT_MATCH_COMPONENT_WEIGHTS['edit_distance'] * edit_distance_score +
        _CAT_MATCH_COMPONENT_WEIGHTS['term_anchor'] * term_anchor_score +
        _CAT_MATCH_COMPONENT_WEIGHTS['number_placeholder'] * number_placeholder_score +
        _CAT_MATCH_COMPONENT_WEIGHTS['context'] * context_score
    )

    display_weight_total = sum(_CAT_MATCH_COMPONENT_WEIGHTS[key] for key in _CAT_MATCH_DISPLAY_COMPONENTS)
    display_raw_score = 0.0
    if display_weight_total > 0:
        display_raw_score = (
            _CAT_MATCH_COMPONENT_WEIGHTS['structured'] * structured_overall +
            _CAT_MATCH_COMPONENT_WEIGHTS['edit_distance'] * edit_distance_score +
            _CAT_MATCH_COMPONENT_WEIGHTS['term_anchor'] * term_anchor_score +
            _CAT_MATCH_COMPONENT_WEIGHTS['number_placeholder'] * number_placeholder_score +
            _CAT_MATCH_COMPONENT_WEIGHTS['context'] * context_score
        ) / display_weight_total

    penalties = _collect_match_penalties(sentence, template, structured)
    ranking_score = min(1.0, max(0.0, raw_score - penalties['score_penalty']))
    overall_score = min(1.0, max(0.0, display_raw_score - penalties['score_penalty']))
    overall_percent = int(round(overall_score * 100))

    return {
        'overall_score': round(overall_score, 4),
        'overall_percent': overall_percent,
        'ranking_score': round(ranking_score, 4),
        'raw_score': round(raw_score, 4),
        'display_raw_score': round(display_raw_score, 4),
        'structured_score': round(structured_overall, 4),
        'edit_distance_score': round(edit_distance_score, 4),
        'term_anchor_score': round(term_anchor_score, 4),
        'number_placeholder_score': round(number_placeholder_score, 4),
        'context_score': round(context_score, 4),
        'segment_scores': segment_scores,
        'penalty_reasons': penalties['reasons'],
        'penalty_score': round(penalties['score_penalty'], 4),
    }


def _score_match_segments(sentence: str, template: str) -> dict:
    cat_score = _cat_match_score(sentence, template)
    overall_percent = cat_score['overall_percent']
    band, label = _match_band_label(overall_percent)

    return {
        'overall_score': cat_score['overall_score'],
        'overall_percent': overall_percent,
        'ranking_score': cat_score['ranking_score'],
        'band': band,
        'label': label,
        'raw_score': cat_score['raw_score'],
        'display_raw_score': cat_score['display_raw_score'],
        'lexical_score': cat_score['edit_distance_score'],
        'structured_score': cat_score['structured_score'],
        'edit_distance_score': cat_score['edit_distance_score'],
        'term_anchor_score': cat_score['term_anchor_score'],
        'number_placeholder_score': cat_score['number_placeholder_score'],
        'context_score': cat_score['context_score'],
        'segment_scores': cat_score['segment_scores'],
        'penalty_reasons': cat_score['penalty_reasons'],
        'penalty_score': cat_score['penalty_score'],
    }


def _build_change_match_detail(
    before: str,
    after: str,
    change_type: str = '',
    rule_name: str = '',
    sentence_guide: str = '',
    is_title: bool = False,
    precomputed_candidates: Optional[list[dict]] = None,
) -> Optional[dict]:
    before_text = _normalize_compare_text(before)
    after_text = _normalize_compare_text(_reapply_sentence_prefix(before, after))
    if not before_text or not after_text or before_text == after_text:
        return None
    if _is_field_style_colon_rewrite(before_text, after_text):
        return None
    if _is_short_label_like_text(before_text) or _is_short_label_like_text(after_text):
        return {
            'overall_score': 0.0,
            'overall_percent': 0,
            'ranking_score': 0.0,
            'band': '50%以下',
            'label': '短标签文本，不参与句式匹配',
            'raw_score': 0.0,
            'display_raw_score': 0.0,
            'lexical_score': 0.0,
            'structured_score': 0.0,
            'edit_distance_score': 0.0,
            'term_anchor_score': 0.0,
            'number_placeholder_score': 0.0,
            'context_score': 0.0,
            'segment_scores': _empty_segment_scores('短标签文本'),
            'penalty_reasons': ['短标签文本'],
            'penalty_score': 0.0,
        }

    normalized_type = str(change_type or '').lower()
    if normalized_type in {'format', 'punctuation'} or rule_name == '基础规范化':
        return None
    if normalized_type in {'terminology', 'term', 'terminology_rule'}:
        detail = _score_match_segments(before_text, after_text)
        detail['label'] = f"{detail.get('label', '较高匹配，建议人工确认')}（术语替换轻量评估）"
        return detail

    if _is_slot_sample_mapping_sentence(before_text) and _is_slot_sample_mapping_sentence(after_text):
        source_pairs = _extract_slot_sample_pairs(before_text)
        target_pairs = _extract_slot_sample_pairs(after_text)
        pair_score = _slot_sample_pair_overlap_score(source_pairs, target_pairs)
        edit_distance_score = _edit_distance_score(before_text, after_text)
        additional_score = _token_overlap_score(_extract_additional_units(before_text), _extract_additional_units(after_text))
        overall_score = round(min(1.0, 0.6 * pair_score + 0.3 * edit_distance_score + 0.1 * additional_score), 4)
        overall_percent = int(round(overall_score * 100))
        band, label = _match_band_label(overall_percent)
        detail = {
            'overall_score': overall_score,
            'overall_percent': overall_percent,
            'ranking_score': overall_score,
            'band': band,
            'label': f'{label}（结构化映射句轻量评估）',
            'raw_score': overall_score,
            'display_raw_score': overall_score,
            'lexical_score': round(edit_distance_score, 4),
            'structured_score': round(pair_score, 4),
            'edit_distance_score': round(edit_distance_score, 4),
            'term_anchor_score': round(pair_score, 4),
            'number_placeholder_score': round(pair_score, 4),
            'context_score': round(additional_score, 4),
            'segment_scores': {
                'condition': _segment_score_entry('condition', [], [], 0.0),
                'action': _segment_score_entry('action', ['对应'], ['对应'], pair_score),
                'object': _segment_score_entry(
                    'object',
                    [f'{slot}->sample {sample}' for slot, sample in source_pairs],
                    [f'{slot}->sample {sample}' for slot, sample in target_pairs],
                    pair_score,
                ),
                'result': _segment_score_entry('result', [], [], 0.0),
                'additional': _segment_score_entry(
                    'additional',
                    _extract_additional_units(before_text),
                    _extract_additional_units(after_text),
                    additional_score,
                ),
            },
            'penalty_reasons': [],
            'penalty_score': 0.0,
        }
    elif _should_skip_expensive_template_match(before_text) and _should_skip_expensive_template_match(after_text):
        detail = _score_match_segments(before_text, after_text)
        detail['label'] = f"{detail.get('label', '较高匹配，建议人工确认')}（操作界面句轻量评估）"
    elif _is_pooling_platform_sentence(before_text) and _is_pooling_platform_sentence(after_text):
        source_markers = _extract_pooling_topic_markers(before_text)
        target_markers = _extract_pooling_topic_markers(after_text)
        marker_score = _pooling_marker_overlap_score(source_markers, target_markers)
        clause_score = _clause_alignment_score(before_text, after_text)
        edit_distance_score = _edit_distance_score(before_text, after_text)
        model_score = _marker_overlap_score(set(_extract_model_markers(before_text).keys()), set(_extract_model_markers(after_text).keys()))
        number_score = _marker_overlap_score(_extract_numeric_markers(before_text), _extract_numeric_markers(after_text))
        overall_score = round(min(1.0, 0.35 * marker_score + 0.25 * clause_score + 0.15 * edit_distance_score + 0.15 * model_score + 0.1 * number_score), 4)
        overall_percent = int(round(overall_score * 100))
        band, label = _match_band_label(overall_percent)
        detail = {
            'overall_score': overall_score,
            'overall_percent': overall_percent,
            'ranking_score': overall_score,
            'band': band,
            'label': f'{label}（混样/平台长句轻量评估）',
            'raw_score': overall_score,
            'display_raw_score': overall_score,
            'lexical_score': round(edit_distance_score, 4),
            'structured_score': round(clause_score, 4),
            'edit_distance_score': round(edit_distance_score, 4),
            'term_anchor_score': round(marker_score, 4),
            'number_placeholder_score': round(number_score, 4),
            'context_score': round(model_score, 4),
            'segment_scores': {
                'condition': _segment_score_entry('condition', _extract_condition_units(before_text), _extract_condition_units(after_text), clause_score),
                'action': _segment_score_entry('action', _extract_sentence_intent(before_text).get('actions', []), _extract_sentence_intent(after_text).get('actions', []), clause_score),
                'object': _segment_score_entry('object', sorted(source_markers), sorted(target_markers), marker_score),
                'result': _segment_score_entry('result', sorted(_extract_numeric_markers(before_text)), sorted(_extract_numeric_markers(after_text)), number_score),
                'additional': _segment_score_entry('additional', sorted(set(_extract_model_markers(before_text).keys())), sorted(set(_extract_model_markers(after_text).keys())), model_score),
            },
            'penalty_reasons': [],
            'penalty_score': 0.0,
        }
    else:
        detail = _score_match_segments(before_text, after_text)
    if sentence_guide and not is_title and not _should_skip_expensive_template_match(before_text):
        candidates = list(precomputed_candidates or _guide_top_template_candidates(
            before_text,
            sentence_guide,
            limit=8,
            single_clause_only=_is_pooling_platform_sentence(before_text),
        ))
        if candidates:
            detail['candidates'] = candidates
    return detail


def _empty_segment_scores(reason: str = '未命中模板') -> dict:
    scores = {}
    for key in _MATCH_SEGMENT_WEIGHTS:
        scores[key] = {
            'key': key,
            'label': _MATCH_SEGMENT_LABELS[key],
            'weight': int(_MATCH_SEGMENT_WEIGHTS[key] * 100),
            'score': 0.0,
            'percent': 0,
            'applicable': False,
            'source_text': '',
            'target_text': '',
            'reason': reason,
        }
    return scores


def _lightweight_no_change_review_detail(label: str = '未触发改动') -> dict:
    return {
        'overall_score': 0.0,
        'overall_percent': 0,
        'ranking_score': 0.0,
        'band': '50%以下',
        'label': label,
        'segment_scores': _empty_segment_scores(label),
        'candidates': [],
        'suggested_text': '',
        'auto_apply_threshold': _DOC_REVIEW_AUTO_APPLY_THRESHOLD,
        'auto_applied': False,
        'review_mode': 'manual',
        'has_change': False,
    }


def _normalize_review_suggestion(
    sentence: str,
    suggested_text: str,
    change_type: str = '',
    rule_name: str = '',
    is_title: bool = False,
) -> str:
    source_text = _normalize_compare_text(sentence)
    effective_is_title = bool(is_title or _looks_like_title_or_noun_phrase(sentence))
    current_suggestion = _normalize_compare_text(_reapply_sentence_prefix(sentence, suggested_text))
    if not current_suggestion or current_suggestion == source_text:
        return current_suggestion
    if _is_low_value_doc_change(sentence, current_suggestion, change_type, rule_name):
        return ''
    if _is_field_style_colon_rewrite(source_text, current_suggestion):
        return ''
    if effective_is_title:
        if _strip_doc_trailing_punctuation(current_suggestion) == _strip_doc_trailing_punctuation(source_text):
            return ''
        source_len = len(_normalize_sentence_for_match(source_text))
        suggestion_len = len(_normalize_sentence_for_match(current_suggestion))
        if source_len and suggestion_len >= source_len * 2:
            return ''
    should_drop_suggestion = _has_conflicting_critical_literals(sentence, current_suggestion)
    if not should_drop_suggestion:
        try:
            from app.utils.instrument_polisher import instrument_polish_engine
            should_drop_suggestion = not instrument_polish_engine.post_protect(sentence, current_suggestion).get('safe', True)
        except Exception:
            should_drop_suggestion = False
    return '' if should_drop_suggestion else current_suggestion


def _should_use_lightweight_review_detail(original_text: str, polished_text: str, is_title: bool = False) -> bool:
    if is_title:
        return False
    normalized = _normalize_sentence_for_match(original_text)
    normalized_polished = _normalize_compare_text(_reapply_sentence_prefix(original_text, polished_text))
    normalized_original = _normalize_compare_text(original_text)
    if polished_text != original_text and normalized_polished != normalized_original:
        return False
    if _should_skip_expensive_template_match(original_text):
        return True
    if _is_pooling_platform_sentence(original_text):
        return True
    if len(normalized) < 40:
        return False
    clause_count = len([part for part in re.split(r'[，。；！？,;!?]+', original_text) if part.strip()])
    return clause_count >= 1


def _is_simple_operation_sentence(sentence: str) -> bool:
    normalized = _normalize_sentence_for_match(sentence)
    if len(normalized) < 12 or len(normalized) > 80:
        return False
    if _is_pooling_platform_sentence(sentence) or _is_slot_sample_mapping_sentence(sentence):
        return False
    actions = [action for action in _extract_sentence_intent(sentence).get('actions', []) if action and action != '等待']
    if not actions:
        return False
    clause_count = len(_split_sentence_clauses(sentence))
    return clause_count >= 2 or len(set(actions)) >= 2


def _build_sentence_review_detail(
    sentence: str,
    suggested_text: str = '',
    sentence_guide: str = '',
    change_type: str = '',
    rule_name: str = '',
    is_title: bool = False,
) -> tuple[dict, str]:
    source_text = _normalize_compare_text(sentence)
    effective_is_title = bool(is_title or _looks_like_title_or_noun_phrase(sentence))
    current_suggestion = _normalize_review_suggestion(sentence, suggested_text, change_type, rule_name, is_title=effective_is_title)
    normalized_change_type = str(change_type or '').lower()
    skip_review_candidates = (
        _should_skip_expensive_template_match(source_text) or
        normalized_change_type in {'terminology', 'term', 'terminology_rule'}
    )
    candidates = []
    if sentence_guide and not effective_is_title and not skip_review_candidates:
        candidates = _guide_top_template_candidates(
            source_text,
            sentence_guide,
            limit=8,
            single_clause_only=_is_pooling_platform_sentence(source_text),
        )

    detail = None
    if current_suggestion and current_suggestion != source_text:
        detail = _build_change_match_detail(
            source_text,
            current_suggestion,
            change_type,
            rule_name,
            sentence_guide,
            is_title=effective_is_title,
            precomputed_candidates=candidates,
        )

    if detail is None and candidates:
        detail = dict(candidates[0])
        detail['candidates'] = candidates

    if detail is None:
        detail = {
            'overall_score': 0.0,
            'overall_percent': 0,
            'band': '50%以下',
            'label': effective_is_title and '标题/短标签，不参与句式匹配' or '低匹配，建议人工编辑',
            'segment_scores': _empty_segment_scores(),
            'candidates': candidates,
        }

    best_guarded_candidate = next((item for item in candidates if item.get('guard_passed', False)), None)
    best_manual_candidate = next(
        (
            item for item in candidates
            if not item.get('guard_passed', False) and int(item.get('overall_percent', 0) or 0) >= _DOC_REVIEW_REFERENCE_CANDIDATE_THRESHOLD
        ),
        None,
    )
    best_reference_candidate = best_guarded_candidate or best_manual_candidate
    best_candidate_detail = dict(best_reference_candidate) if best_reference_candidate else None
    current_percent = int(detail.get('overall_percent', 0) or 0)
    best_candidate_percent = int((best_candidate_detail or {}).get('overall_percent', 0) or 0)
    should_prefer_candidate = (
        best_candidate_detail is not None and
        best_candidate_percent >= 85 and
        current_percent < 50
    )
    if should_prefer_candidate:
        detail = best_candidate_detail
        detail['candidates'] = candidates

    candidate_text = ''
    if best_reference_candidate:
        candidate_text = _normalize_compare_text(best_reference_candidate.get('candidate_text') or best_reference_candidate.get('template') or '')

    current_is_low_value = bool(current_suggestion) and (
        normalized_change_type in {'format', 'punctuation'} or
        _strip_doc_trailing_punctuation(current_suggestion) == _strip_doc_trailing_punctuation(source_text)
    )
    if should_prefer_candidate and candidate_text:
        resolved_suggestion = _reapply_sentence_prefix(sentence, candidate_text)
    elif current_is_low_value and candidate_text and candidate_text != source_text:
        resolved_suggestion = _reapply_sentence_prefix(sentence, candidate_text)
    elif not current_suggestion and best_manual_candidate and candidate_text and candidate_text != source_text:
        resolved_suggestion = _reapply_sentence_prefix(sentence, candidate_text)
    else:
        resolved_suggestion = current_suggestion or _reapply_sentence_prefix(sentence, candidate_text) or source_text
    auto_applied = detail.get('ranking_score', 0) * 100 >= _DOC_REVIEW_AUTO_APPLY_THRESHOLD and resolved_suggestion != source_text
    detail['suggested_text'] = resolved_suggestion
    detail['auto_apply_threshold'] = _DOC_REVIEW_AUTO_APPLY_THRESHOLD
    detail['auto_applied'] = auto_applied
    detail['review_mode'] = 'auto' if auto_applied else 'manual'
    detail['has_change'] = resolved_suggestion != source_text
    return detail, resolved_suggestion


def _should_emit_reference_review_change(original_text: str, polished_text: str, review_after: str, change_type: str, rule_name: str) -> bool:
    suggested = _normalize_compare_text(review_after)
    source = _normalize_compare_text(original_text)
    if not suggested or suggested == source:
        return False
    polished = _normalize_compare_text(polished_text)
    if polished == source:
        return True
    return _is_low_value_doc_change(original_text, polished_text, change_type, rule_name)


def _normalize_ui_button_label(raw: str) -> str:
    text = (raw or '').strip()
    normalized = _normalize_sentence_for_match(text)
    if not normalized:
        return '按钮'

    if 'run' in text.lower() or '运行' in text:
        return '运行按钮'
    if 'start' in text.lower() or '启动' in text:
        return '启动按钮'

    bracket_match = re.search(r'【([^】]+)】', text)
    if bracket_match:
        return f"【{bracket_match.group(1).strip()}】"

    if '下拉箭头' in text:
        return '下拉箭头'
    if any(token in text for token in ('界面', '页面', '窗口', '系统')):
        return text
    if '按钮' in text:
        return text
    if len(normalized) <= 8:
        return text
    return '按钮'


def _rewrite_ui_navigation_sentence(sentence: str) -> str:
    core, figure_ref = _extract_trailing_figure_ref(sentence)
    condition = ''
    rest = core
    condition_match = re.match(r'^(?P<condition>[^，。；！？]*?(?:后|时))(?:[，,])?(?P<rest>.+)$', core)
    if condition_match:
        condition = condition_match.group('condition').strip('，, ')
        rest = condition_match.group('rest').strip()

    nav_match = re.search(r'^(?P<prefix>.*?)(?:点击)(?P<button>.*?)(?:[，,])?进入(?P<target>[^，。；！？]+?(?:界面|页面|窗口))$', rest)
    if not nav_match:
        return sentence

    target = nav_match.group('target').strip()
    if not target:
        return sentence

    prefix = nav_match.group('prefix').strip()
    button_label = _normalize_ui_button_label(nav_match.group('button'))
    context_match = re.search(r'(在[^，。；]*?(?:界面|页面|窗口))', prefix)
    context = context_match.group(1).strip() if context_match else ''

    body = f'{context}点击{button_label}，进入{target}' if context else f'点击{button_label}，进入{target}'
    if condition:
        body = f'{condition}，{body}'
    if figure_ref:
        body = f'{body}{figure_ref}'
    return f'{body}。'


def _rewrite_ui_run_sentence(sentence: str) -> str:
    core, figure_ref = _extract_trailing_figure_ref(sentence)
    condition = ''
    rest = core
    condition_match = re.match(r'^(?P<condition>[^，。；！？]*?(?:后|时))(?:[，,])?(?P<rest>.+)$', core)
    if condition_match:
        condition = condition_match.group('condition').strip('，, ')
        rest = condition_match.group('rest').strip()

    run_match = re.search(r'^(?:点击[^，。；！？]*?(?:按钮|按键|】)?[，,])?点击\s*(?P<label>run|Run|RUN|start|Start|START|运行|启动)\s*(?P<tail>开始实验|开始操作)?$', rest)
    if not run_match:
        return sentence

    button_label = _normalize_ui_button_label(run_match.group('label'))
    tail = run_match.group('tail') or '开始操作'
    body = f'点击{button_label}，{tail}'
    if condition:
        body = f'{condition}，{body}'
    if figure_ref:
        body = f'{body}{figure_ref}'
    return f'{body}。'


def _lightweight_operation_review_suggestion(sentence: str) -> str:
    if not _should_skip_expensive_template_match(sentence):
        return ''
    suggested = str(sentence or '')
    suggested = re.sub(r'录入(?=(?:界面|页面|窗口|信息))', '输入', suggested)
    suggested = re.sub(r'填写(?=(?:界面|页面|窗口|信息))', '输入', suggested)
    if suggested == sentence:
        return ''
    return suggested


def _apply_safe_structural_rewrite(sentence: str) -> str:
    for rewriter in (_rewrite_ui_run_sentence, _rewrite_ui_navigation_sentence):
        rewritten = rewriter(sentence)
        if rewritten != sentence:
            return rewritten
    return sentence


def _normalize_operation_sentence_for_match(sentence: str) -> str:
    raw = str(sentence or '').strip()
    if not raw:
        return ''

    normalized = raw
    normalized = re.sub(r'参数确认无误后(?=点击)', '参数确认无误后，', normalized)
    normalized = re.sub(r'(?<![，。；！？,;!?])并在弹窗中点击', '，在弹窗中点击', normalized)
    normalized = re.sub(r'(?<![，。；！？,;!?])后点击', '后，点击', normalized)
    normalized = re.sub(r'点击\s*(确定|取消|运行|启动)(?![按钮】])', lambda m: f"点击【{m.group(1)}】" if m.group(1) in {'确定', '取消'} else f"点击{m.group(1)}按钮", normalized)
    normalized = re.sub(r'】(?=开始)', '】，', normalized)
    normalized = re.sub(r'按钮(?=开始)', '按钮，', normalized)
    normalized = re.sub(r'(?<![，。；！？,;!?])开始建库', '，开始建库', normalized)
    normalized = re.sub(r'[，,]{2,}', '，', normalized)
    return normalized.strip('，, ')


def _apply_constraint_polish(sentence: str) -> str:
    """对未匹配到句式的句子进行约束润色。"""
    result = _apply_safe_structural_rewrite(sentence)
    # 术语标准化
    for non_std, std in _CONSTRAINT_TERMINOLOGY.items():
        result = result.replace(non_std, std)
    # 删除口语化冗余
    for pattern, replacement in _COLLOQUIAL_PATTERNS:
        result = re.sub(pattern, replacement, result)
    # 句式优化
    for pattern, replacement in _SYNTAX_OPTIMIZATIONS:
        result = re.sub(pattern, replacement, result)
    return result


def _should_skip_expensive_template_match(sentence: str) -> bool:
    normalized = _normalize_operation_sentence_for_match(sentence)
    if len(_normalize_sentence_for_match(normalized)) < 24:
        return False

    skip_patterns = [
        r'点击.+界面上\s*按钮[，,]\s*进入.+(?:界面|页面|窗口)',
        r'打开.+电源[，,].*输入用户名和密码.*进入.+(?:界面|页面|窗口)',
        r'扫码枪扫描.+二维码[，,].*(?:输入|录入).+信息',
    ]
    return any(re.search(pattern, normalized) for pattern in skip_patterns)


# 匹配策略配置
_POLISH_MATCH_CONFIG = {
    "l2_auto_replace": True,      # L2 模糊匹配自动替换
    "l3_auto_replace": True,      # L3 语义匹配自动替换
    "l1_confidence": 1.0,         # L1 精确匹配置信度
    "l2_min_confidence": 0.85,    # L2 自动替换置信度
    "l2_review_confidence": 0.70, # L2 手动审阅阈值 (0.70-0.85)
    "l2_hint_confidence": 0.55,   # L2 作为 AI hint 阈值 (0.55-0.70)
    "l3_min_confidence": 0.90,    # L3 语义匹配最低置信度
    "l3_auto_confidence": 0.65,   # L3 普通规则匹配阈值（preferred_sentences）
}

_AI_TEMPLATE_RERANK_LIMIT = 4
_ai_template_rerank_cache: dict = {}

_CAT_MATCH_COMPONENT_WEIGHTS = {
    'structured': 0.45,
    'edit_distance': 0.15,
    'term_anchor': 0.20,
    'number_placeholder': 0.15,
    'context': 0.05,
}

_CAT_MATCH_DISPLAY_COMPONENTS = ('structured', 'edit_distance', 'term_anchor', 'number_placeholder', 'context')


def _is_valid_platform_feedback_sentence(text: str) -> bool:
    raw = str(text or '').strip()
    if not raw:
        return False
    normalized = _normalize_sentence_for_match(raw)
    if len(normalized) < 8:
        return False
    if any(token in raw for token in ('例如', '不要写成', '此类标签', '此类参数', '句尾不加句号')):
        return False
    if re.match(r'^(有效期|存储条件|规格)\s*[：:]', raw):
        return False
    if _has_inconsistent_model_spacing(raw, _protect_model_numbers(raw)):
        return False
    if len(normalized) >= 18 and not re.search(r'[。！？.!?]$', raw):
        return False
    return True


def _sanitize_platform_feedback_guide(content: str) -> str:
    sanitized_lines = []
    for line in (content or '').splitlines():
        stripped = line.strip()
        if stripped.startswith('- '):
            sentence = stripped[2:].strip()
            if not _is_valid_platform_feedback_sentence(sentence):
                continue
            sanitized_lines.append(f'- {_protect_model_numbers(sentence)}')
            continue
        sanitized_lines.append(line)
    return '\n'.join(sanitized_lines)


def _prepare_sentence_guide_content(file_name: str, content: str) -> str:
    if not content:
        return ''
    if (file_name or '').strip() == PLATFORM_FEEDBACK_FILENAME:
        return _sanitize_platform_feedback_guide(_unescape_markdown_inline(content))
    return _unescape_markdown_inline(content)


def _fallback_guide_entries_from_text(guide_text: str) -> list:
    entries = []
    seen = set()
    for raw_line in str(guide_text or '').splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith('#') or line.startswith('|') or re.fullmatch(r'[-:| ]+', line):
            continue
        line = re.sub(r'^(?:[-*•·]+\s*|\d+[.)、]\s*|[一二三四五六七八九十]+[、.]\s*)', '', line).strip()
        if not line:
            continue
        sentences = re.findall(r'[^。！？；!?;]+[。！？；!?;]', line)
        if not sentences and len(_normalize_sentence_for_match(line)) >= 8:
            sentences = [line]
        for sentence in sentences:
            text = _protect_model_numbers(_unescape_markdown_inline(sentence.strip()))
            normalized = _normalize_sentence_for_match(text)
            if len(normalized) < 8:
                continue
            if len(normalized) >= 15 and not re.search(r'[。！？；!?;]$', text):
                continue
            if text in seen:
                continue
            seen.add(text)
            entries.append(text)
    return entries

def _normalize_document_product_type(product_type: Optional[str]) -> Optional[str]:
    value = str(product_type or '').strip()
    return value if value in DOCUMENT_PRODUCT_TYPES else None


def _collect_sentence_guide_folder_ids(db: Session, root_folder_ids: set[int]) -> set[int]:
    all_folder_ids = set(root_folder_ids)
    stack = list(root_folder_ids)
    while stack:
        fid = stack.pop()
        subfolders = db.query(Folder).filter(Folder.parent_id == fid).all()
        for sf in subfolders:
            if sf.id in all_folder_ids:
                continue
            all_folder_ids.add(sf.id)
            stack.append(sf.id)
    return all_folder_ids


def _resolve_product_root_folder_ids(db: Session, base_folder_ids: set[int], product_type: Optional[str] = None) -> set[int]:
    normalized_product_type = _normalize_document_product_type(product_type)
    if not normalized_product_type:
        return set(base_folder_ids)

    folder_ids = {
        folder.id
        for folder in db.query(Folder).filter(
            Folder.parent_id.in_(base_folder_ids),
            Folder.name == normalized_product_type,
        ).all()
    }
    return folder_ids


def _resolve_sentence_guide_root_folder_ids(db: Session, product_type: Optional[str] = None) -> set[int]:
    return _resolve_product_root_folder_ids(db, set(SENTENCE_GUIDE_FOLDER_IDS), product_type)


def _resolve_terminology_root_folder_ids(db: Session, product_type: Optional[str] = None) -> set[int]:
    return _resolve_product_root_folder_ids(db, set(TERMINOLOGY_FOLDER_IDS), product_type)


def _load_scoped_terminology_source(db: Session, terminology_id: int = None, product_type: Optional[str] = None) -> Optional[str]:
    source = _load_terminology_source(db, terminology_id)
    if source or terminology_id:
        return source

    normalized_product_type = _normalize_document_product_type(product_type)
    if not normalized_product_type:
        return None

    root_folder_ids = _resolve_terminology_root_folder_ids(db, normalized_product_type)
    all_folder_ids = _collect_sentence_guide_folder_ids(db, root_folder_ids)
    if not all_folder_ids:
        return None

    files = db.query(KnowledgeFile).filter(KnowledgeFile.folder_id.in_(all_folder_ids)).order_by(KnowledgeFile.id.asc()).all()
    sources = []
    loaded_paths = set()
    for term_file in files:
        path = str(term_file.file_path or '').strip()
        if not path or not os.path.exists(path) or path in loaded_paths:
            continue
        loaded_paths.add(path)
        if path.lower().endswith('.xlsx'):
            sources.append(path)
        else:
            content = _read_file_safe(path)
            if content:
                sources.append(content)

    return "\n\n".join(part for part in sources if part) if sources else None


def _load_sentence_guides(db: Session, style_guide_id: int = None, product_type: Optional[str] = None) -> str:
    """加载句式清单内容（带缓存）。

    若指定了 style_guide_id，加载该文件并拼接平台反馈句式清单；
    否则递归加载句式清单文件夹下所有 .md 文件。
    """
    normalized_product_type = _normalize_document_product_type(product_type)
    if style_guide_id:
        cache_key = ('file', int(style_guide_id))
    elif normalized_product_type:
        cache_key = ('product_type', normalized_product_type)
    else:
        cache_key = '__all__'
    if cache_key in _sentence_guide_cache:
        return _sentence_guide_cache[cache_key]

    if style_guide_id:
        guides = []
        loaded_paths = set()
        style_file = db.query(KnowledgeFile).filter(KnowledgeFile.id == style_guide_id).first()
        if style_file and style_file.file_path and os.path.exists(style_file.file_path):
            try:
                content = _prepare_sentence_guide_content(style_file.name, _read_knowledge_text_content(style_file.file_path))
                if content.strip():
                    guides.append(content)
                    loaded_paths.add(style_file.file_path)
            except Exception as e:
                print(f"加载句式文件失败 (id={style_guide_id}): {e}")

        for platform_file in _get_platform_feedback_targets(db, 1):
            if not platform_file or not platform_file.file_path or platform_file.file_path in loaded_paths:
                continue
            if not os.path.exists(platform_file.file_path):
                continue
            try:
                content = _prepare_sentence_guide_content(platform_file.name, _read_knowledge_text_content(platform_file.file_path))
                if content.strip():
                    guides.append(content)
                    loaded_paths.add(platform_file.file_path)
            except Exception:
                pass

        result = "\n\n".join(guides) if guides else None
        _sentence_guide_cache[cache_key] = result
        return result

    # 未指定文件，递归加载句式清单文件夹下所有 .md
    guides = []
    # 一次性收集所有相关文件夹 ID 及其后代
    root_folder_ids = _resolve_sentence_guide_root_folder_ids(db, normalized_product_type)
    all_folder_ids = _collect_sentence_guide_folder_ids(db, root_folder_ids)

    files = []
    if all_folder_ids:
        files = db.query(KnowledgeFile).filter(
            KnowledgeFile.folder_id.in_(all_folder_ids),
            KnowledgeFile.file_type == "md"
        ).all()

    def _guide_priority(kf: KnowledgeFile) -> tuple[int, str, int]:
        name = (kf.name or '').strip()
        if name in PRIMARY_SENTENCE_GUIDE_FILENAMES:
            return (0, name, kf.id or 0)
        if name == PLATFORM_FEEDBACK_FILENAME:
            return (2, name, kf.id or 0)
        if name in REFERENCE_SENTENCE_GUIDE_FILENAMES:
            return (3, name, kf.id or 0)
        return (1, name, kf.id or 0)

    files.sort(key=_guide_priority)

    loaded_paths = set()

    for kf in files:
        if kf.file_path and os.path.exists(kf.file_path):
            try:
                content = _prepare_sentence_guide_content(kf.name, _read_knowledge_text_content(kf.file_path))
                if content.strip():
                    guides.append(content)
                    loaded_paths.add(kf.file_path)
            except Exception:
                pass

    for platform_file in _get_platform_feedback_targets(db, 1):
        if not platform_file or not platform_file.file_path or platform_file.file_path in loaded_paths:
            continue
        if not os.path.exists(platform_file.file_path):
            continue
        try:
            content = _prepare_sentence_guide_content(platform_file.name, _read_knowledge_text_content(platform_file.file_path))
            if content.strip():
                guides.append(content)
        except Exception:
            pass

    app_root = os.path.dirname(os.path.dirname(__file__))
    for relative_path in BUNDLED_SENTENCE_GUIDE_RELATIVE_PATHS:
        bundled_path = os.path.join(app_root, relative_path)
        if bundled_path in loaded_paths or not os.path.exists(bundled_path):
            continue
        try:
            content = _prepare_sentence_guide_content(os.path.basename(bundled_path), _read_knowledge_text_content(bundled_path))
            if content.strip():
                guides.append(content)
                loaded_paths.add(bundled_path)
        except Exception:
            pass

    result = "\n\n".join(guides) if guides else None
    _sentence_guide_cache[cache_key] = result
    return result


def _ensure_platform_feedback_sentence_file(db: Session, user_id: int) -> KnowledgeFile:
    """确保平台反馈句式清单存在于知识库中。"""
    folder_id = None
    knowledge_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "knowledge")
    if not os.path.exists(knowledge_dir):
        os.makedirs(knowledge_dir)

    file_path = os.path.join(knowledge_dir, PLATFORM_FEEDBACK_SENTENCE_RELATIVE_PATH)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    initial_content = "# 平台反馈的句式清单\n\n## 用户反馈修正\n\n"

    feedback_file = db.query(KnowledgeFile).filter(
        KnowledgeFile.folder_id == folder_id,
        KnowledgeFile.name == PLATFORM_FEEDBACK_FILENAME,
        KnowledgeFile.file_path == file_path
    ).first()

    if not feedback_file:
        feedback_file = db.query(KnowledgeFile).filter(
            KnowledgeFile.folder_id == folder_id,
            KnowledgeFile.name == PLATFORM_FEEDBACK_FILENAME
        ).order_by(KnowledgeFile.id.asc()).first()
        if feedback_file:
            feedback_file.file_path = file_path
            feedback_file.filename = PLATFORM_FEEDBACK_FILENAME
            feedback_file.file_type = 'md'
            feedback_file.file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            db.commit()
            db.refresh(feedback_file)

    if feedback_file:
        if not os.path.exists(feedback_file.file_path):
            with open(feedback_file.file_path, 'w', encoding='utf-8') as f:
                f.write(initial_content)
            feedback_file.file_size = os.path.getsize(feedback_file.file_path)
            db.commit()
            db.refresh(feedback_file)
        return feedback_file

    if not os.path.exists(file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(initial_content)

    feedback_file = KnowledgeFile(
        folder_id=folder_id,
        name=PLATFORM_FEEDBACK_FILENAME,
        filename=PLATFORM_FEEDBACK_FILENAME,
        file_path=file_path,
        file_size=os.path.getsize(file_path),
        file_type='md',
        created_by=user_id
    )
    db.add(feedback_file)
    db.commit()
    db.refresh(feedback_file)
    return feedback_file


def _ensure_platform_feedback_terminology_file(db: Session, user_id: int) -> KnowledgeFile:
    """确保平台反馈术语对照表存在于知识库中。"""
    folder_id = None

    knowledge_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "knowledge")
    if not os.path.exists(knowledge_dir):
        os.makedirs(knowledge_dir)

    file_path = os.path.join(knowledge_dir, PLATFORM_FEEDBACK_TERMINOLOGY_RELATIVE_PATH)
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    initial_content = "# 平台反馈的术语对照表\n\n| 非标准词 | 标准词 |\n| --- | --- |\n"

    feedback_file = db.query(KnowledgeFile).filter(
        KnowledgeFile.folder_id == folder_id,
        KnowledgeFile.name == PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME,
        KnowledgeFile.file_path == file_path
    ).first()

    if not feedback_file:
        feedback_file = db.query(KnowledgeFile).filter(
            KnowledgeFile.folder_id == folder_id,
            KnowledgeFile.name == PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME
        ).order_by(KnowledgeFile.id.asc()).first()
        if feedback_file:
            feedback_file.file_path = file_path
            feedback_file.filename = PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME
            feedback_file.file_type = 'md'
            feedback_file.file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
            db.commit()
            db.refresh(feedback_file)

    if feedback_file:
        if not os.path.exists(feedback_file.file_path):
            with open(feedback_file.file_path, 'w', encoding='utf-8') as f:
                f.write(initial_content)
            feedback_file.file_size = os.path.getsize(feedback_file.file_path)
            db.commit()
            db.refresh(feedback_file)
        return feedback_file

    if not os.path.exists(file_path):
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(initial_content)

    feedback_file = KnowledgeFile(
        folder_id=folder_id,
        name=PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME,
        filename=PLATFORM_FEEDBACK_TERMINOLOGY_FILENAME,
        file_path=file_path,
        file_size=os.path.getsize(file_path),
        file_type='md',
        created_by=user_id
    )
    db.add(feedback_file)
    db.commit()
    db.refresh(feedback_file)
    return feedback_file


def _get_platform_feedback_targets(db: Session, user_id: int) -> list[KnowledgeFile]:
    """返回平台反馈句式清单固定主文件。"""
    primary_file = _ensure_platform_feedback_sentence_file(db, user_id)
    return [primary_file] if primary_file else []


def _get_platform_feedback_terminology_targets(db: Session, user_id: int) -> list[KnowledgeFile]:
    """返回平台反馈术语对照表固定主文件。"""
    primary_file = _ensure_platform_feedback_terminology_file(db, user_id)
    return [primary_file] if primary_file else []

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "polish-lab")

# ============================================================
# 模型定义
# ============================================================

def _get_date_subfolder_id(db, folder_id: int, user_id: int) -> tuple:
    """返回物理目录路径和 None 作为 folder_id"""
    from datetime import datetime
    date_str = datetime.now().strftime("%Y%m%d")
    return None, date_str


class TextPolishInput(BaseModel):
    text: str
    style_guide_id: Optional[int] = None
    terminology_id: Optional[int] = None
    product_type: Optional[str] = None


class SkillPolishInput(BaseModel):
    text: str
    style_guide_id: int = 1
    terminology_id: Optional[int] = None


class FeedbackInput(BaseModel):
    original_text: str
    polished_text: str
    accuracy: int              # 0-100
    corrections: str = ""       # 用户修正内容，每行一条 "非标准 → 标准"
    target: str = "terminology" # "terminology" 或 "sentence_guide"
    terminology_file_id: Optional[int] = None
    sentence_file_id: Optional[int] = None


class DocumentFeedbackItem(BaseModel):
    before: str = ""
    after: str = ""
    type: str = ""
    accepted: bool = True
    status: str = ""
    paragraph: Optional[int] = None


class DocumentFeedbackInput(BaseModel):
    document_id: Optional[int] = None
    source_filename: str = ""
    items: List[DocumentFeedbackItem] = []


class PolishRuleMatch(BaseModel):
    rule_name: str
    before: str
    after: str
    type: str
    paragraph: Optional[int] = None
    match_detail: Optional[dict] = None


class CatCandidate(BaseModel):
    """单个候选模板（展示给前端）。"""
    template_text: str = ""
    template_id: Optional[str] = None
    raw_template_text: Optional[str] = None
    normalized_template_text: Optional[str] = None
    string_score: float = 0.0
    semantic_score: Optional[float] = None
    effective_semantic_score: Optional[float] = None
    filtered_semantic_score: Optional[float] = None
    ai_reason: Optional[str] = None
    match_tier: str = "reference"
    rule_source: Optional[str] = None
    needs_review: bool = False
    review_tags: List[str] = []
    drop_reason: Optional[str] = None


class CatAnalyzeItem(BaseModel):
    """一个段落的匹配结果。"""
    paragraph_index: int = 0
    sentence_index: int = 0
    source_paragraph_index: int = 0
    source_paragraph_text: str = ""
    original_text: str = ""
    has_candidates: bool = False
    candidates: List[CatCandidate] = []


class CatDecision(BaseModel):
    """用户对单个段落的决策。"""
    paragraph_index: int = 0
    sentence_index: int = 0
    source_paragraph_index: Optional[int] = None
    source_paragraph_text: str = ""
    source_sentence_text: str = ""
    action: str = "pending"
    original_text: str = ""
    accepted_template: Optional[str] = None
    accepted_template_id: Optional[str] = None
    modified_text: Optional[str] = None
    rejected_template: Optional[str] = None
    rejected_template_id: Optional[str] = None
    string_score: float = 0.0
    semantic_score: Optional[float] = None
    ai_reason: Optional[str] = None
    category: Optional[str] = None
    severity: Optional[str] = None
    candidate_text: Optional[str] = None
    rule_source: Optional[str] = None


class CatAnalyzeRequest(BaseModel):
    """第一阶段请求参数。"""
    file_id: Optional[int] = None
    sentence_file_id: Optional[int] = None
    terminology_file_id: Optional[int] = None
    requirements: Optional[str] = None
    min_match_threshold: float = 0.34
    fuzzy_lower_bound: float = 0.70
    ai_semantic_scoring: bool = True
    ai_reason_max_chars: int = 15


class CatApplyRequest(BaseModel):
    """第二阶段请求参数。"""
    analyze_id: str = ""
    file_id: Optional[int] = None
    source_filename: str = ""
    decisions: List[CatDecision] = []
    sentence_file_id: Optional[int] = None
    sentence_file_name: Optional[str] = None
    ai_semantic_scoring: Optional[bool] = None
    ai_semantic_scoring: Optional[bool] = None


_DOC_REVIEW_AUTO_APPLY_THRESHOLD = 95


_MATCH_SEGMENT_WEIGHTS = {
    "condition": 0.15,
    "action": 0.30,
    "object": 0.30,
    "result": 0.20,
    "additional": 0.05,
}

_MATCH_SEGMENT_LABELS = {
    "condition": "条件段",
    "action": "动作段",
    "object": "对象段",
    "result": "结果/判断段",
    "additional": "附加信息段",
}


def _normalize_compare_text(text: str) -> str:
    if text is None:
        return ""
    return re.sub(r'\s+', ' ', str(text)).strip()


def _is_short_label_like_text(text: str) -> bool:
    value = _normalize_compare_text(text)
    if not value:
        return False
    if re.fullmatch(r'[A-Za-z]\d{1,3}', value):
        return True
    if re.fullmatch(r'[A-Za-z]{1,3}', value):
        return True
    if re.fullmatch(r'\d{1,3}', value):
        return True
    normalized = _normalize_sentence_for_match(value)
    return len(normalized) <= 3 and not re.search(r'[\u4e00-\u9fff]', value)


def _strip_doc_trailing_punctuation(text: str) -> str:
    return re.sub(r'[。.!！？?，,;；:：]+$', '', _normalize_visible_compare_text(text))


def _is_field_style_colon_rewrite(before: str, after: str) -> bool:
    before_text = _normalize_visible_compare_text(before)
    after_text = _normalize_visible_compare_text(after)
    if not before_text or not after_text or before_text == after_text:
        return False
    if '为' not in before_text or not re.search(r'[：:]', after_text):
        return False
    compact_before = re.sub(r'\s+', '', before_text)
    compact_after = re.sub(r'\s+', '', after_text)
    reverted = re.sub(r'[：:]', '为', compact_after)
    return reverted == compact_before


def _is_low_value_doc_change(before: str, after: str, change_type: str = '', rule_name: str = '') -> bool:
    before_text = _normalize_visible_compare_text(before)
    after_text = _normalize_visible_compare_text(after)
    normalized_type = str(change_type or '').lower()
    normalized_rule = str(rule_name or '')
    punctuation_only_change = (
        before_text
        and after_text
        and before_text.rstrip('。.!！？?，,;；:：') == after_text.rstrip('。.!！？?，,;；:：')
        and before_text != after_text
    )
    if punctuation_only_change and _should_add_contextual_terminal_punctuation(before_text):
        return False

    before_core = _strip_doc_trailing_punctuation(before)
    after_core = _strip_doc_trailing_punctuation(after)
    if not before_core or not after_core:
        return False
    if before_core == after_core:
        return True
    if _is_field_style_colon_rewrite(before, after):
        return True

    short_limit = 4
    if len(before_core) <= short_limit and after_core == f'请{before_core}':
        return True
    if before_core.startswith('请') and len(before_core) <= short_limit and after_core == before_core[1:]:
        return True

    is_format_like = normalized_type in {'format', 'punctuation'} or normalized_rule == '基础规范化'
    if is_format_like and len(before_core) <= short_limit and after_core.startswith(before_core):
        return True
    return False


def _build_visible_change_entry(line: int, original: str, polished: str, change_type: str, rule_name: str = '', sentence_guide: str = '') -> Optional[dict]:
    visible_polished = _reapply_sentence_prefix(original, polished)
    if _is_low_value_doc_change(original, visible_polished, change_type, rule_name):
        return None

    entry = {
        "line": line,
        "original": original[:200],
        "polished": visible_polished[:200],
        "type": change_type,
    }
    if change_type != 'ai':
        entry["match_detail"] = _build_change_match_detail(original, visible_polished, change_type, rule_name, sentence_guide)
    return entry


def _existing_change_match_detail(item) -> Optional[dict]:
    if isinstance(item, dict):
        detail = item.get('match_detail')
    else:
        detail = getattr(item, 'match_detail', None)
    return detail if isinstance(detail, dict) else None


def _doc_change_memory_key(before: str, after: str) -> str:
    return f"{_normalize_compare_text(before)}\u0001{_normalize_compare_text(after)}"


def _load_rejected_doc_change_keys(db: Session) -> set[tuple[str, str]]:
    if db is None:
        return set()
    rows = db.query(PolishFeedback.original_text, PolishFeedback.polished_text).filter(
        PolishFeedback.target == 'document_rejected_change'
    ).all()
    return {(_normalize_compare_text(before), _normalize_compare_text(after)) for before, after in rows if before or after}


def _is_rejected_doc_change(before: str, after: str, rejected_keys: set[tuple[str, str]] = None) -> bool:
    if not rejected_keys:
        return False
    current_before = _normalize_compare_text(before)
    current_after = _normalize_compare_text(after)
    if not current_before or not current_after:
        return False
    for rejected_before, rejected_after in rejected_keys:
        if not rejected_before or not rejected_after:
            continue
        before_matches = current_before == rejected_before or current_before.startswith(rejected_before) or rejected_before.startswith(current_before)
        after_matches = current_after == rejected_after or current_after.startswith(rejected_after) or rejected_after.startswith(current_after)
        if before_matches and after_matches:
            return True
    return False


def _doc_change_display_priority(change_type: str) -> int:
    priority_map = {
        'style': 100,
        'preferred_sentences': 100,
        'sentence_applicability_rule': 100,
        'terminology': 90,
        'term': 90,
        'terminology_rule': 90,
        'forbidden': 80,
        'forbidden_rule': 80,
        'forbidden_words': 80,
        'imperative': 70,
        'imperative_rule': 70,
        'ai': 40,
        'format': 10,
        'punctuation': 10,
    }
    return priority_map.get(str(change_type or ''), 30)


def _is_displayable_doc_change(change_type: str) -> bool:
    return _doc_change_display_priority(change_type) >= 70


def _dedupe_visible_doc_changes(changes: list[dict]) -> list[dict]:
    selected = {}
    order = []
    for item in changes:
        before = _normalize_compare_text(item.get('before', ''))
        after = _normalize_compare_text(item.get('after', ''))
        key = before or after
        if not key:
            continue
        score = _doc_change_display_priority(item.get('type', ''))
        existing = selected.get(key)
        if existing is None:
            selected[key] = item
            order.append(key)
            continue
        existing_score = _doc_change_display_priority(existing.get('type', ''))
        if score > existing_score:
            selected[key] = item
    return [selected[key] for key in order if key in selected]


def _pick_visible_doc_changes(changes: list[dict]) -> list[dict]:
    deduped = _dedupe_visible_doc_changes(changes)
    preferred = [item for item in deduped if _is_displayable_doc_change(item.get('type', ''))]
    return preferred if preferred else deduped


def _build_doc_change_details(original_text: str, polished_text: str, changes: list, rejected_keys: set[str] = None, sentence_guide: str = '') -> list[dict]:
    original_lines = [line.strip() for line in (original_text or '').split('\n') if line.strip()]
    polished_lines = [line.strip() for line in (polished_text or '').split('\n') if line.strip()]

    matcher = SequenceMatcher(None, original_lines, polished_lines)
    diff_rows = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue

        before_lines = original_lines[i1:i2]
        after_lines = polished_lines[j1:j2]
        max_len = max(len(before_lines), len(after_lines), 1)

        for idx in range(max_len):
            before = before_lines[idx] if idx < len(before_lines) else ''
            after = _protect_model_numbers(after_lines[idx]) if idx < len(after_lines) else ''
            if _normalize_compare_text(before) == _normalize_compare_text(after):
                continue
            if _is_low_value_doc_change(before, after, 'ai', ''):
                continue
            if _is_rejected_doc_change(before, after, rejected_keys):
                continue
            diff_rows.append({
                "before": before,
                "after": after,
                "type": "ai" if before and after else "format"
            })

    if not diff_rows:
        return []

    normalized_changes = []
    for item in changes or []:
        existing_detail = _existing_change_match_detail(item)
        if isinstance(item, dict):
            before = item.get('before', '')
            after = _protect_model_numbers(item.get('after', ''))
            change_type = item.get('type', '')
            paragraph = item.get('paragraph') or item.get('paragraph_index')
        else:
            before = getattr(item, 'before', '')
            after = _protect_model_numbers(getattr(item, 'after', ''))
            change_type = getattr(item, 'type', '')
            paragraph = getattr(item, 'paragraph', None)

        normalized_before = _normalize_compare_text(before)
        normalized_after = _normalize_compare_text(after)
        if normalized_before == normalized_after:
            continue
        if _is_low_value_doc_change(before, after, change_type, getattr(item, 'rule_name', '') if not isinstance(item, dict) else item.get('rule_name', '')):
            continue
        if _is_rejected_doc_change(before, after, rejected_keys):
            continue
        normalized_changes.append({
            "before": before,
            "after": after,
            "type": change_type,
            "paragraph": paragraph,
            "rule_name": getattr(item, 'rule_name', '') if not isinstance(item, dict) else item.get('rule_name', ''),
            "match_detail": existing_detail or _build_change_match_detail(
                before,
                after,
                change_type,
                getattr(item, 'rule_name', '') if not isinstance(item, dict) else item.get('rule_name', ''),
                sentence_guide,
            ),
        })

    for row in diff_rows:
        row_before = _normalize_compare_text(row['before'])
        row_after = _normalize_compare_text(row['after'])
        matched_type = row['type']
        for item in normalized_changes:
            item_before = _normalize_compare_text(item['before'])
            item_after = _normalize_compare_text(item['after'])
            if item_before and item_before in row_before:
                matched_type = item['type'] or matched_type
                row['paragraph'] = item.get('paragraph')
                row['rule_name'] = item.get('rule_name', '')
                row['match_detail'] = item.get('match_detail')
                break
            if item_after and item_after in row_after:
                matched_type = item['type'] or matched_type
                row['paragraph'] = item.get('paragraph')
                row['rule_name'] = item.get('rule_name', '')
                row['match_detail'] = item.get('match_detail')
                break
        row['type'] = matched_type
        row['match_detail'] = row.get('match_detail') or _build_change_match_detail(
            row.get('before', ''),
            row.get('after', ''),
            matched_type,
            row.get('rule_name', ''),
            sentence_guide,
        )

    return _pick_visible_doc_changes(diff_rows)


def _filter_visible_doc_changes(changes: list, rejected_keys: set[str] = None, sentence_guide: str = '') -> list[dict]:
    visible = []
    for item in changes or []:
        existing_detail = _existing_change_match_detail(item)
        if isinstance(item, dict):
            before = item.get('before', '')
            after = _protect_model_numbers(item.get('after', ''))
            change_type = item.get('type', '')
            rule_name = item.get('rule_name', '')
            paragraph = item.get('paragraph') or item.get('paragraph_index')
        else:
            before = getattr(item, 'before', '')
            after = _protect_model_numbers(getattr(item, 'after', ''))
            change_type = getattr(item, 'type', '')
            rule_name = getattr(item, 'rule_name', '')
            paragraph = getattr(item, 'paragraph', None)

        if _normalize_compare_text(before) == _normalize_compare_text(after):
            continue

        if not _normalize_compare_text(before) and not _normalize_compare_text(after):
            continue

        if _is_low_value_doc_change(before, after, change_type, rule_name):
            continue

        if _is_rejected_doc_change(before, after, rejected_keys):
            continue

        visible.append({
            'before': before,
            'after': after,
            'type': change_type,
            'rule_name': rule_name,
            'paragraph': paragraph,
            'match_detail': existing_detail or _build_change_match_detail(before, after, change_type, rule_name, sentence_guide),
        })
    return _pick_visible_doc_changes(visible)


def _doc_change_pair_key(before: str, after: str) -> tuple[str, str]:
    return (_normalize_compare_text(before), _normalize_compare_text(after))


def _doc_change_before_key(before: str) -> str:
    return _normalize_compare_text(before)


def _build_previous_doc_change_keys(previous_doc: Optional[PolishedDocument], rejected_keys: set[str] = None) -> tuple[set[tuple[str, str]], set[str]]:
    if previous_doc is None:
        return set(), set()
    previous_original = previous_doc.original_content or ''
    previous_polished = previous_doc.polished_content or ''
    if not previous_original or not previous_polished:
        return set(), set()
    previous_changes = _build_doc_change_details(previous_original, previous_polished, [], rejected_keys, '')
    pair_keys = {
        _doc_change_pair_key(item.get('before', ''), item.get('after', ''))
        for item in previous_changes
        if _normalize_compare_text(item.get('before', '')) and _normalize_compare_text(item.get('after', ''))
    }
    before_keys = {
        _doc_change_before_key(item.get('before', ''))
        for item in previous_changes
        if _normalize_compare_text(item.get('before', '')) and _normalize_compare_text(item.get('after', ''))
    }
    if not before_keys:
        original_lines = {
            _normalize_compare_text(line)
            for line in (previous_original or '').split('\n')
            if _normalize_compare_text(line)
        }
        polished_lines = {
            _normalize_compare_text(line)
            for line in (previous_polished or '').split('\n')
            if _normalize_compare_text(line)
        }
        before_keys = {
            line for line in original_lines
            if line and line not in polished_lines
        }
    return pair_keys, before_keys


def _mark_new_doc_review_items(review_items: list[dict], previous_change_keys: set[tuple[str, str]], previous_before_keys: set[str]) -> int:
    new_count = 0
    for item in review_items or []:
        change_key = _doc_change_pair_key(item.get('before', ''), item.get('after', ''))
        before_key = _doc_change_before_key(item.get('before', ''))
        is_new = bool(before_key) and before_key not in previous_before_keys and change_key not in previous_change_keys
        item['is_new_since_last_polish'] = is_new
        if is_new:
            new_count += 1
    return new_count




def _apply_term_only(text: str, term_dict: dict) -> tuple[str, list[PolishRuleMatch]]:
    """仅按术语库逐行替换（不触发样式规则），用于 AI 已润色后的术语修正。"""
    if not term_dict:
        return text, []
    changes = []
    lines = text.split('\n')
    result_lines = []
    for line in lines:
        new_line = line
        for old_term, new_term in term_dict.items():
            if old_term in new_line:
                new_line = new_line.replace(old_term, new_term)
                changes.append(PolishRuleMatch(
                    rule_name="术语替换",
                    before=old_term,
                    after=new_term,
                    type="terminology"
                ))
        result_lines.append(new_line)
    return '\n'.join(result_lines), changes


def _apply_skill_polish(
    text: str, 
    skill_rules: dict, 
    db: Session,
    sentence_guide: str = None,
    terminology: str = None,
    requirements: str = None,
    is_title: bool = False,
    db_terminology: dict = None
) -> tuple[str, list[PolishRuleMatch]]:
    """应用skill规则进行润色。is_title=True 时跳过尾部标点规范化。"""

    def _is_noun_phrase(text: str) -> bool:
        """判断文本是否为纯名词短语（标题、标签、参数说明等），不需要追加标点。"""
        t = text.strip()
        if not t:
            return False
        # 1. 极短文本（<=4字）通常是标签/名称
        if len(t) <= 4:
            return True
        # 2. 以冒号结尾，通常是字段标签（如"试剂名称："）
        if t.endswith(('：', ':')):
            return True
        # 3. 包含编号/列表标记，通常是标题（如"1. 概述"、"第2章"）
        if re.match(r'^[\d一二三四五六七八九十]+[\.、\s]', t):
            return True
        # 4. 不包含任何谓语动词标记，判定为名词短语
        verb_markers = [
            '将', '请', '按', '点击', '选择', '输入', '打开', '关闭',
            '启动', '停止', '设置', '检查', '确认', '安装', '连接',
            '使用', '进行', '可以', '需要', '应该', '必须', '确保',
            '按下', '旋转', '调节', '插入', '取出', '放置', '点击',
            '执行', '访问', '查看', '显示', '支持', '提供', '包含',
            '通过', '根据', '按照', '用于', '适用于', '分为',
        ]
        if not any(marker in t for marker in verb_markers):
            return True
        return False

    def _should_show_basic_change(before: str, after: str, change_type: str) -> bool:
        before_text = (before or '').strip()
        after_text = (after or '').strip()
        if not before_text or before_text == after_text:
            return False
        if change_type == 'punctuation':
            return False
        if len(before_text) <= 4:
            return False
        if _is_noun_phrase(before_text):
            return False
        before_core = before_text.rstrip('。.!！？?，,;；:：')
        after_core = after_text.rstrip('。.!！？?，,;；:：')
        return before_core != after_core
    if _use_ai_only():
        return text, []

    sentence_only_mode = False
    changes = []
    lines = text.split('\n')
    polished_lines = []
    triggered_rule_ids = []
    triggered_engine_keys = []
    
    style_rules = []
    if sentence_guide:
        style_rules = _extract_style_rules(sentence_guide)
        guide_cache_key = _guide_cache_key(sentence_guide)
        if guide_cache_key not in _logged_sentence_guide_keys:
            logger.warning(
                "polish sentence guide parsed: text_len=%s rules=%s rule_types=%s",
                len(sentence_guide or ''),
                len(style_rules),
                [rule.get('type') for rule in style_rules[:10]],
            )
            _logged_sentence_guide_keys.add(guide_cache_key)
    
    term_dict = {}
    # 先解析文件中的术语替换（支持中英文多列表，自动语言过滤）
    if terminology and not sentence_only_mode:
        try:
            parsed = _parse_terminology(terminology)
            if parsed:
                lang = _detect_language(text)
                term_dict = _filter_terms_by_lang(parsed, lang)
        except Exception:
            pass
    # 合并数据库术语库（优先级高于文件术语）
    if db_terminology and not sentence_only_mode:
        term_dict.update(db_terminology)
    
    engine_enabled_rules = [] if sentence_only_mode else (get_enabled_engine_keys(db) if db else None)
    if is_title and engine_enabled_rules:
        engine_enabled_rules = [key for key in engine_enabled_rules if key != 'punctuation']
    custom_rules = [] if sentence_only_mode else (get_enabled_custom_rules(db) if db else [])

    def _rule_type(rule) -> str:
        return str(getattr(rule, 'rule_type', '') or '')

    sentence_custom_rules = [rule for rule in custom_rules if _rule_type(rule) == 'sentence_applicability_rule']
    term_custom_rules = [rule for rule in custom_rules if _rule_type(rule) == 'replacement_rule']
    other_custom_rules = [
        rule for rule in custom_rules
        if _rule_type(rule) not in {'sentence_applicability_rule', 'replacement_rule'}
    ]

    term_enabled_rules = None
    other_engine_rules = None
    if engine_enabled_rules is not None:
        term_enabled_rules = [key for key in engine_enabled_rules if key == 'termReplace']
        other_engine_rules = [key for key in engine_enabled_rules if key != 'termReplace']
    else:
        term_enabled_rules = ['termReplace']
        other_engine_rules = ['typoCheck', 'imperativePlease', 'numberSpace', 'punctuation']
    if is_title:
        other_engine_rules = [key for key in other_engine_rules if key != 'punctuation']

    typo_dict = _load_typo_dict(db) if db else get_default_typo_dict()

    def _append_custom_issues(issues: list[dict]):
        nonlocal has_changes
        if not issues:
            return
        for issue in issues:
            if issue.get('rule_id'):
                triggered_rule_ids.append(issue.get('rule_id'))
            changes.append(PolishRuleMatch(
                rule_name=issue.get('rule_name', '自定义规则'),
                before=issue.get('before', issue.get('original', ''))[:80],
                after=issue.get('after', issue.get('replacement', ''))[:80],
                type=issue.get('type', 'custom')
            ))
        has_changes = True

    def _append_engine_issues(issues: list[dict]):
        nonlocal has_changes
        if not issues:
            return
        for issue in issues:
            if issue.get('engine_key'):
                triggered_engine_keys.append(issue.get('engine_key'))
            changes.append(PolishRuleMatch(
                rule_name=issue.get('rule_name', '规则检测'),
                before=issue.get('original', ''),
                after=issue.get('replacement', ''),
                type=issue.get('type', 'format')
            ))
        has_changes = True

    for line in lines:
        original = line
        new_line = line.strip()
        # 剥离编号前缀（如 "5. "、"1、"、"3) "），正文独立匹配
        step_prefix = ""
        step_match = re.match(r'^((?:\d+[.、)]?)+(?:\s+|(?=[\u4e00-\u9fffA-Za-z(（])))\s*(.+)$', new_line)
        if step_match:
            step_prefix = step_match.group(1)
            new_line = step_match.group(2)
        
        has_changes = False
        matched_preferred_sentence = False
        non_tmpl_rules = [r for r in style_rules if r.get("type") != "preferred_sentences"] if style_rules else []
        template_match_line = _normalize_operation_sentence_for_match(new_line)
        
        if style_rules and not is_title and _should_skip_expensive_template_match(template_match_line):
            original_before = new_line
            if non_tmpl_rules:
                new_line, rule_changes = _apply_style_rules(new_line, non_tmpl_rules)
                if rule_changes:
                    changes.extend(rule_changes)
                    has_changes = True
            if not has_changes:
                fallback_line = _apply_constraint_polish(new_line)
                if fallback_line != new_line:
                    changes.append(PolishRuleMatch(
                        rule_name="约束润色",
                        before=new_line[:80],
                        after=fallback_line[:80],
                        type="style"
                    ))
                    new_line = fallback_line
                    has_changes = True
            logger.warning(
                "polish sentence guide skip heavy template match: before=%r after=%r",
                original_before[:120],
                new_line[:120],
            )
        elif style_rules and not is_title:
            preferred_entries = _preferred_entries_for_sentence(template_match_line, sentence_guide)
            pooling_sentence_mode = _is_pooling_platform_sentence(template_match_line)
            # 拆分子句检测匹配，匹配到后在全文级执行替换以避免子句边界重复
            clauses = [c.strip() for c in re.split(r'[，。；！？,;!?]+', new_line) if c.strip()]
            normalized_clauses = [c.strip() for c in re.split(r'[，。；！？,;!?]+', template_match_line) if c.strip()]
            if len(clauses) > 1:
                whole_template, whole_score, whole_level = (None, 0.0, 'NONE')
                if not pooling_sentence_mode:
                    whole_template, whole_score, whole_level = _best_guarded_match(template_match_line, preferred_entries)
                if whole_template and whole_score >= _POLISH_MATCH_CONFIG["l3_auto_confidence"]:
                    original_before = new_line
                    new_line = replace_with_context(new_line, whole_template)
                    display_after = _reapply_sentence_prefix(original, new_line)
                    changes.append(PolishRuleMatch(
                        rule_name="句式模板匹配",
                        before=original_before[:80],
                        after=display_after[:80],
                        type="style"
                    ))
                    logger.warning(
                        "polish sentence guide matched: before=%r after=%r template=%r score=%s level=%s",
                        original_before[:120],
                        new_line[:120],
                        whole_template[:80],
                        round(whole_score, 4),
                        whole_level,
                    )
                    matched_preferred_sentence = True
                    has_changes = True
                else:
                    original_before = new_line
                    clause_matched_templates = []
                    if pooling_sentence_mode:
                        updated_line = new_line
                        for window in _split_pooling_sentence_windows(new_line):
                            normalized_window = _normalize_operation_sentence_for_match(window)
                            window_templates = _preferred_entries_for_sentence(normalized_window, sentence_guide, single_clause_only=False)
                            tmpl, score, level = _best_pooling_clause_template(normalized_window, window_templates)
                            if not tmpl:
                                continue
                            replaced_window = replace_with_context(window, tmpl)
                            if replaced_window == window:
                                continue
                            updated_line = updated_line.replace(window, replaced_window, 1)
                            clause_matched_templates.append(tmpl)
                            matched_preferred_sentence = True
                        new_line = _normalize_internal_sentence_punctuation(updated_line)
                    else:
                        clause_parts = re.split(r'([，。；！？,;!?]+)', new_line)
                        rebuilt_parts = []
                        for index, part in enumerate(clause_parts):
                            if index % 2 == 1:
                                rebuilt_parts.append(part)
                                continue
                            clause = part.strip()
                            normalized_clause = normalized_clauses[index // 2] if index // 2 < len(normalized_clauses) else _normalize_operation_sentence_for_match(clause)
                            if not clause:
                                rebuilt_parts.append(part)
                                continue
                            best_template = None
                            best_rank = None
                            clause_templates = _preferred_entries_for_sentence(normalized_clause, sentence_guide, single_clause_only=True)
                            if clause_templates:
                                tmpl, score, level = _best_guarded_match(normalized_clause, clause_templates)
                                if not tmpl:
                                    pass
                                else:
                                    clause_score = _clause_alignment_score(normalized_clause, tmpl)
                                    keyword_score = _compare_semantic_keywords(normalized_clause, tmpl)
                                    structured_score = _structured_match_score(normalized_clause, tmpl)
                                    rank = (
                                        score,
                                        structured_score['pair'],
                                        structured_score['condition'],
                                        structured_score['action'],
                                        clause_score,
                                        keyword_score,
                                        len(_normalize_sentence_for_match(tmpl)),
                                    )
                                    if best_rank is None or rank > best_rank:
                                        best_rank = rank
                                        best_template = tmpl
                            replaced_clause = clause
                            if best_template:
                                next_separator = clause_parts[index + 1] if index + 1 < len(clause_parts) else ''
                                replaced_clause = _replace_clause_directly(clause, best_template, next_separator)
                                clause_matched_templates.append(best_template)
                                matched_preferred_sentence = True
                            else:
                                replaced_clause = _apply_constraint_polish(clause)
                            rebuilt_parts.append(part.replace(clause, replaced_clause, 1))
                        new_line = ''.join(rebuilt_parts)
                    if new_line != original_before:
                        display_after = _reapply_sentence_prefix(original, new_line)
                        changes.append(PolishRuleMatch(
                            rule_name="句式模板匹配" if clause_matched_templates else "约束润色",
                            before=original_before[:80],
                            after=display_after[:80],
                            type="style"
                        ))
                        if clause_matched_templates:
                            logger.warning(
                                "polish sentence guide matched: before=%r after=%r templates=%s",
                                original_before[:120],
                                new_line[:120],
                                [template[:80] for template in clause_matched_templates],
                            )
                        has_changes = True
                    # 阶段3: 应用非句式规则（禁用词、被动语态、约束润色等）
                    if non_tmpl_rules and not matched_preferred_sentence:
                        new_line, rule_changes = _apply_style_rules(new_line, non_tmpl_rules)
                        if rule_changes:
                            changes.extend(rule_changes)
                            has_changes = True
                    if not matched_preferred_sentence and not has_changes:
                        fallback_line = _apply_constraint_polish(new_line)
                        if fallback_line != new_line:
                            changes.append(PolishRuleMatch(
                                rule_name="约束润色",
                                before=new_line[:80],
                                after=fallback_line[:80],
                                type="style"
                            ))
                            new_line = fallback_line
                            has_changes = True
            else:
                whole_template, whole_score, whole_level = _best_guarded_match(template_match_line, preferred_entries)
                if whole_template and whole_score >= _POLISH_MATCH_CONFIG["l3_auto_confidence"]:
                    original_before = new_line
                    new_line = replace_with_context(new_line, whole_template)
                    display_after = _reapply_sentence_prefix(original, new_line)
                    changes.append(PolishRuleMatch(
                        rule_name="句式模板匹配",
                        before=original_before[:80],
                        after=display_after[:80],
                        type="style"
                    ))
                    logger.warning(
                        "polish sentence guide matched: before=%r after=%r template=%r score=%s level=%s",
                        original_before[:120],
                        new_line[:120],
                        whole_template[:80],
                        round(whole_score, 4),
                        whole_level,
                    )
                    has_changes = True
                    matched_preferred_sentence = True
                elif non_tmpl_rules:
                    new_line, rule_changes = _apply_style_rules(new_line, non_tmpl_rules)
                    if rule_changes:
                        logger.warning(
                            "polish sentence guide matched: before=%r after=%r matches=%s",
                            original[:120],
                            new_line[:120],
                            [change.rule_name for change in rule_changes],
                        )
                        changes.extend(rule_changes)
                        has_changes = True

        if sentence_custom_rules and not is_title:
            new_line, custom_issues = apply_custom_rules(new_line, sentence_custom_rules)
            if custom_issues:
                logger.warning(
                    "polish sentence custom matched: before=%r after=%r matches=%s",
                    original[:120],
                    new_line[:120],
                    [issue.get('rule_name') for issue in custom_issues],
                )
            _append_custom_issues(custom_issues)

        if not sentence_only_mode:
            term_line, term_issues = apply_all_rules(
                new_line,
                term_dict=term_dict,
                typo_dict=typo_dict,
                enabled_rules=term_enabled_rules,
                context_text=text
            )
            _append_engine_issues(term_issues)
            new_line = term_line

        if term_custom_rules:
            new_line, custom_issues = apply_custom_rules(new_line, term_custom_rules)
            _append_custom_issues(custom_issues)

        if other_custom_rules and not matched_preferred_sentence:
            new_line, custom_issues = apply_custom_rules(new_line, other_custom_rules)
            _append_custom_issues(custom_issues)

        # ── 应用其余系统规则（从 DB 读取启用的规则） ──
        if not sentence_only_mode and not matched_preferred_sentence:
            polished_line, engine_issues = apply_all_rules(
                new_line,
                term_dict={},
                typo_dict=typo_dict,
                enabled_rules=other_engine_rules,
                context_text=text
            )
            _append_engine_issues(engine_issues)
            new_line = polished_line
        
        if not sentence_only_mode:
            new_line = re.sub(r'\s+', ' ', new_line)
            new_line = _protect_model_numbers(new_line)
        
        # 标题、表标题、图标题等不加句号，也不做空间距规整
        if not is_title and not sentence_only_mode:
            if _should_add_contextual_terminal_punctuation(new_line):
                new_line = new_line.rstrip('，,;；：:') + '。'
            new_line = re.sub(rf'(?<=\d)\s*(?=({_NUMBER_SPACE_UNITS})(?![A-Za-z]))', ' ', new_line)
            new_line = _protect_model_numbers(new_line)
        elif sentence_only_mode and not has_changes:
            new_line = original
        
        if not sentence_only_mode and (new_line != original or has_changes):
            change_type = "format"
            if '。' in new_line[-2:] and original[-1] not in '。.!！？?':
                change_type = "punctuation"
            elif terminology and term_dict:
                change_type = "terminology"
            elif style_rules:
                change_type = "style"
            
            if _should_show_basic_change(original, new_line, change_type):
                changes.append(PolishRuleMatch(
                    rule_name="基础规范化",
                    before=original[:80],
                    after=new_line[:80],
                    type=change_type
                ))
        
        if step_prefix and not new_line.startswith(step_prefix):
            new_line = step_prefix + new_line
        polished_lines.append(new_line)
    
    if db and (triggered_rule_ids or triggered_engine_keys):
        record_rule_triggers(db, triggered_rule_ids, triggered_engine_keys)

    return '\n'.join(polished_lines), changes


def _parse_table_sentence_templates(section: str) -> list[str]:
    """从 Markdown 表格中提取句式模板列的内容。
    
    优先提取"示例"列（包含真实句子），没有示例列时回退到"句式模板"列。
    """
    sentences = []
    table_blocks = re.findall(r'(\|[^\n]+\|\n\|[-:\s|]+\|\n(?:\|[^\n]+\|\n?)+)', section)
    for block in table_blocks:
        lines = [l for l in block.strip().split('\n') if l.strip()]
        if len(lines) < 2:
            continue
        header_row = [c.strip() for c in lines[0].strip('|').split('|')]
        # 优先查找"示例"列（包含真实例句，适合规则匹配）
        example_col_idx = None
        template_col_idx = None
        for i, h in enumerate(header_row):
            if any(kw in h for kw in ['示例', '例子']):
                example_col_idx = i
            if any(kw in h for kw in ['句式模板', '句式', '模板', '句型', '标准句式']):
                template_col_idx = i
        # 优先用示例列，回退到模板列，再回退到第二列
        col_idx = example_col_idx if example_col_idx is not None else template_col_idx
        if col_idx is None and len(header_row) >= 2:
            col_idx = 1
        if col_idx is None:
            continue
        for line in lines[2:]:
            cells = [c.strip() for c in line.strip('|').split('|')]
            if col_idx < len(cells):
                val = cells[col_idx].strip().strip('"\'""''').strip()
                # 去掉常见 Markdown 转义，避免模板文本把反斜杠带进最终替换结果。
                val = _unescape_markdown_inline(val)
                if val and not re.match(r'^\d+$', val) and val not in ('...', '....'):
                    sentences.append(val)
    return sentences


def _normalize_template_header(header: str) -> str:
    value = re.sub(r'[`\s_]+', '', str(header or '').strip().lower())
    return value.replace('（', '(').replace('）', ')')


def _template_header_key(header: str) -> str | None:
    value = _normalize_template_header(header)
    mapping = {
        '模板编号': 'template_id',
        'templateid': 'template_id',
        '标准句式': 'template_text',
        '句式模板': 'template_text',
        '模板句式': 'template_text',
        'template_text': 'template_text',
        'templatetext': 'template_text',
        'scene': 'scene',
        '场景': 'scene',
        '所属章节': 'section',
        'section': 'section',
        '来源步骤号': 'source_step',
        'sourcestep': 'source_step',
        '句子类型': 'intent_type',
        'intenttype': 'intent_type',
        '动作词': 'actions',
        'actions': 'actions',
        '核心对象': 'objects',
        'objects': 'objects',
        '前置条件': 'conditions',
        '条件': 'conditions',
        'conditions': 'conditions',
        '结果或状态': 'result',
        '结果': 'result',
        'result': 'result',
        '同义表达': 'synonyms',
        '同义词': 'synonyms',
        'synonyms': 'synonyms',
        '是否自动替换': 'auto_apply',
        'autoapply': 'auto_apply',
        '守卫说明': 'guard_notes',
        'guardnotes': 'guard_notes',
    }
    return mapping.get(value)


def _split_template_metadata_values(value: str) -> list[str]:
    return [part.strip() for part in re.split(r'[;；\n]+', str(value or '')) if part.strip()]


def _parse_template_synonym_pairs(value: str) -> list[tuple[str, str]]:
    pairs = []
    for item in _split_template_metadata_values(value):
        match = re.match(r'(.+?)\s*[=＝]\s*(.+)', item)
        if not match:
            continue
        source = match.group(1).strip()
        target = match.group(2).strip()
        if source and target and source != target:
            pairs.append((source, target))
    return pairs


def _parse_structured_sentence_templates(section: str) -> list[dict]:
    templates = []
    table_blocks = re.findall(r'(\|[^\n]+\|\n\|[-:\s|]+\|\n(?:\|[^\n]+\|\n?)+)', section)
    for block in table_blocks:
        lines = [line for line in block.strip().split('\n') if line.strip()]
        if len(lines) < 3:
            continue
        header_row = [cell.strip() for cell in lines[0].strip('|').split('|')]
        header_map = {}
        for idx, header in enumerate(header_row):
            key = _template_header_key(header)
            if key and key not in header_map:
                header_map[key] = idx
        example_col_idx = None
        for idx, header in enumerate(header_row):
            if any(keyword in header for keyword in ['示例', '例子']):
                example_col_idx = idx
                break
        if 'template_text' not in header_map:
            continue
        for line in lines[2:]:
            cells = [cell.strip() for cell in line.strip('|').split('|')]
            template_text = ''
            if example_col_idx is not None and example_col_idx < len(cells):
                template_text = cells[example_col_idx].strip()
            if not template_text and header_map['template_text'] < len(cells):
                template_text = cells[header_map['template_text']].strip()
            template_text = _unescape_markdown_inline(template_text).strip('"\'“”')
            if not template_text or re.match(r'^\d+$', template_text):
                continue
            template = {
                'template_id': cells[header_map['template_id']].strip() if header_map.get('template_id', -1) < len(cells) and header_map.get('template_id') is not None else '',
                'template_text': template_text,
                'scene': cells[header_map['scene']].strip() if header_map.get('scene', -1) < len(cells) and header_map.get('scene') is not None else '',
                'section': cells[header_map['section']].strip() if header_map.get('section', -1) < len(cells) and header_map.get('section') is not None else '',
                'source_step': cells[header_map['source_step']].strip() if header_map.get('source_step', -1) < len(cells) and header_map.get('source_step') is not None else '',
                'intent_type': cells[header_map['intent_type']].strip() if header_map.get('intent_type', -1) < len(cells) and header_map.get('intent_type') is not None else '',
                'actions': _split_template_metadata_values(cells[header_map['actions']]) if header_map.get('actions', -1) < len(cells) and header_map.get('actions') is not None else [],
                'objects': _split_template_metadata_values(cells[header_map['objects']]) if header_map.get('objects', -1) < len(cells) and header_map.get('objects') is not None else [],
                'conditions': _split_template_metadata_values(cells[header_map['conditions']]) if header_map.get('conditions', -1) < len(cells) and header_map.get('conditions') is not None else [],
                'result': cells[header_map['result']].strip() if header_map.get('result', -1) < len(cells) and header_map.get('result') is not None else '',
                'synonym_pairs': _parse_template_synonym_pairs(cells[header_map['synonyms']]) if header_map.get('synonyms', -1) < len(cells) and header_map.get('synonyms') is not None else [],
                'auto_apply': (cells[header_map['auto_apply']].strip().lower() not in {'no', 'false', '0', '否'}) if header_map.get('auto_apply', -1) < len(cells) and header_map.get('auto_apply') is not None else True,
                'guard_notes': cells[header_map['guard_notes']].strip() if header_map.get('guard_notes', -1) < len(cells) and header_map.get('guard_notes') is not None else '',
            }
            templates.append(template)
    return templates


def _template_entry_text(template) -> str:
    if isinstance(template, dict):
        return _protect_model_numbers(_unescape_markdown_inline(str(template.get('template_text', '') or '').strip()))
    return _protect_model_numbers(_unescape_markdown_inline(str(template or '').strip()))


def _template_entry_candidates(template) -> list[str]:
    text = _template_entry_text(template)
    if not text:
        return []
    if not isinstance(template, dict):
        return [text]
    candidates = [text]
    for source, target in template.get('synonym_pairs', []) or []:
        new_candidates = []
        for base in candidates:
            if target in base:
                replaced = base.replace(target, source)
                if replaced and replaced not in candidates and replaced not in new_candidates:
                    new_candidates.append(replaced)
        candidates.extend(new_candidates)
    return candidates


def _structured_template_metadata_score(sentence: str, template) -> dict:
    if not isinstance(template, dict):
        return {'condition': 0.0, 'action': 0.0, 'object': 0.0, 'total': 0.0}
    sentence_intent = _extract_sentence_intent(sentence)
    condition_tokens = [token for token in template.get('conditions', []) if token]
    action_tokens = [token for token in template.get('actions', []) if token]
    object_tokens = [token for token in template.get('objects', []) if token]
    condition_score = _token_overlap_score(sentence_intent.get('conditions', []), condition_tokens) if condition_tokens else 0.0
    action_score = _token_overlap_score(sentence_intent.get('actions', []), action_tokens) if action_tokens else 0.0
    object_score = _token_overlap_score(sentence_intent.get('objects', []), object_tokens) if object_tokens else 0.0
    total = 0.4 * condition_score + 0.25 * action_score + 0.35 * object_score
    return {
        'condition': condition_score,
        'action': action_score,
        'object': object_score,
        'total': total,
    }


def _preferred_sentence_entries(rule: dict, single_clause_only: bool = False) -> list:
    entries = rule.get('templates') or rule.get('sentences') or []
    expanded = []
    seen = set()
    for entry in entries:
        text = _template_entry_text(entry)
        if not text:
            continue
        if not single_clause_only:
            key = text
            if key not in seen:
                expanded.append(entry)
                seen.add(key)
        clauses = _split_sentence_clauses(text)
        if len(clauses) <= 1:
            if single_clause_only and text not in seen:
                expanded.append(entry)
                seen.add(text)
            continue
        for clause in clauses:
            clause_text = clause.strip()
            if len(_normalize_sentence_for_match(clause_text)) < 8 or clause_text in seen:
                continue
            expanded.append(clause_text)
            seen.add(clause_text)
    return expanded


def _is_candidate_recall_rule(rule: dict) -> bool:
    if rule.get('type') != 'preferred_sentences':
        return False
    if rule.get('templates'):
        return True
    sentences = [sentence for sentence in (rule.get('sentences') or []) if _template_entry_text(sentence)]
    if sentences:
        return True
    return '句子分类汇总' in str(rule.get('name') or '')


def _extract_candidate_topic_markers(text: str) -> set[str]:
    raw = str(text or '')
    markers = set()
    if not raw:
        return markers

    for marker in _extract_reference_markers(raw):
        if len(marker) >= 3:
            markers.add(marker)
    for marker in _extract_model_markers(raw).keys():
        if len(marker) >= 3:
            markers.add(marker)
    for marker in _extract_numeric_markers(raw):
        if len(marker) >= 2:
            markers.add(marker)
    for marker in _extract_ui_markers(raw):
        if marker and marker not in _GENERIC_UI_MARKERS:
            markers.add(marker)

    sentence_intent = _extract_sentence_intent(raw)
    candidate_units = (
        sentence_intent.get('objects', []) +
        sentence_intent.get('conditions', []) +
        _extract_result_units(raw) +
        _extract_additional_units(raw)
    )
    for unit in candidate_units:
        normalized = _normalize_structure_object(unit)
        if len(normalized) >= 2 and normalized not in _GENERIC_UI_MARKERS:
            markers.add(normalized)

    for match in re.findall(r'[\u4e00-\u9fffA-Za-z0-9\-]{2,16}(?:试剂|样本|载台|制备卡|产物|孔|仓门|视频|说明书|界面|按钮|文库|模块|仪器|系统|软件|缓冲液)', raw):
        normalized = _normalize_structure_text(match)
        if len(normalized) >= 2 and normalized not in _GENERIC_UI_MARKERS:
            markers.add(normalized)
    return markers


def _template_topic_markers(template) -> set[str]:
    markers = _extract_candidate_topic_markers(_template_entry_text(template))
    if not isinstance(template, dict):
        return markers
    for key in ('scene', 'section', 'source_step', 'intent_type', 'result'):
        normalized = _normalize_structure_text(template.get(key, ''))
        if len(normalized) >= 2:
            markers.add(normalized)
    for key in ('actions', 'objects', 'conditions'):
        for value in template.get(key, []) or []:
            normalized = _normalize_structure_text(value)
            if len(normalized) >= 2 and normalized not in _GENERIC_UI_MARKERS:
                markers.add(normalized)
    return markers


def _filter_candidate_templates(sentence: str, templates: list) -> list:
    if not templates:
        return []
    _, sentence_body = _split_step_prefix(sentence)
    source_sentence = sentence_body or sentence
    if _is_short_label_like_text(source_sentence) or _looks_like_title_or_noun_phrase(source_sentence):
        return []

    source_markers = _extract_candidate_topic_markers(source_sentence)
    source_ui_markers = _extract_ui_markers(source_sentence)
    source_model_markers = set(_extract_model_markers(source_sentence).keys())
    source_number_markers = _extract_numeric_markers(source_sentence)
    source_actions = set(_extract_sentence_intent(source_sentence).get('actions', []))
    source_identity_markers = _extract_action_object_identity_markers(source_sentence)
    source_struct = _extract_sentence_structure(source_sentence)
    source_slot_sample_pairs = _extract_slot_sample_pairs(source_sentence)
    slot_sample_mode = len(source_slot_sample_pairs) >= 3
    ranked = []

    for index, template in enumerate(templates):
        template_text = _template_entry_text(template)
        if not template_text:
            continue
        pair_overlap = 0.0
        if slot_sample_mode:
            template_pairs = _extract_slot_sample_pairs(template_text)
            has_mapping_hint = bool(re.search(r'对应\s*(?:sample|样本)', template_text, re.IGNORECASE))
            if not template_pairs and not has_mapping_hint:
                continue
            pair_overlap = _slot_sample_pair_overlap_score(source_slot_sample_pairs, template_pairs)
            if pair_overlap <= 0 and '孔位' not in template_text:
                continue
        template_markers = _template_topic_markers(template)
        template_actions = set(_extract_sentence_intent(template_text).get('actions', []))
        template_identity_markers = _extract_action_object_identity_markers(template_text)
        if source_identity_markers and template_identity_markers and not (source_identity_markers & template_identity_markers):
            continue
        marker_overlap = len(source_markers & template_markers)
        ui_overlap = len(source_ui_markers & _extract_ui_markers(template_text))
        model_overlap = len(source_model_markers & set(_extract_model_markers(template_text).keys()))
        number_overlap = len(source_number_markers & _extract_numeric_markers(template_text))
        action_overlap = len(source_actions & template_actions)
        similarity = _sentence_similarity(source_sentence, template_text)
        clause_score = _clause_alignment_score(source_sentence, template_text)
        tech_overlap = _tech_term_overlap(source_sentence, template_text)
        if (
            source_actions and not template_actions and
            marker_overlap < 2 and similarity < 0.20 and clause_score < 0.22
        ):
            continue
        score = (
            marker_overlap * 4 +
            ui_overlap * 3 +
            model_overlap * 3 +
            number_overlap * 2 +
            action_overlap * 2 +
            pair_overlap * 8 +
            similarity * 8 +
            clause_score * 10 +
            tech_overlap * 4
        )
        if source_markers and score < 1.2 and action_overlap <= 0 and clause_score < 0.18:
            continue
        ranked.append((score, clause_score, similarity, pair_overlap, marker_overlap, action_overlap, ui_overlap, -index, template))

    if not ranked:
        return [] if source_markers else templates[:40]

    if slot_sample_mode:
        ranked.sort(reverse=True)
        return [item[-1] for item in ranked[:32]]

    if len(ranked) < 24:
        selected_texts = {_template_entry_text(item[-1]) for item in ranked}
        fallback_ranked = []
        for index, template in enumerate(templates):
            template_text = _template_entry_text(template)
            if not template_text or template_text in selected_texts:
                continue
            template_actions = set(_extract_sentence_intent(template_text).get('actions', []))
            action_overlap = len(source_actions & template_actions)
            similarity = _sentence_similarity(source_sentence, template_text)
            clause_score = _clause_alignment_score(source_sentence, template_text)
            if similarity < 0.12 and clause_score < 0.15 and action_overlap <= 0:
                continue
            fallback_score = similarity * 100 + clause_score * 15 + action_overlap * 6
            fallback_ranked.append((fallback_score, clause_score, action_overlap, -index, template))
        fallback_ranked.sort(reverse=True)
        for fallback in fallback_ranked[:24 - len(ranked)]:
            ranked.append((fallback[0], fallback[1], 0, 0, 0, fallback[2], 0, fallback[3], fallback[-1]))

    ranked.sort(reverse=True)
    return [item[-1] for item in ranked[:120]]


def _preferred_entries_from_guide(guide_text: str, single_clause_only: bool = False) -> list:
    cache_key = (_guide_cache_key(guide_text), single_clause_only)
    cached_entries = _preferred_entries_cache.get(cache_key)
    if cached_entries is not None:
        return cached_entries

    structured_entries = []
    scoped_guide = _candidate_recall_guide_text(guide_text)
    for rule in _extract_style_rules(scoped_guide):
        if not _is_candidate_recall_rule(rule):
            continue
        structured_entries.extend(_preferred_sentence_entries(rule, single_clause_only=single_clause_only))
    fallback_entries = _fallback_guide_entries_from_text(scoped_guide)
    entries = []
    seen = set()
    for entry in structured_entries:
        text = _template_entry_text(entry) if isinstance(entry, dict) else str(entry or '').strip()
        if not text or text in seen:
            continue
        seen.add(text)
        entries.append(entry)
    for entry in fallback_entries:
        text = _template_entry_text(entry) if isinstance(entry, dict) else str(entry or '').strip()
        if not text or text in seen or _looks_like_title_or_noun_phrase(text):
            continue
        seen.add(text)
        entries.append(entry)
    _preferred_entries_cache[cache_key] = entries
    return entries


def _sentence_guide_cache_key(sentence: str, guide_text: str, single_clause_only: bool = False, limit: int = 8) -> tuple:
    return (
        _guide_cache_key(guide_text),
        _normalize_operation_sentence_for_match(sentence),
        bool(single_clause_only),
        int(limit),
    )


def _guide_candidate_pool(sentence: str, guide_text: str, single_clause_only: bool = False) -> list:
    cache_key = _sentence_guide_cache_key(sentence, guide_text, single_clause_only=single_clause_only, limit=0)
    cached_pool = _sentence_candidate_pool_cache.get(cache_key)
    if cached_pool is not None:
        return cached_pool

    normalized_sentence = _normalize_operation_sentence_for_match(sentence)
    entries = _preferred_entries_from_guide(guide_text, single_clause_only=single_clause_only)
    candidate_pool = _filter_candidate_templates(normalized_sentence, entries)
    _sentence_candidate_pool_cache[cache_key] = candidate_pool
    return candidate_pool


def _guide_top_template_candidates(sentence: str, guide_text: str, limit: int = 8, single_clause_only: bool = False) -> list[dict]:
    cache_key = _sentence_guide_cache_key(sentence, guide_text, single_clause_only=single_clause_only, limit=limit)
    cached_candidates = _sentence_top_candidates_cache.get(cache_key)
    if cached_candidates is not None:
        return cached_candidates

    normalized_sentence = _normalize_operation_sentence_for_match(sentence)
    candidate_pool = _guide_candidate_pool(normalized_sentence, guide_text, single_clause_only=single_clause_only)
    candidates = _top_template_candidates(normalized_sentence, candidate_pool, limit=limit)
    _sentence_top_candidates_cache[cache_key] = candidates
    return candidates


def _preferred_entries_for_sentence(sentence: str, guide_text: str, single_clause_only: bool = False) -> list:
    candidate_pool = _guide_candidate_pool(sentence, guide_text, single_clause_only=single_clause_only)
    if candidate_pool:
        return candidate_pool
    return _preferred_entries_from_guide(guide_text, single_clause_only=single_clause_only)


def _extract_style_rules(guide_text: str) -> list[dict]:
    """从句式指南中提取可执行的检测规则，解析 Markdown 内容"""
    cache_key = _guide_cache_key(guide_text)
    cached_rules = _style_rules_cache.get(cache_key)
    if cached_rules is not None:
        return cached_rules

    rules = []

    def _extract_numbered_sentences(text: str) -> list[str]:
        sentences = []
        for line in (text or '').splitlines():
            value = line.strip()
            if not value:
                continue
            match = re.match(r'^(?:\d+[.)]|[一二三四五六七八九十]+[、.])\s+(.+)$', value)
            if not match:
                continue
            sentence = match.group(1).strip()
            if len(sentence) < 6:
                continue
            if sentence in {'---'}:
                continue
            if sentence not in sentences:
                sentences.append(sentence)
        return sentences
    
    # 从 guide_text 中解析实际规则，否则使用默认规则
    if guide_text and guide_text.strip():
        # 解析 ## 开头的章节作为规则类别
        sections = re.split(r'\n(?=##\s)', guide_text)
        for section in sections:
            header_match = re.match(r'##\s+(.+)', section)
            if not header_match:
                continue
            header = header_match.group(1).strip()
            
            # 提取列表项作为具体规则，兼容无序列表和编号列表
            items = re.findall(r'(?:^|\n)(?:[-*]|\d+[.)])\s+(.+)', section)
            
            # 同时解析 Markdown 表格中的句式模板
            structured_templates = _parse_structured_sentence_templates(section)
            table_sentences = [template.get('template_text', '') for template in structured_templates if template.get('template_text')]
            if not table_sentences:
                table_sentences = _parse_table_sentence_templates(section)
            
            if not items and not table_sentences:
                continue
            
            # 判断规则类型
            lower_header = header.lower()
            if any(kw in lower_header for kw in ['禁用', '禁止', 'forbidden', '避免使用', '不要用']):
                # 从列表中提取短语
                phrases = []
                replacements = {}
                for item in items:
                    # 匹配 "A → B" 或 "A -> B" 或 "A：B" 格式
                    arrow_match = re.match(r'(.+?)\s*(?:→|->|：)\s*(.+)', item)
                    if arrow_match:
                        old_phrase = arrow_match.group(1).strip().strip('"\'「」""''')
                        new_phrase = arrow_match.group(2).strip().strip('"\'「」""''')
                        phrases.append(old_phrase)
                        replacements[old_phrase] = new_phrase
                    else:
                        phrase = item.strip().strip('"\'「」""''')
                        if phrase:
                            phrases.append(phrase)
                
                if phrases:
                    rules.append({
                        "type": "forbidden_words",
                        "name": header,
                        "patterns": phrases,
                        "replacements": replacements if replacements else None,
                        "fix": "替换为更规范的表达"
                    })
            
            elif any(kw in lower_header for kw in ['被动', 'passive', '语态']):
                patterns = []
                for item in items:
                    item = item.strip().strip('`')
                    if item:
                        patterns.append((item, "主动语态"))
                if patterns:
                    rules.append({
                        "type": "passive_voice",
                        "name": header,
                        "patterns": patterns,
                        "fix": "改用主动语态"
                    })
            
            elif any(kw in lower_header for kw in ['双重否定', 'double negative']):
                patterns = []
                for item in items:
                    item = item.strip().strip('`')
                    if item:
                        patterns.append((item, "双重否定"))
                if patterns:
                    rules.append({
                        "type": "double_negative",
                        "name": header,
                        "patterns": patterns,
                        "fix": "改用肯定表达"
                    })
            
            elif any(kw in lower_header for kw in ['非正式', 'informal', '口语', '俚语']):
                patterns = []
                for item in items:
                    item = item.strip().strip('`')
                    if item:
                        patterns.append((item, "非正式语言"))
                if patterns:
                    rules.append({
                        "type": "informal",
                        "name": header,
                        "patterns": patterns,
                        "fix": "使用正式表达"
                    })
            
            elif any(kw in lower_header for kw in ['句子长度', 'sentence length', '长句', '字数']):
                max_chars = 100
                for item in items:
                    num_match = re.search(r'(\d+)', item)
                    if num_match:
                        max_chars = int(num_match.group(1))
                        break
                rules.append({
                    "type": "sentence_length",
                    "name": header,
                    "max_chars": max_chars,
                    "fix": "拆分过长句子"
                })
            
            elif any(kw in lower_header for kw in ['代词', 'pronoun', '指代']):
                patterns = []
                for item in items:
                    item = item.strip().strip('`')
                    if item:
                        patterns.append((item, "代词指代不明确"))
                if patterns:
                    rules.append({
                        "type": "pronoun_reference",
                        "name": header,
                        "patterns": patterns,
                        "fix": "明确指代对象"
                    })
            
            elif any(kw in lower_header for kw in ['术语', 'terminology', '词汇']):
                term_map = {}
                for item in items:
                    arrow_match = re.match(r'(.+?)\s*(?:→|->|：)\s*(.+)', item)
                    if arrow_match:
                        old_term = arrow_match.group(1).strip().strip('"\'「」""''')
                        new_term = arrow_match.group(2).strip().strip('"\'「」""''')
                        term_map[old_term] = new_term
                if term_map:
                    rules.append({
                        "type": "terminology_rule",
                        "name": header,
                        "term_map": term_map,
                        "fix": "统一术语"
                    })

            elif any(kw in lower_header for kw in ['用户反馈修正', '推荐句式', '优先句式', '反馈句式', '句式', '模板', '句型']):
                preferred_sentences = []
                for item in items:
                    sentence = item.strip().strip('`').strip()
                    if sentence:
                        preferred_sentences.append(sentence)
                # 如果表格中也有句式模板，合并进来
                if table_sentences:
                    preferred_sentences.extend(table_sentences)
                if preferred_sentences:
                    rules.append({
                        "type": "preferred_sentences",
                        "name": header,
                        "sentences": preferred_sentences,
                        "templates": structured_templates or None,
                        "fix": "优先采用用户确认过的句式"
                    })
            
            # 表格句式模板：从句式参考表中提取句型（跳过统计表）
            elif table_sentences and '统计' not in header:
                rules.append({
                    "type": "preferred_sentences",
                    "name": header,
                    "sentences": table_sentences,
                    "templates": structured_templates or None,
                    "fix": "采用规范句式"
                })
    
    if '句子分类汇总' in (guide_text or ''):
        numbered_sentences = _extract_numbered_sentences(guide_text)
        if numbered_sentences:
            rules.append({
                "type": "preferred_sentences",
                "name": "句子分类汇总",
                "sentences": numbered_sentences,
                "fix": "优先采用汇总文件中的规范句式"
            })

    _style_rules_cache[cache_key] = rules
    return rules


def _default_style_rules() -> list[dict]:
    """默认句式规则（当 KB 中无自定义规则时使用）"""
    return [
        {
            "type": "forbidden_words",
            "name": "禁用强调词",
            "patterns": ["最佳", "最好", "最著名", "最新技术", "最高水平", "最先进水平", "最高技术", "非常", "极其"],
            "fix": "删除或替换为客观描述"
        },
        {
            "type": "passive_voice",
            "name": "被动转主动",
            "patterns": [(r'被(用于|应用于|设计|制造|创建|安装|设置|配置|提供|调用|使用)', "主动语态"), (r'由.+?提供(了)?', "主动语态")],
            "fix": "改用主动语态"
        },
        {
            "type": "double_negative",
            "name": "避免双重否定",
            "patterns": [(r'不(能|得|可|允许).+?不(能|得|可|允许)', "双重否定"), (r'没(有)?.+?不(能|得|可)', "双重否定"), (r'非.+?不', "双重否定")],
            "fix": "改用肯定表达"
        },
        {
            "type": "informal",
            "name": "避免非正式语言",
            "patterns": [(r'[酷毙|爽翻|给力|碉堡|牛逼]', "非正式语言"), (r'！{2,}', "过度感叹"), (r'～{2,}', "过度波浪号")],
            "fix": "使用正式表达"
        },
        {
            "type": "sentence_length",
            "name": "句子长度控制",
            "max_chars": 100,
            "fix": "拆分长句"
        },
        {
            "type": "pronoun_reference",
            "name": "代词指代明确",
            "patterns": [(r'其[中他它们她]', "代词指代不明"), (r'该.+?(?!系统|设备|产品|方法|技术|功能|模块|参数|配置)', "代词指代不明")],
            "fix": "明确指代对象"
        },
    ]


# 同义词归一化组：每组第一个词为标准形式
_SYNONYM_GROUPS = [
    ('点击', ['单击', '按下']),
    ('选择', ['选中', '勾选']),
    ('打开', ['开启', '启动']),
    ('关闭', ['停止', '关停']),
    ('进入', ['切换到', '跳转至']),
    ('放入', ['放置于', '置于']),
    ('取出', ['拿出']),
    ('输入', ['录入', '填写']),
    ('确认', ['确定', '核实']),
    ('安装', ['装载', '装配']),
]

_SYNONYM_CANONICAL = {}
for _canonical, _syns in _SYNONYM_GROUPS:
    _SYNONYM_CANONICAL[_canonical] = _canonical
    for _syn in _syns:
        _SYNONYM_CANONICAL[_syn] = _canonical

# 模块级句子匹配缓存：同一句子不重复匹配
_polish_sentence_cache: dict = {}
_preferred_entries_cache: dict = {}
_style_rules_cache: dict = {}
_sentence_candidate_pool_cache: dict = {}
_sentence_top_candidates_cache: dict = {}
_logged_sentence_guide_keys: set[str] = set()


def _guide_cache_key(text: str) -> str:
    if not text:
        return ''
    return hashlib.md5(text.encode('utf-8')).hexdigest()


def _apply_synonym_normalization(text: str) -> str:
    """同义词归一化：将同义词统一为标准形式，用于 L1.5 匹配。"""
    result = text
    for _word, _canonical in _SYNONYM_CANONICAL.items():
        result = result.replace(_word, _canonical)
    return result


def _three_tier_match(sentence: str, templates: list[str]) -> tuple:
    """
    三层匹配策略：L1 精确 → L1.5 同义词 → L2 模糊 → L3 语义。
    返回: (best_template, best_score, match_level)
    """
    normalized = _normalize_sentence_for_match(sentence)
    if not normalized:
        return None, 0.0, "NONE"

    # ===== L1: 精确匹配（含变体） =====
    variants = _generate_sentence_variants(sentence)
    for tmpl in templates:
        tmpl_normalized = _normalize_sentence_for_match(tmpl)
        if not tmpl_normalized:
            continue
        # 检查输入变体 → 模板
        for variant in variants:
            var_normalized = _normalize_sentence_for_match(variant)
            if var_normalized == tmpl_normalized:
                return tmpl, 1.0, "L1"
        # 检查输入变体 → 模板变体
        tmpl_variants = _generate_sentence_variants(tmpl)
        for tv in tmpl_variants:
            tv_normalized = _normalize_sentence_for_match(tv)
            for variant in variants:
                var_normalized = _normalize_sentence_for_match(variant)
                if var_normalized == tv_normalized:
                    return tmpl, 1.0, "L1"

    # ===== L1.5: 同义词归一化后精确匹配 =====
    sent_syn = _apply_synonym_normalization(normalized)
    if sent_syn != normalized:
        for tmpl in templates:
            tmpl_normalized = _normalize_sentence_for_match(tmpl)
            if not tmpl_normalized:
                continue
            tmpl_syn = _apply_synonym_normalization(tmpl_normalized)
            if sent_syn == tmpl_syn:
                return tmpl, 0.95, "L1.5"

    # ===== L2: 模糊匹配（bigram+LCS+结构+关键词 高阈值） =====
    best_tmpl = None
    best_score = 0.0
    best_rank = None
    for tmpl in templates:
        tmpl_normalized = _normalize_sentence_for_match(tmpl)
        if not tmpl_normalized:
            continue
        structured_score = _structured_match_score(sentence, tmpl)
        cat_score = _cat_match_score(sentence, tmpl, structured_score=structured_score)
        final_score = cat_score.get('ranking_score', cat_score['overall_score'])
        rank = (
            final_score,
            cat_score['term_anchor_score'],
            cat_score['number_placeholder_score'],
            cat_score['context_score'],
            structured_score['ui_target'],
            structured_score['sequence'],
            structured_score['pair'],
            structured_score['condition'],
            structured_score['action'],
            len(tmpl_normalized),
        )
        if best_rank is None or rank > best_rank:
            best_rank = rank
            best_score = final_score
            best_tmpl = tmpl
    
    if best_tmpl and best_score >= _POLISH_MATCH_CONFIG["l2_min_confidence"]:
        return best_tmpl, best_score, "L2"
    
    # ===== L3: 语义匹配（低阈值，仅做提示） =====
    l3_threshold = _POLISH_MATCH_CONFIG["l3_auto_confidence"]
    if best_tmpl and best_score >= l3_threshold:
        return best_tmpl, best_score, "L3"
    
    return best_tmpl, best_score, "NONE"


def _ai_semantic_rerank_candidates(sentence: str, ranked: list[dict]) -> list[dict]:
    if not _ai_template_rerank_enabled.get():
        return ranked

    shortlist = []
    seen_keys = set()
    for item in ranked or []:
        key = (item.get('template', ''), item.get('candidate_text', ''))
        if key in seen_keys:
            continue
        seen_keys.add(key)

        final_score = float(item.get('final_score', 0.0) or 0.0)
        overall_percent = int(item.get('overall_percent', 0) or 0)
        if final_score >= 0.1 or overall_percent >= 20:
            shortlist.append(item)
        if len(shortlist) >= _AI_TEMPLATE_RERANK_LIMIT:
            break

    if len(shortlist) < 2:
        shortlist = []
        seen_keys.clear()
        for item in ranked or []:
            key = (item.get('template', ''), item.get('candidate_text', ''))
            if key in seen_keys:
                continue
            seen_keys.add(key)
            shortlist.append(item)
            if len(shortlist) >= min(_AI_TEMPLATE_RERANK_LIMIT, 2):
                break

    if len(shortlist) < 2:
        return ranked

    cache_key = hashlib.sha1(
        (
            _normalize_compare_text(sentence) + '\n' +
            '\n'.join(f"{item.get('template', '')}\u0001{item.get('candidate_text', '')}" for item in shortlist)
        ).encode('utf-8')
    ).hexdigest()
    cached = _ai_template_rerank_cache.get(cache_key)
    if cached is None:
        try:
            from app.utils.ai_client import ai_client
        except Exception:
            ai_client = None
        if not ai_client or not ai_client.has_any_client:
            return ranked

        candidates_text = '\n'.join(
            f"[{index}] 候选句式：{item.get('candidate_text') or item.get('template', '')}"
            for index, item in enumerate(shortlist, 1)
        )
        prompt = f"""你是技术文档句式匹配专家。请根据原句判断下面候选句式的语义匹配度和改写适用性。

原句：
{sentence}

候选句式：
{candidates_text}

判断要求：
- score 取值 0-100，表示候选句式与原句在语义、动作、对象、数字、表图引用上的综合匹配度。
- recommended=true 表示该候选句式适合直接作为润色建议。
- 产品名、型号、数字、数量、表图引用、关键对象不能丢失或改错。
- 仅做轻微空格修正的句子，如果候选句式表达更完整且保留关键事实，可以推荐。
- 只输出 JSON，不要输出其他文字。

输出格式：
{{
  "items": [
    {{"index": 1, "score": 92, "recommended": true, "reason": "简短理由"}}
  ]
}}"""
        try:
            from app.utils.cat_diagnose import lab_ai_chat
            result = lab_ai_chat(
                [{"role": "user", "content": prompt}],
                max_tokens=1200,
                temperature=0.1,
                request_label="polish.template.rerank",
            )
            parsed = json.loads(re.sub(r'^```[a-zA-Z]*\n?|\n?```$', '', str(result or '').strip())) if result else {"items": []}
        except Exception:
            parsed = {"items": []}
        cached = {}
        for item in parsed.get('items', []) if isinstance(parsed, dict) else []:
            if not isinstance(item, dict):
                continue
            index = int(item.get('index', 0) or 0)
            if index < 1 or index > len(shortlist):
                continue
            key = (shortlist[index - 1].get('template', ''), shortlist[index - 1].get('candidate_text', ''))
            score = max(0, min(100, int(item.get('score', 0) or 0)))
            cached[key] = {
                'score': score,
                'recommended': bool(item.get('recommended')),
                'reason': str(item.get('reason', '') or '').strip(),
            }
        _ai_template_rerank_cache[cache_key] = cached

    if not cached:
        return ranked

    reranked = []
    for item in ranked:
        next_item = dict(item)
        ai_info = cached.get((item.get('template', ''), item.get('candidate_text', '')))
        base_score = float(item.get('final_score', 0.0) or 0.0)
        if ai_info:
            ai_score = round(ai_info['score'] / 100.0, 4)
            next_item['ai_semantic_score'] = ai_info['score']
            next_item['ai_semantic_recommended'] = ai_info['recommended']
            next_item['ai_semantic_reason'] = ai_info['reason']
            next_item['ai_rank_score'] = round(base_score * 0.7 + ai_score * 0.3, 4)
        else:
            next_item['ai_semantic_score'] = None
            next_item['ai_semantic_recommended'] = False
            next_item['ai_semantic_reason'] = ''
            next_item['ai_rank_score'] = round(base_score, 4)
        reranked.append(next_item)

    reranked.sort(
        key=lambda item: (
            item.get('ai_semantic_recommended', False),
            item.get('guard_passed', False),
            item.get('ai_rank_score', 0.0),
            item.get('final_score', 0.0),
            item.get('overall_score', 0.0),
            len(_normalize_sentence_for_match(item.get('template', ''))),
        ),
        reverse=True,
    )
    return reranked


_CAT_SEMANTIC_SCORING_PROMPT = """你是语义匹配评分专家。对每个原句，判断其与各候选句在语义上的匹配度。

{sentences}

输出要求：
1. 严格输出 JSON 数组，不要输出其他文字、不要用 markdown 代码块包裹。
2. 每个候选项给出：
   - semantic_score: 0.0~1.0，基于语义相似度判断（术语/动作/对象/数字是否一致）
   - reason: {reason_max}字以内，说明推荐或不推荐的理由
3. 只评分和给理由，绝对不要改写任何句子。
4. 原句与候选的内容、关键术语、数字、单位必须保留一致。

输出格式：
[
  {{"sentence_index": 0, "candidates": [{{"index": 0, "semantic_score": 0.92, "reason": "术语匹配，仅数字不同"}}]}},
  {{"sentence_index": 1, "candidates": [{{"index": 0, "semantic_score": 0.45, "reason": "句式结构差异大"}}]}}
]
"""


async def _batch_ai_semantic_score(
    items: list,
    reason_max_chars: int = 15,
) -> dict:
    """
    批量调用 AI，对所有有候选的行打语义分 + 给推荐理由。
    直接修改 items 中每个元素的 candidates 列表（in-place）。
    """
    scored_items = []
    for item in items:
        candidates = item.candidates if hasattr(item, 'candidates') else item.get('candidates', [])
        if not candidates:
            continue
        scored_items.append(item)

    if not scored_items:
        return {
            "status": "skipped",
            "error": "无候选句子需要评分",
            "scored_count": 0,
        }

    try:
        from app.utils.ai_client import ai_client
    except Exception:
        ai_client = None

    if not ai_client or not ai_client.has_any_client:
        return {
            "status": "no_api_key",
            "error": "未配置可用的 AI 客户端",
            "scored_count": 0,
        }

    batch_size = 8
    scored_count = 0
    partial_failure = False
    last_error = None

    for batch_start in range(0, len(scored_items), batch_size):
        batch = scored_items[batch_start:batch_start + batch_size]
        parts = []
        total_candidates = 0
        for i, item in enumerate(batch):
            original = item.original_text if hasattr(item, 'original_text') else item.get('original_text', '')
            candidates = item.candidates if hasattr(item, 'candidates') else item.get('candidates', [])
            parts.append(f"[原句{i}] {original}")
            for j, c in enumerate(candidates):
                tpl_text = c.get('template_text', '') if isinstance(c, dict) else getattr(c, 'template_text', '')
                str_score = c.get('string_score', 0) if isinstance(c, dict) else getattr(c, 'string_score', 0)
                parts.append(f"  候选{j}(字符串匹配度:{str_score:.2f}): {tpl_text}")
                total_candidates += 1

        prompt = _CAT_SEMANTIC_SCORING_PROMPT.format(
            sentences="\n".join(parts),
            reason_max=reason_max_chars,
        )
        dynamic_max_tokens = min(max(1200, total_candidates * 50), 5000)

        try:
            from app.utils.cat_diagnose import lab_ai_chat
            result = lab_ai_chat(
                [{"role": "user", "content": prompt}],
                max_tokens=dynamic_max_tokens,
                temperature=0.1,
                request_label="polish.cat.semantic_score",
                timeout=90,
            )
        except Exception as e:
            partial_failure = True
            last_error = str(e)
            logger.warning(
                "[CAT_SCORING] AI 调用失败(batch=%s, provider_errors=%s): %s",
                batch_start // batch_size,
                getattr(ai_client, 'last_chat_errors', []),
                e,
            )
            continue

        if not result:
            partial_failure = True
            provider_errors = getattr(ai_client, 'last_chat_errors', [])
            last_error = provider_errors[-1] if provider_errors else "AI 未返回评分结果"
            logger.warning(
                "[CAT_SCORING] AI 未返回评分结果(batch=%s, provider_errors=%s)",
                batch_start // batch_size,
                provider_errors,
            )
            continue

        try:
            cleaned = re.sub(r'^```[a-zA-Z]*\n?|\n?```$', '', str(result).strip())
            parsed = json.loads(cleaned)
        except Exception as e:
            partial_failure = True
            last_error = str(e)
            logger.warning("[CAT_SCORING] JSON 解析失败(batch=%s): %s, raw=%s", batch_start // batch_size, e, str(result)[:200])
            continue

        if not isinstance(parsed, list):
            partial_failure = True
            last_error = "AI 返回结果格式不正确"
            continue

        for entry in parsed:
            if not isinstance(entry, dict):
                continue
            idx = int(entry.get("sentence_index", -1))
            if idx < 0 or idx >= len(batch):
                continue
            item = batch[idx]
            candidates = item.candidates if hasattr(item, 'candidates') else item.get('candidates', [])

            for c_result in entry.get("candidates", []):
                if not isinstance(c_result, dict):
                    continue
                c_idx = int(c_result.get("index", -1))
                if c_idx < 0 or c_idx >= len(candidates):
                    continue
                c = candidates[c_idx]
                score = round(max(0.0, min(1.0, float(c_result.get("semantic_score", 0)))), 4)
                reason = str(c_result.get("reason", "")).strip()[:reason_max_chars]
                if isinstance(c, dict):
                    c["semantic_score"] = score
                    c["ai_reason"] = reason
                else:
                    c.semantic_score = score
                    c.ai_reason = reason
                scored_count += 1

        if batch_start + batch_size < len(scored_items):
            await asyncio.sleep(0.5)

    if scored_count > 0:
        return {
            "status": "completed",
            "error": last_error if partial_failure else None,
            "scored_count": scored_count,
        }

    return {
        "status": "error" if partial_failure else "failed",
        "error": last_error or "AI 调用完成但未返回有效评分",
        "scored_count": 0,
    }


def _normalized_edit_distance(a: str, b: str) -> float:
    """
    编辑距离归一化相似度：1 - lev_distance(a,b) / max(len(a), len(b))
    返回 0.0~1.0，1.0 表示完全相同。
    """
    if not a and not b:
        return 1.0
    max_len = max(len(a), len(b))
    if max_len == 0:
        return 1.0
    return SequenceMatcher(None, a, b).ratio()


_CAT_STOP_WORDS = set(
    "的 了 和 与 在 是 有 对 为 从 向 把 被 让 给 跟 同 以 按 "
    "请 确保 需要 必须 应 应当 可以 能够 已经 将 会 要 都 也 还 "
    "这 那 这个 那个 这些 那些 其 该 此 本".split()
)

_CAT_AI_SEMANTIC_SHOW_THRESHOLD = 0.76
_CAT_AI_SEMANTIC_PENDING_THRESHOLD = 0.64
_CAT_RULE_ONLY_SHOW_THRESHOLD = 0.65
_CAT_RULE_ONLY_PENDING_THRESHOLD = 0.50
_CAT_RULE_ONLY_LEGACY_PENDING_THRESHOLD = 0.54

_CAT_KEYWORD_ENTITY_PATTERNS = [
    (re.compile(r'\b(?:货号|型号|编号|序列号|No\.?|Part\s*#|CAT\s*#)?\s*[:：]?\s*([A-Za-z0-9\-]{6,})\b', re.IGNORECASE), "编号/货号/型号"),
    (re.compile(r'(?:套件|试剂盒|Kit|制备套件)\s*([A-Z])\b'), "套件版本"),
    (re.compile(r'\b[Vv]\s*(\d+(?:\.\d+)*)\b'), "版本号"),
    (re.compile(r'\b([A-Z]{2,}SEQ-[A-Za-z0-9]+(?:\s+[A-Z]{2}(?![A-Za-z]))?)'), "测序平台"),
    (re.compile(r'\b(DNBelab-[A-Za-z0-9]+(?:\s+[A-Z]{2}(?![A-Za-z]))?)'), "仪器型号"),
    (re.compile(r'([A-Za-z0-9\-]*[\u4e00-\u9fffA-Za-z0-9\-（）()]{0,24}试剂套装)'), "试剂套装名"),
    (re.compile(r'@([A-Za-z0-9.-]+\.[A-Za-z]{2,})', re.IGNORECASE), "邮箱域名"),
    (re.compile(r'(\d+(?:\.\d+)?)\s*(μL|µL|uL|mL|ng|μg|µg|mg|°C|℃|°)(?![A-Za-z])', re.IGNORECASE), "数值"),
]

_CAT_ENTITY_CHANGE_SCORE_FACTOR = 0.82

_CAT_ENTITY_BACKFILL_LABELS = {
    '测序平台',
    '仪器型号',
    '试剂套装名',
}
_CAT_MODEL_BACKFILL_LABELS = {'测序平台', '仪器型号'}
_CN_CARDINAL_DIGITS = {
    '零': 0, '一': 1, '二': 2, '两': 2, '三': 3, '四': 4,
    '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
}

_CAT_KEY_TERM_ANCHOR_GROUPS = {
    'temperature': (
        ('室温', re.compile(r'室温')),
        ('冰上', re.compile(r'冰上')),
        ('4c', re.compile(r'(?:4\s*[℃°cC]|冷藏)')),
        ('-20c', re.compile(r'-\s*20\s*[℃°cC]')),
        ('-80c', re.compile(r'-\s*80\s*[℃°cC]')),
        ('-15c', re.compile(r'-\s*15\s*[℃°cC]')),
        ('37c', re.compile(r'37\s*[℃°cC]')),
        ('65c', re.compile(r'65\s*[℃°cC]')),
    ),
    'molecule': (
        ('totalrna', re.compile(r'total\s*rna', re.IGNORECASE)),
        ('mrna', re.compile(r'm\s*rna', re.IGNORECASE)),
        ('rrna', re.compile(r'r\s*rna', re.IGNORECASE)),
        ('trna', re.compile(r't\s*rna', re.IGNORECASE)),
        ('cdna', re.compile(r'c\s*dna', re.IGNORECASE)),
        ('gdna', re.compile(r'g\s*dna', re.IGNORECASE)),
        ('cfdna', re.compile(r'cf\s*dna', re.IGNORECASE)),
        ('rna', re.compile(r'(?<![a-z])rna(?![a-z])', re.IGNORECASE)),
        ('dna', re.compile(r'(?<![a-z])dna(?![a-z])', re.IGNORECASE)),
    ),
    'organism': (
        ('人', re.compile(r'人')),
        ('鼠', re.compile(r'(?<!小)鼠')),
        ('小鼠', re.compile(r'小鼠')),
        ('大肠杆菌', re.compile(r'大肠杆菌')),
        ('酵母', re.compile(r'酵母')),
        ('病原微生物', re.compile(r'病原微生物')),
        ('细菌', re.compile(r'细菌')),
        ('真菌', re.compile(r'真菌')),
        ('病毒', re.compile(r'病毒')),
    ),
    'workflow': (
        ('样本', re.compile(r'样本')),
        ('检测', re.compile(r'检测')),
        ('建库', re.compile(r'建库|文库制备')),
        ('提取', re.compile(r'提取')),
        ('扩增', re.compile(r'扩增')),
        ('酶切', re.compile(r'酶切')),
        ('逆转录', re.compile(r'逆转录')),
        ('纯化', re.compile(r'纯化')),
    ),
    'container': (
        ('试剂套装', re.compile(r'试剂套装')),
        ('试剂盒', re.compile(r'试剂盒')),
        ('试剂管', re.compile(r'试剂管')),
        ('试剂板', re.compile(r'试剂板')),
        ('制备卡', re.compile(r'制备卡')),
        ('样本制备卡', re.compile(r'样本制备卡')),
        ('芯片', re.compile(r'芯片')),
        ('孔板', re.compile(r'孔板')),
        ('孔位', re.compile(r'孔位')),
    ),
    'condition_action': (
        ('室温解冻', re.compile(r'室温解冻')),
        ('冰上解冻', re.compile(r'冰上解冻')),
        ('常温解冻', re.compile(r'常温解冻')),
        ('水浴解冻', re.compile(r'水浴解冻')),
        ('置于冰上', re.compile(r'置于冰上')),
        ('冰上备用', re.compile(r'冰上备用')),
        ('室温放置', re.compile(r'室温放置')),
        ('离心', re.compile(r'离心')),
        ('混匀', re.compile(r'混匀')),
    ),
    'operation_mode': (
        ('手动', re.compile(r'手动')),
        ('自动', re.compile(r'自动')),
    ),
}

_CAT_NO_CHANGE_REASON_KEYWORDS = (
    '完全一致',
    '仅标点差异',
    '仅标点不同',
    '无需修改',
    '无需润色',
    '表达一致',
)

_CAT_NEGATIVE_REASON_KEYWORDS = (
    '错误',
    '冲突',
    '背离',
    '不等价',
    '不相同',
    '仅提取',
    '片段',
    '缺',
    '丢失',
    '缺少',
    '改变',
    '不一致',
    '不推荐',
    '不建议',
    'vs',
)

_CAT_TECH_TERM_PATTERNS = [
    re.compile(r'pooling', re.IGNORECASE),
    re.compile(r'测序'),
    re.compile(r'投入量'),
    re.compile(r'出库浓度'),
    re.compile(r'样本'),
    re.compile(r'制备卡'),
    re.compile(r'凹口'),
    re.compile(r'凸出位置'),
    re.compile(r'仪器载台'),
    re.compile(r'图\s*\d+'),
]

_CAT_CONDITION_ACTION_CONFLICT_GROUPS = {
    'thaw_method': {'室温解冻', '常温解冻', '冰上解冻', '水浴解冻'},
    'post_thaw_placement': {'置于冰上', '冰上备用', '室温放置'},
}


def _tokenize_cat_text(text: str) -> list[str]:
    normalized = re.sub(r'[^\w\u4e00-\u9fff]+', ' ', str(text or '')).strip()
    if not normalized:
        return []
    try:
        import jieba  # type: ignore
        tokens = jieba.lcut(normalized)
    except Exception:
        chars = re.findall(r'[\u4e00-\u9fff]', normalized)
        tokens = [chars[i] + chars[i + 1] for i in range(len(chars) - 1)]
        tokens.extend(re.findall(r'[A-Za-z0-9]+', normalized))
    return [token.strip().lower() for token in tokens if token and token.strip() and token.strip() not in _CAT_STOP_WORDS]


def _word_ngram_similarity(a: str, b: str, n: int = 2) -> float:
    tokens_a = _tokenize_cat_text(a)
    tokens_b = _tokenize_cat_text(b)
    if not tokens_a or not tokens_b:
        return 0.0
    if len(tokens_a) < n or len(tokens_b) < n:
        set_a = set(tokens_a)
        set_b = set(tokens_b)
        union = set_a | set_b
        return len(set_a & set_b) / len(union) if union else 0.0

    def _word_ngrams(tokens: list[str], size: int) -> set[str]:
        return {" ".join(tokens[i:i + size]) for i in range(len(tokens) - size + 1)}

    grams_a = _word_ngrams(tokens_a, n)
    grams_b = _word_ngrams(tokens_b, n)
    union = grams_a | grams_b
    return len(grams_a & grams_b) / len(union) if union else 0.0


def _word_set_overlap(a: str, b: str) -> float:
    tokens_a = set(_tokenize_cat_text(a))
    tokens_b = set(_tokenize_cat_text(b))
    if not tokens_a or not tokens_b:
        return 0.0
    longer = max(len(tokens_a), len(tokens_b))
    return len(tokens_a & tokens_b) / longer if longer else 0.0


def _ngram_similarity(a: str, b: str, n: int = 2) -> float:
    """
    n-gram Jaccard 相似度：两个字符串的 n-gram 集合的交集/并集。
    返回 0.0~1.0。
    """
    if not a or not b:
        return 0.0
    if len(a) < n or len(b) < n:
        return _normalized_edit_distance(a, b)

    def _ngrams(s: str, n: int) -> set:
        return {s[i:i+n] for i in range(len(s) - n + 1)}

    ga = _ngrams(a, n)
    gb = _ngrams(b, n)
    if not ga or not gb:
        return 0.0
    intersection = ga & gb
    union = ga | gb
    return len(intersection) / len(union) if union else 0.0


def _calc_similarity(a: str, b: str) -> float:
    """
    综合字符串相似度 = 0.3 × 编辑距离 + 0.3 × 词级 n-gram + 0.4 × 词集合重叠。
    返回 0.0~1.0。
    """
    if not a or not b:
        return 0.0
    edit_sim = _normalized_edit_distance(a, b)
    word_ngram_sim = _word_ngram_similarity(a, b, n=2)
    word_overlap = _word_set_overlap(a, b)
    return round(0.3 * edit_sim + 0.3 * word_ngram_sim + 0.4 * word_overlap, 4)


def _extract_critical_entities(text: str) -> dict[str, list[str]]:
    entities: dict[str, list[str]] = {}
    for label, raw, _, _ in _extract_critical_entity_spans(text):
        value = _normalize_entity_value(raw)
        if value:
            entities.setdefault(label, []).append(value)
    return entities


def _critical_entities_compatible(sentence: str, candidate: str) -> tuple[bool, str]:
    changes = _critical_entity_changes(sentence, candidate)
    if not changes:
        return True, ''
    return False, changes[0]


def _critical_entity_changes(sentence: str, candidate: str) -> list[str]:
    sentence_entities = _extract_critical_entities(sentence)
    candidate_entities = _extract_critical_entities(candidate)
    if not sentence_entities:
        return []
    changes: list[str] = []
    for label, sentence_values in sentence_entities.items():
        candidate_values = candidate_entities.get(label, [])
        if not candidate_values:
            changes.append(f'{label}缺失')
        elif not (set(sentence_values) & set(candidate_values)):
            changes.append(f'{label}不一致')
    return changes


def _apply_cat_entity_change_penalty(source_text: str, candidate: dict) -> dict:
    if not isinstance(candidate, dict):
        return candidate
    display = str(candidate.get('template_text') or '')
    changes = _critical_entity_changes(source_text, display)
    if not changes:
        return candidate
    original_score = float(candidate.get('string_score') or 0.0)
    candidate['unpenalized_string_score'] = original_score
    candidate['string_score'] = round(original_score * _CAT_ENTITY_CHANGE_SCORE_FACTOR, 4)
    candidate['entity_penalty_reason'] = '；'.join(changes)
    tags = list(candidate.get('review_tags') or [])
    if 'entity_changed' not in tags:
        tags.append('entity_changed')
    candidate['review_tags'] = tags
    return candidate


def _normalize_entity_value(value: str) -> str:
    return re.sub(r'\s+', '', str(value or '')).upper()


def _is_model_token_extension(shorter: str, longer: str) -> bool:
    if not shorter or not longer or len(longer) <= len(shorter):
        return False
    if not longer.startswith(shorter):
        return False
    suffix = longer[len(shorter):]
    return bool(suffix) and suffix.isalnum()


def _extract_critical_entity_spans(text: str) -> list[tuple[str, str, int, int]]:
    value = str(text or '')
    spans: list[tuple[str, str, int, int]] = []
    for pattern, label in _CAT_KEYWORD_ENTITY_PATTERNS:
        for match in pattern.finditer(value):
            raw = str(match.group(0) or '').strip()
            if not raw:
                continue
            start = match.start()
            end = match.end()
            if label == '试剂套装名':
                trimmed = re.search(r'[A-Z][A-Za-z0-9\-]*(?:[\u4e00-\u9fffA-Za-z0-9\-（）()]{0,20})?试剂套装$', raw)
                if not trimmed:
                    continue
                start += trimmed.start()
                raw = trimmed.group(0)
            spans.append((label, raw, start, end))
    spans.sort(key=lambda item: (item[2], -(item[3] - item[2])))
    return spans


def _backfill_critical_entities(source_text: str, candidate_text: str) -> str:
    source_text = str(source_text or '')
    result = str(candidate_text or '')
    if not source_text or not result or source_text == result:
        return result
    source_spans = _extract_critical_entity_spans(source_text)
    candidate_spans = _extract_critical_entity_spans(result)
    if not source_spans or not candidate_spans:
        return result

    source_by_label: dict[str, list[str]] = {}
    for label, raw, _, _ in source_spans:
        source_by_label.setdefault(label, []).append(raw)

    candidate_by_label: dict[str, list[tuple[str, int, int]]] = {}
    for label, raw, start, end in candidate_spans:
        candidate_by_label.setdefault(label, []).append((raw, start, end))

    labels = sorted(
        candidate_by_label.keys(),
        key=lambda label: -max((end - start) for _, start, end in candidate_by_label[label]),
    )
    used_source: dict[str, set[int]] = {}
    occupied: list[tuple[int, int]] = []
    planned: list[tuple[int, int, str]] = []
    for label in labels:
        if label not in _CAT_ENTITY_BACKFILL_LABELS:
            continue
        source_vals = source_by_label.get(label) or []
        if not source_vals:
            continue
        source_norms = [_normalize_entity_value(item) for item in source_vals]
        for raw, start, end in candidate_by_label[label]:
            cand_norm = _normalize_entity_value(raw)
            if cand_norm in source_norms:
                used_source.setdefault(label, set()).add(source_norms.index(cand_norm))
                continue
            replacement = None
            for index, source_raw in enumerate(source_vals):
                if index in used_source.get(label, set()):
                    continue
                source_norm = source_norms[index]
                if label in _CAT_MODEL_BACKFILL_LABELS:
                    if _is_model_token_extension(source_norm, cand_norm):
                        used_source.setdefault(label, set()).add(index)
                        replacement = None
                        break
                    if _is_model_token_extension(cand_norm, source_norm):
                        replacement = source_raw
                        used_source.setdefault(label, set()).add(index)
                        break
                replacement = source_raw
                used_source.setdefault(label, set()).add(index)
                break
            if not replacement:
                continue
            if any(not (end <= occupied_start or start >= occupied_end) for occupied_start, occupied_end in occupied):
                continue
            planned.append((start, end, replacement))
            occupied.append((start, end))

    planned.sort(key=lambda item: item[0], reverse=True)
    for start, end, replacement in planned:
        result = f'{result[:start]}{replacement}{result[end:]}'
    return result


def _normalize_cat_artifact_text(text: str) -> str:
    value = str(text or '').strip()
    _, body = _split_step_prefix(value)
    value = (body or value).lstrip('.。').strip()
    value = value.replace('的', '')
    value = re.sub(r'界面上\s+按钮', '界面上按钮', value)
    return re.sub(r'\s+', ' ', value).strip()


def _is_trivial_cat_artifact_edit(source_text: str, candidate_text: str) -> bool:
    source_norm = _normalize_cat_artifact_text(source_text)
    candidate_norm = _normalize_cat_artifact_text(candidate_text)
    if not source_norm or not candidate_norm:
        return False
    return source_norm == candidate_norm


def _has_missing_icon_button_name(text: str) -> bool:
    """True when extraction dropped the icon button name: 界面上[的]按钮."""
    return bool(re.search(r'界面上\s*的?\s*按钮', str(text or '')))


def _should_skip_cat_artifact_sentence(text: str) -> bool:
    return _has_missing_icon_button_name(text)


def _is_dropped_cat_artifact_revision(original: str, revised: str) -> bool:
    original_text = str(original or '').strip()
    revised_text = str(revised or '').strip()
    if not revised_text:
        return False
    if _is_trivial_cat_artifact_edit(original_text, revised_text):
        return True
    if _has_missing_icon_button_name(original_text) and _has_missing_icon_button_name(revised_text):
        return True
    return False


def _filter_cat_artifact_diagnose_pool(sentence_items: Optional[list]) -> list:
    kept = []
    for item in sentence_items or []:
        if not isinstance(item, dict):
            continue
        text = str(
            item.get('source_sentence_text')
            or item.get('text')
            or item.get('original_text')
            or ''
        )
        if _should_skip_cat_artifact_sentence(text):
            continue
        kept.append(item)
    return kept


def _filter_cat_artifact_diagnoses(diagnoses: Optional[list], sentence_items: Optional[list] = None) -> list:
    sentence_by_index = {
        item.get('sentence_index'): item
        for item in sentence_items or []
        if isinstance(item, dict)
    }
    kept = []
    for diag in diagnoses or []:
        if not isinstance(diag, dict):
            continue
        sent = sentence_by_index.get(diag.get('sentence_index')) or {}
        original = str(
            sent.get('source_sentence_text')
            or sent.get('text')
            or sent.get('original_text')
            or diag.get('quote')
            or ''
        )
        if _should_skip_cat_artifact_sentence(original):
            continue
        revised = str(diag.get('revised') or '')
        quote = str(diag.get('quote') or original)
        if revised and (
            _is_dropped_cat_artifact_revision(original, revised)
            or _is_dropped_cat_artifact_revision(quote, revised)
        ):
            continue
        kept.append(diag)
    return kept


def _filter_cat_artifact_diagnose_items(diagnose_items: Optional[list]) -> list:
    kept = []
    for item in diagnose_items or []:
        if not isinstance(item, dict):
            continue
        original = str(item.get('original_text') or item.get('source_sentence_text') or '')
        if _should_skip_cat_artifact_sentence(original):
            continue
        candidates = []
        for cand in item.get('candidates') or []:
            if not isinstance(cand, dict):
                continue
            revised = str(cand.get('template_text') or cand.get('text') or '')
            if revised and _is_dropped_cat_artifact_revision(original, revised):
                continue
            candidates.append(cand)
        if not candidates:
            continue
        next_item = dict(item)
        next_item['candidates'] = candidates
        kept.append(next_item)
    return kept


def _is_concatenated_context_merge(original: str, template: str, merged: str) -> bool:
    original_text = str(original or '').strip()
    template_text = str(template or '').strip()
    merged_text = str(merged or '').strip()
    if not original_text or not template_text or not merged_text:
        return False
    if len(merged_text) <= max(len(original_text), len(template_text)) + 6:
        return False
    matcher = SequenceMatcher(None, original_text, template_text, autojunk=False)
    original_only = []
    template_only = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag in ('delete', 'replace') and (i2 - i1) >= 8:
            original_only.append(original_text[i1:i2])
        if tag in ('insert', 'replace') and (j2 - j1) >= 8:
            template_only.append(template_text[j1:j2])
    original_kept = any(chunk in merged_text for chunk in original_only)
    template_kept = any(chunk in merged_text for chunk in template_only)
    return (
        original_kept
        and template_kept
        and len(merged_text) > len(original_text) + 8
        and len(merged_text) > len(template_text) + 8
    )


def _finalize_composed_cat_candidate(original: str, candidate: str, composed: str) -> str:
    source_text = str(original or '').strip()
    candidate_text = str(candidate or '').strip()
    result = _backfill_critical_entities(source_text, str(composed or '').strip())
    source_core = _strip_cat_match_prefix(source_text)
    candidate_core = _strip_cat_match_prefix(candidate_text)
    result_core = _strip_cat_match_prefix(result)
    if _is_concatenated_context_merge(source_core, candidate_core, result_core):
        result = _reapply_sentence_prefix(source_text, candidate_text)
        result = _backfill_critical_entities(source_text, result)
    result = _append_uncovered_trailing_clauses(source_text, result)
    result = _restore_source_quantity_forms(source_text, result)
    return result


def _extract_key_term_anchor_groups(text: str) -> dict[str, set[str]]:
    raw = str(text or '')
    grouped: dict[str, set[str]] = {}
    if not raw:
        return grouped
    for group, rules in _CAT_KEY_TERM_ANCHOR_GROUPS.items():
        values = set()
        for canonical, pattern in rules:
            if pattern.search(raw):
                values.add(canonical)
        if values:
            grouped[group] = values
    return grouped


def _key_term_anchor_consistent(sentence: str, template: str) -> tuple[bool, str]:
    sentence_groups = _extract_key_term_anchor_groups(sentence)
    template_groups = _extract_key_term_anchor_groups(template)
    if not sentence_groups or not template_groups:
        action_object_ok, action_object_reason = _action_object_pairs_compatible(sentence, template)
        if not action_object_ok:
            return False, action_object_reason
        return True, ''

    for group, sentence_values in sentence_groups.items():
        template_values = template_groups.get(group)
        if not template_values:
            continue
        overlap = _marker_overlap_score(sentence_values, template_values)
        sentence_only = sentence_values - template_values
        template_only = template_values - sentence_values

        if group in {'temperature', 'molecule', 'operation_mode'} and sentence_only and template_only:
            if group == 'temperature' and re.search(r'-\s*\d+\s*[℃°cC]\s*[~～\-至]\s*-?\s*\d+\s*[℃°cC]', sentence):
                continue
            return False, f'{group}冲突'

        if group == 'condition_action':
            for conflict_label, conflict_markers in _CAT_CONDITION_ACTION_CONFLICT_GROUPS.items():
                source_conflicts = sentence_values & conflict_markers
                target_conflicts = template_values & conflict_markers
                if source_conflicts and target_conflicts and not (source_conflicts & target_conflicts):
                    return False, f'{group}:{conflict_label}冲突'

        if group in {'container', 'condition_action'} and overlap == 0.0 and sentence_only and template_only:
            return False, f'{group}冲突'

        if group == 'workflow' and overlap < 0.5 and sentence_only and template_only:
            return False, 'workflow冲突'

        if group == 'organism':
            if overlap == 0.0 and sentence_only and template_only:
                return False, 'organism冲突'
            if len(sentence_values) >= 2 and len(template_values) >= 2 and overlap < 0.75 and sentence_only and template_only:
                return False, 'organism冲突'

    action_object_ok, action_object_reason = _action_object_pairs_compatible(sentence, template)
    if not action_object_ok:
        return False, action_object_reason
    return True, ''


def _action_object_pairs_compatible(sentence: str, template: str) -> tuple[bool, str]:
    sentence_pairs = _extract_sentence_intent(sentence).get('pairs', []) or []
    template_pairs = _extract_sentence_intent(template).get('pairs', []) or []
    if not sentence_pairs or not template_pairs:
        return True, ''

    sentence_by_verb: dict[str, set[str]] = defaultdict(set)
    template_by_verb: dict[str, set[str]] = defaultdict(set)
    for pair in sentence_pairs:
        verb = str(pair.get('verb', '') or '').strip()
        obj = _normalize_structure_object(pair.get('object', '') or '')
        if verb and obj:
            sentence_by_verb[verb].add(obj)
    for pair in template_pairs:
        verb = str(pair.get('verb', '') or '').strip()
        obj = _normalize_structure_object(pair.get('object', '') or '')
        if verb and obj:
            template_by_verb[verb].add(obj)

    for verb in set(sentence_by_verb) & set(template_by_verb):
        sentence_objects = sentence_by_verb.get(verb, set())
        template_objects = template_by_verb.get(verb, set())
        if not sentence_objects or not template_objects:
            continue
        compatible = False
        for sentence_object in sentence_objects:
            for template_object in template_objects:
                if _same_action_object_compatible(sentence_object, template_object):
                    compatible = True
                    break
            if compatible:
                break
        if compatible:
            continue
        if len(sentence_objects) == 1 and len(template_objects) == 1:
            return False, f'action_object:{verb}'
    return True, ''


def _normalize_action_object_aliases(text: str) -> str:
    value = _normalize_structure_object(text)
    if not value:
        return ''
    replacements = {
        '样本制备系统': '样本制备卡系统',
        '机器推板': '仪器载台',
        '机器': '仪器',
        '推板': '载台',
        '平置': '水平放置',
        '对应': '一致',
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def _extract_action_object_identity_markers(text: str) -> set[str]:
    raw = str(text or '')
    markers = set()
    for match in re.findall(r'套件\s*([A-Z])', raw):
        markers.add(f'kit:{match.upper()}')
    for match in re.findall(r'\b([A-Z]\d+)\b', raw, flags=re.IGNORECASE):
        markers.add(re.sub(r'\s+', '', match).upper())
    return markers


def _same_action_object_compatible(source_object: str, template_object: str) -> bool:
    source_value = _normalize_structure_object(source_object)
    template_value = _normalize_structure_object(template_object)
    if not source_value or not template_value:
        return bool(source_value == template_value)
    if source_value == template_value:
        return True
    if source_value in template_value or template_value in source_value:
        return True

    source_identity = _extract_action_object_identity_markers(source_value)
    template_identity = _extract_action_object_identity_markers(template_value)
    if source_identity and template_identity and not (source_identity & template_identity):
        return False

    source_alias = _normalize_action_object_aliases(source_value)
    template_alias = _normalize_action_object_aliases(template_value)
    if source_alias == template_alias:
        return True
    if source_alias in template_alias or template_alias in source_alias:
        return True

    clause_score = _clause_similarity(source_alias, template_alias)
    semantic_score = _compare_semantic_keywords(source_alias, template_alias)
    return max(clause_score, semantic_score) >= 0.42


def _candidate_effective_semantic_score(candidate: dict) -> float:
    semantic_score = candidate.get('semantic_score') if isinstance(candidate, dict) else None
    if semantic_score is None:
        semantic_score = candidate.get('string_score', 0.0) if isinstance(candidate, dict) else 0.0
    return round(max(0.0, min(1.0, float(semantic_score or 0.0))), 4)


def _candidate_filtered_semantic_score(original_text: str, candidate: dict, ai_semantic_active: bool = True) -> float:
    score = _candidate_effective_semantic_score(candidate)
    reason = (candidate or {}).get('ai_reason', '') if isinstance(candidate, dict) else ''
    if ai_semantic_active and _has_negative_reason(reason):
        score = max(0.0, score - 0.22)
    tech_overlap = _tech_term_overlap(original_text, str((candidate or {}).get('template_text', '') or ''))
    tech_term_count = len(_extract_tech_terms(original_text))
    if tech_term_count >= 2:
        if tech_overlap < 0.35:
            score = max(0.0, score - (0.18 if ai_semantic_active else 0.12))
        elif tech_overlap < 0.60:
            score = max(0.0, score - (0.08 if ai_semantic_active else 0.05))
    return round(score, 4)


def _apply_cat_terms(text: str, term_dict: Optional[dict]) -> str:
    value = str(text or '')
    if not value or not term_dict:
        return value
    result = value
    for old_term, new_term in term_dict.items():
        old_text = str(old_term or '').strip()
        new_text = str(new_term or '').strip()
        if old_text and new_text and old_text in result:
            result = result.replace(old_text, new_text)
    return result


def _apply_cat_surface_rules(text: str, term_dict: Optional[dict] = None, typo_dict: Optional[dict] = None) -> str:
    value = str(text or '').strip()
    if not value:
        return value
    value = _apply_cat_terms(value, term_dict)
    value = _apply_cat_terms(value, typo_dict)
    value = value.replace('按要求', '按指示')
    value = re.sub(rf'(?<=\d)\s*(?=({_NUMBER_SPACE_UNITS})(?![A-Za-z]))', ' ', value)
    value = re.sub(r'\s+', ' ', value).strip()
    if _should_add_contextual_terminal_punctuation(value):
        value = value.rstrip('，,;；：:') + '。'
    return _protect_model_numbers(value).strip()


def _strip_cat_match_prefix(text: str) -> str:
    value = str(text or '').strip()
    if not value:
        return value
    _, value = _split_list_marker_prefix(value)
    _, value = _split_notice_prefix(value)
    _, body = _split_step_prefix(value)
    return (body or value).strip()


def _prepare_cat_match_text(text: str, term_dict: Optional[dict] = None, typo_dict: Optional[dict] = None) -> str:
    return _apply_cat_surface_rules(_strip_cat_match_prefix(text), term_dict, typo_dict)


def _normalize_cat_duplicate_text(text: str) -> str:
    value = _normalize_visible_compare_text(text)
    value = value.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'")
    value = _normalize_terminal_sentence_punctuation(value)
    value = re.sub(r'\s+', ' ', value).strip()
    return value.lstrip('.。').strip()


def _cat_exact_duplicate(source_text: str, candidate_text: str) -> bool:
    return _normalize_cat_duplicate_text(source_text) == _normalize_cat_duplicate_text(candidate_text)


def _extract_ui_destination_markers(text: str) -> set[str]:
    return {
        marker for marker in _extract_ui_markers(text)
        if marker.endswith(('界面', '页面', '窗口'))
    }


def _has_conflicting_ui_destination(sentence: str, template: str) -> bool:
    source_destinations = _extract_ui_destination_markers(sentence)
    template_destinations = _extract_ui_destination_markers(template)
    source_actions = set(_extract_sentence_intent(sentence).get('actions', []))
    template_actions = set(_extract_sentence_intent(template).get('actions', []))
    return _has_conflicting_ui_destination_with_markers(
        source_destinations,
        template_destinations,
        source_actions,
        template_actions,
    )


def _has_conflicting_ui_destination_with_markers(
    source_destinations: set[str],
    template_destinations: set[str],
    source_actions: set[str],
    template_actions: set[str],
) -> bool:
    if not template_destinations:
        if source_destinations and (source_actions & {'进入', '打开'}) and not (template_actions & {'进入', '打开'}):
            return True
        return False
    if not source_destinations:
        return True
    return not bool(source_destinations & template_destinations)


def _extract_ui_operation_markers(text: str) -> set[str]:
    raw = str(text or '')
    markers = set()
    if re.search(r'\brun\b', raw, flags=re.IGNORECASE):
        markers.add('run')
    for marker, pattern in [
        ('运行', r'运行'),
        ('开始实验', r'开始实验'),
        ('开始建库', r'开始建库'),
        ('确定', r'确定'),
        ('登录', r'登录'),
    ]:
        if re.search(pattern, raw):
            markers.add(marker)
    return markers


def _has_missing_ui_operation_markers(sentence: str, template: str) -> bool:
    source_markers = _extract_ui_operation_markers(sentence)
    template_markers = _extract_ui_operation_markers(template)
    return _has_missing_ui_operation_markers_with_markers(source_markers, template_markers)


def _has_missing_ui_operation_markers_with_markers(source_markers: set[str], template_markers: set[str]) -> bool:
    if not source_markers:
        return False
    critical_source_markers = source_markers & {'run', '运行', '开始实验', '开始建库', '登录'}
    if not critical_source_markers:
        return False
    return not bool(critical_source_markers & template_markers)


def _has_injected_ui_bracket_labels(source: str, candidate: str) -> bool:
    labels = re.findall(r'【([^】]{1,20})】', str(candidate or ''))
    if not labels:
        return False
    source_text = str(source or '')
    for label in labels:
        if label and label in source_text:
            continue
        return True
    return False


def _clause_covered_in_text(clause: str, text: str, threshold: float = 0.45) -> bool:
    clause_text = str(clause or '').strip()
    haystack = str(text or '').strip()
    if not clause_text or not haystack:
        return False
    compact_clause = re.sub(r'\s+', '', clause_text)
    compact_text = re.sub(r'\s+', '', haystack)
    if compact_clause and compact_clause in compact_text:
        return True
    if _clause_similarity(clause_text, haystack) >= threshold:
        return True
    return any(_clause_similarity(clause_text, other) >= threshold for other in _split_sentence_clauses(haystack))


def _uncovered_trailing_source_clauses(source_core: str, candidate_core: str) -> list[str]:
    source_clauses = _split_sentence_clauses(source_core)
    if len(source_clauses) < 2:
        return []
    trailing = []
    seen_covered = False
    for clause in reversed(source_clauses):
        if _clause_covered_in_text(clause, candidate_core):
            seen_covered = True
            break
        if len(_normalize_sentence_for_match(clause)) < 4:
            continue
        trailing.append(clause)
    if not seen_covered:
        return []
    trailing.reverse()
    return trailing


def _append_uncovered_trailing_clauses(source_text: str, result: str) -> str:
    source_core = _strip_cat_match_prefix(source_text)
    result_text = str(result or '').strip()
    trailing = _uncovered_trailing_source_clauses(source_core, result_text)
    if not trailing or not result_text:
        return result_text
    terminal = ''
    terminal_match = re.search(r'([。.!！？?]+)\s*$', result_text)
    if terminal_match:
        terminal = terminal_match.group(1)[-1]
        result_text = result_text[:terminal_match.start()].rstrip('，,;；')
    extra = '，'.join(item.rstrip('。.!！？?，,;；') for item in trailing if item.strip())
    if not extra:
        return f'{result_text}{terminal}'
    return f'{result_text}，{extra}{terminal}'


def _restore_source_quantity_forms(source_text: str, candidate_text: str) -> str:
    source = str(source_text or '')
    result = str(candidate_text or '')
    if not source or not result:
        return result
    classifier = r'[个张份条项次轮孔管板卡]'
    for match in re.finditer(rf'([{"".join(_CN_CARDINAL_DIGITS.keys())}]+)({classifier})', source):
        digits, unit = match.group(1), match.group(2)
        source_form = match.group(0)
        if len(digits) == 1:
            arabic = _CN_CARDINAL_DIGITS.get(digits)
        elif digits == '十':
            arabic = 10
        else:
            continue
        if arabic is None:
            continue
        arabic_form = f'{arabic}{unit}'
        if arabic_form in result:
            result = result.replace(arabic_form, source_form, 1)
    for match in re.finditer(rf'([{"".join(_CN_CARDINAL_DIGITS.keys())}])({classifier})([\u4e00-\u9fff]{{2,6}})', source):
        source_form = match.group(0)
        if source_form in result:
            continue
        unit, noun = match.group(2), match.group(3)
        pattern = rf'每{re.escape(unit)}(?:样本)?{re.escape(noun)}'
        result = re.sub(pattern, source_form, result, count=1)
    return result


def _trim_redundant_cat_suffix(text: str) -> str:
    value = str(text or '').strip()
    if not value:
        return value
    terminal = ''
    terminal_match = re.search(r'([。.!！？?；;]+)$', value)
    if terminal_match:
        terminal = terminal_match.group(1)
        value = value[:-len(terminal)].rstrip()
    match = re.search(r'(和|与|及|并)([^，。；！？,.;:：]{2,12})$', value)
    if not match:
        return f'{value}{terminal}' if value else terminal
    repeated_fragment = match.group(2).strip()
    if not repeated_fragment:
        return f'{value}{terminal}' if value else terminal
    body = value[:-len(match.group(0))]
    # Only drop duplicated conjuncts (A与A). Shared substrings in earlier modifiers stay.
    if body.endswith(repeated_fragment):
        value = body.rstrip('，,；;：: ')
    if value == str(text or '').strip().rstrip(terminal):
        for fragment_length in range(min(8, len(value) // 2), 1, -1):
            fragment = value[-fragment_length:]
            prefix = value[:-fragment_length]
            if not fragment.strip() or not prefix.endswith(('和', '与', '及', '并')):
                continue
            head = prefix[:-1].rstrip('，,；;：: ')
            if not head.endswith(fragment):
                continue
            value = head
            break
    return f'{value}{terminal}' if value else terminal


def _normalize_cat_length_text(text: str) -> str:
    return re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', str(text or ''))


def _is_fragment_candidate(source: str, candidate: str, min_ratio: float = 0.65, max_ratio: float = 1.55) -> bool:
    source_norm = _normalize_cat_length_text(source)
    candidate_norm = _normalize_cat_length_text(candidate)
    if not source_norm or not candidate_norm:
        return True
    ratio = len(candidate_norm) / max(len(source_norm), 1)
    if len(source_norm) <= 18:
        max_ratio = 2.35
    elif len(source_norm) <= 26:
        max_ratio = max(max_ratio, 1.9)
    if ratio < min_ratio or ratio > max_ratio:
        return True
    if ratio < 0.78 and (candidate_norm in source_norm or source_norm in candidate_norm):
        return True
    return False


def _is_valid_cat_template(text: str) -> bool:
    value = str(text or '').strip()
    if re.search(r'(?:孔位\s*)?[A-Za-z]\d\s*对应\s*(?:sample|样本)\s*[A-Za-z]?\d+', value, re.IGNORECASE):
        return True
    if len(_normalize_cat_length_text(value)) < 15:
        return False
    if re.search(r'[。！？；!?;]$', value):
        return True
    if ('：' in value or ':' in value) and re.search(r'@[A-Za-z0-9.-]+', value):
        return True
    return False


def _merge_split_templates(templates: list[str]) -> list[str]:
    merged = []
    index = 0
    while index < len(templates or []):
        current = str(templates[index] or '').strip()
        if not current:
            index += 1
            continue

        if index + 1 < len(templates):
            next_text = str(templates[index + 1] or '').strip()
            if next_text and not re.search(r'[。！？；!?;]$', current):
                joiner = '' if re.search(r'[，,；;：:]$', current) else '，'
                combined = f'{current}{joiner}{next_text}'.strip()
                if _is_valid_cat_template(combined):
                    current_norm = _normalize_cat_length_text(current)
                    next_norm = _normalize_cat_length_text(next_text)
                    combined_norm = _normalize_cat_length_text(combined)
                    if len(combined_norm) > max(len(current_norm), len(next_norm)) + 5:
                        merged.append(combined)
                        index += 2
                        continue

        merged.append(current)
        index += 1
    return merged


def _email_domains(text: str) -> set[str]:
    cleaned = str(text or '').replace('\\.', '.')
    return {item.lower() for item in re.findall(r'@([A-Za-z0-9.-]+\.[A-Za-z]{2,})', cleaned, flags=re.IGNORECASE)}


def _should_skip_context_merge(source_core: str, candidate_core: str, overall_similarity: float) -> bool:
    if overall_similarity >= 0.75:
        return not _uncovered_trailing_source_clauses(source_core, candidate_core)
    source_domains = _email_domains(source_core)
    candidate_domains = _email_domains(candidate_core)
    return bool(source_domains and candidate_domains and source_domains != candidate_domains)


def _trim_duplicated_leading_clause(text: str) -> str:
    value = str(text or '').strip()
    split_match = re.search(r'[，,；;、]', value)
    if not split_match:
        return value
    first = value[:split_match.start()].strip()
    rest = value[split_match.end():].strip()
    if len(first) < 12 or len(rest) < 12:
        return value
    max_stem = min(len(first), 40)
    for length in range(max_stem, 11, -1):
        if rest.startswith(first[:length]):
            return rest
    return value



def _restore_stripped_sentence_terminal(original: str, core: str) -> str:
    core_text = str(core or '').strip()
    raw = str(original or '').strip()
    if not core_text or not raw:
        return core_text
    if re.search(r'[。.!！？?：:;；]\s*$', core_text):
        return core_text
    match = re.search(r'([。.!！？?：:;；]+)\s*$', raw)
    if not match:
        return core_text
    return f'{core_text}{match.group(1)[-1]}'


def _has_conflicting_list_intro(source: str, template: str) -> bool:
    source_text = str(source or '').strip()
    template_text = str(template or '').strip()
    if not source_text or not template_text:
        return False
    source_is_list_intro = bool(re.search(r'[：:]\s*$', source_text))
    template_has_colon = ('：' in template_text) or (':' in template_text)
    template_is_complete = bool(re.search(r'[。.!！？?]\s*$', template_text))
    return source_is_list_intro and template_is_complete and not template_has_colon


def _has_conflicting_list_step_template(source: str, template: str) -> bool:
    source_text = str(source or '').strip()
    template_text = str(template or '').strip()
    if not source_text or not template_text:
        return False
    if not re.search(r'[。.!！？?]\s*$', source_text):
        return False
    if not re.search(r'[；;]\s*$', template_text):
        return False
    if _sentence_similarity(source_text, template_text) >= 0.75:
        return False
    extra_numbers = _extract_numeric_markers(template_text) - _extract_numeric_markers(source_text)
    if extra_numbers:
        return True
    return len(_split_sentence_clauses(source_text)) >= 2 and len(_split_sentence_clauses(template_text)) == 1


def _source_clause_repeats_in_candidate(clause: str, candidate_core: str) -> bool:
    clause_text = str(clause or '').strip()
    candidate_text = str(candidate_core or '').strip()
    if not clause_text or not candidate_text:
        return False
    if _clause_covered_in_text(clause_text, candidate_text, threshold=0.42):
        return True
    actions = [item for item in _extract_sentence_intent(clause_text).get('actions', []) if item and len(item) >= 2]
    return any(action in candidate_text for action in actions)


def _is_duplicated_clause_splice(source: str, composed: str, template: str) -> bool:
    source_clauses = _split_sentence_clauses(source)
    if len(source_clauses) < 2:
        return False
    prefix = source_clauses[0]
    composed_text = str(composed or '').strip()
    template_text = str(template or '').strip()
    if not prefix or not composed_text.startswith(prefix):
        return False
    if prefix in template_text:
        return False
    template_core = re.sub(r'[；;。.!！？?\s]+$', '', template_text)
    return bool(template_core) and template_core in composed_text


def _finish_cat_composed_text(
    source_text: str,
    merged_text: str,
    trailing_figure_ref: str = '',
    source_terminal: str = '',
) -> str:
    merged_text = _trim_duplicated_leading_clause(str(merged_text or '').strip())
    merged_text = _trim_redundant_cat_suffix(merged_text)
    if trailing_figure_ref and trailing_figure_ref not in merged_text:
        merged_text = f'{merged_text}{trailing_figure_ref}'
    if source_terminal:
        merged_text = re.sub(r'[；;]+\s*$', '', merged_text).rstrip()
        if not re.search(r'[。.!！？?]\s*$', merged_text):
            merged_text = f'{merged_text}{source_terminal}'
        merged_text = re.sub(r'[；;]+(?=[。.!！？?]\s*$)', '', merged_text)
    return _reapply_sentence_prefix(source_text, merged_text)


def _compose_cat_candidate_text(original: str, candidate: str) -> str:
    source_text = str(original or '').strip()
    candidate_text = str(candidate or '').strip()
    if not source_text or not candidate_text:
        return candidate_text
    return _finalize_composed_cat_candidate(
        source_text,
        candidate_text,
        _compose_cat_candidate_body(source_text, candidate_text),
    )


def _compose_cat_candidate_body(original: str, candidate: str) -> str:
    source_text = str(original or '').strip()
    candidate_text = _unescape_markdown_inline(str(candidate or '').strip())
    if not source_text or not candidate_text:
        return candidate_text

    anchored_candidate = _merge_generic_subject_with_source(source_text, candidate_text)
    if anchored_candidate:
        candidate_text = anchored_candidate

    source_terminal_match = re.search(r'([。.!！？?]+)\s*$', source_text)
    source_terminal = source_terminal_match.group(1)[-1] if source_terminal_match else ''

    source_core = _strip_cat_match_prefix(source_text)
    candidate_core = _strip_cat_match_prefix(candidate_text)
    source_core_raw = source_core
    candidate_core_raw = candidate_core
    source_core, source_figure_ref = _extract_trailing_figure_ref(source_core)
    candidate_core, candidate_figure_ref = _extract_trailing_figure_ref(candidate_core)
    if not source_figure_ref:
        source_core = _restore_stripped_sentence_terminal(source_core_raw, source_core)
    if not candidate_figure_ref:
        candidate_core = _restore_stripped_sentence_terminal(candidate_core_raw, candidate_core)
    if not source_core or not candidate_core:
        return _reapply_sentence_prefix(source_text, candidate_text)

    source_clauses = _split_sentence_clauses(source_core)
    candidate_clauses = _split_sentence_clauses(candidate_core)
    overall_similarity = _sentence_similarity(source_core, candidate_core)
    best_clause_similarity = max((_clause_similarity(clause, candidate_core) for clause in source_clauses), default=0.0)
    action_overlap = len(
        set(_extract_sentence_intent(source_core).get('actions', [])) &
        set(_extract_sentence_intent(candidate_core).get('actions', []))
    )

    trailing_figure_ref = candidate_figure_ref or source_figure_ref

    candidate_core = candidate_core.lstrip('.。')
    if _should_skip_context_merge(source_core, candidate_core, overall_similarity):
        return _finish_cat_composed_text(source_text, candidate_core, trailing_figure_ref, source_terminal)

    if _has_conflicting_list_intro(source_core, candidate_core):
        return _finish_cat_composed_text(source_text, candidate_core, trailing_figure_ref, source_terminal)

    if len(source_clauses) >= 3 and len(candidate_clauses) == 1:
        trailing_clause_matches = sum(1 for clause in source_clauses[1:] if _clause_similarity(clause, candidate_core) >= 0.5)
        leading_clause_similarity = _clause_similarity(source_clauses[0], candidate_core) if source_clauses else 0.0
        if trailing_clause_matches >= 2 and leading_clause_similarity < 0.2 and overall_similarity < 0.55:
            return _finish_cat_composed_text(source_text, candidate_core, trailing_figure_ref, source_terminal)

    if overall_similarity < 0.3 and best_clause_similarity < 0.38:
        return _reapply_sentence_prefix(source_text, candidate_text)

    if len(source_clauses) <= 1 or len(candidate_clauses) > 1:
        merged_text = replace_with_context(source_core, candidate_core)
        return _finish_cat_composed_text(source_text, merged_text, trailing_figure_ref, source_terminal)

    clause_parts = re.split(r'([，。；！？,;!?]+)', source_core)
    best_index = None
    best_score = 0.0
    for index, part in enumerate(clause_parts):
        if index % 2 == 1:
            continue
        clause = part.strip()
        if not clause:
            continue
        score = _clause_similarity(clause, candidate_core)
        if score > best_score:
            best_score = score
            best_index = index

    if best_index is None or best_score < 0.45:
        if len(source_clauses) > 1 and len(candidate_clauses) == 1 and overall_similarity < 0.48 and action_overlap <= 0:
            return _finish_cat_composed_text(source_text, candidate_core, trailing_figure_ref, source_terminal)
        merged_text = replace_with_context(source_core, candidate_core)
        return _finish_cat_composed_text(source_text, merged_text, trailing_figure_ref, source_terminal)

    kept_repeats_candidate = False
    for index, part in enumerate(clause_parts):
        if index % 2 == 1 or index == best_index:
            continue
        if _source_clause_repeats_in_candidate(part.strip(), candidate_core):
            kept_repeats_candidate = True
            break
    if kept_repeats_candidate:
        merged_text = replace_with_context(source_core, candidate_core)
        return _finish_cat_composed_text(source_text, merged_text, trailing_figure_ref, source_terminal)

    rebuilt_parts = []
    for index, part in enumerate(clause_parts):
        if index != best_index:
            rebuilt_parts.append(part)
            continue
        clause = part.strip()
        next_separator = clause_parts[index + 1] if index + 1 < len(clause_parts) else ''
        rebuilt_parts.append(part.replace(clause, _replace_clause_directly(clause, candidate_core, next_separator), 1))

    rebuilt_text = ''.join(rebuilt_parts)
    return _finish_cat_composed_text(source_text, rebuilt_text, trailing_figure_ref, source_terminal)


async def _collect_text_polish_cat_items(
    original_text: str,
    db: Session,
    sentence_guide: str,
    terminology_md: str,
    style_guide_id: Optional[int] = None,
    sentence_file_id: Optional[int] = None,
    min_match_threshold: float = 0.34,
    fuzzy_lower_bound: float = 0.70,
) -> tuple[str, list[dict]]:
    selected_sentence_guide = _load_selected_sentence_guide_with_feedback(db, sentence_file_id)
    candidate_recall_guide = _candidate_recall_guide_text(sentence_guide or '')
    if sentence_file_id and selected_sentence_guide:
        candidate_recall_guide = selected_sentence_guide

    guide_entries = _preferred_entries_from_guide(candidate_recall_guide)
    guide_templates = []
    seen_template_texts = set()
    for entry in guide_entries or []:
        template_text = _template_entry_text(entry)
        if not template_text or not _is_valid_cat_template(template_text):
            continue
        dedupe_candidate = re.sub(r'\s+', '', re.sub(r'[，。！？!?；;：:、,\.\s]+$', '', template_text))
        if not dedupe_candidate or dedupe_candidate in seen_template_texts:
            continue
        seen_template_texts.add(dedupe_candidate)
        template_id = str(entry.get('id', '')) if isinstance(entry, dict) else ''
        guide_templates.append({'text': template_text, 'id': template_id})

    if not guide_templates:
        return []

    terminology = terminology_md or _load_scoped_terminology_source(db, None, None)
    resolved_terms = _resolve_terminology(db, terminology, original_text) if terminology else {}
    original_paragraphs = original_text.split('\n')
    sentence_items = _split_cat_sentences(original_paragraphs, source_paragraphs=original_paragraphs)
    if not sentence_items:
        return []

    cat_items = []
    for sentence_item in sentence_items:
        line_stripped = sentence_item['text'].strip()
        source_line_stripped = str(sentence_item.get('source_sentence_text') or line_stripped).strip()
        local_entries = _preferred_entries_for_sentence(line_stripped, candidate_recall_guide)
        local_templates = []
        local_seen = set()
        for entry in local_entries or []:
            template_text = _template_entry_text(entry)
            if not template_text or not _is_valid_cat_template(template_text):
                continue
            dedupe = re.sub(r'\s+', '', re.sub(r'[，。！？!?；;：:、,\.\s]+$', '', template_text))
            if not dedupe or dedupe in local_seen:
                continue
            local_seen.add(dedupe)
            template_id = str(entry.get('id', '')) if isinstance(entry, dict) else ''
            local_templates.append({'text': template_text, 'id': template_id})
            if len(local_templates) >= 120:
                break
        if not local_templates:
            local_templates = guide_templates[:120]

        simple_match_debug = {}
        candidates = _simple_match(
            line_stripped,
            local_templates,
            min_threshold=min_match_threshold,
            fuzzy_lower=fuzzy_lower_bound,
            term_dict=resolved_terms,
            source_sentence=source_line_stripped,
            debug_stats=simple_match_debug,
        )
        if not candidates and local_templates is not guide_templates:
            candidates = _simple_match(
                line_stripped,
                guide_templates,
                min_threshold=min_match_threshold,
                fuzzy_lower=fuzzy_lower_bound,
                term_dict=resolved_terms,
                source_sentence=source_line_stripped,
                debug_stats=simple_match_debug,
            )
        cat_items.append({
            'original_text': source_line_stripped,
            'candidates': candidates,
            'paragraph_index': sentence_item['source_paragraph_index'],
            'sentence_index': sentence_item['sentence_index'],
        })

    ai_score_result = await _batch_ai_semantic_score(cat_items, reason_max_chars=15)
    ai_semantic_active = (ai_score_result or {}).get('status') == 'completed'

    if not ai_semantic_active:
        for item in cat_items:
            for candidate in item.get('candidates', []):
                if candidate.get('semantic_score') is None:
                    candidate['semantic_score'] = candidate.get('string_score', 0.0)
                    candidate['ai_reason'] = 'AI未启用，使用字符串匹配分'

    for item in cat_items:
        filtered = []
        for candidate in item.get('candidates', []):
            candidate['effective_semantic_score'] = _candidate_effective_semantic_score(candidate)
            if _should_keep_text_manual_candidate(item.get('original_text', ''), candidate, ai_semantic_active=ai_semantic_active):
                filtered.append(candidate)
        filtered.sort(
            key=lambda c: (
                float(c.get('filtered_semantic_score', 0.0) or 0.0),
                float(c.get('effective_semantic_score', 0.0) or 0.0),
                float(c.get('semantic_score', 0.0) or 0.0),
                float(c.get('string_score', 0.0) or 0.0),
            ),
            reverse=True,
        )
        item['candidates'] = filtered[:5]
        item['has_candidates'] = bool(item['candidates'])
    return [item for item in cat_items if item.get('has_candidates')]


def _extract_sentence_subject_predicate(text: str) -> tuple[str, str, str]:
    value = str(text or '').strip()
    if not value:
        return '', '', ''
    match = re.match(r'^(.{2,40}?)(由|为|是|分为|包括|包含)(.+)$', value)
    if not match:
        return '', '', ''
    subject = match.group(1).strip('，,；;：: ')
    anchor = match.group(2).strip()
    predicate = match.group(3).strip()
    if not subject or not predicate:
        return '', '', ''
    return subject, anchor, predicate


def _merge_generic_subject_with_source(source_text: str, candidate_text: str) -> str:
    source_subject, source_anchor, source_predicate = _extract_sentence_subject_predicate(_strip_cat_match_prefix(source_text))
    candidate_subject, candidate_anchor, candidate_predicate = _extract_sentence_subject_predicate(_strip_cat_match_prefix(candidate_text))
    if not source_subject or not candidate_subject or source_anchor != candidate_anchor:
        return ''
    if len(source_subject) <= len(candidate_subject):
        return ''

    predicate_similarity = max(
        _clause_similarity(source_predicate, candidate_predicate),
        _sentence_similarity(source_predicate, candidate_predicate),
    )
    if predicate_similarity < 0.45:
        return ''

    source_entities = _extract_critical_entities(source_subject)
    candidate_entities = _extract_critical_entities(candidate_subject)
    subject_similarity = _sentence_similarity(source_subject, candidate_subject)
    if not source_entities and subject_similarity >= 0.55:
        return ''
    if candidate_entities and source_entities == candidate_entities:
        return ''

    merged_text = f'{source_subject}{candidate_anchor}{candidate_predicate}'
    merged_text = _protect_model_numbers(merged_text)
    return _reapply_sentence_prefix(source_text, merged_text)


def _has_negative_reason(reason: str) -> bool:
    normalized_reason = str(reason or '').strip()
    if not normalized_reason:
        return False
    return any(keyword in normalized_reason for keyword in _CAT_NEGATIVE_REASON_KEYWORDS)


def _tech_term_overlap(source: str, candidate: str) -> float:
    source_terms = _extract_tech_terms(source)
    candidate_terms = _extract_tech_terms(candidate)
    if not source_terms:
        return 1.0
    return len(source_terms & candidate_terms) / max(len(source_terms), 1)


def _extract_tech_terms(text: str) -> set[str]:
    terms = set()
    for pattern in _CAT_TECH_TERM_PATTERNS:
        terms.update(match.group(0).lower() for match in pattern.finditer(str(text or '')))
    return terms


def _cat_reason_indicates_no_change(reason: str) -> bool:
    normalized_reason = str(reason or '').strip()
    if not normalized_reason:
        return False
    return any(keyword in normalized_reason for keyword in _CAT_NO_CHANGE_REASON_KEYWORDS)


def _should_keep_cat_candidate(original_text: str, candidate: dict, ai_semantic_active: bool = True) -> bool:
    template_text = str((candidate or {}).get('template_text', '') or '').strip()
    if isinstance(candidate, dict):
        preserved_tags = [tag for tag in (candidate.get('review_tags') or []) if tag == 'entity_changed']
        candidate['review_tags'] = preserved_tags
        candidate['drop_reason'] = None
    if not template_text:
        if isinstance(candidate, dict):
            candidate['drop_reason'] = 'empty_template'
        return False

    if _cat_exact_duplicate(original_text, template_text):
        if isinstance(candidate, dict):
            candidate['drop_reason'] = 'exact_duplicate'
        return False

    if _is_trivial_cat_artifact_edit(original_text, template_text):
        if isinstance(candidate, dict):
            candidate['drop_reason'] = 'trivial_artifact'
        return False

    if isinstance(candidate, dict) and candidate.get('rule_source') == 'surface_rules':
        candidate['filtered_semantic_score'] = max(
            float(candidate.get('effective_semantic_score', 0.0) or 0.0),
            float(candidate.get('semantic_score', 0.0) or 0.0),
            0.95,
        )
        candidate['needs_review'] = False
        candidate['review_tags'] = []
        return True

    if _cat_reason_indicates_no_change((candidate or {}).get('ai_reason', '')):
        if isinstance(candidate, dict):
            candidate['drop_reason'] = 'ai_no_change'
        return False

    show_threshold = _CAT_AI_SEMANTIC_SHOW_THRESHOLD if ai_semantic_active else _CAT_RULE_ONLY_SHOW_THRESHOLD
    pending_threshold = _CAT_AI_SEMANTIC_PENDING_THRESHOLD if ai_semantic_active else _CAT_RULE_ONLY_PENDING_THRESHOLD
    filtered_score = _candidate_filtered_semantic_score(original_text, candidate, ai_semantic_active=ai_semantic_active)
    rescue_penalty = None
    rescue_penalty_reasons: list[str] = []
    is_rule_rescue_band = (not ai_semantic_active) and filtered_score < _CAT_RULE_ONLY_LEGACY_PENDING_THRESHOLD and filtered_score >= pending_threshold
    if is_rule_rescue_band:
        rescue_penalty = _collect_match_penalties(original_text, template_text)
        rescue_penalty_reasons = list((rescue_penalty or {}).get('reasons', []) or [])
    if isinstance(candidate, dict):
        review_tags = list(candidate.get('review_tags') or [])
        if ai_semantic_active and _has_negative_reason(candidate.get('ai_reason', '')):
            review_tags.append('negative_ai_reason')
        if _tech_term_overlap(original_text, template_text) < 0.6:
            review_tags.append('tech_term_overlap_low')
        if is_rule_rescue_band:
            severe_penalty_reasons = {'对象不一致', '动作对象冲突', '数字/单位不一致', '主题锚点覆盖不足'}
            if (rescue_penalty or {}).get('score_penalty', 0.0) >= 0.18 or severe_penalty_reasons.intersection(rescue_penalty_reasons):
                candidate['drop_reason'] = 'rule_rescue_penalty'
                candidate['review_tags'] = review_tags + ['rule_rescue_penalty']
                candidate['filtered_semantic_score'] = filtered_score
                candidate['needs_review'] = True
                return False
            review_tags.append('rule_threshold_rescue')
        candidate['filtered_semantic_score'] = filtered_score
        candidate['needs_review'] = filtered_score < show_threshold
        if candidate['needs_review']:
            if 'needs_review' not in review_tags:
                review_tags.append('needs_review')
        candidate['review_tags'] = review_tags
        keep_score = filtered_score
        if 'entity_changed' in review_tags:
            unpenalized = float(candidate.get('unpenalized_string_score') or 0.0)
            if unpenalized:
                keep_score = max(filtered_score, unpenalized)
        if keep_score < pending_threshold:
            candidate['drop_reason'] = 'low_filtered_score'
        return keep_score >= pending_threshold
    return filtered_score >= pending_threshold


def _should_keep_text_manual_candidate(original_text: str, candidate: dict, ai_semantic_active: bool = True) -> bool:
    if _should_keep_cat_candidate(original_text, candidate, ai_semantic_active=ai_semantic_active):
        return True
    if not isinstance(candidate, dict):
        return False
    template_text = str(candidate.get('template_text', '') or '').strip()
    if not template_text:
        return False
    if candidate.get('drop_reason') in {'exact_duplicate', 'ai_no_change', 'empty_template', 'trivial_artifact'}:
        return False
    string_score = float(candidate.get('string_score', 0.0) or 0.0)
    filtered_score = float(candidate.get('filtered_semantic_score', 0.0) or 0.0)
    source_norm = _normalize_cat_length_text(original_text)
    template_norm = _normalize_cat_length_text(template_text)
    if not source_norm or not template_norm:
        return False
    if len(template_norm) < len(source_norm):
        return False
    keyword_overlap = _compare_semantic_keywords(original_text, template_text)
    return string_score >= 0.38 and filtered_score >= 0.38 and keyword_overlap >= 0.2


def _increment_debug_reason(debug_stats: Optional[dict], key: str) -> None:
    if not isinstance(debug_stats, dict):
        return
    reason_map = debug_stats.setdefault('dropped_by_reason', {})
    reason_map[key] = int(reason_map.get(key, 0) or 0) + 1


def _finalize_simple_match_debug(debug_stats: Optional[dict]) -> dict:
    if not isinstance(debug_stats, dict):
        return {}
    debug_stats['considered_templates'] = int(debug_stats.get('considered_templates', 0) or 0)
    debug_stats['matched_templates'] = int(debug_stats.get('matched_templates', 0) or 0)
    debug_stats['returned_candidates'] = int(debug_stats.get('returned_candidates', 0) or 0)
    debug_stats['surface_rule_candidates'] = int(debug_stats.get('surface_rule_candidates', 0) or 0)
    debug_stats['dropped_by_reason'] = dict(debug_stats.get('dropped_by_reason', {}) or {})
    return debug_stats


def _simple_match(
    sentence: str,
    templates: list[dict],
    min_threshold: float = 0.34,
    fuzzy_lower: float = 0.70,
    term_dict: Optional[dict] = None,
    typo_dict: Optional[dict] = None,
    source_sentence: Optional[str] = None,
    debug_stats: Optional[dict] = None,
) -> list[dict]:
    """
    CAT 式简化匹配：编辑距离 + n-gram，返回所有 >= min_threshold 的候选，按匹配度降序。
    """
    if not sentence or not sentence.strip():
        return []

    if isinstance(debug_stats, dict):
        debug_stats.setdefault('considered_templates', 0)
        debug_stats.setdefault('matched_templates', 0)
        debug_stats.setdefault('returned_candidates', 0)
        debug_stats.setdefault('surface_rule_candidates', 0)
        debug_stats.setdefault('dropped_by_reason', {})

    best_by_template = {}
    sentence_text = sentence.strip()
    display_source_text = str(source_sentence or sentence_text or '').strip() or sentence_text
    normalized_sentence_text = _prepare_cat_match_text(sentence_text, term_dict, typo_dict).strip()
    source_match_text = normalized_sentence_text or sentence_text
    source_intent = _extract_sentence_intent(source_match_text)
    source_actions = set(source_intent.get('actions', []))
    source_action_count = len(source_actions)
    source_slot_sample_mode = _is_slot_sample_mapping_sentence(source_match_text)
    source_destinations = _extract_ui_destination_markers(source_match_text)
    source_ui_operation_markers = _extract_ui_operation_markers(source_match_text)
    if normalized_sentence_text:
        surface_display_text = _compose_cat_candidate_text(display_source_text, normalized_sentence_text)
        if (
            not _cat_exact_duplicate(display_source_text, surface_display_text)
            and not _is_trivial_cat_artifact_edit(display_source_text, surface_display_text)
        ):
            from app.utils.cat_diagnose import classify_surface_edit
            surface_kind = classify_surface_edit(display_source_text, surface_display_text)
            surface_candidate = {
                "template_text": surface_display_text,
                "raw_template_text": display_source_text,
                "normalized_template_text": normalized_sentence_text,
                "template_id": "surface_rule",
                "string_score": 0.98,
                "semantic_score": 0.95,
                "effective_semantic_score": 0.95,
                "match_tier": "surface_rule",
                "rule_source": "surface_rules",
                "type": surface_kind,
            }
            best_by_template['__surface_rule__'] = _apply_cat_entity_change_penalty(display_source_text, surface_candidate)
            if isinstance(debug_stats, dict):
                debug_stats['surface_rule_candidates'] += 1
    sentence_len = len(re.sub(r'\s+', '', source_match_text))
    for tpl in templates or []:
        tpl_text = tpl.get("text", "") if isinstance(tpl, dict) else str(tpl)
        tpl_id = tpl.get("id", "") if isinstance(tpl, dict) else ""
        if not tpl_text or not tpl_text.strip():
            _increment_debug_reason(debug_stats, 'empty_template')
            continue
        if isinstance(debug_stats, dict):
            debug_stats['considered_templates'] += 1
        if not _is_valid_cat_template(tpl_text):
            _increment_debug_reason(debug_stats, 'invalid_template')
            continue
        tpl_text = tpl_text.strip()
        normalized_tpl_match_text = _prepare_cat_match_text(tpl_text, term_dict, typo_dict).strip()
        template_match_text = normalized_tpl_match_text or tpl_text
        raw_tpl_len = len(re.sub(r'\s+', '', normalized_tpl_match_text or tpl_text))
        if sentence_len > 0 and raw_tpl_len > 0:
            raw_len_ratio = min(sentence_len, raw_tpl_len) / max(sentence_len, raw_tpl_len)
            template_slot_sample_mode = bool(_extract_slot_sample_pairs(template_match_text))
            if raw_len_ratio < 0.45 and not (source_slot_sample_mode and template_slot_sample_mode):
                _increment_debug_reason(debug_stats, 'pre_length_window')
                continue
        if _has_conflicting_list_intro(display_source_text, tpl_text) or _has_conflicting_list_intro(source_match_text, template_match_text):
            _increment_debug_reason(debug_stats, 'list_intro_conflict')
            continue
        if (
            _has_conflicting_list_step_template(display_source_text, tpl_text)
            or _has_conflicting_list_step_template(source_match_text, template_match_text)
        ):
            _increment_debug_reason(debug_stats, 'list_step_conflict')
            continue
        display_tpl_text = _compose_cat_candidate_text(display_source_text, normalized_tpl_match_text or tpl_text)
        display_tpl_match_text = _prepare_cat_match_text(display_tpl_text, term_dict, typo_dict).strip() or display_tpl_text
        template_intent = _extract_sentence_intent(template_match_text)
        template_actions = set(template_intent.get('actions', []))
        template_destinations = _extract_ui_destination_markers(template_match_text)
        template_ui_operation_markers = _extract_ui_operation_markers(template_match_text)
        if _cat_exact_duplicate(display_source_text, display_tpl_text):
            _increment_debug_reason(debug_stats, 'exact_duplicate')
            continue

        if _is_trivial_cat_artifact_edit(display_source_text, display_tpl_text):
            _increment_debug_reason(debug_stats, 'trivial_artifact')
            continue

        if _is_duplicated_clause_splice(display_source_text, display_tpl_text, tpl_text):
            _increment_debug_reason(debug_stats, 'duplicated_clause_splice')
            continue

        if (
            _has_injected_ui_bracket_labels(display_source_text, display_tpl_text)
            or _has_injected_ui_bracket_labels(display_source_text, tpl_text)
        ):
            _increment_debug_reason(debug_stats, 'injected_ui_label')
            continue

        if _is_fragment_candidate(source_match_text, display_tpl_match_text):
            _increment_debug_reason(debug_stats, 'fragment_or_length_window')
            continue
        score = _calc_similarity(source_match_text, display_tpl_match_text)
        raw_action_overlap = len(source_actions & template_actions)
        raw_clause_score = _clause_alignment_score(source_match_text, template_match_text)
        if source_action_count >= 2 and raw_action_overlap <= 0:
            if raw_clause_score < 0.32:
                score *= 0.72
            elif raw_clause_score < 0.40:
                score *= 0.85
        tpl_len = len(re.sub(r'\s+', '', display_tpl_match_text))
        len_ratio = 0.0
        if sentence_len > 0 and tpl_len > 0:
            len_ratio = min(sentence_len, tpl_len) / max(sentence_len, tpl_len)
            if len_ratio < 0.10:
                _increment_debug_reason(debug_stats, 'length_ratio_too_low')
                continue
            if len_ratio < 0.20:
                score *= max(len_ratio / 0.20, 0.50)
            elif len_ratio < 0.35:
                score *= max(len_ratio / 0.35, 0.75)

        effective_threshold = round(min_threshold, 4)

        if score < effective_threshold:
            _increment_debug_reason(debug_stats, 'below_threshold')
            continue

        if _has_conflicting_ui_destination_with_markers(
            source_destinations,
            template_destinations,
            source_actions,
            template_actions,
        ):
            _increment_debug_reason(debug_stats, 'ui_destination_conflict')
            continue

        if _has_missing_ui_operation_markers_with_markers(
            source_ui_operation_markers,
            template_ui_operation_markers,
        ):
            _increment_debug_reason(debug_stats, 'ui_operation_marker_missing')
            continue

        anchor_ok, anchor_reason = _key_term_anchor_consistent(source_match_text, display_tpl_match_text)
        if not anchor_ok:
            normalized_reason = re.sub(r'\s+', '_', str(anchor_reason or 'other')).strip('_') or 'other'
            _increment_debug_reason(debug_stats, f'anchor_mismatch:{normalized_reason}')
            continue

        if isinstance(debug_stats, dict):
            debug_stats['matched_templates'] += 1

        if score >= 1.0:
            tier = "exact"
        elif score >= fuzzy_lower:
            tier = "fuzzy"
        else:
            tier = "reference"

        candidate = {
            "template_text": display_tpl_text,
            "raw_template_text": tpl_text,
            "normalized_template_text": normalized_tpl_match_text or tpl_text,
            "template_id": str(tpl_id),
            "string_score": round(score, 4),
            "match_tier": tier,
            "rule_source": "sentence_guide",
        }
        candidate = _apply_cat_entity_change_penalty(display_source_text, candidate)
        if candidate.get('entity_penalty_reason'):
            _increment_debug_reason(debug_stats, 'entity_changed')
        dedupe_key = re.sub(r'\s+', '', re.sub(r'[，。！？!?；;：:、,\.\s]+$', '', normalized_tpl_match_text or tpl_text)).strip()
        existing = best_by_template.get(dedupe_key)
        if existing is None or candidate["string_score"] > existing["string_score"]:
            best_by_template[dedupe_key] = candidate

    results = list(best_by_template.values())
    results.sort(key=lambda x: x["string_score"], reverse=True)
    if isinstance(debug_stats, dict):
        debug_stats['returned_candidates'] = len(results)
        _finalize_simple_match_debug(debug_stats)
    return results


def _best_guarded_match(sentence: str, templates: list[str]) -> tuple:
    ranked = []
    for template in templates or []:
        entry_text = _template_entry_text(template)
        if not entry_text:
            continue
        metadata_score = _structured_template_metadata_score(sentence, template)
        for candidate in _template_entry_candidates(template):
            matched_template, score, level = _three_tier_match(sentence, [candidate])
            if not matched_template:
                continue
            if not _template_replace_guard(sentence, matched_template, level, score):
                continue
            structured_score = _structured_match_score(sentence, matched_template)
            cat_score = _cat_match_score(sentence, matched_template, template_metadata=metadata_score, structured_score=structured_score)
            final_score = cat_score.get('ranking_score', cat_score['overall_score'])
            ranked.append({
                'template': entry_text,
                'candidate_text': matched_template,
                'match_level': level,
                'guard_passed': True,
                'final_score': round(final_score, 4),
                'overall_score': round(cat_score.get('overall_score', 0.0), 4),
                'term_anchor_score': cat_score['term_anchor_score'],
                'number_placeholder_score': cat_score['number_placeholder_score'],
                'context_score': cat_score['context_score'],
            })

    if not ranked:
        return None, 0.0, 'NONE'

    if not ranked or ranked[0].get('final_score', 0) < 0.90:
        ranked = _ai_semantic_rerank_candidates(sentence, ranked)
    best = ranked[0]
    return best.get('template'), float(best.get('final_score', 0.0) or 0.0), best.get('match_level', 'NONE')


def _split_cat_sentences(paragraphs: list[str], source_paragraphs: Optional[list[str]] = None) -> list[dict]:
    sentence_items = []
    for para_idx, raw_paragraph in enumerate(paragraphs or []):
        paragraph_text = str(raw_paragraph or '').strip()
        source_raw = raw_paragraph
        if source_paragraphs and para_idx < len(source_paragraphs):
            source_raw = source_paragraphs[para_idx]
        source_paragraph_text = str(source_raw or '').strip()
        if not paragraph_text:
            continue
        compact_paragraph = re.sub(r'\s+', '', paragraph_text)
        if '\t' in paragraph_text and re.search(r'\t\s*\d+\s*$', paragraph_text):
            continue
        if re.fullmatch(r'\d+(?:\.\d+){0,4}[\u4e00-\u9fffA-Za-z（）()\-_/：:、\s]{0,24}', compact_paragraph):
            continue
        if re.match(r'^(?:表|图)\s*\d+[A-Za-z0-9.\-]*\s*', paragraph_text):
            continue
        if _is_short_label_like_text(paragraph_text) or _looks_like_title_or_noun_phrase(paragraph_text):
            continue
        chunks = re.split(r'(?<=[。！？!?；;])', paragraph_text)
        source_chunks = re.split(r'(?<=[。！？!?；;])', source_paragraph_text or paragraph_text)
        for chunk_index, chunk in enumerate(chunks):
            sentence_text = str(chunk or '').strip()
            source_sentence_text = str(source_chunks[chunk_index] or sentence_text).strip() if chunk_index < len(source_chunks) else sentence_text
            sentence_text = re.sub(r'^\d{1,3}(?:\.\d{1,3}){1,3}[.。、)）]?\s*', '', sentence_text.strip())
            if not sentence_text:
                continue
            normalized = re.sub(r'\s+', '', sentence_text)
            if re.match(r'^(?:表|图)\s*\d+[A-Za-z0-9.\-]*\s*', sentence_text):
                continue
            if _is_short_label_like_text(sentence_text) or _looks_like_title_or_noun_phrase(sentence_text):
                continue
            if len(normalized) <= 12:
                continue
            if not re.search(r'[。！？!?；;]', sentence_text) and len(normalized) < 18:
                continue
            if re.fullmatch(r'[\dA-Za-z\-_.()（）/]+', normalized):
                continue
            sentence_items.append({
                'sentence_index': len(sentence_items),
                'source_paragraph_index': para_idx,
                'source_paragraph_text': source_paragraph_text or paragraph_text,
                'text': sentence_text,
                'source_sentence_text': source_sentence_text or sentence_text,
            })
    return sentence_items


def _top_template_candidates(sentence: str, templates: list, limit: int = 8) -> list[dict]:
    ranked = []
    for template in templates or []:
        entry_text = _template_entry_text(template)
        if not entry_text:
            continue
        metadata_score = _structured_template_metadata_score(sentence, template)
        best_candidate_score = None
        best_candidate_detail = None
        best_candidate_level = 'NONE'
        best_candidate_guard = False
        best_candidate_text = entry_text

        for candidate in _template_entry_candidates(template):
            matched_template, score, level = _three_tier_match(sentence, [candidate])
            if not matched_template:
                continue
            guard_passed = _template_replace_guard(sentence, matched_template, level, score)
            structured_score = _structured_match_score(sentence, matched_template)
            cat_score = _cat_match_score(sentence, matched_template, template_metadata=metadata_score, structured_score=structured_score)
            final_score = cat_score.get('ranking_score', cat_score['overall_score'])
            if best_candidate_score is None or final_score > best_candidate_score:
                best_candidate_score = final_score
                best_candidate_detail = cat_score
                best_candidate_level = level
                best_candidate_guard = guard_passed
                best_candidate_text = matched_template

        if best_candidate_score is None:
            continue

        detail = dict(best_candidate_detail or _score_match_segments(sentence, best_candidate_text))
        final_percent = int(round(float(detail.get('overall_score', 0.0) or 0.0) * 100))
        band, label = _match_band_label(final_percent)
        detail['overall_score'] = round(float(detail.get('overall_score', 0.0) or 0.0), 4)
        detail['overall_percent'] = final_percent
        detail['band'] = band
        detail['label'] = label
        detail['match_level'] = best_candidate_level
        detail['guard_passed'] = best_candidate_guard
        detail['template'] = entry_text
        detail['candidate_text'] = best_candidate_text
        detail['final_score'] = round(best_candidate_score, 4)
        ranked.append(detail)

    if not ranked or ranked[0].get('final_score', 0) < 0.90:
        ranked = _ai_semantic_rerank_candidates(sentence, ranked)
    guarded_ranked = [item for item in ranked if item.get('guard_passed', False)]
    manual_ranked = [
        item for item in ranked
        if not item.get('guard_passed', False) and int(item.get('overall_percent', 0) or 0) >= 15
    ]

    primary = guarded_ranked[:limit]
    selected_keys = {
        (
            item.get('template', ''),
            item.get('candidate_text', ''),
            item.get('match_level', 'NONE'),
        )
        for item in primary
    }

    exploratory = []
    for item in guarded_ranked:
        key = (
            item.get('template', ''),
            item.get('candidate_text', ''),
            item.get('match_level', 'NONE'),
        )
        if key in selected_keys:
            continue
        if item.get('final_score', 0.0) < 0.5:
            exploratory.append(item)
            selected_keys.add(key)
        if len(exploratory) >= 4:
            break

    manual_candidates = []
    for item in manual_ranked:
        key = (
            item.get('template', ''),
            item.get('candidate_text', ''),
            item.get('match_level', 'NONE'),
        )
        if key in selected_keys:
            continue
        manual_candidates.append(item)
        selected_keys.add(key)
        if len(manual_candidates) >= 6:
            break

    return primary + exploratory + manual_candidates


def _should_keep_original_delete_span(text: str) -> bool:
    value = str(text or '').strip()
    if not value:
        return False
    if re.search(r'[（(][A-Za-z][A-Za-z0-9-]{3,}[）)]', value):
        return True
    if re.fullmatch(r'[A-Za-z][A-Za-z0-9-]{4,}', value):
        return True
    return False


def replace_with_context(original: str, template: str) -> str:
    """将原文与模板逐块融合。

    前缀/后缀上下文保留，匹配块之间的内部间隙移除。
    """
    s = SequenceMatcher(None, original, template)
    opcodes = list(s.get_opcodes())
    # 标记内部 delete 为 replace：位于匹配块之间的原文多余内容应移除
    for idx, (op, i1, i2, j1, j2) in enumerate(opcodes):
        if op == 'delete':
            has_before = any(o[0] in ('equal', 'replace') for o in opcodes[:idx])
            has_after = any(o[0] in ('equal', 'replace') for o in opcodes[idx+1:])
            if has_before and has_after:
                deleted = original[i1:i2]
                if not _should_keep_original_delete_span(deleted):
                    opcodes[idx] = ('replace', i1, i2, j1, j2)
    chars = list(original)
    offset = 0
    for op, i1, i2, j1, j2 in opcodes:
        if op in ('equal', 'delete'):
            continue
        elif op == 'replace':
            chars[i1 + offset : i2 + offset] = list(template[j1:j2])
            offset += (j2 - j1) - (i2 - i1)
        elif op == 'insert':
            chars[i1 + offset : i1 + offset] = list(template[j1:j2])
            offset += j2 - j1
    merged = ''.join(chars)
    if _is_concatenated_context_merge(original, template, merged):
        return template
    return merged


def _replace_clause_directly(clause: str, template: str, next_separator: str = '') -> str:
    prefix_match = re.match(r'^(\s*[*\-•·]+\s*)', clause or '')
    prefix = prefix_match.group(1) if prefix_match else ''
    replaced = (template or '').strip()
    if next_separator and next_separator[0] in '，,；;' and replaced.endswith(('。', '！', '？', '.', '!', '?', '；', ';')):
        replaced = replaced[:-1]
    replaced = _reapply_sentence_prefix(clause, replaced)
    return f'{prefix}{replaced}' if replaced else clause


def _apply_style_rules(text: str, rules: list[dict]) -> tuple[str, list[PolishRuleMatch]]:
    """对单句应用句式风格规则"""
    changes = []
    result = text
    
    forbidden_words_map = {
        "最佳": "较优",
        "最好": "较优", 
        "最著名": "知名",
        "最新技术": "先进技术",
        "最高水平": "先进水平",
        "最先进水平": "先进水平",
        "最高技术": "先进技术",
        "非常": "",
        "极其": "",
        "最": "较"
    }
    
    for rule in rules:
        if rule["type"] == "forbidden_words":
            # 优先使用规则中指定的替换表，回退到硬编码映射
            rule_replacements = rule.get("replacements", {}) or {}
            for phrase in rule["patterns"]:
                if phrase in result:
                    original = result
                    replacement = rule_replacements.get(phrase) or forbidden_words_map.get(phrase, "")
                    result = result.replace(phrase, replacement)
                    result = re.sub(r'\s+', ' ', result).strip()
                    if result != original:
                        changes.append(PolishRuleMatch(
                            rule_name=rule["name"],
                            before=f"...{phrase}...",
                            after=result[:50],
                            type="style"
                        ))

        elif rule["type"] == "preferred_sentences":
            best_sentence, best_score, match_level = _best_guarded_match(result, _preferred_sentence_entries(rule))
            if not best_sentence:
                continue

            if result.strip() != best_sentence.strip():
                original = result
                result = replace_with_context(original, best_sentence)
                changes.append(PolishRuleMatch(
                    rule_name=rule["name"],
                    before=original[:50],
                    after=result[:50],
                    type="style"
                ))
            elif result.strip() == best_sentence.strip():
                # 完全匹配无变化，记录为空匹配
                pass
        
        elif rule["type"] == "passive_voice":
            for pattern, issue in rule["patterns"]:
                match = re.search(pattern, result)
                if match:
                    changes.append(PolishRuleMatch(
                        rule_name=rule["name"],
                        before=match.group()[:50],
                        after="建议改用主动语态",
                        type="style"
                    ))
        
        elif rule["type"] == "double_negative":
            for pattern, issue in rule["patterns"]:
                match = re.search(pattern, result)
                if match:
                    changes.append(PolishRuleMatch(
                        rule_name=rule["name"],
                        before=match.group()[:50],
                        after="建议改用肯定表达",
                        type="style"
                    ))
        
        elif rule["type"] == "informal":
            informal_replacements = {
                "牛逼": "出色",
                "酷毙": "高效",
                "给力": "有效",
                "碉堡": "优异",
                "！": "。"
            }
            for pattern, issue in rule["patterns"]:
                match = re.search(pattern, result)
                if match:
                    original = result
                    for informal, formal in informal_replacements.items():
                        if informal in result:
                            result = result.replace(informal, formal)
                    result = re.sub(r'！+', '。', result)
                    result = result.strip()
                    if result != original:
                        changes.append(PolishRuleMatch(
                            rule_name=rule["name"],
                            before=match.group()[:50],
                            after=result[:50],
                            type="style"
                        ))
        
        elif rule["type"] == "sentence_length":
            clean = re.sub(r'[，。；！？、,.!?;、]', '', result)
            if len(clean) > rule["max_chars"]:
                changes.append(PolishRuleMatch(
                    rule_name=rule["name"],
                    before=f"句子长度{len(clean)}字",
                    after=f"建议控制在{rule['max_chars']}字以内",
                    type="style"
                ))
        
        elif rule["type"] == "pronoun_reference":
            for pattern, issue in rule["patterns"]:
                match = re.search(pattern, result)
                if match:
                    changes.append(PolishRuleMatch(
                        rule_name=rule["name"],
                        before=match.group()[:50],
                        after="需明确指代对象",
                        type="style"
                    ))
    
    # 约束润色：无匹配时进行术语标准化 + 去口语化
    if not changes or all(c.type != 'style' for c in changes):
        polished = _apply_constraint_polish(result)
        if polished != result:
            changes.append(PolishRuleMatch(
                rule_name="约束润色",
                before=result[:50],
                after=polished[:50],
                type="style"
            ))
            result = polished
    
    return result, changes


@router.post("/text")
async def polish_text_endpoint(input_data: TextPolishInput, db: Session = Depends(get_db)):
    """基础文本润色（自动加载句式清单和术语库）"""
    terminology_md = _load_scoped_terminology_source(
        db,
        input_data.terminology_id,
        input_data.product_type,
    )
    sentence_guide = _load_sentence_guides(
        db,
        style_guide_id=input_data.style_guide_id,
        product_type=input_data.product_type,
    )
    resolved_terminology = _resolve_terminology(db, terminology_md, input_data.text)
    ai_polished = input_data.text
    try:
        from app.utils.ai_client import ai_client
        from app.utils.cat_diagnose import use_lab_ai_provider
        with use_lab_ai_provider():
            result = ai_client.polish_text(
                input_data.text,
                style_guide=sentence_guide,
                terminology=resolved_terminology if resolved_terminology else None,
                request_label="polish.text",
            )
        ai_polished = _reapply_sentence_prefix(
            input_data.text,
            _protect_model_numbers(result.get("polished", input_data.text))
        )
    except Exception:
        pass

    cat_items = await _collect_text_polish_cat_items(
        input_data.text,
        db,
        sentence_guide,
        terminology_md,
        style_guide_id=input_data.style_guide_id,
        sentence_file_id=None,
    )
    if not isinstance(cat_items, list):
        cat_items = []
    diagnose_items = []
    try:
        from app.utils.cat_diagnose import (
            annotate_cat_candidates,
            is_ai_diagnose_enabled,
            merge_local_and_diagnoses,
            open_diagnose_sentences,
        )
        annotate_cat_candidates(cat_items)
        if is_ai_diagnose_enabled():
            try:
                sentence_items = _filter_cat_artifact_diagnose_pool(
                    _split_cat_sentences(input_data.text.split('\n'))
                )
                ai_diag = await open_diagnose_sentences(
                    sentence_items,
                    resolved_terminology or {},
                    sentence_guide,
                    input_data.product_type or "",
                )
                ai_diag = _filter_cat_artifact_diagnoses(ai_diag, sentence_items)
                _persist_lab_diagnoses(db, ai_diag, sentence_items, source="text", source_name="文本润色")
                cat_items, diagnose_items = merge_local_and_diagnoses(
                    cat_items,
                    ai_diag,
                    sentence_items,
                )
                diagnose_items = _filter_cat_artifact_diagnose_items(diagnose_items)
            except Exception:
                diagnose_items = []
    except Exception:
        diagnose_items = []

    polished_text, rule_changes = _apply_skill_polish(
        ai_polished,
        {},
        db,
        sentence_guide,
        terminology_md,
        None,
        db_terminology=resolved_terminology if resolved_terminology else None,
    )
    changes = []
    if ai_polished != input_data.text:
        ai_entry = _build_visible_change_entry(1, input_data.text, ai_polished, "ai", "ai", sentence_guide)
        if ai_entry:
            changes.append(ai_entry)
    for change in rule_changes:
        entry = _build_visible_change_entry(
            1,
            change.before,
            change.after,
            change.type,
            change.rule_name,
            sentence_guide,
        )
        if entry:
            entry["match_detail"] = change.match_detail or entry.get("match_detail")
            changes.append(entry)
    final_polished = _reapply_sentence_prefix(input_data.text, polished_text)
    if not changes and _is_low_value_doc_change(input_data.text, final_polished, 'style', '基础规范化'):
        final_polished = input_data.text
    return {
        "original": input_data.text,
        "base_polished": final_polished,
        "polished": final_polished,
        "changes": changes,
        "cat_items": cat_items,
        "diagnose_items": diagnose_items,
    }


@router.post("/skill")
async def polish_with_skill(
    input_data: SkillPolishInput,
    db: Session = Depends(get_db)
):
    """使用内置润色规则进行润色（V2 四层规则体系）"""
    sentence_guide = _load_sentence_guides(db, style_guide_id=input_data.style_guide_id)
    if not sentence_guide:
        sentence_guide = _build_document_polish_guide(db)

    # 加载术语：文件术语优先，否则回退数据库术语
    terminology_md = _load_terminology_source(db, input_data.terminology_id)

    # 规则引擎预处理
    try:
        from app.utils.instrument_polisher import instrument_polish_engine
        pre_text = instrument_polish_engine.pre_polish(input_data.text)
    except Exception:
        pre_text = input_data.text

    ai_polished = pre_text
    ai_changes = []
    try:
        from app.utils.ai_client import ai_client
        from app.utils.cat_diagnose import use_lab_ai_provider
        resolved_terminology = _resolve_terminology(db, terminology_md, pre_text)
        with use_lab_ai_provider():
            ai_result = ai_client.polish_text(
                pre_text,
                style_guide=sentence_guide,
                terminology=resolved_terminology if resolved_terminology else None,
                request_label="polish.skill",
            )
        if ai_result and ai_result.get("polished"):
            ai_polished = _reapply_sentence_prefix(pre_text, _protect_model_numbers(ai_result["polished"]))
            if ai_polished != pre_text:
                ai_changes = [change for change in [_build_visible_change_entry(1, input_data.text, ai_polished, 'ai', 'ai', sentence_guide)] if change]
    except Exception:
        pass

    # 规则引擎后置保护
    try:
        if ai_polished != pre_text:
            protect_result = instrument_polish_engine.post_protect(pre_text, ai_polished)
            if not protect_result.get("safe"):
                ai_polished = protect_result.get("suggested", pre_text)
    except Exception:
        pass

    polished_text, rule_changes = _apply_skill_polish(
        ai_polished,
        {},
        db,
        sentence_guide,
        terminology_md,
        None,
    )
    all_changes = ai_changes + rule_changes
    final_polished = _reapply_sentence_prefix(input_data.text, polished_text)
    if not all_changes and _is_low_value_doc_change(input_data.text, final_polished, 'style', '基础规范化'):
        final_polished = input_data.text
    
    return {
        "original": input_data.text,
        "polished": final_polished,
        "changes": all_changes,
        "skill_name": "技术文档智能润色 (V2)",
        "rules_applied": {
            "术语替换": True,
            "句式规范": True,
            "风格规范": True,
            "微调优化": True
        }
    }



# ============================================================
# DOCX 润色与批注注入
# ============================================================

def _xml_local_name(element) -> str:
    tag = getattr(element, 'tag', '')
    if not tag:
        return ''
    return tag.split('}')[-1] if '}' in tag else tag


def _read_docx_document_root(docx_path: str):
    import zipfile
    from lxml import etree

    with zipfile.ZipFile(docx_path, 'r') as zin:
        document_xml = zin.read('word/document.xml')
    return etree.fromstring(document_xml)


def _iter_docx_paragraphs_in_order(container):
    from docx.document import Document as DocxDocument
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    if isinstance(container, DocxDocument):
        parent_elm = container.element.body
        parent = container
    elif isinstance(container, _Cell):
        parent_elm = container._tc
        parent = container
    elif isinstance(container, Table):
        for row in container.rows:
            for cell in row.cells:
                yield from _iter_docx_paragraphs_in_order(cell)
        return
    else:
        return

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield from _iter_docx_paragraphs_in_order(Table(child, parent))


def _iter_docx_body_paragraphs(container):
    """只遍历文档正文段落，跳过表格内容。
    CAT 分析和写回都基于该顺序，避免表格单元格文本污染候选与索引。"""
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph
    from docx.oxml.text.paragraph import CT_P

    if not isinstance(container, DocxDocument):
        return

    for child in container.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, container)


def _is_simple_revision_paragraph(p_element, w_ns: str) -> bool:
    """判断段落是否为纯文本段落，可安全写入修订标记。
    允许纯文本多 run 段落，每个 run 必须仅包含 rPr + 单个 t。
    仅跳过包含绘图、域代码等复杂结构的段落。"""
    allowed_children = {'pPr', 'r'}
    allowed_run_children = {'rPr', 't'}

    run_count = 0
    for child in list(p_element):
        child_name = _xml_local_name(child)
        if child_name not in allowed_children:
            return False
        if child_name == 'r':
            run_count += 1
            text_child_count = 0
            for run_child in list(child):
                run_child_name = _xml_local_name(run_child)
                if run_child_name not in allowed_run_children:
                    return False
                if run_child_name == 't':
                    text_child_count += 1
            if text_child_count != 1:
                return False

    return run_count > 0


def _extract_run_visible_text(run_element, w_ns: str) -> str:
    parts = []
    for child in list(run_element):
        child_name = _xml_local_name(child)
        if child_name in {'t', 'delText'}:
            parts.append(child.text or '')
        elif child_name == 'tab':
            parts.append('\t')
        elif child_name in {'br', 'cr'}:
            parts.append('\n')
    return ''.join(parts)


def _split_text_by_run_lengths(text: str, run_lengths: list[int]) -> list[str]:
    if not run_lengths:
        return [text]
    if len(run_lengths) == 1:
        return [text]

    total = sum(max(length, 0) for length in run_lengths)
    if total <= 0:
        return [text] + [''] * (len(run_lengths) - 1)

    chunks = []
    consumed = 0
    cumulative = 0
    text_length = len(text)
    for index, run_length in enumerate(run_lengths):
        if index == len(run_lengths) - 1:
            chunks.append(text[consumed:])
            break
        cumulative += max(run_length, 0)
        next_consumed = round(text_length * cumulative / total)
        chunks.append(text[consumed:next_consumed])
        consumed = next_consumed
    return chunks


def _strip_revision_display_props(rpr_element, w_ns: str):
    if rpr_element is None:
        return
    for child in list(rpr_element):
        if _xml_local_name(child) in {'color', 'highlight', 'shd'}:
            rpr_element.remove(child)


def _apply_paragraph_revision_xml(p_element, polished_text: str, author: str, now: str, rid_del: str, rid_ins: str, w_ns: str):
    from copy import deepcopy
    from lxml import etree

    children = list(p_element)
    run_indexes = [index for index, child in enumerate(children) if _xml_local_name(child) == 'r']
    if not run_indexes:
        return False

    first_run = children[run_indexes[0]]
    is_single_run = len(run_indexes) == 1

    # ── <w:del>: 保留所有原始 run 的完整格式 ──
    del_element = etree.Element(f'{{{w_ns}}}del')
    del_element.set(f'{{{w_ns}}}id', rid_del)
    del_element.set(f'{{{w_ns}}}author', author)
    del_element.set(f'{{{w_ns}}}date', now)

    for run_index in run_indexes:
        cloned_run = deepcopy(children[run_index])
        for text_element in cloned_run.findall(f'.//{{{w_ns}}}t'):
            text_element.tag = f'{{{w_ns}}}delText'
        del_element.append(cloned_run)

    # ── <w:ins>: 单 run 切分文本，多 run 仅写一个干净 run ──
    ins_element = etree.Element(f'{{{w_ns}}}ins')
    ins_element.set(f'{{{w_ns}}}id', rid_ins)
    ins_element.set(f'{{{w_ns}}}author', author)
    ins_element.set(f'{{{w_ns}}}date', now)

    if is_single_run:
        original_runs = [children[index] for index in run_indexes]
        run_lengths = [len(_extract_run_visible_text(run, w_ns)) for run in original_runs]
        text_chunks = _split_text_by_run_lengths(polished_text, run_lengths)

        appended = False
        for run_element, text_chunk in zip(original_runs, text_chunks):
            if not text_chunk and appended:
                continue
            inserted_run = etree.Element(f'{{{w_ns}}}r')
            first_rpr = run_element.find(f'{{{w_ns}}}rPr')
            if first_rpr is not None:
                cloned_rpr = deepcopy(first_rpr)
                _strip_revision_display_props(cloned_rpr, w_ns)
                inserted_run.append(cloned_rpr)
            text_element = etree.SubElement(inserted_run, f'{{{w_ns}}}t')
            if text_chunk[:1].isspace() or text_chunk[-1:].isspace():
                text_element.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            text_element.text = text_chunk or ''
            ins_element.append(inserted_run)
            appended = True

        if not appended:
            inserted_run = etree.Element(f'{{{w_ns}}}r')
            first_rpr = first_run.find(f'{{{w_ns}}}rPr')
            if first_rpr is not None:
                cloned_rpr = deepcopy(first_rpr)
                _strip_revision_display_props(cloned_rpr, w_ns)
                inserted_run.append(cloned_rpr)
            text_element = etree.SubElement(inserted_run, f'{{{w_ns}}}t')
            text_element.text = polished_text
            ins_element.append(inserted_run)
    else:
        # 多 run 段落：仅插入一个干净 run，不再切分文本以避免格式漂移
        inserted_run = etree.Element(f'{{{w_ns}}}r')
        first_rpr = first_run.find(f'{{{w_ns}}}rPr')
        if first_rpr is not None:
            cloned_rpr = deepcopy(first_rpr)
            _strip_revision_display_props(cloned_rpr, w_ns)
            inserted_run.append(cloned_rpr)
        text_element = etree.SubElement(inserted_run, f'{{{w_ns}}}t')
        if polished_text[:1].isspace() or polished_text[-1:].isspace():
            text_element.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        text_element.text = polished_text
        ins_element.append(inserted_run)

    # ── 重建段落子元素 ──
    first_run_index = run_indexes[0]
    last_run_index = run_indexes[-1]
    new_children = []
    for index, child in enumerate(children):
        if index == first_run_index:
            new_children.append(del_element)
            new_children.append(ins_element)
        if first_run_index <= index <= last_run_index and _xml_local_name(child) == 'r':
            continue
        new_children.append(child)

    for child in list(p_element):
        p_element.remove(child)
    for child in new_children:
        p_element.append(child)
    return True

def _polish_docx_with_comments(
    docx_path: str,
    output_path: str,
    skill_rules: dict,
    db: Session,
    sentence_guide: str = None,
    terminology: str = None,
    requirements: str = None,
    ai_lines: list = None,
    db_terminology: dict = None
) -> tuple[list[PolishRuleMatch], list[dict]]:
    """对 DOCX 文件进行润色，保留排版并添加批注。
    
    ai_lines: AI 预润色后的文本行列表，与原始段落逐行对应。用于句式清单匹配润色。
    """
    from docx import Document
    from lxml import etree
    
    all_changes = []
    review_items = []
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    
    doc = Document(docx_path)
    document_root = _read_docx_document_root(docx_path)
    xml_paragraphs = document_root.findall(f'.//{{{w_ns}}}body//{{{w_ns}}}p')
    
    term_dict = {}
    if terminology:
        try:
            parsed = _parse_terminology(terminology)
            if parsed:
                all_text = '\n'.join([p.text for p in _iter_docx_paragraphs_in_order(doc) if p.text and p.text.strip()])
                lang = _detect_language(all_text)
                term_dict = _filter_terms_by_lang(parsed, lang)
        except Exception:
            pass
    if db_terminology:
        term_dict.update(db_terminology)
    
    author = "技术文档智能润色助手"
    revision_id = 0
    rejected_change_keys = _load_rejected_doc_change_keys(db)
    rule_polish_elapsed = 0.0
    review_detail_elapsed = 0.0
    revision_write_elapsed = 0.0
    
    toc_prefixes = ("TOC", "Table of Contents", "目录")
    
    ordered_paragraphs = list(_iter_docx_paragraphs_in_order(doc))
    non_empty_paras = [(idx, para) for idx, para in enumerate(ordered_paragraphs) 
                       if para.text and para.text.strip()]
    non_empty_ai_lines = [l.strip() for l in ai_lines if l.strip()] if ai_lines else []
    if non_empty_ai_lines and len(non_empty_ai_lines) != len(non_empty_paras):
        # 行数不一致时无法可靠映射到 Word 段落，避免标题/图注错位被覆盖。
        non_empty_ai_lines = []

    rerank_token = _ai_template_rerank_enabled.set(False)
    try:
        for i, (para_idx, para) in enumerate(non_empty_paras):
            original_text = para.text.strip()
            paragraph_started = time.perf_counter()

            style_name = (para.style.name or "").lower()

            # 仅跳过确认为目录的段落（样式名以 "toc" 开头或为 "toc heading"）
            is_toc = style_name.startswith('toc') or style_name == 'table of contents'

            if is_toc:
                continue

            title_keywords = ['heading', 'title', '目录', '标题', 'toc', '表', '图', 'table', 'figure', 'caption']
            is_title = any(kw in (style_name or "").lower() for kw in title_keywords) if style_name else False
            if not is_title and original_text.strip():
                if re.match(r'^(表|图|Table|Figure)\s*\d', original_text.strip()):
                    is_title = True

            # Step 1: AI polish
            intermediate_text = original_text
            para_ai_change = None
            if (not is_title) and i < len(non_empty_ai_lines):
                ai_line = _normalize_doc_ai_line(original_text, non_empty_ai_lines[i])
                if ai_line and ai_line != original_text:
                    intermediate_text = ai_line
                    para_ai_change = PolishRuleMatch(
                        rule_name="ai", before=original_text[:100], after=intermediate_text[:100], type="ai"
                    )

            # Step 2: Rule polish + 术语替换
            rule_polish_started = time.perf_counter()
            polished_text, rule_changes = _apply_skill_polish(
                intermediate_text, skill_rules, db, sentence_guide, terminology, requirements,
                is_title=is_title, db_terminology=db_terminology
            )
            para_changes = ([para_ai_change] if para_ai_change else []) + rule_changes

            # 最终术语替换（无论 AI 是否已改，确保术语库强制生效）
            term_changes = []
            if term_dict:
                polished_text, term_changes = _apply_term_only(polished_text, term_dict)
                if term_changes:
                    para_changes.extend(term_changes)
            rule_polish_elapsed += time.perf_counter() - rule_polish_started
            paragraph_rule_polish_elapsed = time.perf_counter() - rule_polish_started

            polished_text = _normalize_doc_polished_text(original_text, polished_text, is_title=is_title)

            primary_change = next((change for change in para_changes if change.type != 'ai'), para_changes[0] if para_changes else None)
            primary_type = primary_change.type if primary_change else ('terminology' if term_changes else 'style')
            primary_rule = primary_change.rule_name if primary_change else '句式评分'

            review_detail_started = time.perf_counter()
            review_suggestion_text = polished_text if polished_text != original_text else ''
            normalized_review_suggestion = ''
            if review_suggestion_text:
                normalized_review_suggestion = _normalize_review_suggestion(
                    original_text,
                    review_suggestion_text,
                    primary_type,
                    primary_rule,
                    is_title=is_title,
                )
            elif _should_skip_expensive_template_match(original_text):
                review_suggestion_text = _lightweight_operation_review_suggestion(original_text)
                if review_suggestion_text:
                    normalized_review_suggestion = _normalize_review_suggestion(
                        original_text,
                        review_suggestion_text,
                        primary_type,
                        primary_rule,
                        is_title=is_title,
                    )
            if (
                (_should_use_lightweight_review_detail(original_text, polished_text, is_title=is_title) and not normalized_review_suggestion) or
                (_is_simple_operation_sentence(original_text) and not normalized_review_suggestion) or
                (_is_pooling_platform_sentence(original_text) and polished_text != original_text and not normalized_review_suggestion)
            ):
                review_detail = _lightweight_no_change_review_detail('长段未改动，跳过候选详情')
                review_after = original_text
            else:
                review_detail, review_after = _build_sentence_review_detail(
                    original_text,
                    review_suggestion_text,
                    sentence_guide,
                    primary_type,
                    primary_rule,
                    is_title=is_title,
                )
            review_detail_elapsed += time.perf_counter() - review_detail_started
            paragraph_review_detail_elapsed = time.perf_counter() - review_detail_started
            review_items.append({
                'before': original_text,
                'after': review_after,
                'type': primary_type,
                'rule_name': primary_rule,
                'paragraph': para_idx + 1,
                'is_title': is_title,
                'match_detail': review_detail,
            })

            if _should_emit_reference_review_change(original_text, polished_text, review_after, primary_type, primary_rule):
                all_changes.append(PolishRuleMatch(
                    rule_name=primary_rule,
                    before=original_text,
                    after=review_after,
                    type=primary_type,
                    paragraph=para_idx + 1,
                    match_detail=review_detail,
                ))

            if polished_text != original_text and (para_changes or polished_text.strip() != original_text.strip()):
                primary_type = primary_change.type if primary_change else ('terminology' if term_changes else 'format')
                primary_rule = primary_change.rule_name if primary_change else 'Word修订标记'
                if _is_low_value_doc_change(original_text, polished_text, primary_type, primary_rule):
                    continue
                if _is_rejected_doc_change(original_text, polished_text, rejected_change_keys):
                    continue

                if para_idx < len(xml_paragraphs):
                    p_element = xml_paragraphs[para_idx]
                    if _is_simple_revision_paragraph(p_element, w_ns):
                        revision_started = time.perf_counter()
                        revision_id += 1
                        rid_del = str(revision_id)
                        now = '2026-06-18T00:00:00Z'
                        revision_id += 1
                        rid_ins = str(revision_id)
                        _apply_paragraph_revision_xml(p_element, polished_text, author, now, rid_del, rid_ins, w_ns)
                        revision_write_elapsed += time.perf_counter() - revision_started

                all_changes.append(PolishRuleMatch(
                    rule_name=primary_rule,
                    before=original_text,
                    after=polished_text,
                    type=primary_type,
                    paragraph=para_idx + 1,
                    match_detail=review_detail,
                ))

            processed_count = i + 1
            paragraph_elapsed = time.perf_counter() - paragraph_started
            if paragraph_elapsed >= 0.8:
                logger.warning(
                    "[POLISH_LAB_SLOW_PARAGRAPH] processed=%s/%s paragraph=%s is_title=%s total_s=%.3f rule_polish_s=%.3f review_detail_s=%.3f text=%r",
                    processed_count,
                    len(non_empty_paras),
                    para_idx + 1,
                    is_title,
                    paragraph_elapsed,
                    paragraph_rule_polish_elapsed,
                    paragraph_review_detail_elapsed,
                    original_text[:80],
                )
            if processed_count % 25 == 0:
                logger.warning(
                    "[POLISH_LAB_PROGRESS] processed=%s/%s rule_polish_s=%.3f review_detail_s=%.3f revision_write_s=%.3f",
                    processed_count,
                    len(non_empty_paras),
                    rule_polish_elapsed,
                    review_detail_elapsed,
                    revision_write_elapsed,
                )
    finally:
        _ai_template_rerank_enabled.reset(rerank_token)

    logger.warning(
        "[POLISH_LAB_TIMING] paragraphs=%s rule_polish_s=%.3f review_detail_s=%.3f revision_write_s=%.3f",
        len(non_empty_paras),
        rule_polish_elapsed,
        review_detail_elapsed,
        revision_write_elapsed,
    )

    document_xml = etree.tostring(document_root, xml_declaration=True, encoding='UTF-8')
    _write_revised_docx(docx_path, output_path, document_xml)
    
    return all_changes, review_items


def _write_revised_docx(source_docx_path: str, output_docx_path: str, document_xml: bytes):
    """基于原始 DOCX 仅替换 document.xml 与 settings.xml，保留 Word 原始排版结构。"""
    import zipfile
    import os as os_mod
    from lxml import etree
    
    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    temp_path = output_docx_path + '.tmp'
    
    try:
        with zipfile.ZipFile(source_docx_path, 'r') as zin, \
             zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            settings_found = False
            
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = document_xml
                elif item.filename == 'word/settings.xml':
                    settings_found = True
                    root = etree.fromstring(data)
                    existing = root.findall(f'{{{w_ns}}}trackRevisions')
                    if not existing:
                        etree.SubElement(root, f'{{{w_ns}}}trackRevisions')
                    revision_view = root.find(f'{{{w_ns}}}revisionView')
                    if revision_view is None:
                        revision_view = etree.SubElement(root, f'{{{w_ns}}}revisionView')
                    revision_view.set(f'{{{w_ns}}}markup', '1')
                    revision_view.set(f'{{{w_ns}}}comments', '0')
                    revision_view.set(f'{{{w_ns}}}insDel', '1')
                    revision_view.set(f'{{{w_ns}}}formatting', '1')
                    data = etree.tostring(root, xml_declaration=True, encoding='UTF-8')
                
                zout.writestr(item, data)
            
            if not settings_found:
                root = etree.Element(f'{{{w_ns}}}settings')
                etree.SubElement(root, f'{{{w_ns}}}trackRevisions')
                revision_view = etree.SubElement(root, f'{{{w_ns}}}revisionView')
                revision_view.set(f'{{{w_ns}}}markup', '1')
                revision_view.set(f'{{{w_ns}}}comments', '0')
                revision_view.set(f'{{{w_ns}}}insDel', '1')
                revision_view.set(f'{{{w_ns}}}formatting', '1')
                data = etree.tostring(root, xml_declaration=True, encoding='UTF-8')
                zi = zipfile.ZipInfo('word/settings.xml')
                zout.writestr(zi, data)
        
        os_mod.replace(temp_path, output_docx_path)
    finally:
        if os_mod.path.exists(temp_path):
            os_mod.remove(temp_path)


def _revision_element_text(element, w_ns: str, text_tags: set[str]) -> str:
    parts = []
    for node in element.iter():
        if _xml_local_name(node) in text_tags:
            parts.append(node.text or '')
    return ''.join(parts)


def _restore_deleted_runs(del_element, w_ns: str):
    from copy import deepcopy

    restored = []
    for run in del_element.findall(f'{{{w_ns}}}r'):
        cloned = deepcopy(run)
        for text_element in cloned.findall(f'.//{{{w_ns}}}delText'):
            text_element.tag = f'{{{w_ns}}}t'
        restored.append(cloned)
    return restored


def _replace_revision_insert_text(ins_element, text: str, w_ns: str):
    text_nodes = [node for node in ins_element.iter() if _xml_local_name(node) == 't']
    if text_nodes:
        text_nodes[0].text = text
        if text[:1].isspace() or text[-1:].isspace():
            text_nodes[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        for node in text_nodes[1:]:
            node.text = ''


def _accepted_revision_runs(ins_element, text: str, w_ns: str):
    from copy import deepcopy
    from lxml import etree

    runs = ins_element.findall(f'{{{w_ns}}}r')
    if runs:
        cloned = deepcopy(runs[0])
        for node in list(cloned):
            if _xml_local_name(node) in {'delText'}:
                cloned.remove(node)
        text_nodes = [node for node in cloned.iter() if _xml_local_name(node) == 't']
        if text_nodes:
            text_nodes[0].text = text
            if text[:1].isspace() or text[-1:].isspace():
                text_nodes[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            for node in text_nodes[1:]:
                node.text = ''
        else:
            text_element = etree.SubElement(cloned, f'{{{w_ns}}}t')
            text_element.text = text
        return [cloned]

    run = etree.Element(f'{{{w_ns}}}r')
    text_element = etree.SubElement(run, f'{{{w_ns}}}t')
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    text_element.text = text
    return [run]


def _paragraph_revision_text(paragraph, w_ns: str) -> str:
    parts = []
    for node in paragraph.iter():
        if _xml_local_name(node) in {'t', 'delText'}:
            parts.append(node.text or '')
    return ''.join(parts)


def _paragraph_visible_text(paragraph, w_ns: str) -> str:
    parts = []
    for node in paragraph.iter():
        if _xml_local_name(node) == 't':
            parts.append(node.text or '')
    return ''.join(parts)


def _replace_paragraph_text_xml(paragraph, text: str, w_ns: str):
    from copy import deepcopy
    from lxml import etree

    ppr = None
    for child in list(paragraph):
        if _xml_local_name(child) == 'pPr':
            ppr = deepcopy(child)
            break

    for child in list(paragraph):
        paragraph.remove(child)
    if ppr is not None:
        paragraph.append(ppr)

    run = etree.SubElement(paragraph, f'{{{w_ns}}}r')
    text_element = etree.SubElement(run, f'{{{w_ns}}}t')
    if text[:1].isspace() or text[-1:].isspace():
        text_element.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    text_element.text = text


def _target_paragraph_by_index(root, paragraph_index: Optional[int], w_ns: str):
    if not paragraph_index or paragraph_index < 1:
        return None
    paragraphs = root.findall(f'.//{{{w_ns}}}body/{{{w_ns}}}p')
    index = paragraph_index - 1
    if index < 0 or index >= len(paragraphs):
        return None
    return paragraphs[index]


def _text_matches_decision(candidate: str, expected: str) -> bool:
    candidate_key = _normalize_compare_text(candidate)
    expected_key = _normalize_compare_text(expected)
    if not candidate_key or not expected_key:
        return False
    return (
        candidate_key == expected_key or
        candidate_key.startswith(expected_key) or
        expected_key.startswith(candidate_key)
    )


def _force_apply_feedback_to_document_xml(root, decisions: list[DocumentFeedbackItem], w_ns: str) -> list[dict]:
    """最终 XML 兜底：按段落文本强制落地接受/拒绝后的正文。"""
    applied = []
    for decision in decisions or []:
        before = (decision.before or '').strip()
        after = (decision.after or '').strip()
        status = (decision.status or '').strip()
        if not before or not after:
            continue
        if decision.accepted:
            match_text = before
            target_text = after
        elif status == 'rejected':
            match_text = after
            target_text = before
        else:
            continue

        target_paragraph = _target_paragraph_by_index(root, decision.paragraph, w_ns)
        if target_paragraph is not None:
            visible_text = _paragraph_visible_text(target_paragraph, w_ns)
            if _normalize_compare_text(visible_text) != _normalize_compare_text(target_text):
                _apply_paragraph_revision_xml(
                    target_paragraph,
                    target_text,
                    "技术文档智能润色助手",
                    datetime.datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'),
                    str(uuid.uuid4()),
                    str(uuid.uuid4()),
                    w_ns,
                )
                applied.append({
                    'before': before,
                    'after': target_text,
                    'type': decision.type,
                    'rule_name': '用户确认',
                    'paragraph': decision.paragraph,
                })
            continue

        for paragraph in root.findall(f'.//{{{w_ns}}}p'):
            visible_text = _paragraph_visible_text(paragraph, w_ns)
            revision_text = _paragraph_revision_text(paragraph, w_ns)
            if _normalize_compare_text(visible_text) == _normalize_compare_text(target_text):
                break
            if _text_matches_decision(visible_text, match_text) or _text_matches_decision(revision_text, match_text):
                _apply_paragraph_revision_xml(
                    paragraph,
                    target_text,
                    "技术文档智能润色助手",
                    datetime.datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ'),
                    str(uuid.uuid4()),
                    str(uuid.uuid4()),
                    w_ns,
                )
                applied.append({
                    'before': before,
                    'after': target_text,
                    'type': decision.type,
                    'rule_name': '用户确认',
                    'paragraph': decision.paragraph,
                })
                break
    return applied


def _apply_feedback_plaintext_fallback(docx_path: str, decisions: list[DocumentFeedbackItem], already_applied: list[dict]) -> list[dict]:
    """兜底写回：按段落文本确认接受/拒绝决策已经落到正文。"""
    if not docx_path or not os.path.exists(docx_path):
        return []

    from docx import Document

    pending = []
    for decision in decisions or []:
        before = (decision.before or '').strip()
        after = (decision.after or '').strip()
        if not before or not after:
            continue
        pending.append(decision)

    if not pending:
        return []

    doc = Document(docx_path)
    fallback_applied = []
    changed = False
    for decision in pending:
        before = (decision.before or '').strip()
        after = (decision.after or '').strip()
        status = (decision.status or '').strip()
        if decision.accepted:
            target_text = after
            match_text = before
        elif status == 'rejected':
            target_text = before
            match_text = after
        else:
            continue

        for paragraph in doc.paragraphs:
            if _normalize_compare_text(paragraph.text) == _normalize_compare_text(target_text):
                break
            if _text_matches_decision(paragraph.text, match_text):
                paragraph.text = target_text
                fallback_applied.append({
                    'before': before,
                    'after': target_text,
                    'type': decision.type,
                    'rule_name': '用户确认',
                })
                changed = True
                break

    if changed:
        doc.save(docx_path)
    return fallback_applied


def _apply_feedback_to_revised_docx(source_docx_path: str, output_docx_path: str, decisions: list[DocumentFeedbackItem]) -> list[dict]:
    """按用户接受/拒绝/自定义决策重写修订版 DOCX。"""
    from lxml import etree

    if not source_docx_path or not output_docx_path or not os.path.exists(source_docx_path):
        return []

    w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    root = _read_docx_document_root(source_docx_path)
    decisions_by_before = {}
    for item in decisions or []:
        before_key = _normalize_compare_text(item.before)
        if not before_key:
            continue
        decisions_by_before.setdefault(before_key, []).append(item)

    def _pop_decision(before_text: str):
        before_key = _normalize_compare_text(before_text)
        candidates = decisions_by_before.get(before_key)
        if candidates:
            return candidates.pop(0)
        for decision_key, decision_items in decisions_by_before.items():
            if not decision_items:
                continue
            if before_key.startswith(decision_key) or decision_key.startswith(before_key):
                return decision_items.pop(0)
        return None

    applied = []
    for paragraph in root.findall(f'.//{{{w_ns}}}p'):
        children = list(paragraph)
        new_children = []
        index = 0
        paragraph_replaced = False
        while index < len(children):
            current = children[index]
            next_child = children[index + 1] if index + 1 < len(children) else None
            if _xml_local_name(current) == 'del' and next_child is not None and _xml_local_name(next_child) == 'ins':
                before = _revision_element_text(current, w_ns, {'delText'})
                original_after = _revision_element_text(next_child, w_ns, {'t'})
                decision = _pop_decision(before)
                decision_status = (decision.status or '').strip() if decision else ''
                if decision and decision.accepted and (decision.after or '').strip():
                    after = (decision.after or '').strip()
                    _replace_revision_insert_text(next_child, after, w_ns)
                    paragraph_replaced = True
                    applied.append({
                        'before': before,
                        'after': after or original_after,
                        'type': decision.type,
                        'rule_name': '用户确认',
                        'paragraph': decision.paragraph,
                    })
                    break
                elif decision_status == 'rejected':
                    new_children.extend(_restore_deleted_runs(current, w_ns))
                else:
                    new_children.append(current)
                    new_children.append(next_child)
                index += 2
                continue
            new_children.append(current)
            index += 1

        if paragraph_replaced:
            continue

        if len(new_children) != len(children):
            for child in list(paragraph):
                paragraph.remove(child)
            for child in new_children:
                paragraph.append(child)

    applied.extend(_force_apply_feedback_to_document_xml(root, decisions, w_ns))
    document_xml = etree.tostring(root, xml_declaration=True, encoding='UTF-8')
    _write_revised_docx(source_docx_path, output_docx_path, document_xml)
    applied.extend(_apply_feedback_plaintext_fallback(output_docx_path, decisions, applied))
    return applied




# ============================================================
# 润色报告生成
# ============================================================

def _generate_polish_report(
    report_path: str,
    original_filename: str,
    changes: list,
    sentence_file_name: str = None,
    terminology_file_name: str = None,
    requirements: str = None
):
    """生成润色报告 DOCX 文件"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    from docx.oxml.ns import qn
    
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    title = doc.add_heading('技术文档润色报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    p = doc.add_paragraph()
    p.add_run('生成日期：').bold = True
    p.add_run(now)
    
    doc.add_heading('1. 文档基本信息', level=2)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Shading Accent 1'
    
    info_data = [
        ('原文件名', original_filename),
        ('文件大小', '待计算'),
        ('句式参考', sentence_file_name or '未指定'),
        ('术语库', terminology_file_name or '未指定'),
        ('润色要求', requirements or '无'),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
        for cell in info_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_heading('2. 改动统计', level=2)
    
    total_changes = len(changes)
    change_types = {}
    for c in changes:
        ct = c.get('type', 'unknown') if isinstance(c, dict) else getattr(c, 'type', 'unknown')
        change_types[ct] = change_types.get(ct, 0) + 1
    
    p = doc.add_paragraph()
    p.add_run('总改动数：').bold = True
    p.add_run(f'{total_changes} 处')
    
    p = doc.add_paragraph()
    p.add_run('改动类型分布：')
    
    type_names = {
        'terminology': '术语替换',
        'format': '格式规范',
        'punctuation': '标点修正',
        'style': '句式检测',
        'ai': 'AI句式润色',
    }
    for ct, count in change_types.items():
        name = type_names.get(ct, ct)
        p = doc.add_paragraph(style='List Bullet')
        p.text = f'{name}：{count} 处'
    
    if total_changes > 0:
        doc.add_heading('3. 主要润色方向', level=2)
        
        direction_table = doc.add_table(rows=1, cols=2)
        direction_table.style = 'Table Grid'
        direction_table.rows[0].cells[0].text = '问题类型'
        direction_table.rows[0].cells[1].text = '修复举例'
        for cell in direction_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        
        seen_types = set()
        for c in changes[:10]:
            ct = c.get('type', 'unknown') if isinstance(c, dict) else getattr(c, 'type', 'unknown')
            if ct not in seen_types:
                seen_types.add(ct)
                before = c.get('before', '')[:40] if isinstance(c, dict) else getattr(c, 'before', '')[:40]
                after = c.get('after', '')[:40] if isinstance(c, dict) else getattr(c, 'after', '')[:40]
                row = direction_table.add_row()
                row.cells[0].text = type_names.get(ct, ct)
                row.cells[1].text = f'{before} → {after}'
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
    
    doc.add_heading('4. 完整改动清单', level=2)
    
    list_table = doc.add_table(rows=1, cols=4)
    list_table.style = 'Table Grid'
    headers = ['序号', '修改前', '修改后', '类型']
    for i, h in enumerate(headers):
        list_table.rows[0].cells[i].text = h
        for run in list_table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    
    for idx, c in enumerate(changes, 1):
        row = list_table.add_row()
        before = c.get('before', '')[:50] if isinstance(c, dict) else getattr(c, 'before', '')[:50]
        after = c.get('after', '')[:50] if isinstance(c, dict) else getattr(c, 'after', '')[:50]
        ct = c.get('type', '') if isinstance(c, dict) else getattr(c, 'type', '')
        
        row.cells[0].text = str(idx)
        row.cells[1].text = before
        row.cells[2].text = after
        row.cells[3].text = type_names.get(ct, ct)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_heading('5. 查看修订的方法', level=2)
    
    steps = [
        '在 Microsoft Word 或 WPS 桌面端打开【修订标记版】文件',
        '点击【审阅】选项卡',
        '确保【修订】按钮处于开启状态',
        '即可在文档中看到所有修改痕迹和批注',
        '如需接受所有修改：点击【接受】→【接受所有修订】',
    ]
    for step in steps:
        p = doc.add_paragraph(style='List Number')
        p.text = step
    
    p = doc.add_paragraph()
    run = p.add_run('注意：飞书在线预览不支持显示 OOXML 修订标记/批注，请下载到本地查看。')
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0)
    
    doc.save(report_path)


@router.post("/analyze-file")
async def analyze_file_endpoint(
    file: UploadFile = File(...),
    product_type: Optional[str] = Form(None),
    sentence_file_id: Optional[int] = Form(None),
    terminology_file_id: Optional[int] = Form(None),
    requirements: Optional[str] = Form(None),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user_optional)
):
    import tempfile
    import docx2txt
    
    task_id = str(uuid.uuid4())
    with _polish_tasks_lock:
        _polish_tasks[task_id] = {"status": "running", "progress": 0, "message": "开始润色..."}
    
    def _update_progress(pct: int, msg: str):
        with _polish_tasks_lock:
            if task_id in _polish_tasks:
                _polish_tasks[task_id] = {"status": "running", "progress": pct, "message": msg}
    
    def _finish_task(result=None, error=None):
        with _polish_tasks_lock:
            if error:
                _polish_tasks[task_id] = {"status": "error", "progress": 100, "message": str(error)}
            else:
                _polish_tasks[task_id] = {"status": "done", "progress": 100, "message": "润色完成", "result": result}
    
    user = current_user or get_default_user(db)
    temp_path = None
    original_temp_path = None
    try:
        filename = file.filename or "unnamed"
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else "txt"
        
        _update_progress(5, "读取文件中...")
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            content_bytes = await file.read()
            tmp.write(content_bytes)
            temp_path = tmp.name

        content = None
        output_docx_path = None
        is_docx = (ext == 'docx')
        
        if is_docx:    
            output_filename = f"【修订标记版】{filename}"
        else:
            output_filename = filename

        _update_progress(10, "解析文本内容...")
        if ext in ['txt', 'md', 'markdown']:
            content = _read_file_safe(temp_path)
        elif ext == 'docx':
            from docx import Document
            doc = Document(temp_path)
            content = '\n'.join([p.text for p in _iter_docx_paragraphs_in_order(doc)])
        
        if content is None or content.strip() == "":
            raise HTTPException(status_code=400, detail="无法提取文本内容")

        _update_progress(15, "加载润色规则...")
        sentence_guide = _build_document_polish_guide(
            db,
            sentence_file_id=sentence_file_id,
            requirements=requirements,
            product_type=product_type,
        )
        candidate_recall_guide = _candidate_recall_guide_text(sentence_guide or '')
        ai_style_guide = _compact_document_ai_style_guide(sentence_guide or '')
        sentence_file_name = _resolve_sentence_scope_name(db, sentence_file_id, product_type)
        
        terminology = _load_scoped_terminology_source(db, terminology_file_id, product_type)
        term_file_name = _resolve_terminology_scope_name(db, terminology_file_id, product_type)
        if terminology and term_file_name:
            print(f"[POLISH_LAB] 已加载术语范围: {term_file_name}")

        db_terms = None if terminology else _load_terms_from_db(db)
        skill_rules = {}

        # 规则引擎预处理：AI 润色前先做确定性术语替换 + 标点规范化
        _update_progress(18, "规则引擎预处理...")
        try:
            from app.utils.instrument_polisher import instrument_polish_engine
            pre_polished = instrument_polish_engine.pre_polish(content)
        except Exception as e:
            print(f"规则引擎预处理失败(使用原文): {e}")
            pre_polished = content

        resolved_terms = _resolve_terminology(db, terminology, pre_polished)
        resolved_terms = _trim_terms_for_ai(resolved_terms, pre_polished)

        _update_progress(20, "AI 智能润色中...")
        ai_polished = pre_polished
        ai_changes = []
        skip_ai, skip_reason = _should_skip_document_ai(pre_polished, ai_style_guide, resolved_terms)
        if skip_ai:
            print(f"[POLISH_LAB] 跳过整篇 AI 预润色: {skip_reason}")
            _update_progress(22, "文档较大，使用规则引擎润色...")
        else:
            # ---- 预匹配：先用句式库匹配，仅未匹配行送 AI ----
            pre_lines = pre_polished.split('\n')
            matched_indices = set()
            template_replacements = {}
            try:
                guide_entries = _preferred_entries_from_guide(candidate_recall_guide)
                guide_templates = [_template_entry_text(e) for e in guide_entries if _template_entry_text(e)]
                if guide_templates:
                    for idx, line in enumerate(pre_lines):
                        if not line.strip() or len(line.strip()) <= 8:
                            continue
                        cache_key = line.strip()
                        if cache_key in _polish_sentence_cache:
                            best_tmpl, best_score, best_level = _polish_sentence_cache[cache_key]
                        else:
                            best_tmpl, best_score, best_level = _three_tier_match(line, guide_templates)
                            _polish_sentence_cache[cache_key] = (best_tmpl, best_score, best_level)
                        if best_tmpl and best_score >= 0.85 and _template_replace_guard(line, best_tmpl, best_level, best_score):
                            template_replacements[idx] = best_tmpl
                            matched_indices.add(idx)
            except Exception as e:
                print(f"[POLISH_LAB] 预匹配阶段异常(跳过): {e}")

            unmatched_lines = [l for i, l in enumerate(pre_lines) if i not in matched_indices]
            unmatched_text = '\n'.join(unmatched_lines)
            match_count = len(matched_indices)
            total_lines = len([l for l in pre_lines if l.strip()])
            if match_count > 0:
                print(f"[POLISH_LAB] 句式预匹配: {match_count}/{total_lines} 行命中句式库，仅未匹配行送 AI")

            if unmatched_text.strip():
                try:
                    from app.utils.ai_client import ai_client
                    from app.utils.cat_diagnose import use_lab_ai_provider
                    with use_lab_ai_provider():
                        ai_result = ai_client.polish_text(unmatched_text, style_guide=ai_style_guide or None, terminology=resolved_terms if resolved_terms else None)
                    if ai_result and ai_result.get("polished"):
                        ai_lines_result = ai_result["polished"].split('\n')
                        merged_lines = list(pre_lines)
                        ai_idx = 0
                        for idx in range(len(merged_lines)):
                            if idx in matched_indices:
                                merged_lines[idx] = template_replacements[idx]
                            elif ai_idx < len(ai_lines_result):
                                merged_lines[idx] = ai_lines_result[ai_idx]
                                ai_idx += 1
                        ai_polished = '\n'.join(merged_lines)
                        ai_changes = ai_result.get("changes") or [{
                            "type": "ai",
                            "summary": f"AI 润色（{match_count} 行句式预匹配）"
                        }]
                except Exception as e:
                    print(f"AI 润色失败(返回原文): {e}")
            elif match_count > 0:
                merged_lines = list(pre_lines)
                for idx in matched_indices:
                    merged_lines[idx] = template_replacements[idx]
                ai_polished = '\n'.join(merged_lines)
                ai_changes = [{"type": "template", "summary": f"句式库匹配 {match_count} 行，跳过 AI"}]
                print(f"[POLISH_LAB] 全部 {match_count} 行命中句式库，跳过 AI 调用")

        logger.warning(
            "[POLISH_LAB_DEBUG] sentence_file_id=%s sentence_file_name=%r terminology_file_id=%s terminology_file_name=%r guide_chars=%s ai_skipped=%s ai_changed=%s",
            sentence_file_id,
            sentence_file_name,
            terminology_file_id,
            term_file_name,
            len(sentence_guide or ""),
            skip_ai,
            ai_polished != pre_polished,
        )

        # 规则引擎后置保护：检查 AI 润色是否丢失专有名词
        try:
            if ai_polished != pre_polished:
                protect_result = instrument_polish_engine.post_protect(pre_polished, ai_polished)
                if not protect_result.get("safe"):
                    ai_polished = protect_result.get("suggested", pre_polished)
                    print(f"[POLISH_LAB] 规则引擎保护触发: {protect_result.get('reason')}")
        except Exception:
            pass

        _update_progress(50, "应用修订标记...")
        if is_docx:
            previous_doc = _find_previous_polished_document(db, filename, getattr(user, 'id', None), 'docx')
            ai_lines = [l for l in ai_polished.split('\n') if l.strip()] if ai_polished != content else None
            _, date_str = _get_date_subfolder_id(db, None, user.id)
            date_dir = os.path.join(UPLOAD_DIR, date_str)
            if not os.path.exists(date_dir):
                os.makedirs(date_dir)
            
            unique_filename = f"【修订标记版】{filename}"
            saved_file_path = os.path.join(date_dir, unique_filename)
            
            _update_progress(60, "生成修订版 DOCX...")
            changes, review_items = _polish_docx_with_comments(
                temp_path, saved_file_path, skill_rules, db,
                sentence_guide, terminology, requirements,
                ai_lines=ai_lines,
                db_terminology=db_terms
            )
            import shutil
            shutil.copy2(saved_file_path, f"{saved_file_path}.all-revisions.docx")
            
            from docx import Document
            polished_doc = Document(saved_file_path)
            # 从修订标记版 DOCX 读取润色后文本（<w:delText> 被忽略，仅读 <w:t> 即插入文本）
            preview_text = '\n'.join([
                p.text for p in _iter_docx_paragraphs_in_order(polished_doc)
            ])
            polished_text = preview_text
            rejected_change_keys = _load_rejected_doc_change_keys(db)
            display_changes = _filter_visible_doc_changes(changes, rejected_change_keys, sentence_guide)
            if not display_changes:
                display_changes = _build_doc_change_details(content, polished_text, changes, rejected_change_keys, sentence_guide)
            previous_change_keys, previous_before_keys = _build_previous_doc_change_keys(previous_doc, rejected_change_keys)
            previous_new_change_count = _mark_new_doc_review_items(review_items, previous_change_keys, previous_before_keys) if previous_doc else 0
            if previous_doc:
                for item in display_changes:
                    before_key = _doc_change_before_key(item.get('before', ''))
                    pair_key = _doc_change_pair_key(item.get('before', ''), item.get('after', ''))
                    item['is_new_since_last_polish'] = bool(before_key) and before_key not in previous_before_keys and pair_key not in previous_change_keys
            debug_info = _build_polish_debug_info(
            sentence_file_id=sentence_file_id,
            sentence_file_name=sentence_file_name,
            sentence_guide=sentence_guide,
            candidate_recall_guide=candidate_recall_guide,
            ai_style_guide=ai_style_guide,
            terminology_file_id=terminology_file_id,
            terminology_file_name=term_file_name,
            skip_ai=skip_ai,
                skip_reason=skip_reason,
                ai_polished=ai_polished,
                pre_polished=pre_polished,
                total_changes=len(changes),
                visible_changes=len(display_changes),
                previous_polish_found=previous_doc is not None,
                previous_new_change_count=previous_new_change_count,
            )
            
            _update_progress(80, "生成润色报告...")
            report_filename = f"【润色报告】{filename.rsplit('.', 1)[0]}.docx"
            report_path = os.path.join(date_dir, report_filename)
            _generate_polish_report(
                report_path, filename, display_changes,
                sentence_file_name,
                term_file_name,
                requirements
            )
            
            _update_progress(90, "保存到知识库...")
            db_doc = create_polished_document(
                db=db,
                name=f"【修订标记版】{filename}",
                filename=unique_filename,
                file_path=saved_file_path,
                file_size=os.path.getsize(saved_file_path),
                file_type="docx",
                original_content=content,
                polished_content=polished_text,
                report_filename=report_filename,
                report_file_path=report_path,
                folder_id=None,
                created_by=user.id
            )
            
            result_data = {
                "task_id": task_id,
                "id": db_doc.id,
                "original": content,
                "polished": polished_text,
                "changes": display_changes,
                "review_items": review_items,
                "report_file": report_filename,
                "download_filename": unique_filename,
                "file_type": "docx",
                "debug_info": debug_info,
            }
            _finish_task(result_data)
            return result_data
        else:    
            _update_progress(70, "整理润色结果...")
            polished_text = _protect_model_numbers(ai_polished)
            changes = []
            if ai_changes and polished_text != content:
                changes.append(PolishRuleMatch(
                    rule_name="ai",
                    before=content[:50],
                    after=polished_text[:50],
                    type="ai"
                ))
            debug_info = _build_polish_debug_info(
                sentence_file_id=sentence_file_id,
                sentence_file_name=sentence_file_name,
                sentence_guide=sentence_guide,
                candidate_recall_guide=candidate_recall_guide,
                ai_style_guide=ai_style_guide,
                terminology_file_id=terminology_file_id,
                terminology_file_name=term_file_name,
                skip_ai=skip_ai,
                skip_reason=skip_reason,
                ai_polished=ai_polished,
                pre_polished=pre_polished,
                total_changes=len(changes),
            )
            _update_progress(90, "保存结果...")
            if not os.path.exists(UPLOAD_DIR):
                os.makedirs(UPLOAD_DIR)
            
            file_extension = f".{ext}" if ext else ".txt"
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            saved_file_path = os.path.join(UPLOAD_DIR, unique_filename)
            
            with open(saved_file_path, "w", encoding="utf-8") as f:
                f.write(polished_text)
            
            db_doc = create_polished_document(
                db=db,
                name=filename,
                filename=unique_filename,
                file_path=saved_file_path,
                file_size=len(content_bytes),
                file_type=ext,
                original_content=content,
                polished_content=polished_text,
                created_by=user.id
            )
            
            result_data = {
                "task_id": task_id,
                "id": db_doc.id,
                "original": content,
                "polished": polished_text,
                "changes": changes,
                "download_filename": unique_filename,
                "file_type": ext,
                "debug_info": debug_info,
            }
            _finish_task(result_data)
            return result_data

    except HTTPException:
        _finish_task(error="请求参数错误")
        raise
    except Exception as e:
        _finish_task(error=str(e))
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except:
                pass
        if original_temp_path and os.path.exists(original_temp_path):
            try:
                os.remove(original_temp_path)
            except:
                pass


@router.get("/progress/{task_id}")
async def get_polish_progress(task_id: str):
    """查询润色任务进度"""
    with _polish_tasks_lock:
        task = _polish_tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return task


# ============================================================
# 润色反馈：准确率评分 + 修正词自动入库
# ============================================================

@router.post("/feedback", response_model=None)
def submit_polish_feedback(
    feedback: FeedbackInput,
    db: Session = Depends(get_db)
):
    """提交润色反馈：记录准确率评分，并将修正词写入选中的术语文件或句式文件。"""
    current_user = get_default_user(db)
    corrections_pairs = _parse_corrections(feedback.corrections)
    raw_lines = [line.strip() for line in (feedback.corrections or '').splitlines() if line.strip()]
    correction_items = _build_feedback_correction_items(feedback, corrections_pairs, raw_lines)
    polish_session_id = str(uuid.uuid4())
    processed_count = 0
    errors = []

    if not raw_lines and feedback.accuracy >= 100:
        processed_count = 0

    elif feedback.target == "terminology":
        # 术语修正 → 固定写入平台反馈术语文件
        try:
            target_files = _get_platform_feedback_terminology_targets(db, current_user.id if current_user else 1)
            new_pairs = []
            for term_file in target_files:
                existing = ""
                if term_file.file_path and os.path.exists(term_file.file_path):
                    with open(term_file.file_path, 'r', encoding='utf-8') as f:
                        existing = f.read()

                file_new_pairs = []
                for old_term, new_term in corrections_pairs:
                    normalized_old = old_term.strip()
                    normalized_new = new_term.strip()
                    if (
                        f'| {normalized_old} | {normalized_new} |' in existing or
                        f'|{normalized_old}|{normalized_new}|' in existing
                    ):
                        continue
                    file_new_pairs.append((normalized_old, normalized_new))

                if not file_new_pairs:
                    term_file.file_size = os.path.getsize(term_file.file_path)
                    continue

                with open(term_file.file_path, 'a', encoding='utf-8') as f:
                    for old_term, new_term in file_new_pairs:
                        f.write(f'| {old_term} | {new_term} |\n')

                term_file.file_size = os.path.getsize(term_file.file_path)
                if not new_pairs:
                    new_pairs = file_new_pairs

            processed_count = len(new_pairs)
            db.commit()
        except Exception as e:
            errors.append(str(e))

    elif feedback.target == "sentence_guide":
        # 句式修正 → 固定写入平台反馈句式文件
        if not raw_lines:
            raise HTTPException(status_code=400, detail="请填写需修正的词语或句子")
        timestamp = datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        try:
            target_files = _get_platform_feedback_targets(db, current_user.id if current_user else 1)
            feedback_file_id = target_files[0].id if target_files else None
            new_lines = []

            for guide_file in target_files:
                existing = ""
                if guide_file.file_path and os.path.exists(guide_file.file_path):
                    with open(guide_file.file_path, 'r', encoding='utf-8') as f:
                        existing = f.read()

                file_new_lines = []
                for line in raw_lines:
                    if f"- {line}\n" in existing:
                        continue
                    file_new_lines.append(line)

                if not file_new_lines:
                    guide_file.file_size = os.path.getsize(guide_file.file_path)
                    continue

                with open(guide_file.file_path, 'a', encoding='utf-8') as f:
                    f.write(f'\n## 用户反馈修正 ({timestamp})\n\n')
                    for line in file_new_lines:
                        f.write(f'- {line}\n')
                    f.write('\n')

                guide_file.file_size = os.path.getsize(guide_file.file_path)
                if not new_lines:
                    new_lines = file_new_lines

            processed_count = len(new_lines)
            db.commit()
            if feedback_file_id is not None:
                _invalidate_sentence_guide_cache(feedback_file_id)
        except Exception as e:
            errors.append(str(e))

    db.add(PolishFeedback(
        original_text=feedback.original_text,
        polished_text=feedback.polished_text,
        accuracy=feedback.accuracy,
        corrections=feedback.corrections,
        target=feedback.target,
        processed_count=processed_count,
        correction_items=_json_dumps(correction_items),
        polish_session_id=polish_session_id,
        created_by=current_user.username if current_user else "guest"
    ))
    db.commit()

    return {
        "message": "反馈已提交",
        "accuracy": feedback.accuracy,
        "corrections_count": len(raw_lines) if feedback.target == "sentence_guide" else len(corrections_pairs),
        "processed_count": processed_count,
        "target": feedback.target,
        "errors": errors if errors else None
    }


@router.post("/feedback/document", response_model=None)
def submit_document_feedback(
    feedback: DocumentFeedbackInput,
    db: Session = Depends(get_db)
):
    """提交文档润色反馈，将接受项写入句式清单并同步修订版 Word。"""
    from sqlalchemy import func

    current_user = get_default_user(db)
    total_items = len(feedback.items or [])
    if total_items == 0:
        raise HTTPException(status_code=400, detail="当前没有可提交的润色结果")

    doc = get_polished_document(db, feedback.document_id) if feedback.document_id else None
    applied_changes = []
    if doc and doc.file_type == 'docx' and doc.file_path:
        import shutil

        source_revision_path = f"{doc.file_path}.all-revisions.docx"
        if not os.path.exists(source_revision_path) and os.path.exists(doc.file_path):
            shutil.copy2(doc.file_path, source_revision_path)
        applied_changes = _apply_feedback_to_revised_docx(source_revision_path, doc.file_path, feedback.items)
        doc.file_size = os.path.getsize(doc.file_path)
        doc.polished_content = '\n'.join(item['after'] for item in applied_changes)
        if doc.report_file_path:
            _generate_polish_report(
                doc.report_file_path,
                feedback.source_filename or doc.name or doc.filename,
                applied_changes,
            )

    accepted_items = [item for item in feedback.items if item.accepted and (item.after or '').strip()]
    rejected_items = [
        item for item in feedback.items
        if (item.status or '').strip() == 'rejected' and (item.before or '').strip() and (item.after or '').strip()
    ]
    accepted_lines = []
    seen_lines = set()
    for item in accepted_items:
        line = item.after.strip()
        if line in seen_lines:
            continue
        seen_lines.add(line)
        accepted_lines.append(line)

    processed_count = 0
    feedback_file_id = None
    source_name = (feedback.source_filename or '').strip() or (doc.name if doc and doc.name else '') or (doc.filename if doc and doc.filename else '') or f"文档{feedback.document_id or ''}".strip()
    feedback_record_key = _build_document_feedback_record_key(doc.id if doc else feedback.document_id, source_name)
    if accepted_lines:
        target_files = _get_platform_feedback_targets(db, current_user.id if current_user else 1)
        feedback_file_id = target_files[0].id if target_files else None
        timestamp = datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        new_lines = []
        for file in target_files:
            existing = ""
            if file.file_path and os.path.exists(file.file_path):
                with open(file.file_path, 'r', encoding='utf-8') as f:
                    existing = f.read()

            file_new_lines = []
            for line in accepted_lines:
                if f"- {line}\n" in existing:
                    continue
                file_new_lines.append(line)

            if not file_new_lines:
                file.file_size = os.path.getsize(file.file_path)
                continue

            with open(file.file_path, 'a', encoding='utf-8') as f:
                f.write(f"\n## 用户反馈修正 ({timestamp} / 来源：{source_name})\n\n")
                for line in file_new_lines:
                    f.write(f"- {line}\n")
                f.write("\n")

            file.file_size = os.path.getsize(file.file_path)
            if not new_lines:
                new_lines = file_new_lines

        processed_count = len(new_lines)
        db.commit()
        if feedback_file_id is not None:
            _invalidate_sentence_guide_cache(feedback_file_id)

    rejected_count = 0
    if rejected_items:
        existing_rejected = _load_rejected_doc_change_keys(db)
        for item in rejected_items:
            key = (_normalize_compare_text(item.before), _normalize_compare_text(item.after))
            if key in existing_rejected:
                continue
            db.add(PolishFeedback(
                original_text=item.before.strip(),
                polished_text=item.after.strip(),
                accuracy=0,
                corrections=item.type or '',
                target='document_rejected_change',
                processed_count=1,
                created_by=current_user.username if current_user else 'guest'
            ))
            existing_rejected.add(key)
            rejected_count += 1

    db.add(PolishFeedback(
        original_text=feedback_record_key,
        polished_text='\n'.join(accepted_lines),
        accuracy=len(accepted_items),
        corrections='\n'.join(accepted_lines),
        target='document_sentence_guide',
        processed_count=total_items,
        created_by=current_user.username if current_user else 'guest'
    ))
    db.commit()

    all_document_feedback = db.query(PolishFeedback).filter(
        PolishFeedback.target == 'document_sentence_guide'
    ).all()
    total_docs = len(_latest_document_feedback_records(all_document_feedback))

    return {
        "message": "文档润色反馈已提交",
        "document_id": doc.id if doc else feedback.document_id,
        "raw_url": f"/api/polish-lab/{doc.id}/raw" if doc else None,
        "processed_count": processed_count,
        "accepted_count": len(accepted_lines),
        "rejected_count": rejected_count,
        "applied_changes": applied_changes,
        "total_count": total_items,
        "feedback_file_id": feedback_file_id,
        "total_docs": total_docs
    }


def _cat_save_rejected_as_learning(
    decisions: list,
    db: Session,
    current_user=None,
) -> int:
    """
    拒绝项 → 存入 PolishFeedback 表（target='document_rejected_change'），
    下次润色时 _load_rejected_doc_change_keys() 会自动跳过。
    """
    existing_keys = _load_rejected_doc_change_keys(db)
    count = 0
    for d in decisions:
        action = d.action if hasattr(d, 'action') else d.get('action', '')
        if action != "reject":
            continue
        original = d.original_text if hasattr(d, 'original_text') else d.get('original_text', '')
        rejected = d.rejected_template if hasattr(d, 'rejected_template') else d.get('rejected_template', '')
        if not original or not rejected:
            continue
        key = (_normalize_compare_text(original), _normalize_compare_text(rejected))
        if key in existing_keys:
            continue
        db.add(PolishFeedback(
            original_text=original.strip(),
            polished_text=rejected.strip(),
            accuracy=0,
            corrections="cat_rejected",
            target='document_rejected_change',
            processed_count=1,
            created_by=(current_user.username if current_user else 'guest'),
        ))
        existing_keys.add(key)
        count += 1
    if count:
        db.commit()
    return count


def _cat_save_modified_to_feedback(
    decisions: list,
    db: Session,
    current_user=None,
    source_filename: str = "",
) -> int:
    """
    手动修改项 → 用户修改后的文本写入"平台反馈的句式清单.md"。
    """
    modified_lines = []
    seen = set()
    for d in decisions:
        action = d.action if hasattr(d, 'action') else d.get('action', '')
        if action != "modify":
            continue
        modified = d.modified_text if hasattr(d, 'modified_text') else d.get('modified_text', '')
        if not modified or not modified.strip():
            continue
        line = modified.strip()
        if line in seen:
            continue
        seen.add(line)
        modified_lines.append(line)

    if not modified_lines:
        return 0

    primary_file = _ensure_platform_feedback_sentence_file(
        db, current_user.id if current_user else 1
    )
    if not primary_file:
        return 0

    timestamp = datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
    source_name = source_filename or "润色结果"
    new_count = 0

    existing = ""
    if primary_file.file_path and os.path.exists(primary_file.file_path):
        with open(primary_file.file_path, 'r', encoding='utf-8') as f:
            existing = f.read()

    file_new = []
    for line in modified_lines:
        if f"- {line}\n" in existing:
            continue
        file_new.append(line)

    if not file_new:
        primary_file.file_size = os.path.getsize(primary_file.file_path) if os.path.exists(primary_file.file_path) else 0
        db.commit()
        db.refresh(primary_file)
        _invalidate_sentence_guide_cache(primary_file.id)
        return 0

    with open(primary_file.file_path, 'a', encoding='utf-8') as f:
        f.write(f"\n## 用户修改反馈 ({timestamp} / 来源：{source_name})\n\n")
        for line in file_new:
            f.write(f"- {line}\n")
        f.write("\n")

    primary_file.file_size = os.path.getsize(primary_file.file_path)
    new_count = len(file_new)

    db.commit()
    db.refresh(primary_file)
    _invalidate_sentence_guide_cache(primary_file.id)
    return new_count


def _cat_calc_file_accuracy(decisions: list, total_paragraphs: int) -> dict:
    """
    单文件润色准确率统计。
    """
    accepted = 0
    rejected = 0
    modified = 0
    has_candidates = 0

    for d in decisions:
        action = d.action if hasattr(d, 'action') else d.get('action', '')
        has_candidates += 1
        if action == "accept":
            accepted += 1
        elif action == "reject":
            rejected += 1
        elif action == "modify":
            modified += 1

    decided = accepted + rejected + modified
    pending = max(0, has_candidates - decided)
    no_match = max(0, total_paragraphs - has_candidates)

    def _pct(n, d):
        return round(n / d * 100, 1) if d > 0 else None

    return {
        "total_items": has_candidates,
        "accepted": accepted,
        "rejected": rejected,
        "modified": modified,
        "pending": pending,
        "no_match": no_match,
        "accuracy_rate": _pct(accepted, decided),
        "rejection_rate": _pct(rejected, decided),
        "modification_rate": _pct(modified, decided),
        "template_coverage": _pct(has_candidates, total_paragraphs),
    }


@router.get("/feedback/stats", response_model=None)
def get_feedback_stats(db: Session = Depends(get_db)):
    """获取润色准确率统计：每条反馈的接受率（accuracy/processed_count*100）取平均。"""
    records = db.query(PolishFeedback).all()
    total = len(records)
    if total == 0:
        return {"total_count": 0, "average_accuracy": 0}
    ratios = []
    for r in records:
        denom = r.processed_count if r.processed_count and r.processed_count > 0 else 1
        ratios.append(min(100.0, (r.accuracy or 0) / denom * 100))
    avg = sum(ratios) / total
    return {
        "total_count": total,
        "average_accuracy": round(avg, 1)
    }


@router.get("/feedback/document-stats", response_model=None)
def get_document_feedback_stats(db: Session = Depends(get_db)):
    """获取文档润色页统计。每个文档仅按最新一次反馈计 1 次提交。"""
    records = db.query(PolishFeedback).filter(
        PolishFeedback.target == 'document_sentence_guide'
    ).all()
    return _document_feedback_stats(records)


@router.get("/stats/text", response_model=None)
def get_polish_text_stats(db: Session = Depends(get_db)):
    records = db.query(PolishFeedback).filter(
        PolishFeedback.target.in_(['terminology', 'sentence_guide'])
    ).order_by(PolishFeedback.created_at.asc(), PolishFeedback.id.asc()).all()
    if not records:
        return {
            "overview": {"total_sessions": 0, "average_accuracy": 0, "total_corrections": 0, "total_accepted": 0},
            "accuracy_trend": [],
            "correction_category_pie": [],
            "category_trend": [],
            "recent_sessions": []
        }

    category_counts = defaultdict(int)
    trend_map = defaultdict(lambda: defaultdict(int))
    trend_dates = []
    recent_sessions = []
    accuracies = []
    total_corrections = 0
    total_accepted = 0

    for record in records:
        items = _json_loads(record.correction_items, [])
        if not items:
            if record.target == 'terminology':
                items = [
                    {"before": old, "after": new, "category": _cat_report_change_category(old, new, 'modify')}
                    for old, new in _parse_corrections(record.corrections)
                ]
            else:
                raw_lines = [line.strip() for line in (record.corrections or '').splitlines() if line.strip()]
                baseline = str(record.polished_text or record.original_text or '').strip()
                items = [
                    {"before": baseline, "after": line, "category": _cat_report_change_category(baseline, line, 'modify')}
                    for line in raw_lines
                ]

        accuracies.append(_feedback_accuracy_percent(record))
        total_corrections += len(items)
        total_accepted += int(record.processed_count or 0)
        day = _normalize_stat_date(record.created_at)
        trend_dates.append(day)
        item_category_counts = defaultdict(int)
        for item in items:
            category = str((item or {}).get('category') or '其他').strip() or '其他'
            category_counts[category] += 1
            trend_map[day][category] += 1
            item_category_counts[category] += 1

        recent_sessions.append({
            "session_id": record.id,
            "created_at": record.created_at.isoformat() if record.created_at else '',
            "accuracy": round(_feedback_accuracy_percent(record), 1),
            "accepted_count": int(record.processed_count or 0),
            "correction_count": len(items),
            "corrections_by_category": dict(item_category_counts),
        })

    top_categories = sorted(category_counts.items(), key=lambda item: (-item[1], item[0]))
    correction_category_pie = [
        {"category": category, "count": count, "percentage": round(count / total_corrections * 100, 1) if total_corrections else 0}
        for category, count in top_categories
    ]
    unique_days = sorted(set(trend_dates))[-14:]
    category_trend = []
    for day in unique_days:
        row = {"date": day}
        for category in [item[0] for item in top_categories[:6]]:
            row[category] = int(trend_map[day].get(category, 0))
        category_trend.append(row)

    return {
        "overview": {
            "total_sessions": len(records),
            "average_accuracy": round(sum(accuracies) / len(accuracies), 1),
            "total_corrections": total_corrections,
            "total_accepted": total_accepted,
        },
        "accuracy_trend": [
            {
                "session_id": record.id,
                "created_at": record.created_at.isoformat() if record.created_at else '',
                "accuracy": round(_feedback_accuracy_percent(record), 1),
            }
            for record in records[-30:]
        ],
        "correction_category_pie": correction_category_pie,
        "category_trend": category_trend,
        "recent_sessions": list(reversed(recent_sessions[-20:])),
    }


@router.get("/stats/document", response_model=None)
def get_polish_document_stats(db: Session = Depends(get_db)):
    sessions = db.query(CatAnalysisSession).order_by(CatAnalysisSession.created_at.asc(), CatAnalysisSession.id.asc()).all()
    if not sessions:
        return {
            "overview": {"total_documents": 0, "total_sessions": 0, "average_accuracy": 0, "total_accepted": 0, "total_modified": 0, "total_rejected": 0, "total_pending": 0},
            "accuracy_trend": [],
            "decision_pie": [],
            "category_distribution": [],
            "guide_comparison": [],
            "documents": [],
        }

    latest_by_doc = {}
    document_groups = defaultdict(list)
    guide_groups = defaultdict(list)
    decision_totals = defaultdict(int)
    category_totals = defaultdict(int)
    accuracy_values = []
    total_accepted = total_modified = total_rejected = total_pending = 0

    for session in sessions:
        doc_key = session.source_filename or f'文档{session.id}'
        document_groups[doc_key].append(session)
        guide_key = session.sentence_file_name or '未指定句式库'
        guide_groups[guide_key].append(session)
        latest_by_doc[doc_key] = max(latest_by_doc.get(doc_key, session), session, key=lambda item: (item.created_at or datetime.datetime.min, item.id or 0))
        if session.accuracy_rate is not None:
            accuracy_values.append(float(session.accuracy_rate or 0))
        total_accepted += int(session.accepted or 0)
        total_modified += int(session.modified or 0)
        total_rejected += int(session.rejected or 0)
        total_pending += int(session.pending or 0)
        decision_totals['接受'] += int(session.accepted or 0)
        decision_totals['自定义'] += int(session.modified or 0)
        decision_totals['拒绝'] += int(session.rejected or 0)
        decision_totals['待处理'] += int(session.pending or 0)
        for category, count in _json_loads(session.category_summary, {}).items():
            category_totals[str(category)] += int(count or 0)

    accuracy_trend = [
        {
            "analyze_id": session.analyze_id,
            "source_filename": session.source_filename,
            "created_at": session.created_at.isoformat() if session.created_at else '',
            "accuracy_rate": float(session.accuracy_rate or 0),
        }
        for session in sessions[-30:]
    ]
    decision_pie = [
        {"action": action, "count": count, "percentage": round(count / max(1, total_accepted + total_modified + total_rejected + total_pending) * 100, 1)}
        for action, count in decision_totals.items()
    ]
    category_distribution = [
        {"category": category, "count": count, "percentage": round(count / max(1, sum(category_totals.values())) * 100, 1)}
        for category, count in sorted(category_totals.items(), key=lambda item: (-item[1], item[0]))
    ]
    guide_comparison = []
    for guide_name, guide_sessions in sorted(guide_groups.items(), key=lambda item: (-sum(float(s.accuracy_rate or 0) for s in item[1]) / max(1, len(item[1])), item[0])):
        doc_names = sorted({session.source_filename for session in guide_sessions if session.source_filename})
        avg_accuracy = round(sum(float(s.accuracy_rate or 0) for s in guide_sessions) / len(guide_sessions), 1)
        guide_comparison.append({
            "sentence_file_name": guide_name if guide_name != '未指定句式库' else None,
            "session_count": len(guide_sessions),
            "avg_accuracy": avg_accuracy,
            "documents": doc_names,
        })

    documents = []
    for doc_name, doc_sessions in sorted(document_groups.items(), key=lambda item: item[0]):
        sorted_sessions = sorted(doc_sessions, key=lambda item: (item.created_at or datetime.datetime.min, item.id or 0))
        latest_session = sorted_sessions[-1]
        documents.append({
            "source_filename": doc_name,
            "session_count": len(doc_sessions),
            "latest_accuracy": round(float(latest_session.accuracy_rate or 0), 1),
            "average_accuracy": round(sum(float(s.accuracy_rate or 0) for s in doc_sessions) / len(doc_sessions), 1),
            "sentence_file_name": latest_session.sentence_file_name,
            "latest_session": {
                "analyze_id": latest_session.analyze_id,
                "created_at": latest_session.created_at.isoformat() if latest_session.created_at else '',
                "accepted": int(latest_session.accepted or 0),
                "modified": int(latest_session.modified or 0),
                "rejected": int(latest_session.rejected or 0),
                "pending": int(latest_session.pending or 0),
                "accuracy_rate": round(float(latest_session.accuracy_rate or 0), 1),
                "category_summary": _json_loads(latest_session.category_summary, {}),
            },
        })

    return {
        "overview": {
            "total_documents": len(document_groups),
            "total_sessions": len(sessions),
            "average_accuracy": round(sum(accuracy_values) / len(accuracy_values), 1) if accuracy_values else 0,
            "total_accepted": total_accepted,
            "total_modified": total_modified,
            "total_rejected": total_rejected,
            "total_pending": total_pending,
        },
        "accuracy_trend": accuracy_trend,
        "decision_pie": decision_pie,
        "category_distribution": category_distribution,
        "guide_comparison": guide_comparison,
        "documents": documents,
    }


@router.get("/stats/text/sessions", response_model=None)
def get_polish_text_session_details(
    user_name: str = Query(None, description="按提交人筛选"),
    start_date: str = Query(None, description="起始日期 YYYY-MM-DD"),
    end_date: str = Query(None, description="结束日期 YYYY-MM-DD"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
    db: Session = Depends(get_db),
    _: object = Depends(require_admin),
):
    query = db.query(PolishFeedback).filter(
        PolishFeedback.target.in_(['terminology', 'sentence_guide'])
    )
    if user_name:
        query = query.filter(PolishFeedback.created_by.contains(user_name.strip()))
    query = _apply_created_at_filters(query, PolishFeedback, start_date, end_date)

    total = query.count()
    records = query.order_by(PolishFeedback.created_at.desc(), PolishFeedback.id.desc()).offset((page - 1) * page_size).limit(page_size).all()
    items = []
    for record in records:
        has_corrections = bool(str(record.corrections or '').strip()) or bool(_json_loads(record.correction_items, []))
        items.append({
            "id": record.id,
            "created_by": record.created_by or 'guest',
            "submitted_text": str(record.original_text or '').strip(),
            "polished_text": str(record.polished_text or '').strip(),
            "match_rate": round(_session_match_rate_percent(record.original_text, record.polished_text), 1),
            "accuracy": round(_feedback_accuracy_percent(record), 1),
            "has_corrections": has_corrections,
            "created_at": record.created_at.isoformat() if record.created_at else '',
        })

    return {
        "items": items,
        "total": total,
        "page": page,
        "page_size": page_size,
    }


@router.get("/stats/document/sessions", response_model=None)
def get_polish_document_session_details(
    user_name: str = Query(None, description="按提交人筛选"),
    start_date: str = Query(None, description="起始日期 YYYY-MM-DD"),
    end_date: str = Query(None, description="结束日期 YYYY-MM-DD"),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=200),
    db: Session = Depends(get_db),
    _: object = Depends(require_admin),
):
    query = db.query(CatAnalysisSession)
    if user_name:
        query = query.filter(CatAnalysisSession.created_by.contains(user_name.strip()))
    query = _apply_created_at_filters(query, CatAnalysisSession, start_date, end_date)

    total = query.count()
    sessions = query.order_by(CatAnalysisSession.created_at.desc(), CatAnalysisSession.id.desc()).offset((page - 1) * page_size).limit(page_size).all()

    feedback_lookup = _build_document_feedback_lookup_for_source_names(
        db,
        [session.source_filename or '' for session in sessions],
    )

    items = []
    for session in sessions:
        feedback_record = feedback_lookup.get(_normalize_document_feedback_key(session.source_filename or ''))
        has_corrections = bool(str((feedback_record.corrections if feedback_record else '') or '').strip())
        accepted = int(session.accepted or 0)
        rejected = int(session.rejected or 0)
        modified = int(session.modified or 0)
        items.append({
            "id": session.id,
            "analyze_id": session.analyze_id,
            "created_by": session.created_by or 'guest',
            "source_filename": session.source_filename or '',
            "accepted": accepted,
            "rejected": rejected,
            "modified": modified,
            "total_polished_count": int(session.total_items or 0),
            "accuracy": round(float(session.accuracy_rate or 0), 1) if session.accuracy_rate is not None else 0,
            "has_corrections": has_corrections,
            "created_at": session.created_at.isoformat() if session.created_at else '',
        })

    return {
        "items": items,
        "total": total,
        "page": page,
        "page_size": page_size,
    }


@router.get("/stats/document/{analyze_id}", response_model=None)
def get_polish_document_stats_detail(analyze_id: str, db: Session = Depends(get_db)):
    session_row = db.query(CatAnalysisSession).filter(CatAnalysisSession.analyze_id == analyze_id).first()
    if not session_row:
        raise HTTPException(status_code=404, detail="会话不存在")
    decision_rows = db.query(CatDecisionRecord).filter(CatDecisionRecord.analyze_id == analyze_id).order_by(CatDecisionRecord.sentence_index.asc(), CatDecisionRecord.id.asc()).all()
    return {
        "session": {
            "analyze_id": session_row.analyze_id,
            "source_filename": session_row.source_filename,
            "sentence_file_name": session_row.sentence_file_name,
            "created_at": session_row.created_at.isoformat() if session_row.created_at else '',
            "accuracy_rate": round(float(session_row.accuracy_rate or 0), 1) if session_row.accuracy_rate is not None else None,
            "accepted": int(session_row.accepted or 0),
            "modified": int(session_row.modified or 0),
            "rejected": int(session_row.rejected or 0),
            "pending": int(session_row.pending or 0),
            "category_summary": _json_loads(session_row.category_summary, {}),
        },
        "decisions": [
            {
                "sentence_index": row.sentence_index,
                "paragraph_index": row.paragraph_index,
                "action": row.action,
                "original_text": row.original_text,
                "replacement_text": row.replacement_text,
                "category": row.category,
                "string_score": row.string_score,
                "semantic_score": row.semantic_score,
            }
            for row in decision_rows
        ],
    }


# ============================================================
# 文档润色端点（历史文档 / 种子导出）
# ============================================================

@router.post("/export-seed")
def export_polished_seed(db: Session = Depends(get_db)):
    """将已润色文档导出到种子目录，用于 Git 团队共享"""
    current_user = get_default_user(db)
    try:
        from seed.polished_seed import export_polished_to_seed
        export_polished_to_seed()
        return {"message": "已润色文档已导出到 seed/polished/ 目录，请执行 git add/commit/push 分享给团队"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/{document_id}")
async def polish_document(document_id: int, db: Session = Depends(get_db)):
    document = get_document(db, document_id=document_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    try:
        from app.utils.ai_client import ai_client
        from app.utils.cat_diagnose import use_lab_ai_provider
        terminology = _resolve_terminology(db, text=document.content or "")
        with use_lab_ai_provider():
            result = ai_client.polish_text(document.content or "", terminology=terminology if terminology else None, request_label="polish.document.quick")
        polished = _protect_model_numbers(result.get("polished", document.content or ""))
        changes = []
        if polished != (document.content or ""):
            changes.append({"line": 1, "original": (document.content or "")[:80], "polished": polished[:80], "type": "ai"})
        return {
            "document_id": document_id,
            "original": result.get("original", document.content or ""),
            "polished": polished,
            "changes": changes
        }
    except Exception:
        return {
            "document_id": document_id,
            "original": document.content or "",
            "polished": document.content or "",
            "changes": []
        }


def _polish_fallback(text: str, db_terminology: dict = None):
    polished, changes = _apply_skill_polish(text, {}, None, None, None, None, db_terminology=db_terminology)
    return {
        "original": text,
        "polished": polished,
        "changes": changes
    }



# ============================================================
# 已润色文档 CRUD
# ============================================================

@router.post("/upload")
async def upload_polished_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    if not os.path.exists(UPLOAD_DIR):
        os.makedirs(UPLOAD_DIR)
    
    file_extension = os.path.splitext(file.filename)[1] if file.filename else ""
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    content = await file.read()
    with open(file_path, "wb") as buffer:
        buffer.write(content)
    
    file_size = os.path.getsize(file_path)
    file_type = file_extension[1:] if file_extension else "unknown"
    
    # Try to read content for text-based files
    original_content = None
    polished_content = None
    if file_type in ["txt", "md", "docx"]:
        try:
            if file_type == "docx":
                import docx2txt
                original_content = docx2txt.process(file_path)
            else:
                with open(file_path, "r", encoding="utf-8") as f:
                    original_content = f.read()
        except:
            pass
    
    db_file = create_polished_document(
        db=db,
        name=file.filename or "unknown",
        filename=unique_filename,
        file_path=file_path,
        file_size=file_size,
        file_type=file_type,
        original_content=original_content,
        polished_content=polished_content,
        created_by=user.id
    )
    
    return {"message": "文件上传成功", "id": db_file.id}


class DiagnoseImportBody(BaseModel):
    match_pattern: Optional[str] = None
    replacement_text: Optional[str] = None
    rule_name: Optional[str] = None
    rule_type: Optional[str] = None


_DIAGNOSE_RULE_TYPES = {
    "term": "replacement_rule",
    "spelling": "format_rule",
    "word": "replacement_rule",
    "register": "imperative_rule",
    "syntax": "sentence_applicability_rule",
    "grammar": "format_rule",
}


def _serialize_diagnose_record(row) -> dict:
    return {
        "id": row.id,
        "analyze_id": row.analyze_id,
        "source": row.source,
        "source_name": getattr(row, "source_name", None) or "",
        "sentence_index": row.sentence_index,
        "original_text": row.original_text,
        "quote": row.quote,
        "category": row.category,
        "severity": row.severity,
        "problem": row.problem,
        "revised": row.revised,
        "rationale": row.rationale,
        "ruleable": bool(row.ruleable),
        "rule_hint": row.rule_hint,
        "status": row.status,
        "imported_rule_id": row.imported_rule_id,
        "created_at": row.created_at.isoformat() if row.created_at else None,
    }


@router.get("/diagnose-candidates")
def list_diagnose_candidates(
    skip: int = Query(0, ge=0),
    limit: int = Query(50, ge=1, le=200),
    ruleable: Optional[bool] = Query(None),
    status: Optional[str] = Query("pending"),
    db: Session = Depends(get_db),
):
    from app.models.cat_diagnose_record_lab import CatDiagnoseRecordLab

    q = db.query(CatDiagnoseRecordLab)
    if ruleable is not None:
        if ruleable:
            from sqlalchemy import or_
            q = q.filter(or_(
                CatDiagnoseRecordLab.ruleable.is_(True),
                (CatDiagnoseRecordLab.revised.isnot(None) & (CatDiagnoseRecordLab.revised != "")),
            ))
        else:
            q = q.filter(CatDiagnoseRecordLab.ruleable.is_(False))
    if status:
        q = q.filter(CatDiagnoseRecordLab.status == status)
    total = q.count()
    rows = q.order_by(CatDiagnoseRecordLab.id.desc()).offset(skip).limit(limit).all()
    _fill_diagnose_source_names(db, rows)
    return {"total": total, "items": [_serialize_diagnose_record(row) for row in rows]}


@router.post("/diagnose-candidates/{record_id}/import")
def import_diagnose_candidate(
    record_id: int,
    body: Optional[DiagnoseImportBody] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user_optional),
):
    from app.models.cat_diagnose_record_lab import CatDiagnoseRecordLab
    from app.crud.polish_learning_rule import create_rule, get_rule_by_key
    from app.schemas.polish_learning_rule import PolishLearningRuleCreate

    row = db.query(CatDiagnoseRecordLab).filter(CatDiagnoseRecordLab.id == record_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="诊断候补不存在")
    payload = body or DiagnoseImportBody()
    match_pattern = (payload.match_pattern or row.rule_hint or row.quote or "").strip()
    if match_pattern:
        try:
            re.compile(match_pattern)
        except re.error:
            match_pattern = re.escape(row.quote or row.original_text or match_pattern)
    else:
        match_pattern = re.escape(row.quote or row.original_text or "")
    if not match_pattern:
        raise HTTPException(status_code=400, detail="缺少可导入的匹配模式")
    from app.utils.cat_diagnose import hint_import_requires_replacement

    provided_replacement = payload.replacement_text if payload.replacement_text is not None else None
    if hint_import_requires_replacement(row.category, row.revised):
        replacement_text = str(provided_replacement or "").strip()
        if not replacement_text:
            raise HTTPException(
                status_code=400,
                detail="logic/missing/ambiguity 诊断 revised 为空，请手填 replacement_text 后再导入",
            )
    else:
        replacement_text = provided_replacement if provided_replacement is not None else (row.revised or "")
    rule_name = (payload.rule_name or row.problem or "AI 诊断规则").strip()[:128]
    rule_type = payload.rule_type or _DIAGNOSE_RULE_TYPES.get(row.category or "", "format_rule")
    rule_key = f"lab-diag-{row.id}-{uuid.uuid4().hex[:8]}"
    while get_rule_by_key(db, rule_key):
        rule_key = f"lab-diag-{row.id}-{uuid.uuid4().hex[:8]}"
    created = create_rule(
        db,
        PolishLearningRuleCreate(
            rule_name=rule_name,
            rule_type=rule_type,
            engine_key=None,
            rule_key=rule_key,
            match_pattern=match_pattern,
            replacement_text=replacement_text,
            description=(row.rationale or row.problem or "").strip() or None,
            priority_level=0,
            enabled=True,
        ),
    )
    row.status = "imported"
    row.imported_rule_id = created.id
    db.commit()
    db.refresh(row)
    _invalidate_typo_cache()
    return {
        "ok": True,
        "rule_id": created.id,
        "rule_key": created.rule_key,
        "item": _serialize_diagnose_record(row),
    }


@router.post("/diagnose-candidates/{record_id}/dismiss")
def dismiss_diagnose_candidate(
    record_id: int,
    db: Session = Depends(get_db),
):
    from app.models.cat_diagnose_record_lab import CatDiagnoseRecordLab

    row = db.query(CatDiagnoseRecordLab).filter(CatDiagnoseRecordLab.id == record_id).first()
    if not row:
        raise HTTPException(status_code=404, detail="诊断候补不存在")
    row.status = "dismissed"
    db.commit()
    db.refresh(row)
    return {"ok": True, "item": _serialize_diagnose_record(row)}


@router.get("/")
async def list_polished_documents(db: Session = Depends(get_db)):
    docs = get_polished_documents(db)
    result = []
    for d in docs:
        created_at_str = None
        if d.created_at:
            created_at_str = (d.created_at + timedelta(hours=8)).strftime("%Y/%m/%d %H:%M:%S")
        result.append({
            "id": d.id,
            "name": d.name,
            "filename": d.filename,
            "file_size": d.file_size,
            "file_type": d.file_type,
            "created_at": created_at_str,
            "has_polished_content": d.polished_content is not None,
            "report_file_path": d.report_file_path or None
        })
    return result


@router.get("/{doc_id}")
async def get_polished_document_info(doc_id: int, db: Session = Depends(get_db)):
    doc = get_polished_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return {
        "id": doc.id,
        "name": doc.name,
        "filename": doc.filename,
        "file_path": doc.file_path,
        "file_size": doc.file_size,
        "file_type": doc.file_type,
        "original_content": doc.original_content,
        "polished_content": doc.polished_content,
        "created_at": (doc.created_at + timedelta(hours=8)).isoformat() if doc.created_at else None
    }


@router.get("/{doc_id}/download")
async def download_polished_file(doc_id: int, db: Session = Depends(get_db)):
    doc = get_polished_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")
    
    return FileResponse(
        path=doc.file_path,
        filename=doc.filename,
        media_type="application/octet-stream"
    )


@router.get("/{doc_id}/download-report")
async def download_polished_file_report(doc_id: int, db: Session = Depends(get_db)):
    doc = get_polished_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if not doc.report_file_path:
        raise HTTPException(status_code=404, detail="润色报告不存在")
    
    if not os.path.exists(doc.report_file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")
    
    return FileResponse(
        path=doc.report_file_path,
        filename=doc.report_filename or "润色报告.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/{doc_id}/preview")
async def preview_polished_file(doc_id: int, db: Session = Depends(get_db)):
    doc = get_polished_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    # 文件在磁盘上不存在时，回退到 DB 中已保存的文字内容
    if not doc.file_path or not os.path.exists(doc.file_path):
        fallback_content = doc.polished_content or doc.original_content or ""
        if fallback_content:
            return {
                "content": fallback_content,
                "type": "text",
                "file_name": doc.filename,
                "polished_content": doc.polished_content,
                "fallback": True
            }
        else:
            raise HTTPException(status_code=404, detail="文件内容不可用")
    
    file_type = doc.file_type.lower()
    
    # Text-based files
    if file_type in ["txt", "md", "markdown", "json", "xml", "html", "css", "js", "py"]:
        with open(doc.file_path, "r", encoding="utf-8") as f:
            content = f.read()
        return {
            "content": content,
            "type": "text",
            "file_name": doc.filename,
            "polished_content": doc.polished_content
        }
    
    # Images
    elif file_type in ["jpg", "jpeg", "png", "gif", "bmp", "svg", "webp"]:
        return {
            "file_path": f"/api/polish-lab/{doc_id}/raw",
            "type": "image",
            "file_name": doc.filename
        }
    
    # PDF - extract text
    elif file_type == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(doc.file_path)
            content = "\n".join([page.extract_text() or "" for page in reader.pages])
            return {"content": content, "type": "text", "file_name": doc.filename}
        except Exception:
            fallback = doc.polished_content or doc.original_content or ""
            return {"content": fallback, "type": "text", "file_name": doc.filename, "fallback": True}
    
    # DOCX - extract text
    elif file_type == "docx":
        try:
            content = docx2txt.process(doc.file_path)
            return {
                "content": content,
                "type": "text",
                "file_name": doc.filename,
                "polished_content": doc.polished_content
            }
        except Exception:
            fallback = doc.polished_content or doc.original_content or ""
            return {"content": fallback, "type": "text", "file_name": doc.filename, "fallback": True}
    
    else:
        return {"content": "此文件类型不支持在线预览，请下载后查看", "type": "unsupported", "file_name": doc.filename}


@router.get("/{doc_id}/raw")
async def get_raw_polished_file(doc_id: int, db: Session = Depends(get_db)):
    doc = get_polished_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if not os.path.exists(doc.file_path):
        raise HTTPException(status_code=404, detail="服务器文件不存在")
    
    media_type = mimetypes.guess_type(doc.file_path)[0] or "application/octet-stream"
    
    return FileResponse(
        path=doc.file_path,
        filename=doc.filename,
        media_type=media_type
    )


@router.delete("/{doc_id}")
async def delete_polished_document_endpoint(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除文件")
    
    doc = get_polished_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    
    # 同时删除关联的润色报告文件
    if doc.report_file_path and os.path.exists(doc.report_file_path):
        os.remove(doc.report_file_path)
    
    delete_polished_document(db, doc_id)
    return {"message": "删除成功"}


class BatchDeleteRequest(BaseModel):
    ids: list[int]


@router.delete("/batch")
async def batch_delete_polished_documents(
    payload: BatchDeleteRequest,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    user = get_default_user(db)
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="仅管理员可删除文件")
    
    deleted = 0
    for doc_id in payload.ids:
        doc = get_polished_document(db, doc_id)
        if not doc:
            continue
        if os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        if doc.report_file_path and os.path.exists(doc.report_file_path):
            os.remove(doc.report_file_path)
        delete_polished_document(db, doc_id)
        deleted += 1
    
    return {"message": f"已删除 {deleted} 个文件", "deleted_count": deleted}


# ============================================================
# 润色反馈：准确率评分 + 修正词自动入库
# ============================================================

def _parse_corrections(text: str) -> list[tuple[str, str]]:
    """解析用户输入的修正内容，返回 [(非标准, 标准), ...] 列表。
    支持格式：'非标准→标准'、'非标准|标准'、'非标准 标准'
    """
    pairs = []
    if not text or not text.strip():
        return pairs
    
    for line in text.split('\n'):
        line = line.strip()
        if not line or line.startswith('#'):
            continue
        
        # 尝试多种分隔符
        for sep in ['→', '->', '|', '\t']:
            if sep in line:
                parts = line.split(sep, 1)
                old = parts[0].strip()
                new = parts[1].strip() if len(parts) > 1 else ''
                if old and new and old != new and len(old) >= 1:
                    pairs.append((old, new))
                break
        else:
            # 空格分隔（取前两个词）
            words = line.split()
            if len(words) >= 2:
                old = words[0].strip()
                new = words[1].strip()
                if old and new and old != new and len(old) >= 1:
                    pairs.append((old, new))
    
    return pairs


_cat_analyze_cache: dict = {}
_cat_download_cache: dict = {}
_cat_cache_timestamps: dict = {}
_CAT_CACHE_TTL_SECONDS = 12 * 3600
_CAT_ANALYZE_CACHE_DIR = os.path.join(UPLOAD_DIR, "cat-analyze-cache")


def _cleanup_cat_cache() -> None:
    return _cleanup_cat_cache_impl()


def _cat_cache_safe_id(analyze_id: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]", "", analyze_id or "")


def _cat_analyze_cache_json_path(analyze_id: str) -> Optional[str]:
    safe = _cat_cache_safe_id(analyze_id)
    if not safe:
        return None
    return os.path.join(_CAT_ANALYZE_CACHE_DIR, f"{safe}.json")


def _filename_from_cat_cache(analyze_id: str) -> str:
    json_path = _cat_analyze_cache_json_path(analyze_id)
    if not json_path or not os.path.exists(json_path):
        return ""
    try:
        with open(json_path, "r", encoding="utf-8") as handle:
            data = json.load(handle)
        payload = data.get("payload") if isinstance(data, dict) and isinstance(data.get("payload"), dict) else data
        file_info = payload.get("file_info") if isinstance(payload, dict) else {}
        return str((file_info or {}).get("filename") or "").strip()
    except Exception:
        return ""


def _fill_diagnose_source_names(db: Session, rows) -> None:
    pending = [row for row in rows if not str(getattr(row, "source_name", "") or "").strip()]
    if not pending:
        return
    analyze_ids = {row.analyze_id for row in pending if row.analyze_id}
    session_names = {}
    if analyze_ids:
        for session in db.query(CatAnalysisSession).filter(CatAnalysisSession.analyze_id.in_(analyze_ids)).all():
            name = str(session.source_filename or "").strip()
            if name:
                session_names[session.analyze_id] = name
    changed = False
    for row in pending:
        name = session_names.get(row.analyze_id) or _filename_from_cat_cache(row.analyze_id)
        if not name:
            continue
        row.source_name = name
        changed = True
    if changed:
        try:
            db.commit()
        except Exception:
            db.rollback()


def _persist_cat_analyze_cache(analyze_id: str, payload: dict) -> None:
    json_path = _cat_analyze_cache_json_path(analyze_id)
    if not json_path:
        return
    try:
        os.makedirs(_CAT_ANALYZE_CACHE_DIR, exist_ok=True)
        payload_copy = json.loads(json.dumps(payload, ensure_ascii=False, default=str))
        file_info = payload_copy.get("file_info") if isinstance(payload_copy.get("file_info"), dict) else {}
        temp_path = file_info.get("temp_path")
        if temp_path and os.path.isfile(temp_path):
            import shutil

            ext = os.path.splitext(temp_path)[1] or ""
            persisted_source = os.path.join(
                _CAT_ANALYZE_CACHE_DIR,
                f"{_cat_cache_safe_id(analyze_id)}.source{ext}",
            )
            shutil.copy2(temp_path, persisted_source)
            file_info["persisted_source_path"] = persisted_source
            payload_copy["file_info"] = file_info
            if isinstance(payload.get("file_info"), dict):
                payload["file_info"]["persisted_source_path"] = persisted_source
        blob = {"saved_at": time.time(), "payload": payload_copy}
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(blob, f, ensure_ascii=False)
    except Exception:
        logger.warning("[CAT_CACHE] 持久化分析结果失败: %s", analyze_id, exc_info=True)


def _restore_cat_source_path(payload: dict) -> dict:
    file_info = payload.get("file_info") if isinstance(payload.get("file_info"), dict) else {}
    temp_path = file_info.get("temp_path")
    persisted = file_info.get("persisted_source_path")
    if temp_path and os.path.isfile(temp_path):
        return payload
    if persisted and os.path.isfile(persisted):
        file_info["temp_path"] = persisted
        payload["file_info"] = file_info
    return payload


def _load_persisted_cat_analyze_cache(analyze_id: str) -> Optional[dict]:
    json_path = _cat_analyze_cache_json_path(analyze_id)
    if not json_path or not os.path.isfile(json_path):
        return None
    try:
        with open(json_path, encoding="utf-8") as f:
            blob = json.load(f)
        saved_at = float((blob or {}).get("saved_at") or 0)
        if time.time() - saved_at > _CAT_CACHE_TTL_SECONDS:
            return None
        payload = (blob or {}).get("payload")
        if not isinstance(payload, dict):
            return None
        return _restore_cat_source_path(payload)
    except Exception:
        logger.warning("[CAT_CACHE] 读取持久化分析结果失败: %s", analyze_id, exc_info=True)
        return None


def _touch_persisted_cat_analyze_cache(analyze_id: str) -> None:
    json_path = _cat_analyze_cache_json_path(analyze_id)
    if not json_path or not os.path.isfile(json_path):
        return
    try:
        with open(json_path, encoding="utf-8") as f:
            blob = json.load(f)
        if not isinstance(blob, dict):
            return
        blob["saved_at"] = time.time()
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(blob, f, ensure_ascii=False)
    except Exception:
        logger.warning("[CAT_CACHE] 续期持久化分析结果失败: %s", analyze_id, exc_info=True)


def _store_cat_analyze_cache(analyze_id: str, payload: dict) -> None:
    _cat_analyze_cache[analyze_id] = payload
    _cat_cache_timestamps[analyze_id] = time.time()
    _persist_cat_analyze_cache(analyze_id, payload)


def _get_cat_analyze_cache(analyze_id: str) -> Optional[dict]:
    now = time.time()
    cached = _cat_analyze_cache.get(analyze_id)
    last_access = _cat_cache_timestamps.get(analyze_id)
    if cached and last_access and now - last_access <= _CAT_CACHE_TTL_SECONDS:
        _cat_cache_timestamps[analyze_id] = now
        _touch_persisted_cat_analyze_cache(analyze_id)
        return _restore_cat_source_path(cached)
    loaded = _load_persisted_cat_analyze_cache(analyze_id)
    if not loaded:
        return None
    _cat_analyze_cache[analyze_id] = loaded
    _cat_cache_timestamps[analyze_id] = now
    _touch_persisted_cat_analyze_cache(analyze_id)
    return loaded


def _cleanup_cat_cache_impl() -> None:
    now = time.time()
    expired_ids = [
        analyze_id
        for analyze_id, last_access in list(_cat_cache_timestamps.items())
        if now - last_access > _CAT_CACHE_TTL_SECONDS
    ]
    for analyze_id in expired_ids:
        cached = _cat_analyze_cache.pop(analyze_id, None)
        _cat_cache_timestamps.pop(analyze_id, None)
        temp_path = ((cached or {}).get("file_info") or {}).get("temp_path")
        persisted_source = ((cached or {}).get("file_info") or {}).get("persisted_source_path")
        if temp_path and temp_path != persisted_source and os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                logger.warning("[CAT_CACHE] 清理临时文件失败: %s", temp_path)

    expired_download_tokens = [
        token
        for token, payload in list(_cat_download_cache.items())
        if now - float((payload or {}).get("created_at") or 0) > _CAT_CACHE_TTL_SECONDS
    ]
    for token in expired_download_tokens:
        _cat_download_cache.pop(token, None)


def _register_cat_download_asset(file_path: str, filename: str = "", media_type: str = "") -> tuple[Optional[str], Optional[str]]:
    if not file_path or not os.path.exists(file_path):
        return None, None
    download_token = str(uuid.uuid4())
    resolved_filename = filename or os.path.basename(file_path)
    _cat_download_cache[download_token] = {
        "path": file_path,
        "filename": resolved_filename,
        "media_type": media_type or mimetypes.guess_type(resolved_filename)[0] or "application/octet-stream",
        "created_at": time.time(),
    }
    return download_token, f"/api/polish-lab/cat/download/{download_token}"


def _generate_cat_html_report(
    report_path: str,
    source_filename: str,
    analyze_id: str,
    decisions: list,
    applied_changes: list,
    accuracy: dict,
    failed_replacements: Optional[list] = None,
    cached_items: Optional[list] = None,
    ai_semantic_scoring: Optional[bool] = None,
):
    return _generate_cat_html_report_impl(
        report_path,
        source_filename,
        analyze_id,
        decisions,
        applied_changes,
        accuracy,
        failed_replacements,
        cached_items,
        ai_semantic_scoring,
    )


_CAT_REPORT_CATEGORY_ORDER = (
    "grammar",
    "word",
    "term",
    "ambiguity",
    "redundancy",
    "logic",
    "missing",
    "risk",
)
_CAT_REPORT_SEVERITY_ORDER = ("high", "medium", "low")
_CAT_REPORT_SEVERITY_LABELS = {
    "high": "高 · 严重问题",
    "medium": "中 · 建议关注",
    "low": "低 · 提示性",
}
_CAT_REPORT_UNLABELED = "未标注"


def _cat_report_category_label(category_key: str) -> str:
    from app.utils.cat_diagnose import CATEGORIES, CATEGORY_LABELS, canonicalize_category

    raw = str(category_key or "").strip()
    if not raw:
        return _CAT_REPORT_UNLABELED
    key = canonicalize_category(raw, "")
    if key in CATEGORIES:
        return CATEGORY_LABELS.get(key, key)
    return CATEGORY_LABELS.get(raw.lower(), raw)


def _cat_report_severity_label(severity_key: str) -> str:
    key = str(severity_key or "").strip().lower()
    return _CAT_REPORT_SEVERITY_LABELS.get(key, _CAT_REPORT_UNLABELED)


def _cat_match_cached_item(decision, cached_items):
    key = _cat_decision_key(
        getattr(decision, "source_paragraph_index", None)
        if getattr(decision, "source_paragraph_index", None) is not None
        else getattr(decision, "paragraph_index", 0),
        getattr(decision, "sentence_index", 0),
        getattr(decision, "original_text", "") or getattr(decision, "source_sentence_text", "") or "",
    )
    for item in cached_items or []:
        if not isinstance(item, dict):
            continue
        item_key = _cat_decision_key(
            item.get("source_paragraph_index", item.get("paragraph_index")),
            item.get("sentence_index"),
            item.get("original_text") or item.get("source_sentence_text") or item.get("text") or "",
        )
        if item_key == key:
            return item
    return None


def _cat_pick_report_candidate(item, decision):
    candidates = [cand for cand in ((item or {}).get("candidates") or []) if isinstance(cand, dict)]
    if not candidates:
        return None
    wanted = (
        str(getattr(decision, "candidate_text", "") or "").strip()
        or str(getattr(decision, "accepted_template", "") or "").strip()
        or str(getattr(decision, "rejected_template", "") or "").strip()
    )
    if wanted:
        for cand in candidates:
            text = str(cand.get("template_text") or cand.get("revised") or "").strip()
            if text == wanted:
                return cand
    return candidates[0]


def _cat_report_is_ai_diagnose(decision, candidate) -> bool:
    rule_source = str(getattr(decision, "rule_source", None) or "").strip() or str((candidate or {}).get("rule_source") or "").strip()
    if rule_source == "ai_diagnose":
        return True
    template_id = str(getattr(decision, "accepted_template_id", None) or getattr(decision, "rejected_template_id", None) or "").strip() or str((candidate or {}).get("template_id") or "").strip()
    if template_id == "ai_diagnose":
        return True
    revised = str((candidate or {}).get("revised") or "").strip()
    problem = str((candidate or {}).get("problem") or "").strip()
    return bool(revised and problem)


def _cat_report_issue_meta(decision, cached_items=None) -> dict:
    from app.utils.cat_diagnose import CATEGORIES, canonicalize_category

    cached_item = _cat_match_cached_item(decision, cached_items)
    candidate = _cat_pick_report_candidate(cached_item, decision)
    category_raw = str(getattr(decision, "category", None) or "").strip() or str((candidate or {}).get("category") or "").strip()
    severity_raw = str(getattr(decision, "severity", None) or "").strip() or str((candidate or {}).get("severity") or "").strip()
    category_key = canonicalize_category(category_raw, str((candidate or {}).get("problem") or "")) if category_raw else ""
    if category_key not in CATEGORIES:
        category_key = ""
    is_ai_diagnose = _cat_report_is_ai_diagnose(decision, candidate)
    if is_ai_diagnose:
        severity_key = severity_raw.lower() if severity_raw.lower() in _CAT_REPORT_SEVERITY_ORDER else ""
        severity_label = _cat_report_severity_label(severity_key) if severity_key else _CAT_REPORT_UNLABELED
    else:
        severity_key = ""
        severity_label = "/"
    candidate_text = str(getattr(decision, "candidate_text", None) or "").strip()
    if not candidate_text:
        candidate_text = str((candidate or {}).get("template_text") or (candidate or {}).get("revised") or "").strip()
    if not candidate_text:
        action = str(getattr(decision, "action", "") or "")
        if action == "accept":
            candidate_text = str(getattr(decision, "accepted_template", "") or "")
        elif action == "reject":
            candidate_text = str(getattr(decision, "rejected_template", "") or "")
    original_text = str(getattr(decision, "original_text", "") or "")
    action = str(getattr(decision, "action", "") or "")
    if action == "accept":
        final_text = str(getattr(decision, "accepted_template", "") or "").strip()
    elif action == "modify":
        final_text = str(getattr(decision, "modified_text", "") or "").strip()
    else:
        final_text = ""
    return {
        "category_key": category_key,
        "category_label": _cat_report_category_label(category_key) if category_key else _CAT_REPORT_UNLABELED,
        "severity_key": severity_key,
        "severity_label": severity_label,
        "is_ai_diagnose": is_ai_diagnose,
        "original_text": original_text,
        "candidate_text": candidate_text,
        "final_text": final_text,
    }


def _cat_report_diff_token_html(text: str) -> str:
    return html_escape(str(text or "")).replace(" ", '<span class="diff-space">&#183;</span>')


def _cat_report_diff_segments(left_text: str, right_text: str) -> tuple[list, list]:
    left = str(left_text or "")
    right = str(right_text or "")
    left_segs = []
    right_segs = []
    matcher = SequenceMatcher(None, left, right, autojunk=False)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            if i1 != i2:
                left_segs.append((left[i1:i2], False))
            if j1 != j2:
                right_segs.append((right[j1:j2], False))
        elif tag == "delete":
            if i1 != i2:
                left_segs.append((left[i1:i2], True))
        elif tag == "insert":
            if j1 != j2:
                right_segs.append((right[j1:j2], True))
        else:
            if i1 != i2:
                left_segs.append((left[i1:i2], True))
            if j1 != j2:
                right_segs.append((right[j1:j2], True))
    return left_segs, right_segs


def _cat_report_render_segments_html(segments: list, changed_class: str) -> str:
    parts = []
    for text, changed in segments:
        content = _cat_report_diff_token_html(text)
        if changed:
            parts.append(f'<span class="{changed_class}">{content}</span>')
        else:
            parts.append(content)
    return "".join(parts)


def _cat_report_render_compare_html(original_text: str, candidate_text: str) -> tuple[str, str]:
    original = str(original_text or "")
    candidate = str(candidate_text or "")
    if not candidate:
        return _cat_report_diff_token_html(original), ""
    if len(original) + len(candidate) > 20000:
        return _cat_report_diff_token_html(original), _cat_report_diff_token_html(candidate)
    left_segs, right_segs = _cat_report_diff_segments(original, candidate)
    return (
        _cat_report_render_segments_html(left_segs, "diff-remove"),
        _cat_report_render_segments_html(right_segs, "diff-add"),
    )


def _cat_report_issue_highlights(category_counts: dict, severity_counts: dict, action_counts: dict, failed_replacements: list) -> list[str]:
    highlights = []
    labeled_categories = [
        (key, category_counts.get(key, 0))
        for key in _CAT_REPORT_CATEGORY_ORDER
        if category_counts.get(key, 0)
    ]
    labeled_categories.sort(key=lambda item: (-item[1], item[0]))
    if labeled_categories:
        top_key, top_count = labeled_categories[0]
        highlights.append(f'本轮问题类别以“{_cat_report_category_label(top_key)}”最多，共 {top_count} 句。')
    labeled_severities = [
        (key, severity_counts.get(key, 0))
        for key in _CAT_REPORT_SEVERITY_ORDER
        if severity_counts.get(key, 0)
    ]
    labeled_severities.sort(key=lambda item: (-item[1], _CAT_REPORT_SEVERITY_ORDER.index(item[0])))
    if labeled_severities:
        top_key, top_count = labeled_severities[0]
        highlights.append(f'本轮 AI 诊断严重程度以“{_cat_report_severity_label(top_key)}”最多，共 {top_count} 句。')
    if action_counts.get("modify"):
        highlights.append(f'共有 {action_counts["modify"]} 条句子采用人工自定义。')
    if action_counts.get("reject"):
        highlights.append(f'共有 {action_counts["reject"]} 条句子被人工驳回。')
    if failed_replacements:
        highlights.append(f'本轮有 {len(failed_replacements)} 条句子未成功定位原文。')
    return highlights or ["本轮没有形成足够的问题分类数据。"]


def _generate_cat_html_report_impl(
    report_path: str,
    source_filename: str,
    analyze_id: str,
    decisions: list,
    applied_changes: list,
    accuracy: dict,
    failed_replacements: Optional[list] = None,
    cached_items: Optional[list] = None,
    ai_semantic_scoring: Optional[bool] = None,
):
    generated_at = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    if ai_semantic_scoring is True:
        ai_semantic_label = "已开启"
    elif ai_semantic_scoring is False:
        ai_semantic_label = "已关闭"
    else:
        ai_semantic_label = "未知"
    action_labels = {
        'accept': '接受候选',
        'modify': '自定义润色',
        'reject': '拒绝候选',
        'pending': '待处理',
    }
    action_counts = {key: 0 for key in action_labels}
    decision_rows = []
    issue_category_counts = {key: 0 for key in _CAT_REPORT_CATEGORY_ORDER}
    issue_severity_counts = {key: 0 for key in _CAT_REPORT_SEVERITY_ORDER}
    unlabeled_category_count = 0
    unlabeled_severity_count = 0
    category_examples = {}
    category_summary_rows = []
    severity_summary_rows = []

    for decision in decisions or []:
        action = str(getattr(decision, 'action', '') or '')
        if action in action_counts:
            action_counts[action] += 1
        issue = _cat_report_issue_meta(decision, cached_items)
        category_key = issue["category_key"]
        severity_key = issue["severity_key"]
        if category_key in issue_category_counts:
            issue_category_counts[category_key] += 1
        else:
            unlabeled_category_count += 1
        if issue.get("is_ai_diagnose"):
            if severity_key in issue_severity_counts:
                issue_severity_counts[severity_key] += 1
            else:
                unlabeled_severity_count += 1
        example_text = issue["final_text"] or issue["candidate_text"] or issue["original_text"]
        if example_text and category_key and category_key not in category_examples:
            category_examples[category_key] = example_text[:80]

        original_html, candidate_html = _cat_report_render_compare_html(
            issue["original_text"],
            issue["candidate_text"],
        )
        decision_rows.append(
            f"<tr>"
            f"<td>{int(getattr(decision, 'sentence_index', 0)) + 1}</td>"
            f"<td>{int(getattr(decision, 'source_paragraph_index', getattr(decision, 'paragraph_index', 0)) or 0) + 1}</td>"
            f"<td>{html_escape(action_labels.get(action, action or '未标记'))}</td>"
            f"<td>{html_escape(issue['category_label'])}</td>"
            f"<td>{html_escape(issue['severity_label'])}</td>"
            f"<td>{original_html}</td>"
            f"<td>{candidate_html}</td>"
            f"<td>{html_escape(issue['final_text'])}</td>"
            f"</tr>"
        )

    total_decisions = max(1, len(decisions or []))
    for category_key in _CAT_REPORT_CATEGORY_ORDER:
        count = issue_category_counts[category_key]
        category_summary_rows.append(
            f"<tr>"
            f"<td>{html_escape(_cat_report_category_label(category_key))}</td>"
            f"<td>{count}</td>"
            f"<td>{round(count / total_decisions * 100, 1)}%</td>"
            f"<td>{html_escape(category_examples.get(category_key, ''))}</td>"
            f"</tr>"
        )
    if unlabeled_category_count:
        category_summary_rows.append(
            f"<tr>"
            f"<td>{html_escape(_CAT_REPORT_UNLABELED)}</td>"
            f"<td>{unlabeled_category_count}</td>"
            f"<td>{round(unlabeled_category_count / total_decisions * 100, 1)}%</td>"
            f"<td></td>"
            f"</tr>"
        )
    for severity_key in _CAT_REPORT_SEVERITY_ORDER:
        count = issue_severity_counts[severity_key]
        severity_summary_rows.append(
            f"<tr>"
            f"<td>{html_escape(_cat_report_severity_label(severity_key))}</td>"
            f"<td>{count}</td>"
            f"<td>{round(count / total_decisions * 100, 1)}%</td>"
            f"</tr>"
        )
    if unlabeled_severity_count:
        severity_summary_rows.append(
            f"<tr>"
            f"<td>{html_escape(_CAT_REPORT_UNLABELED)}</td>"
            f"<td>{unlabeled_severity_count}</td>"
            f"<td>{round(unlabeled_severity_count / total_decisions * 100, 1)}%</td>"
            f"</tr>"
        )

    report_highlights = _cat_report_issue_highlights(
        issue_category_counts,
        issue_severity_counts,
        action_counts,
        failed_replacements or [],
    )

    applied_rows = []
    for index, change in enumerate(applied_changes or [], start=1):
        action = str(change.get('action', '') or '')
        applied_rows.append(
            f"<tr>"
            f"<td>{index}</td>"
            f"<td>{html_escape(str(change.get('paragraph', '')))}</td>"
            f"<td>{html_escape(action_labels.get(action, action or '未标记'))}</td>"
            f"<td>{html_escape(str(change.get('before', '') or ''))}</td>"
            f"<td>{html_escape(str(change.get('after', '') or ''))}</td>"
            f"</tr>"
        )

    failed_rows = []
    for index, failure in enumerate(failed_replacements or [], start=1):
        failed_rows.append(
            f"<tr>"
            f"<td>{index}</td>"
            f"<td>{html_escape(str(failure.get('paragraph', '')))}</td>"
            f"<td>{html_escape(str(failure.get('action', '')))}</td>"
            f"<td>{html_escape(str(failure.get('source_sentence', '') or ''))}</td>"
            f"</tr>"
        )

    html = f"""<!DOCTYPE html>
<html lang=\"zh-CN\">
<head>
  <meta charset=\"UTF-8\" />
  <title>润色报告</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Microsoft YaHei', sans-serif; margin: 0; padding: 32px; color: #0f172a; background: #f8fafc; }}
    .page {{ max-width: 1120px; margin: 0 auto; background: #ffffff; border-radius: 20px; padding: 32px; box-shadow: 0 16px 40px rgba(15, 23, 42, 0.08); }}
    h1 {{ margin: 0 0 12px; font-size: 30px; }}
    h2 {{ margin: 32px 0 12px; font-size: 20px; }}
    p {{ margin: 6px 0; line-height: 1.7; color: #334155; }}
    .meta {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 8px 24px; margin-top: 16px; }}
    .summary {{ display: grid; grid-template-columns: repeat(6, minmax(0, 1fr)); gap: 12px; margin-top: 24px; }}
    .insight-list {{ margin: 12px 0 0; padding-left: 20px; color: #334155; line-height: 1.8; }}
    .card {{ background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 14px; padding: 16px; }}
    .label {{ display: block; font-size: 12px; color: #64748b; margin-bottom: 6px; }}
    .value {{ font-size: 24px; font-weight: 700; color: #0f172a; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 12px; table-layout: fixed; }}
    th, td {{ border: 1px solid #e2e8f0; padding: 10px 12px; text-align: left; vertical-align: top; font-size: 13px; line-height: 1.6; word-break: break-word; }}
    th {{ background: #f1f5f9; font-weight: 700; }}
    .diff-add {{ background: #fecaca; color: #991b1b; padding: 1px 4px; border-radius: 4px; }}
    .diff-remove {{ background: #bbf7d0; color: #166534; padding: 1px 4px; border-radius: 4px; }}
    .diff-space {{ display: inline-block; min-width: 0.72em; text-align: center; color: inherit; font-size: 0.9em; }}
    .note {{ margin-top: 24px; padding: 14px 16px; border-radius: 12px; background: #eff6ff; color: #1e3a8a; }}
  </style>
</head>
<body>
  <div class=\"page\">
    <h1>润色报告</h1>
    <p>本报告整理本次句式辅助润色的确认结果与实际写回内容，便于归档、复核与二次审校。</p>
    <div class=\"meta\">
      <p><strong>原始文件：</strong>{html_escape(source_filename or '')}</p>
      <p><strong>分析编号：</strong>{html_escape(analyze_id or '')}</p>
      <p><strong>生成时间：</strong>{generated_at}</p>
      <p><strong>准确率：</strong>{html_escape('待评估' if (accuracy or {}).get('accuracy_rate') is None else str((accuracy or {}).get('accuracy_rate')) + '%')}</p>
      <p><strong>AI语义：</strong>{html_escape(ai_semantic_label)}</p>
    </div>
    <div class=\"summary\">
      <div class=\"card\"><span class=\"label\">总句子决策</span><span class=\"value\">{len(decisions or [])}</span></div>
      <div class=\"card\"><span class=\"label\">接受候选</span><span class=\"value\">{action_counts['accept']}</span></div>
      <div class=\"card\"><span class=\"label\">自定义</span><span class=\"value\">{action_counts['modify']}</span></div>
      <div class=\"card\"><span class=\"label\">拒绝</span><span class=\"value\">{action_counts['reject']}</span></div>
      <div class=\"card\"><span class=\"label\">待处理</span><span class=\"value\">{action_counts['pending']}</span></div>
      <div class=\"card\"><span class=\"label\">实际写回句子</span><span class=\"value\">{len(applied_changes or [])}</span></div>
    </div>

    <h2>本轮润色特点分析</h2>
    <ul class=\"insight-list\">
      {''.join(f'<li>{html_escape(item)}</li>' for item in report_highlights)}
    </ul>

    <h2>问题类别分类</h2>
    <table>
      <thead>
        <tr><th style=\"width:120px\">问题类别</th><th style=\"width:70px\">数量</th><th style=\"width:90px\">占比</th><th>代表句</th></tr>
      </thead>
      <tbody>
        {''.join(category_summary_rows) if category_summary_rows else '<tr><td colspan="4">本次没有可分析的问题类别。</td></tr>'}
      </tbody>
    </table>

    <h2>AI诊断严重程度分类</h2>
    <table>
      <thead>
        <tr><th style=\"width:220px\">严重程度</th><th style=\"width:70px\">数量</th><th style=\"width:90px\">占比</th></tr>
      </thead>
      <tbody>
        {''.join(severity_summary_rows) if severity_summary_rows else '<tr><td colspan="3">本次没有可分析的严重程度。</td></tr>'}
      </tbody>
    </table>

    <h2>已写回文档的润色结果</h2>
    <table>
      <thead>
        <tr><th style=\"width:60px\">序号</th><th style=\"width:90px\">段落</th><th style=\"width:110px\">动作</th><th>修改前</th><th>修改后</th></tr>
      </thead>
      <tbody>
        {''.join(applied_rows) if applied_rows else '<tr><td colspan="5">本次没有写回到文档的修订内容。</td></tr>'}
      </tbody>
    </table>

    <h2>逐句确认明细</h2>
    <table>
      <thead>
        <tr><th style=\"width:50px\">句子</th><th style=\"width:50px\">段落</th><th style=\"width:80px\">动作</th><th style=\"width:80px\">问题类别</th><th style=\"width:110px\">严重程度</th><th>原文</th><th>候选</th><th>最终文本</th></tr>
      </thead>
      <tbody>
        {''.join(decision_rows) if decision_rows else '<tr><td colspan="8">本次没有可用的句子决策记录。</td></tr>'}
      </tbody>
    </table>

    <h2>未成功写回的句子</h2>
    <table>
      <thead>
        <tr><th style=\"width:60px\">序号</th><th style=\"width:70px\">段落</th><th style=\"width:90px\">动作</th><th>原句定位信息</th></tr>
      </thead>
      <tbody>
        {''.join(failed_rows) if failed_rows else '<tr><td colspan="4">本次所有需要写回的句子都已完成原文定位。</td></tr>'}
      </tbody>
    </table>

    <div class=\"note\">报告中的“实际写回句子”以生成的修订版 DOCX 为准；“逐句确认明细”完整记录用户在句式候选页面上的接受、自定义、拒绝和待处理状态。</div>
  </div>
</body>
</html>"""

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)


_CAT_REPLACE_NORMALIZE_PATTERN = re.compile(r'[\s\u3000，。；：！？,;:!?]')


def _normalize_cat_replace_text(text: str) -> str:
    return _CAT_REPLACE_NORMALIZE_PATTERN.sub('', str(text or ''))


def _cat_report_change_category(before: str, after: str, action: str) -> str:
    if action == 'reject':
        return '人工驳回'
    if action == 'pending':
        return '待人工确认'
    if not after:
        return '未输出结果'

    normalized_before = _normalize_cat_replace_text(before)
    normalized_after = _normalize_cat_replace_text(after)
    if before != after and normalized_before == normalized_after:
        return '标点格式统一'
    if len(normalized_after) <= max(4, int(len(normalized_before) * 0.85)):
        return '表达精简'
    if len(normalized_after) >= max(len(normalized_before) + 6, int(len(normalized_before) * 1.15)):
        return '信息补全'
    if re.search(r'(请|应|需|必须|不得|确保)', after) and not re.search(r'(请|应|需|必须|不得|确保)', before):
        return '操作指令规范'
    if re.search(r'[A-Za-z]{2,}|[0-9]+', before + after):
        return '术语措辞统一'
    return '句式重组'


def _cat_report_change_summary(before: str, after: str, action: str) -> str:
    if action == 'reject':
        return '本句保留人工判断，未写回候选句式。'
    if action == 'pending':
        return '本句仍处于待处理状态。'
    if not after:
        return '当前没有形成最终写回文本。'

    before_len = len(_normalize_cat_replace_text(before))
    after_len = len(_normalize_cat_replace_text(after))
    delta = after_len - before_len
    if before != after and before_len == after_len:
        return '主要调整了标点、空格或版式表达。'
    if delta < 0:
        return f'句子净减少 {abs(delta)} 个字符，整体更精简。'
    if delta > 0:
        return f'句子净增加 {delta} 个字符，补充了说明信息。'
    return '保留原意的同时重组了句式表达。'


def _cat_report_category_description(category: str) -> str:
    descriptions = {
        '表达精简': '压缩冗余措辞，让句子更短、更直接。',
        '信息补全': '补入限定条件、结果或操作说明。',
        '操作指令规范': '把动作句改成更规范的说明书表达。',
        '术语措辞统一': '统一术语、型号、英文缩写或专业措辞。',
        '标点格式统一': '统一全角半角、停顿和格式写法。',
        '句式重组': '保持原意，重排语序和结构。',
        '人工驳回': '用户明确驳回候选句式。',
        '待人工确认': '本句尚未形成最终决策。',
        '未输出结果': '没有生成可写回的最终文本。',
    }
    return descriptions.get(category, '本类修改体现了本轮润色的主要表达倾向。')


def _cat_report_highlights(action_counts: dict, sorted_categories: list, effective_changes: int, total_delta: int, failed_replacements: list) -> list[str]:
    highlights = []
    if sorted_categories:
        top_category, top_count = sorted_categories[0]
        highlights.append(f'本轮最集中的润色类型是“{top_category}”，共 {top_count} 句。')
    if effective_changes > 0:
        direction = '补充信息' if total_delta > 0 else '压缩表达' if total_delta < 0 else '重组句式'
        highlights.append(f'本轮实际生效的句子共 {effective_changes} 条，整体倾向于{direction}。')
    if action_counts.get('modify'):
        highlights.append(f'共有 {action_counts["modify"]} 条句子采用人工自定义，说明模板仍需继续贴近真实写作习惯。')
    if action_counts.get('reject'):
        highlights.append(f'共有 {action_counts["reject"]} 条句子被人工驳回，适合回查对应模板的召回条件。')
    if failed_replacements:
        highlights.append(f'本轮有 {len(failed_replacements)} 条句子未成功定位原文，生成前建议再次核对原句与候选句。')
    return highlights or ['本轮没有形成足够的有效润色句子，建议继续补充人工决策后再生成报告。']


def _replace_by_normalized_match(text: str, target: str, replacement: str) -> tuple[str, bool]:
    normalized_text = _normalize_cat_replace_text(text)
    normalized_target = _normalize_cat_replace_text(target)
    if not normalized_text or not normalized_target:
        return text, False

    start = normalized_text.find(normalized_target)
    if start < 0:
        return text, False

    index_map = []
    for index, char in enumerate(str(text or '')):
        if _CAT_REPLACE_NORMALIZE_PATTERN.match(char):
            continue
        index_map.append(index)

    if start >= len(index_map):
        return text, False

    end = start + len(normalized_target) - 1
    if end >= len(index_map):
        return text, False

    source_start = index_map[start]
    source_end = index_map[end] + 1
    return f"{text[:source_start]}{replacement}{text[source_end:]}", True


def _normalized_text_index_map(text: str) -> tuple[str, list[int]]:
    normalized_chars = []
    index_map = []
    for index, char in enumerate(str(text or '')):
        if _CAT_REPLACE_NORMALIZE_PATTERN.match(char):
            continue
        normalized_chars.append(char)
        index_map.append(index)
    return ''.join(normalized_chars), index_map


def _merge_replacement_with_source_tail(source_sentence: str, replacement: str) -> tuple[str, bool]:
    source_text = str(source_sentence or '')
    replacement_text = str(replacement or '')
    if not source_text or not replacement_text:
        return replacement_text, False

    normalized_source, source_index_map = _normalized_text_index_map(source_text)
    normalized_replacement, _ = _normalized_text_index_map(replacement_text)
    if not normalized_source or not normalized_replacement:
        return replacement_text, False
    if len(normalized_replacement) >= len(normalized_source):
        return replacement_text, False

    prefix_len = 0
    max_prefix = min(len(normalized_source), len(normalized_replacement))
    while prefix_len < max_prefix and normalized_source[prefix_len] == normalized_replacement[prefix_len]:
        prefix_len += 1

    if prefix_len < max(6, int(len(normalized_replacement) * 0.7)):
        return replacement_text, False
    if prefix_len >= len(normalized_source):
        return replacement_text, False
    if prefix_len >= len(source_index_map):
        return replacement_text, False

    suffix_start = source_index_map[prefix_len]
    source_suffix = source_text[suffix_start:]
    if not source_suffix.strip():
        return replacement_text, False

    merged = replacement_text.rstrip()
    if merged.endswith(('。', '！', '？', '.', '!', '?', ';', '；')) and source_suffix.startswith(('，', ',', '、', '；', ';', '：', ':')):
        merged = merged[:-1].rstrip()
    return f"{merged}{source_suffix}", True


def _resolve_cat_paragraph_index(paragraph_texts: list[str], decision: CatDecision) -> Optional[int]:
    candidate_indexes = []
    for raw_index in [decision.source_paragraph_index, decision.paragraph_index]:
        if raw_index is None:
            continue
        try:
            normalized_index = int(raw_index)
        except Exception:
            continue
        if 0 <= normalized_index < len(paragraph_texts):
            candidate_indexes.append(normalized_index)

    source_paragraph_text = str(getattr(decision, 'source_paragraph_text', '') or '').strip()
    source_sentence_text = str(getattr(decision, 'source_sentence_text', '') or getattr(decision, 'original_text', '') or '').strip()

    for idx in candidate_indexes:
        paragraph_text = str(paragraph_texts[idx] or '')
        if source_paragraph_text and paragraph_text.strip() == source_paragraph_text:
            return idx
        if source_sentence_text and source_sentence_text in paragraph_text:
            return idx

    if source_paragraph_text:
        for idx, paragraph_text in enumerate(paragraph_texts):
            if str(paragraph_text or '').strip() == source_paragraph_text:
                return idx

    if source_sentence_text:
        for idx, paragraph_text in enumerate(paragraph_texts):
            if source_sentence_text in str(paragraph_text or ''):
                return idx

    return candidate_indexes[0] if candidate_indexes else None


def _cat_decision_key(paragraph_index: Optional[int], sentence_index: Optional[int], original_text: str) -> tuple[int, int, str]:
    return (
        int(paragraph_index or 0),
        int(sentence_index or 0),
        _normalize_compare_text(original_text),
    )


@router.post("/cat/analyze", response_model=None)
async def cat_analyze(
    file: UploadFile = File(...),
    product_type: Optional[str] = Form(None),
    sentence_file_id: Optional[int] = Form(None),
    terminology_file_id: Optional[int] = Form(None),
    requirements: Optional[str] = Form(None),
    min_match_threshold: float = Form(0.34),
    fuzzy_lower_bound: float = Form(0.70),
    ai_semantic_scoring: bool = Form(True),
    ai_reason_max_chars: int = Form(15),
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user_optional),
):
    """
    第一阶段：CAT 式润色分析。
    上传 DOCX → 规则预处理 → 简化匹配 → AI 批量语义评分 → 返回候选列表。
    """
    import tempfile

    user = current_user or get_default_user(db)
    temp_path = None

    try:
        filename = file.filename or "unnamed"
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else "txt"

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp:
            content_bytes = await file.read()
            tmp.write(content_bytes)
            temp_path = tmp.name

        original_lines = None
        if ext in ['txt', 'md', 'markdown']:
            content = _read_file_safe(temp_path)
            original_lines = content.split('\n')
        elif ext == 'docx':
            from docx import Document
            doc = Document(temp_path)
            original_lines = [p.text for p in _iter_docx_body_paragraphs(doc)]
            content = '\n'.join(original_lines)
        else:
            content = _read_file_safe(temp_path)
            original_lines = content.split('\n')

        if not content or not content.strip():
            raise HTTPException(status_code=400, detail="无法提取文本内容")

        line_pre_polish_failed = False
        try:
            from app.utils.instrument_polisher import instrument_polish_engine
            pre_polished = instrument_polish_engine.pre_polish(content)
        except Exception as e:
            logger.warning("[CAT_ANALYZE] 规则预处理失败: %s", e)
            pre_polished = content
            line_pre_polish_failed = True

        selected_sentence_guide = _load_selected_sentence_guide_with_feedback(db, sentence_file_id)
        sentence_guide = _build_document_polish_guide(
            db,
            sentence_file_id=sentence_file_id,
            requirements=requirements,
            product_type=product_type,
        )
        candidate_recall_guide = _candidate_recall_guide_text(sentence_guide or '')
        if sentence_file_id and selected_sentence_guide:
            candidate_recall_guide = selected_sentence_guide
        guide_entries = _preferred_entries_from_guide(candidate_recall_guide)
        guide_templates = []
        seen_template_texts = set()
        for entry in guide_entries or []:
            template_text = _template_entry_text(entry)
            if not template_text or not _is_valid_cat_template(template_text):
                continue
            dedupe_candidate = re.sub(r'\s+', '', re.sub(r'[，。！？!?；;：:、,\.\s]+$', '', template_text))
            if not dedupe_candidate or dedupe_candidate in seen_template_texts:
                continue
            seen_template_texts.add(dedupe_candidate)
            template_id = str(entry.get("id", "")) if isinstance(entry, dict) else ""
            guide_templates.append({"text": template_text, "id": template_id})

        if not guide_templates:
            raise HTTPException(status_code=400, detail="句式库为空，请先选择或导入句式库")

        terminology = _load_scoped_terminology_source(db, terminology_file_id, product_type)
        resolved_terms = _resolve_terminology(db, terminology, pre_polished) if terminology else {}
        resolved_typo_dict = _load_typo_dict(db)

        lines = pre_polished.split('\n')
        if original_lines is None:
            original_lines = content.split('\n')

        if len(lines) != len(original_lines):
            if not line_pre_polish_failed:
                try:
                    from app.utils.instrument_polisher import instrument_polish_engine
                    lines = [instrument_polish_engine.pre_polish(line) if line else line for line in original_lines]
                except Exception as e:
                    logger.warning("[CAT_ANALYZE] 逐段预处理失败，回退原文段落: %s", e)
                    lines = list(original_lines)
            else:
                lines = list(original_lines)

        sentence_items = _split_cat_sentences(lines, source_paragraphs=original_lines)
        items = []
        simple_match_debug_summary = {
            "template_pool_size": len(guide_templates),
            "templates_considered": 0,
            "templates_matched": 0,
            "returned_candidates_before_ai": 0,
            "surface_rule_candidates": 0,
            "dropped_by_reason": {},
        }
        for sentence_item in sentence_items:
            line_stripped = sentence_item["text"].strip()
            source_line_stripped = str(sentence_item.get("source_sentence_text") or line_stripped).strip()
            local_entries = _preferred_entries_for_sentence(line_stripped, candidate_recall_guide)
            local_templates = []
            local_seen_template_texts = set()
            for entry in local_entries or []:
                template_text = _template_entry_text(entry)
                if not template_text or not _is_valid_cat_template(template_text):
                    continue
                dedupe_candidate = re.sub(r'\s+', '', re.sub(r'[，。！？!?；;：:、,\.\s]+$', '', template_text))
                if not dedupe_candidate or dedupe_candidate in local_seen_template_texts:
                    continue
                local_seen_template_texts.add(dedupe_candidate)
                template_id = str(entry.get("id", "")) if isinstance(entry, dict) else ""
                local_templates.append({"text": template_text, "id": template_id})
                if len(local_templates) >= 120:
                    break
            if not local_templates:
                local_templates = guide_templates[:120]
            simple_match_debug = {}
            candidates = _simple_match(
                line_stripped,
                local_templates,
                min_threshold=min_match_threshold,
                fuzzy_lower=fuzzy_lower_bound,
                term_dict=resolved_terms,
                typo_dict=resolved_typo_dict,
                source_sentence=source_line_stripped,
                debug_stats=simple_match_debug,
            )
            simple_match_debug = _finalize_simple_match_debug(simple_match_debug)
            simple_match_debug_summary["templates_considered"] += int(simple_match_debug.get("considered_templates", 0) or 0)
            simple_match_debug_summary["templates_matched"] += int(simple_match_debug.get("matched_templates", 0) or 0)
            simple_match_debug_summary["returned_candidates_before_ai"] += int(simple_match_debug.get("returned_candidates", 0) or 0)
            simple_match_debug_summary["surface_rule_candidates"] += int(simple_match_debug.get("surface_rule_candidates", 0) or 0)
            for reason, count in (simple_match_debug.get("dropped_by_reason", {}) or {}).items():
                simple_match_debug_summary["dropped_by_reason"][reason] = simple_match_debug_summary["dropped_by_reason"].get(reason, 0) + int(count or 0)
            items.append({
                "paragraph_index": sentence_item["source_paragraph_index"],
                "sentence_index": sentence_item["sentence_index"],
                "source_paragraph_index": sentence_item["source_paragraph_index"],
                "source_paragraph_text": sentence_item["source_paragraph_text"],
                "source_sentence_text": source_line_stripped,
                "original_text": source_line_stripped,
                "has_candidates": len(candidates) > 0,
                "candidates": candidates[:5],
                "simple_match_debug": simple_match_debug,
            })

        ai_scoring_status = "skipped"
        ai_scoring_error = None
        if ai_semantic_scoring:
            ai_score_result = await _batch_ai_semantic_score(
                items,
                reason_max_chars=ai_reason_max_chars,
            )
            ai_scoring_status = (ai_score_result or {}).get("status", "skipped")
            ai_scoring_error = (ai_score_result or {}).get("error")

        if ai_scoring_status != "completed":
            for item in items:
                for candidate in item.get("candidates", []):
                    if candidate.get("semantic_score") is None:
                        candidate["semantic_score"] = candidate.get("string_score", 0.0)
                        candidate["ai_reason"] = "AI未启用，使用字符串匹配分"

        candidate_debug_summary = {
            "template_pool_size": len(guide_templates),
            "templates_considered": simple_match_debug_summary["templates_considered"],
            "templates_matched": simple_match_debug_summary["templates_matched"],
            "returned_candidates_before_ai": simple_match_debug_summary["returned_candidates_before_ai"],
            "surface_rule_candidates": simple_match_debug_summary["surface_rule_candidates"],
            "simple_match_dropped_by_reason": simple_match_debug_summary["dropped_by_reason"],
            "total_before_filter": 0,
            "total_after_filter": 0,
            "needs_review_count": 0,
            "dropped_by_reason": {},
            "unmatched_sentence_count": 0,
            "diagnose_item_count": 0,
            "diagnose_hint_count": 0,
            "diagnose_rewrite_count": 0,
        }
        ai_semantic_active = ai_semantic_scoring and ai_scoring_status == "completed"

        for item in items:
            filtered_candidates = []
            for candidate in item.get("candidates", []):
                candidate["effective_semantic_score"] = _candidate_effective_semantic_score(candidate)
                candidate_debug_summary["total_before_filter"] += 1
                if _should_keep_cat_candidate(
                    item.get("original_text", ""),
                    candidate,
                    ai_semantic_active=ai_semantic_active,
                ):
                    if candidate.get("needs_review"):
                        candidate_debug_summary["needs_review_count"] += 1
                    filtered_candidates.append(candidate)
                else:
                    drop_reason = str(candidate.get("drop_reason") or "unknown")
                    candidate_debug_summary["dropped_by_reason"][drop_reason] = candidate_debug_summary["dropped_by_reason"].get(drop_reason, 0) + 1
                    if drop_reason == "low_filtered_score":
                        logger.info(
                            "[CAT_ANALYZE] drop stage=low_filtered_score quote=%s score=%s template=%s",
                            item.get("original_text") or "",
                            candidate.get("filtered_semantic_score"),
                            candidate.get("template_text") or "",
                        )
            filtered_candidates.sort(
                key=lambda candidate: (
                    float(candidate.get("filtered_semantic_score", 0.0) or 0.0),
                    float(candidate.get("effective_semantic_score", 0.0) or 0.0),
                    float(candidate.get("semantic_score", 0.0) or 0.0),
                    float(candidate.get("string_score", 0.0) or 0.0),
                ),
                reverse=True,
            )
            item["candidates"] = filtered_candidates[:5]
            item["has_candidates"] = len(item["candidates"]) > 0
            candidate_debug_summary["total_after_filter"] += len(item["candidates"])

        diagnose_items = []
        unmatched_sentence_count = sum(1 for item in items if not item.get("has_candidates"))
        candidate_debug_summary["unmatched_sentence_count"] = unmatched_sentence_count
        diagnose_status = "skipped"
        diagnose_error = ""
        analyze_id = str(uuid.uuid4())
        try:
            from app.utils.cat_diagnose import (
                annotate_cat_candidates,
                diagnoses_to_cat_items,
                is_ai_diagnose_enabled,
                open_diagnose_sentences,
            )
            annotate_cat_candidates(items)
            if is_ai_diagnose_enabled():
                diagnose_status = "completed"
                unmatched = _filter_cat_artifact_diagnose_pool([
                    {
                        "text": item.get("original_text") or "",
                        "paragraph_index": item.get("paragraph_index"),
                        "sentence_index": item.get("sentence_index"),
                        "source_paragraph_index": item.get("source_paragraph_index", item.get("paragraph_index")),
                        "source_paragraph_text": item.get("source_paragraph_text") or "",
                        "source_sentence_text": item.get("source_sentence_text") or item.get("original_text") or "",
                        "original_text": item.get("original_text") or "",
                    }
                    for item in items
                    if not item.get("has_candidates")
                ])
                try:
                    ai_diag = await open_diagnose_sentences(
                        unmatched,
                        resolved_terms or {},
                        sentence_guide,
                        product_type or "",
                    )
                    ai_diag = _filter_cat_artifact_diagnoses(ai_diag, unmatched)
                    _persist_lab_diagnoses(
                        db,
                        ai_diag,
                        unmatched,
                        analyze_id=analyze_id,
                        source="document",
                        source_name=filename,
                    )
                    diagnose_items = _filter_cat_artifact_diagnose_items(
                        diagnoses_to_cat_items(ai_diag, unmatched)
                    )
                except Exception as e:
                    logger.warning("[CAT_ANALYZE] AI 诊断失败: %s", e, exc_info=True)
                    diagnose_items = []
                    diagnose_status = "failed"
                    diagnose_error = str(e)
        except Exception as e:
            logger.warning("[CAT_ANALYZE] 诊断阶段失败: %s", e, exc_info=True)
            diagnose_items = []
            diagnose_status = "failed"
            diagnose_error = str(e)

        candidate_debug_summary["diagnose_item_count"] = len(diagnose_items)
        from app.utils.cat_diagnose import HINT_ONLY_CATEGORIES as _HINT_ONLY_CATEGORIES
        diagnose_hint_count = 0
        diagnose_rewrite_count = 0
        for item in diagnose_items:
            for candidate in item.get("candidates") or []:
                if not isinstance(candidate, dict):
                    continue
                is_hint = (
                    candidate.get("category") in _HINT_ONLY_CATEGORIES
                    and not str(candidate.get("revised") or "").strip()
                )
                if is_hint:
                    diagnose_hint_count += 1
                else:
                    diagnose_rewrite_count += 1
        candidate_debug_summary["diagnose_hint_count"] = diagnose_hint_count
        candidate_debug_summary["diagnose_rewrite_count"] = diagnose_rewrite_count
        _cleanup_cat_cache()
        _store_cat_analyze_cache(analyze_id, {
            "items": list(items) + list(diagnose_items),
            "templates": guide_templates,
            "file_info": {
                "filename": filename,
                "total_paragraphs": len(items),
                "content": pre_polished,
                "paragraph_texts": list(original_lines),
                "polished_lines": list(lines),
                "temp_path": temp_path,
                "resolved_terms": resolved_terms,
                "resolved_term_count": len(resolved_terms or {}),
                "user_id": user.id if user else None,
                "sentence_file_id": sentence_file_id,
                "sentence_file_name": _resolve_sentence_file_name(db, sentence_file_id),
                "ai_semantic_scoring": bool(ai_semantic_scoring),
                "ai_scoring_status": ai_scoring_status,
            },
        })

        total_with_candidates = sum(1 for i in items if i.get("has_candidates")) + sum(
            1 for i in diagnose_items if i.get("has_candidates")
        )
        total_paragraphs = len(items)

        return {
            "analyze_id": analyze_id,
            "total_paragraphs": total_paragraphs,
            "total_with_candidates": total_with_candidates,
            "template_coverage": round(total_with_candidates / total_paragraphs * 100, 1) if total_paragraphs else 0,
            "ai_scoring_status": ai_scoring_status,
            "ai_scoring_error": ai_scoring_error,
            "diagnose_status": diagnose_status,
            "diagnose_error": diagnose_error,
            "resolved_term_count": len(resolved_terms or {}),
            "candidate_debug_summary": candidate_debug_summary,
            "items": items,
            "diagnose_items": diagnose_items,
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error("[CAT_ANALYZE] 分析失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"分析失败: {e}")


@router.post("/cat/apply", response_model=None)
async def cat_apply(
    request: CatApplyRequest,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user_optional),
):
    """
    第二阶段：用户决策 + 写入 DOCX + 反馈。
    """
    user = current_user or get_default_user(db)
    cached = _get_cat_analyze_cache(request.analyze_id)
    if not cached:
        raise HTTPException(status_code=404, detail="分析结果已过期，请重新上传分析")

    file_info = cached["file_info"]
    temp_path = file_info.get("temp_path")
    content = file_info.get("content", "")
    paragraph_texts = list(file_info.get("paragraph_texts") or content.split('\n'))
    filename = file_info.get("filename", "润色文档.docx")
    visible_item_keys = {
        _cat_decision_key(
            item.get("source_paragraph_index"),
            item.get("sentence_index"),
            item.get("original_text", ""),
        )
        for item in (cached.get("items") or [])
        if item.get("has_candidates") and item.get("candidates")
    }

    if not temp_path or not os.path.exists(temp_path):
        raise HTTPException(status_code=404, detail="临时文件已丢失，请重新上传分析")

    decisions = [
        d for d in request.decisions
        if _cat_decision_key(
            getattr(d, "source_paragraph_index", None),
            getattr(d, "sentence_index", None),
            getattr(d, "original_text", ""),
        ) in visible_item_keys
    ]

    rejected_count = _cat_save_rejected_as_learning(decisions, db, user)

    modified_count = _cat_save_modified_to_feedback(
        decisions, db, user, source_filename=request.source_filename or filename,
    )

    paragraph_sentence_replacements = defaultdict(list)
    for d in decisions:
        action = d.action
        para_idx = _resolve_cat_paragraph_index(paragraph_texts, d)
        if para_idx is None or para_idx < 0 or para_idx >= len(paragraph_texts):
            continue
        if action not in {"accept", "modify"}:
            continue
        replacement_text = d.accepted_template if action == "accept" else d.modified_text
        if not replacement_text:
            continue
        source_sentence = d.source_sentence_text or d.original_text or ""
        normalized_replacement_text = replacement_text
        replacement_completed_from_source = False
        if source_sentence:
            normalized_replacement_text, replacement_completed_from_source = _merge_replacement_with_source_tail(
                source_sentence,
                replacement_text,
            )
        paragraph_sentence_replacements[para_idx].append({
            "source_sentence": source_sentence,
            "fallback_sentence": d.original_text or "",
            "replacement_text": normalized_replacement_text,
            "action": action,
            "replacement_completed_from_source": replacement_completed_from_source,
        })

    paragraph_revisions = {}
    paragraph_revision_actions = {}
    failed_replacements = []
    applied_changes = []
    applied_accept_count = 0
    applied_modify_count = 0
    for para_idx, replacements in paragraph_sentence_replacements.items():
        paragraph_text = paragraph_texts[para_idx]
        updated_text = paragraph_text
        paragraph_actions = sorted({
            replacement.get("action")
            for replacement in replacements
            if replacement.get("action") in {"accept", "modify"}
        })
        for replacement in replacements:
            source_sentence = replacement.get("source_sentence", "")
            fallback_sentence = replacement.get("fallback_sentence", "")
            replacement_text = replacement.get("replacement_text", "")
            replacement_completed_from_source = bool(replacement.get("replacement_completed_from_source"))
            matched = False
            if source_sentence and source_sentence in updated_text:
                updated_text = updated_text.replace(source_sentence, replacement_text, 1)
                matched = True
            elif source_sentence:
                updated_text, matched = _replace_by_normalized_match(updated_text, source_sentence, replacement_text)
            if not matched and fallback_sentence and fallback_sentence in updated_text:
                updated_text = updated_text.replace(fallback_sentence, replacement_text, 1)
                matched = True
            if not matched and fallback_sentence:
                updated_text, matched = _replace_by_normalized_match(updated_text, fallback_sentence, replacement_text)
            if matched:
                if replacement.get("action") == "accept":
                    applied_accept_count += 1
                elif replacement.get("action") == "modify":
                    applied_modify_count += 1
                applied_changes.append({
                    "paragraph": para_idx + 1,
                    "before": (source_sentence or fallback_sentence or '')[:200],
                    "after": replacement_text[:200],
                    "action": replacement.get("action") or "",
                })
            if not matched:
                failed_replacements.append({
                    "paragraph": para_idx + 1,
                    "action": replacement.get("action") or "",
                    "source_sentence": (source_sentence or fallback_sentence or '')[:120],
                })
            elif replacement_completed_from_source:
                logger.info("[CAT_APPLY] 已补回原文尾部: paragraph=%s source=%s", para_idx + 1, (source_sentence or fallback_sentence or '')[:80])
        if updated_text.strip() != paragraph_text.strip():
            paragraph_revisions[para_idx] = updated_text
            paragraph_revision_actions[para_idx] = paragraph_actions

    output_path = None
    report_path = None

    if temp_path.lower().endswith('.docx') and paragraph_revisions:
        output_dir = os.path.dirname(temp_path)
        output_filename = f"【润色版】{filename}" if not filename.startswith("【") else filename
        output_path = os.path.join(output_dir, output_filename)

        try:
            from lxml import etree
            from docx import Document as DocxDocument

            doc = DocxDocument(temp_path)
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            xml_paragraphs = [paragraph._p for paragraph in _iter_docx_body_paragraphs(doc)]

            revision_id = 100
            author = user.username if user else "润色结果"
            now_str = datetime.datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')

            for para_idx, new_text in paragraph_revisions.items():
                if para_idx >= len(xml_paragraphs):
                    continue
                p_element = xml_paragraphs[para_idx]
                original_text = paragraph_texts[para_idx] if para_idx < len(paragraph_texts) else ""
                if original_text.strip() == new_text.strip():
                    continue

                revision_id += 1
                rid_del = str(revision_id)
                revision_id += 1
                rid_ins = str(revision_id)
                _apply_paragraph_revision_xml(
                    p_element, new_text, author, now_str,
                    rid_del, rid_ins, w_ns,
                )

            document_root = doc.element
            final_xml = etree.tostring(document_root, xml_declaration=True, encoding='UTF-8')
            _write_revised_docx(temp_path, output_path, final_xml)

        except Exception as e:
            logger.error("[CAT_APPLY] DOCX 修订标记写入失败: %s", e, exc_info=True)
            raise HTTPException(status_code=500, detail=f"DOCX 写入失败: {e}")
    else:
        output_path = temp_path

    total_paragraphs = file_info.get("total_paragraphs", len(cached.get("items", [])))
    accuracy = _cat_calc_file_accuracy(decisions, total_paragraphs)
    rejected_decisions = sum(1 for d in decisions if getattr(d, "action", "") == "reject")
    pending_decisions = sum(1 for d in decisions if getattr(d, "action", "") == "pending")
    effective_decided = applied_accept_count + applied_modify_count + rejected_decisions + len(failed_replacements)
    accuracy.update({
        "accepted": applied_accept_count,
        "modified": applied_modify_count,
        "rejected": rejected_decisions,
        "pending": pending_decisions,
        "failed": len(failed_replacements),
        "accuracy_rate": round(applied_accept_count / effective_decided * 100, 1) if effective_decided > 0 else None,
        "rejection_rate": round(rejected_decisions / effective_decided * 100, 1) if effective_decided > 0 else None,
        "modification_rate": round(applied_modify_count / effective_decided * 100, 1) if effective_decided > 0 else None,
    })

    sentence_file_id = request.sentence_file_id or file_info.get("sentence_file_id")
    sentence_file_name = request.sentence_file_name or file_info.get("sentence_file_name")
    category_summary = {}
    decision_records = []
    for decision in decisions:
        replacement_text = _cat_decision_replacement_text(decision)
        issue = _cat_report_issue_meta(decision, cached.get("items") or [])
        category = issue["category_label"]
        category_summary[category] = category_summary.get(category, 0) + 1
        decision_records.append(CatDecisionRecord(
            analyze_id=request.analyze_id,
            paragraph_index=getattr(decision, 'paragraph_index', 0),
            sentence_index=getattr(decision, 'sentence_index', 0),
            action=getattr(decision, 'action', 'pending') or 'pending',
            original_text=getattr(decision, 'original_text', '') or '',
            replacement_text=replacement_text,
            category=category,
            string_score=float(getattr(decision, 'string_score', 0.0) or 0.0),
            semantic_score=getattr(decision, 'semantic_score', None),
        ))

    db.query(CatDecisionRecord).filter(CatDecisionRecord.analyze_id == request.analyze_id).delete(synchronize_session=False)
    db.query(CatAnalysisSession).filter(CatAnalysisSession.analyze_id == request.analyze_id).delete(synchronize_session=False)
    session_row = CatAnalysisSession(
        analyze_id=request.analyze_id,
        source_filename=request.source_filename or filename,
        sentence_file_id=sentence_file_id,
        sentence_file_name=sentence_file_name,
        total_paragraphs=total_paragraphs,
        total_items=len(decisions),
        accepted=applied_accept_count,
        rejected=rejected_decisions,
        modified=applied_modify_count,
        pending=pending_decisions,
        accuracy_rate=accuracy.get("accuracy_rate"),
        rejection_rate=accuracy.get("rejection_rate"),
        modification_rate=accuracy.get("modification_rate"),
        template_coverage=accuracy.get("template_coverage"),
        category_summary=_json_dumps(category_summary),
        created_by=user.username if user else "guest",
    )
    db.add(session_row)
    db.flush()
    for record in decision_records:
        record.session_id = session_row.id
        db.add(record)
    db.commit()

    report_dir = os.path.dirname(temp_path)
    report_base = os.path.splitext(filename)[0] if filename else "润色文档"
    report_filename = f"【润色报告】{report_base}.html"
    report_path = os.path.join(report_dir, report_filename)
    _generate_cat_html_report(
        report_path=report_path,
        source_filename=request.source_filename or filename,
        analyze_id=request.analyze_id,
        decisions=decisions,
        applied_changes=applied_changes,
        accuracy=accuracy,
        failed_replacements=failed_replacements,
        cached_items=cached.get("items") or [],
        ai_semantic_scoring=(
            request.ai_semantic_scoring
            if request.ai_semantic_scoring is not None
            else (cached.get("file_info") or {}).get("ai_semantic_scoring")
        ),
    )

    download_url = None
    download_filename = None
    if output_path and os.path.exists(output_path):
        download_filename = os.path.basename(output_path)
        _, download_url = _register_cat_download_asset(output_path, download_filename)

    report_download_url = None
    report_download_filename = None
    if report_path and os.path.exists(report_path):
        report_download_filename = os.path.basename(report_path)
        _, report_download_url = _register_cat_download_asset(report_path, report_download_filename, "text/html; charset=utf-8")

    doc_id = None
    preview_url = None
    if output_path and os.path.exists(output_path):
        try:
            import shutil

            os.makedirs(UPLOAD_DIR, exist_ok=True)
            persistent_filename = f"{uuid.uuid4()}.docx"
            persistent_path = os.path.join(UPLOAD_DIR, persistent_filename)
            shutil.copy2(output_path, persistent_path)

            persistent_report_filename = None
            persistent_report_path = None
            if report_path and os.path.exists(report_path):
                persistent_report_filename = f"{uuid.uuid4()}.html"
                persistent_report_path = os.path.join(UPLOAD_DIR, persistent_report_filename)
                shutil.copy2(report_path, persistent_report_path)

            polished_content = ""
            try:
                from docx import Document as DocxReader

                polished_doc = DocxReader(persistent_path)
                polished_content = '\n'.join([p.text for p in _iter_docx_body_paragraphs(polished_doc)])
            except Exception as preview_error:
                logger.warning("[CAT_APPLY] 提取预览文本失败: %s", preview_error)

            db_doc = create_polished_document(
                db=db,
                name=download_filename or os.path.basename(output_path),
                filename=persistent_filename,
                file_path=persistent_path,
                file_size=os.path.getsize(persistent_path),
                file_type="docx",
                original_content='\n'.join(paragraph_texts),
                polished_content=polished_content,
                created_by=user.id if user else None,
                report_filename=persistent_report_filename,
                report_file_path=persistent_report_path,
            )
            doc_id = db_doc.id
            preview_url = f"/polish/preview/{doc_id}"
        except Exception as persist_error:
            logger.error("[CAT_APPLY] 保存预览记录失败: %s", persist_error, exc_info=True)

    _cat_analyze_cache.pop(request.analyze_id, None)
    _cat_cache_timestamps.pop(request.analyze_id, None)

    return {
        "message": "润色完成",
        "output_file": output_path,
        "download_url": download_url,
        "download_filename": download_filename,
        "report_download_url": report_download_url,
        "report_download_filename": report_download_filename,
        "applied_changes": applied_changes,
        "applied_count": len(applied_changes),
        "failed_count": len(failed_replacements),
        "failed_replacements": failed_replacements,
        "accuracy": accuracy,
        "doc_id": doc_id,
        "preview_url": preview_url,
        "feedback": {
            "rejected_saved": rejected_count,
            "modified_saved": modified_count,
        },
    }


@router.get("/cat/stats/{analyze_id}", response_model=None)
def cat_get_stats(
    analyze_id: str,
    db: Session = Depends(get_db),
):
    """查询单次分析的准确率（在用户决策提交后可用）。"""
    cached = _get_cat_analyze_cache(analyze_id)
    if not cached:
        raise HTTPException(status_code=404, detail="分析结果已过期")
    items = cached.get("items", [])
    total_paragraphs = cached.get("file_info", {}).get("total_paragraphs", 0)
    total_with_candidates = sum(1 for i in items if i.get("has_candidates"))
    return {
        "analyze_id": analyze_id,
        "total_paragraphs": total_paragraphs,
        "total_with_candidates": total_with_candidates,
        "template_coverage": round(total_with_candidates / total_paragraphs * 100, 1) if total_paragraphs else 0,
    }


@router.get("/cat/download/{download_token}")
def cat_download_file(download_token: str):
    payload = _cat_download_cache.get(download_token)
    if not payload:
        raise HTTPException(status_code=404, detail="下载链接已失效")

    file_path = payload.get("path")
    filename = payload.get("filename") or "polished.docx"
    if not file_path or not os.path.exists(file_path):
        _cat_download_cache.pop(download_token, None)
        raise HTTPException(status_code=404, detail="输出文件不存在")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type=payload.get("media_type") or "application/octet-stream",
    )
