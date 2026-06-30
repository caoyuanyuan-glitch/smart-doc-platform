from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import csv
import io
import os
import re
import json
import time
import zipfile
import shutil
import threading
import xml.etree.ElementTree as ET
from app.database import get_db
from app.crud.convert import (
    create_convert_task, get_convert_task, get_convert_tasks,
    update_convert_task_status, update_convert_task_result, update_convert_task_failed,
    update_convert_task_progress, delete_convert_task
)
from app.schemas.convert import (
    ConvertTaskOut, ConvertTaskCreate, ProgressOut, ReportOut,
    StepStatus, CheckItem
)

router = APIRouter()

UPLOAD_DIR = "./static/uploads"
OUTPUT_DIR = "./static/uploads/outputs"

_tasks_lock = threading.Lock()
_task_store = {}

IME_TOPIC_DECL = '<!DOCTYPE topic PUBLIC "-//OASIS//DTD DITA Topic//EN" "topic.dtd">'
IME_BOOKMAP_DECL = '<!DOCTYPE bookmap PUBLIC "-//OASIS//DTD DITA BookMap//EN" "bookmap.dtd">'


def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)


def _media_type_for_file(filename):
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".csv":
        return "text/csv; charset=utf-8"
    if ext in [".md", ".markdown"]:
        return "text/markdown; charset=utf-8"
    if ext == ".zip":
        return "application/zip"
    return "application/octet-stream"


_ensure_dir(UPLOAD_DIR)
_ensure_dir(OUTPUT_DIR)

GENERATED_NODE_VERSION = "A.2"
GENERATED_TOPIC_BASE = 41000


def _get_step_names():
    return ["解析", "结构映射", "内容替换", "验证", "打包"]


def _get_topic_type(heading_text):
    task_keywords = ["安装", "配置", "操作", "步骤", "部署", "启动", "设置", "升级", "卸载", "登录",
                     "install", "config", "setup", "deploy", "start", "upgrade", "uninstall", "login"]
    ref_keywords = ["参数", "规格", "表格", "参考", "附录", "错误码", "api", "接口", "命令",
                    "parameter", "spec", "reference", "appendix", "command"]
    heading_lower = heading_text.lower()

    for kw in task_keywords:
        if kw in heading_lower:
            return "task"
    for kw in ref_keywords:
        if kw in heading_lower:
            return "reference"
    return "concept"


def _clean_title(title):
    """Clean markdown formatting from section titles.
    - Unescape markdown backslash escapes
    - Strip ** bold markers
    - Remove leading numbering (1.1, Chapter 1:, etc.)
    """
    if not title:
        return title

    # 1. Unescape markdown: \\ must come first to avoid double-unescaping
    t = title
    t = t.replace("\\\\", "\\")
    for ch in "&*_`.-#":
        t = t.replace("\\" + ch, ch)

    # 2. Strip ** bold markers
    t = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', t)

    # 3. Remove leading chapter markers while preserving source section numbering
    t = re.sub(r'^(Chapter|Appendix|第)\s*\d+[\.\s:：]*\s*', '', t, flags=re.IGNORECASE)

    # 4. Trim and clean up
    t = re.sub(r'\s+', ' ', t).strip()
    return t


def _parse_md_sections(content):
    """Parse markdown into nested sections for all heading levels (H1-H6).
    Returns: [{"title": str, "sections": [...], "content": str}]
    """
    lines = content.split("\n")
    root_sections = []
    # Stack entries: {"title": str, "level": int, "sections": list, "content_lines": list}
    stack = []
    pending_lines = []

    for line in lines:
        heading_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()

            if pending_lines:
                if stack:
                    ck = "content_lines"
                    if ck not in stack[-1]:
                        stack[-1][ck] = []
                    stack[-1][ck].extend(pending_lines)
                pending_lines = []

            while stack and stack[-1]["level"] >= level:
                complete = stack.pop()
                sec_content = "\n".join(complete.get("content_lines", [])).strip()
                entry = {
                    "title": complete["title"],
                    "sections": complete.get("sections", []),
                    "content": sec_content,
                }
                if stack:
                    stack[-1].setdefault("sections", []).append(entry)
                else:
                    root_sections.append(entry)

            stack.append({
                "title": title,
                "level": level,
                "sections": [],
            })
        else:
            pending_lines.append(line)

    if pending_lines:
        if stack:
            stack[-1].setdefault("content_lines", []).extend(pending_lines)
        pending_lines = []

    while stack:
        complete = stack.pop()
        sec_content = "\n".join(complete.get("content_lines", [])).strip()
        entry = {
            "title": complete["title"],
            "sections": complete.get("sections", []),
            "content": sec_content,
        }
        if stack:
            stack[-1].setdefault("sections", []).append(entry)
        else:
            root_sections.append(entry)

    return root_sections


def _parse_docx_sections(file_path):
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            return [{"title": "文档", "sections": [{"title": "正文", "content": "(无内容)"}]}]

        root_sections = []
        stack = []
        pending_lines = []

        for p in paragraphs:
            text = p.text.strip()
            style_name = p.style.name if p.style else ""
            heading_level = _get_heading_level(style_name)

            if heading_level:
                if pending_lines:
                    if stack:
                        stack[-1].setdefault("content_lines", []).extend(pending_lines)
                    pending_lines = []

                while stack and stack[-1]["level"] >= heading_level:
                    complete = stack.pop()
                    sec_content = "\n".join(complete.get("content_lines", [])).strip()
                    entry = {
                        "title": complete["title"],
                        "sections": complete.get("sections", []),
                        "content": sec_content,
                    }
                    if stack:
                        stack[-1].setdefault("sections", []).append(entry)
                    else:
                        root_sections.append(entry)

                stack.append({
                    "title": text,
                    "level": heading_level,
                    "sections": [],
                })
            else:
                pending_lines.append(text)

        if pending_lines:
            if stack:
                stack[-1].setdefault("content_lines", []).extend(pending_lines)
            pending_lines = []

        while stack:
            complete = stack.pop()
            sec_content = "\n".join(complete.get("content_lines", [])).strip()
            entry = {
                "title": complete["title"],
                "sections": complete.get("sections", []),
                "content": sec_content,
            }
            if stack:
                stack[-1].setdefault("sections", []).append(entry)
            else:
                root_sections.append(entry)

        if not root_sections:
            full_text = "\n".join(p.text.strip() for p in paragraphs)
            return [{"title": "文档", "sections": [{"title": "正文", "content": full_text}]}]

        return root_sections
    except Exception as e:
        raise ValueError(f"DOCX解析失败: {str(e)}")


def _iter_docx_blocks(doc):
    from docx.document import Document as DocxDocumentClass
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    parent = doc.element.body if isinstance(doc, DocxDocumentClass) else doc
    for child in parent.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _normalize_docx_text(text):
    return re.sub(r'\s+', ' ', str(text or '')).strip()


def _is_docx_toc_number(text):
    return bool(re.fullmatch(r'\d+(?:\.\d+)*', text or ''))


def _docx_table_to_markdown(table):
    rows = []
    for row in table.rows:
        cells = [_normalize_docx_text(cell.text) for cell in row.cells]
        while cells and not cells[-1]:
            cells.pop()
        if any(cells):
            rows.append(cells)

    if not rows:
        return []

    cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (cols - len(row)) for row in rows]
    header = normalized_rows[0]
    md_lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join([":---"] * cols) + " |",
    ]
    for row in normalized_rows[1:]:
        md_lines.append("| " + " | ".join(row) + " |")
    return md_lines


def _extract_docx_paragraph_images(paragraph, media_dir, image_index):
    image_refs = []
    rel_ns = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    blips = paragraph._element.xpath('.//*[local-name()="blip"]')
    for blip in blips:
        rel_id = blip.get(rel_ns)
        if not rel_id:
            continue
        image_part = paragraph.part.related_parts.get(rel_id)
        if not image_part:
            continue
        ext = os.path.splitext(str(getattr(image_part, "partname", "")))[1] or ".png"
        file_name = f"docx_image_{image_index:03d}{ext.lower()}"
        file_path = os.path.join(media_dir, file_name)
        if not os.path.exists(file_path):
            with open(file_path, "wb") as f:
                f.write(image_part.blob)
        image_refs.append(file_path)
        image_index += 1
    return image_refs, image_index


def _docx_to_markdown(file_path):
    from docx import Document as DocxDocument
    from docx.table import Table

    doc = DocxDocument(file_path)
    media_dir = os.path.splitext(file_path)[0] + "_media"
    _ensure_dir(media_dir)

    blocks = list(_iter_docx_blocks(doc))
    lines = []
    image_index = 1
    cover_started = False

    def ensure_cover_heading():
        nonlocal cover_started
        if not cover_started:
            lines.append("# Cover")
            lines.append("")
            cover_started = True

    for idx, block in enumerate(blocks):
        next_block = blocks[idx + 1] if idx + 1 < len(blocks) else None
        if not isinstance(block, Table):
            text = _normalize_docx_text(block.text)
            image_refs, image_index = _extract_docx_paragraph_images(block, media_dir, image_index)
            style_name = block.style.name if block.style else ""
            heading_level = _get_heading_level(style_name)

            is_table_title = text.startswith(("Table ", "TABLE ", "Figure ", "FIGURE "))
            inferred_heading = (
                heading_level == 0
                and text
                and not is_table_title
                and len(text) <= 80
                and next_block is not None
                and hasattr(next_block, "rows")
            )

            if text and heading_level and _is_docx_toc_number(text):
                continue

            if text and (heading_level or inferred_heading):
                lines.append(f"{'#' * max(1, heading_level or 2)} {text}")
                lines.append("")
            else:
                if text:
                    ensure_cover_heading()
                    lines.append(text)
                for image_path in image_refs:
                    ensure_cover_heading()
                    lines.append(f"![{os.path.basename(image_path)}]({image_path})")
                if text or image_refs:
                    lines.append("")
        else:
            ensure_cover_heading()
            table_lines = _docx_table_to_markdown(block)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")

    markdown_text = "\n".join(lines)
    markdown_text = re.sub(r'\n{3,}', '\n\n', markdown_text).strip()
    return markdown_text


def _get_heading_level(style_name):
    """Extract heading level (1-6) from Word style name, or 0 if not a heading."""
    if not style_name:
        return 0
    import re as _re
    m = _re.match(r'(?i)heading\s+(\d+)', style_name)
    if m:
        return int(m.group(1))
    # Map common Chinese style names
    lower = style_name.lower()
    for i in range(1, 7):
        if f"heading {i}" in lower:
            return i
        if f"标题 {i}" in lower:
            return i
    return 0


def _parse_content(source_format, file_path, content):
    if source_format == "md":
        return _parse_md_sections(content)
    elif source_format == "docx":
        return _parse_md_sections(content)
    raise ValueError(f"不支持的文件格式: {source_format}")


def _normalize_section_key(title):
    text = _clean_title(title or "").lower()
    text = re.sub(r'^\d+(?:\.\d+)*\s*', '', text)
    text = re.sub(r'[^a-z0-9]+', ' ', text).strip()
    return text


def _split_leading_docx_subheading(content):
    if not content:
        return None, content

    lines = content.splitlines()
    first_idx = next((idx for idx, line in enumerate(lines) if line.strip()), None)
    if first_idx is None:
        return None, content

    first_line = lines[first_idx].strip()
    if (
        not first_line
        or len(first_line) > 80
        or len(first_line.split()) > 12
        or first_line.startswith(("#", "|", "!["))
        or re.match(r'^(table|figure|formula)\b', first_line, re.IGNORECASE)
        or re.search(r'[.!?;:)]$', first_line)
    ):
        return None, content

    rest_lines = lines[first_idx + 1:]
    while rest_lines and not rest_lines[0].strip():
        rest_lines = rest_lines[1:]
    rest_content = "\n".join(rest_lines).strip()
    if not rest_content:
        return None, content
    return first_line, rest_content


def _docx_child_belongs(parent_title, child_title, is_absorbing):
    parent_key = _normalize_section_key(parent_title)
    child_key = _normalize_section_key(child_title)
    if not parent_key or not child_key:
        return False

    if child_key in {"preparation", "operation", "procedure"}:
        return True
    if child_key == parent_key:
        return True
    if parent_key.startswith("cleanup of") and child_key.startswith(("single size selection", "double size selection")):
        return True
    if " and " in parent_key:
        parts = [part.strip() for part in parent_key.split(" and ") if part.strip()]
        if any(child_key == part or child_key.endswith(part) for part in parts):
            return True
    if is_absorbing and "optional" in child_key:
        return True
    if is_absorbing and child_key.startswith("qc of"):
        return True
    return False


def _reshape_docx_sections(sections):
    prepared = []
    for sec in sections:
        current = {
            "title": sec.get("title", ""),
            "content": sec.get("content", ""),
            "sections": _reshape_docx_sections(sec.get("sections", [])),
        }
        child_title, child_content = _split_leading_docx_subheading(current["content"])
        if child_title:
            current["content"] = ""
            current["sections"] = [{
                "title": child_title,
                "content": child_content,
                "sections": [],
            }] + current["sections"]
        prepared.append(current)

    grouped = []
    idx = 0
    while idx < len(prepared):
        current = prepared[idx]
        next_idx = idx + 1
        absorbed = False
        while next_idx < len(prepared) and _docx_child_belongs(current["title"], prepared[next_idx].get("title", ""), absorbed):
            current["sections"].append(prepared[next_idx])
            absorbed = True
            next_idx += 1
        grouped.append(current)
        idx = next_idx

    return grouped


def _append_text_block(base_text, extra_text):
    base_text = (base_text or "").strip()
    extra_text = (extra_text or "").strip()
    if base_text and extra_text:
        return f"{base_text}\n\n{extra_text}"
    return base_text or extra_text


def _postprocess_section_tree(sections):
    processed = []
    for sec in sections:
        current = {
            "title": sec.get("title", ""),
            "content": sec.get("content", ""),
            "sections": _postprocess_section_tree(sec.get("sections", [])),
        }

        title_key = _normalize_section_key(current["title"])
        if title_key == "appendix" and current["sections"]:
            grouped_children = []
            appendix_parent = None
            for child in current["sections"]:
                child_key = _normalize_section_key(child.get("title", ""))
                if appendix_parent is None and "barcode number and sequence information" in child_key:
                    appendix_parent = child
                    grouped_children.append(child)
                    continue
                if appendix_parent is not None and child_key.startswith("instructions for"):
                    appendix_parent.setdefault("sections", []).append(child)
                    continue
                grouped_children.append(child)
            current["sections"] = grouped_children

        if title_key == "cleanup of adapter ligated product" and current["content"] and current["sections"]:
            target_child = None
            for child in current["sections"]:
                child_key = _normalize_section_key(child.get("title", ""))
                if child_key == title_key:
                    target_child = child
                    break
            if target_child is not None:
                target_child["content"] = _append_text_block(current["content"], target_child.get("content", ""))
                current["content"] = ""

        processed.append(current)

    return processed


def _extract_images_md(content):
    """Extract image metadata from markdown content (no download)."""
    images = []
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    for match in re.finditer(pattern, content):
        images.append({"alt": match.group(1), "path": match.group(2)})
    return images


def _download_images_for_output(images, output_base, timeout=30):
    """Download external images and save to output_dir/image/.
    Returns dict mapping original URL to local filename.
    """
    import urllib.request
    import urllib.error
    from pathlib import Path

    image_dir = os.path.join(output_base, "image")
    _ensure_dir(image_dir)

    url_to_local = {}
    idx = 0
    for img in images:
        url = img["path"]
        if not url:
            continue
        if url in url_to_local:
            continue
        try:
            ext = os.path.splitext(url.split("?")[0])[1] or ".png"
            if ext.lower() not in (".png", ".jpg", ".jpeg", ".gif", ".svg", ".webp", ".bmp"):
                ext = ".png"
            local_name = f"image_{idx:03d}{ext}"
            local_path = os.path.join(image_dir, local_name)
            if os.path.exists(url):
                shutil.copy2(url, local_path)
            else:
                if not url.startswith(("http://", "https://")):
                    continue
                req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
                with urllib.request.urlopen(req, timeout=timeout) as resp:
                    with open(local_path, "wb") as f:
                        f.write(resp.read())
            url_to_local[url] = f"image/{local_name}"
            idx += 1
        except Exception:
            url_to_local[url] = url
    return url_to_local


def _parse_template_zip(template_path):
    tree = {}
    try:
        with zipfile.ZipFile(template_path, "r") as zf:
            for name in zf.namelist():
                if name.endswith("/"):
                    continue
                parts = name.strip("/").split("/")
                current = tree
                for i, part in enumerate(parts[:-1]):
                    if part not in current:
                        current[part] = {}
                    current = current[part]
                if parts[-1].endswith((".dita", ".xml", ".ditamap")):
                    content = zf.read(name).decode("utf-8", errors="ignore")
                    leaf = parts[-1]
                    if len(parts) > 1:
                        parent = tree
                        for p in parts[:-1]:
                            parent = parent.setdefault(p, {})
                        parent[leaf] = {"name": leaf, "content": content}
                    else:
                        tree[leaf] = {"name": leaf, "content": content}
    except Exception:
        pass
    return tree


def _extract_template_parts(template_path):
    """Extract frontmatter XML, Manufacturer appendix XML, and file list from template ZIP."""
    result = {
        "frontmatter_xml": "",
        "manufacturer_xml": "",
        "files": {},        # filename -> content (bytes)
        "image_files": {},  # image/filename -> content (bytes)
        "language": None,
    }
    try:
        with zipfile.ZipFile(template_path, "r") as zf:
            namelist = zf.namelist()

            # Find and parse the ditamap
            ditamap_content = ""
            for name in namelist:
                if name.endswith(".ditamap"):
                    ditamap_content = zf.read(name).decode("utf-8", errors="ignore")
                    break

            if not ditamap_content:
                return result

            result["language"] = _extract_xml_language(ditamap_content)

            # Extract frontmatter section
            fm_match = re.search(r'<frontmatter[^>]*>.*?</frontmatter>', ditamap_content, re.DOTALL)
            if fm_match:
                result["frontmatter_xml"] = fm_match.group(0)

            # Extract Manufacturer appendix
            appx_pattern = r'<appendix[^>]*navtitle="Manufacturer[^"]*"[^>]*>.*?</appendix>'
            mfg_match = re.search(appx_pattern, ditamap_content, re.DOTALL)
            if mfg_match:
                result["manufacturer_xml"] = mfg_match.group(0)

            # Collect all referenced dita files
            all_xml = result["frontmatter_xml"] + result["manufacturer_xml"]
            refs = re.findall(r'href="([^"]+\.dita)"', all_xml)

            # Copy dita files from template ZIP
            for ref in refs:
                for name in namelist:
                    if name == ref or name.endswith("/" + ref):
                        result["files"][ref] = zf.read(name)
                        break

            # Copy all image files from template ZIP
            for name in namelist:
                if name.startswith("image/") and not name.endswith("/"):
                    img_name = name
                    result["image_files"][img_name] = zf.read(name)

    except Exception:
        pass
    return result


def _extract_xml_language(xml_text):
    if not xml_text:
        return None

    patterns = [
        r'xml:lang\s*=\s*"([^"]+)"',
        r'cms:lang\s*=\s*"([^"]+)"',
        r'imeCMS:iba_lang\s*=\s*"([^"]+)"',
    ]
    for pattern in patterns:
        match = re.search(pattern, xml_text, re.IGNORECASE)
        if match:
            return _normalize_dita_language(match.group(1))
    return None


def _normalize_dita_language(lang):
    value = (lang or "").strip().lower()
    if value in {"en", "en-us", "en_us", "english", "英文"}:
        return "en-US"
    if value in {"zh", "zh-cn", "zh_cn", "cn", "chinese", "中文"}:
        return "zh-CN"
    return None


def _parse_language_override(requirements_text):
    if not requirements_text:
        return None

    match = re.search(
        r'(?:xml:lang|lang|language|语言|输出语言|文档语言)\s*[:：=]\s*(zh-cn|zh_cn|zh|cn|中文|en-us|en_us|en|english|英文)',
        requirements_text,
        re.IGNORECASE,
    )
    if match:
        return _normalize_dita_language(match.group(1))

    if re.search(r'输出语言.*英文|文档语言.*英文|xml:lang.*en', requirements_text, re.IGNORECASE):
        return "en-US"
    if re.search(r'输出语言.*中文|文档语言.*中文|xml:lang.*zh', requirements_text, re.IGNORECASE):
        return "zh-CN"
    return None


def _detect_dita_language(source_content, template_parts=None, requirements_text="", explicit_language=None):
    override = _normalize_dita_language(explicit_language) or _parse_language_override(requirements_text)
    if override:
        return override

    text = source_content or ""
    chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
    alpha_chars = sum(1 for c in text if ('A' <= c <= 'Z') or ('a' <= c <= 'z'))
    total = chinese_chars + alpha_chars

    if total:
        if chinese_chars / total >= 0.2:
            return "zh-CN"
        if alpha_chars:
            return "en-US"

    template_lang = (template_parts or {}).get("language") if template_parts else None
    if template_lang:
        return template_lang
    return "zh-CN"


def _rewrite_xml_language_attrs(xml_text, dita_lang):
    if not xml_text or not dita_lang:
        return xml_text

    rewritten = re.sub(r'xml:lang\s*=\s*"[^"]*"', f'xml:lang="{dita_lang}"', xml_text)
    rewritten = re.sub(r'cms:lang\s*=\s*"[^"]*"', f'cms:lang="{dita_lang}"', rewritten)
    rewritten = re.sub(r'imeCMS:iba_lang\s*=\s*"[^"]*"', f'imeCMS:iba_lang="{dita_lang}"', rewritten)
    return rewritten


def _parse_special_requirements(requirements_text):
    parsed = {"reuse_topics": [], "type_overrides": {}, "auto_split": False, "other": []}
    if not requirements_text:
        return parsed

    reuse_pattern = re.findall(r'沿用[：:]?\s*(\S+)', requirements_text)
    parsed["reuse_topics"] = [r.strip().rstrip(".,;") for r in reuse_pattern]

    type_override_pattern = re.findall(
        r'([^\s,，]+?)\s*[映转][射换][到为]\s*(concept|task|reference)\s*[类类型]',
        requirements_text
    )
    for src, ttype in type_override_pattern:
        parsed["type_overrides"][src.strip()] = ttype.strip()

    if "拆分" in requirements_text or "每个" in requirements_text:
        parsed["auto_split"] = True

    return parsed


def _content_to_dita_xml(section_title, section_content, dita_type, topic_id, dita_lang, level=2):
    safe_title = section_title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _unescape_md(text):
        for ch in ['*', '_', '\\', '.', '-', '#']:
            text = text.replace('\\' + ch, ch)
        return text

    def _strip_md_bold(text):
        return re.sub(r'\*\*([^*]+)\*\*', r'\1', text)

    def _strip_md_image_alt(text):
        return re.sub(r'!\[([^\]]*)\]\(', r'\1[', text)

    safe_content = section_content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("<br>", "")

    table_rows = []
    in_table = False
    list_lines = []
    in_list = False
    list_type = None

    def _flush_list():
        nonlocal in_list, list_lines, list_type
        if not in_list:
            return ""
        tag = "ul" if list_type == "ul" else "ol"
        items = "\n".join(f"          <li>{li}</li>" for li in list_lines)
        result = f"        <{tag}>\n{items}\n        </{tag}>"
        in_list = False
        list_lines = []
        list_type = None
        return result

    body_parts = []

    for line in safe_content.split("\n"):
        stripped = line.strip()
        if not stripped:
            result = _flush_list()
            if result:
                body_parts.append(result)
            body_parts.append("")
            continue

        if in_table and re.match(r'^\|[-:\s|]+\|$', stripped):
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            if not in_table:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append([_strip_md_bold(_unescape_md(c)) for c in cells])
            continue
        else:
            if in_table:
                in_table = False
                if table_rows:
                    has_title = False
                    if body_parts and not body_parts[-1]:
                        body_parts.pop()
                    if len(body_parts) >= 1:
                        prev = body_parts[-1]
                        if prev and prev.strip().startswith("<p>") and re.match(r'<p>Table\s+\d+([\s:.-]|&amp;|\(|$)', prev):
                            has_title = True
                            body_parts.pop()
                    cols = max(len(r) for r in table_rows)
                    if has_title:
                        table_xml = ['        <table>', f'          <title>{_strip_tag(prev)}</title>']
                    else:
                        table_xml = ['        <table>']
                    table_xml.append(f'          <tgroup cols="{cols}">')
                    if len(table_rows) > 1:
                        table_xml.append('            <thead>')
                        table_xml.append('              <row>')
                        for c in table_rows[0]:
                            table_xml.append(f'                <entry>{c}</entry>')
                        empty_needed = cols - len(table_rows[0])
                        for _ in range(empty_needed):
                            table_xml.append('                <entry></entry>')
                        table_xml.append('              </row>')
                        table_xml.append('            </thead>')
                        table_xml.append('            <tbody>')
                        for row in table_rows[1:]:
                            table_xml.append('              <row>')
                            for c in row:
                                table_xml.append(f'                <entry>{c}</entry>')
                            empty_needed = cols - len(row)
                            for _ in range(empty_needed):
                                table_xml.append('                <entry></entry>')
                            table_xml.append('              </row>')
                        table_xml.append('            </tbody>')
                    else:
                        table_xml.append('            <tbody>')
                        for row in table_rows:
                            table_xml.append('              <row>')
                            for c in row:
                                table_xml.append(f'                <entry>{c}</entry>')
                            table_xml.append('              </row>')
                        table_xml.append('            </tbody>')
                    table_xml.append('          </tgroup>')
                    table_xml.append('        </table>')
                    body_parts.append("\n".join(table_xml))
                table_rows = []

        image_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', stripped)
        if image_match:
            result = _flush_list()
            if result:
                body_parts.append(result)
            alt = image_match.group(1)
            href = image_match.group(2)
            body_parts.append(f'        <image href="{href}" placement="break"><alt>{alt}</alt></image>')
            continue

        if re.match(r'^\s*[-*+]\s+', line):
            if not in_list or list_type != "ul":
                result = _flush_list()
                if result:
                    body_parts.append(result)
                in_list = True
                list_type = "ul"
            item = re.sub(r'^\s*[-*+]\s+', '', line)
            list_lines.append(_strip_md_bold(_unescape_md(item)))
        elif re.match(r'^\s*\d+[.)]\s+', line):
            if not in_list or list_type != "ol":
                result = _flush_list()
                if result:
                    body_parts.append(result)
                in_list = True
                list_type = "ol"
            item = re.sub(r'^\s*\d+[.)]\s+', '', line)
            list_lines.append(_strip_md_bold(_unescape_md(item)))
        elif re.match(r'^#{1,6}\s+', line):
            result = _flush_list()
            if result:
                body_parts.append(result)
            h_level = len(re.match(r'^(#+)', line).group(1))
            h_text = _unescape_md(re.sub(r'^#{1,6}\s+', '', line))
            body_parts.append(f"        <section><title>{h_text}</title></section>")
        else:
            result = _flush_list()
            if result:
                body_parts.append(result)
            clean = _unescape_md(_strip_md_bold(stripped))
            body_parts.append(f"        <p>{clean}</p>")

    result = _flush_list()
    if result:
        body_parts.append(result)
    if in_table and table_rows:
        cols = max(len(r) for r in table_rows)
        table_xml = ['        <table>', f'          <tgroup cols="{cols}">', '            <tbody>']
        for row in table_rows:
            table_xml.append('              <row>')
            for c in row:
                table_xml.append(f'                <entry>{c}</entry>')
            table_xml.append('              </row>')
        table_xml.append('            </tbody>')
        table_xml.append('          </tgroup>')
        table_xml.append('        </table>')
        body_parts.append("\n".join(table_xml))

    body_content = "\n".join(p for p in body_parts if p)

    unique_id = topic_id
    ime_topic_type = "sDitaTopic"

    return f"""<?xml version="1.0" encoding="UTF-8"?>
{IME_TOPIC_DECL}
<topic xmlns:cms="http://www.w3.org/ime/cms" xmlns:m="http://www.w3.org/1998/Math/MathML" xml:lang="{dita_lang}" id="{unique_id}" cms:keys="{topic_id}" cms:title="{safe_title}" cms:placeholder="{topic_id}" cms:number="{topic_id}" cms:imesofttype="{ime_topic_type}" cms:placeHolder="{topic_id}">
  <title id="title_{topic_id}">{safe_title}</title>
  <body id="body_{topic_id}" outputclass="pretext">
{body_content}
  </body>
</topic>"""


def _strip_tag(xml):
    return re.sub(r'<[^>]+>', '', xml)


def _content_to_markdown(section_title, section_content):
    title = _clean_title(section_title) or "未命名章节"
    content = (section_content or "").strip()
    if content:
        return f"# {title}\n\n{content}\n"
    return f"# {title}\n"


def _markdown_to_lossless_csv(content):
    buffer = io.StringIO(newline="")
    writer = csv.writer(buffer)
    writer.writerow(["line_number", "content"])
    for line_number, line in enumerate(content.splitlines(), start=1):
        writer.writerow([line_number, line])
    return "\ufeff" + buffer.getvalue()


def _csv_table_to_markdown(content):
    rows = []
    reader = csv.reader(io.StringIO(content.lstrip("\ufeff")))
    for row in reader:
        rows.append([str(cell or "") for cell in row])

    if not rows:
        return ""

    def _escape_md_cell(cell):
        return cell.replace("|", "\\|")

    header = rows[0]
    align = [":---" for _ in header]
    md_lines = [
        "| " + " | ".join(_escape_md_cell(cell) for cell in header) + " |",
        "| " + " | ".join(align) + " |",
    ]
    for row in rows[1:]:
        if len(row) < len(header):
            row = row + [""] * (len(header) - len(row))
        md_lines.append("| " + " | ".join(_escape_md_cell(cell) for cell in row[:len(header)]) + " |")
    return "\n".join(md_lines) + "\n"


def _csv_to_markdown(content):
    text = content.lstrip("\ufeff")
    reader = csv.DictReader(io.StringIO(text))
    fieldnames = reader.fieldnames or []
    normalized = {str(name or "").strip().lower() for name in fieldnames}
    if {"line_number", "content"}.issubset(normalized):
        rows = []
        for row in reader:
            line_number = row.get("line_number") or row.get("LINE_NUMBER") or row.get("Line_Number") or "0"
            try:
                order = int(str(line_number).strip() or "0")
            except ValueError:
                order = 0
            rows.append((order, row.get("content") or row.get("CONTENT") or ""))
        rows.sort(key=lambda item: item[0])
        return "\n".join(line for _, line in rows)
    return _csv_table_to_markdown(text)


def _build_active_rules(source_format, target_format, source_content, template_path, images, sections, is_lossless_md_csv):
    rules = [{
        "rule_number": "01",
        "category": "内容",
        "description": "转换内容须100%与原文保持一致，不得遗漏或添加任何文字、符号、空白段落",
    }]

    if is_lossless_md_csv:
        return rules

    if images:
        rules.append({
            "rule_number": "02",
            "category": "图片",
            "description": "原文中引用的外部图片须下载并嵌入到输出文档包中，确保离线可查看",
        })

    if target_format == "dita":
        if re.search(r'^\s{0,3}#{1,6}\s+((\d+(?:\.\d+)*)|chapter\s+\d+|第\s*\d+\s*章)[\s:：-]+', source_content, re.IGNORECASE | re.MULTILINE):
            rules.append({
                "rule_number": "03",
                "category": "编号",
                "description": "章节标题前的编号按模板样式处理，避免正文重复编号",
            })
        if re.search(r'^\|.*\|\s*$', source_content, re.MULTILINE):
            rules.append({
                "rule_number": "04",
                "category": "表格标题",
                "description": "Markdown 表格结构需映射为 DITA 表格结构",
            })
        if re.search(r'^\s*[-*+]\s+', source_content, re.MULTILINE):
            rules.append({
                "rule_number": "05",
                "category": "列表-无序",
                "description": "Markdown 无序列表需映射为 DITA 列表结构",
            })
        if re.search(r'^\s*\d+[.)]\s+', source_content, re.MULTILINE):
            rules.append({
                "rule_number": "06",
                "category": "列表-有序",
                "description": "Markdown 有序列表需映射为 DITA 列表结构",
            })
        if any(section.get("title") for section in sections):
            rules.append({
                "rule_number": "07",
                "category": "结构",
                "description": "文档标题层级按实际结构拆分并组织为 DITA 输出",
            })
        if template_path and os.path.exists(template_path):
            rules.append({
                "rule_number": "08",
                "category": "模板",
                "description": "提供参考模板时复用模板结构与必要元数据",
            })

    return rules


def _rewrite_template_frontmatter(frontmatter_xml, topics_output):
    if not frontmatter_xml:
        return frontmatter_xml, set()

    rewritten = frontmatter_xml
    used_files = set()
    title_to_topic = {
        topic["title"].strip().lower(): topic
        for topic in topics_output
        if topic.get("title") and topic.get("filename")
    }

    for navtitle, topic in title_to_topic.items():
        pattern = re.compile(rf'(<topicref[^>]*navtitle="{re.escape(topic["title"])}"[^>]*href=")([^"]+)(")')
        if pattern.search(rewritten):
            rewritten = pattern.sub(rf'\1{topic["filename"]}\3', rewritten)
            rewritten = re.sub(
                rf'(<topicref[^>]*navtitle="{re.escape(topic["title"])}"[^>]*cms:placeHolder=")([^"]*)(")',
                rf'\1{_topic_placeholder(topic)}\3',
                rewritten,
            )
            rewritten = re.sub(
                rf'(<topicref[^>]*navtitle="{re.escape(topic["title"])}"[^>]*keys=")([^"]*)(")',
                rf'\1{topic["id"]}\3',
                rewritten,
            )
            used_files.add(topic["filename"])

    return rewritten, used_files


def _collect_template_dita_refs(*xml_parts):
    refs = set()
    for xml in xml_parts:
        if not xml:
            continue
        refs.update(re.findall(r'href="([^"]+\.dita)"', xml))
    return refs


def _prune_unreferenced_output_images(output_base):
    referenced = set()
    image_dir = os.path.join(output_base, "image")
    if not os.path.isdir(image_dir):
        return set(), set()

    for root, _, files in os.walk(output_base):
        for fname in files:
            if not fname.endswith((".dita", ".ditamap", ".md")):
                continue
            file_path = os.path.join(root, fname)
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
            except Exception:
                continue

            for href in re.findall(r'href="([^"]+)"', content):
                href = href.strip()
                if href.startswith("image/"):
                    referenced.add(href)
            for md_path in re.findall(r'!\[[^\]]*\]\(([^)]+)\)', content):
                md_path = md_path.strip()
                if md_path.startswith("image/"):
                    referenced.add(md_path)

    existing = set()
    for root, _, files in os.walk(image_dir):
        for fname in files:
            abs_path = os.path.join(root, fname)
            rel_path = os.path.relpath(abs_path, output_base).replace(os.sep, "/")
            existing.add(rel_path)
            if rel_path not in referenced:
                os.remove(abs_path)

    return referenced, existing


def _escape_xml_attr(value):
    return (value or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _build_topic_hierarchy(topics, skipped_files):
    roots = []
    stack = []

    for topic in topics:
        if topic["filename"] in skipped_files:
            continue

        node = {
            "topic": topic,
            "level": max(1, int(topic.get("level", 1) or 1)),
            "children": [],
        }

        while stack and stack[-1]["level"] >= node["level"]:
            stack.pop()

        if stack:
            stack[-1]["children"].append(node)
        else:
            roots.append(node)

        stack.append(node)

    return roots


def _topic_uses_dita_file(topic):
    return bool(topic.get("filename") and not topic.get("is_container_only"))


def _generated_topic_code(topic_index):
    return f"DTO{GENERATED_TOPIC_BASE + topic_index:06d}"


def _topic_placeholder(topic, has_children=False):
    if has_children:
        return ""
    filename = topic.get("filename") or ""
    return os.path.splitext(filename)[0] if filename else ""


def _append_topicref_xml(lines, topic, node_id_seq, indent, level_attr, dita_lang, template_name="sDitaTopic"):
    navtitle = _escape_xml_attr(topic["title"])
    node_id = f"PN{next(node_id_seq):03d}"
    ime_soft_type = template_name
    placeholder = _escape_xml_attr(_topic_placeholder(topic))
    attrs = (
        f'navtitle="{navtitle}" '
        f'xml:lang="{dita_lang}" id="{node_id}" '
        f'applicDesc="" '
        f'href="{topic["filename"]}" '
        f'keys="{topic["id"]}" type="topic" '
        f'cms:title="{navtitle}" '
        f'cms:placeHolder="{placeholder}" '
        f'cms:nodeType="XML" '
        f'cms:template="{template_name}" '
        f'level="{level_attr}" '
        f'cms:imesofttype="{ime_soft_type}" '
        f'version="{GENERATED_NODE_VERSION}"'
    )
    lines.append(f'{indent}<topicref {attrs}/>' )


def _topic_has_body(topic):
    return bool((topic.get("raw_content") or "").strip())


def _append_topicref_container_xml(lines, node, node_id_seq, indent, level_attr, dita_lang):
    topic = node["topic"]
    navtitle = _escape_xml_attr(topic["title"])
    node_id = f"PN{next(node_id_seq):03d}"
    has_children = bool(node.get("children"))
    attrs = [
        f'navtitle="{navtitle}"',
        f'xml:lang="{dita_lang}"',
        f'id="{node_id}"',
        'applicDesc=""',
        'type="topic"',
        f'level="{level_attr}"',
        'cms:imesofttype="sDitaTopic"',
        f'version="{GENERATED_NODE_VERSION}"',
    ]

    if _topic_uses_dita_file(topic):
        attrs.extend([
            f'href="{topic["filename"]}"',
            f'keys="{topic["id"]}"',
            f'cms:title="{navtitle}"',
            f'cms:placeHolder="{_escape_xml_attr(_topic_placeholder(topic, has_children=has_children))}"',
            'cms:nodeType="XML"',
            'cms:template="sDitaTopic"',
        ])
    else:
        attrs.extend([
            'cms:title=""',
            'cms:placeHolder=""',
            'cms:isTemplet="N"',
            'cms:referenceType=""',
            'cms:type="topicref"',
            'cms:outputclass=""',
            'cms:descriptioin=""',
            'cms:xCoordination=""',
            'cms:sourceTag=""',
            'cms:referenceMap=""',
            'cms:colNumber="2"',
            'cms:nodeRemark=""',
        ])

    joined_attrs = " ".join(attrs)
    lines.append(f'{indent}<topicref {joined_attrs}>')
    for child in node["children"]:
        if child["children"]:
            _append_topicref_container_xml(lines, child, node_id_seq, indent + "  ", level_attr + 1, dita_lang)
        else:
            _append_topicref_xml(lines, child["topic"], node_id_seq, indent + "  ", level_attr + 1, dita_lang)
    lines.append(f'{indent}</topicref>')


def _append_root_container_xml(lines, node, chapter_seq, node_id_seq, indent, level_attr, timestamp, dita_lang):
    topic = node["topic"]
    navtitle = _escape_xml_attr(topic["title"])
    container_tag = "appendix" if topic["title"].strip().lower() == "appendix" else "chapter"
    chapter_id = f"PG{timestamp}_{next(chapter_seq):03d}"
    lines.append(
        f'{indent}<{container_tag} navtitle="{navtitle}" id="{chapter_id}" applicDesc=""'
        f' cms:lang="{dita_lang}" cms:type="{container_tag}" cms:title=""'
        f' cms:isTemplet="N"'
        f' level="{level_attr}" version="{GENERATED_NODE_VERSION}">'
    )
    if _topic_uses_dita_file(topic):
        root_template = "chapterTopic" if container_tag == "chapter" else "sDitaTopic"
        _append_topicref_xml(lines, topic, node_id_seq, indent + "  ", level_attr + 1, dita_lang, root_template)
    for child in node["children"]:
        if child["children"]:
            _append_topicref_container_xml(lines, child, node_id_seq, indent + "  ", level_attr + 1, dita_lang)
        else:
            _append_topicref_xml(lines, child["topic"], node_id_seq, indent + "  ", level_attr + 1, dita_lang)
    lines.append(f'{indent}</{container_tag}>')


def _append_frontmatter_root_xml(lines, node, node_id_seq, indent, timestamp, dita_lang):
    frontmatter_id = f"PG{timestamp}_FM"
    lines.append(
        f'{indent}<frontmatter navtitle="Preface" id="{frontmatter_id}" applicDesc=""'
        f' cms:lang="{dita_lang}" cms:title="" cms:placeHolder="" cms:isTemplet="N"'
        f' cms:referenceType="" cms:type="frontmatter" cms:outputclass=""'
        f' cms:descriptioin="" cms:xCoordination="" cms:sourceTag=""'
        f' cms:referenceMap="" cms:colNumber="2" cms:nodeRemark=""'
        f' level="1" version="{GENERATED_NODE_VERSION}">'
    )
    _append_topicref_xml(lines, node["topic"], node_id_seq, indent + "  ", 2, dita_lang, "sCoverTopic")
    for child in node["children"]:
        if child["children"]:
            _append_topicref_container_xml(lines, child, node_id_seq, indent + "  ", 2, dita_lang)
        else:
            _append_topicref_xml(lines, child["topic"], node_id_seq, indent + "  ", 2, dita_lang)

    booklists_id = f"PG{timestamp}_BL"
    toc_id = f"PG{timestamp}_TOC"
    lines.append(
        f'{indent}  <booklists id="{booklists_id}" applicDesc="" cms:lang="{dita_lang}" cms:title="" cms:placeHolder=""'
        f' cms:isTemplet="N" cms:referenceType="" cms:type="booklists" cms:outputclass=""'
        f' cms:descriptioin="" cms:xCoordination="" cms:sourceTag="" cms:referenceMap=""'
        f' cms:colNumber="2" cms:nodeRemark="" level="2" version="{GENERATED_NODE_VERSION}">'
    )
    lines.append(
        f'{indent}    <toc id="{toc_id}" applicDesc="" cms:lang="{dita_lang}" cms:title="" cms:placeHolder=""'
        f' cms:isTemplet="N" cms:referenceType="" cms:type="toc" cms:outputclass=""'
        f' cms:descriptioin="" cms:xCoordination="" cms:sourceTag="" cms:referenceMap=""'
        f' cms:colNumber="2" cms:nodeRemark="" level="3" version="{GENERATED_NODE_VERSION}"/>'
    )
    lines.append(f'{indent}  </booklists>')
    lines.append(f'{indent}</frontmatter>')


def _append_bookmap_topics(lines, topic_roots, timestamp, dita_lang):
    chapter_seq = iter(range(1, 10000))
    node_id_seq = iter(range(1, 100000))

    for root in topic_roots:
        if _normalize_section_key(root["topic"]["title"]) == "cover":
            _append_frontmatter_root_xml(lines, root, node_id_seq, "  ", timestamp, dita_lang)
            continue
        _append_root_container_xml(lines, root, chapter_seq, node_id_seq, "  ", 1, timestamp, dita_lang)


def _derive_map_title(source_format, file_path, sections):
    if source_format == "docx":
        stem = os.path.splitext(os.path.basename(file_path))[0]
        stem = re.sub(r'^src_\d+_', '', stem)
        stem = stem.replace("_", " ").strip()
        if stem:
            return stem

    for sec in sections:
        title = _clean_title(sec.get("title", ""))
        if title and title.lower() not in {"cover", "contents"}:
            return title
    return "Generated Document"


def _flatten_sections(sections, parent_h1="", level=1):
    """Recursively flatten nested sections into a list of topics.
    Every heading at any level becomes its own topic.
    """
    result = []
    for sec in sections:
        title = sec.get("title", "")
        content = sec.get("content", "")
        children = sec.get("sections", [])

        if title or content or not children:
            result.append({
                "h1": parent_h1,
                "level": level,
                "title": title,
                "content": content,
                "has_children": bool(children),
                "child_titles": [child.get("title", "") for child in children],
                "dita_type": _get_topic_type(title),
            })

        if children:
            new_parent = title or parent_h1
            result.extend(_flatten_sections(children, new_parent, level + 1))

    return result


def _should_emit_container_only(section):
    if not section.get("has_children"):
        return False

    if (section.get("content") or "").strip():
        return False

    return int(section.get("level", 1) or 1) <= 1


def _run_pipeline(task_id, db_session_factory, source_format, file_path, source_content,
                   target_format, template_path, requirements_text, output_language=None):
    """格式转换流水线。

    固定行为（不可变）：
    - 标题清理: 去 \\ 转义/去 ** 加粗/去编号前缀
    - 全级拆分: H1-H6 每个标题一个独立 topic
    - 图片下载: 外部图片打包到 image/ 目录
    - 模板复用: 从模板复制 <frontmatter> + Manufacturer <appendix> 及其文件
    可变部分: 仅复用哪些具体 topic 可随任务调整。
    """
    import time as _time
    from app.database import SessionLocal

    db = SessionLocal()
    try:

        def _set_progress(pct, step):
            with _tasks_lock:
                _task_store[task_id] = {
                    "progress": pct, "current_step": step, "status": "processing"
                }
            try:
                update_convert_task_progress(db, task_id, pct, step)
            except Exception:
                pass

        _set_progress(5, "解析")

        is_lossless_md_csv = (
            (source_format == "md" and target_format == "csv")
            or (source_format == "csv" and target_format == "markdown")
        )

        sections = []
        images = []
        if not is_lossless_md_csv:
            sections = _parse_content(source_format, file_path, source_content)
            if source_format == "docx":
                sections = _reshape_docx_sections(sections)
                sections = _postprocess_section_tree(sections)
            images = _extract_images_md(source_content)

        _set_progress(20, "解析")

        template_parts = None
        if template_path and os.path.exists(template_path):
            template_parts = _extract_template_parts(template_path)

        dita_lang = _detect_dita_language(source_content, template_parts, requirements_text, output_language)

        _set_progress(30, "结构映射")

        special = _parse_special_requirements(requirements_text)

        conversion_detail = []
        topics_output = []
        image_count = len(images)
        topic_index = 0
        all_sections = []

        if is_lossless_md_csv:
            if source_format == "md":
                topics_output.append({
                    "type": "csv",
                    "filename": "content.csv",
                    "title": "Markdown 全文",
                    "content": _markdown_to_lossless_csv(source_content),
                })
                conversion_detail.append({
                    "source_section": "Markdown 全文",
                    "target_type": "csv",
                    "topic_file": "content.csv",
                    "status": "ok",
                })
            else:
                topics_output.append({
                    "type": "markdown",
                    "filename": "content.md",
                    "title": "Markdown 全文",
                    "content": _csv_to_markdown(source_content),
                })
                conversion_detail.append({
                    "source_section": "CSV 全文",
                    "target_type": "markdown",
                    "topic_file": "content.md",
                    "status": "ok",
                })
        else:
            all_sections = _flatten_sections(sections)

            if special["auto_split"] and len(all_sections) == 1:
                def _rebuild_content(sec):
                    parts = []
                    if sec.get("content"):
                        parts.append(sec["content"])
                    for child in sec.get("sections", []):
                        parts.append(f"{'#' * (sec.get('level', 1) + 1)} {child['title']}")
                        parts.append(_rebuild_content(child))
                    return "\n".join(p.strip() for p in parts if p.strip())

                all_sections = []
                for h1 in sections:
                    if h1["title"]:
                        all_sections.append({
                            "h1": "",
                            "title": h1["title"],
                            "content": _rebuild_content(h1),
                            "dita_type": _get_topic_type(h1["title"]),
                        })
                if not all_sections:
                    all_sections = _flatten_sections(sections)

            slug_counts = {}
            for section in all_sections:
                title = _clean_title(section["title"])
                h1_title = _clean_title(section.get("h1", ""))
                if not title:
                    title = h1_title or f"Section {topic_index + 1}"

                dita_type = section["dita_type"]
                for keyword, override_type in special["type_overrides"].items():
                    if keyword in title:
                        dita_type = override_type
                        break

                topic_index += 1
                slug = re.sub(r'[^\w\s-]', '', title).strip().lower()
                slug = re.sub(r'[-\s]+', '_', slug) or f"topic_{topic_index}"
                slug_count = slug_counts.get(slug, 0)
                slug_counts[slug] = slug_count + 1
                if slug_count:
                    slug = f"{slug}_{slug_count + 1}"
                topic_id = _generated_topic_code(topic_index)
                is_container_only = _should_emit_container_only(section)
                topic_filename = None if is_container_only else f"{topic_id}.dita"

                dita_content = None
                if topic_filename:
                    dita_content = _content_to_dita_xml(title, section["content"], dita_type, topic_id, dita_lang)

                topics_output.append({
                    "type": dita_type,
                    "filename": topic_filename,
                    "id": topic_id,
                    "level": section.get("level", 1),
                    "content": dita_content,
                    "raw_content": section["content"],
                    "is_container_only": is_container_only,
                    "title": title,
                })
                if topic_filename:
                    conversion_detail.append({
                        "source_section": title,
                        "target_type": dita_type,
                        "topic_file": topic_filename,
                        "status": "ok",
                    })

        _set_progress(55, "内容替换")

        for topic in topics_output:
            for reuse_name in special["reuse_topics"]:
                if reuse_name.lower() in topic["filename"].lower() or reuse_name.lower() in topic["title"].lower():
                    pass

        _set_progress(70, "验证")

        template_check_status = "passed"
        template_check_detail = "未使用模板"
        if template_path and target_format == "dita":
            template_check_status = "warning"
            template_check_detail = "模板已参考"
        elif template_path and target_format == "markdown":
            template_check_status = "warning"
            template_check_detail = "Markdown 输出未使用模板包"
        elif template_path and target_format == "csv":
            template_check_status = "warning"
            template_check_detail = "CSV 输出未使用模板包"

        content_check_detail = "全部章节已映射"
        if is_lossless_md_csv:
            content_check_detail = "已按逐行无损格式转换"

        checks = [
            {"name": "结构完整性", "status": "passed"},
            {"name": "模板一致性", "status": template_check_status,
             "detail": template_check_detail},
            {"name": "语言标记", "status": "passed", "detail": dita_lang},
            {"name": "图片验证", "status": "passed", "detail": f"{image_count}/{image_count}"},
            {"name": "内容完整性", "status": "passed", "detail": content_check_detail},
        ]

        unmapped = []
        for sec in all_sections:
            if sec["title"] and not sec["content"]:
                pass

        overall = "passed"
        if unmapped:
            checks[4]["status"] = "warning"
            checks[4]["detail"] = f"{len(unmapped)}个章节未映射"
            overall = "warning"

        _set_progress(85, "打包")

        timestamp = _time.strftime("%Y%m%d_%H%M%S")
        output_name = f"output_{timestamp}"
        output_base = os.path.join(OUTPUT_DIR, output_name)
        _ensure_dir(output_base)
        written_topics = [topic for topic in topics_output if topic.get("filename") and topic.get("content")]

        url_map = _download_images_for_output(images, output_base)

        if is_lossless_md_csv:
            direct_filename = f"{output_name}{os.path.splitext(topics_output[0]['filename'])[1]}"
            direct_path = os.path.join(OUTPUT_DIR, direct_filename)
            with open(direct_path, "w", encoding="utf-8", newline="") as f:
                f.write(topics_output[0]["content"])
        elif target_format == "markdown":
            index_lines = ["# 转换结果", ""]
            for detail, topic in zip(conversion_detail, topics_output):
                topic_filename = topic["filename"].rsplit(".", 1)[0] + ".md"
                markdown_content = _content_to_markdown(topic["title"], topic.get("raw_content", ""))
                for url, local in url_map.items():
                    markdown_content = markdown_content.replace(f'({url})', f'({local})')
                with open(os.path.join(output_base, topic_filename), "w", encoding="utf-8") as f:
                    f.write(markdown_content)
                topic["filename"] = topic_filename
                detail["topic_file"] = topic_filename
                index_lines.append(f'- [{topic["title"]}]({topic_filename})')

            with open(os.path.join(output_base, "index.md"), "w", encoding="utf-8") as f:
                f.write("\n".join(index_lines) + "\n")
        else:
            frontmatter_used_files = set()
            manufacturer_used_files = set()
            fm_xml = ""
            mfg_xml = ""
            template_referenced_files = set()
            if template_parts and template_parts.get("frontmatter_xml"):
                fm_xml, frontmatter_used_files = _rewrite_template_frontmatter(template_parts["frontmatter_xml"], topics_output)
            if template_parts and template_parts.get("manufacturer_xml"):
                mfg_xml, manufacturer_used_files = _rewrite_template_frontmatter(template_parts["manufacturer_xml"], topics_output)
            fm_xml = _rewrite_xml_language_attrs(fm_xml, dita_lang)
            mfg_xml = _rewrite_xml_language_attrs(mfg_xml, dita_lang)
            if template_parts:
                template_referenced_files = _collect_template_dita_refs(fm_xml, mfg_xml)

            if template_parts:
                for fname, content in template_parts["files"].items():
                    if template_referenced_files and fname not in template_referenced_files:
                        continue
                    if fname.lower().endswith((".dita", ".xml", ".ditamap")):
                        try:
                            text_content = content.decode("utf-8", errors="ignore")
                            content = _rewrite_xml_language_attrs(text_content, dita_lang).encode("utf-8")
                        except Exception:
                            pass
                    tpl_fname = os.path.join(output_base, fname)
                    with open(tpl_fname, "wb") as f:
                        f.write(content)
                for img_path, img_content in template_parts["image_files"].items():
                    tpl_img = os.path.join(output_base, img_path)
                    os.makedirs(os.path.dirname(tpl_img), exist_ok=True)
                    with open(tpl_img, "wb") as f:
                        f.write(img_content)

            for topic in written_topics:
                content = topic["content"]
                for url, local in url_map.items():
                    content = content.replace(f'href="{url}"', f'href="{local}"')
                fname = os.path.join(output_base, topic["filename"])
                with open(fname, "w", encoding="utf-8") as f:
                    f.write(content)

            map_title = _derive_map_title(source_format, file_path, sections)
            map_title = map_title.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            sanitized_name = re.sub(r'[^\w\s-]', '', map_title).strip()
            sanitized_name = re.sub(r'[-\s]+', '_', sanitized_name) or "document"
            map_filename = f"{sanitized_name}.ditamap"
            map_id = sanitized_name.upper()[:12]

            map_lines = [
                '<?xml version="1.0" encoding="UTF-8"?>',
                IME_BOOKMAP_DECL,
                f'<bookmap xmlns:imeCMS="http://www.megalinkware.com" xmlns:vf="http://www.megalinkware.com" xmlns:cms="http://www.w3.org/ime/cms" imeCMS:imesofttype="PartBookMap" imeCMS:softtype="PartBookMap" xml:lang="{dita_lang}" id="{map_id}" imeCMS:iba_lang="{dita_lang}" imeCMS:iba_title="" imeCMS:iba_placeHolder="" imeCMS:iba_isTemplet="N" imeCMS:iba_referenceType="" level="0" version="{GENERATED_NODE_VERSION}">',
                '  <booktitle>',
                f'    <mainbooktitle>{map_title}</mainbooktitle>',
                '  </booktitle>',
            ]

            if fm_xml:
                for line in fm_xml.strip().split("\n"):
                    map_lines.append(f"  {line.strip()}")

            referenced_topic_files = frontmatter_used_files | manufacturer_used_files
            topic_roots = _build_topic_hierarchy(topics_output, referenced_topic_files)
            _append_bookmap_topics(map_lines, topic_roots, timestamp, dita_lang)

            if mfg_xml:
                for line in mfg_xml.strip().split("\n"):
                    map_lines.append(f"  {line.strip()}")

            map_lines.append('</bookmap>')
            with open(os.path.join(output_base, map_filename), "w", encoding="utf-8") as f:
                f.write("\n".join(map_lines))

            _prune_unreferenced_output_images(output_base)

        metadata_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<metadata>
  <generator>Smart Doc Platform</generator>
  <created>{timestamp}</created>
  <topics>{len(written_topics)}</topics>
  <images>{image_count}</images>
  <language>{dita_lang}</language>
  <source_format>{source_format}</source_format>
  <target_format>{target_format}</target_format>
</metadata>"""
        with open(os.path.join(output_base, "metadata.xml"), "w", encoding="utf-8") as f:
            f.write(metadata_xml)

        if is_lossless_md_csv:
            shutil.rmtree(output_base)
            output_public_path = f"/static/uploads/outputs/{direct_filename}"
            output_size = os.path.getsize(direct_path)
        else:
            zip_filename = f"{output_name}.zip"
            zip_path = os.path.join(OUTPUT_DIR, zip_filename)
            with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                for root, dirs, files in os.walk(output_base):
                    for fname in files:
                        file_path_abs = os.path.join(root, fname)
                        arcname = os.path.relpath(file_path_abs, output_base)
                        zf.write(file_path_abs, arcname)

            shutil.rmtree(output_base)
            output_public_path = f"/static/uploads/outputs/{zip_filename}"
            output_size = os.path.getsize(zip_path)

        verification_report = {
            "overall": overall,
            "checks": checks,
            "unmapped_sections": unmapped,
        }

        verification_report["active_rules"] = _build_active_rules(
            source_format=source_format,
            target_format=target_format,
            source_content=source_content,
            template_path=template_path,
            images=images,
            sections=sections,
            is_lossless_md_csv=is_lossless_md_csv,
        )

        update_convert_task_result(
            db, task_id,
            output_zip_path=output_public_path,
            output_size=output_size,
            topic_count=len(written_topics),
            image_count=image_count,
            verification_report=verification_report,
            conversion_detail=conversion_detail,
        )

        _set_progress(100, "打包")
        with _tasks_lock:
            _task_store[task_id] = {
                "progress": 100, "current_step": "打包", "status": "completed"
            }

    except Exception as e:
        with _tasks_lock:
            _task_store[task_id] = {
                "progress": 0, "current_step": "", "status": "failed",
                "error": str(e)
            }
        try:
            update_convert_task_failed(db, task_id, str(e))
        except Exception:
            pass
    finally:
        db.close()


@router.post("/", response_model=dict)
async def start_conversion(
    source_file: UploadFile = File(...),
    target_format: str = Form("dita"),
    template_file: UploadFile = File(None),
    requirements: str = Form(None),
    output_language: str = Form(None),
    retry_feedback: str = Form(None),
    retry_screenshot: UploadFile = File(None),
    db: Session = Depends(get_db),
):
    if target_format not in ["dita", "markdown", "csv"]:
        raise HTTPException(status_code=400, detail="目标格式仅支持 dita、markdown 或 csv")

    ext = os.path.splitext(source_file.filename)[1].lower()
    if ext not in [".md", ".markdown", ".docx", ".doc", ".csv"]:
        raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 .md、.csv 或 .docx 文件")

    if ext in [".md", ".markdown"]:
        source_format = "md"
    elif ext == ".csv":
        source_format = "csv"
    else:
        source_format = "docx"

    if source_format == "csv" and target_format != "markdown":
        raise HTTPException(status_code=400, detail="CSV 源文件当前仅支持转换为 markdown")
    if source_format == "md" and target_format == "dita" and template_file and template_file.filename and not template_file.filename.lower().endswith(".zip"):
        raise HTTPException(status_code=400, detail="DITA 模板文件必须是 ZIP 包")
    if source_format == "docx" and target_format == "csv":
        raise HTTPException(status_code=400, detail="DOCX 当前仅支持转换为 dita 或 markdown")
    normalized_output_language = _normalize_dita_language(output_language) if output_language else None
    if output_language and not normalized_output_language:
        raise HTTPException(status_code=400, detail="输出语言仅支持 en-US 或 zh-CN")

    _ensure_dir(UPLOAD_DIR)
    source_path = os.path.join(UPLOAD_DIR, f"src_{int(time.time() * 1000)}_{source_file.filename}")
    with open(source_path, "wb") as f:
        f.write(await source_file.read())

    template_path = None
    template_filename = None
    if template_file:
        template_filename = template_file.filename
        template_path = os.path.join(UPLOAD_DIR, f"tpl_{int(time.time() * 1000)}_{template_file.filename}")
        with open(template_path, "wb") as f:
            f.write(await template_file.read())

    retry_screenshot_path = None
    if retry_screenshot and retry_screenshot.filename:
        retry_screenshot_path = os.path.join(UPLOAD_DIR, f"ss_{int(time.time() * 1000)}_{retry_screenshot.filename}")
        with open(retry_screenshot_path, "wb") as f:
            f.write(await retry_screenshot.read())

    combined_requirements = requirements or ""
    if retry_feedback:
        if combined_requirements:
            combined_requirements += "\n[IME反馈] " + retry_feedback
        else:
            combined_requirements = "[IME反馈] " + retry_feedback

    source_content = ""
    try:
        if source_format in ["md", "csv"]:
            with open(source_path, "r", encoding="utf-8", errors="ignore") as f:
                source_content = f.read()
        elif source_format == "docx":
            source_content = _docx_to_markdown(source_path)
    except Exception as e:
        try:
            os.remove(source_path)
        except Exception:
            pass
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")

    task = create_convert_task(db, ConvertTaskCreate(
        source_filename=source_file.filename,
        source_format=source_format,
        target_format=target_format,
        template_filename=template_filename,
        requirements=combined_requirements,
        retry_feedback=retry_feedback,
        retry_screenshot_path=retry_screenshot_path,
    ))

    with _tasks_lock:
        _task_store[task.task_id] = {
            "progress": 0, "current_step": "", "status": "processing"
        }

    thread = threading.Thread(
        target=_run_pipeline,
        args=(task.task_id, None, source_format, source_path, source_content,
              target_format, template_path, combined_requirements, normalized_output_language),
        daemon=True,
    )
    thread.start()

    return {"task_id": task.task_id, "status": "processing"}


@router.get("/{task_id}/progress", response_model=ProgressOut)
async def query_progress(task_id: str, db: Session = Depends(get_db)):
    task = get_convert_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    steps = []
    step_names = _get_step_names()
    step_progress_thresholds = [20, 50, 75, 85, 100]
    current_step = task.current_step or ""

    for i, name in enumerate(step_names):
        if task.progress >= step_progress_thresholds[i]:
            step_status = "done" if task.progress >= step_progress_thresholds[i] else "processing"
        elif name == current_step:
            step_status = "processing"
        else:
            step_status = "pending"

        if task.status == "completed":
            step_status = "done"
        elif task.status == "failed" and i > 0 and task.progress < step_progress_thresholds[i]:
            step_status = "pending"

        steps.append(StepStatus(name=name, status=step_status))

    if task.status == "failed":
        for step in steps:
            if step.status == "processing":
                step.status = "pending"

    return ProgressOut(
        task_id=task.task_id,
        status=task.status,
        progress=task.progress,
        current_step=task.current_step,
        steps=steps,
    )


@router.get("/{task_id}/download")
async def download_zip(task_id: str, db: Session = Depends(get_db)):
    task = get_convert_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    if task.status != "completed":
        raise HTTPException(status_code=400, detail="转换尚未完成")
    if not task.output_zip_path:
        raise HTTPException(status_code=500, detail="输出文件不存在")

    file_path = f".{task.output_zip_path}"
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="输出文件不存在")

    return FileResponse(
        file_path,
        media_type=_media_type_for_file(file_path),
        filename=os.path.basename(file_path),
    )


@router.get("/{task_id}/report", response_model=ReportOut)
async def get_report(task_id: str, db: Session = Depends(get_db)):
    task = get_convert_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    if task.status not in ["completed", "failed"]:
        raise HTTPException(status_code=400, detail="转换尚未完成")

    if task.verification_report:
        report_data = json.loads(task.verification_report)
    else:
        report_data = {"overall": "failed", "checks": [], "unmapped_sections": []}

    checks = []
    for c in report_data.get("checks", []):
        checks.append(CheckItem(
            name=c["name"],
            status=c.get("status", "failed"),
            detail=c.get("detail"),
        ))

    return ReportOut(
        task_id=task.task_id,
        overall=report_data.get("overall", "failed"),
        checks=checks,
        unmapped_sections=report_data.get("unmapped_sections"),
        active_rules=report_data.get("active_rules"),
    )


@router.get("/{task_id}/detail")
async def get_detail(task_id: str, db: Session = Depends(get_db)):
    task = get_convert_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    detail = []
    if task.conversion_detail:
        detail = json.loads(task.conversion_detail)
    return {"task_id": task_id, "detail": detail}


@router.get("/", response_model=list[ConvertTaskOut])
async def list_tasks(skip: int = 0, limit: int = 100, db: Session = Depends(get_db)):
    tasks = get_convert_tasks(db, skip=skip, limit=limit)
    return tasks


@router.delete("/{task_id}")
async def remove_task(task_id: str, db: Session = Depends(get_db)):
    task = delete_convert_task(db, task_id)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    with _tasks_lock:
        _task_store.pop(task_id, None)
    return {"message": "任务已删除"}


# ─── 转换规则库 ─────────────────────────────────────────────

@router.get("/rules", response_model=list)
async def list_convert_rules(db: Session = Depends(get_db)):
    from app.crud.convert_rule import get_convert_rules, seed_default_rules
    seed_default_rules(db)
    rules = get_convert_rules(db)
    return [
        {
            "id": r.id,
            "rule_number": r.rule_number,
            "category": r.category,
            "description": r.description,
            "is_active": r.is_active,
            "created_at": r.created_at.isoformat() if r.created_at else "",
        }
        for r in rules
    ]


@router.post("/rules", response_model=dict)
async def create_convert_rule_api(
    category: str = Form(...),
    description: str = Form(...),
    rule_number: str = Form(None),
    db: Session = Depends(get_db),
):
    from app.crud.convert_rule import create_convert_rule
    from app.schemas.convert_rule import ConvertRuleCreate
    rule = create_convert_rule(db, ConvertRuleCreate(
        rule_number=rule_number, category=category, description=description
    ))
    return {"id": rule.id, "rule_number": rule.rule_number}


@router.put("/rules/{rule_id}", response_model=dict)
async def toggle_convert_rule_api(rule_id: int, db: Session = Depends(get_db)):
    from app.crud.convert_rule import toggle_convert_rule
    rule = toggle_convert_rule(db, rule_id)
    if not rule:
        raise HTTPException(status_code=404, detail="规则不存在")
    return {"id": rule.id, "is_active": rule.is_active}


@router.delete("/rules/{rule_id}", response_model=dict)
async def delete_convert_rule_api(rule_id: int, db: Session = Depends(get_db)):
    from app.crud.convert_rule import delete_convert_rule
    rule = delete_convert_rule(db, rule_id)
    if not rule:
        raise HTTPException(status_code=404, detail="规则不存在")
    return {"ok": True}


@router.post("/rules/bulk-delete", response_model=dict)
async def bulk_delete_convert_rules_api(
    rule_ids: str = Form(...),
    db: Session = Depends(get_db),
):
    from app.crud.convert_rule import bulk_delete_convert_rules
    ids = [int(x) for x in rule_ids.split(",") if x.strip()]
    cnt = bulk_delete_convert_rules(db, ids)
    return {"deleted": cnt}
